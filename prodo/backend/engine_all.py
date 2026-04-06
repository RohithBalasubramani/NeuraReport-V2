from __future__ import annotations

"""Merged engine module: core + pipelines."""

"""Engine core — types, errors, events, results, domain models, orchestration."""

# CORE

"""Core module - types, errors, result, and event bus.

Consolidated from core/types.py, core/errors.py, core/result.py, core/events.py.
"""

import asyncio
import concurrent.futures
import logging
import time
import uuid
from abc import ABC
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import (
    Any,
    Awaitable,
    Callable,
    Dict,
    Generic,
    List,
    NewType,
    Optional,
    Protocol,
    TypeVar,
    Union,
)

# Types (from core/types.py)

EntityId = NewType("EntityId", str)
Timestamp = NewType("Timestamp", datetime)

JSON = Union[Dict[str, Any], List[Any], str, int, float, bool, None]
JSONObject = Dict[str, Any]
JSONArray = List[Any]

# Errors (from core/errors.py)

class NeuraError(Exception):
    """Base error type for all NeuraReport errors."""

    code = "error"

    def __init__(
        self,
        message: str,
        *,
        details: Optional[Dict[str, Any]] = None,
        cause: Optional[Exception] = None,
    ) -> None:
        super().__init__(message)
        self.message = message
        self.details = details
        self.cause = cause

    def __str__(self) -> str:
        if self.details:
            return f"[{self.code}] {self.message} - {self.details}"
        return f"[{self.code}] {self.message}"

    def to_dict(self) -> Dict[str, Any]:
        result = {"code": self.code, "message": self.message}
        if self.details:
            result["details"] = self.details
        return result

class ValidationError(NeuraError):
    """Raised when input validation fails."""

    code = "validation_error"

class NotFoundError(NeuraError):
    """Raised when a requested resource does not exist."""

    code = "not_found"

class ConflictError(NeuraError):
    """Raised when an operation conflicts with current state."""

    code = "conflict"

class ExternalServiceError(NeuraError):
    """Raised when an external service (LLM, email, etc.) fails."""

    code = "external_service_error"

    def __init__(
        self,
        message: str,
        *,
        service: str = "unknown",
        details: Optional[Dict[str, Any]] = None,
        cause: Optional[Exception] = None,
    ) -> None:
        super().__init__(message, details=details, cause=cause)
        self.service = service

class ConfigurationError(NeuraError):
    """Raised when system configuration is invalid."""

    code = "configuration_error"

class PipelineError(NeuraError):
    """Raised when a pipeline step fails."""

    code = "pipeline_error"

    def __init__(
        self,
        message: str,
        *,
        step: str = "unknown",
        details: Optional[Dict[str, Any]] = None,
        cause: Optional[Exception] = None,
    ) -> None:
        super().__init__(message, details=details, cause=cause)
        self.step = step

class DataSourceError(NeuraError):
    """Raised when data source operations fail."""

    code = "data_source_error"

class RenderError(NeuraError):
    """Raised when rendering fails."""

    code = "render_error"

    def __init__(
        self,
        message: str,
        *,
        format: str = "unknown",
        details: Optional[Dict[str, Any]] = None,
        cause: Optional[Exception] = None,
    ) -> None:
        super().__init__(message, details=details, cause=cause)
        self.format = format

# Result (from core/result.py)

T = TypeVar("T")
E = TypeVar("E", bound=NeuraError)
U = TypeVar("U")

@dataclass(frozen=True, slots=True)
class Ok(Generic[T]):
    """Success variant of Result."""

    value: T

    def is_ok(self) -> bool:
        return True

    def is_err(self) -> bool:
        return False

    def unwrap(self) -> T:
        return self.value

    def unwrap_or(self, default: T) -> T:
        return self.value

    def unwrap_or_else(self, f: Callable[[NeuraError], T]) -> T:
        return self.value

    def map(self, f: Callable[[T], U]) -> Result[U, NeuraError]:
        return Ok(f(self.value))

    def map_err(self, f: Callable[[NeuraError], NeuraError]) -> Result[T, NeuraError]:
        return self

    def and_then(self, f: Callable[[T], Result[U, NeuraError]]) -> Result[U, NeuraError]:
        return f(self.value)

    def or_else(self, f: Callable[[NeuraError], Result[T, NeuraError]]) -> Result[T, NeuraError]:
        return self

@dataclass(frozen=True, slots=True)
class Err(Generic[E]):
    """Error variant of Result."""

    error: E

    def is_ok(self) -> bool:
        return False

    def is_err(self) -> bool:
        return True

    def unwrap(self) -> T:
        raise self.error

    def unwrap_or(self, default: T) -> T:
        return default

    def unwrap_or_else(self, f: Callable[[E], T]) -> T:
        return f(self.error)

    def map(self, f: Callable[[T], U]) -> Result[U, E]:
        return self

    def map_err(self, f: Callable[[E], NeuraError]) -> Result[T, NeuraError]:
        return Err(f(self.error))

    def and_then(self, f: Callable[[T], Result[U, E]]) -> Result[U, E]:
        return self

    def or_else(self, f: Callable[[E], Result[T, E]]) -> Result[T, E]:
        return f(self.error)

Result = Union[Ok[T], Err[E]]

def result_from_exception(f: Callable[[], T]) -> Result[T, NeuraError]:
    """Execute a function and wrap its result in a Result type."""
    try:
        return Ok(f())
    except NeuraError as e:
        return Err(e)
    except Exception as e:
        return Err(NeuraError(code="unexpected_error", message=str(e), cause=e))

def collect_results(results: list[Result[T, E]]) -> Result[list[T], E]:
    """Collect a list of Results into a Result of list."""
    values = []
    for r in results:
        if r.is_err():
            return r
        values.append(r.unwrap())
    return Ok(values)

# Events (from core/events.py)

_events_logger = logging.getLogger("neura.events")

@dataclass(frozen=True)
class Event:
    """Immutable event with metadata."""

    name: str
    payload: Dict[str, Any]
    event_id: str = field(default_factory=lambda: uuid.uuid4().hex)
    correlation_id: Optional[str] = None
    timestamp: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    source: str = field(default="backend")

    def to_dict(self) -> Dict[str, Any]:
        return {
            "event_id": self.event_id,
            "name": self.name,
            "payload": self.payload,
            "correlation_id": self.correlation_id,
            "timestamp": self.timestamp.isoformat(),
            "source": self.source,
        }

EventHandler = Callable[[Event], Awaitable[None]]
EventMiddleware = Callable[[Event, Callable[[Event], Awaitable[None]]], Awaitable[None]]

class EventPersistence(Protocol):
    """Protocol for event persistence backends."""

    async def persist(self, event: Event) -> None: ...

    async def get_events(
        self,
        *,
        name: Optional[str] = None,
        correlation_id: Optional[str] = None,
        since: Optional[datetime] = None,
        limit: int = 100,
    ) -> List[Event]: ...

class EventBus:
    """Async event bus with middleware and persistence support."""

    def __init__(
        self,
        *,
        middlewares: Optional[List[EventMiddleware]] = None,
        persistence: Optional[EventPersistence] = None,
    ) -> None:
        self._handlers: Dict[str, List[EventHandler]] = {}
        self._middlewares = middlewares or []
        self._persistence = persistence
        self._wildcard_handlers: List[EventHandler] = []

    def subscribe(self, event_name: str, handler: EventHandler) -> Callable[[], None]:
        """Subscribe a handler to an event. Returns unsubscribe function."""
        if event_name == "*":
            self._wildcard_handlers.append(handler)
            return lambda: self._wildcard_handlers.remove(handler)

        if event_name not in self._handlers:
            self._handlers[event_name] = []
        self._handlers[event_name].append(handler)
        return lambda: self._handlers[event_name].remove(handler)

    async def publish(self, event: Event) -> None:
        """Publish an event to all subscribers."""
        async def dispatch(evt: Event) -> None:
            await self._dispatch(evt)

        handler = dispatch
        for middleware in reversed(self._middlewares):
            prev_handler = handler

            async def make_handler(m: EventMiddleware, h: Callable) -> Callable:
                async def wrapped(e: Event) -> None:
                    await m(e, h)
                return wrapped

            handler = await make_handler(middleware, prev_handler)

        await handler(event)

        if self._persistence:
            try:
                await self._persistence.persist(event)
            except Exception:
                _events_logger.exception("event_persist_failed", extra={"event": event.name})

    async def _dispatch(self, event: Event) -> None:
        """Dispatch event to registered handlers."""
        handlers = list(self._handlers.get(event.name, []))
        handlers.extend(self._wildcard_handlers)

        if not handlers:
            return

        results = await asyncio.gather(
            *[self._safe_call(handler, event) for handler in handlers],
            return_exceptions=True,
        )

        for result in results:
            if isinstance(result, Exception):
                _events_logger.exception(
                    "event_handler_failed",
                    extra={"event": event.name, "error": str(result)},
                )

    async def _safe_call(self, handler: EventHandler, event: Event) -> None:
        """Safely call a handler, catching exceptions."""
        try:
            await handler(event)
        except Exception as e:
            _events_logger.exception(
                "event_handler_error",
                extra={"event": event.name, "handler": handler.__name__},
            )
            raise

def logging_middleware(log: logging.Logger) -> EventMiddleware:
    """Middleware that logs all events."""

    async def middleware(event: Event, next_handler: Callable[[Event], Awaitable[None]]) -> None:
        start = time.perf_counter()
        log.info(
            "event_published",
            extra={
                "event": event.name,
                "event_id": event.event_id,
                "correlation_id": event.correlation_id,
            },
        )
        try:
            await next_handler(event)
        finally:
            elapsed = (time.perf_counter() - start) * 1000
            log.debug(
                "event_handled",
                extra={
                    "event": event.name,
                    "event_id": event.event_id,
                    "elapsed_ms": elapsed,
                },
            )

    return middleware

def metrics_middleware(log: logging.Logger) -> EventMiddleware:
    """Middleware that tracks event metrics."""

    _counts: Dict[str, int] = {}

    async def middleware(event: Event, next_handler: Callable[[Event], Awaitable[None]]) -> None:
        _counts[event.name] = _counts.get(event.name, 0) + 1
        await next_handler(event)

    return middleware

_global_bus: Optional[EventBus] = None

def get_event_bus() -> EventBus:
    """Get or create the global event bus."""
    global _global_bus
    if _global_bus is None:
        _global_bus = EventBus(
            middlewares=[
                logging_middleware(_events_logger),
                metrics_middleware(_events_logger),
            ]
        )
    return _global_bus

def publish_sync(event: Event) -> None:
    """Publish an event synchronously (for use in sync code)."""
    bus = get_event_bus()
    try:
        loop = asyncio.get_running_loop()
        asyncio.run_coroutine_threadsafe(bus.publish(event), loop)
    except RuntimeError:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            pool.submit(asyncio.run, bus.publish(event)).result()

__all__ = [
    # Types
    "EntityId", "Timestamp", "JSON", "JSONObject", "JSONArray",
    # Errors
    "NeuraError", "ValidationError", "NotFoundError", "ConflictError",
    "ExternalServiceError", "ConfigurationError", "PipelineError",
    "DataSourceError", "RenderError",
    # Result
    "Result", "Ok", "Err", "result_from_exception", "collect_results",
    # Events
    "Event", "EventBus", "EventHandler", "EventMiddleware", "EventPersistence",
    "get_event_bus", "publish_sync", "logging_middleware", "metrics_middleware",
]

# DOMAIN

"""Domain layer - pure business logic with no IO dependencies.

Consolidated from domain/connections.py, domain/contracts.py, domain/jobs.py,
domain/reports.py, domain/templates.py.
"""

import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

# Connections (from domain/connections.py)

class ConnectionType(str, Enum):
    """Types of database connections."""

    SQLITE = "sqlite"
    POSTGRES = "postgres"
    MYSQL = "mysql"

class ConnectionStatus(str, Enum):
    """Health status of a connection."""

    UNKNOWN = "unknown"
    HEALTHY = "healthy"
    DEGRADED = "degraded"
    UNAVAILABLE = "unavailable"

@dataclass(frozen=True)
class TableInfo:
    """Information about a database table."""

    name: str
    columns: List[str]
    row_count: Optional[int] = None
    primary_key: Optional[str] = None

@dataclass(frozen=True)
class SchemaInfo:
    """Database schema information."""

    tables: List[TableInfo]
    catalog: List[str]

    @property
    def table_names(self) -> List[str]:
        return [t.name for t in self.tables]

@dataclass
class ConnectionTest:
    """Result of testing a connection."""

    success: bool
    latency_ms: Optional[float] = None
    error: Optional[str] = None
    tested_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    table_count: Optional[int] = None
    total_rows: Optional[int] = None

@dataclass
class Connection:
    """A database connection configuration."""

    connection_id: str
    name: str
    connection_type: ConnectionType
    path: Path
    status: ConnectionStatus = ConnectionStatus.UNKNOWN
    schema_info: Optional[SchemaInfo] = None
    last_test: Optional[ConnectionTest] = None
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    last_used_at: Optional[datetime] = None
    last_used_template: Optional[str] = None
    description: Optional[str] = None
    tags: List[str] = field(default_factory=list)

    @classmethod
    def create_sqlite(
        cls,
        name: str,
        path: Path,
        connection_id: Optional[str] = None,
        **kwargs: Any,
    ) -> Connection:
        return cls(
            connection_id=connection_id or str(uuid.uuid4()),
            name=name,
            connection_type=ConnectionType.SQLITE,
            path=path.resolve(),
            **kwargs,
        )

    def record_test(self, test: ConnectionTest) -> None:
        self.last_test = test
        self.status = (
            ConnectionStatus.HEALTHY
            if test.success
            else ConnectionStatus.UNAVAILABLE
        )
        self.updated_at = datetime.now(timezone.utc)

    def record_use(self, template_id: str) -> None:
        self.last_used_at = datetime.now(timezone.utc)
        self.last_used_template = template_id
        self.updated_at = datetime.now(timezone.utc)

    def update_schema(self, schema: SchemaInfo) -> None:
        self.schema_info = schema
        self.updated_at = datetime.now(timezone.utc)

    @property
    def catalog(self) -> List[str]:
        if self.schema_info:
            return self.schema_info.catalog
        return []

    def to_dict(self) -> Dict[str, Any]:
        return {
            "connection_id": self.connection_id,
            "name": self.name,
            "connection_type": self.connection_type.value,
            "path": str(self.path),
            "status": self.status.value,
            "schema_info": {
                "tables": [
                    {
                        "name": t.name,
                        "columns": t.columns,
                        "row_count": t.row_count,
                        "primary_key": t.primary_key,
                    }
                    for t in self.schema_info.tables
                ],
                "catalog": self.schema_info.catalog,
            }
            if self.schema_info
            else None,
            "last_test": {
                "success": self.last_test.success,
                "latency_ms": self.last_test.latency_ms,
                "error": self.last_test.error,
                "tested_at": self.last_test.tested_at.isoformat(),
                "table_count": self.last_test.table_count,
            }
            if self.last_test
            else None,
            "created_at": self.created_at.isoformat(),
            "updated_at": self.updated_at.isoformat(),
            "last_used_at": self.last_used_at.isoformat() if self.last_used_at else None,
            "last_used_template": self.last_used_template,
            "description": self.description,
            "tags": self.tags,
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> Connection:
        schema_data = data.get("schema_info")
        schema_info = None
        if schema_data:
            tables = [
                TableInfo(
                    name=t["name"],
                    columns=t["columns"],
                    row_count=t.get("row_count"),
                    primary_key=t.get("primary_key"),
                )
                for t in schema_data.get("tables", [])
            ]
            schema_info = SchemaInfo(
                tables=tables,
                catalog=schema_data.get("catalog", []),
            )

        test_data = data.get("last_test")
        last_test = None
        if test_data:
            tested_at = test_data.get("tested_at")
            if isinstance(tested_at, str):
                tested_at = datetime.fromisoformat(tested_at)
            last_test = ConnectionTest(
                success=test_data["success"],
                latency_ms=test_data.get("latency_ms"),
                error=test_data.get("error"),
                tested_at=tested_at or datetime.now(timezone.utc),
                table_count=test_data.get("table_count"),
            )

        created_at = data.get("created_at")
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at)

        updated_at = data.get("updated_at")
        if isinstance(updated_at, str):
            updated_at = datetime.fromisoformat(updated_at)

        last_used_at = data.get("last_used_at")
        if isinstance(last_used_at, str):
            last_used_at = datetime.fromisoformat(last_used_at)

        return cls(
            connection_id=data["connection_id"],
            name=data["name"],
            connection_type=ConnectionType(data.get("connection_type", "sqlite")),
            path=Path(data["path"]),
            status=ConnectionStatus(data.get("status", "unknown")),
            schema_info=schema_info,
            last_test=last_test,
            created_at=created_at or datetime.now(timezone.utc),
            updated_at=updated_at or datetime.now(timezone.utc),
            last_used_at=last_used_at,
            last_used_template=data.get("last_used_template"),
            description=data.get("description"),
            tags=data.get("tags", []),
        )

# Contracts (from domain/contracts.py)

class TokenType(str, Enum):
    """Types of tokens in a contract."""

    SCALAR = "scalar"
    ROW = "row"
    TOTAL = "total"

@dataclass(frozen=True)
class Token:
    """A single token/placeholder in a contract."""

    name: str
    token_type: TokenType
    expression: Optional[str] = None
    description: Optional[str] = None

    def __post_init__(self) -> None:
        if not self.name or not self.name.strip():
            raise ValueError("Token name cannot be empty")

@dataclass(frozen=True)
class TokenSet:
    """Collection of tokens organized by type."""

    scalars: List[str] = field(default_factory=list)
    row_tokens: List[str] = field(default_factory=list)
    totals: List[str] = field(default_factory=list)

    def all_tokens(self) -> List[str]:
        return [*self.scalars, *self.row_tokens, *self.totals]

    def __contains__(self, token: str) -> bool:
        return token in self.scalars or token in self.row_tokens or token in self.totals

@dataclass(frozen=True)
class Mapping:
    """Mapping from token name to SQL expression."""

    token: str
    expression: str
    source_table: Optional[str] = None
    source_column: Optional[str] = None

    def __post_init__(self) -> None:
        if not self.token.strip():
            raise ValueError("Mapping token cannot be empty")
        if not self.expression.strip():
            raise ValueError(f"Mapping expression for '{self.token}' cannot be empty")

@dataclass(frozen=True)
class ReshapeColumn:
    """A column in a reshape rule."""

    alias: str
    sources: List[str]

@dataclass(frozen=True)
class ReshapeRule:
    """Rule for reshaping/transforming data before rendering."""

    purpose: str
    strategy: str
    columns: List[ReshapeColumn]
    order_by: List[str] = field(default_factory=list)
    filters: Optional[str] = None
    group_by: Optional[List[str]] = None
    explain: Optional[str] = None

@dataclass(frozen=True)
class JoinSpec:
    """Specification for joining tables."""

    parent_table: str
    parent_key: str
    child_table: str
    child_key: str

    def is_valid(self) -> bool:
        return bool(self.parent_table and self.parent_key)

@dataclass(frozen=True)
class OrderSpec:
    """Ordering specification for rows."""

    rows: List[str] = field(default_factory=list)

@dataclass
class Contract:
    """Complete contract for report generation."""

    contract_id: str
    template_id: str
    tokens: TokenSet
    mappings: Dict[str, str]
    reshape_rules: List[ReshapeRule] = field(default_factory=list)
    join: Optional[JoinSpec] = None
    date_columns: Dict[str, str] = field(default_factory=dict)
    order_by: OrderSpec = field(default_factory=OrderSpec)
    row_order: List[str] = field(default_factory=lambda: ["ROWID"])
    literals: Dict[str, Any] = field(default_factory=dict)
    totals_math: Dict[str, str] = field(default_factory=dict)
    row_computed: Dict[str, str] = field(default_factory=dict)
    created_at: Optional[datetime] = None
    version: str = "v2"

    @property
    def header_tokens(self) -> List[str]:
        return list(self.tokens.scalars)

    @property
    def row_tokens(self) -> List[str]:
        return list(self.tokens.row_tokens)

    @property
    def totals(self) -> Dict[str, str]:
        return {tok: self.mappings.get(tok, "") for tok in self.tokens.totals}

    def get_mapping(self, token: str) -> Optional[str]:
        return self.mappings.get(token)

    def validate(self) -> List[str]:
        """Validate contract and return list of issues."""
        issues = []
        for token in self.tokens.all_tokens():
            if token not in self.mappings:
                issues.append(f"Token '{token}' has no mapping")
        for rule in self.reshape_rules:
            if not rule.columns:
                issues.append(f"Reshape rule '{rule.purpose}' has no columns")
        if self.join and not self.join.is_valid():
            issues.append("Join specification is incomplete")
        return issues

    def to_dict(self) -> Dict[str, Any]:
        """Serialize to dict for persistence."""
        return {
            "contract_id": self.contract_id,
            "template_id": self.template_id,
            "tokens": {
                "scalars": list(self.tokens.scalars),
                "row_tokens": list(self.tokens.row_tokens),
                "totals": list(self.tokens.totals),
            },
            "mapping": self.mappings,
            "reshape_rules": [
                {
                    "purpose": r.purpose,
                    "strategy": r.strategy,
                    "columns": [{"as": c.alias, "from": c.sources} for c in r.columns],
                    "order_by": r.order_by,
                    "filters": r.filters,
                    "group_by": r.group_by,
                    "explain": r.explain,
                }
                for r in self.reshape_rules
            ],
            "join": {
                "parent_table": self.join.parent_table,
                "parent_key": self.join.parent_key,
                "child_table": self.join.child_table,
                "child_key": self.join.child_key,
            }
            if self.join
            else None,
            "date_columns": dict(self.date_columns),
            "order_by": {"rows": list(self.order_by.rows)},
            "row_order": list(self.row_order),
            "literals": self.literals,
            "totals_math": self.totals_math,
            "row_computed": self.row_computed,
            "version": self.version,
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any], template_id: str) -> Contract:
        """Deserialize from dict."""
        tokens_data = data.get("tokens", {})
        tokens = TokenSet(
            scalars=list(tokens_data.get("scalars", [])),
            row_tokens=list(tokens_data.get("row_tokens", [])),
            totals=list(tokens_data.get("totals", [])),
        )

        reshape_rules = []
        for rule_data in data.get("reshape_rules", []):
            columns = [
                ReshapeColumn(alias=c.get("as", ""), sources=c.get("from", []))
                for c in rule_data.get("columns", [])
            ]
            reshape_rules.append(
                ReshapeRule(
                    purpose=rule_data.get("purpose", ""),
                    strategy=rule_data.get("strategy", "SELECT"),
                    columns=columns,
                    order_by=rule_data.get("order_by", []),
                    filters=rule_data.get("filters"),
                    group_by=rule_data.get("group_by"),
                    explain=rule_data.get("explain"),
                )
            )

        join_data = data.get("join")
        join = None
        if join_data and isinstance(join_data, dict):
            join = JoinSpec(
                parent_table=join_data.get("parent_table", ""),
                parent_key=join_data.get("parent_key", ""),
                child_table=join_data.get("child_table", ""),
                child_key=join_data.get("child_key", ""),
            )

        order_data = data.get("order_by", {})
        order_by = OrderSpec(
            rows=order_data.get("rows", []) if isinstance(order_data, dict) else []
        )
        date_columns_raw = data.get("date_columns", {})
        date_columns = (
            {str(k): str(v) for k, v in date_columns_raw.items()}
            if isinstance(date_columns_raw, dict)
            else {}
        )

        return cls(
            contract_id=data.get("contract_id", f"contract-{template_id}"),
            template_id=template_id,
            tokens=tokens,
            mappings=dict(data.get("mapping", {})),
            reshape_rules=reshape_rules,
            join=join,
            date_columns=date_columns,
            order_by=order_by,
            row_order=list(data.get("row_order", ["ROWID"])),
            literals=dict(data.get("literals", {})),
            totals_math=dict(data.get("totals_math", {})),
            row_computed=dict(data.get("row_computed", {})),
            version=data.get("version", "v2"),
        )

# Jobs (from domain/jobs.py)

class JobStatus(str, Enum):
    """Status of a job."""

    PENDING = "pending"
    QUEUED = "queued"
    RUNNING = "running"
    SUCCEEDED = "succeeded"
    FAILED = "failed"
    CANCELLED = "cancelled"

    @property
    def is_terminal(self) -> bool:
        return self in (JobStatus.SUCCEEDED, JobStatus.FAILED, JobStatus.CANCELLED)

    @property
    def is_active(self) -> bool:
        return self in (JobStatus.PENDING, JobStatus.QUEUED, JobStatus.RUNNING)

class StepStatus(str, Enum):
    """Status of a job step."""

    PENDING = "pending"
    RUNNING = "running"
    SUCCEEDED = "succeeded"
    FAILED = "failed"
    SKIPPED = "skipped"

@dataclass
class JobStep:
    """A step within a job."""

    name: str
    label: str
    status: StepStatus = StepStatus.PENDING
    error: Optional[str] = None
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    progress_weight: float = 0.0

    def start(self) -> None:
        self.status = StepStatus.RUNNING
        self.started_at = datetime.now(timezone.utc)

    def succeed(self) -> None:
        self.status = StepStatus.SUCCEEDED
        self.completed_at = datetime.now(timezone.utc)

    def fail(self, error: str) -> None:
        self.status = StepStatus.FAILED
        self.error = error
        self.completed_at = datetime.now(timezone.utc)

    def skip(self) -> None:
        self.status = StepStatus.SKIPPED
        self.completed_at = datetime.now(timezone.utc)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "label": self.label,
            "status": self.status.value,
            "error": self.error,
            "started_at": self.started_at.isoformat() if self.started_at else None,
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
        }

class JobType(str, Enum):
    """Types of jobs."""

    REPORT_GENERATION = "run_report"
    TEMPLATE_IMPORT = "import_template"
    TEMPLATE_ANALYSIS = "analyze_template"
    CONTRACT_BUILD = "build_contract"
    SCHEMA_DISCOVERY = "discover_schema"

@dataclass
class Job:
    """A tracked job/operation."""

    job_id: str
    job_type: JobType
    status: JobStatus
    steps: List[JobStep] = field(default_factory=list)
    progress: float = 0.0
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    template_id: Optional[str] = None
    template_name: Optional[str] = None
    template_kind: Optional[str] = None
    connection_id: Optional[str] = None
    schedule_id: Optional[str] = None
    correlation_id: Optional[str] = None
    meta: Dict[str, Any] = field(default_factory=dict)
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None

    @classmethod
    def create(
        cls,
        job_type: JobType,
        steps: Optional[List[JobStep]] = None,
        **kwargs: Any,
    ) -> Job:
        return cls(
            job_id=str(uuid.uuid4()),
            job_type=job_type,
            status=JobStatus.PENDING,
            steps=steps or [],
            **kwargs,
        )

    def can_transition_to(self, new_status: JobStatus) -> bool:
        valid_transitions = {
            JobStatus.PENDING: {JobStatus.QUEUED, JobStatus.RUNNING, JobStatus.CANCELLED},
            JobStatus.QUEUED: {JobStatus.RUNNING, JobStatus.CANCELLED},
            JobStatus.RUNNING: {JobStatus.SUCCEEDED, JobStatus.FAILED, JobStatus.CANCELLED},
            JobStatus.SUCCEEDED: set(),
            JobStatus.FAILED: set(),
            JobStatus.CANCELLED: set(),
        }
        return new_status in valid_transitions.get(self.status, set())

    def start(self) -> None:
        if not self.can_transition_to(JobStatus.RUNNING):
            return
        self.status = JobStatus.RUNNING
        self.started_at = datetime.now(timezone.utc)

    def succeed(self, result: Optional[Dict[str, Any]] = None) -> None:
        if not self.can_transition_to(JobStatus.SUCCEEDED):
            return
        self.status = JobStatus.SUCCEEDED
        self.progress = 100.0
        self.result = result
        self.completed_at = datetime.now(timezone.utc)

    def fail(self, error: str) -> None:
        if not self.can_transition_to(JobStatus.FAILED):
            return
        self.status = JobStatus.FAILED
        self.error = error
        self.completed_at = datetime.now(timezone.utc)

    def cancel(self) -> None:
        if not self.can_transition_to(JobStatus.CANCELLED):
            return
        self.status = JobStatus.CANCELLED
        self.completed_at = datetime.now(timezone.utc)

    def update_progress(self, progress: float) -> None:
        self.progress = max(0.0, min(100.0, progress))

    def get_step(self, name: str) -> Optional[JobStep]:
        for step in self.steps:
            if step.name == name:
                return step
        return None

    def step_running(self, name: str) -> None:
        step = self.get_step(name)
        if step:
            step.start()

    def step_succeeded(self, name: str, progress: Optional[float] = None) -> None:
        step = self.get_step(name)
        if step:
            step.succeed()
            if progress is not None:
                self.update_progress(progress)

    def step_failed(self, name: str, error: str) -> None:
        step = self.get_step(name)
        if step:
            step.fail(error)

    @property
    def duration_seconds(self) -> Optional[float]:
        if not self.started_at:
            return None
        end = self.completed_at or datetime.now(timezone.utc)
        return (end - self.started_at).total_seconds()

    def to_dict(self) -> Dict[str, Any]:
        return {
            "job_id": self.job_id,
            "job_type": self.job_type.value,
            "status": self.status.value,
            "steps": [s.to_dict() for s in self.steps],
            "progress": self.progress,
            "result": self.result,
            "error": self.error,
            "template_id": self.template_id,
            "template_name": self.template_name,
            "template_kind": self.template_kind,
            "connection_id": self.connection_id,
            "schedule_id": self.schedule_id,
            "correlation_id": self.correlation_id,
            "meta": self.meta,
            "created_at": self.created_at.isoformat(),
            "started_at": self.started_at.isoformat() if self.started_at else None,
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
        }

@dataclass
class Schedule:
    """A recurring job schedule."""

    schedule_id: str
    name: str
    template_id: str
    connection_id: str
    interval_minutes: int
    active: bool = True
    template_name: Optional[str] = None
    template_kind: str = "pdf"
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    batch_ids: Optional[List[str]] = None
    key_values: Optional[Dict[str, str]] = None
    docx: bool = False
    xlsx: bool = False
    email_recipients: Optional[List[str]] = None
    email_subject: Optional[str] = None
    email_message: Optional[str] = None
    next_run_at: Optional[datetime] = None
    last_run_at: Optional[datetime] = None
    last_run_status: Optional[str] = None
    last_run_error: Optional[str] = None
    run_count: int = 0
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))

    @classmethod
    def create(
        cls,
        name: str,
        template_id: str,
        connection_id: str,
        interval_minutes: int,
        **kwargs: Any,
    ) -> Schedule:
        now = datetime.now(timezone.utc)
        return cls(
            schedule_id=str(uuid.uuid4()),
            name=name,
            template_id=template_id,
            connection_id=connection_id,
            interval_minutes=max(1, interval_minutes),
            next_run_at=now + timedelta(minutes=max(1, interval_minutes)),
            **kwargs,
        )

    def is_due(self, now: Optional[datetime] = None) -> bool:
        if not self.active or not self.next_run_at:
            return False
        now = now or datetime.now(timezone.utc)
        return self.next_run_at <= now

    def record_run(
        self,
        status: str,
        error: Optional[str] = None,
        finished_at: Optional[datetime] = None,
    ) -> None:
        finished = finished_at or datetime.now(timezone.utc)
        self.last_run_at = finished
        self.last_run_status = status
        self.last_run_error = error
        self.run_count += 1
        self.next_run_at = finished + timedelta(minutes=self.interval_minutes)
        self.updated_at = datetime.now(timezone.utc)

    def pause(self) -> None:
        self.active = False
        self.updated_at = datetime.now(timezone.utc)

    def resume(self) -> None:
        self.active = True
        if not self.next_run_at or self.next_run_at < datetime.now(timezone.utc):
            self.next_run_at = datetime.now(timezone.utc) + timedelta(
                minutes=self.interval_minutes
            )
        self.updated_at = datetime.now(timezone.utc)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "schedule_id": self.schedule_id,
            "id": self.schedule_id,
            "name": self.name,
            "template_id": self.template_id,
            "connection_id": self.connection_id,
            "interval_minutes": self.interval_minutes,
            "active": self.active,
            "template_name": self.template_name,
            "template_kind": self.template_kind,
            "start_date": self.start_date,
            "end_date": self.end_date,
            "batch_ids": self.batch_ids,
            "key_values": self.key_values,
            "docx": self.docx,
            "xlsx": self.xlsx,
            "email_recipients": self.email_recipients,
            "email_subject": self.email_subject,
            "email_message": self.email_message,
            "next_run_at": self.next_run_at.isoformat() if self.next_run_at else None,
            "last_run_at": self.last_run_at.isoformat() if self.last_run_at else None,
            "last_run_status": self.last_run_status,
            "last_run_error": self.last_run_error,
            "run_count": self.run_count,
            "created_at": self.created_at.isoformat(),
            "updated_at": self.updated_at.isoformat(),
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> Schedule:
        def parse_dt(val: Any) -> Optional[datetime]:
            if val is None:
                return None
            if isinstance(val, datetime):
                return val
            if isinstance(val, str):
                return datetime.fromisoformat(val)
            return None

        return cls(
            schedule_id=data.get("schedule_id") or data.get("id", str(uuid.uuid4())),
            name=data["name"],
            template_id=data["template_id"],
            connection_id=data["connection_id"],
            interval_minutes=int(data.get("interval_minutes", 60)),
            active=bool(data.get("active", True)),
            template_name=data.get("template_name"),
            template_kind=data.get("template_kind", "pdf"),
            start_date=data.get("start_date"),
            end_date=data.get("end_date"),
            batch_ids=data.get("batch_ids"),
            key_values=data.get("key_values"),
            docx=bool(data.get("docx")),
            xlsx=bool(data.get("xlsx")),
            email_recipients=data.get("email_recipients"),
            email_subject=data.get("email_subject"),
            email_message=data.get("email_message"),
            next_run_at=parse_dt(data.get("next_run_at")),
            last_run_at=parse_dt(data.get("last_run_at")),
            last_run_status=data.get("last_run_status"),
            last_run_error=data.get("last_run_error"),
            run_count=int(data.get("run_count", 0)),
            created_at=parse_dt(data.get("created_at")) or datetime.now(timezone.utc),
            updated_at=parse_dt(data.get("updated_at")) or datetime.now(timezone.utc),
        )

# Reports (from domain/reports.py)

class OutputFormat(str, Enum):
    """Supported output formats."""

    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"

class ReportStatus(str, Enum):
    """Status of a report generation."""

    PENDING = "pending"
    RUNNING = "running"
    SUCCEEDED = "succeeded"
    FAILED = "failed"
    CANCELLED = "cancelled"

@dataclass(frozen=True)
class Batch:
    """A batch of data for report generation."""

    batch_id: str
    row_count: int
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass(frozen=True)
class KeyValue:
    """A key-value filter for report generation."""

    key: str
    value: str

    def to_sql_condition(self, table: Optional[str] = None) -> str:
        if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', self.key):
            raise ValueError(f"Invalid SQL identifier in key: {self.key}")
        if table and not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', table):
            raise ValueError(f"Invalid SQL identifier in table: {table}")
        safe_key = f"[{self.key}]"
        col = f"[{table}].{safe_key}" if table else safe_key
        safe_value = self.value.replace("'", "''")
        return f"{col} = '{safe_value}'"

@dataclass
class RenderRequest:
    """Request to render a report."""

    template_id: str
    connection_id: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    batch_ids: Optional[List[str]] = None
    key_values: Optional[List[KeyValue]] = None
    output_formats: List[OutputFormat] = field(
        default_factory=lambda: [OutputFormat.HTML, OutputFormat.PDF]
    )
    email_recipients: Optional[List[str]] = None
    email_subject: Optional[str] = None
    email_message: Optional[str] = None
    correlation_id: Optional[str] = None

    def __post_init__(self) -> None:
        if not self.template_id:
            raise ValueError("template_id is required")
        if not self.connection_id:
            raise ValueError("connection_id is required")

@dataclass
class RenderOutput:
    """Output artifact from rendering."""

    format: OutputFormat
    path: Path
    size_bytes: int
    checksum: Optional[str] = None
    url: Optional[str] = None

@dataclass
class Report:
    """A generated report."""

    report_id: str
    template_id: str
    template_name: str
    connection_id: str
    connection_name: Optional[str]
    status: ReportStatus
    outputs: List[RenderOutput] = field(default_factory=list)
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    batch_ids: Optional[List[str]] = None
    key_values: Optional[List[KeyValue]] = None
    error: Optional[str] = None
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    correlation_id: Optional[str] = None
    schedule_id: Optional[str] = None
    schedule_name: Optional[str] = None

    @classmethod
    def create(
        cls,
        template_id: str,
        template_name: str,
        connection_id: str,
        connection_name: Optional[str] = None,
        **kwargs: Any,
    ) -> Report:
        return cls(
            report_id=str(uuid.uuid4()),
            template_id=template_id,
            template_name=template_name,
            connection_id=connection_id,
            connection_name=connection_name,
            status=ReportStatus.PENDING,
            **kwargs,
        )

    def start(self) -> None:
        self.status = ReportStatus.RUNNING
        self.started_at = datetime.now(timezone.utc)

    def succeed(self, outputs: List[RenderOutput]) -> None:
        self.status = ReportStatus.SUCCEEDED
        self.outputs = outputs
        self.completed_at = datetime.now(timezone.utc)

    def fail(self, error: str) -> None:
        self.status = ReportStatus.FAILED
        self.error = error
        self.completed_at = datetime.now(timezone.utc)

    def cancel(self) -> None:
        self.status = ReportStatus.CANCELLED
        self.completed_at = datetime.now(timezone.utc)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "report_id": self.report_id,
            "template_id": self.template_id,
            "template_name": self.template_name,
            "connection_id": self.connection_id,
            "connection_name": self.connection_name,
            "status": self.status.value,
            "outputs": [
                {
                    "format": o.format.value,
                    "path": str(o.path),
                    "size_bytes": o.size_bytes,
                    "url": o.url,
                }
                for o in self.outputs
            ],
            "start_date": self.start_date,
            "end_date": self.end_date,
            "batch_ids": self.batch_ids,
            "error": self.error,
            "started_at": self.started_at.isoformat() if self.started_at else None,
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
            "correlation_id": self.correlation_id,
        }

@dataclass(frozen=True)
class DataWindow:
    """A window of data for report generation."""

    start_date: Optional[str]
    end_date: Optional[str]
    filters: List[KeyValue] = field(default_factory=list)

    def to_sql_conditions(self, date_column: str = "date") -> List[str]:
        conditions = []
        if self.start_date:
            conditions.append(f"{date_column} >= '{self.start_date}'")
        if self.end_date:
            conditions.append(f"{date_column} <= '{self.end_date}'")
        for kv in self.filters:
            conditions.append(kv.to_sql_condition())
        return conditions

# Templates (from domain/templates.py)

class TemplateKind(str, Enum):
    """Types of templates."""

    PDF = "pdf"
    EXCEL = "excel"

class TemplateStatus(str, Enum):
    """Template lifecycle status."""

    DRAFT = "draft"
    ANALYZING = "analyzing"
    MAPPED = "mapped"
    APPROVED = "approved"
    FAILED = "failed"

@dataclass(frozen=True)
class Artifact:
    """A file artifact associated with a template."""

    name: str
    path: Path
    artifact_type: str
    size_bytes: Optional[int] = None
    checksum: Optional[str] = None
    created_at: Optional[datetime] = None

@dataclass
class TemplateSchema:
    """Schema extracted from a template."""

    scalars: List[str] = field(default_factory=list)
    row_tokens: List[str] = field(default_factory=list)
    totals: List[str] = field(default_factory=list)
    tables_detected: List[str] = field(default_factory=list)
    placeholders_found: int = 0

@dataclass
class Template:
    """A report template."""

    template_id: str
    name: str
    kind: TemplateKind
    status: TemplateStatus
    schema: Optional[TemplateSchema] = None
    contract_id: Optional[str] = None
    artifacts: List[Artifact] = field(default_factory=list)
    source_file: Optional[str] = None
    description: Optional[str] = None
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    last_run_at: Optional[datetime] = None
    run_count: int = 0
    tags: List[str] = field(default_factory=list)

    @classmethod
    def create(
        cls,
        name: str,
        kind: TemplateKind = TemplateKind.PDF,
        template_id: Optional[str] = None,
        **kwargs: Any,
    ) -> Template:
        return cls(
            template_id=template_id or str(uuid.uuid4()),
            name=name,
            kind=kind,
            status=TemplateStatus.DRAFT,
            **kwargs,
        )

    def record_run(self) -> None:
        self.run_count += 1
        self.last_run_at = datetime.now(timezone.utc)
        self.updated_at = datetime.now(timezone.utc)

    def transition_to(self, status: TemplateStatus) -> None:
        self.status = status
        self.updated_at = datetime.now(timezone.utc)

    def add_artifact(self, artifact: Artifact) -> None:
        self.artifacts = [a for a in self.artifacts if a.name != artifact.name]
        self.artifacts.append(artifact)
        self.updated_at = datetime.now(timezone.utc)

    def get_artifact(self, name: str) -> Optional[Artifact]:
        for artifact in self.artifacts:
            if artifact.name == name:
                return artifact
        return None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "template_id": self.template_id,
            "name": self.name,
            "kind": self.kind.value,
            "status": self.status.value,
            "schema": {
                "scalars": self.schema.scalars,
                "row_tokens": self.schema.row_tokens,
                "totals": self.schema.totals,
                "tables_detected": self.schema.tables_detected,
                "placeholders_found": self.schema.placeholders_found,
            }
            if self.schema
            else None,
            "contract_id": self.contract_id,
            "artifacts": [
                {
                    "name": a.name,
                    "path": str(a.path),
                    "artifact_type": a.artifact_type,
                    "size_bytes": a.size_bytes,
                }
                for a in self.artifacts
            ],
            "source_file": self.source_file,
            "description": self.description,
            "created_at": self.created_at.isoformat(),
            "updated_at": self.updated_at.isoformat(),
            "last_run_at": self.last_run_at.isoformat() if self.last_run_at else None,
            "run_count": self.run_count,
            "tags": self.tags,
        }

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> Template:
        schema_data = data.get("schema")
        schema = None
        if schema_data:
            schema = TemplateSchema(
                scalars=schema_data.get("scalars", []),
                row_tokens=schema_data.get("row_tokens", []),
                totals=schema_data.get("totals", []),
                tables_detected=schema_data.get("tables_detected", []),
                placeholders_found=schema_data.get("placeholders_found", 0),
            )

        artifacts = []
        for a in data.get("artifacts", []):
            artifacts.append(
                Artifact(
                    name=a["name"],
                    path=Path(a["path"]),
                    artifact_type=a["artifact_type"],
                    size_bytes=a.get("size_bytes"),
                )
            )

        created_at = data.get("created_at")
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at)

        updated_at = data.get("updated_at")
        if isinstance(updated_at, str):
            updated_at = datetime.fromisoformat(updated_at)

        last_run_at = data.get("last_run_at")
        if isinstance(last_run_at, str):
            last_run_at = datetime.fromisoformat(last_run_at)

        return cls(
            template_id=data["template_id"],
            name=data["name"],
            kind=TemplateKind(data.get("kind", "pdf")),
            status=TemplateStatus(data.get("status", "draft")),
            schema=schema,
            contract_id=data.get("contract_id"),
            artifacts=artifacts,
            source_file=data.get("source_file"),
            description=data.get("description"),
            created_at=created_at or datetime.now(timezone.utc),
            updated_at=updated_at or datetime.now(timezone.utc),
            last_run_at=last_run_at,
            run_count=data.get("run_count", 0),
            tags=data.get("tags", []),
        )

__all__ = [
    # Connections
    "ConnectionType", "ConnectionStatus", "TableInfo", "SchemaInfo",
    "ConnectionTest", "Connection",
    # Contracts
    "TokenType", "Token", "TokenSet", "Mapping", "ReshapeColumn",
    "ReshapeRule", "JoinSpec", "OrderSpec", "Contract",
    # Jobs
    "JobStatus", "StepStatus", "JobStep", "JobType", "Job", "Schedule",
    # Reports
    "OutputFormat", "ReportStatus", "Batch", "KeyValue", "RenderRequest",
    "RenderOutput", "Report", "DataWindow",
    # Templates
    "TemplateKind", "TemplateStatus", "Artifact", "TemplateSchema", "Template",
]

# ORCHESTRATION

"""Job orchestration - executor, scheduler, and worker pool.

Consolidated from orchestration/executor.py, orchestration/scheduler.py, orchestration/worker.py.
"""

import asyncio
import logging
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Callable, Dict, List, Optional, Protocol

from backend.engine_all import Event, publish_sync
from backend.engine_all import Schedule, Job, JobStatus, JobType, JobStep

# Executor (from orchestration/executor.py)

_executor_logger = logging.getLogger("neura.orchestration.executor")

@dataclass
class ExecutorConfig:
    """Configuration for job executor."""

    max_workers: int = 4
    default_timeout_seconds: float = 600.0
    enable_thread_injection_cancel: bool = False

@dataclass
class JobExecution:
    """Tracks a single job execution."""

    job: Job
    future: Optional[asyncio.Future] = None
    thread_id: Optional[int] = None
    started_at: Optional[datetime] = None
    child_pids: set = field(default_factory=set)

JobRunner = Callable[[Job, "JobExecutor"], Any]

class JobExecutor:
    """Executes jobs with tracking and cancellation support."""

    def __init__(self, config: Optional[ExecutorConfig] = None) -> None:
        self._config = config or ExecutorConfig()
        self._pool = ThreadPoolExecutor(
            max_workers=self._config.max_workers,
            thread_name_prefix="job-executor",
        )
        self._executions: Dict[str, JobExecution] = {}
        self._runners: Dict[JobType, JobRunner] = {}
        self._shutdown = False

    def register_runner(self, job_type: JobType, runner: JobRunner) -> None:
        """Register a runner function for a job type."""
        self._runners[job_type] = runner

    async def submit(self, job: Job) -> None:
        """Submit a job for execution."""
        if self._shutdown:
            raise RuntimeError("Executor is shutting down")

        if job.job_id in self._executions:
            raise ValueError(f"Job {job.job_id} is already running")

        runner = self._runners.get(job.job_type)
        if not runner:
            raise ValueError(f"No runner registered for job type: {job.job_type}")

        execution = JobExecution(job=job)
        self._executions[job.job_id] = execution

        publish_sync(
            Event(
                name="job.submitted",
                payload={
                    "job_id": job.job_id,
                    "job_type": job.job_type.value,
                },
                correlation_id=job.correlation_id,
            )
        )

        loop = asyncio.get_running_loop()
        future = loop.run_in_executor(
            self._pool,
            self._run_job,
            job,
            runner,
            execution,
        )
        execution.future = future

    def _run_job(
        self,
        job: Job,
        runner: JobRunner,
        execution: JobExecution,
    ) -> Any:
        """Run a job in a worker thread."""
        execution.thread_id = threading.get_ident()
        execution.started_at = datetime.now(timezone.utc)

        job.start()

        publish_sync(
            Event(
                name="job.started",
                payload={"job_id": job.job_id},
                correlation_id=job.correlation_id,
            )
        )

        _executor_logger.info(
            "job_started",
            extra={
                "job_id": job.job_id,
                "job_type": job.job_type.value,
                "correlation_id": job.correlation_id,
            },
        )

        start = time.perf_counter()

        try:
            result = runner(job, self)

            elapsed = (time.perf_counter() - start) * 1000
            job.succeed({"result": result} if result else None)

            publish_sync(
                Event(
                    name="job.completed",
                    payload={
                        "job_id": job.job_id,
                        "status": "succeeded",
                        "duration_ms": elapsed,
                    },
                    correlation_id=job.correlation_id,
                )
            )

            _executor_logger.info(
                "job_completed",
                extra={
                    "job_id": job.job_id,
                    "status": "succeeded",
                    "duration_ms": elapsed,
                    "correlation_id": job.correlation_id,
                },
            )

            return result

        except asyncio.CancelledError:
            job.cancel()
            publish_sync(
                Event(
                    name="job.cancelled",
                    payload={"job_id": job.job_id},
                    correlation_id=job.correlation_id,
                )
            )
            _executor_logger.info(
                "job_cancelled",
                extra={
                    "job_id": job.job_id,
                    "correlation_id": job.correlation_id,
                },
            )
            raise

        except Exception as e:
            elapsed = (time.perf_counter() - start) * 1000
            job.fail("Job execution failed")

            publish_sync(
                Event(
                    name="job.failed",
                    payload={
                        "job_id": job.job_id,
                        "error": "Job execution failed",
                        "duration_ms": elapsed,
                    },
                    correlation_id=job.correlation_id,
                )
            )

            _executor_logger.exception(
                "job_failed",
                extra={
                    "job_id": job.job_id,
                    "error": str(e),
                    "duration_ms": elapsed,
                    "correlation_id": job.correlation_id,
                },
            )

            raise

        finally:
            self._executions.pop(job.job_id, None)

    def cancel(self, job_id: str, *, force: bool = False) -> bool:
        """Cancel a running job."""
        execution = self._executions.get(job_id)
        if not execution:
            return False

        execution.job.cancel()

        if execution.future and not execution.future.done():
            cancelled = execution.future.cancel()
            if cancelled:
                return True

        if force and self._config.enable_thread_injection_cancel:
            return self._inject_cancel(execution)

        return False

    def _inject_cancel(self, execution: JobExecution) -> bool:
        """Attempt to cancel via thread exception injection."""
        if not execution.thread_id:
            return False

        try:
            import ctypes

            res = ctypes.pythonapi.PyThreadState_SetAsyncExc(
                ctypes.c_long(execution.thread_id),
                ctypes.py_object(asyncio.CancelledError),
            )
            return res == 1
        except Exception:
            return False

    def get_status(self, job_id: str) -> Optional[JobStatus]:
        """Get the status of a job."""
        execution = self._executions.get(job_id)
        if execution:
            return execution.job.status
        return None

    def get_active_jobs(self) -> list[str]:
        """Get IDs of all active jobs."""
        return list(self._executions.keys())

    async def shutdown(self, wait: bool = True) -> None:
        """Shutdown the executor."""
        self._shutdown = True

        if wait:
            futures = [
                e.future for e in self._executions.values()
                if e.future and not e.future.done()
            ]
            if futures:
                await asyncio.gather(*futures, return_exceptions=True)

        self._pool.shutdown(wait=wait)

_executor: Optional[JobExecutor] = None

def get_executor() -> JobExecutor:
    """Get or create the global job executor."""
    global _executor
    if _executor is None:
        _executor = JobExecutor()
    return _executor

# Scheduler (from orchestration/scheduler.py)

_scheduler_logger = logging.getLogger("neura.orchestration.scheduler")

class ScheduleRepository(Protocol):
    """Protocol for schedule storage."""

    def find_due(self, now: Optional[datetime] = None) -> List[Schedule]: ...

    def save(self, schedule: Schedule) -> Schedule: ...

class JobSubmitter(Protocol):
    """Protocol for job submission."""

    async def submit(self, job: Job) -> None: ...

class Scheduler:
    """Scheduler for recurring jobs."""

    def __init__(
        self,
        schedule_repo: ScheduleRepository,
        job_submitter: JobSubmitter,
        *,
        poll_interval_seconds: int = 60,
    ) -> None:
        self._repo = schedule_repo
        self._submitter = job_submitter
        self._poll_interval = max(poll_interval_seconds, 5)
        self._task: Optional[asyncio.Task] = None
        self._stop_event = asyncio.Event()
        self._inflight: set[str] = set()

    async def start(self) -> None:
        """Start the scheduler."""
        if self._task and not self._task.done():
            return

        self._stop_event.clear()
        self._task = asyncio.create_task(
            self._run_loop(),
            name="scheduler-loop",
        )

        _scheduler_logger.info("scheduler_started", extra={"event": "scheduler_started"})

        publish_sync(
            Event(
                name="scheduler.started",
                payload={"poll_interval": self._poll_interval},
            )
        )

    async def stop(self) -> None:
        """Stop the scheduler."""
        if not self._task:
            return

        self._stop_event.set()
        self._task.cancel()

        try:
            await self._task
        except asyncio.CancelledError:
            pass
        finally:
            self._task = None

        _scheduler_logger.info("scheduler_stopped", extra={"event": "scheduler_stopped"})

        publish_sync(Event(name="scheduler.stopped", payload={}))

    async def _run_loop(self) -> None:
        """Main scheduler loop."""
        try:
            while not self._stop_event.is_set():
                try:
                    await self._check_and_dispatch()
                except Exception:
                    _scheduler_logger.exception(
                        "scheduler_tick_failed",
                        extra={"event": "scheduler_tick_failed"},
                    )

                try:
                    await asyncio.wait_for(
                        self._stop_event.wait(),
                        timeout=self._poll_interval,
                    )
                except asyncio.TimeoutError:
                    continue
        except asyncio.CancelledError:
            raise

    async def _check_and_dispatch(self) -> None:
        """Check for due schedules and dispatch jobs."""
        now = datetime.now(timezone.utc)
        due_schedules = self._repo.find_due(now)

        for schedule in due_schedules:
            if not schedule.active:
                continue

            if schedule.schedule_id in self._inflight:
                continue

            self._inflight.add(schedule.schedule_id)

            try:
                await self._dispatch_schedule(schedule)
            except Exception:
                _scheduler_logger.exception(
                    "schedule_dispatch_failed",
                    extra={
                        "schedule_id": schedule.schedule_id,
                        "event": "schedule_dispatch_failed",
                    },
                )
            finally:
                self._inflight.discard(schedule.schedule_id)

    async def _dispatch_schedule(self, schedule: Schedule) -> None:
        """Create and submit a job for a schedule."""
        correlation_id = f"sched-{schedule.schedule_id}-{int(datetime.now(timezone.utc).timestamp())}"

        job = Job.create(
            job_type=JobType.REPORT_GENERATION,
            template_id=schedule.template_id,
            template_name=schedule.template_name,
            template_kind=schedule.template_kind,
            connection_id=schedule.connection_id,
            schedule_id=schedule.schedule_id,
            correlation_id=correlation_id,
            steps=[
                JobStep(name="dataLoad", label="Load database"),
                JobStep(name="contractCheck", label="Prepare contract"),
                JobStep(name="renderPdf", label="Render PDF"),
                JobStep(name="finalize", label="Finalize"),
            ],
            meta={
                "start_date": schedule.start_date,
                "end_date": schedule.end_date,
                "schedule_name": schedule.name,
                "docx": schedule.docx,
                "xlsx": schedule.xlsx,
            },
        )

        _scheduler_logger.info(
            "schedule_job_created",
            extra={
                "schedule_id": schedule.schedule_id,
                "job_id": job.job_id,
                "correlation_id": correlation_id,
                "event": "schedule_job_created",
            },
        )

        publish_sync(
            Event(
                name="schedule.triggered",
                payload={
                    "schedule_id": schedule.schedule_id,
                    "job_id": job.job_id,
                },
                correlation_id=correlation_id,
            )
        )

        try:
            await self._submitter.submit(job)
            schedule.record_run("triggered")
            self._repo.save(schedule)
        except Exception as e:
            _scheduler_logger.exception("Schedule trigger failed for %s", schedule.schedule_id)
            schedule.record_run("failed", error="Schedule trigger failed")
            self._repo.save(schedule)
            raise

    @property
    def is_running(self) -> bool:
        """Check if scheduler is running."""
        return self._task is not None and not self._task.done()

# Worker Pool (from orchestration/worker.py)

_worker_logger = logging.getLogger("neura.orchestration.worker")

@dataclass
class WorkerPoolConfig:
    """Configuration for worker pool."""

    num_workers: int = 4
    queue_size: int = 100
    shutdown_timeout_seconds: float = 30.0

class WorkerPool:
    """Pool of workers for processing jobs."""

    def __init__(
        self,
        executor: JobExecutor,
        config: Optional[WorkerPoolConfig] = None,
    ) -> None:
        self._executor = executor
        self._config = config or WorkerPoolConfig()
        self._queue: asyncio.Queue = asyncio.Queue(maxsize=self._config.queue_size)
        self._workers: list[asyncio.Task] = []
        self._shutdown = False

    async def start(self) -> None:
        """Start the worker pool."""
        if self._workers:
            return

        self._shutdown = False

        for i in range(self._config.num_workers):
            worker = asyncio.create_task(
                self._worker_loop(i),
                name=f"worker-{i}",
            )
            self._workers.append(worker)

        _worker_logger.info(
            "worker_pool_started",
            extra={
                "event": "worker_pool_started",
                "num_workers": len(self._workers),
            },
        )

    async def stop(self) -> None:
        """Stop the worker pool gracefully."""
        self._shutdown = True

        for worker in self._workers:
            worker.cancel()

        if self._workers:
            await asyncio.gather(*self._workers, return_exceptions=True)

        self._workers = []

        _worker_logger.info("worker_pool_stopped", extra={"event": "worker_pool_stopped"})

    async def submit(self, job) -> None:
        """Submit a job to the pool."""
        if self._shutdown:
            raise RuntimeError("Worker pool is shutting down")

        await self._queue.put(job)

        _worker_logger.debug(
            "job_queued",
            extra={
                "job_id": job.job_id,
                "queue_size": self._queue.qsize(),
            },
        )

    async def _worker_loop(self, worker_id: int) -> None:
        """Worker loop that processes jobs."""
        _worker_logger.debug(
            "worker_started",
            extra={"worker_id": worker_id, "event": "worker_started"},
        )

        try:
            while not self._shutdown:
                try:
                    job = await asyncio.wait_for(
                        self._queue.get(),
                        timeout=1.0,
                    )
                except asyncio.TimeoutError:
                    continue

                try:
                    await self._executor.submit(job)
                except Exception:
                    _worker_logger.exception(
                        "worker_job_failed",
                        extra={
                            "worker_id": worker_id,
                            "job_id": job.job_id,
                        },
                    )
                finally:
                    self._queue.task_done()

        except asyncio.CancelledError:
            pass
        finally:
            _worker_logger.debug(
                "worker_stopped",
                extra={"worker_id": worker_id, "event": "worker_stopped"},
            )

    @property
    def queue_size(self) -> int:
        """Get current queue size."""
        return self._queue.qsize()

    @property
    def is_running(self) -> bool:
        """Check if pool is running."""
        return bool(self._workers) and not self._shutdown

__all__ = [
    "ExecutorConfig", "JobExecution", "JobRunner", "JobExecutor", "get_executor",
    "ScheduleRepository", "JobSubmitter", "Scheduler",
    "WorkerPoolConfig", "WorkerPool",
]

"""Engine adapters and pipelines — all adapter implementations + pipeline logic."""

# ADAPTERS

"""Database adapters for querying data sources.

Consolidated from databases/base.py, databases/sqlite.py.
The dataframes/ subpackage remains separate.
"""

import logging
import threading
import time
from abc import ABC, abstractmethod
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Protocol, Sequence

import pandas as pd

from backend.engine_all import ConnectionTest, SchemaInfo, TableInfo

logger = logging.getLogger("neura.adapters.sqlite")

# Base (from databases/base.py)

@dataclass(frozen=True)
class QueryResult:
    """Result of a database query."""

    columns: List[str]
    rows: List[List[Any]]
    row_count: int
    query: str
    execution_time_ms: float

    def to_dataframe(self) -> pd.DataFrame:
        return pd.DataFrame(self.rows, columns=self.columns)

    def to_dicts(self) -> List[Dict[str, Any]]:
        return [dict(zip(self.columns, row)) for row in self.rows]

class DataSource(Protocol):
    """Interface for data source access."""

    @property
    def path(self) -> Path: ...
    def test_connection(self) -> ConnectionTest: ...
    def discover_schema(self) -> SchemaInfo: ...
    def execute_query(self, query: str, parameters: Optional[Sequence[Any]] = None) -> QueryResult: ...
    def stream_query(self, query: str, parameters: Optional[Sequence[Any]] = None, batch_size: int = 1000) -> Iterator[QueryResult]: ...
    def get_table_columns(self, table_name: str) -> List[str]: ...
    def get_row_count(self, table_name: str) -> int: ...

class SchemaDiscovery(ABC):
    """Abstract base for schema discovery implementations."""

    @abstractmethod
    def discover_tables(self) -> List[TableInfo]:
        pass

    @abstractmethod
    def discover_columns(self, table_name: str) -> List[str]:
        pass

    @abstractmethod
    def build_catalog(self, tables: List[TableInfo]) -> List[str]:
        pass

    def discover(self) -> SchemaInfo:
        tables = self.discover_tables()
        catalog = self.build_catalog(tables)
        return SchemaInfo(tables=tables, catalog=catalog)

# SQLite / DataFrame Adapter (from databases/sqlite.py)

class DataFrameConnectionPool:
    """DataFrame-based connection pool using sqlite_shim."""

    def __init__(self, path: Path, readonly: bool = True, pool_size: int = 5, max_overflow: int = 10) -> None:
        self._path = path.resolve()
        self._readonly = readonly
        self._pool_size = pool_size
        self._max_overflow = max_overflow
        self._lock = threading.Lock()
        self._closed = False
        self._active_count = 0

        # (same file) # get_loader
        self._loader = get_loader(self._path)
        self._loader.frames()
        logger.info(f"Loaded {len(self._loader.table_names())} tables into DataFrames for {self._path}")

    def _create_connection(self) -> sqlite_shim.DataFrameConnection:
        return sqlite_shim.connect(str(self._path))

    @contextmanager
    def acquire(self):
        if self._closed:
            raise RuntimeError("Connection pool is closed")
        with self._lock:
            self._active_count += 1
        conn = self._create_connection()
        try:
            yield conn
        except Exception:
            raise
        finally:
            try:
                conn.close()
            except Exception as e:
                logger.debug("Connection close failed: %s", e)
            with self._lock:
                self._active_count = max(0, self._active_count - 1)

    def close(self) -> None:
        self._closed = True

    def status(self) -> Dict[str, Any]:
        return {
            "pool_size": self._pool_size,
            "active_connections": self._active_count,
            "tables_loaded": len(self._loader.table_names()),
            "closed": self._closed,
        }

SQLiteConnectionPool = DataFrameConnectionPool

class DataFrameSchemaDiscovery(SchemaDiscovery):
    """Schema discovery using DataFrames."""

    def __init__(self, loader: SQLiteDataFrameLoader) -> None:
        self._loader = loader
        self._table_info_cache: Dict[str, List[dict]] = {}

    def discover_tables(self) -> List[TableInfo]:
        table_names = self._loader.table_names()
        if not table_names:
            return []
        self._prefetch_table_info(table_names)
        row_counts = self._batch_get_row_counts(table_names)
        tables = []
        for table_name in table_names:
            columns, pk = self._get_cached_table_info(table_name)
            tables.append(TableInfo(name=table_name, columns=columns, row_count=row_counts.get(table_name, 0), primary_key=pk))
        return tables

    def _prefetch_table_info(self, table_names: List[str]) -> None:
        for table_name in table_names:
            try:
                info = self._loader.pragma_table_info(table_name)
                self._table_info_cache[table_name] = info
            except Exception as e:
                logger.debug(f"Failed to get table info for {table_name}: {e}")
                self._table_info_cache[table_name] = []

    def _get_cached_table_info(self, table_name: str) -> tuple[List[str], Optional[str]]:
        info = self._table_info_cache.get(table_name, [])
        columns = [row.get("name", "") for row in info]
        pk = None
        for row in info:
            if row.get("pk"):
                pk = row.get("name")
                break
        return columns, pk

    def _batch_get_row_counts(self, table_names: List[str]) -> Dict[str, int]:
        counts = {}
        for table_name in table_names:
            try:
                frame = self._loader.frame(table_name)
                counts[table_name] = len(frame)
            except Exception as e:
                logger.debug(f"Failed to get row count for {table_name}: {e}")
                counts[table_name] = 0
        return counts

    def discover_columns(self, table_name: str) -> List[str]:
        if table_name in self._table_info_cache:
            return [row.get("name", "") for row in self._table_info_cache[table_name]]
        try:
            frame = self._loader.frame(table_name)
            return list(frame.columns)
        except Exception:
            return []

    def build_catalog(self, tables: List[TableInfo]) -> List[str]:
        catalog = []
        for table in tables:
            for column in table.columns:
                catalog.append(f"{table.name}.{column}")
        return catalog

    def _get_row_count(self, table_name: str) -> int:
        try:
            frame = self._loader.frame(table_name)
            return len(frame)
        except Exception as e:
            logger.debug(f"Failed to get row count for {table_name}: {e}")
            return 0

    def _get_primary_key(self, table_name: str) -> Optional[str]:
        try:
            info = self._loader.pragma_table_info(table_name)
            for row in info:
                if row.get("pk"):
                    return row.get("name")
            return None
        except Exception as e:
            logger.debug(f"Failed to get primary key for {table_name}: {e}")
            return None

SQLiteSchemaDiscovery = DataFrameSchemaDiscovery

class DataFrameDataSource:
    """DataFrame-based implementation of DataSource interface."""

    def __init__(self, path: Path, *, readonly: bool = True, use_pool: bool = False, pool_size: int = 5) -> None:
        self._path = path.resolve()
        self._readonly = readonly
        self._use_pool = use_pool
        self._pool: Optional[DataFrameConnectionPool] = None

        # (same file) # get_loader
        self._loader = get_loader(self._path)
        if eager_load_enabled():
            self._loader.frames()

        if use_pool:
            self._pool = DataFrameConnectionPool(path=self._path, readonly=readonly, pool_size=pool_size)

        logger.info(f"DataFrameDataSource initialized with {len(self._loader.table_names())} tables")

    @property
    def path(self) -> Path:
        return self._path

    @contextmanager
    def _get_connection(self):
        if self._pool is not None:
            with self._pool.acquire() as conn:
                yield conn
            return
        conn = sqlite_shim.connect(str(self._path))
        try:
            yield conn
        finally:
            conn.close()

    def test_connection(self) -> ConnectionTest:
        start = time.perf_counter()
        try:
            table_names = self._loader.table_names()
            table_count = len(table_names)
            if table_names:
                _ = self._loader.frame(table_names[0])
            latency = (time.perf_counter() - start) * 1000
            return ConnectionTest(success=True, latency_ms=latency, tested_at=datetime.now(timezone.utc), table_count=table_count)
        except Exception as e:
            latency = (time.perf_counter() - start) * 1000
            logger.exception("SQLite connection test failed")
            return ConnectionTest(success=False, latency_ms=latency, error="Connection test failed", tested_at=datetime.now(timezone.utc))

    def discover_schema(self) -> SchemaInfo:
        discovery = DataFrameSchemaDiscovery(self._loader)
        return discovery.discover()

    def execute_query(self, query: str, parameters: Optional[Sequence[Any]] = None) -> QueryResult:
        start = time.perf_counter()
        with self._get_connection() as conn:
            conn.row_factory = sqlite_shim.Row
            cursor = conn.execute(query, parameters or ())
            rows_raw = cursor.fetchall()
            columns = list(rows_raw[0].keys()) if rows_raw else []
            rows = [list(row) for row in rows_raw]
            execution_time = (time.perf_counter() - start) * 1000
        return QueryResult(columns=columns, rows=rows, row_count=len(rows), query=query, execution_time_ms=execution_time)

    def stream_query(self, query: str, parameters: Optional[Sequence[Any]] = None, batch_size: int = 1000) -> Iterator[QueryResult]:
        start = time.perf_counter()
        with self._get_connection() as conn:
            conn.row_factory = sqlite_shim.Row
            cursor = conn.execute(query, parameters or ())
            first_batch = cursor.fetchmany(batch_size)
            columns = list(first_batch[0].keys()) if first_batch else []
            batch = first_batch
            while batch:
                execution_time = (time.perf_counter() - start) * 1000
                yield QueryResult(columns=columns, rows=[list(row) for row in batch], row_count=len(batch), query=query, execution_time_ms=execution_time)
                batch = cursor.fetchmany(batch_size)

    def get_table_columns(self, table_name: str) -> List[str]:
        try:
            frame = self._loader.frame(table_name)
            return list(frame.columns)
        except Exception:
            return []

    def get_row_count(self, table_name: str) -> int:
        try:
            frame = self._loader.frame(table_name)
            return len(frame)
        except Exception:
            return 0

    def close(self) -> None:
        if self._pool:
            self._pool.close()
            self._pool = None

    def pool_status(self) -> Optional[Dict[str, Any]]:
        if self._pool:
            return self._pool.status()
        return {"tables_loaded": len(self._loader.table_names()), "pooling_enabled": False}

    def __del__(self) -> None:
        self.close()

SQLiteDataSource = DataFrameDataSource

__all__ = [
    "QueryResult", "DataSource", "SchemaDiscovery",
    "DataFrameConnectionPool", "SQLiteConnectionPool",
    "DataFrameSchemaDiscovery", "SQLiteSchemaDiscovery",
    "DataFrameDataSource", "SQLiteDataSource",
]

# --- databases/sqlite.py (re-exports already defined above) ---

# --- databases/dataframes/store.py ---
"""
Centralized DataFrame Store for connection-based DataFrame management.

This module provides a singleton store that:
1. Automatically loads all database tables as DataFrames when a connection is used
2. Caches DataFrames per connection_id for reuse
3. Provides a unified interface to access DataFrames across all services
4. Eliminates direct database access after initial load
"""

import logging
import os
import threading
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd

logger = logging.getLogger("neura.dataframes.store")

class DataFrameStore:
    """
    Centralized store for managing DataFrames by connection.

    All database interactions go through this store, ensuring:
    - Tables are loaded once as DataFrames and cached
    - All queries run against in-memory DataFrames via DuckDB
    - No direct database connections after initial load
    """

    _instance: Optional["DataFrameStore"] = None
    _lock = threading.Lock()

    def __new__(cls) -> "DataFrameStore":
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(self) -> None:
        if self._initialized:
            return
        self._loaders: Dict[str, SQLiteDataFrameLoader] = {}
        self._frames_cache: Dict[str, Dict[str, pd.DataFrame]] = {}
        self._db_paths: Dict[str, Path] = {}
        self._query_engines: Dict[str, DuckDBDataFrameQuery] = {}
        self._store_lock = threading.Lock()
        self._initialized = True
        logger.info("DataFrameStore initialized")

    def register_connection(self, connection_id: str, db_path: Path) -> None:
        """
        Register a database connection and load all tables as DataFrames.

        This should be called when a connection is established. All tables
        will be loaded into memory as DataFrames for subsequent queries.
        """
        db_path = Path(db_path).resolve()

        with self._store_lock:
            # Check if already registered with same path
            existing_path = self._db_paths.get(connection_id)
            if existing_path and existing_path == db_path:
                # Already registered, check if file modified
                loader = self._loaders.get(connection_id)
                if loader:
                    current_mtime = os.path.getmtime(db_path) if db_path.exists() else 0.0
                    if loader._mtime == current_mtime:
                        logger.debug(f"Connection {connection_id} already registered and up to date")
                        return

            logger.info(f"Loading DataFrames for connection {connection_id} from {db_path}")
            loader = get_loader(db_path)
            eager = eager_load_enabled()
            frames = loader.frames() if eager else {}

            # Close existing query engine if any
            existing_engine = self._query_engines.get(connection_id)
            if existing_engine:
                try:
                    existing_engine.close()
                except Exception as e:
                    logger.debug("Engine close failed: %s", e)

            # Store everything
            self._loaders[connection_id] = loader
            self._frames_cache[connection_id] = frames if frames else {}
            self._db_paths[connection_id] = db_path
            self._query_engines[connection_id] = DuckDBDataFrameQuery(frames, loader=loader)

            logger.info(
                f"Loaded {len(frames)} tables for connection {connection_id}: {list(frames.keys())}"
                if frames else f"Registered connection {connection_id} for lazy DataFrame loading"
            )

    def get_loader(self, connection_id: str) -> Optional[SQLiteDataFrameLoader]:
        """Get the loader for a connection."""
        with self._store_lock:
            return self._loaders.get(connection_id)

    def get_frames(self, connection_id: str) -> Dict[str, pd.DataFrame]:
        """Get all DataFrames for a connection."""
        with self._store_lock:
            return self._frames_cache.get(connection_id, {})

    def get_frame(self, connection_id: str, table_name: str) -> Optional[pd.DataFrame]:
        """Get a specific DataFrame for a connection."""
        frames = self.get_frames(connection_id)
        return frames.get(table_name)

    def get_query_engine(self, connection_id: str) -> Optional[DuckDBDataFrameQuery]:
        """Get the DuckDB query engine for a connection."""
        with self._store_lock:
            return self._query_engines.get(connection_id)

    def execute_query(
        self,
        connection_id: str,
        sql: str,
        params: Any = None,
    ) -> pd.DataFrame:
        """
        Execute a SQL query against the DataFrames for a connection.

        Returns results as a DataFrame.
        """
        engine = self.get_query_engine(connection_id)
        if engine is None:
            raise ValueError(f"Connection {connection_id} not registered in DataFrameStore")
        return engine.execute(sql, params)

    def execute_query_to_dicts(
        self,
        connection_id: str,
        sql: str,
        params: Any = None,
    ) -> list[dict[str, Any]]:
        """
        Execute a SQL query and return results as list of dicts.

        This is the preferred method for API responses.
        """
        df = self.execute_query(connection_id, sql, params)
        return df.to_dict("records")

    def get_table_names(self, connection_id: str) -> list[str]:
        """Get list of table names for a connection."""
        loader = self.get_loader(connection_id)
        if loader is None:
            return []
        return loader.table_names()

    def get_table_info(self, connection_id: str, table_name: str) -> list[dict[str, Any]]:
        """Get PRAGMA table_info equivalent for a table."""
        loader = self.get_loader(connection_id)
        if loader is None:
            return []
        return loader.pragma_table_info(table_name)

    def get_foreign_keys(self, connection_id: str, table_name: str) -> list[dict[str, Any]]:
        """Get foreign keys for a table."""
        loader = self.get_loader(connection_id)
        if loader is None:
            return []
        return loader.foreign_keys(table_name)

    def invalidate_connection(self, connection_id: str) -> None:
        """
        Invalidate cached DataFrames for a connection.

        Call this when the underlying database changes.
        """
        with self._store_lock:
            engine = self._query_engines.pop(connection_id, None)
            if engine:
                try:
                    engine.close()
                except Exception as e:
                    logger.debug("Engine close failed: %s", e)
            self._loaders.pop(connection_id, None)
            self._frames_cache.pop(connection_id, None)
            self._db_paths.pop(connection_id, None)
            logger.info(f"Invalidated DataFrames for connection {connection_id}")

    def is_registered(self, connection_id: str) -> bool:
        """Check if a connection is registered."""
        with self._store_lock:
            return connection_id in self._loaders

    def get_db_path(self, connection_id: str) -> Optional[Path]:
        """Get the database path for a connection."""
        with self._store_lock:
            return self._db_paths.get(connection_id)

    def status(self) -> dict[str, Any]:
        """Get store status."""
        with self._store_lock:
            return {
                "connections": list(self._loaders.keys()),
                "total_connections": len(self._loaders),
                "tables_per_connection": {
                    conn_id: len(frames)
                    for conn_id, frames in self._frames_cache.items()
                },
            }

    def clear(self) -> None:
        """Clear all cached DataFrames."""
        with self._store_lock:
            for engine in self._query_engines.values():
                try:
                    engine.close()
                except Exception as e:
                    logger.debug("Engine close failed: %s", e)
            self._loaders.clear()
            self._frames_cache.clear()
            self._db_paths.clear()
            self._query_engines.clear()
            logger.info("DataFrameStore cleared")

# Singleton instance
dataframe_store = DataFrameStore()

def get_dataframe_store() -> DataFrameStore:
    """Get the singleton DataFrameStore instance."""
    return dataframe_store

def ensure_connection_loaded(connection_id: str, db_path: Path) -> DataFrameStore:
    """
    Ensure a connection's DataFrames are loaded in the store.

    Convenience function that registers the connection if needed and returns the store.
    """
    store = get_dataframe_store()
    if not store.is_registered(connection_id):
        store.register_connection(connection_id, db_path)
    return store

# --- databases/dataframes/sqlite_loader.py ---

import os
import re
import sqlite3
import threading
from pathlib import Path
from typing import Any, Mapping, Sequence

import duckdb
import pandas as pd

class SQLiteDataFrameLoader:
    """Load SQLite tables into cached pandas DataFrames."""

    def __init__(self, db_path: Path):
        self.db_path = Path(db_path)
        self._table_names: list[str] | None = None
        self._frames: dict[str, pd.DataFrame] = {}
        self._lock = threading.Lock()
        self._mtime = os.path.getmtime(self.db_path) if self.db_path.exists() else 0.0
        self._table_info_cache: dict[str, list[dict[str, Any]]] = {}
        self._foreign_keys_cache: dict[str, list[dict[str, Any]]] = {}

    def table_names(self) -> list[str]:
        """Return a cached list of user tables in the database."""
        with self._lock:
            if self._table_names is not None:
                return list(self._table_names)

            with sqlite3.connect(str(self.db_path)) as con:
                cur = con.execute(
                    "SELECT name FROM sqlite_master "
                    "WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name;"
                )
                tables = [str(row[0]) for row in cur.fetchall() if row and row[0]]
            self._table_names = tables
            return list(self._table_names)

    def _assert_table(self, table_name: str) -> str:
        clean = str(table_name or "").strip()
        if not clean:
            raise ValueError("table_name must be a non-empty string")
        if clean not in self.table_names():
            raise RuntimeError(f"Table {clean!r} not found in {self.db_path}")
        return clean

    def frame(self, table_name: str) -> pd.DataFrame:
        """
        Return the cached DataFrame for `table_name`, loading it from disk
        on first access. Callers should treat the returned DataFrame as
        read-only because it is shared across consumers.
        """
        clean = self._assert_table(table_name)
        with self._lock:
            cached = self._frames.get(clean)
            if cached is not None:
                return cached
        df = self._read_table(clean)
        with self._lock:
            self._frames[clean] = df
        return df

    def frames(self) -> dict[str, pd.DataFrame]:
        """Eagerly load and return all user tables as DataFrames."""
        for name in self.table_names():
            self.frame(name)
        with self._lock:
            return dict(self._frames)

    def _read_table(self, table_name: str) -> pd.DataFrame:
        quoted = table_name.replace('"', '""')
        try:
            with sqlite3.connect(str(self.db_path)) as con:
                df = pd.read_sql_query(f'SELECT rowid AS "__rowid__", * FROM "{quoted}"', con)
        except Exception as exc:  # pragma: no cover - surfaced to caller
            raise RuntimeError(f"Failed loading table {table_name!r} into DataFrame: {exc}") from exc
        if "__rowid__" in df.columns:
            rowid_series = df["__rowid__"].copy()
            if "rowid" not in df.columns:
                df.insert(0, "rowid", rowid_series)
        return df

    def column_type(self, table_name: str, column_name: str) -> str:
        table = self.frame(table_name)
        if column_name not in table.columns:
            return ""
        series = table[column_name]
        if pd.api.types.is_datetime64_any_dtype(series):
            return "DATETIME"
        if pd.api.types.is_integer_dtype(series):
            return "INTEGER"
        if pd.api.types.is_float_dtype(series):
            return "REAL"
        if pd.api.types.is_bool_dtype(series):
            return "INTEGER"
        return "TEXT"

    def table_info(self, table_name: str) -> list[tuple[str, str]]:
        table = self.frame(table_name)
        return [(col, str(table[col].dtype)) for col in table.columns]

    def _load_table_metadata(
        self, table_name: str
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        clean = self._assert_table(table_name)
        with self._lock:
            cached_info = self._table_info_cache.get(clean)
            cached_fks = self._foreign_keys_cache.get(clean)
            if cached_info is not None and cached_fks is not None:
                return cached_info, cached_fks

        quoted = clean.replace("'", "''")
        info_rows: list[dict[str, Any]] = []
        fk_rows: list[dict[str, Any]] = []
        try:
            with sqlite3.connect(str(self.db_path)) as con:
                cur = con.execute(f"PRAGMA table_info('{quoted}')")
                info_rows = [
                    {
                        "cid": int(row[0]),
                        "name": str(row[1]),
                        "type": str(row[2] or ""),
                        "notnull": int(row[3] or 0),
                        "dflt_value": row[4],
                        "pk": int(row[5] or 0),
                    }
                    for row in cur.fetchall()
                ]
                cur = con.execute(f"PRAGMA foreign_key_list('{quoted}')")
                fk_rows = [
                    {
                        "id": int(row[0]),
                        "seq": int(row[1]),
                        "table": str(row[2] or ""),
                        "from": str(row[3] or ""),
                        "to": str(row[4] or ""),
                        "on_update": str(row[5] or ""),
                        "on_delete": str(row[6] or ""),
                        "match": str(row[7] or ""),
                    }
                    for row in cur.fetchall()
                ]
        except Exception as exc:  # pragma: no cover - surfaced to caller
            raise RuntimeError(f"Failed loading metadata for table {clean!r}: {exc}") from exc

        with self._lock:
            self._table_info_cache[clean] = info_rows
            self._foreign_keys_cache[clean] = fk_rows
        return info_rows, fk_rows

    def pragma_table_info(self, table_name: str) -> list[dict[str, Any]]:
        info, _ = self._load_table_metadata(table_name)
        return list(info)

    def foreign_keys(self, table_name: str) -> list[dict[str, Any]]:
        _, fks = self._load_table_metadata(table_name)
        return list(fks)

_LOADER_CACHE: dict[str, SQLiteDataFrameLoader] = {}
_LOADER_CACHE_LOCK = threading.Lock()

def get_loader(db_path: Path) -> SQLiteDataFrameLoader:
    key = str(Path(db_path).resolve())
    with _LOADER_CACHE_LOCK:
        loader = _LOADER_CACHE.get(key)
        mtime = os.path.getmtime(key) if os.path.exists(key) else 0.0
        if loader is None or loader._mtime != mtime:
            loader = SQLiteDataFrameLoader(Path(key))
            loader._mtime = mtime
            _LOADER_CACHE[key] = loader
    return loader

_PARAM_PATTERN = re.compile(r":([A-Za-z_][A-Za-z0-9_]*)")
_SQLITE_DATETIME_RE = re.compile(r"(?i)(?<!sqlite_)\bdatetime\s*\(")
_SQLITE_STRFTIME_RE = re.compile(r"(?i)(?<!sqlite_)\bstrftime\s*\(")

def _normalize_params(sql: str, params: Any | None) -> tuple[str, Sequence[Any]]:
    if params is None:
        return sql, ()
    if isinstance(params, Mapping):
        ordered: list[Any] = []

        def _repl(match: re.Match[str]) -> str:
            name = match.group(1)
            if name not in params:
                raise KeyError(f"Missing SQL parameter: {name}")
            ordered.append(params[name])
            return "?"

        prepared = _PARAM_PATTERN.sub(_repl, sql)
        return prepared, tuple(ordered)
    if isinstance(params, (list, tuple)):
        return sql, tuple(params)
    return sql, (params,)

def _rewrite_sql(sql: str) -> str:
    """Apply lightweight rewrites so legacy SQLite SQL runs in DuckDB."""
    updated = _SQLITE_DATETIME_RE.sub("sqlite_datetime(", sql)
    updated = _SQLITE_STRFTIME_RE.sub("sqlite_strftime(", updated)
    return updated

_MISSING_TABLE_RE = re.compile(r'(?:Table|Relation) with name "?(?P<table>[^"\s]+)"? does not exist', re.I)

class DuckDBDataFrameQuery:
    """
    Execute SQL statements against in-memory pandas DataFrames by delegating
    evaluation to DuckDB. This keeps the contract SQL assets unchanged while
    avoiding a live SQLite database dependency.
    """

    def __init__(self, frames: Mapping[str, pd.DataFrame], loader: SQLiteDataFrameLoader | None = None):
        self._conn = duckdb.connect(database=":memory:")
        self._loader = loader
        self._registered: set[str] = set()
        self._register_frames(frames)
        self._register_sqlite_macros()

    def _register_frames(self, frames: Mapping[str, pd.DataFrame]) -> None:
        for name, frame in frames.items():
            if frame is None:
                continue
            self._register_frame(name, frame)

    def _register_frame(self, name: str, frame: pd.DataFrame) -> None:
        if not name or name in self._registered:
            return
        self._conn.register(name, frame)
        self._registered.add(name)

    def _try_register_missing_table(self, exc: Exception) -> bool:
        if self._loader is None:
            return False
        match = _MISSING_TABLE_RE.search(str(exc))
        if not match:
            return False
        table = match.group("table")
        if not table or table in self._registered:
            return False
        try:
            frame = self._loader.frame(table)
        except Exception:
            return False
        self._register_frame(table, frame)
        return True

    def _register_sqlite_macros(self) -> None:
        """Install lightweight SQLite compatibility macros for DuckDB execution."""
        self._conn.execute(
            """
            CREATE MACRO IF NOT EXISTS sqlite_datetime(x) AS (
                CASE
                    WHEN x IS NULL THEN NULL
                    WHEN LOWER(CAST(x AS VARCHAR)) = 'now' THEN CURRENT_TIMESTAMP
                    WHEN TRY_CAST(x AS DOUBLE) IS NOT NULL THEN TO_TIMESTAMP(CAST(x AS DOUBLE))
                    ELSE TRY_CAST(x AS TIMESTAMP)
                END
            )
            """
        )
        self._conn.execute(
            """
            CREATE MACRO IF NOT EXISTS sqlite_strftime(fmt, value, modifier := NULL) AS (
                CASE
                    WHEN value IS NULL THEN NULL
                    ELSE STRFTIME(sqlite_datetime(value), fmt)
                END
            )
            """
        )
        # Alias the macros to the SQLite function names so legacy SQL keeps working.
        self._conn.execute("CREATE MACRO IF NOT EXISTS datetime(x) AS sqlite_datetime(x)")
        self._conn.execute(
            "CREATE MACRO IF NOT EXISTS strftime(fmt, value, modifier := NULL) AS sqlite_strftime(fmt, value, modifier)"
        )

    def execute(self, sql: str, params: Any | None = None) -> pd.DataFrame:
        prepared_sql, ordered_params = _normalize_params(sql, params)
        rewritten_sql = _rewrite_sql(prepared_sql)
        attempts = 0
        while True:
            try:
                result = self._conn.execute(rewritten_sql, ordered_params)
                return result.fetchdf()
            except duckdb.Error as exc:  # pragma: no cover - surfaced to caller
                if attempts < 5 and self._try_register_missing_table(exc):
                    attempts += 1
                    continue
                raise RuntimeError(f"DuckDB execution failed: {exc}") from exc

    def close(self) -> None:
        self._conn.close()

def eager_load_enabled() -> bool:
    """Return True if DataFrame eager loading is enabled."""
    flag = os.getenv("NEURA_DATAFRAME_EAGER_LOAD", "false")
    return str(flag).strip().lower() in {"1", "true", "yes"}

# --- databases/dataframes/sqlite_shim.py ---

import re
from pathlib import Path
from typing import Any, Sequence

import pandas as pd

class Error(Exception):
    """Base error matching sqlite3.Error."""

class OperationalError(Error):
    """Raised when SQL execution fails."""

class Row:
    """Lightweight sqlite3.Row replacement supporting dict + index access."""

    __slots__ = ("_columns", "_values", "_mapping")

    def __init__(self, columns: Sequence[str], values: Sequence[Any]):
        self._columns = list(columns)
        self._values = tuple(values)
        self._mapping = {col: value for col, value in zip(self._columns, self._values)}

    def __getitem__(self, key: int | str) -> Any:
        if isinstance(key, int):
            return self._values[key]
        return self._mapping[key]

    def keys(self) -> list[str]:
        return list(self._columns)

    def __iter__(self):
        return iter(self._values)

    def __len__(self) -> int:
        return len(self._values)

    def __repr__(self) -> str:
        items = ", ".join(f"{col}={self._mapping[col]!r}" for col in self._columns)
        return f"Row({items})"

def _apply_row_factory(columns: list[str], values: Sequence[Any], factory: Any | None):
    if factory is None:
        return tuple(values)
    if factory is Row:
        return Row(columns, values)
    return factory(columns, values)

class DataFrameCursor:
    def __init__(self, connection: "DataFrameConnection"):
        self.connection = connection
        self._df: pd.DataFrame | None = None
        self._columns: list[str] = []
        self._pos = 0

    def execute(self, sql: str, params: Any | None = None) -> "DataFrameCursor":
        meta_df = self._try_meta_query(sql)
        if meta_df is not None:
            self._df = meta_df
            self._columns = list(meta_df.columns)
            self._pos = 0
            return self
        try:
            df = self.connection._query.execute(sql, params)
        except Exception as exc:  # pragma: no cover - propagated to caller
            raise OperationalError(str(exc)) from exc
        self._df = df
        self._columns = list(df.columns)
        self._pos = 0
        return self

    def _try_meta_query(self, sql: str) -> pd.DataFrame | None:
        sql_clean = (sql or "").strip()
        pragma_match = re.match(r"(?is)^PRAGMA\s+table_info\(['\"]?(?P<table>[^'\")]+)['\"]?\)\s*;?$", sql_clean)
        if pragma_match:
            table_name = pragma_match.group("table")
            try:
                rows = self.connection._loader.pragma_table_info(table_name)
            except Exception:
                rows = []
            data = [
                (
                    int(row.get("cid", 0)),
                    str(row.get("name", "")),
                    str(row.get("type", "")),
                    int(row.get("notnull", 0)),
                    row.get("dflt_value"),
                    int(row.get("pk", 0)),
                )
                for row in rows
            ]
            return pd.DataFrame(data, columns=["cid", "name", "type", "notnull", "dflt_value", "pk"])

        fk_match = re.match(r"(?is)^PRAGMA\s+foreign_key_list\(['\"]?(?P<table>[^'\")]+)['\"]?\)\s*;?$", sql_clean)
        if fk_match:
            table_name = fk_match.group("table")
            try:
                rows = self.connection._loader.foreign_keys(table_name)
            except Exception:
                rows = []
            data = [
                (
                    int(row.get("id", 0)),
                    int(row.get("seq", 0)),
                    str(row.get("table", "")),
                    str(row.get("from", "")),
                    str(row.get("to", "")),
                    str(row.get("on_update", "")),
                    str(row.get("on_delete", "")),
                    str(row.get("match", "")),
                )
                for row in rows
            ]
            return pd.DataFrame(
                data, columns=["id", "seq", "table", "from", "to", "on_update", "on_delete", "match"]
            )

        if "sqlite_master" in sql_clean.lower():
            names = [name for name in self.connection._loader.table_names() if not name.lower().startswith("sqlite_")]
            data = [(name,) for name in sorted(names)]
            return pd.DataFrame(data, columns=["name"])
        return None

    def fetchone(self):
        if self._df is None:
            return None
        if self._pos >= len(self._df):
            return None
        row = self._df.iloc[self._pos].tolist()
        self._pos += 1
        return _apply_row_factory(self._columns, row, self.connection.row_factory)

    def fetchall(self):
        rows = []
        while True:
            row = self.fetchone()
            if row is None:
                break
            rows.append(row)
        return rows

    def fetchmany(self, size: int = 1):
        if size is None:
            size = 1
        rows = []
        for _ in range(max(0, int(size))):
            row = self.fetchone()
            if row is None:
                break
            rows.append(row)
        return rows

    def __iter__(self):
        while True:
            row = self.fetchone()
            if row is None:
                return
            yield row

class DataFrameConnection:
    def __init__(self, db_path: Path):
        self.db_path = Path(db_path)
        self._loader = get_loader(self.db_path)
        if eager_load_enabled():
            frames = self._loader.frames()
        else:
            frames = {}
        self._query = DuckDBDataFrameQuery(frames, loader=self._loader)
        self.row_factory: Any | None = None

    def cursor(self) -> DataFrameCursor:
        return DataFrameCursor(self)

    def execute(self, sql: str, params: Any | None = None) -> DataFrameCursor:
        return self.cursor().execute(sql, params)

    def close(self) -> None:
        self._query.close()

    def commit(self) -> None:  # pragma: no cover - compatibility no-op
        return None

    def rollback(self) -> None:  # pragma: no cover - compatibility no-op
        return None

    def __enter__(self) -> "DataFrameConnection":
        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        self.close()

def connect(db_path: str | Path, **_kwargs) -> DataFrameConnection:
    """sqlite3.connect-compatible entrypoint backed by pandas DataFrames."""
    return DataFrameConnection(Path(db_path))

# EXTRACTION (merged from extraction/__init__.py)

"""Document extraction adapters.

Consolidated from extraction/base.py, extraction/excel.py, extraction/pdf.py.
"""

import logging
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Protocol

# Base (from extraction/base.py)

@dataclass
class ExtractedTable:
    """A table extracted from a document."""

    page_number: int
    table_index: int
    headers: List[str]
    rows: List[List[Any]]
    confidence: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def row_count(self) -> int:
        return len(self.rows)

    @property
    def column_count(self) -> int:
        return len(self.headers)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "page_number": self.page_number,
            "table_index": self.table_index,
            "headers": self.headers,
            "rows": self.rows,
            "confidence": self.confidence,
            "row_count": self.row_count,
            "column_count": self.column_count,
        }

@dataclass
class ExtractedText:
    """Text extracted from a document."""

    page_number: int
    content: str
    bbox: Optional[tuple] = None
    font_info: Optional[Dict[str, Any]] = None

@dataclass
class ExtractionResult:
    """Result of document extraction."""

    source_path: Path
    page_count: int
    tables: List[ExtractedTable] = field(default_factory=list)
    text_blocks: List[ExtractedText] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)
    extraction_time_ms: float = 0.0

    @property
    def success(self) -> bool:
        return len(self.errors) == 0

    @property
    def table_count(self) -> int:
        return len(self.tables)

class Extractor(Protocol):
    """Interface for document extractors."""

    def extract(self, path: Path) -> ExtractionResult: ...
    def extract_tables(self, path: Path) -> List[ExtractedTable]: ...
    def supports(self, path: Path) -> bool: ...

class BaseExtractor(ABC):
    """Abstract base for extractors with common functionality."""

    @abstractmethod
    def extract(self, path: Path) -> ExtractionResult:
        pass

    @abstractmethod
    def extract_tables(self, path: Path) -> List[ExtractedTable]:
        pass

    @abstractmethod
    def supports(self, path: Path) -> bool:
        pass

    def _validate_path(self, path: Path) -> None:
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if not path.is_file():
            raise ValueError(f"Path is not a file: {path}")

# Excel Extractor (from extraction/excel.py)

_excel_logger = logging.getLogger("neura.adapters.extraction.excel")

class ExcelExtractor(BaseExtractor):
    """Extract data from Excel documents."""

    SUPPORTED_EXTENSIONS = {".xlsx", ".xls", ".xlsm", ".xlsb"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def extract(self, path: Path) -> ExtractionResult:
        self._validate_path(path)
        start = time.perf_counter()
        errors: List[str] = []
        tables: List[ExtractedTable] = []
        metadata: Dict[str, Any] = {}

        try:
            tables = self.extract_tables(path)
            metadata = self._get_workbook_metadata(path)
        except Exception as e:
            errors.append(f"Excel extraction failed: {e}")
            _excel_logger.exception("excel_extraction_failed")

        extraction_time = (time.perf_counter() - start) * 1000

        return ExtractionResult(
            source_path=path,
            page_count=len(tables),
            tables=tables,
            text_blocks=[],
            metadata=metadata,
            errors=errors,
            extraction_time_ms=extraction_time,
        )

    def extract_tables(self, path: Path) -> List[ExtractedTable]:
        self._validate_path(path)

        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required. Install with: pip install openpyxl")

        tables: List[ExtractedTable] = []
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            table = self._extract_sheet_table(sheet, sheet_idx, sheet_name)
            if table:
                tables.append(table)

        wb.close()
        return tables

    def _extract_sheet_table(self, sheet, sheet_idx: int, sheet_name: str) -> ExtractedTable | None:
        rows_data: List[List[Any]] = []

        for row in sheet.iter_rows(values_only=True):
            row_values = [self._cell_to_string(cell) for cell in row]
            if any(v for v in row_values):
                rows_data.append(row_values)

        if not rows_data:
            return None

        headers = rows_data[0]
        rows = rows_data[1:] if len(rows_data) > 1 else []

        return ExtractedTable(
            page_number=sheet_idx + 1,
            table_index=0,
            headers=headers,
            rows=rows,
            confidence=0.9,
            metadata={"sheet_name": sheet_name},
        )

    def _cell_to_string(self, cell: Any) -> str:
        if cell is None:
            return ""
        if isinstance(cell, (int, float)):
            if isinstance(cell, float) and cell.is_integer():
                return str(int(cell))
            return str(cell)
        return str(cell)

    def _get_workbook_metadata(self, path: Path) -> Dict[str, Any]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True)
            metadata = {
                "sheet_names": wb.sheetnames,
                "sheet_count": len(wb.sheetnames),
            }
            if wb.properties:
                metadata.update({
                    "title": wb.properties.title,
                    "creator": wb.properties.creator,
                    "created": str(wb.properties.created) if wb.properties.created else None,
                    "modified": str(wb.properties.modified) if wb.properties.modified else None,
                })
            wb.close()
            return metadata
        except Exception:
            return {}

# PDF Extractor (from extraction/pdf.py)

_pdf_logger = logging.getLogger("neura.adapters.extraction.pdf")

class PDFExtractor(BaseExtractor):
    """Extract data from PDF documents."""

    def __init__(self, *, prefer_backend: str = "pdfplumber", extract_text: bool = True) -> None:
        self._prefer_backend = prefer_backend
        self._extract_text = extract_text

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".pdf"

    def extract(self, path: Path) -> ExtractionResult:
        self._validate_path(path)
        start = time.perf_counter()
        errors: List[str] = []
        tables: List[ExtractedTable] = []
        text_blocks: List[ExtractedText] = []
        page_count = 0
        metadata: Dict[str, Any] = {}

        try:
            page_count, metadata = self._get_pdf_info(path)
        except Exception as e:
            errors.append(f"Failed to read PDF info: {e}")

        try:
            tables = self.extract_tables(path)
        except Exception as e:
            errors.append(f"Table extraction failed: {e}")

        if self._extract_text:
            try:
                text_blocks = self._extract_text_blocks(path)
            except Exception as e:
                errors.append(f"Text extraction failed: {e}")

        extraction_time = (time.perf_counter() - start) * 1000

        return ExtractionResult(
            source_path=path,
            page_count=page_count,
            tables=tables,
            text_blocks=text_blocks,
            metadata=metadata,
            errors=errors,
            extraction_time_ms=extraction_time,
        )

    def extract_tables(self, path: Path) -> List[ExtractedTable]:
        self._validate_path(path)

        if self._prefer_backend == "pdfplumber":
            return self._extract_with_pdfplumber(path)
        elif self._prefer_backend == "tabula":
            return self._extract_with_tabula(path)
        else:
            try:
                return self._extract_with_pdfplumber(path)
            except Exception:
                return self._extract_with_tabula(path)

    def _get_pdf_info(self, path: Path) -> tuple[int, Dict[str, Any]]:
        try:
            import fitz
            doc = fitz.open(str(path))
            try:
                page_count = len(doc)
                metadata = dict(doc.metadata) if doc.metadata else {}
                return page_count, metadata
            finally:
                doc.close()
        except ImportError:
            pass

        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return len(pdf.pages), {}
        except ImportError:
            pass

        return 0, {}

    def _extract_with_pdfplumber(self, path: Path) -> List[ExtractedTable]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required. Install with: pip install pdfplumber")

        tables: List[ExtractedTable] = []

        with pdfplumber.open(path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if not table or len(table) < 2:
                        continue

                    headers = [str(h or "") for h in table[0]]
                    rows = [[str(c or "") for c in row] for row in table[1:]]

                    tables.append(
                        ExtractedTable(
                            page_number=page_idx + 1,
                            table_index=table_idx,
                            headers=headers,
                            rows=rows,
                            confidence=0.8,
                        )
                    )

        return tables

    def _extract_with_tabula(self, path: Path) -> List[ExtractedTable]:
        try:
            import tabula
        except ImportError:
            raise ImportError("tabula-py is required. Install with: pip install tabula-py")

        tables: List[ExtractedTable] = []
        dfs = tabula.read_pdf(str(path), pages="all", multiple_tables=True)

        for table_idx, df in enumerate(dfs):
            if df.empty:
                continue

            headers = [str(c) for c in df.columns.tolist()]
            rows = df.fillna("").astype(str).values.tolist()

            tables.append(
                ExtractedTable(
                    page_number=1,
                    table_index=table_idx,
                    headers=headers,
                    rows=rows,
                    confidence=0.7,
                )
            )

        return tables

    def _extract_text_blocks(self, path: Path) -> List[ExtractedText]:
        try:
            import fitz
        except ImportError:
            _pdf_logger.warning("PyMuPDF not available for text extraction")
            return []

        text_blocks: List[ExtractedText] = []
        doc = fitz.open(str(path))

        for page_idx, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") == 0:
                    text = ""
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text += span.get("text", "")
                        text += "\n"

                    if text.strip():
                        text_blocks.append(
                            ExtractedText(
                                page_number=page_idx + 1,
                                content=text.strip(),
                                bbox=tuple(block.get("bbox", [])),
                            )
                        )

        doc.close()
        return text_blocks

__all__ = [
    "ExtractedTable", "ExtractedText", "ExtractionResult",
    "Extractor", "BaseExtractor",
    "ExcelExtractor", "PDFExtractor",
]

# LLM (merged from llm/__init__.py)

"""LLM adapters for AI-powered features.

Consolidated from llm/base.py, llm/openai.py.
"""

import logging
import os
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Protocol

from backend.engine_all import ExternalServiceError

_logger = logging.getLogger("neura.adapters.llm.openai")

# Base (from llm/base.py)

class LLMRole(str, Enum):
    """Role in a conversation."""

    SYSTEM = "system"
    USER = "user"
    ASSISTANT = "assistant"

@dataclass(frozen=True)
class LLMMessage:
    """A message in an LLM conversation."""

    role: LLMRole
    content: str

    def to_dict(self) -> Dict[str, str]:
        return {"role": self.role.value, "content": self.content}

@dataclass
class LLMResponse:
    """Response from an LLM call."""

    content: str
    model: str
    usage: Dict[str, int] = field(default_factory=dict)
    finish_reason: Optional[str] = None
    raw_response: Optional[Any] = None

    @property
    def prompt_tokens(self) -> int:
        return self.usage.get("prompt_tokens", 0)

    @property
    def completion_tokens(self) -> int:
        return self.usage.get("completion_tokens", 0)

    @property
    def total_tokens(self) -> int:
        return self.usage.get("total_tokens", 0)

class LLMClient(Protocol):
    """Interface for LLM clients."""

    def complete(
        self, messages: List[LLMMessage], *,
        model: Optional[str] = None, temperature: float = 0.0,
        max_tokens: Optional[int] = None, json_mode: bool = False,
    ) -> LLMResponse: ...

    async def complete_async(
        self, messages: List[LLMMessage], *,
        model: Optional[str] = None, temperature: float = 0.0,
        max_tokens: Optional[int] = None, json_mode: bool = False,
    ) -> LLMResponse: ...

class BaseLLMClient(ABC):
    """Abstract base for LLM clients with common functionality."""

    def __init__(self, *, default_model: str, max_retries: int = 3, timeout_seconds: float = 60.0) -> None:
        self._default_model = default_model
        self._max_retries = max_retries
        self._timeout = timeout_seconds

    @abstractmethod
    def complete(self, messages: List[LLMMessage], *, model: Optional[str] = None, temperature: float = 0.0, max_tokens: Optional[int] = None, json_mode: bool = False) -> LLMResponse:
        pass

    @abstractmethod
    async def complete_async(self, messages: List[LLMMessage], *, model: Optional[str] = None, temperature: float = 0.0, max_tokens: Optional[int] = None, json_mode: bool = False) -> LLMResponse:
        pass

    def _prepare_messages(self, messages: List[LLMMessage]) -> List[Dict[str, str]]:
        return [m.to_dict() for m in messages]

@dataclass
class PromptTemplate:
    """A reusable prompt template with variable substitution."""

    template: str
    system_prompt: Optional[str] = None
    variables: List[str] = field(default_factory=list)

    def render(self, **kwargs: Any) -> List[LLMMessage]:
        messages = []
        if self.system_prompt:
            messages.append(LLMMessage(role=LLMRole.SYSTEM, content=self.system_prompt))
        content = self.template
        for var in self.variables:
            if var in kwargs:
                content = content.replace(f"{{{{{var}}}}}", str(kwargs[var]))
        messages.append(LLMMessage(role=LLMRole.USER, content=content))
        return messages

# OpenAI Client (from llm/openai.py)

_FORCE_GPT5 = os.getenv("NEURA_FORCE_GPT5", "true").lower() in {"1", "true", "yes"}

class OpenAIClient(BaseLLMClient):
    """OpenAI API client implementation."""

    def __init__(
        self, *, api_key: Optional[str] = None, default_model: str = "gpt-5",
        max_retries: int = 3, timeout_seconds: float = 60.0,
        base_url: Optional[str] = None,
    ) -> None:
        super().__init__(default_model=default_model, max_retries=max_retries, timeout_seconds=timeout_seconds)
        self._api_key = api_key or os.getenv("OPENAI_API_KEY")
        self._base_url = base_url
        self._client = None

    def _get_client(self):
        if self._client is None:
            try:
                from openai import OpenAI
            except ImportError:
                raise ImportError("openai package is required. Install with: pip install openai")

            if not self._api_key:
                raise ValueError("OpenAI API key is required. Set OPENAI_API_KEY environment variable.")
            if not self._api_key.startswith(("sk-", "sess-")):
                _logger.warning("OpenAI API key may be invalid (expected 'sk-' or 'sess-' prefix)", extra={"event": "api_key_format_warning"})

            kwargs: Dict[str, Any] = {"api_key": self._api_key, "timeout": self._timeout, "max_retries": self._max_retries}
            if self._base_url:
                kwargs["base_url"] = self._base_url
            self._client = OpenAI(**kwargs)
        return self._client

    def complete(self, messages: List[LLMMessage], *, model: Optional[str] = None, temperature: float = 0.0, max_tokens: Optional[int] = None, json_mode: bool = False) -> LLMResponse:
        client = self._get_client()
        model_name = _force_gpt5(model or self._default_model)

        kwargs: Dict[str, Any] = {"model": model_name, "messages": self._prepare_messages(messages), "temperature": temperature}
        if max_tokens:
            kwargs["max_tokens"] = max_tokens
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}

        start = time.perf_counter()
        try:
            if _use_responses_model(model_name):
                payload = _prepare_responses_payload(kwargs)
                try:
                    response = client.responses.create(**payload)
                except AttributeError as exc:
                    raise ExternalServiceError(message="OpenAI Responses API is required for gpt-5. Upgrade the openai package to >=1.0.0.", service="openai", cause=exc)
                content = _response_output_text(response)
                usage = _response_usage(response)
                elapsed = (time.perf_counter() - start) * 1000
                _logger.info("llm_completion_success", extra={"event": "llm_completion_success", "model": model_name, "elapsed_ms": elapsed, "tokens": usage.get("total_tokens", 0), "endpoint": "responses"})
                return LLMResponse(content=content, model=response.model if hasattr(response, "model") else model_name, usage=usage, finish_reason="stop", raw_response=response)

            response = client.chat.completions.create(**kwargs)
            elapsed = (time.perf_counter() - start) * 1000
            _logger.info("llm_completion_success", extra={"event": "llm_completion_success", "model": model_name, "elapsed_ms": elapsed, "tokens": response.usage.total_tokens if response.usage else 0})
            return LLMResponse(
                content=response.choices[0].message.content or "",
                model=response.model,
                usage={"prompt_tokens": response.usage.prompt_tokens if response.usage else 0, "completion_tokens": response.usage.completion_tokens if response.usage else 0, "total_tokens": response.usage.total_tokens if response.usage else 0},
                finish_reason=response.choices[0].finish_reason,
                raw_response=response,
            )
        except Exception as e:
            elapsed = (time.perf_counter() - start) * 1000
            _logger.exception("llm_completion_failed", extra={"event": "llm_completion_failed", "model": model_name, "elapsed_ms": elapsed, "error": str(e)})
            raise ExternalServiceError(message="OpenAI API call failed", service="openai", cause=e)

    async def complete_async(self, messages: List[LLMMessage], *, model: Optional[str] = None, temperature: float = 0.0, max_tokens: Optional[int] = None, json_mode: bool = False) -> LLMResponse:
        try:
            from openai import AsyncOpenAI
        except ImportError:
            raise ImportError("openai package is required. Install with: pip install openai")

        kwargs: Dict[str, Any] = {"api_key": self._api_key, "timeout": self._timeout, "max_retries": self._max_retries}
        if self._base_url:
            kwargs["base_url"] = self._base_url

        client = AsyncOpenAI(**kwargs)
        model_name = _force_gpt5(model or self._default_model)

        request_kwargs: Dict[str, Any] = {"model": model_name, "messages": self._prepare_messages(messages), "temperature": temperature}
        if max_tokens:
            request_kwargs["max_tokens"] = max_tokens
        if json_mode:
            request_kwargs["response_format"] = {"type": "json_object"}

        start = time.perf_counter()
        try:
            if _use_responses_model(model_name):
                payload = _prepare_responses_payload(request_kwargs)
                try:
                    response = await client.responses.create(**payload)
                except AttributeError as exc:
                    raise ExternalServiceError(message="OpenAI Responses API is required for gpt-5. Upgrade the openai package to >=1.0.0.", service="openai", cause=exc)
                content = _response_output_text(response)
                usage = _response_usage(response)
                elapsed = (time.perf_counter() - start) * 1000
                _logger.info("llm_completion_success_async", extra={"event": "llm_completion_success_async", "model": model_name, "elapsed_ms": elapsed, "tokens": usage.get("total_tokens", 0), "endpoint": "responses"})
                return LLMResponse(content=content, model=response.model if hasattr(response, "model") else model_name, usage=usage, finish_reason="stop", raw_response=response)

            response = await client.chat.completions.create(**request_kwargs)
            elapsed = (time.perf_counter() - start) * 1000
            _logger.info("llm_completion_success_async", extra={"event": "llm_completion_success_async", "model": model_name, "elapsed_ms": elapsed, "tokens": response.usage.total_tokens if response.usage else 0})
            return LLMResponse(
                content=response.choices[0].message.content or "",
                model=response.model,
                usage={"prompt_tokens": response.usage.prompt_tokens if response.usage else 0, "completion_tokens": response.usage.completion_tokens if response.usage else 0, "total_tokens": response.usage.total_tokens if response.usage else 0},
                finish_reason=response.choices[0].finish_reason,
                raw_response=response,
            )
        except Exception as e:
            elapsed = (time.perf_counter() - start) * 1000
            _logger.exception("llm_completion_failed_async", extra={"event": "llm_completion_failed_async", "model": model_name, "elapsed_ms": elapsed, "error": str(e)})
            raise ExternalServiceError(message="OpenAI API call failed", service="openai", cause=e)

# Helper functions (from llm/openai.py)

def _use_responses_model(model_name: Optional[str]) -> bool:
    force = os.getenv("OPENAI_USE_RESPONSES", "").lower() in {"1", "true", "yes"}
    return force or str(model_name or "").lower().startswith("gpt-5")

def _force_gpt5(model_name: Optional[str]) -> str:
    if not _FORCE_GPT5:
        return str(model_name or "gpt-5").strip() or "gpt-5"
    normalized = str(model_name or "").strip()
    if normalized.lower().startswith("gpt-5"):
        return normalized
    if normalized:
        _logger.warning("llm_model_overridden", extra={"event": "llm_model_overridden", "requested": normalized, "forced": "gpt-5"})
    return "gpt-5"

def _prepare_responses_payload(request_kwargs: Dict[str, Any]) -> Dict[str, Any]:
    payload = dict(request_kwargs)
    messages = payload.pop("messages", [])
    payload["input"] = _messages_to_responses_input(messages)
    if "max_tokens" in payload and "max_output_tokens" not in payload:
        payload["max_output_tokens"] = payload.pop("max_tokens")
    return payload

def _messages_to_responses_input(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    converted: List[Dict[str, Any]] = []
    for message in messages:
        if not isinstance(message, dict):
            continue
        role = message.get("role") or "user"
        content = message.get("content", "")
        if isinstance(content, list):
            parts: List[Dict[str, Any]] = []
            for part in content:
                if isinstance(part, dict):
                    part_type = part.get("type")
                    if part_type == "text":
                        parts.append({"type": "input_text", "text": part.get("text", "")})
                        continue
                    if part_type == "image_url":
                        image_url = part.get("image_url")
                        if isinstance(image_url, dict):
                            image_url = image_url.get("url") or image_url.get("image_url")
                        parts.append({"type": "input_image", "image_url": image_url})
                        continue
                    parts.append(part)
                else:
                    parts.append({"type": "input_text", "text": str(part)})
            content = parts
        converted.append({"role": role, "content": content})
    return converted

def _response_output_text(response: Any) -> str:
    if isinstance(response, dict):
        output_text = response.get("output_text")
        if isinstance(output_text, str) and output_text.strip():
            return output_text
        output = response.get("output")
    else:
        output_text = getattr(response, "output_text", None)
        if isinstance(output_text, str) and output_text.strip():
            return output_text
        output = getattr(response, "output", None)

    if isinstance(output, list):
        texts: List[str] = []
        for item in output:
            if isinstance(item, dict):
                item_type = item.get("type")
                content = item.get("content") or []
            else:
                item_type = getattr(item, "type", None)
                content = getattr(item, "content", None) or []
            if item_type != "message":
                continue
            for segment in content:
                if isinstance(segment, dict):
                    seg_type = segment.get("type")
                    text = segment.get("text")
                else:
                    seg_type = getattr(segment, "type", None)
                    text = getattr(segment, "text", None)
                if seg_type in {"output_text", "text"} and isinstance(text, str):
                    texts.append(text)
        if texts:
            return "\n".join(texts)
    return ""

def _response_usage(response: Any) -> Dict[str, int]:
    usage = response.get("usage") if isinstance(response, dict) else getattr(response, "usage", None)
    if usage is None:
        return {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    if isinstance(usage, dict):
        input_tokens = usage.get("input_tokens") or usage.get("prompt_tokens") or 0
        output_tokens = usage.get("output_tokens") or usage.get("completion_tokens") or 0
    else:
        input_tokens = getattr(usage, "input_tokens", None)
        if input_tokens is None:
            input_tokens = getattr(usage, "prompt_tokens", 0)
        output_tokens = getattr(usage, "output_tokens", None)
        if output_tokens is None:
            output_tokens = getattr(usage, "completion_tokens", 0)
    total_tokens = int(input_tokens or 0) + int(output_tokens or 0)
    return {"prompt_tokens": int(input_tokens or 0), "completion_tokens": int(output_tokens or 0), "total_tokens": int(total_tokens)}

__all__ = [
    "LLMRole", "LLMMessage", "LLMResponse", "LLMClient", "BaseLLMClient",
    "PromptTemplate", "OpenAIClient",
]

# PERSISTENCE (merged from persistence/__init__.py)

"""Persistence adapters - storage for domain entities.

Consolidated from persistence/base.py, persistence/repositories.py.
"""

from abc import ABC, abstractmethod
from datetime import datetime
from typing import Generic, List, Optional, Protocol, TypeVar, runtime_checkable

from backend.engine_all import (
    Template, Connection, Job, JobStatus, Schedule, Report,
)

# Base (from persistence/base.py)

T = TypeVar("T")
ID = TypeVar("ID")

@runtime_checkable
class Repository(Protocol[T, ID]):
    """Generic repository interface for CRUD operations."""

    def get(self, id: ID) -> Optional[T]: ...
    def get_all(self) -> List[T]: ...
    def save(self, entity: T) -> T: ...
    def delete(self, id: ID) -> bool: ...
    def exists(self, id: ID) -> bool: ...

class UnitOfWork(ABC):
    """Unit of Work pattern for transaction management."""

    @abstractmethod
    def __enter__(self) -> UnitOfWork: ...

    @abstractmethod
    def __exit__(self, exc_type, exc_val, exc_tb) -> None: ...

    @abstractmethod
    def commit(self) -> None: ...

    @abstractmethod
    def rollback(self) -> None: ...

class BaseRepository(Generic[T, ID], ABC):
    """Abstract base class for repositories with common functionality."""

    @abstractmethod
    def get(self, id: ID) -> Optional[T]:
        pass

    @abstractmethod
    def get_all(self) -> List[T]:
        pass

    @abstractmethod
    def save(self, entity: T) -> T:
        pass

    @abstractmethod
    def delete(self, id: ID) -> bool:
        pass

    def exists(self, id: ID) -> bool:
        return self.get(id) is not None

    def count(self) -> int:
        return len(self.get_all())

# Repository Interfaces (from persistence/repositories.py)

class TemplateRepository(Protocol):
    """Repository for Template entities."""

    def get(self, template_id: str) -> Optional[Template]: ...
    def get_all(self) -> List[Template]: ...
    def save(self, template: Template) -> Template: ...
    def delete(self, template_id: str) -> bool: ...
    def exists(self, template_id: str) -> bool: ...
    def find_by_kind(self, kind: str) -> List[Template]: ...
    def find_by_name(self, name: str) -> Optional[Template]: ...

class ConnectionRepository(Protocol):
    """Repository for Connection entities."""

    def get(self, connection_id: str) -> Optional[Connection]: ...
    def get_all(self) -> List[Connection]: ...
    def save(self, connection: Connection) -> Connection: ...
    def delete(self, connection_id: str) -> bool: ...
    def exists(self, connection_id: str) -> bool: ...
    def find_by_name(self, name: str) -> Optional[Connection]: ...
    def get_default(self) -> Optional[Connection]: ...

class JobRepository(Protocol):
    """Repository for Job entities."""

    def get(self, job_id: str) -> Optional[Job]: ...
    def get_all(self) -> List[Job]: ...
    def save(self, job: Job) -> Job: ...
    def delete(self, job_id: str) -> bool: ...
    def find_by_status(self, status: JobStatus) -> List[Job]: ...
    def find_by_template(self, template_id: str) -> List[Job]: ...
    def find_active(self) -> List[Job]: ...
    def find_recent(self, limit: int = 50) -> List[Job]: ...

class ScheduleRepository(Protocol):
    """Repository for Schedule entities."""

    def get(self, schedule_id: str) -> Optional[Schedule]: ...
    def get_all(self) -> List[Schedule]: ...
    def save(self, schedule: Schedule) -> Schedule: ...
    def delete(self, schedule_id: str) -> bool: ...
    def find_active(self) -> List[Schedule]: ...
    def find_due(self, now: Optional[datetime] = None) -> List[Schedule]: ...
    def find_by_template(self, template_id: str) -> List[Schedule]: ...

class ReportRepository(Protocol):
    """Repository for Report entities (run history)."""

    def get(self, report_id: str) -> Optional[Report]: ...
    def get_all(self) -> List[Report]: ...
    def save(self, report: Report) -> Report: ...
    def delete(self, report_id: str) -> bool: ...
    def find_by_template(self, template_id: str, limit: int = 50) -> List[Report]: ...
    def find_by_connection(self, connection_id: str, limit: int = 50) -> List[Report]: ...
    def find_by_schedule(self, schedule_id: str, limit: int = 50) -> List[Report]: ...
    def find_recent(self, limit: int = 50) -> List[Report]: ...

__all__ = [
    "Repository", "UnitOfWork", "BaseRepository",
    "TemplateRepository", "ConnectionRepository", "JobRepository",
    "ScheduleRepository", "ReportRepository",
]

# RENDERING (merged from rendering/__init__.py)

"""Rendering adapters for document generation.

Consolidated from rendering/base.py, rendering/html.py, rendering/pdf.py,
rendering/docx.py, rendering/xlsx.py.
"""

import asyncio
import html as html_mod
import logging
import re
import time
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, List, Optional, Protocol

from backend.engine_all import OutputFormat

_base_logger = logging.getLogger("neura.adapters.rendering.base")
_html_logger = logging.getLogger("neura.adapters.rendering.html")
_pdf_logger = logging.getLogger("neura.adapters.rendering.pdf")
_docx_logger = logging.getLogger("neura.adapters.rendering.docx")
_xlsx_logger = logging.getLogger("neura.adapters.rendering.xlsx")

# Base (from rendering/base.py)

@dataclass
class RenderContext:
    """Context passed to renderers."""

    template_html: str
    data: Dict[str, Any]
    output_format: OutputFormat
    output_path: Path
    metadata: Dict[str, Any] = field(default_factory=dict)
    landscape: bool = False
    font_scale: Optional[float] = None
    page_size: str = "A4"
    margins: Optional[Dict[str, str]] = None

@dataclass(frozen=True)
class RenderResult:
    """Result of a render operation."""

    success: bool
    output_path: Optional[Path]
    format: OutputFormat
    size_bytes: int = 0
    error: Optional[str] = None
    warnings: list[str] = field(default_factory=list)
    render_time_ms: float = 0.0

class Renderer(Protocol):
    """Interface for document renderers."""

    @property
    def output_format(self) -> OutputFormat: ...

    def render(self, context: RenderContext) -> RenderResult: ...

    def supports(self, format: OutputFormat) -> bool: ...

class BaseRenderer(ABC):
    """Abstract base class for renderers with common functionality."""

    @property
    @abstractmethod
    def output_format(self) -> OutputFormat:
        pass

    @abstractmethod
    def render(self, context: RenderContext) -> RenderResult:
        pass

    def supports(self, format: OutputFormat) -> bool:
        return format == self.output_format

    def _ensure_output_dir(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)

    def _get_file_size(self, path: Path) -> int:
        try:
            return path.stat().st_size
        except Exception as e:
            _base_logger.debug("Failed to get file size: %s", e)
            return 0

# HTML Renderer (from rendering/html.py)

class HTMLRenderer(BaseRenderer):
    """Renderer that produces HTML output with token substitution."""

    @property
    def output_format(self) -> OutputFormat:
        return OutputFormat.HTML

    def render(self, context: RenderContext) -> RenderResult:
        start = time.perf_counter()
        warnings: List[str] = []

        try:
            self._ensure_output_dir(context.output_path)
            html = self._substitute_tokens(context.template_html, context.data, warnings)
            context.output_path.write_text(html, encoding="utf-8")

            render_time = (time.perf_counter() - start) * 1000
            return RenderResult(
                success=True,
                output_path=context.output_path,
                format=OutputFormat.HTML,
                size_bytes=self._get_file_size(context.output_path),
                warnings=warnings,
                render_time_ms=render_time,
            )
        except Exception as e:
            _html_logger.exception("html_render_failed")
            return RenderResult(
                success=False,
                output_path=None,
                format=OutputFormat.HTML,
                error="HTML rendering failed",
                render_time_ms=(time.perf_counter() - start) * 1000,
            )

    def _substitute_tokens(self, html: str, data: Dict[str, Any], warnings: List[str]) -> str:
        def replace_token(match: re.Match) -> str:
            token = match.group(1).strip()
            if token in data:
                value = data[token]
                if value is None:
                    return ""
                return html_mod.escape(str(value))
            else:
                warnings.append(f"Token '{token}' not found in data")
                return match.group(0)

        pattern = r"\{\{([^}]+)\}\}"
        return re.sub(pattern, replace_token, html)

class TokenEngine:
    """Engine for processing tokens in HTML templates."""

    SCALAR_PATTERN = re.compile(r"\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}")
    ROW_PATTERN = re.compile(
        r"<!--\s*BEGIN_ROW\s*-->(.*?)<!--\s*END_ROW\s*-->",
        re.DOTALL,
    )
    CONDITIONAL_PATTERN = re.compile(
        r"<!--\s*IF\s+(\w+)\s*-->(.*?)<!--\s*ENDIF\s*-->",
        re.DOTALL,
    )

    def __init__(self) -> None:
        self._missing_tokens: List[str] = []

    @property
    def missing_tokens(self) -> List[str]:
        return list(self._missing_tokens)

    def process(self, html: str, scalars: Dict[str, Any], rows: List[Dict[str, Any]], totals: Dict[str, Any]) -> str:
        self._missing_tokens = []
        html = self._process_conditionals(html, scalars)
        html = self._process_rows(html, rows)
        html = self._process_scalars(html, {**scalars, **totals})
        return html

    def _process_scalars(self, html: str, data: Dict[str, Any]) -> str:
        def replace(match: re.Match) -> str:
            token = match.group(1)
            if token in data:
                value = data[token]
                return "" if value is None else html_mod.escape(str(value))
            self._missing_tokens.append(token)
            return match.group(0)
        return self.SCALAR_PATTERN.sub(replace, html)

    def _process_rows(self, html: str, rows: List[Dict[str, Any]]) -> str:
        def expand_row(match: re.Match) -> str:
            template = match.group(1)
            expanded = []
            for i, row_data in enumerate(rows):
                row_html = template
                row_data = {**row_data, "ROWID": i + 1, "ROW_INDEX": i}
                row_html = self._process_scalars(row_html, row_data)
                expanded.append(row_html)
            return "".join(expanded)
        return self.ROW_PATTERN.sub(expand_row, html)

    def _process_conditionals(self, html: str, data: Dict[str, Any]) -> str:
        def evaluate(match: re.Match) -> str:
            condition = match.group(1)
            content = match.group(2)
            value = data.get(condition)
            if value:
                return content
            return ""
        return self.CONDITIONAL_PATTERN.sub(evaluate, html)

# PDF Renderer (from rendering/pdf.py)

class PDFRenderer(BaseRenderer):
    """Renderer that produces PDF output using Playwright."""

    def __init__(self, *, browser_type: str = "chromium", headless: bool = True) -> None:
        self._browser_type = browser_type
        self._headless = headless

    @property
    def output_format(self) -> OutputFormat:
        return OutputFormat.PDF

    def render(self, context: RenderContext) -> RenderResult:
        start = time.perf_counter()

        try:
            self._ensure_output_dir(context.output_path)

            try:
                loop = asyncio.get_running_loop()
            except RuntimeError:
                loop = None
            if loop and loop.is_running():
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                    pool.submit(asyncio.run, self._render_async(context)).result()
            else:
                asyncio.run(self._render_async(context))

            if not context.output_path.exists():
                return RenderResult(
                    success=False, output_path=None, format=OutputFormat.PDF,
                    error="PDF file was not created",
                    render_time_ms=(time.perf_counter() - start) * 1000,
                )

            render_time = (time.perf_counter() - start) * 1000
            return RenderResult(
                success=True, output_path=context.output_path,
                format=OutputFormat.PDF,
                size_bytes=self._get_file_size(context.output_path),
                render_time_ms=render_time,
            )
        except Exception as e:
            _pdf_logger.exception("pdf_render_failed")
            return RenderResult(
                success=False, output_path=None, format=OutputFormat.PDF,
                error="PDF rendering failed",
                render_time_ms=(time.perf_counter() - start) * 1000,
            )

    async def _render_async(self, context: RenderContext) -> None:
        try:
            from playwright.async_api import async_playwright
        except ImportError:
            raise ImportError(
                "Playwright is required for PDF rendering. "
                "Install with: pip install playwright && playwright install chromium"
            )

        async with async_playwright() as p:
            browser = await getattr(p, self._browser_type).launch(headless=self._headless)
            try:
                page = await browser.new_page()
                await page.set_content(context.template_html, wait_until="networkidle")

                pdf_options = {
                    "path": str(context.output_path),
                    "format": context.page_size,
                    "print_background": True,
                    "landscape": context.landscape,
                }

                if context.margins:
                    pdf_options["margin"] = context.margins

                await page.pdf(**pdf_options)
            finally:
                await browser.close()

class PDFRendererSync:
    """Synchronous wrapper for PDF rendering."""

    def __init__(self, renderer: Optional[PDFRenderer] = None) -> None:
        self._renderer = renderer or PDFRenderer()

    def render_from_html_file(self, html_path: Path, output_path: Path, *, landscape: bool = False, page_size: str = "A4") -> RenderResult:
        html_content = html_path.read_text(encoding="utf-8")
        context = RenderContext(template_html=html_content, data={}, output_format=OutputFormat.PDF, output_path=output_path, landscape=landscape, page_size=page_size)
        return self._renderer.render(context)

    def render_from_html_string(self, html: str, output_path: Path, *, landscape: bool = False, page_size: str = "A4") -> RenderResult:
        context = RenderContext(template_html=html, data={}, output_format=OutputFormat.PDF, output_path=output_path, landscape=landscape, page_size=page_size)
        return self._renderer.render(context)

# DOCX Renderer (from rendering/docx.py)

class DOCXRenderer(BaseRenderer):
    """Renderer that produces DOCX output from HTML."""

    def __init__(self, *, default_font_scale: float = 1.0) -> None:
        self._default_font_scale = default_font_scale

    @property
    def output_format(self) -> OutputFormat:
        return OutputFormat.DOCX

    def render(self, context: RenderContext) -> RenderResult:
        start = time.perf_counter()

        try:
            self._ensure_output_dir(context.output_path)

            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.section import WD_ORIENT
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX rendering. "
                    "Install with: pip install python-docx"
                )

            doc = Document()

            if context.landscape:
                for section in doc.sections:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    new_width = section.page_height
                    new_height = section.page_width
                    section.page_width = new_width
                    section.page_height = new_height

            self._html_to_docx(doc, context.template_html, context.font_scale)
            doc.save(str(context.output_path))

            render_time = (time.perf_counter() - start) * 1000
            return RenderResult(
                success=True, output_path=context.output_path,
                format=OutputFormat.DOCX,
                size_bytes=self._get_file_size(context.output_path),
                render_time_ms=render_time,
            )
        except Exception as e:
            _docx_logger.exception("docx_render_failed")
            return RenderResult(
                success=False, output_path=None, format=OutputFormat.DOCX,
                error=str(e),
                render_time_ms=(time.perf_counter() - start) * 1000,
            )

    def _html_to_docx(self, doc, html: str, font_scale: Optional[float]) -> None:
        try:
            from html2docx import html2docx
            html2docx(html, doc)
            return
        except ImportError:
            pass

        class SimpleHTMLParser(HTMLParser):
            def __init__(self, document):
                super().__init__()
                self.doc = document
                self.current_para = None
                self.in_table = False
                self.table_data = []
                self.current_row = []

            def handle_starttag(self, tag, attrs):
                if tag in ("p", "div"):
                    self.current_para = self.doc.add_paragraph()
                elif tag == "table":
                    self.in_table = True
                    self.table_data = []
                elif tag == "tr":
                    self.current_row = []
                elif tag == "br":
                    if self.current_para:
                        self.current_para.add_run("\n")
                elif tag in ("h1", "h2", "h3"):
                    level = int(tag[1])
                    self.current_para = self.doc.add_heading(level=level)

            def handle_endtag(self, tag):
                if tag == "table" and self.table_data:
                    self._create_table()
                    self.in_table = False
                elif tag == "tr" and self.current_row:
                    self.table_data.append(self.current_row)

            def handle_data(self, data):
                data = data.strip()
                if not data:
                    return
                if self.in_table:
                    self.current_row.append(data)
                elif self.current_para:
                    self.current_para.add_run(data)
                else:
                    self.current_para = self.doc.add_paragraph(data)

            def _create_table(self):
                if not self.table_data:
                    return
                rows = len(self.table_data)
                cols = max(len(row) for row in self.table_data)
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"
                for i, row_data in enumerate(self.table_data):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_data)

        parser = SimpleHTMLParser(doc)
        parser.feed(html)

def render_docx_from_html(html_path: Path, output_path: Path, *, landscape: bool = False, font_scale: Optional[float] = None) -> RenderResult:
    """Convenience function to render DOCX from HTML file."""
    html_content = html_path.read_text(encoding="utf-8")
    renderer = DOCXRenderer()
    context = RenderContext(template_html=html_content, data={}, output_format=OutputFormat.DOCX, output_path=output_path, landscape=landscape, font_scale=font_scale)
    return renderer.render(context)

# XLSX Renderer (from rendering/xlsx.py)

class XLSXRenderer(BaseRenderer):
    """Renderer that produces XLSX output from HTML tables."""

    @property
    def output_format(self) -> OutputFormat:
        return OutputFormat.XLSX

    def render(self, context: RenderContext) -> RenderResult:
        start = time.perf_counter()

        try:
            self._ensure_output_dir(context.output_path)

            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            except ImportError:
                raise ImportError("openpyxl is required for XLSX rendering. Install with: pip install openpyxl")

            wb = Workbook()
            ws = wb.active
            ws.title = "Report"

            tables = self._extract_tables_from_html(context.template_html)

            if tables:
                self._write_tables_to_worksheet(ws, tables)
            else:
                if context.data:
                    self._write_data_to_worksheet(ws, context.data)

            self._auto_fit_columns(ws)
            wb.save(str(context.output_path))

            render_time = (time.perf_counter() - start) * 1000
            return RenderResult(
                success=True, output_path=context.output_path,
                format=OutputFormat.XLSX,
                size_bytes=self._get_file_size(context.output_path),
                render_time_ms=render_time,
            )
        except Exception as e:
            _xlsx_logger.exception("xlsx_render_failed")
            return RenderResult(
                success=False, output_path=None, format=OutputFormat.XLSX,
                error=str(e),
                render_time_ms=(time.perf_counter() - start) * 1000,
            )

    def _extract_tables_from_html(self, html: str) -> List[List[List[str]]]:
        class TableParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.tables = []
                self.current_table = []
                self.current_row = []
                self.current_cell = ""
                self.in_table = False
                self.in_cell = False

            def handle_starttag(self, tag, attrs):
                if tag == "table":
                    self.in_table = True
                    self.current_table = []
                elif tag == "tr" and self.in_table:
                    self.current_row = []
                elif tag in ("td", "th") and self.in_table:
                    self.in_cell = True
                    self.current_cell = ""

            def handle_endtag(self, tag):
                if tag == "table":
                    if self.current_table:
                        self.tables.append(self.current_table)
                    self.in_table = False
                elif tag == "tr" and self.in_table:
                    if self.current_row:
                        self.current_table.append(self.current_row)
                elif tag in ("td", "th") and self.in_table:
                    self.current_row.append(self.current_cell.strip())
                    self.in_cell = False

            def handle_data(self, data):
                if self.in_cell:
                    self.current_cell += data

        parser = TableParser()
        parser.feed(html)
        return parser.tables

    def _write_tables_to_worksheet(self, ws, tables: List[List[List[str]]]) -> None:
        from openpyxl.styles import Font, Border, Side, PatternFill

        current_row = 1

        for table_idx, table in enumerate(tables):
            if table_idx > 0:
                current_row += 2

            for row_idx, row in enumerate(table):
                for col_idx, cell_value in enumerate(row):
                    cell = ws.cell(row=current_row, column=col_idx + 1, value=cell_value)

                    if row_idx == 0:
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

                    thin_border = Border(
                        left=Side(style="thin"), right=Side(style="thin"),
                        top=Side(style="thin"), bottom=Side(style="thin"),
                    )
                    cell.border = thin_border

                current_row += 1

    def _write_data_to_worksheet(self, ws, data: Dict[str, Any]) -> None:
        from openpyxl.styles import Font

        current_row = 1

        if "scalars" in data:
            for key, value in data["scalars"].items():
                ws.cell(row=current_row, column=1, value=key).font = Font(bold=True)
                ws.cell(row=current_row, column=2, value=str(value))
                current_row += 1
            current_row += 1

        if "rows" in data and data["rows"]:
            rows = data["rows"]
            if rows:
                for col_idx, key in enumerate(rows[0].keys()):
                    cell = ws.cell(row=current_row, column=col_idx + 1, value=key)
                    cell.font = Font(bold=True)
                current_row += 1

                for row in rows:
                    for col_idx, value in enumerate(row.values()):
                        ws.cell(row=current_row, column=col_idx + 1, value=str(value))
                    current_row += 1

    def _auto_fit_columns(self, ws) -> None:
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except Exception:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

def render_xlsx_from_html(html_path: Path, output_path: Path) -> RenderResult:
    """Convenience function to render XLSX from HTML file."""
    html_content = html_path.read_text(encoding="utf-8")
    renderer = XLSXRenderer()
    context = RenderContext(template_html=html_content, data={}, output_format=OutputFormat.XLSX, output_path=output_path)
    return renderer.render(context)

__all__ = [
    "RenderContext", "RenderResult", "Renderer", "BaseRenderer",
    "HTMLRenderer", "TokenEngine",
    "PDFRenderer", "PDFRendererSync",
    "DOCXRenderer", "render_docx_from_html",
    "XLSXRenderer", "render_xlsx_from_html",
]

# PIPELINES

"""Pipeline framework and concrete pipelines.

Consolidated from pipelines/base.py, pipelines/import_pipeline.py, pipelines/report_pipeline.py.
"""

import ast
import asyncio
import json
import logging
import os
import re
import shutil
import time
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import (
    Any,
    Awaitable,
    Callable,
    Dict,
    Generic,
    List,
    Optional,
    TypeVar,
    Union,
)

from backend.engine_all import PipelineError, ValidationError, NotFoundError, Event, publish_sync
from backend.engine_all import (
    Contract, TokenSet, ReshapeColumn, ReshapeRule, JoinSpec, OrderSpec,
    OutputFormat, RenderRequest, RenderOutput, Report,
    Template, TemplateKind, TemplateStatus, Artifact, TemplateSchema,
)

from prefect import flow, task
try:
    from prefect.task_runners import SequentialTaskRunner
except ImportError:
    SequentialTaskRunner = None

_base_logger = logging.getLogger("neura.pipelines")
_import_logger = logging.getLogger("neura.pipelines.import")
_report_logger = logging.getLogger("neura.pipelines.report")

T = TypeVar("T")
C = TypeVar("C", bound="PipelineContext")

# Base Pipeline Framework (from pipelines/base.py)

class _StepStatus(str, Enum):
    """Status of a pipeline step."""

    PENDING = "pending"
    RUNNING = "running"
    SUCCEEDED = "succeeded"
    FAILED = "failed"
    SKIPPED = "skipped"

# Alias to avoid collision with domain StepStatus
PipelineStepStatus = _StepStatus

@dataclass
class StepResult(Generic[T]):
    """Result of executing a step."""

    status: _StepStatus
    data: Optional[T] = None
    error: Optional[str] = None
    duration_ms: float = 0.0
    skipped_reason: Optional[str] = None

    @property
    def success(self) -> bool:
        return self.status == _StepStatus.SUCCEEDED

    @classmethod
    def ok(cls, data: T, duration_ms: float = 0.0) -> StepResult[T]:
        return cls(status=_StepStatus.SUCCEEDED, data=data, duration_ms=duration_ms)

    @classmethod
    def fail(cls, error: str, duration_ms: float = 0.0) -> StepResult[T]:
        return cls(status=_StepStatus.FAILED, error=error, duration_ms=duration_ms)

    @classmethod
    def skip(cls, reason: str) -> StepResult[T]:
        return cls(status=_StepStatus.SKIPPED, skipped_reason=reason)

@dataclass
class PipelineContext:
    """Context shared between pipeline steps."""

    correlation_id: str
    inputs: Dict[str, Any] = field(default_factory=dict)
    results: Dict[str, Any] = field(default_factory=dict)
    resources: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    cancelled: bool = False
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None

    def set(self, key: str, value: Any) -> None:
        self.results[key] = value

    def get(self, key: str, default: Any = None) -> Any:
        return self.results.get(key, default)

    def cancel(self) -> None:
        self.cancelled = True

    def is_cancelled(self) -> bool:
        return self.cancelled

    def get_resource(self, name: str) -> Any:
        return self.resources.get(name)

StepFunction = Union[
    Callable[[PipelineContext], Any],
    Callable[[PipelineContext], Awaitable[Any]],
]
GuardFunction = Callable[[PipelineContext], bool]

@dataclass
class Step:
    """A single step in a pipeline."""

    name: str
    fn: StepFunction
    label: Optional[str] = None
    guard: Optional[GuardFunction] = None
    retries: int = 0
    retry_delay_seconds: float = 1.0
    timeout_seconds: Optional[float] = None
    on_error: Optional[Callable[[Exception, PipelineContext], None]] = None

    @property
    def display_name(self) -> str:
        return self.label or self.name

    def should_run(self, ctx: PipelineContext) -> bool:
        if self.guard is None:
            return True
        try:
            return self.guard(ctx)
        except Exception:
            return True

    async def execute(self, ctx: PipelineContext) -> StepResult:
        """Execute this step with retries and timeout."""
        if ctx.is_cancelled():
            return StepResult.skip("Pipeline cancelled")

        if not self.should_run(ctx):
            return StepResult.skip(f"Guard returned false for {self.name}")

        last_error: Optional[Exception] = None
        attempts = self.retries + 1

        for attempt in range(attempts):
            if ctx.is_cancelled():
                return StepResult.skip("Pipeline cancelled")

            start = time.perf_counter()
            try:
                if asyncio.iscoroutinefunction(self.fn):
                    if self.timeout_seconds:
                        result = await asyncio.wait_for(
                            self.fn(ctx),
                            timeout=self.timeout_seconds,
                        )
                    else:
                        result = await self.fn(ctx)
                else:
                    if self.timeout_seconds:
                        result = await asyncio.wait_for(
                            asyncio.to_thread(self.fn, ctx),
                            timeout=self.timeout_seconds,
                        )
                    else:
                        result = self.fn(ctx)

                duration = (time.perf_counter() - start) * 1000
                return StepResult.ok(result, duration_ms=duration)

            except asyncio.TimeoutError:
                duration = (time.perf_counter() - start) * 1000
                return StepResult.fail(
                    f"Step {self.name} timed out after {self.timeout_seconds}s",
                    duration_ms=duration,
                )
            except Exception as e:
                last_error = e
                duration = (time.perf_counter() - start) * 1000

                if self.on_error:
                    try:
                        self.on_error(e, ctx)
                    except Exception:
                        _base_logger.debug("callback_failed", exc_info=True)

                if attempt < attempts - 1:
                    _base_logger.warning(
                        "step_retry",
                        extra={
                            "step": self.name,
                            "attempt": attempt + 1,
                            "max_attempts": attempts,
                            "error": str(e),
                        },
                    )
                    await asyncio.sleep(self.retry_delay_seconds)
                else:
                    return StepResult.fail(str(e), duration_ms=duration)

        return StepResult.fail(str(last_error) if last_error else "Unknown error")

@dataclass
class PipelineResult:
    """Result of executing a pipeline."""

    success: bool
    steps: Dict[str, StepResult]
    context: PipelineContext
    error: Optional[str] = None
    duration_ms: float = 0.0

    def get_step_result(self, name: str) -> Optional[StepResult]:
        return self.steps.get(name)

class Pipeline:
    """A sequence of steps that process data."""

    def __init__(
        self,
        name: str,
        steps: List[Step],
        *,
        on_error: Optional[Callable[[str, Exception, PipelineContext], None]] = None,
        on_success: Optional[Callable[[PipelineContext], None]] = None,
        on_step_complete: Optional[Callable[[str, StepResult, PipelineContext], None]] = None,
    ) -> None:
        self.name = name
        self.steps = steps
        self._on_error = on_error
        self._on_success = on_success
        self._on_step_complete = on_step_complete

    async def execute(self, ctx: PipelineContext) -> PipelineResult:
        """Execute all steps in order."""
        ctx.started_at = datetime.now(timezone.utc)
        start = time.perf_counter()
        step_results: Dict[str, StepResult] = {}

        publish_sync(
            Event(
                name="pipeline.started",
                payload={"pipeline": self.name, "steps": [s.name for s in self.steps]},
                correlation_id=ctx.correlation_id,
            )
        )

        try:
            for step in self.steps:
                if ctx.is_cancelled():
                    step_results[step.name] = StepResult.skip("Pipeline cancelled")
                    continue

                _base_logger.info(
                    "step_started",
                    extra={
                        "pipeline": self.name,
                        "step": step.name,
                        "correlation_id": ctx.correlation_id,
                    },
                )

                publish_sync(
                    Event(
                        name="pipeline.step_started",
                        payload={"pipeline": self.name, "step": step.name},
                        correlation_id=ctx.correlation_id,
                    )
                )

                result = await step.execute(ctx)
                step_results[step.name] = result

                if self._on_step_complete:
                    try:
                        self._on_step_complete(step.name, result, ctx)
                    except Exception:
                        _base_logger.debug("callback_failed", exc_info=True)

                _base_logger.info(
                    "step_completed",
                    extra={
                        "pipeline": self.name,
                        "step": step.name,
                        "status": result.status.value,
                        "duration_ms": result.duration_ms,
                        "correlation_id": ctx.correlation_id,
                    },
                )

                publish_sync(
                    Event(
                        name="pipeline.step_completed",
                        payload={
                            "pipeline": self.name,
                            "step": step.name,
                            "status": result.status.value,
                        },
                        correlation_id=ctx.correlation_id,
                    )
                )

                if result.status == _StepStatus.FAILED:
                    if self._on_error:
                        try:
                            self._on_error(
                                step.name,
                                PipelineError(
                                    message=result.error or "Step failed",
                                    step=step.name,
                                ),
                                ctx,
                            )
                        except Exception:
                            _base_logger.debug("callback_failed", exc_info=True)

                    duration = (time.perf_counter() - start) * 1000
                    ctx.completed_at = datetime.now(timezone.utc)

                    publish_sync(
                        Event(
                            name="pipeline.failed",
                            payload={
                                "pipeline": self.name,
                                "failed_step": step.name,
                                "error": result.error,
                            },
                            correlation_id=ctx.correlation_id,
                        )
                    )

                    return PipelineResult(
                        success=False,
                        steps=step_results,
                        context=ctx,
                        error=f"Step {step.name} failed: {result.error}",
                        duration_ms=duration,
                    )

            duration = (time.perf_counter() - start) * 1000
            ctx.completed_at = datetime.now(timezone.utc)

            if self._on_success:
                try:
                    self._on_success(ctx)
                except Exception:
                    _base_logger.debug("callback_failed", exc_info=True)

            publish_sync(
                Event(
                    name="pipeline.completed",
                    payload={"pipeline": self.name},
                    correlation_id=ctx.correlation_id,
                )
            )

            return PipelineResult(
                success=True,
                steps=step_results,
                context=ctx,
                duration_ms=duration,
            )

        except Exception as e:
            duration = (time.perf_counter() - start) * 1000
            ctx.completed_at = datetime.now(timezone.utc)

            _base_logger.exception(
                "pipeline_error",
                extra={
                    "pipeline": self.name,
                    "correlation_id": ctx.correlation_id,
                },
            )

            publish_sync(
                Event(
                    name="pipeline.error",
                    payload={"pipeline": self.name, "error": str(e)},
                    correlation_id=ctx.correlation_id,
                )
            )

            return PipelineResult(
                success=False,
                steps=step_results,
                context=ctx,
                error=str(e),
                duration_ms=duration,
            )

    def execute_sync(self, ctx: PipelineContext) -> PipelineResult:
        """Execute pipeline synchronously."""
        try:
            loop = asyncio.get_running_loop()
        except RuntimeError:
            loop = None
        if loop and loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                return pool.submit(asyncio.run, self.execute(ctx)).result()
        return asyncio.run(self.execute(ctx))

def step(
    name: str,
    *,
    label: Optional[str] = None,
    guard: Optional[GuardFunction] = None,
    retries: int = 0,
    timeout: Optional[float] = None,
) -> Callable[[StepFunction], Step]:
    """Decorator to create a step from a function."""

    def decorator(fn: StepFunction) -> Step:
        return Step(
            name=name,
            fn=fn,
            label=label or name,
            guard=guard,
            retries=retries,
            timeout_seconds=timeout,
        )

    return decorator

# Import Pipeline (from pipelines/import_pipeline.py)

@dataclass
class ImportPipelineContext(PipelineContext):
    """Context specific to template import."""

    source_path: Optional[Path] = None
    template_name: Optional[str] = None
    template_kind: TemplateKind = TemplateKind.PDF
    output_dir: Optional[Path] = None
    extracted_dir: Optional[Path] = None
    source_files: List[Path] = field(default_factory=list)
    html_content: Optional[str] = None
    extracted_tables: List[Dict[str, Any]] = field(default_factory=list)
    detected_tokens: List[str] = field(default_factory=list)
    template: Optional[Template] = None
    artifacts: List[Artifact] = field(default_factory=list)

def validate_import(ctx: ImportPipelineContext) -> None:
    """Validate the import request."""
    if not ctx.source_path:
        raise ValidationError(message="No source path provided")

    if not ctx.source_path.exists():
        raise ValidationError(message=f"Source file not found: {ctx.source_path}")

    if not ctx.template_name:
        ctx.template_name = ctx.source_path.stem

    suffix = ctx.source_path.suffix.lower()
    if suffix in (".xlsx", ".xls", ".xlsm"):
        ctx.template_kind = TemplateKind.EXCEL
    elif suffix == ".pdf":
        ctx.template_kind = TemplateKind.PDF
    elif suffix == ".zip":
        pass

    _import_logger.info(
        "import_validated",
        extra={
            "source": str(ctx.source_path),
            "kind": ctx.template_kind.value,
            "correlation_id": ctx.correlation_id,
        },
    )

def extract_archive(ctx: ImportPipelineContext) -> Path:
    """Extract ZIP archive if needed."""
    if ctx.source_path.suffix.lower() != ".zip":
        ctx.source_files = [ctx.source_path]
        return ctx.source_path

    extract_dir = ctx.output_dir / f"import_{uuid.uuid4().hex[:8]}"
    extract_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(ctx.source_path, "r") as zf:
        zf.extractall(extract_dir)

    ctx.extracted_dir = extract_dir
    ctx.source_files = list(extract_dir.glob("**/*"))
    ctx.source_files = [f for f in ctx.source_files if f.is_file()]

    for f in ctx.source_files:
        if f.suffix.lower() in (".xlsx", ".xls"):
            ctx.template_kind = TemplateKind.EXCEL
            break
        elif f.suffix.lower() == ".pdf":
            ctx.template_kind = TemplateKind.PDF
            break
        elif f.suffix.lower() == ".html":
            ctx.template_kind = TemplateKind.PDF
            break

    _import_logger.info(
        "archive_extracted",
        extra={
            "files": len(ctx.source_files),
            "kind": ctx.template_kind.value,
            "correlation_id": ctx.correlation_id,
        },
    )

    return extract_dir

def extract_content(ctx: ImportPipelineContext) -> Dict[str, Any]:
    """Extract content from source files."""
    if ctx.template_kind == TemplateKind.EXCEL:
        return _extract_excel_content(ctx)
    else:
        return _extract_pdf_content(ctx)

def _extract_pdf_content(ctx: ImportPipelineContext) -> Dict[str, Any]:
    """Extract content from PDF source."""
    html_files = [f for f in ctx.source_files if f.suffix.lower() == ".html"]
    if html_files:
        ctx.html_content = html_files[0].read_text(encoding="utf-8")
        return {"html": ctx.html_content}

    pdf_files = [f for f in ctx.source_files if f.suffix.lower() == ".pdf"]
    if not pdf_files:
        raise ValidationError(message="No PDF or HTML file found")

    from backend.engine_all import PDFExtractor
    extractor = PDFExtractor()
    result = extractor.extract(pdf_files[0])

    ctx.extracted_tables = [t.to_dict() for t in result.tables]
    ctx.html_content = _build_html_from_extraction(result)

    return {
        "html": ctx.html_content,
        "tables": ctx.extracted_tables,
        "page_count": result.page_count,
    }

def _extract_excel_content(ctx: ImportPipelineContext) -> Dict[str, Any]:
    """Extract content from Excel source."""
    excel_files = [
        f for f in ctx.source_files
        if f.suffix.lower() in (".xlsx", ".xls", ".xlsm")
    ]
    if not excel_files:
        raise ValidationError(message="No Excel file found")

    from backend.engine_all import ExcelExtractor
    extractor = ExcelExtractor()
    result = extractor.extract(excel_files[0])

    ctx.extracted_tables = [t.to_dict() for t in result.tables]
    ctx.html_content = _build_html_from_extraction(result)

    return {
        "html": ctx.html_content,
        "tables": ctx.extracted_tables,
        "sheet_count": len(result.tables),
    }

def _build_html_from_extraction(result) -> str:
    """Build HTML document from extraction result."""
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        '<meta charset="UTF-8">',
        "<style>",
        "table { border-collapse: collapse; width: 100%; }",
        "th, td { border: 1px solid black; padding: 8px; text-align: left; }",
        "th { background-color: #f2f2f2; }",
        "</style>",
        "</head>",
        "<body>",
    ]

    for table in result.tables:
        html_parts.append("<table>")
        html_parts.append("<thead><tr>")
        for header in table.headers:
            html_parts.append(f"<th>{header}</th>")
        html_parts.append("</tr></thead>")
        html_parts.append("<tbody>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row:
                html_parts.append(f"<td>{cell}</td>")
            html_parts.append("</tr>")
        html_parts.append("</tbody>")
        html_parts.append("</table>")
        html_parts.append("<br>")

    html_parts.extend(["</body>", "</html>"])
    return "\n".join(html_parts)

def detect_tokens(ctx: ImportPipelineContext) -> List[str]:
    """Detect placeholder tokens in the template."""
    if not ctx.html_content:
        return []

    pattern = r"\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}"
    matches = re.findall(pattern, ctx.html_content)

    ctx.detected_tokens = list(set(matches))

    _import_logger.info(
        "tokens_detected",
        extra={
            "count": len(ctx.detected_tokens),
            "tokens": ctx.detected_tokens[:10],
            "correlation_id": ctx.correlation_id,
        },
    )

    return ctx.detected_tokens

def create_template_record(ctx: ImportPipelineContext) -> Template:
    """Create the template record."""
    template_id = str(uuid.uuid4())

    template_dir = ctx.output_dir / template_id
    template_dir.mkdir(parents=True, exist_ok=True)

    html_path = template_dir / "template_p1.html"
    if ctx.html_content:
        html_path.write_text(ctx.html_content, encoding="utf-8")
        ctx.artifacts.append(
            Artifact(
                name="template_p1.html",
                path=html_path,
                artifact_type="html",
                size_bytes=html_path.stat().st_size,
                created_at=datetime.now(timezone.utc),
            )
        )

    if ctx.source_path and ctx.source_path.exists():
        source_copy = template_dir / f"original{ctx.source_path.suffix}"
        shutil.copy2(ctx.source_path, source_copy)
        ctx.artifacts.append(
            Artifact(
                name=source_copy.name,
                path=source_copy,
                artifact_type="source",
                size_bytes=source_copy.stat().st_size,
            )
        )

    ctx.template = Template(
        template_id=template_id,
        name=ctx.template_name,
        kind=ctx.template_kind,
        status=TemplateStatus.DRAFT,
        schema=TemplateSchema(
            scalars=[t for t in ctx.detected_tokens if not t.startswith("row_")],
            row_tokens=[t for t in ctx.detected_tokens if t.startswith("row_")],
            totals=[t for t in ctx.detected_tokens if t.startswith("total_")],
            placeholders_found=len(ctx.detected_tokens),
        ),
        artifacts=ctx.artifacts,
        source_file=ctx.source_path.name if ctx.source_path else None,
    )

    meta_path = template_dir / "template_meta.json"
    meta_path.write_text(
        json.dumps(ctx.template.to_dict(), indent=2, default=str),
        encoding="utf-8",
    )

    _import_logger.info(
        "template_created",
        extra={
            "template_id": template_id,
            "name": ctx.template_name,
            "kind": ctx.template_kind.value,
            "artifacts": len(ctx.artifacts),
            "correlation_id": ctx.correlation_id,
        },
    )

    return ctx.template

# Prefect tasks for import pipeline

@task(name="validate_import")
def validate_import_task(ctx: ImportPipelineContext) -> None:
    validate_import(ctx)

@task(name="extract_archive")
def extract_archive_task(ctx: ImportPipelineContext) -> Path:
    return extract_archive(ctx)

@task(name="extract_content")
def extract_content_task(ctx: ImportPipelineContext) -> Dict[str, Any]:
    return extract_content(ctx)

@task(name="detect_tokens")
def detect_tokens_task(ctx: ImportPipelineContext) -> List[str]:
    return detect_tokens(ctx)

@task(name="create_template")
def create_template_task(ctx: ImportPipelineContext) -> Template:
    return create_template_record(ctx)

_import_flow_kwargs = {"name": "template_import"}
if SequentialTaskRunner is not None:
    _import_flow_kwargs["task_runner"] = SequentialTaskRunner()

@flow(**_import_flow_kwargs)
def import_template_flow(
    source_path: Path,
    output_dir: Path,
    *,
    template_name: Optional[str] = None,
    correlation_id: Optional[str] = None,
) -> Template:
    ctx = ImportPipelineContext(
        correlation_id=correlation_id or str(uuid.uuid4()),
        source_path=source_path,
        template_name=template_name,
        output_dir=output_dir,
    )
    validate_import_task(ctx)
    extract_archive_task(ctx)
    extract_content_task(ctx)
    detect_tokens_task(ctx)
    create_template_task(ctx)
    return ctx.template

class ImportPipeline:
    """Template import pipeline wrapper."""

    def __init__(self, output_dir: Path) -> None:
        self._output_dir = output_dir
        self._pipeline = create_import_pipeline()

    def execute(
        self,
        source_path: Path,
        *,
        template_name: Optional[str] = None,
        correlation_id: Optional[str] = None,
    ) -> Template:
        engine = os.getenv("NEURA_PIPELINE_ENGINE", "prefect").strip().lower()
        if engine == "prefect":
            return import_template_flow(
                source_path=source_path,
                output_dir=self._output_dir,
                template_name=template_name,
                correlation_id=correlation_id,
            )

        ctx = ImportPipelineContext(
            correlation_id=correlation_id or str(uuid.uuid4()),
            source_path=source_path,
            template_name=template_name,
            output_dir=self._output_dir,
        )
        result = self._pipeline.execute_sync(ctx)
        if not result.success:
            raise Exception(f"Import pipeline failed: {result.error}")
        return ctx.template

def create_import_pipeline() -> Pipeline:
    """Create the template import pipeline."""
    return Pipeline(
        name="template_import",
        steps=[
            Step(name="validate", fn=validate_import, label="Validate import"),
            Step(name="extract_archive", fn=extract_archive, label="Extract archive"),
            Step(name="extract_content", fn=extract_content, label="Extract content"),
            Step(name="detect_tokens", fn=detect_tokens, label="Detect tokens"),
            Step(name="create_template", fn=create_template_record, label="Create template"),
        ],
    )

# Report Pipeline (from pipelines/report_pipeline.py)

@dataclass
class ReportPipelineContext(PipelineContext):
    """Context specific to report generation."""

    request: Optional[RenderRequest] = None
    template_path: Optional[Path] = None
    contract_path: Optional[Path] = None
    db_path: Optional[Path] = None
    contract: Optional[Contract] = None
    template_html: Optional[str] = None
    data_scalars: Dict[str, Any] = field(default_factory=dict)
    data_rows: List[Dict[str, Any]] = field(default_factory=list)
    data_totals: Dict[str, Any] = field(default_factory=dict)
    filled_html: Optional[str] = None
    outputs: List[RenderOutput] = field(default_factory=list)
    report: Optional[Report] = None

def validate_request(ctx: ReportPipelineContext) -> None:
    """Validate the render request."""
    if not ctx.request:
        raise ValidationError(message="No render request provided")
    if not ctx.request.template_id:
        raise ValidationError(message="template_id is required")
    if not ctx.request.connection_id:
        raise ValidationError(message="connection_id is required")

    _report_logger.info(
        "report_request_validated",
        extra={
            "template_id": ctx.request.template_id,
            "connection_id": ctx.request.connection_id,
            "correlation_id": ctx.correlation_id,
        },
    )

def load_contract(ctx: ReportPipelineContext) -> Contract:
    """Load and parse the contract file."""
    if not ctx.contract_path or not ctx.contract_path.exists():
        raise NotFoundError(message="Contract file not found")

    contract_data = json.loads(ctx.contract_path.read_text(encoding="utf-8"))
    contract = Contract.from_dict(contract_data, ctx.request.template_id)

    issues = contract.validate()
    if issues:
        _report_logger.warning(
            "contract_validation_warnings",
            extra={"issues": issues, "correlation_id": ctx.correlation_id},
        )

    ctx.contract = contract
    _report_logger.info(
        "contract_loaded",
        extra={
            "template_id": ctx.request.template_id,
            "tokens": len(contract.tokens.all_tokens()),
            "correlation_id": ctx.correlation_id,
        },
    )
    return contract

def load_template(ctx: ReportPipelineContext) -> str:
    """Load the template HTML."""
    if not ctx.template_path or not ctx.template_path.exists():
        raise NotFoundError(message="Template HTML file not found")

    ctx.template_html = ctx.template_path.read_text(encoding="utf-8")
    _report_logger.info(
        "template_loaded",
        extra={
            "path": str(ctx.template_path),
            "size_bytes": len(ctx.template_html),
            "correlation_id": ctx.correlation_id,
        },
    )
    return ctx.template_html

def load_data(ctx: ReportPipelineContext) -> Dict[str, Any]:
    """Load data from the database using the contract."""
    if not ctx.db_path or not ctx.db_path.exists():
        raise NotFoundError(message="Database file not found")
    if not ctx.contract:
        raise ValidationError(message="Contract not loaded")

    from backend.engine_all import SQLiteDataSource
    datasource = SQLiteDataSource(ctx.db_path)

    try:
        ctx.data_scalars = _load_scalars(datasource, ctx.contract, ctx.request)
        ctx.data_rows = _load_rows(datasource, ctx.contract, ctx.request)
        ctx.data_totals = _calculate_totals(ctx.data_rows, ctx.contract)

        _report_logger.info(
            "data_loaded",
            extra={
                "scalars": len(ctx.data_scalars),
                "rows": len(ctx.data_rows),
                "totals": len(ctx.data_totals),
                "correlation_id": ctx.correlation_id,
            },
        )

        return {
            "scalars": ctx.data_scalars,
            "rows": ctx.data_rows,
            "totals": ctx.data_totals,
        }
    finally:
        datasource.close()

def _load_scalars(datasource, contract: Contract, request: RenderRequest) -> Dict[str, Any]:
    """Load scalar values from database."""
    scalars: dict[str, Any] = {}

    select_parts: list[str] = []
    ordered_tokens: list[str] = []
    for token in contract.tokens.scalars:
        expr = contract.get_mapping(token)
        if not expr:
            continue
        select_parts.append(f"{expr} AS [{token}]")
        ordered_tokens.append(token)

    if select_parts:
        try:
            result = datasource.execute_query(f"SELECT {', '.join(select_parts)}")
            if result.rows:
                row = result.rows[0]
                for idx, token in enumerate(ordered_tokens):
                    try:
                        scalars[token] = row[idx]
                    except Exception:
                        scalars[token] = None
        except Exception as exc:
            _report_logger.warning(
                "scalar_batch_load_failed",
                extra={"error": str(exc), "token_count": len(ordered_tokens)},
            )
            for token in ordered_tokens:
                expr = contract.get_mapping(token)
                if not expr:
                    continue
                try:
                    result = datasource.execute_query(f"SELECT {expr} AS value")
                    if result.rows:
                        scalars[token] = result.rows[0][0]
                except Exception as e:
                    _report_logger.warning(
                        "scalar_load_failed",
                        extra={"token": token, "error": str(e)},
                    )
                    scalars[token] = None

    if request.start_date:
        scalars["START_DATE"] = request.start_date
    if request.end_date:
        scalars["END_DATE"] = request.end_date

    return scalars

def _load_rows(datasource, contract: Contract, request: RenderRequest) -> List[Dict[str, Any]]:
    """Load row data from database."""
    if not contract.tokens.row_tokens:
        return []

    select_parts = []
    for token in contract.tokens.row_tokens:
        expr = contract.get_mapping(token)
        if expr:
            select_parts.append(f"{expr} AS [{token}]")

    if not select_parts:
        return []

    def _extract_table(expr: str) -> str:
        if not expr:
            return ""
        match = re.search(r"(?:^|[^A-Za-z0-9_])\[?([A-Za-z_][\w]*)\]?\s*\.", str(expr))
        return match.group(1) if match else ""

    def _infer_row_tables() -> List[str]:
        tables: List[str] = []
        seen: set[str] = set()
        for token in contract.tokens.row_tokens:
            expr = contract.get_mapping(token) or ""
            table = _extract_table(expr)
            if table and table not in seen:
                seen.add(table)
                tables.append(table)
        return tables

    join = contract.join
    parent_table = join.parent_table if join else ""
    child_table = join.child_table if join else ""
    parent_key = join.parent_key if join else ""
    child_key = join.child_key if join else ""

    row_tables = _infer_row_tables()

    if not parent_table:
        if len(row_tables) == 1:
            parent_table = row_tables[0]
        elif len(row_tables) > 1:
            _report_logger.warning(
                "row_load_ambiguous_tables",
                extra={"template_id": contract.template_id, "tables": row_tables},
            )
            return []
    if not parent_table:
        _report_logger.warning(
            "row_load_missing_parent_table",
            extra={"template_id": contract.template_id},
        )
        return []

    query = f"SELECT {', '.join(select_parts)} FROM [{parent_table}]"
    if child_table and parent_key and child_key:
        query += (
            f" LEFT JOIN [{child_table}] ON [{parent_table}].[{parent_key}]"
            f" = [{child_table}].[{child_key}]"
        )

    conditions = []
    date_column = ""
    date_table = ""
    if contract.date_columns:
        row_date_candidates = [
            table for table in row_tables if contract.date_columns.get(table)
        ]
        if len(row_date_candidates) == 1:
            date_table = row_date_candidates[0]
            date_column = contract.date_columns.get(date_table, "")
        elif len(row_date_candidates) > 1:
            _report_logger.warning(
                "row_load_ambiguous_date_tables",
                extra={"template_id": contract.template_id, "tables": row_date_candidates},
            )
        else:
            parent_date = contract.date_columns.get(parent_table, "")
            child_date = contract.date_columns.get(child_table, "") if child_table else ""
            if parent_date:
                date_column = parent_date
                date_table = parent_table
            elif child_date:
                date_column = child_date
                date_table = child_table
    _DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
    if request.start_date or request.end_date:
        if date_column:
            date_expr = date_column
            if "." not in date_expr and date_table:
                date_expr = f"{date_table}.{date_expr}"
            if request.start_date:
                if not _DATE_RE.match(request.start_date):
                    raise ValidationError(message="Invalid start_date format: expected YYYY-MM-DD")
                conditions.append(f"date({date_expr}) >= date('{request.start_date}')")
            if request.end_date:
                if not _DATE_RE.match(request.end_date):
                    raise ValidationError(message="Invalid end_date format: expected YYYY-MM-DD")
                conditions.append(f"date({date_expr}) <= date('{request.end_date}')")
        else:
            _report_logger.warning(
                "row_load_missing_date_column",
                extra={
                    "template_id": contract.template_id,
                    "start_date": request.start_date,
                    "end_date": request.end_date,
                },
            )

    if conditions:
        query += " WHERE " + " AND ".join(conditions)

    if contract.row_order:
        safe_order_terms = []
        for term in contract.row_order:
            stripped = term.strip()
            parts = stripped.rsplit(None, 1)
            if len(parts) == 2 and parts[1].upper() in ("ASC", "DESC"):
                safe_order_terms.append(f"[{parts[0]}] {parts[1]}")
            else:
                safe_order_terms.append(f"[{stripped}]")
        query += f" ORDER BY {', '.join(safe_order_terms)}"

    try:
        result = datasource.execute_query(query)
        return result.to_dicts()
    except Exception as e:
        _report_logger.warning("row_load_failed", extra={"error": str(e)})
        return []

def _calculate_totals(rows: List[Dict[str, Any]], contract: Contract) -> Dict[str, Any]:
    """Calculate totals from row data."""
    totals = {}

    def _is_simple_row_ref(expr: str) -> Optional[str]:
        normalized = expr.strip()
        if not normalized:
            return None
        if normalized.lower().startswith("rows."):
            candidate = normalized.split(".", 1)[1]
            return candidate if candidate in contract.tokens.row_tokens else None
        if normalized in contract.tokens.row_tokens:
            return normalized
        return None

    for token in contract.tokens.totals:
        expr = (contract.totals_math.get(token) or "").strip()
        mapping_expr = (contract.get_mapping(token) or "").strip()

        if not expr and mapping_expr:
            expr = mapping_expr

        if expr:
            row_ref = _is_simple_row_ref(expr)
            if row_ref:
                try:
                    totals[token] = sum(float(row.get(row_ref, 0) or 0) for row in rows)
                except (ValueError, TypeError):
                    totals[token] = 0
            else:
                totals[token] = _eval_total_expr(expr, rows, totals)
            continue

        _report_logger.warning("totals_missing_math", extra={"total_token": token})

    return totals

def _eval_total_expr(expr: str, rows: List[Dict[str, Any]], totals: Dict[str, Any]) -> Any:
    """Evaluate a totals expression safely."""
    if not expr:
        return None

    normalized = re.sub(r"\brows\.", "", expr.strip(), flags=re.IGNORECASE)
    normalized = re.sub(r"\btotals\.", "", normalized, flags=re.IGNORECASE)

    if "CASE" in normalized.upper():
        _report_logger.warning("totals_expr_unsupported_case", extra={"expr": expr})
        return None

    replacements = {
        "SUM": "sum_", "COUNT": "count_", "AVG": "avg_",
        "MIN": "min_", "MAX": "max_", "NULLIF": "nullif", "COALESCE": "coalesce",
    }
    for sql_name, py_name in replacements.items():
        normalized = re.sub(rf"\b{sql_name}\b", py_name, normalized, flags=re.IGNORECASE)

    column_values: Dict[str, List[Any]] = {}
    for row in rows:
        for key, value in row.items():
            column_values.setdefault(key, []).append(value)

    def _as_number(value: Any) -> Optional[float]:
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    def _numeric(values: List[Any]) -> List[float]:
        return [num for num in (_as_number(val) for val in values) if num is not None]

    def sum_(values: List[Any]) -> float:
        nums = _numeric(values)
        return float(sum(nums)) if nums else 0.0

    def count_(values: List[Any]) -> int:
        return len([val for val in values if val is not None])

    def avg_(values: List[Any]) -> float:
        nums = _numeric(values)
        return float(sum(nums) / len(nums)) if nums else 0.0

    def min_(values: List[Any]) -> float:
        nums = _numeric(values)
        return float(min(nums)) if nums else 0.0

    def max_(values: List[Any]) -> float:
        nums = _numeric(values)
        return float(max(nums)) if nums else 0.0

    def nullif(a: Any, b: Any) -> Any:
        return None if a == b else a

    def coalesce(*args: Any) -> Any:
        for item in args:
            if item is not None:
                return item
        return None

    allowed_funcs = {
        "sum_": sum_, "count_": count_, "avg_": avg_,
        "min_": min_, "max_": max_, "nullif": nullif, "coalesce": coalesce,
    }
    allowed_names = set(column_values.keys()) | set(totals.keys())
    allowed_nodes = (
        ast.Expression, ast.BinOp, ast.UnaryOp, ast.Call, ast.Name, ast.Load,
        ast.Constant, ast.Add, ast.Sub, ast.Mult, ast.Div, ast.Mod, ast.Pow,
        ast.USub, ast.UAdd,
    )

    try:
        tree = ast.parse(normalized, mode="eval")
        for node in ast.walk(tree):
            if not isinstance(node, allowed_nodes):
                raise ValueError(f"Unsupported expression node: {type(node).__name__}")
            if isinstance(node, ast.Call):
                if not isinstance(node.func, ast.Name) or node.func.id not in allowed_funcs:
                    raise ValueError("Unsupported function in totals expression")
            if isinstance(node, ast.Name):
                if node.id not in allowed_funcs and node.id not in allowed_names:
                    raise ValueError(f"Unknown name '{node.id}' in totals expression")
        context = {**allowed_funcs, **column_values, **totals}
        # SECURITY: eval is sandboxed via AST whitelist above and __builtins__={}.
        return eval(compile(tree, "<totals_expr>", "eval"), {"__builtins__": {}}, context)  # noqa: S307
    except (ValueError, TypeError, ZeroDivisionError, SyntaxError) as exc:
        _report_logger.warning("totals_expr_eval_failed", extra={"expr": expr, "error": str(exc)})
        return None

def render_html(ctx: ReportPipelineContext) -> Path:
    """Render filled HTML from template and data."""
    if not ctx.template_html:
        raise ValidationError(message="Template HTML not loaded")

    from backend.engine_all import HTMLRenderer, RenderContext
    renderer = HTMLRenderer()
    output_dir = ctx.template_path.parent
    output_path = output_dir / f"filled_{int(datetime.now(timezone.utc).timestamp())}.html"

    all_data = {
        **ctx.data_scalars,
        **{f"row_{i}_{k}": v for i, row in enumerate(ctx.data_rows) for k, v in row.items()},
        **ctx.data_totals,
        "ROW_COUNT": len(ctx.data_rows),
        "GENERATED_AT": datetime.now(timezone.utc).isoformat(),
    }

    render_ctx = RenderContext(
        template_html=ctx.template_html,
        data=all_data,
        output_format=OutputFormat.HTML,
        output_path=output_path,
    )

    result = renderer.render(render_ctx)
    if not result.success:
        raise Exception(f"HTML rendering failed: {result.error}")

    ctx.filled_html = output_path.read_text(encoding="utf-8")
    ctx.outputs.append(
        RenderOutput(format=OutputFormat.HTML, path=output_path, size_bytes=result.size_bytes)
    )

    _report_logger.info(
        "html_rendered",
        extra={"path": str(output_path), "size_bytes": result.size_bytes, "correlation_id": ctx.correlation_id},
    )

    return output_path

def render_pdf(ctx: ReportPipelineContext) -> Optional[Path]:
    """Render PDF from HTML."""
    if not ctx.filled_html:
        raise ValidationError(message="HTML not rendered")
    if OutputFormat.PDF not in ctx.request.output_formats:
        return None

    from backend.engine_all import PDFRenderer, RenderContext
    renderer = PDFRenderer()
    output_dir = ctx.template_path.parent
    output_path = output_dir / f"filled_{int(datetime.now(timezone.utc).timestamp())}.pdf"

    render_ctx = RenderContext(
        template_html=ctx.filled_html, data={},
        output_format=OutputFormat.PDF, output_path=output_path,
    )

    result = renderer.render(render_ctx)
    if not result.success:
        _report_logger.error("pdf_render_failed", extra={"error": result.error})
        return None

    ctx.outputs.append(RenderOutput(format=OutputFormat.PDF, path=output_path, size_bytes=result.size_bytes))
    _report_logger.info("pdf_rendered", extra={"path": str(output_path), "size_bytes": result.size_bytes, "correlation_id": ctx.correlation_id})
    return output_path

def render_docx(ctx: ReportPipelineContext) -> Optional[Path]:
    """Render DOCX from HTML."""
    if OutputFormat.DOCX not in ctx.request.output_formats:
        return None
    if not ctx.filled_html:
        return None

    from backend.engine_all import DOCXRenderer, RenderContext
    renderer = DOCXRenderer()
    output_dir = ctx.template_path.parent
    output_path = output_dir / f"filled_{int(datetime.now(timezone.utc).timestamp())}.docx"

    render_ctx = RenderContext(
        template_html=ctx.filled_html, data={},
        output_format=OutputFormat.DOCX, output_path=output_path,
    )

    result = renderer.render(render_ctx)
    if not result.success:
        _report_logger.warning("docx_render_failed", extra={"error": result.error})
        return None

    ctx.outputs.append(RenderOutput(format=OutputFormat.DOCX, path=output_path, size_bytes=result.size_bytes))
    return output_path

def render_xlsx(ctx: ReportPipelineContext) -> Optional[Path]:
    """Render XLSX from HTML tables."""
    if OutputFormat.XLSX not in ctx.request.output_formats:
        return None
    if not ctx.filled_html:
        return None

    from backend.engine_all import XLSXRenderer, RenderContext
    renderer = XLSXRenderer()
    output_dir = ctx.template_path.parent
    output_path = output_dir / f"filled_{int(datetime.now(timezone.utc).timestamp())}.xlsx"

    render_ctx = RenderContext(
        template_html=ctx.filled_html,
        data={"scalars": ctx.data_scalars, "rows": ctx.data_rows},
        output_format=OutputFormat.XLSX, output_path=output_path,
    )

    result = renderer.render(render_ctx)
    if not result.success:
        _report_logger.warning("xlsx_render_failed", extra={"error": result.error})
        return None

    ctx.outputs.append(RenderOutput(format=OutputFormat.XLSX, path=output_path, size_bytes=result.size_bytes))
    return output_path

def finalize_report(ctx: ReportPipelineContext) -> Report:
    """Create the final report record."""
    report = Report(
        report_id=str(uuid.uuid4()),
        template_id=ctx.request.template_id,
        template_name=ctx.request.template_id,
        connection_id=ctx.request.connection_id,
        connection_name=None,
        status="succeeded",
        outputs=ctx.outputs,
        start_date=ctx.request.start_date,
        end_date=ctx.request.end_date,
        correlation_id=ctx.correlation_id,
        started_at=ctx.started_at,
        completed_at=datetime.now(timezone.utc),
    )
    ctx.report = report
    return report

# Prefect tasks for report pipeline

@task(name="validate_request")
def validate_request_task(ctx: ReportPipelineContext) -> None:
    validate_request(ctx)

@task(name="load_contract")
def load_contract_task(ctx: ReportPipelineContext) -> Contract:
    return load_contract(ctx)

@task(name="load_template")
def load_template_task(ctx: ReportPipelineContext) -> str:
    return load_template(ctx)

@task(name="load_data", retries=1, retry_delay_seconds=1)
def load_data_task(ctx: ReportPipelineContext) -> Dict[str, Any]:
    return load_data(ctx)

@task(name="render_html")
def render_html_task(ctx: ReportPipelineContext) -> Path:
    return render_html(ctx)

@task(name="render_pdf", timeout_seconds=120)
def render_pdf_task(ctx: ReportPipelineContext) -> Optional[Path]:
    return render_pdf(ctx)

@task(name="render_docx")
def render_docx_task(ctx: ReportPipelineContext) -> Optional[Path]:
    return render_docx(ctx)

@task(name="render_xlsx")
def render_xlsx_task(ctx: ReportPipelineContext) -> Optional[Path]:
    return render_xlsx(ctx)

@task(name="finalize")
def finalize_report_task(ctx: ReportPipelineContext) -> Report:
    return finalize_report(ctx)

_report_flow_kwargs = {"name": "report_generation"}
if SequentialTaskRunner is not None:
    _report_flow_kwargs["task_runner"] = SequentialTaskRunner()

@flow(**_report_flow_kwargs)
def report_generation_flow(
    request: RenderRequest,
    template_path: Path,
    contract_path: Path,
    db_path: Path,
    correlation_id: Optional[str] = None,
) -> Report:
    ctx = ReportPipelineContext(
        correlation_id=correlation_id or str(uuid.uuid4()),
        request=request,
        template_path=template_path,
        contract_path=contract_path,
        db_path=db_path,
    )
    validate_request_task(ctx)
    load_contract_task(ctx)
    load_template_task(ctx)
    load_data_task(ctx)
    render_html_task(ctx)
    if OutputFormat.PDF in request.output_formats:
        render_pdf_task(ctx)
    if OutputFormat.DOCX in request.output_formats:
        render_docx_task(ctx)
    if OutputFormat.XLSX in request.output_formats:
        render_xlsx_task(ctx)
    finalize_report_task(ctx)
    return ctx.report

class ReportPipeline:
    """Report generation pipeline wrapper."""

    def __init__(self) -> None:
        self._pipeline = create_report_pipeline()

    def execute(
        self,
        request: RenderRequest,
        template_path: Path,
        contract_path: Path,
        db_path: Path,
        *,
        correlation_id: Optional[str] = None,
    ) -> Report:
        engine = os.getenv("NEURA_PIPELINE_ENGINE", "prefect").strip().lower()
        if engine == "prefect":
            return report_generation_flow(
                request=request,
                template_path=template_path,
                contract_path=contract_path,
                db_path=db_path,
                correlation_id=correlation_id,
            )

        ctx = ReportPipelineContext(
            correlation_id=correlation_id or str(uuid.uuid4()),
            request=request,
            template_path=template_path,
            contract_path=contract_path,
            db_path=db_path,
        )
        result = self._pipeline.execute_sync(ctx)
        if not result.success:
            raise Exception(f"Report pipeline failed: {result.error}")
        return ctx.report

def create_report_pipeline() -> Pipeline:
    """Create the report generation pipeline."""
    return Pipeline(
        name="report_generation",
        steps=[
            Step(name="validate", fn=validate_request, label="Validate request"),
            Step(name="load_contract", fn=load_contract, label="Load contract"),
            Step(name="load_template", fn=load_template, label="Load template"),
            Step(name="load_data", fn=load_data, label="Load data", retries=1),
            Step(name="render_html", fn=render_html, label="Render HTML"),
            Step(name="render_pdf", fn=render_pdf, label="Render PDF", timeout_seconds=120.0,
                 guard=lambda ctx: OutputFormat.PDF in ctx.request.output_formats),
            Step(name="render_docx", fn=render_docx, label="Render DOCX",
                 guard=lambda ctx: OutputFormat.DOCX in ctx.request.output_formats),
            Step(name="render_xlsx", fn=render_xlsx, label="Render XLSX",
                 guard=lambda ctx: OutputFormat.XLSX in ctx.request.output_formats),
            Step(name="finalize", fn=finalize_report, label="Finalize report"),
        ],
    )

__all__ = [
    # Base
    "Pipeline", "Step", "PipelineContext", "StepResult", "PipelineResult",
    "PipelineStepStatus", "StepFunction", "GuardFunction", "step",
    # Import
    "ImportPipeline", "ImportPipelineContext", "create_import_pipeline",
    "import_template_flow",
    # Report
    "ReportPipeline", "ReportPipelineContext", "create_report_pipeline",
    "report_generation_flow",
]
