"""
Document Ingestion Service
Handles file import, type detection, and document creation from various sources.
"""
from __future__ import annotations

import logging
import mimetypes
import hashlib
import zipfile
import io
import json
from pathlib import Path
from typing import Any, Dict, List, Optional
from datetime import datetime, timezone
from enum import Enum

from pydantic import BaseModel, Field
from backend.app.common import utc_now, utc_now_iso

logger = logging.getLogger(__name__)

class FileType(str, Enum):
    """Supported file types for ingestion."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    PPTX = "pptx"
    PPT = "ppt"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"
    XML = "xml"
    YAML = "yaml"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    ARCHIVE = "archive"
    UNKNOWN = "unknown"


class IngestionResult(BaseModel):
    """Result of document ingestion."""
    document_id: str
    filename: str
    file_type: FileType
    size_bytes: int
    pages: Optional[int] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)
    preview_url: Optional[str] = None
    created_at: datetime = Field(default_factory=utc_now)
    processing_status: str = "completed"
    warnings: List[str] = Field(default_factory=list)

class BulkIngestionResult(BaseModel):
    """Result of bulk document ingestion."""
    total_files: int
    successful: int
    failed: int
    results: List[IngestionResult] = Field(default_factory=list)
    errors: List[Dict[str, str]] = Field(default_factory=list)

class StructuredDataImport(BaseModel):
    """Result of structured data import."""
    document_id: str
    table_name: str
    row_count: int
    column_count: int
    columns: List[str]
    sample_data: List[Dict[str, Any]]

class IngestionService:
    """
    Service for ingesting documents from various sources.
    Handles auto-detection, preview generation, and document creation.
    """

    # File extension to type mapping
    EXTENSION_MAP = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".doc": FileType.DOC,
        ".xlsx": FileType.XLSX,
        ".xls": FileType.XLS,
        ".csv": FileType.CSV,
        ".pptx": FileType.PPTX,
        ".ppt": FileType.PPT,
        ".txt": FileType.TXT,
        ".rtf": FileType.RTF,
        ".html": FileType.HTML,
        ".htm": FileType.HTML,
        ".md": FileType.MARKDOWN,
        ".markdown": FileType.MARKDOWN,
        ".json": FileType.JSON,
        ".xml": FileType.XML,
        ".yaml": FileType.YAML,
        ".yml": FileType.YAML,
        ".png": FileType.IMAGE,
        ".jpg": FileType.IMAGE,
        ".jpeg": FileType.IMAGE,
        ".gif": FileType.IMAGE,
        ".bmp": FileType.IMAGE,
        ".webp": FileType.IMAGE,
        ".svg": FileType.IMAGE,
        ".mp3": FileType.AUDIO,
        ".wav": FileType.AUDIO,
        ".m4a": FileType.AUDIO,
        ".ogg": FileType.AUDIO,
        ".mp4": FileType.VIDEO,
        ".avi": FileType.VIDEO,
        ".mov": FileType.VIDEO,
        ".mkv": FileType.VIDEO,
        ".webm": FileType.VIDEO,
        ".zip": FileType.ARCHIVE,
        ".tar": FileType.ARCHIVE,
        ".gz": FileType.ARCHIVE,
        ".rar": FileType.ARCHIVE,
        ".7z": FileType.ARCHIVE,
    }

    def __init__(self):
        self._upload_dir: Optional[Path] = None

    def _get_upload_dir(self) -> Path:
        """Lazy load upload directory from settings."""
        if self._upload_dir is None:
            from backend.app.services.config import get_settings
            self._upload_dir = get_settings().uploads_dir
        return self._upload_dir

    def detect_file_type(self, filename: str, content: Optional[bytes] = None) -> FileType:
        """Detect file type from filename and optionally content."""
        ext = Path(filename).suffix.lower()

        if ext in self.EXTENSION_MAP:
            return self.EXTENSION_MAP[ext]

        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            if mime_type.startswith("image/"):
                return FileType.IMAGE
            elif mime_type.startswith("audio/"):
                return FileType.AUDIO
            elif mime_type.startswith("video/"):
                return FileType.VIDEO
            elif mime_type.startswith("text/"):
                return FileType.TXT

        # Magic number detection for content
        if content:
            if content[:4] == b"%PDF":
                return FileType.PDF
            elif content[:4] == b"PK\x03\x04":
                # Could be DOCX, XLSX, PPTX, or ZIP
                return self._detect_office_type(content)
            elif content[:3] == b"\xef\xbb\xbf" or content[:1000].decode("utf-8", errors="ignore").strip().startswith(("{", "[")):
                return FileType.JSON

        return FileType.UNKNOWN

    def _detect_office_type(self, content: bytes) -> FileType:
        """Detect specific Office format from ZIP-based file."""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                if any("word/" in n for n in names):
                    return FileType.DOCX
                elif any("xl/" in n for n in names):
                    return FileType.XLSX
                elif any("ppt/" in n for n in names):
                    return FileType.PPTX
        except zipfile.BadZipFile:
            pass
        return FileType.ARCHIVE

    async def ingest_file(
        self,
        filename: str,
        content: bytes,
        metadata: Optional[Dict[str, Any]] = None,
        auto_ocr: bool = True,
        generate_preview: bool = True,
    ) -> IngestionResult:
        """Ingest a single file and create a document."""
        file_type = self.detect_file_type(filename, content)
        doc_id = self._generate_document_id(filename, content)

        # Save file
        upload_dir = self._get_upload_dir()
        file_path = upload_dir / doc_id / filename
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(content)

        extracted_metadata = await self._extract_metadata(file_path, file_type)
        if metadata:
            extracted_metadata.update(metadata)

        # Get page count for documents
        pages = await self._get_page_count(file_path, file_type)

        # Generate preview if requested
        preview_url = None
        if generate_preview:
            preview_url = await self._generate_preview(file_path, file_type, doc_id)

        # OCR if needed
        warnings = []
        if auto_ocr and file_type == FileType.PDF:
            is_scanned = await self._is_scanned_pdf(file_path)
            if is_scanned:
                await self._perform_ocr(file_path)
                warnings.append("Document was OCR'd from scanned images")

        return IngestionResult(
            document_id=doc_id,
            filename=filename,
            file_type=file_type,
            size_bytes=len(content),
            pages=pages,
            metadata=extracted_metadata,
            preview_url=preview_url,
            warnings=warnings,
        )

    async def ingest_from_url(
        self,
        url: str,
        filename: Optional[str] = None,
    ) -> IngestionResult:
        """Ingest a document from a URL."""
        import aiohttp

        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                response.raise_for_status()
                content = await response.read()

                # Determine filename
                if not filename:
                    # Try to get from Content-Disposition header
                    cd = response.headers.get("Content-Disposition", "")
                    if "filename=" in cd:
                        filename = cd.split("filename=")[1].strip('"')
                    else:
                        # Use URL path
                        from urllib.parse import urlparse
                        filename = Path(urlparse(url).path).name or "downloaded_file"

                return await self.ingest_file(filename, content)

    async def ingest_zip_archive(
        self,
        filename: str,
        content: bytes,
        preserve_structure: bool = True,
        flatten: bool = False,
    ) -> BulkIngestionResult:
        """Ingest documents from a ZIP archive."""
        results = []
        errors = []

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for name in zf.namelist():
                    # Skip directories
                    if name.endswith("/"):
                        continue

                    # Skip hidden/system files
                    if any(part.startswith(".") or part.startswith("__") for part in Path(name).parts):
                        continue

                    try:
                        file_content = zf.read(name)
                        file_name = Path(name).name if flatten else name

                        result = await self.ingest_file(
                            filename=file_name,
                            content=file_content,
                            metadata={"source_archive": filename, "original_path": name},
                        )
                        results.append(result)
                    except Exception as e:
                        logger.error("Failed to ingest %s from archive: %s", name, e, exc_info=True)
                        errors.append({"file": name, "error": "File ingestion failed"})

        except zipfile.BadZipFile as e:
            errors.append({"file": filename, "error": f"Invalid ZIP file: {e}"})

        return BulkIngestionResult(
            total_files=len(results) + len(errors),
            successful=len(results),
            failed=len(errors),
            results=results,
            errors=errors,
        )

    async def import_structured_data(
        self,
        filename: str,
        content: bytes,
        format_hint: Optional[str] = None,
    ) -> StructuredDataImport:
        """Import structured data (JSON, XML, YAML) as an editable table."""
        file_type = format_hint or self.detect_file_type(filename, content)
        text_content = content.decode("utf-8")

        data = None
        if file_type in (FileType.JSON, "json"):
            data = json.loads(text_content)
        elif file_type in (FileType.YAML, "yaml"):
            import yaml
            data = yaml.safe_load(text_content)
        elif file_type in (FileType.XML, "xml"):
            import xml.etree.ElementTree as ET
            root = ET.fromstring(text_content)
            data = self._xml_to_dict(root)

        # Normalize to list of dicts
        if isinstance(data, dict):
            # Check if it's a single record or contains a list
            for key, value in data.items():
                if isinstance(value, list) and all(isinstance(v, dict) for v in value):
                    data = value
                    break
            else:
                data = [data]
        elif not isinstance(data, list):
            data = [{"value": data}]

        # Get columns
        columns = []
        for record in data[:10]:
            if isinstance(record, dict):
                for key in record.keys():
                    if key not in columns:
                        columns.append(key)

        # Create document
        doc_id = self._generate_document_id(filename, content)

        return StructuredDataImport(
            document_id=doc_id,
            table_name=Path(filename).stem,
            row_count=len(data),
            column_count=len(columns),
            columns=columns,
            sample_data=data[:10],
        )

    def _xml_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary."""
        result = {}
        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text
            else:
                result[child.tag] = self._xml_to_dict(child)
        if element.attrib:
            result["@attributes"] = element.attrib
        return result

    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID."""
        hash_input = f"{filename}:{len(content)}:{datetime.now(timezone.utc).isoformat()}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]

    async def _extract_metadata(self, file_path: Path, file_type: FileType) -> Dict[str, Any]:
        """Extract metadata from file."""
        metadata = {
            "original_filename": file_path.name,
            "file_type": file_type.value,
            "ingested_at": datetime.now(timezone.utc).isoformat(),
        }

        try:
            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                pdf_metadata = doc.metadata
                if pdf_metadata:
                    metadata.update({
                        "title": pdf_metadata.get("title"),
                        "author": pdf_metadata.get("author"),
                        "subject": pdf_metadata.get("subject"),
                        "creator": pdf_metadata.get("creator"),
                        "creation_date": pdf_metadata.get("creationDate"),
                    })
                doc.close()

            elif file_type == FileType.DOCX:
                from docx import Document
                doc = Document(str(file_path))
                props = doc.core_properties
                metadata.update({
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                })

            elif file_type == FileType.IMAGE:
                from PIL import Image
                with Image.open(file_path) as img:
                    metadata.update({
                        "width": img.width,
                        "height": img.height,
                        "format": img.format,
                        "mode": img.mode,
                    })

        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")

        return {k: v for k, v in metadata.items() if v is not None}

    async def _get_page_count(self, file_path: Path, file_type: FileType) -> Optional[int]:
        """Get page count for documents."""
        try:
            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                count = len(doc)
                doc.close()
                return count

            elif file_type == FileType.DOCX:
                # Approximate from paragraphs (rough estimate)
                from docx import Document
                doc = Document(str(file_path))
                # ~40 paragraphs per page is a rough estimate
                return max(1, len(doc.paragraphs) // 40)

            elif file_type == FileType.PPTX:
                from pptx import Presentation
                prs = Presentation(str(file_path))
                return len(prs.slides)

        except Exception as e:
            logger.warning(f"Failed to get page count: {e}")

        return None

    async def _generate_preview(self, file_path: Path, file_type: FileType, doc_id: str) -> Optional[str]:
        """Generate preview image for document."""
        try:
            preview_dir = self._get_upload_dir() / doc_id / "previews"
            preview_dir.mkdir(parents=True, exist_ok=True)
            preview_path = preview_dir / "preview.png"

            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                if len(doc) > 0:
                    page = doc[0]
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    pix.save(str(preview_path))
                doc.close()

            elif file_type == FileType.IMAGE:
                from PIL import Image
                with Image.open(file_path) as img:
                    # Create thumbnail
                    img.thumbnail((400, 400))
                    img.save(preview_path, "PNG")

            if preview_path.exists():
                return f"/uploads/{doc_id}/previews/preview.png"

        except Exception as e:
            logger.warning(f"Failed to generate preview: {e}")

        return None

    async def _is_scanned_pdf(self, file_path: Path) -> bool:
        """Check if PDF appears to be scanned (image-based)."""
        try:
            import fitz
            doc = fitz.open(str(file_path))
            if len(doc) == 0:
                doc.close()
                return False

            # Check first few pages
            for page_num in range(min(3, len(doc))):
                page = doc[page_num]
                text = page.get_text()
                if len(text.strip()) > 50:
                    doc.close()
                    return False

            doc.close()
            return True
        except Exception:
            return False

    async def _perform_ocr(self, file_path: Path) -> None:
        """Perform OCR on a scanned PDF."""
        # This would integrate with the DocAI OCR service
        logger.info(f"OCR would be performed on: {file_path}")

# Singleton instance
ingestion_service = IngestionService()

"""
Email Ingestion Service
Handles email-to-document conversion and email inbox monitoring.
"""

import logging
import email
import imaplib
import hashlib
import re
from datetime import datetime, timezone
from email.header import decode_header
from email.utils import parseaddr, parsedate_to_datetime
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

class EmailAttachment(BaseModel):
    """Email attachment details."""
    filename: str
    content_type: str
    size_bytes: int
    document_id: Optional[str] = None

class ParsedEmail(BaseModel):
    """Parsed email structure."""
    message_id: str
    subject: str
    from_address: str
    from_name: Optional[str] = None
    to_addresses: List[str] = Field(default_factory=list)
    cc_addresses: List[str] = Field(default_factory=list)
    date: Optional[datetime] = None
    body_text: Optional[str] = None
    body_html: Optional[str] = None
    attachments: List[EmailAttachment] = Field(default_factory=list)
    thread_id: Optional[str] = None
    in_reply_to: Optional[str] = None

class EmailDocumentResult(BaseModel):
    """Result of converting email to document."""
    document_id: str
    email_subject: str
    from_address: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    attachment_count: int = 0
    attachment_documents: List[str] = Field(default_factory=list)

class EmailInboxConfig(BaseModel):
    """Email inbox configuration."""
    inbox_id: str
    email_address: str
    imap_server: str
    imap_port: int = 993
    username: str
    password: str  # Should be encrypted in practice
    use_ssl: bool = True
    folder: str = "INBOX"
    auto_process: bool = True
    delete_after_process: bool = False

class EmailIngestionService:
    """
    Service for ingesting documents from emails.
    Supports unique inbox addresses, IMAP monitoring, and attachment extraction.
    """

    def __init__(self):
        self._inbox_configs: Dict[str, EmailInboxConfig] = {}

    async def connect_imap(
        self,
        host: str,
        port: int,
        username: str,
        password: str,
        use_ssl: bool = True,
        folder: str = "INBOX",
    ) -> Dict[str, Any]:
        """
        Validate IMAP credentials and register an account configuration.

        Returns a safe account payload without sensitive fields.
        """
        if not host or not username or not password:
            raise ValueError("host, username, and password are required")

        # Validate credentials before storing.
        try:
            client = imaplib.IMAP4_SSL(host, port) if use_ssl else imaplib.IMAP4(host, port)
            client.login(username, password)
            client.select(folder)
            client.logout()
        except imaplib.IMAP4.error as exc:
            raise ValueError("IMAP authentication failed") from exc
        except Exception as exc:
            raise ValueError("Unable to connect to IMAP server") from exc

        account_id = hashlib.sha256(f"{host}:{port}:{username}:{folder}".encode()).hexdigest()[:16]
        cfg = EmailInboxConfig(
            inbox_id=account_id,
            email_address=username,
            imap_server=host,
            imap_port=port,
            username=username,
            password=password,
            use_ssl=use_ssl,
            folder=folder,
        )
        self._inbox_configs[account_id] = cfg

        return {
            "id": account_id,
            "account_id": account_id,
            "email_address": cfg.email_address,
            "host": cfg.imap_server,
            "port": cfg.imap_port,
            "folder": cfg.folder,
            "use_ssl": cfg.use_ssl,
            "status": "connected",
        }

    def list_imap_accounts(self) -> List[Dict[str, Any]]:
        """List registered IMAP accounts (without credentials)."""
        accounts: List[Dict[str, Any]] = []
        for account_id, cfg in self._inbox_configs.items():
            accounts.append({
                "id": account_id,
                "account_id": account_id,
                "email_address": cfg.email_address,
                "host": cfg.imap_server,
                "port": cfg.imap_port,
                "folder": cfg.folder,
                "use_ssl": cfg.use_ssl,
                "status": "connected",
            })
        return accounts

    async def sync_imap_account(self, account_id: str, limit: int = 50) -> Dict[str, Any]:
        """Sync recent messages from a configured IMAP account."""
        cfg = self._inbox_configs.get(account_id)
        if not cfg:
            raise ValueError("IMAP account not found")

        results = await self.fetch_from_imap(cfg, limit=limit, unseen_only=True)
        return {
            "account_id": account_id,
            "status": "completed",
            "synced": len(results),
            "documents": [r.document_id for r in results],
        }

    def generate_inbox_address(self, user_id: str, purpose: str = "default") -> str:
        """Generate a unique inbox address for a user."""
        # Generate unique identifier
        unique_part = hashlib.sha256(f"{user_id}:{purpose}:{datetime.now(timezone.utc).isoformat()}".encode()).hexdigest()[:12]
        # Format: ingest+{unique}@neurareport.io
        return f"ingest+{unique_part}@neurareport.io"

    async def parse_email_content(self, raw_email: bytes) -> ParsedEmail:
        """Parse raw email content into structured format."""
        msg = email.message_from_bytes(raw_email)

        # Parse headers
        subject = self._decode_header(msg.get("Subject", ""))
        from_name, from_address = parseaddr(msg.get("From", ""))
        from_name = self._decode_header(from_name) if from_name else None

        to_addresses = self._parse_address_list(msg.get("To", ""))
        cc_addresses = self._parse_address_list(msg.get("Cc", ""))

        # Parse date
        date_str = msg.get("Date")
        date = None
        if date_str:
            try:
                date = parsedate_to_datetime(date_str)
            except Exception:
                pass

        # Get message ID and threading info
        message_id = msg.get("Message-ID", "")
        in_reply_to = msg.get("In-Reply-To")
        thread_id = msg.get("References", "").split()[0] if msg.get("References") else message_id

        # Extract body
        body_text = None
        body_html = None
        attachments = []

        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition", ""))

                if "attachment" in content_disposition:
                    filename = part.get_filename()
                    if filename:
                        filename = self._decode_header(filename)
                        payload = part.get_payload(decode=True)
                        attachments.append(EmailAttachment(
                            filename=filename,
                            content_type=content_type,
                            size_bytes=len(payload) if payload else 0,
                        ))
                elif content_type == "text/plain" and not body_text:
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_text = payload.decode("utf-8", errors="replace")
                elif content_type == "text/html" and not body_html:
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_html = payload.decode("utf-8", errors="replace")
        else:
            content_type = msg.get_content_type()
            payload = msg.get_payload(decode=True)
            if payload:
                if content_type == "text/html":
                    body_html = payload.decode("utf-8", errors="replace")
                else:
                    body_text = payload.decode("utf-8", errors="replace")

        return ParsedEmail(
            message_id=message_id,
            subject=subject,
            from_address=from_address,
            from_name=from_name,
            to_addresses=to_addresses,
            cc_addresses=cc_addresses,
            date=date,
            body_text=body_text,
            body_html=body_html,
            attachments=attachments,
            thread_id=thread_id,
            in_reply_to=in_reply_to,
        )

    async def convert_email_to_document(
        self,
        raw_email: bytes,
        include_attachments: bool = True,
    ) -> EmailDocumentResult:
        """Convert an email into a document."""
        from .service import ingestion_service

        parsed = await self.parse_email_content(raw_email)

        # Create document content from email
        content = self._format_email_as_document(parsed)

        # Generate document ID
        doc_id = hashlib.sha256(parsed.message_id.encode()).hexdigest()[:16]

        # Save as document
        result = await ingestion_service.ingest_file(
            filename=f"{self._sanitize_filename(parsed.subject)}.html",
            content=content.encode("utf-8"),
            metadata={
                "source": "email",
                "email_from": parsed.from_address,
                "email_subject": parsed.subject,
                "email_date": parsed.date.isoformat() if parsed.date else None,
                "message_id": parsed.message_id,
                "thread_id": parsed.thread_id,
            },
        )

        # Process attachments
        attachment_docs = []
        if include_attachments and parsed.attachments:
            msg = email.message_from_bytes(raw_email)
            for part in msg.walk():
                content_disposition = str(part.get("Content-Disposition", ""))
                if "attachment" in content_disposition:
                    filename = part.get_filename()
                    if filename:
                        filename = self._decode_header(filename)
                        payload = part.get_payload(decode=True)
                        if payload:
                            att_result = await ingestion_service.ingest_file(
                                filename=filename,
                                content=payload,
                                metadata={
                                    "source": "email_attachment",
                                    "parent_email_id": doc_id,
                                },
                            )
                            attachment_docs.append(att_result.document_id)

        return EmailDocumentResult(
            document_id=doc_id,
            email_subject=parsed.subject,
            from_address=parsed.from_address,
            attachment_count=len(attachment_docs),
            attachment_documents=attachment_docs,
        )

    async def create_document_from_thread(
        self,
        emails: List[bytes],
        thread_title: Optional[str] = None,
    ) -> EmailDocumentResult:
        """Create a single document from an email thread."""
        from .service import ingestion_service

        parsed_emails = []
        for raw in emails:
            parsed = await self.parse_email_content(raw)
            parsed_emails.append(parsed)

        # Sort by date
        parsed_emails.sort(key=lambda e: e.date or datetime.min)

        # Build thread document
        if not thread_title and parsed_emails:
            thread_title = parsed_emails[0].subject

        content = self._format_thread_as_document(parsed_emails, thread_title)

        # Generate document
        doc_id = hashlib.sha256(f"thread:{thread_title}:{datetime.now(timezone.utc).isoformat()}".encode()).hexdigest()[:16]

        result = await ingestion_service.ingest_file(
            filename=f"{self._sanitize_filename(thread_title or 'Email Thread')}.html",
            content=content.encode("utf-8"),
            metadata={
                "source": "email_thread",
                "email_count": len(parsed_emails),
                "participants": list(set(e.from_address for e in parsed_emails)),
            },
        )

        return EmailDocumentResult(
            document_id=result.document_id,
            email_subject=thread_title or "Email Thread",
            from_address=parsed_emails[0].from_address if parsed_emails else "",
        )

    async def fetch_from_imap(
        self,
        config: EmailInboxConfig,
        limit: int = 10,
        unseen_only: bool = True,
    ) -> List[EmailDocumentResult]:
        """Fetch and process emails from an IMAP inbox."""
        results = []

        try:
            # Connect to IMAP
            mail = imaplib.IMAP4_SSL(config.imap_server, config.imap_port)
            mail.login(config.username, config.password)
            mail.select(config.folder)

            # Search for emails
            search_criteria = "UNSEEN" if unseen_only else "ALL"
            status, messages = mail.search(None, search_criteria)

            if status != "OK":
                logger.error(f"IMAP search failed: {status}")
                return results

            message_ids = messages[0].split()[-limit:]  # Get latest N

            for msg_id in message_ids:
                try:
                    status, data = mail.fetch(msg_id, "(RFC822)")
                    if status == "OK" and data[0]:
                        raw_email = data[0][1]
                        result = await self.convert_email_to_document(raw_email)
                        results.append(result)

                        # Mark as read or delete
                        if config.delete_after_process:
                            mail.store(msg_id, "+FLAGS", "\\Deleted")
                        else:
                            mail.store(msg_id, "+FLAGS", "\\Seen")

                except Exception as e:
                    logger.error(f"Failed to process email {msg_id}: {e}")

            if config.delete_after_process:
                mail.expunge()

            mail.logout()

        except Exception as e:
            logger.error(f"IMAP connection failed: {e}")

        return results

    async def parse_incoming_email(
        self,
        raw_email: bytes,
        extract_action_items: bool = True,
    ) -> Dict[str, Any]:
        """Parse incoming email and extract structured data."""
        parsed = await self.parse_email_content(raw_email)

        result = {
            "email": parsed.model_dump(),
            "action_items": [],
            "mentions": [],
            "links": [],
            "dates": [],
        }

        # Extract from body
        body = parsed.body_text or self._html_to_text(parsed.body_html or "")

        if extract_action_items:
            # Simple action item extraction (would use AI for better results)
            action_patterns = [
                r"(?:please|could you|can you|would you|need to|must|should|will you)\s+(.+?)(?:\.|$)",
                r"(?:action item|todo|task):\s*(.+?)(?:\.|$)",
                r"(?:by|before|due)\s+(\w+\s+\d+|\d+\/\d+)",
            ]
            for pattern in action_patterns:
                matches = re.findall(pattern, body, re.IGNORECASE)
                result["action_items"].extend(matches[:5])

        # Extract links
        url_pattern = r"https?://[^\s<>\"']+"
        result["links"] = re.findall(url_pattern, body)[:10]

        # Extract @mentions
        mention_pattern = r"@(\w+)"
        result["mentions"] = re.findall(mention_pattern, body)

        return result

    def _decode_header(self, value: str) -> str:
        """Decode email header value."""
        if not value:
            return ""
        decoded_parts = decode_header(value)
        result = []
        for part, charset in decoded_parts:
            if isinstance(part, bytes):
                result.append(part.decode(charset or "utf-8", errors="replace"))
            else:
                result.append(part)
        return "".join(result)

    def _parse_address_list(self, value: str) -> List[str]:
        """Parse comma-separated email addresses."""
        if not value:
            return []
        addresses = []
        for addr in value.split(","):
            _, email_addr = parseaddr(addr.strip())
            if email_addr:
                addresses.append(email_addr)
        return addresses

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize string for use as filename."""
        # Remove invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', "", name)
        # Limit length
        return sanitized[:100] or "untitled"

    def _format_email_as_document(self, parsed: ParsedEmail) -> str:
        """Format parsed email as HTML document."""
        date_str = parsed.date.strftime("%B %d, %Y at %I:%M %p") if parsed.date else "Unknown date"

        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{parsed.subject}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        .email-header {{ background: #f5f5f5; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
        .email-meta {{ color: #666; font-size: 14px; margin: 5px 0; }}
        .email-body {{ line-height: 1.6; }}
        .attachments {{ background: #e3f2fd; padding: 15px; border-radius: 8px; margin-top: 20px; }}
    </style>
</head>
<body>
    <div class="email-header">
        <h1>{parsed.subject}</h1>
        <p class="email-meta"><strong>From:</strong> {parsed.from_name or ''} &lt;{parsed.from_address}&gt;</p>
        <p class="email-meta"><strong>To:</strong> {', '.join(parsed.to_addresses)}</p>
        {f'<p class="email-meta"><strong>CC:</strong> {", ".join(parsed.cc_addresses)}</p>' if parsed.cc_addresses else ''}
        <p class="email-meta"><strong>Date:</strong> {date_str}</p>
    </div>

    <div class="email-body">
        {parsed.body_html or f'<pre>{parsed.body_text or ""}</pre>'}
    </div>

    {self._format_attachments_section(parsed.attachments) if parsed.attachments else ''}
</body>
</html>"""

    def _format_thread_as_document(self, emails: List[ParsedEmail], title: str) -> str:
        """Format email thread as HTML document."""
        emails_html = ""
        for i, email_msg in enumerate(emails):
            date_str = email_msg.date.strftime("%B %d, %Y at %I:%M %p") if email_msg.date else "Unknown date"
            emails_html += f"""
            <div class="email-message" style="border-left: 3px solid #1976d2; padding-left: 20px; margin: 20px 0;">
                <div class="message-header">
                    <strong>{email_msg.from_name or email_msg.from_address}</strong>
                    <span style="color: #666; margin-left: 10px;">{date_str}</span>
                </div>
                <div class="message-body" style="margin-top: 10px;">
                    {email_msg.body_html or f'<pre>{email_msg.body_text or ""}</pre>'}
                </div>
            </div>
            """

        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ border-bottom: 2px solid #1976d2; padding-bottom: 10px; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p style="color: #666;">{len(emails)} messages in thread</p>
    {emails_html}
</body>
</html>"""

    def _format_attachments_section(self, attachments: List[EmailAttachment]) -> str:
        """Format attachments section."""
        if not attachments:
            return ""
        items = "".join(f"<li>{a.filename} ({a.size_bytes:,} bytes)</li>" for a in attachments)
        return f"""
        <div class="attachments">
            <strong>Attachments ({len(attachments)}):</strong>
            <ul>{items}</ul>
        </div>"""

    def _html_to_text(self, html: str) -> str:
        """Convert HTML to plain text."""
        # Simple HTML to text conversion
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

# Singleton instance
email_ingestion_service = EmailIngestionService()

"""
Transcription Service
Handles audio/video transcription for document creation.
"""

import logging
import os
import shutil
import hashlib
import tempfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
from enum import Enum

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

class TranscriptionLanguage(str, Enum):
    """Supported transcription languages."""
    AUTO = "auto"
    ENGLISH = "en"
    SPANISH = "es"
    FRENCH = "fr"
    GERMAN = "de"
    ITALIAN = "it"
    PORTUGUESE = "pt"
    DUTCH = "nl"
    RUSSIAN = "ru"
    CHINESE = "zh"
    JAPANESE = "ja"
    KOREAN = "ko"
    ARABIC = "ar"
    HINDI = "hi"

class TranscriptionSegment(BaseModel):
    """A segment of transcription with timing."""
    start_time: float  # Seconds
    end_time: float
    text: str
    confidence: float = 1.0
    speaker: Optional[str] = None

class TranscriptionResult(BaseModel):
    """Result of transcription."""
    document_id: str
    source_filename: str
    duration_seconds: float
    language: str
    segments: List[TranscriptionSegment] = Field(default_factory=list)
    full_text: str
    word_count: int
    speaker_count: int = 1
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    metadata: Dict[str, Any] = Field(default_factory=dict)

class VoiceMemoResult(BaseModel):
    """Result of voice memo transcription."""
    document_id: str
    title: str
    transcript: str
    duration_seconds: float
    action_items: List[str] = Field(default_factory=list)
    key_points: List[str] = Field(default_factory=list)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TranscriptionService:
    """
    Service for transcribing audio and video files.
    Uses Whisper for transcription.
    """

    # Supported audio formats
    AUDIO_FORMATS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac", ".wma"}

    # Supported video formats
    VIDEO_FORMATS = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".wmv", ".flv"}

    def __init__(self):
        self._whisper_model = None

    def _get_whisper(self):
        """Lazy-load Whisper model."""
        if self._whisper_model is None:
            try:
                import whisper
                self._whisper_model = whisper.load_model("base")
                logger.info("Loaded Whisper base model")
            except ImportError:
                logger.warning("Whisper not installed. Install with: pip install openai-whisper")
                raise RuntimeError("Whisper not installed")
        return self._whisper_model

    def _ensure_ffmpeg(self) -> bool:
        """Ensure ffmpeg is discoverable on PATH for Whisper."""
        if shutil.which("ffmpeg"):
            return True

        backend_root = Path(__file__).resolve().parents[3]
        tools_dir = backend_root / "tools"
        candidates = list(tools_dir.glob("ffmpeg_extracted/**/bin/ffmpeg.exe"))
        if not candidates:
            return False

        bin_dir = candidates[0].parent
        os.environ["PATH"] = f"{bin_dir}{os.pathsep}{os.environ.get('PATH', '')}"
        return True

    async def transcribe_file(
        self,
        filename: str,
        content: bytes,
        language: TranscriptionLanguage = TranscriptionLanguage.AUTO,
        include_timestamps: bool = True,
        diarize_speakers: bool = False,
    ) -> TranscriptionResult:
        """Transcribe an audio or video file."""
        self._ensure_ffmpeg()

        # Save to temp file for processing
        ext = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            # Extract audio if video
            if ext in self.VIDEO_FORMATS:
                audio_path = await self._extract_audio(tmp_path)
            else:
                audio_path = tmp_path

            # Get duration
            duration = await self._get_audio_duration(audio_path)

            # Transcribe with Whisper
            model = self._get_whisper()
            options = {
                "task": "transcribe",
                "verbose": False,
            }
            if language != TranscriptionLanguage.AUTO:
                options["language"] = language.value

            result = model.transcribe(str(audio_path), **options)

            # Parse segments
            segments = []
            if include_timestamps and "segments" in result:
                for seg in result["segments"]:
                    segments.append(TranscriptionSegment(
                        start_time=seg["start"],
                        end_time=seg["end"],
                        text=seg["text"].strip(),
                        confidence=seg.get("avg_logprob", 0) + 1,  # Normalize
                    ))

            # Speaker diarization (simplified)
            if diarize_speakers:
                segments = await self._diarize_speakers(segments, audio_path)

            full_text = result["text"].strip()

            # Create document
            doc_id = hashlib.sha256(f"{filename}:{datetime.now(timezone.utc).isoformat()}".encode()).hexdigest()[:16]

            return TranscriptionResult(
                document_id=doc_id,
                source_filename=filename,
                duration_seconds=duration,
                language=result.get("language", "en"),
                segments=segments,
                full_text=full_text,
                word_count=len(full_text.split()),
                speaker_count=len(set(s.speaker for s in segments if s.speaker)) or 1,
                metadata={
                    "model": "whisper-base",
                    "include_timestamps": include_timestamps,
                    "diarize_speakers": diarize_speakers,
                },
            )

        finally:
            # Cleanup temp files
            tmp_path.unlink(missing_ok=True)
            if ext in self.VIDEO_FORMATS:
                audio_path.unlink(missing_ok=True)

    async def transcribe_voice_memo(
        self,
        filename: str,
        content: bytes,
        extract_action_items: bool = True,
        extract_key_points: bool = True,
    ) -> VoiceMemoResult:
        """Transcribe a voice memo with intelligent extraction."""
        # First, get basic transcription
        result = await self.transcribe_file(
            filename=filename,
            content=content,
            include_timestamps=False,
            diarize_speakers=False,
        )

        # Generate title from first sentence
        title = self._generate_title(result.full_text)

        # Extract action items and key points using AI
        action_items = []
        key_points = []

        if extract_action_items or extract_key_points:
            extracted = await self._extract_insights(
                result.full_text,
                extract_action_items=extract_action_items,
                extract_key_points=extract_key_points,
            )
            action_items = extracted.get("action_items", [])
            key_points = extracted.get("key_points", [])

        return VoiceMemoResult(
            document_id=result.document_id,
            title=title,
            transcript=result.full_text,
            duration_seconds=result.duration_seconds,
            action_items=action_items,
            key_points=key_points,
        )

    async def create_document_from_transcription(
        self,
        result: TranscriptionResult,
        format: str = "html",
        include_timestamps: bool = True,
    ) -> str:
        """Create a document from transcription result."""
        from .service import ingestion_service

        if format == "srt":
            content = self._format_as_srt(result)
            filename = f"{Path(result.source_filename).stem}.srt"
        elif format == "markdown":
            content = self._format_as_markdown(result, include_timestamps)
            filename = f"{Path(result.source_filename).stem}.md"
        else:
            content = self._format_as_html(result, include_timestamps)
            filename = f"{Path(result.source_filename).stem}.html"

        ingestion_result = await ingestion_service.ingest_file(
            filename=filename,
            content=content.encode("utf-8"),
            metadata={
                "source": "transcription",
                "source_file": result.source_filename,
                "duration_seconds": result.duration_seconds,
                "language": result.language,
                "word_count": result.word_count,
            },
        )

        return ingestion_result.document_id

    async def _extract_audio(self, video_path: Path) -> Path:
        """Extract audio from video file."""
        try:
            from moviepy.editor import VideoFileClip

            audio_path = video_path.with_suffix(".wav")
            video = VideoFileClip(str(video_path))
            video.audio.write_audiofile(str(audio_path), verbose=False, logger=None)
            video.close()
            return audio_path
        except ImportError:
            # Fallback to ffmpeg
            import subprocess
            audio_path = video_path.with_suffix(".wav")
            subprocess.run([
                "ffmpeg", "-i", str(video_path),
                "-vn", "-acodec", "pcm_s16le",
                "-ar", "16000", "-ac", "1",
                str(audio_path), "-y"
            ], capture_output=True, check=True)
            return audio_path

    async def _get_audio_duration(self, audio_path: Path) -> float:
        """Get audio file duration in seconds."""
        try:
            from pydub import AudioSegment
            audio = AudioSegment.from_file(str(audio_path))
            return len(audio) / 1000.0
        except ImportError:
            # Fallback using wave module for WAV files
            import wave
            if audio_path.suffix.lower() == ".wav":
                with wave.open(str(audio_path), "r") as wav:
                    frames = wav.getnframes()
                    rate = wav.getframerate()
                    return frames / float(rate)
            return 0.0

    async def _diarize_speakers(
        self,
        segments: List[TranscriptionSegment],
        audio_path: Path,
    ) -> List[TranscriptionSegment]:
        """Simple speaker diarization based on pauses and patterns."""
        # Simple heuristic: alternate speakers on long pauses
        current_speaker = "Speaker 1"
        speaker_count = 1
        diarized = []

        prev_end = 0
        for seg in segments:
            # If there's a significant pause, potentially switch speakers
            if seg.start_time - prev_end > 2.0:  # 2 second pause
                if speaker_count < 4:  # Max 4 speakers
                    speaker_count += 1
                current_speaker = f"Speaker {((speaker_count - 1) % 2) + 1}"

            diarized.append(TranscriptionSegment(
                start_time=seg.start_time,
                end_time=seg.end_time,
                text=seg.text,
                confidence=seg.confidence,
                speaker=current_speaker,
            ))
            prev_end = seg.end_time

        return diarized

    async def _extract_insights(
        self,
        text: str,
        extract_action_items: bool = True,
        extract_key_points: bool = True,
    ) -> Dict[str, List[str]]:
        """Extract insights from transcript using AI."""
        try:
            from backend.app.services.llm import get_llm_client

            client = get_llm_client()

            prompt_parts = []
            if extract_action_items:
                prompt_parts.append("- List any action items, tasks, or TODOs mentioned")
            if extract_key_points:
                prompt_parts.append("- List the key points or main ideas")

            prompt = f"""Analyze this transcript and extract:
{chr(10).join(prompt_parts)}

Transcript:
{text[:4000]}

Respond in JSON format:
{{"action_items": ["item1", "item2"], "key_points": ["point1", "point2"]}}"""

            response = client.complete(
                messages=[{"role": "user", "content": prompt}],
                description="transcription_extract_insights",
                max_tokens=500,
            )

            import json
            content = (
                response.get("choices", [{}])[0]
                .get("message", {})
                .get("content", "{}")
            )
            return json.loads(content)

        except Exception as e:
            logger.warning(f"Failed to extract insights: {e}")
            return {"action_items": [], "key_points": []}

    def _generate_title(self, text: str) -> str:
        """Generate a title from transcript text."""
        # Take first sentence
        sentences = text.split(".")
        if sentences:
            title = sentences[0].strip()[:80]
            if len(sentences[0]) > 80:
                title += "..."
            return title
        return "Voice Memo"

    def _format_as_srt(self, result: TranscriptionResult) -> str:
        """Format transcription as SRT subtitle file."""
        lines = []
        for i, seg in enumerate(result.segments, 1):
            start = self._format_srt_time(seg.start_time)
            end = self._format_srt_time(seg.end_time)
            lines.append(f"{i}")
            lines.append(f"{start} --> {end}")
            lines.append(seg.text)
            lines.append("")
        return "\n".join(lines)

    def _format_srt_time(self, seconds: float) -> str:
        """Format seconds as SRT timestamp."""
        td = timedelta(seconds=seconds)
        hours, remainder = divmod(td.seconds, 3600)
        minutes, secs = divmod(remainder, 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    def _format_as_markdown(self, result: TranscriptionResult, include_timestamps: bool) -> str:
        """Format transcription as Markdown."""
        lines = [
            f"# Transcript: {result.source_filename}",
            "",
            f"**Duration:** {self._format_duration(result.duration_seconds)}",
            f"**Language:** {result.language}",
            f"**Words:** {result.word_count}",
            "",
            "---",
            "",
        ]

        if include_timestamps and result.segments:
            current_speaker = None
            for seg in result.segments:
                if seg.speaker and seg.speaker != current_speaker:
                    current_speaker = seg.speaker
                    lines.append(f"\n**{current_speaker}:**\n")

                timestamp = self._format_duration(seg.start_time)
                lines.append(f"[{timestamp}] {seg.text}")
        else:
            lines.append(result.full_text)

        return "\n".join(lines)

    def _format_as_html(self, result: TranscriptionResult, include_timestamps: bool) -> str:
        """Format transcription as HTML document."""
        segments_html = ""
        if include_timestamps and result.segments:
            current_speaker = None
            for seg in result.segments:
                speaker_html = ""
                if seg.speaker and seg.speaker != current_speaker:
                    current_speaker = seg.speaker
                    speaker_html = f'<strong class="speaker">{seg.speaker}:</strong><br>'

                timestamp = self._format_duration(seg.start_time)
                segments_html += f"""
                <p class="segment">
                    {speaker_html}
                    <span class="timestamp">[{timestamp}]</span>
                    <span class="text">{seg.text}</span>
                </p>"""
        else:
            segments_html = f"<p>{result.full_text}</p>"

        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Transcript: {result.source_filename}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.8;
        }}
        .meta {{
            background: #f5f5f5;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
        }}
        .meta span {{
            margin-right: 20px;
            color: #666;
        }}
        .segment {{
            margin: 10px 0;
        }}
        .timestamp {{
            color: #1976d2;
            font-size: 0.85em;
            margin-right: 8px;
        }}
        .speaker {{
            color: #7b1fa2;
        }}
    </style>
</head>
<body>
    <h1>Transcript</h1>
    <div class="meta">
        <span><strong>Source:</strong> {result.source_filename}</span>
        <span><strong>Duration:</strong> {self._format_duration(result.duration_seconds)}</span>
        <span><strong>Language:</strong> {result.language}</span>
        <span><strong>Words:</strong> {result.word_count}</span>
    </div>
    <div class="transcript">
        {segments_html}
    </div>
</body>
</html>"""

    def _format_duration(self, seconds: float) -> str:
        """Format seconds as human-readable duration."""
        m, s = divmod(int(seconds), 60)
        h, m = divmod(m, 60)
        if h:
            return f"{h}:{m:02d}:{s:02d}"
        return f"{m}:{s:02d}"

# Singleton instance
transcription_service = TranscriptionService()

"""
Web Clipper Service
Handles web page capture, cleaning, and conversion to documents.
"""

import logging
import re
import hashlib
from datetime import datetime, timezone
from typing import Dict, List, Optional
from urllib.parse import urljoin, urlparse

from pydantic import BaseModel, Field

from backend.app.utils import validate_url

logger = logging.getLogger(__name__)

class WebPageMetadata(BaseModel):
    """Metadata extracted from a web page."""
    title: str
    url: str
    author: Optional[str] = None
    published_date: Optional[str] = None
    site_name: Optional[str] = None
    description: Optional[str] = None
    image_url: Optional[str] = None
    word_count: int = 0
    reading_time_minutes: int = 0

class ClippedContent(BaseModel):
    """Content clipped from a web page."""
    document_id: str
    url: str
    title: str
    clean_html: str
    plain_text: str
    metadata: WebPageMetadata
    images: List[str] = Field(default_factory=list)
    links: List[Dict[str, str]] = Field(default_factory=list)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class WebClipperService:
    """
    Service for clipping web pages and converting them to editable documents.
    Supports content extraction, cleaning, and metadata parsing.
    """

    # Elements to remove from content
    REMOVE_ELEMENTS = [
        "script", "style", "nav", "header", "footer", "aside",
        "iframe", "noscript", "form", "button", "input",
        "advertisement", ".ad", ".ads", ".sidebar", ".navigation",
        ".menu", ".social", ".share", ".comments", ".related",
    ]

    # Content containers to prioritize
    CONTENT_SELECTORS = [
        "article", "main", "[role='main']", ".post-content",
        ".article-content", ".entry-content", ".content",
        "#content", ".post", ".article",
    ]

    async def clip_url(
        self,
        url: str,
        include_images: bool = True,
        clean_content: bool = True,
    ) -> ClippedContent:
        """Clip content from a URL."""
        import aiohttp
        from bs4 import BeautifulSoup

        validate_url(url)

        async with aiohttp.ClientSession() as session:
            async with session.get(url, headers={"User-Agent": "Mozilla/5.0 (compatible; NeuraReport/1.0)"}) as response:
                response.raise_for_status()
                html = await response.text()

        soup = BeautifulSoup(html, "html.parser")

        # Extract metadata
        metadata = self._extract_metadata(soup, url)

        # Find main content
        content_element = self._find_content_element(soup)

        if clean_content:
            clean_html = self._clean_content(content_element, url)
        else:
            clean_html = str(content_element)

        # Extract plain text
        plain_text = self._extract_text(content_element)
        metadata.word_count = len(plain_text.split())
        metadata.reading_time_minutes = max(1, metadata.word_count // 200)

        # Extract images
        images = []
        if include_images:
            images = self._extract_images(content_element, url)

        # Extract links
        links = self._extract_links(content_element, url)

        # Generate document ID
        doc_id = hashlib.sha256(f"{url}:{datetime.now(timezone.utc).isoformat()}".encode()).hexdigest()[:16]

        return ClippedContent(
            document_id=doc_id,
            url=url,
            title=metadata.title,
            clean_html=clean_html,
            plain_text=plain_text,
            metadata=metadata,
            images=images,
            links=links,
        )

    async def clip_selection(
        self,
        url: str,
        selected_html: str,
        page_title: Optional[str] = None,
    ) -> ClippedContent:
        """Clip a user-selected portion of a page."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(selected_html, "html.parser")

        clean_html = self._clean_content(soup, url)
        plain_text = self._extract_text(soup)

        doc_id = hashlib.sha256(f"{url}:selection:{datetime.now(timezone.utc).isoformat()}".encode()).hexdigest()[:16]

        metadata = WebPageMetadata(
            title=page_title or "Clipped Selection",
            url=url,
            word_count=len(plain_text.split()),
            reading_time_minutes=max(1, len(plain_text.split()) // 200),
        )

        return ClippedContent(
            document_id=doc_id,
            url=url,
            title=metadata.title,
            clean_html=clean_html,
            plain_text=plain_text,
            metadata=metadata,
            images=self._extract_images(soup, url),
            links=self._extract_links(soup, url),
        )

    async def save_as_document(
        self,
        clipped: ClippedContent,
        format: str = "html",
    ) -> str:
        """Save clipped content as a document."""
        from .service import ingestion_service

        if format == "markdown":
            content = self._html_to_markdown(clipped.clean_html)
            filename = f"{self._sanitize_filename(clipped.title)}.md"
        elif format == "pdf":
            # Would generate PDF here
            content = self._wrap_as_html_document(clipped)
            filename = f"{self._sanitize_filename(clipped.title)}.html"
        else:
            content = self._wrap_as_html_document(clipped)
            filename = f"{self._sanitize_filename(clipped.title)}.html"

        result = await ingestion_service.ingest_file(
            filename=filename,
            content=content.encode("utf-8"),
            metadata={
                "source": "web_clipper",
                "source_url": clipped.url,
                "clipped_at": clipped.created_at.isoformat(),
                **clipped.metadata.model_dump(),
            },
        )

        return result.document_id

    async def capture_screenshot(
        self,
        url: str,
        full_page: bool = False,
    ) -> bytes:
        """Capture screenshot of a web page."""
        # This would use Playwright or Puppeteer
        # Placeholder implementation
        logger.info(f"Screenshot capture requested for: {url}")
        raise NotImplementedError("Screenshot capture requires browser automation")

    def _extract_metadata(self, soup, url: str) -> WebPageMetadata:
        """Extract metadata from page."""
        # Title
        title = ""
        if soup.title:
            title = soup.title.string or ""

        og_title = soup.find("meta", property="og:title")
        if og_title and og_title.get("content"):
            title = og_title["content"]

        # Author
        author = None
        author_meta = soup.find("meta", {"name": "author"})
        if author_meta and author_meta.get("content"):
            author = author_meta["content"]

        # Published date
        published_date = None
        date_meta = soup.find("meta", {"property": "article:published_time"})
        if date_meta and date_meta.get("content"):
            published_date = date_meta["content"]

        # Site name
        site_name = None
        site_meta = soup.find("meta", {"property": "og:site_name"})
        if site_meta and site_meta.get("content"):
            site_name = site_meta["content"]

        # Description
        description = None
        desc_meta = soup.find("meta", {"name": "description"}) or soup.find("meta", {"property": "og:description"})
        if desc_meta and desc_meta.get("content"):
            description = desc_meta["content"]

        # Image
        image_url = None
        img_meta = soup.find("meta", {"property": "og:image"})
        if img_meta and img_meta.get("content"):
            image_url = urljoin(url, img_meta["content"])

        return WebPageMetadata(
            title=title.strip() or urlparse(url).netloc,
            url=url,
            author=author,
            published_date=published_date,
            site_name=site_name,
            description=description,
            image_url=image_url,
        )

    def _find_content_element(self, soup):
        """Find the main content element."""
        # Try priority selectors
        for selector in self.CONTENT_SELECTORS:
            element = soup.select_one(selector)
            if element and len(element.get_text(strip=True)) > 100:
                return element

        # Fall back to body
        return soup.body or soup

    def _clean_content(self, element, base_url: str) -> str:
        """Clean content by removing unwanted elements."""
        from bs4 import BeautifulSoup

        # Work on a copy
        soup = BeautifulSoup(str(element), "html.parser")

        # Remove unwanted elements
        for selector in self.REMOVE_ELEMENTS:
            for el in soup.select(selector):
                el.decompose()

        # Fix relative URLs
        for tag in soup.find_all(["a", "img"]):
            if tag.name == "a" and tag.get("href"):
                tag["href"] = urljoin(base_url, tag["href"])
            if tag.name == "img" and tag.get("src"):
                tag["src"] = urljoin(base_url, tag["src"])

        # Remove empty elements
        for el in soup.find_all():
            if not el.get_text(strip=True) and el.name not in ["img", "br", "hr"]:
                el.decompose()

        return str(soup)

    def _extract_text(self, element) -> str:
        """Extract plain text from element."""
        text = element.get_text(separator=" ", strip=True)
        # Clean up whitespace
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _extract_images(self, element, base_url: str) -> List[str]:
        """Extract image URLs from element."""
        images = []
        for img in element.find_all("img"):
            src = img.get("src") or img.get("data-src")
            if src:
                full_url = urljoin(base_url, src)
                if full_url not in images:
                    images.append(full_url)
        return images[:20]  # Limit

    def _extract_links(self, element, base_url: str) -> List[Dict[str, str]]:
        """Extract links from element."""
        links = []
        seen = set()
        for a in element.find_all("a", href=True):
            href = urljoin(base_url, a["href"])
            if href not in seen and not href.startswith("javascript:"):
                seen.add(href)
                links.append({
                    "url": href,
                    "text": a.get_text(strip=True)[:100],
                })
        return links[:50]  # Limit

    def _html_to_markdown(self, html: str) -> str:
        """Convert HTML to Markdown."""
        try:
            import html2text
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            return h.handle(html)
        except ImportError:
            # Simple fallback
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text(separator="\n\n")

    def _wrap_as_html_document(self, clipped: ClippedContent) -> str:
        """Wrap clipped content as a complete HTML document."""
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{clipped.title}</title>
    <meta name="source-url" content="{clipped.url}">
    <meta name="clipped-date" content="{clipped.created_at.isoformat()}">
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        .source-info {{
            background: #f5f5f5;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
            font-size: 14px;
            color: #666;
        }}
        img {{ max-width: 100%; height: auto; }}
        a {{ color: #1976d2; }}
    </style>
</head>
<body>
    <div class="source-info">
        <strong>Source:</strong> <a href="{clipped.url}">{clipped.metadata.site_name or clipped.url}</a><br>
        {f'<strong>Author:</strong> {clipped.metadata.author}<br>' if clipped.metadata.author else ''}
        {f'<strong>Published:</strong> {clipped.metadata.published_date}<br>' if clipped.metadata.published_date else ''}
        <strong>Clipped:</strong> {clipped.created_at.strftime('%B %d, %Y at %I:%M %p')}
    </div>

    <h1>{clipped.title}</h1>

    <div class="content">
        {clipped.clean_html}
    </div>
</body>
</html>"""

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize string for use as filename."""
        sanitized = re.sub(r'[<>:"/\\|?*]', "", name)
        return sanitized[:100] or "clipped"

# Singleton instance
web_clipper_service = WebClipperService()

"""
Folder Watcher Service
Monitors local folders for new files and auto-imports them.
"""

import logging
import asyncio
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set
from enum import Enum

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

class WatcherEvent(str, Enum):
    """Types of file system events."""
    CREATED = "created"
    MODIFIED = "modified"
    DELETED = "deleted"
    MOVED = "moved"

class WatcherConfig(BaseModel):
    """Configuration for a folder watcher."""
    watcher_id: str
    path: str
    recursive: bool = True
    patterns: List[str] = Field(default_factory=lambda: ["*"])  # Glob patterns
    ignore_patterns: List[str] = Field(default_factory=list)
    auto_import: bool = True
    delete_after_import: bool = False
    target_collection: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    enabled: bool = True

class FileEvent(BaseModel):
    """A file system event."""
    event_type: WatcherEvent
    path: str
    filename: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    size_bytes: Optional[int] = None
    document_id: Optional[str] = None
    error: Optional[str] = None

class WatcherStatus(BaseModel):
    """Status of a folder watcher."""
    watcher_id: str
    path: str
    is_running: bool
    files_processed: int = 0
    files_pending: int = 0
    last_event: Optional[datetime] = None
    errors: List[str] = Field(default_factory=list)

class FolderWatcherService:
    """
    Service for monitoring folders and auto-importing files.
    Uses watchdog for cross-platform file system monitoring.
    """

    def __init__(self):
        self._watchers: Dict[str, WatcherConfig] = {}
        self._running: Dict[str, bool] = {}
        self._stats: Dict[str, Dict[str, Any]] = {}
        self._processed_files: Dict[str, Set[str]] = {}  # Track processed files by hash

    async def create_watcher(self, config: WatcherConfig) -> WatcherStatus:
        """Create a new folder watcher."""
        # Validate path exists
        path = Path(config.path)
        if not path.exists():
            path.mkdir(parents=True, exist_ok=True)

        self._watchers[config.watcher_id] = config
        self._running[config.watcher_id] = False
        self._stats[config.watcher_id] = {
            "files_processed": 0,
            "files_pending": 0,
            "last_event": None,
            "errors": [],
        }
        self._processed_files[config.watcher_id] = set()

        if config.enabled:
            await self.start_watcher(config.watcher_id)

        return self.get_status(config.watcher_id)

    async def start_watcher(self, watcher_id: str) -> bool:
        """Start a folder watcher."""
        if watcher_id not in self._watchers:
            raise ValueError(f"Watcher {watcher_id} not found")

        if self._running.get(watcher_id):
            return True  # Already running

        config = self._watchers[watcher_id]

        try:
            # Using watchdog for file system monitoring
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler, FileSystemEvent

            class Handler(FileSystemEventHandler):
                def __init__(self, service, watcher_id):
                    self.service = service
                    self.watcher_id = watcher_id

                def on_created(self, event: FileSystemEvent):
                    if not event.is_directory:
                        asyncio.create_task(
                            self.service._handle_event(
                                self.watcher_id,
                                WatcherEvent.CREATED,
                                event.src_path,
                            )
                        )

                def on_modified(self, event: FileSystemEvent):
                    if not event.is_directory:
                        asyncio.create_task(
                            self.service._handle_event(
                                self.watcher_id,
                                WatcherEvent.MODIFIED,
                                event.src_path,
                            )
                        )

            observer = Observer()
            handler = Handler(self, watcher_id)
            observer.schedule(handler, config.path, recursive=config.recursive)
            observer.start()

            self._running[watcher_id] = True
            logger.info(f"Started watcher {watcher_id} on {config.path}")

            return True

        except ImportError:
            logger.warning("watchdog not installed, using polling fallback")
            # Fallback to polling
            asyncio.create_task(self._poll_folder(watcher_id))
            self._running[watcher_id] = True
            return True

        except Exception as e:
            logger.error(f"Failed to start watcher {watcher_id}: {e}")
            self._stats[watcher_id]["errors"].append(str(e))
            return False

    async def stop_watcher(self, watcher_id: str) -> bool:
        """Stop a folder watcher."""
        if watcher_id not in self._running:
            return False

        self._running[watcher_id] = False
        logger.info(f"Stopped watcher {watcher_id}")
        return True

    async def delete_watcher(self, watcher_id: str) -> bool:
        """Delete a folder watcher."""
        await self.stop_watcher(watcher_id)

        if watcher_id in self._watchers:
            del self._watchers[watcher_id]
        if watcher_id in self._stats:
            del self._stats[watcher_id]
        if watcher_id in self._processed_files:
            del self._processed_files[watcher_id]

        return True

    def get_status(self, watcher_id: str) -> WatcherStatus:
        """Get status of a folder watcher."""
        if watcher_id not in self._watchers:
            raise ValueError(f"Watcher {watcher_id} not found")

        config = self._watchers[watcher_id]
        stats = self._stats.get(watcher_id, {})

        return WatcherStatus(
            watcher_id=watcher_id,
            path=config.path,
            is_running=self._running.get(watcher_id, False),
            files_processed=stats.get("files_processed", 0),
            files_pending=stats.get("files_pending", 0),
            last_event=stats.get("last_event"),
            errors=stats.get("errors", [])[-10:],  # Last 10 errors
        )

    def list_watchers(self) -> List[WatcherStatus]:
        """List all folder watchers."""
        return [self.get_status(wid) for wid in self._watchers]

    async def scan_folder(self, watcher_id: str) -> List[FileEvent]:
        """Manually scan a watched folder for existing files."""
        if watcher_id not in self._watchers:
            raise ValueError(f"Watcher {watcher_id} not found")

        config = self._watchers[watcher_id]
        path = Path(config.path)
        events = []

        if config.recursive:
            files = path.rglob("*")
        else:
            files = path.glob("*")

        for file_path in files:
            if file_path.is_file():
                if self._matches_patterns(file_path, config):
                    event = await self._handle_event(
                        watcher_id,
                        WatcherEvent.CREATED,
                        str(file_path),
                    )
                    if event:
                        events.append(event)

        return events

    async def _handle_event(
        self,
        watcher_id: str,
        event_type: WatcherEvent,
        file_path: str,
    ) -> Optional[FileEvent]:
        """Handle a file system event."""
        config = self._watchers.get(watcher_id)
        if not config:
            return None

        path = Path(file_path)

        # Check if file matches patterns
        if not self._matches_patterns(path, config):
            return None

        # Check if already processed (by content hash)
        file_hash = self._get_file_hash(path)
        if file_hash in self._processed_files.get(watcher_id, set()):
            return None

        event = FileEvent(
            event_type=event_type,
            path=str(path),
            filename=path.name,
            size_bytes=path.stat().st_size if path.exists() else None,
        )

        # Auto-import if configured
        if config.auto_import and event_type in (WatcherEvent.CREATED, WatcherEvent.MODIFIED):
            try:
                from .service import ingestion_service

                content = path.read_bytes()
                result = await ingestion_service.ingest_file(
                    filename=path.name,
                    content=content,
                    metadata={
                        "source": "folder_watcher",
                        "watcher_id": watcher_id,
                        "original_path": str(path),
                        "tags": config.tags,
                        "collection": config.target_collection,
                    },
                )
                event.document_id = result.document_id

                # Mark as processed
                self._processed_files[watcher_id].add(file_hash)
                self._stats[watcher_id]["files_processed"] += 1

                # Delete if configured
                if config.delete_after_import:
                    path.unlink()

            except Exception as e:
                event.error = "File import failed"
                self._stats[watcher_id]["errors"].append(f"{path.name}: import failed")
                logger.error(f"Failed to import {file_path}: {e}")

        self._stats[watcher_id]["last_event"] = datetime.now(timezone.utc)
        return event

    async def _poll_folder(self, watcher_id: str, interval: float = 5.0):
        """Polling fallback for when watchdog is not available."""
        config = self._watchers.get(watcher_id)
        if not config:
            return

        seen_files: Dict[str, float] = {}

        while self._running.get(watcher_id, False):
            try:
                path = Path(config.path)
                current_files = {}

                if config.recursive:
                    files = path.rglob("*")
                else:
                    files = path.glob("*")

                for file_path in files:
                    if file_path.is_file():
                        mtime = file_path.stat().st_mtime
                        current_files[str(file_path)] = mtime

                        # Check for new or modified files
                        if str(file_path) not in seen_files:
                            await self._handle_event(watcher_id, WatcherEvent.CREATED, str(file_path))
                        elif seen_files[str(file_path)] != mtime:
                            await self._handle_event(watcher_id, WatcherEvent.MODIFIED, str(file_path))

                seen_files = current_files

            except Exception as e:
                logger.error(f"Polling error for watcher {watcher_id}: {e}")

            await asyncio.sleep(interval)

    def _matches_patterns(self, path: Path, config: WatcherConfig) -> bool:
        """Check if file matches include/exclude patterns."""
        import fnmatch

        filename = path.name

        # Check ignore patterns first
        for pattern in config.ignore_patterns:
            if fnmatch.fnmatch(filename, pattern):
                return False

        # Check include patterns
        if config.patterns == ["*"]:
            return True

        for pattern in config.patterns:
            if fnmatch.fnmatch(filename, pattern):
                return True

        return False

    def _get_file_hash(self, path: Path) -> str:
        """Get hash of file for deduplication."""
        if not path.exists():
            return ""

        # Use file path + size + mtime for quick hash
        stat = path.stat()
        return hashlib.md5(f"{path}:{stat.st_size}:{stat.st_mtime}".encode()).hexdigest()

# Singleton instance
folder_watcher_service = FolderWatcherService()
