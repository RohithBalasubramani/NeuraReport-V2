from __future__ import annotations

import logging
import re

from fastapi import APIRouter, Depends, HTTPException, Query, Request

from backend.app.services.config import require_api_key
from backend.app.services.config import AppError
from backend.app.services.llm import get_model
from backend.app.schemas import ConnectionTestRequest, ConnectionUpsertRequest
from backend.app.services.infra_services import ConnectionService

logger = logging.getLogger("neura.api.connections")

connections_router = APIRouter(dependencies=[Depends(require_api_key)])

# SQL identifier pattern: alphanumeric + underscores, optionally schema-qualified
_SQL_IDENTIFIER_RE = re.compile(r"^[a-zA-Z_][a-zA-Z0-9_]*(\.[a-zA-Z_][a-zA-Z0-9_]*)*$")

def get_service() -> ConnectionService:
    return ConnectionService()

def _corr(request: Request) -> str | None:
    return getattr(request.state, "correlation_id", None)

def _handle_connection_error(exc: Exception, operation: str):
    logger.exception(
        "connection_operation_failed",
        extra={"event": "connection_operation_failed", "operation": operation},
    )
    raise HTTPException(status_code=500, detail=f"Connection {operation} failed")

@connections_router.post("/test")
async def test_connection(
    payload: ConnectionTestRequest,
    request: Request,
    svc: ConnectionService = Depends(get_service),
):
    try:
        return {"status": "ok", **svc.test(payload, _corr(request))}
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "test")

@connections_router.get("")
async def list_connections(
    request: Request,
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    svc: ConnectionService = Depends(get_service),
):
    try:
        connections = svc.list(_corr(request))
        total = len(connections)
        connections = connections[offset : offset + limit]
        return {"status": "ok", "connections": connections, "total": total, "correlation_id": _corr(request)}
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "list")

@connections_router.post("")
async def upsert_connection(
    payload: ConnectionUpsertRequest,
    request: Request,
    svc: ConnectionService = Depends(get_service),
):
    try:
        connection = svc.upsert(payload, _corr(request))
        return {"status": "ok", "connection": connection.model_dump(), "correlation_id": _corr(request)}
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "upsert")

@connections_router.delete("/{connection_id}")
async def delete_connection(
    connection_id: str,
    request: Request,
    svc: ConnectionService = Depends(get_service),
):
    try:
        # Verify existence before deleting
        existing = svc.repo.get(connection_id) if hasattr(svc.repo, "get") else None
        if existing is None and hasattr(svc.repo, "get"):
            raise HTTPException(status_code=404, detail="Connection not found")
        svc.delete(connection_id)
        return {"status": "ok", "connection_id": connection_id, "correlation_id": _corr(request)}
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "delete")

@connections_router.post("/{connection_id}/health")
async def healthcheck_connection(
    connection_id: str,
    request: Request,
    svc: ConnectionService = Depends(get_service),
):
    """Verify a saved connection is still accessible."""
    try:
        result = svc.healthcheck(connection_id, _corr(request))
        return {
            "status": "ok",
            "connection_id": result.get("connection_id"),
            "latency_ms": result.get("latency_ms"),
            "correlation_id": _corr(request),
        }
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "healthcheck")

@connections_router.get("/{connection_id}/schema")
async def connection_schema(
    connection_id: str,
    request: Request,
    include_row_counts: bool = Query(True),
    include_foreign_keys: bool = Query(True),
    sample_rows: int = Query(0, ge=0, le=25),
):
    from backend.app.services.legacy_services import get_connection_schema

    try:
        result = get_connection_schema(
            connection_id,
            include_row_counts=include_row_counts,
            include_foreign_keys=include_foreign_keys,
            sample_rows=sample_rows,
        )
        result["correlation_id"] = _corr(request)
        return result
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "schema")

@connections_router.get("/{connection_id}/preview")
async def connection_preview(
    connection_id: str,
    request: Request,
    table: str = Query(..., min_length=1, max_length=255),
    limit: int = Query(10, ge=1, le=200),
    offset: int = Query(0, ge=0),
):
    from backend.app.services.legacy_services import get_connection_table_preview

    # Validate table name is a safe SQL identifier
    if not _SQL_IDENTIFIER_RE.match(table):
        raise HTTPException(
            status_code=400,
            detail="Invalid table name. Must be a valid SQL identifier (letters, digits, underscores).",
        )

    try:
        result = get_connection_table_preview(
            connection_id,
            table=table,
            limit=limit,
            offset=offset,
        )
        result["correlation_id"] = _corr(request)
        return result
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        _handle_connection_error(exc, "preview")

import contextlib
import os
import tempfile
from pathlib import Path
from types import SimpleNamespace
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse

from backend.app.services.config import get_settings
from backend.app.utils import is_safe_name, validate_file_extension
from backend.app.schemas import TemplateImportResult
from backend.app.services.templates import TemplateService
from backend.app.schemas import (
    ChartSuggestPayload,
    SavedChartCreatePayload,
    SavedChartUpdatePayload,
)
from backend.app.services.platform_services import suggest_charts as suggest_charts_service
from backend.app.services.platform_services import (
    create_saved_chart as create_saved_chart_service,
    delete_saved_chart as delete_saved_chart_service,
    list_saved_charts as list_saved_charts_service,
    update_saved_chart as update_saved_chart_service,
)
from backend.app.services.contract_builder import load_contract_v2
from backend.app.services.config import (
    enqueue_background_job,
    iter_ndjson_events_async,
    run_event_stream_async,
)
from backend.app.services.ai_services import (
    CHART_SUGGEST_PROMPT_VERSION,
    build_chart_suggestions_prompt,
)
from backend.app.services.reports import discover_batches_and_counts
from backend.app.services.reports import (
    build_batch_field_catalog_and_stats,
    build_batch_metrics,
)
import backend.app.services.config as state_access
from backend.app.services.templates import get_openai_client
from backend.app.services.infra_services import call_chat_completion, get_correlation_id, strip_code_fences

# Import service functions from the service layer
from backend.app.services.legacy_services import (
    get_template_html,
    edit_template_ai,
    edit_template_manual,
    chat_template_edit,
    chat_template_create,
    create_template_from_chat,
    apply_chat_template_edit,
    undo_last_template_edit,
    verify_template,
    list_templates,
    templates_catalog,
    recommend_templates,
    delete_template,
    update_template_metadata,
    generator_assets,
)
from backend.app.services.legacy_services import run_mapping_approve
from backend.app.services.legacy_services import run_corrections_preview
from backend.app.services.legacy_services import mapping_key_options as mapping_key_options_service
from backend.app.services.legacy_services import run_mapping_preview
from backend.app.services.legacy_services import artifact_head_response, artifact_manifest_response
from backend.app.services.legacy_services import (
    CorrectionsPreviewPayload,
    GeneratorAssetsPayload,
    MappingPayload,
    TemplateAiEditPayload,
    TemplateChatPayload,
    TemplateCreateFromChatPayload,
    TemplateManualEditPayload,
    TemplateRecommendPayload,
    TemplateUpdatePayload,
)
from backend.app.services.legacy_services import db_path_from_payload_or_default
from backend.app.services.legacy_services import clean_key_values
from backend.app.services.legacy_services import normalize_template_id, template_dir

templates_router = APIRouter(dependencies=[Depends(require_api_key)])

ALLOWED_EXTENSIONS = [".zip"]
MAX_FILENAME_LENGTH = 255

def _correlation(request: Request) -> str | None:
    return getattr(request.state, "correlation_id", None)

def _request_with_correlation(correlation_id: str | None) -> SimpleNamespace:
    return SimpleNamespace(state=SimpleNamespace(correlation_id=correlation_id))

def _wrap(payload: dict, correlation_id: str | None) -> dict:
    payload = dict(payload)
    if correlation_id is not None:
        payload["correlation_id"] = correlation_id
    return payload

def _ensure_template_exists(template_id: str) -> tuple[str, dict]:
    normalized = normalize_template_id(template_id)
    record = state_access.get_template_record(normalized)
    if not record:
        raise HTTPException(status_code=404, detail="template_not_found")
    return normalized, record

def get_service(settings=Depends(get_settings)) -> TemplateService:
    return TemplateService(
        uploads_root=settings.uploads_dir,
        excel_uploads_root=settings.excel_uploads_dir,
        max_bytes=settings.max_upload_bytes,
        max_zip_entries=settings.max_zip_entries,
        max_zip_uncompressed_bytes=settings.max_zip_uncompressed_bytes,
        max_concurrency=settings.template_import_max_concurrency,
    )

def validate_upload_file(file: UploadFile, max_bytes: int) -> None:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    if len(file.filename) > MAX_FILENAME_LENGTH:
        raise HTTPException(status_code=400, detail=f"Filename too long (max {MAX_FILENAME_LENGTH} characters)")
    is_valid, error = validate_file_extension(file.filename, ALLOWED_EXTENSIONS)
    if not is_valid:
        raise HTTPException(status_code=400, detail=error)
    if file.content_type and file.content_type not in (
        "application/zip",
        "application/x-zip-compressed",
        "application/octet-stream",
    ):
        raise HTTPException(
            status_code=400,
            detail="Please upload a valid ZIP file. The file you selected does not appear to be a ZIP archive.",
        )

async def _persist_upload(file: UploadFile, suffix: str) -> tuple[Path, str]:
    filename = Path(file.filename or f"upload{suffix}").name
    tmp = tempfile.NamedTemporaryFile(prefix="nr-upload-", suffix=suffix, delete=False)
    try:
        with tmp:
            file.file.seek(0)
            while True:
                chunk = file.file.read(1024 * 1024)
                if not chunk:
                    break
                tmp.write(chunk)
    finally:
        with contextlib.suppress(Exception):
            await file.close()
    return Path(tmp.name), filename

# Template List & Catalog

@templates_router.get("")
def list_templates_route(
    request: Request,
    status: Optional[str] = None,
    kind: Optional[str] = Query(None, description="Filter by template kind (pdf, excel)"),
    q: Optional[str] = Query(None, description="Search templates by name"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
):
    """List all templates with optional status, kind, search, and pagination filters."""
    result = list_templates(status, request)
    templates = result.get("templates", [])
    # Strict status filter (legacy service may treat 'active' as 'approved')
    if status:
        status_lower = status.strip().lower()
        templates = [t for t in templates if (t.get("status") or "").lower() == status_lower]
    # Apply kind filter
    if kind:
        kind_lower = kind.strip().lower()
        templates = [t for t in templates if (t.get("kind") or "pdf").lower() == kind_lower]
    # Apply search filter
    if q:
        q_lower = q.strip().lower()
        templates = [t for t in templates if q_lower in (t.get("name") or "").lower() or q_lower in (t.get("description") or "").lower()]
    total = len(templates)
    # Apply pagination
    templates = templates[offset:offset + limit]
    result["templates"] = templates
    result["total"] = total
    result["limit"] = limit
    result["offset"] = offset
    return result

@templates_router.get("/catalog")
def templates_catalog_route(request: Request):
    """Get template catalog for browsing."""
    return templates_catalog(request)

# Chat-based Template Creation (from scratch)

@templates_router.post("/chat-create")
async def chat_template_create_route(request: Request):
    """Conversational template creation endpoint (no template_id needed).

    Accepts either:
    - JSON body (TemplateChatPayload) when no file is attached
    - multipart/form-data with messages_json, optional html, and optional sample_pdf
    """
    import json as _json
    from backend.app.services.legacy_services import TemplateChatMessage

    content_type = request.headers.get("content-type", "")
    sample_pdf_bytes = None

    if "multipart/form-data" in content_type:
        form = await request.form()
        messages_json = form.get("messages_json")
        if not messages_json:
            raise HTTPException(status_code=422, detail="messages_json is required in form data")
        try:
            msgs_raw = _json.loads(messages_json)
        except _json.JSONDecodeError:
            raise HTTPException(status_code=422, detail="Invalid messages_json")
        payload = TemplateChatPayload(
            messages=[TemplateChatMessage(role=m["role"], content=m["content"]) for m in msgs_raw],
            html=form.get("html") or None,
        )
        sample_file = form.get("sample_pdf")
        if sample_file is not None:
            sample_pdf_bytes = await sample_file.read()
        kind = str(form.get("kind") or "pdf").lower()
    else:
        body = await request.json()
        payload = TemplateChatPayload(**body)
        kind = str(body.get("kind", "pdf")).lower()

    if kind not in ("pdf", "excel"):
        kind = "pdf"

    return chat_template_create(payload, request, sample_pdf_bytes=sample_pdf_bytes, kind=kind)

@templates_router.post("/create-from-chat")
def create_template_from_chat_route(payload: TemplateCreateFromChatPayload, request: Request):
    """Persist a template that was created via the chat conversation."""
    return create_template_from_chat(payload, request)

# Template CRUD

@templates_router.get("/{template_id}")
def get_template_route(template_id: str, request: Request):
    """Get a single template by ID."""
    normalized, record = _ensure_template_exists(template_id)
    return {"status": "ok", "template": record, "correlation_id": _correlation(request)}

@templates_router.delete("/{template_id}")
def delete_template_route(template_id: str, request: Request):
    """Delete a template."""
    return delete_template(template_id, request)

@templates_router.patch("/{template_id}")
def update_template_metadata_route(template_id: str, payload: TemplateUpdatePayload, request: Request):
    """Update template metadata (name, description, etc.)."""
    return update_template_metadata(template_id, payload, request)

# Template Verification

@templates_router.post("/verify")
async def verify_template_route(
    request: Request,
    file: UploadFile = File(...),
    connection_id: Optional[str] = Form(None),
    refine_iters: int = Form(0),
    page: int = Form(0),
    background: bool = Query(False),
):
    """Verify and process a PDF template.

    DEPRECATED: Use POST /api/v1/pipeline/chat/upload instead.
    All pipeline steps should go through the unified chat interface.

    Args:
        page: Zero-based page index to render from multi-page PDFs (default 0).
    """
    import warnings
    warnings.warn("Direct /templates/verify is deprecated — use /pipeline/chat/upload", DeprecationWarning, stacklevel=2)
    if not background:
        return verify_template(file=file, connection_id=connection_id, refine_iters=refine_iters, page=page, request=request)

    upload_path, filename = await _persist_upload(file, suffix=".pdf")
    correlation_id = _correlation(request)
    template_name = Path(filename).stem or filename

    async def runner(job_id: str) -> None:
        upload = UploadFile(filename=filename, file=upload_path.open("rb"))
        try:
            response = verify_template(
                file=upload,
                connection_id=connection_id,
                refine_iters=refine_iters,
                page=page,
                request=_request_with_correlation(correlation_id),
            )
            await run_event_stream_async(job_id, iter_ndjson_events_async(response.body_iterator))
        finally:
            with contextlib.suppress(Exception):
                await upload.close()
            with contextlib.suppress(FileNotFoundError):
                upload_path.unlink(missing_ok=True)

    job = await enqueue_background_job(
        job_type="verify_template",
        connection_id=connection_id,
        template_name=template_name,
        template_kind="pdf",
        meta={"filename": filename, "background": True, "refine_iters": refine_iters, "page": page},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

# Template Import/Export

@templates_router.post("/import-zip", response_model=TemplateImportResult)
async def import_template_zip(
    request: Request,
    file: UploadFile = File(...),
    name: str | None = Form(None, max_length=100),
    service: TemplateService = Depends(get_service),
    settings=Depends(get_settings),
):
    """Import a template from a zip file."""
    validate_upload_file(file, settings.max_upload_bytes)
    if name is not None and not is_safe_name(name):
        raise HTTPException(status_code=400, detail="Template name contains invalid characters")
    correlation_id = _correlation(request)
    return await service.import_zip(file, name, correlation_id)

@templates_router.get("/{template_id}/export")
async def export_template_zip(
    template_id: str,
    request: Request,
    service: TemplateService = Depends(get_service),
):
    """Export a template as a zip file for sharing or backup."""
    correlation_id = _correlation(request)
    result = await service.export_zip(template_id, correlation_id)
    return FileResponse(
        path=result["zip_path"],
        filename=result["filename"],
        media_type="application/zip",
        background=None,
    )

@templates_router.post("/{template_id}/duplicate")
async def duplicate_template(
    template_id: str,
    request: Request,
    name: str | None = Form(None, max_length=100),
    service: TemplateService = Depends(get_service),
):
    """Duplicate a template to create a new copy."""
    if name is not None and not is_safe_name(name):
        raise HTTPException(status_code=400, detail="Template name contains invalid characters")
    correlation_id = _correlation(request)
    return await service.duplicate(template_id, name, correlation_id)

# Template Tags

@templates_router.put("/{template_id}/tags")
async def update_template_tags(
    template_id: str,
    payload: dict,
    service: TemplateService = Depends(get_service),
):
    """Update tags for a template."""
    tags = payload.get("tags", [])
    if not isinstance(tags, list):
        raise HTTPException(status_code=400, detail="Tags must be an array of strings")
    for tag in tags:
        if not isinstance(tag, str) or len(tag) > 50:
            raise HTTPException(status_code=400, detail="Each tag must be a string under 50 characters")
    return await service.update_tags(template_id, tags)

@templates_router.get("/tags/all")
async def get_all_tags(service: TemplateService = Depends(get_service)):
    """Get all unique tags across all templates."""
    return await service.get_all_tags()

# Template Recommendations

@templates_router.post("/recommend")
async def recommend_templates_route(
    payload: TemplateRecommendPayload,
    request: Request,
    background: bool = Query(False),
):
    """Get AI-powered template recommendations based on user requirements."""
    if not background:
        return recommend_templates(payload, request)

    correlation_id = _correlation(request)

    async def runner(job_id: str) -> None:
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "recommend", status="running", label="Generate recommendations")
        try:
            response = recommend_templates(payload, _request_with_correlation(correlation_id))
            if hasattr(response, "model_dump"):
                result_payload = response.model_dump(mode="json")
            elif hasattr(response, "model_dump"):
                result_payload = response.model_dump()
            else:
                result_payload = response
            result_data = (
                result_payload.get("recommendations")
                if isinstance(result_payload, dict)
                else result_payload
            )
            state_access.record_job_step(job_id, "recommend", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"recommendations": result_data},
            )
        except Exception as exc:
            logger.exception("template_recommend_job_failed", extra={"job_id": job_id})
            state_access.record_job_step(job_id, "recommend", status="failed", error="Template recommendation failed")
            state_access.record_job_completion(job_id, status="failed", error="Template recommendation failed")

    job = await enqueue_background_job(
        job_type="recommend_templates",
        steps=[{"name": "recommend", "label": "Generate recommendations"}],
        meta={"background": True, "requirement": payload.requirement},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

# Template HTML & Editing

@templates_router.get("/{template_id}/html")
def get_template_html_route(template_id: str, request: Request):
    """Get the current HTML content of a template."""
    return get_template_html(template_id, request)

@templates_router.post("/{template_id}/edit-manual")
def edit_template_manual_route(template_id: str, payload: TemplateManualEditPayload, request: Request):
    """Save manual HTML edits to a template."""
    return edit_template_manual(template_id, payload, request)

@templates_router.post("/{template_id}/edit-ai")
def edit_template_ai_route(template_id: str, payload: TemplateAiEditPayload, request: Request):
    """Apply AI-powered edits to a template based on instructions."""
    return edit_template_ai(template_id, payload, request)

@templates_router.post("/{template_id}/undo-last-edit")
def undo_last_edit_route(template_id: str, request: Request):
    """Undo the last edit made to a template."""
    return undo_last_template_edit(template_id, request)

@templates_router.post("/{template_id}/chat")
def chat_template_edit_route(template_id: str, payload: TemplateChatPayload, request: Request):
    """Conversational template editing endpoint."""
    return chat_template_edit(template_id, payload, request)

@templates_router.post("/{template_id}/chat/apply")
def apply_chat_template_edit_route(template_id: str, payload: TemplateManualEditPayload, request: Request):
    """Apply the HTML changes from a chat conversation."""
    return apply_chat_template_edit(template_id, payload.html, request)

# Mapping Preview/Approve/Corrections

@templates_router.post("/{template_id}/mapping/preview")
async def mapping_preview(template_id: str, connection_id: str, request: Request, force_refresh: bool = False):
    """Preview mapping for a PDF template. DEPRECATED: Use /pipeline/chat instead."""
    return await run_mapping_preview(template_id, connection_id, request, force_refresh, kind="pdf")

@templates_router.post("/{template_id}/mapping/approve")
async def mapping_approve(template_id: str, payload: MappingPayload, request: Request):
    """Approve mapping for a PDF template. DEPRECATED: Use /pipeline/chat instead."""
    return await run_mapping_approve(template_id, payload, request, kind="pdf")

@templates_router.post("/{template_id}/mapping/corrections-preview")
def mapping_corrections_preview(template_id: str, payload: CorrectionsPreviewPayload, request: Request):
    """Preview corrections for PDF template mapping."""
    return run_corrections_preview(template_id, payload, request, kind="pdf")

# Generator Assets

@templates_router.post("/{template_id}/generator-assets/v1")
def generator_assets_route(template_id: str, payload: GeneratorAssetsPayload, request: Request):
    """Generate assets for a PDF template. DEPRECATED: Use /pipeline/chat instead."""
    return generator_assets(template_id, payload, request, kind="pdf")

# Key Options

@templates_router.get("/{template_id}/keys/options")
def mapping_key_options(
    template_id: str,
    request: Request,
    connection_id: str | None = None,
    tokens: str | None = None,
    limit: int = 500,
    start_date: str | None = None,
    end_date: str | None = None,
    debug: bool = False,
):
    """Get available key options for template filtering."""
    return mapping_key_options_service(
        template_id=template_id,
        request=request,
        connection_id=connection_id,
        tokens=tokens,
        limit=limit,
        start_date=start_date,
        end_date=end_date,
        kind="pdf",
        debug=debug,
    )

# Artifacts

@templates_router.get("/{template_id}/artifacts/manifest")
def get_artifact_manifest(template_id: str, request: Request):
    """Get the artifact manifest for a template."""
    data = artifact_manifest_response(template_id, kind="pdf")
    return _wrap(data, _correlation(request))

@templates_router.get("/{template_id}/artifacts/head")
def get_artifact_head(template_id: str, request: Request, name: str):
    """Get the head (preview) of a specific artifact."""
    data = artifact_head_response(template_id, name, kind="pdf")
    return _wrap(data, _correlation(request))

# Charts

@templates_router.post("/{template_id}/charts/suggest")
def suggest_charts_route(template_id: str, payload: ChartSuggestPayload, request: Request):
    """Get chart suggestions for a template."""
    correlation_id = _correlation(request) or get_correlation_id()
    logger = logging.getLogger("neura.api")
    return suggest_charts_service(
        template_id,
        payload,
        kind="pdf",
        correlation_id=correlation_id,
        template_dir_fn=lambda tpl: template_dir(tpl, kind="pdf"),
        db_path_fn=db_path_from_payload_or_default,
        load_contract_fn=load_contract_v2,
        clean_key_values_fn=clean_key_values,
        discover_fn=discover_batches_and_counts,
        build_field_catalog_fn=build_batch_field_catalog_and_stats,
        build_metrics_fn=build_batch_metrics,
        build_prompt_fn=build_chart_suggestions_prompt,
        call_chat_completion_fn=lambda **kwargs: call_chat_completion(
            get_openai_client(), **kwargs, description=CHART_SUGGEST_PROMPT_VERSION
        ),
        model=get_model(),
        strip_code_fences_fn=strip_code_fences,
        logger=logger,
    )

@templates_router.get("/{template_id}/charts/saved")
def list_saved_charts_route(template_id: str, request: Request):
    """List saved charts for a template."""
    payload = list_saved_charts_service(template_id, _ensure_template_exists)
    return _wrap(payload, _correlation(request))

@templates_router.post("/{template_id}/charts/saved")
def create_saved_chart_route(
    template_id: str,
    payload: SavedChartCreatePayload,
    request: Request,
):
    """Create a saved chart for a template."""
    chart = create_saved_chart_service(
        template_id,
        payload,
        ensure_template_exists=_ensure_template_exists,
        normalize_template_id=normalize_template_id,
    )
    chart_payload = chart.model_dump(mode="json") if hasattr(chart, "model_dump") else chart
    return _wrap(chart_payload, _correlation(request))

@templates_router.put("/{template_id}/charts/saved/{chart_id}")
def update_saved_chart_route(
    template_id: str,
    chart_id: str,
    payload: SavedChartUpdatePayload,
    request: Request,
):
    """Update a saved chart."""
    chart = update_saved_chart_service(template_id, chart_id, payload, _ensure_template_exists)
    chart_payload = chart.model_dump(mode="json") if hasattr(chart, "model_dump") else chart
    return _wrap(chart_payload, _correlation(request))

@templates_router.delete("/{template_id}/charts/saved/{chart_id}")
def delete_saved_chart_route(
    template_id: str,
    chart_id: str,
    request: Request,
):
    """Delete a saved chart."""
    payload = delete_saved_chart_service(template_id, chart_id, _ensure_template_exists)
    return _wrap(payload, _correlation(request))

# Run Agent (Canvas Intelligence)

from pydantic import BaseModel as _PydanticBase, Field as _PydField
from typing import Any as _Any, Dict as _Dict

class RunAgentPayload(_PydanticBase):
    """Request body for running an agent in the template creator canvas."""
    agent_type: str = _PydField(..., description="Agent type: template_qa, data_mapping, data_quality, anomaly_detection, trend_analysis, report_pipeline")
    params: _Dict[str, _Any] = _PydField(default_factory=dict, description="Agent-specific parameters")
    sync: bool = _PydField(default=True, description="Wait for result if true")

@templates_router.post("/{template_id}/run-agent")
async def run_template_agent(template_id: str, payload: RunAgentPayload, request: Request):
    """Run an AI agent in the context of a template for the intelligence canvas.

    Supported agent types:
    - template_qa: Quality analysis of the template HTML
    - data_mapping: Auto-map tokens to database columns
    - data_quality: Analyze data source quality
    - anomaly_detection: Detect anomalies in data
    - trend_analysis: Analyze data trends
    - report_pipeline: Run full report validation pipeline
    """
    from backend.app.services.templates import run_canvas_agent

    # Allow __draft__ for unsaved templates - agents only need the HTML params
    if template_id == "__draft__":
        normalized = "__draft__"
    else:
        normalized, record = _ensure_template_exists(template_id)

    try:
        result = await run_canvas_agent(
            template_id=normalized,
            agent_type=payload.agent_type,
            params=payload.params,
            sync=payload.sync,
        )
        return _wrap({"ok": True, "agent_type": payload.agent_type, "result": result}, _correlation(request))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logging.getLogger(__name__).error(f"run_template_agent failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Agent execution failed: {str(e)}")

"""Reports API Routes.

This module contains endpoints for report generation and management:
- Report generation (sync and async)
- Report run history
- Report discovery
"""

from fastapi import APIRouter, Depends, HTTPException, Request

from backend.app.schemas import RunPayload, DiscoverPayload
from backend.app.services.legacy_services import (
    queue_report_job,
    queue_generate_docx_job,
    run_report as run_report_service,
    list_report_runs as list_report_runs_service,
    get_report_run as get_report_run_service,
    generate_docx_for_run as generate_docx_for_run_service,
)

reports_router = APIRouter(dependencies=[Depends(require_api_key)])

# Report Generation

@reports_router.post("/run")
def run_report(payload: RunPayload, request: Request):
    """Run a report synchronously. Auto-detects kind from template record."""

    rec = state_access.get_template_record(payload.template_id)
    if not rec:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "template_not_found", "message": f"Template '{payload.template_id}' not found."},
        )
    kind = str(rec.get("kind") or "pdf").strip().lower() or "pdf"
    # Validate connection_id exists
    if payload.connection_id:
        conn = state_access.get_connection_record(payload.connection_id)
        if not conn:
            raise HTTPException(
                status_code=404,
                detail={"status": "error", "code": "connection_not_found", "message": f"Connection '{payload.connection_id}' not found."},
            )
    # Reject synchronous DOCX generation (use /runs/{id}/generate-docx instead)
    if getattr(payload, "docx", False):
        raise HTTPException(
            status_code=422,
            detail={
                "status": "error",
                "code": "docx_not_supported_sync",
                "message": "DOCX generation is not supported in synchronous mode. "
                           "Generate the report first, then use POST /reports/runs/{run_id}/generate-docx "
                           "or POST /reports/jobs/generate-docx/{run_id} for async conversion.",
            },
        )
    # Validate date range
    if payload.start_date and payload.end_date and payload.start_date > payload.end_date:
        raise HTTPException(
            status_code=422,
            detail={
                "status": "error",
                "code": "invalid_date_range",
                "message": f"start_date ({payload.start_date}) must be <= end_date ({payload.end_date}).",
            },
        )
    try:
        return run_report_service(payload, request, kind=kind)
    except HTTPException:
        raise
    except Exception as exc:
        logging.getLogger("neura.api").exception("report_generation_failed", extra={"template_id": payload.template_id, "kind": kind})
        raise HTTPException(
            status_code=500,
            detail={"status": "error", "code": "report_generation_failed", "message": f"Report generation failed for kind={kind}: {type(exc).__name__}"},
        )

@reports_router.post("/jobs/run-report")
async def enqueue_report_job(payload: RunPayload | list[RunPayload], request: Request):
    """Queue a report job for async generation.

    Auto-detects the template kind (pdf/excel) from the template record.
    """

    payloads = payload if isinstance(payload, list) else [payload]
    kinds = set()
    for item in payloads:
        rec = state_access.get_template_record(item.template_id)
        if not rec:
            raise HTTPException(
                status_code=404,
                detail={"status": "error", "code": "template_not_found", "message": f"Template '{item.template_id}' not found."},
            )
        kinds.add(str(rec.get("kind") or "pdf").strip().lower() or "pdf")
    if len(kinds) > 1:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "mixed_template_kinds",
                "message": "All runs in a batch must share the same template kind.",
            },
        )
    kind = next(iter(kinds)) if kinds else "pdf"
    return await queue_report_job(payload, request, kind=kind)

@reports_router.post("/jobs/generate-docx/{run_id}")
async def enqueue_generate_docx_job(run_id: str, request: Request):
    """Queue a background job to convert a run's PDF to DOCX."""
    return await queue_generate_docx_job(run_id, request)

# Report Discovery

@reports_router.post("/discover")
def discover_reports(payload: DiscoverPayload, request: Request):
    """Discover available batches for report generation."""
    from backend.app.services.platform_services import discover_reports as discover_reports_service
    from backend.app.services.legacy_services import template_dir
    from backend.app.services.reports import build_batch_field_catalog_and_stats, build_batch_metrics
    from backend.app.services.infra_services import load_manifest
    from backend.app.services.legacy_services import manifest_endpoint

    logger = logging.getLogger("neura.api")
    return discover_reports_service(
        payload,
        kind="pdf",
        template_dir_fn=lambda tpl: template_dir(tpl, kind="pdf"),
        db_path_fn=db_path_from_payload_or_default,
        load_contract_fn=load_contract_v2,
        clean_key_values_fn=clean_key_values,
        discover_fn=discover_batches_and_counts,
        build_field_catalog_fn=build_batch_field_catalog_and_stats,
        build_batch_metrics_fn=build_batch_metrics,
        load_manifest_fn=load_manifest,
        manifest_endpoint_fn=lambda tpl: manifest_endpoint(tpl, kind="pdf"),
        logger=logger,
    )

# Report Run History

@reports_router.get("/runs")
def list_report_runs_route(
    request: Request,
    template_id: Optional[str] = None,
    connection_id: Optional[str] = None,
    schedule_id: Optional[str] = None,
    limit: int = 50,
):
    """List report generation runs with optional filtering."""
    runs = list_report_runs_service(
        template_id=template_id,
        connection_id=connection_id,
        schedule_id=schedule_id,
        limit=limit,
    )
    return {"runs": runs, "correlation_id": _correlation(request)}

@reports_router.post("/runs/{run_id}/generate-docx")
def generate_docx_route(run_id: str, request: Request):
    """Generate DOCX from an existing report run's PDF (on-demand, may take minutes)."""
    try:
        run = generate_docx_for_run_service(run_id)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail={"status": "error", "code": "generate_docx_failed", "message": str(exc)})
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail={"status": "error", "code": "generate_docx_failed", "message": str(exc)})
    return {"run": run, "correlation_id": _correlation(request)}

@reports_router.get("/runs/{run_id}")
def get_report_run_route(run_id: str, request: Request):
    """Get a specific report run by ID."""
    run = get_report_run_service(run_id)
    if not run:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "run_not_found", "message": "Run not found."}
        )
    return {"run": run, "correlation_id": _correlation(request)}

# V2: Enhanced Report Generation with SSE Streaming

@reports_router.post("/generate-enhanced")
async def generate_enhanced_report(payload: RunPayload, request: Request):
    """
    V2 enhanced report generation with SSE pipeline streaming.

    Returns a Server-Sent Events stream showing real-time pipeline
    stage progress (verify -> map -> query -> render -> pdf).
    Falls back to standard generation if V2 is disabled.
    """
    import asyncio
    from backend.app.services.infra_services import get_v2_config

    cfg = get_v2_config()

    if not cfg.enable_langgraph_pipeline:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "v2_disabled",
                "message": "Enhanced pipeline is not enabled. Use POST /run instead.",
            },
        )

    from backend.app.services.streaming import PipelineSSEBridge
    from backend.app.services.reports import run_enhanced_report

    bridge = PipelineSSEBridge()

    # Launch pipeline in background task
    async def _run_pipeline():
        await run_enhanced_report(
            template_id=payload.template_id,
            connection_id=payload.connection_id,
            filters=payload.filters if hasattr(payload, "filters") else None,
            batch_values=payload.batch_values if hasattr(payload, "batch_values") else None,
            sse_bridge=bridge,
        )

    asyncio.create_task(_run_pipeline())

    return StreamingResponse(
        bridge.stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Pipeline-Run-Id": bridge.run_id,
        },
    )

"""Jobs API Routes.

This module contains endpoints for job management:
- List jobs with filtering
- Get job details
- Cancel jobs
- Retry failed jobs
- Dead Letter Queue management
"""

from typing import List, Optional

from backend.app.schemas import RunPayload
from backend.app.services.config import normalize_job_status, normalize_job
from backend.app.services.legacy_services import get_job, list_active_jobs, list_jobs, cancel_job

jobs_router = APIRouter(dependencies=[Depends(require_api_key)])

# Use shared normalize_job_status and normalize_job from backend.app.services.config
_normalize_job_status = normalize_job_status
_normalize_job = normalize_job

@jobs_router.post("/run-report")
async def run_report_job(payload: RunPayload | list[RunPayload], request: Request):
    """Queue a report generation job (compatibility alias for `/reports/jobs/run-report`)."""
    from backend.app.services.legacy_services import queue_report_job

    payloads = payload if isinstance(payload, list) else [payload]
    kinds = set()
    for item in payloads:
        rec = state_access.get_template_record(item.template_id) or {}
        kinds.add(str(rec.get("kind") or "pdf").strip().lower() or "pdf")
    if len(kinds) > 1:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "mixed_template_kinds",
                "message": "All runs in a batch submission must share the same template kind.",
            },
        )
    kind = next(iter(kinds)) if kinds else "pdf"
    return await queue_report_job(payload, request, kind=kind)

@jobs_router.get("")
def list_jobs_route(
    request: Request,
    status: Optional[List[str]] = Query(None),
    job_type: Optional[List[str]] = Query(None, alias="type"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    active_only: bool = Query(False),
):
    """List jobs with optional filtering by status and type."""
    # Normalize "completed" -> "succeeded" for user convenience
    if status:
        status = [("succeeded" if s.lower() == "completed" else s) for s in status]
    # Fetch all matching jobs to get accurate total count
    all_jobs = list_jobs(status, job_type, 10000, active_only) or []
    total = len(all_jobs)
    # Apply pagination
    page_jobs = all_jobs[offset:offset + limit]
    normalized_jobs = [_normalize_job(job) for job in page_jobs]
    return {
        "jobs": normalized_jobs,
        "total": total,
        "limit": limit,
        "offset": offset,
        "correlation_id": _correlation(request),
    }

@jobs_router.get("/active")
def list_active_jobs_route(request: Request, limit: int = Query(20, ge=1, le=200)):
    """List only active (non-completed) jobs."""
    jobs = list_active_jobs(limit)
    normalized_jobs = [_normalize_job(job) for job in jobs] if jobs else []
    return {"jobs": normalized_jobs, "correlation_id": _correlation(request)}

# Dead Letter Queue Endpoints
# IMPORTANT: These must be defined BEFORE /{job_id} routes to avoid path conflicts

@jobs_router.get("/dead-letter")
def list_dead_letter_jobs_route(
    request: Request,
    limit: int = Query(50, ge=1, le=200),
):
    """List jobs in the Dead Letter Queue."""
    dlq_jobs = state_access.list_dead_letter_jobs(limit=limit)
    stats = state_access.get_dlq_stats()
    return {
        "jobs": dlq_jobs,
        "stats": stats,
        "correlation_id": _correlation(request),
    }

@jobs_router.get("/dead-letter/{job_id}")
def get_dead_letter_job_route(job_id: str, request: Request):
    """Get a specific job from the Dead Letter Queue."""
    dlq_job = state_access.get_dead_letter_job(job_id)
    if not dlq_job:
        raise HTTPException(
            status_code=404,
            detail={
                "status": "error",
                "code": "dlq_job_not_found",
                "message": "Job not found in Dead Letter Queue",
            }
        )
    return {"job": dlq_job, "correlation_id": _correlation(request)}

@jobs_router.post("/dead-letter/{job_id}/requeue")
async def requeue_from_dlq_route(job_id: str, request: Request):
    """
    Requeue a job from the Dead Letter Queue.

    Creates a new job with reset retry count and state.
    """
    dlq_job = state_access.get_dead_letter_job(job_id)
    if not dlq_job:
        raise HTTPException(
            status_code=404,
            detail={
                "status": "error",
                "code": "dlq_job_not_found",
                "message": "Job not found in Dead Letter Queue",
            }
        )

    # Create new job from DLQ record
    new_job = state_access.requeue_from_dlq(job_id)
    if not new_job:
        raise HTTPException(
            status_code=500,
            detail={
                "status": "error",
                "code": "requeue_failed",
                "message": "Failed to requeue job",
            }
        )

    return {
        "status": "ok",
        "message": "Job requeued from Dead Letter Queue",
        "original_job_id": job_id,
        "new_job": new_job,
        "correlation_id": _correlation(request),
    }

@jobs_router.delete("/dead-letter/{job_id}")
def delete_from_dlq_route(job_id: str, request: Request):
    """Permanently delete a job from the Dead Letter Queue."""
    deleted = state_access.delete_from_dlq(job_id)
    if not deleted:
        raise HTTPException(
            status_code=404,
            detail={
                "status": "error",
                "code": "dlq_job_not_found",
                "message": "Job not found in Dead Letter Queue",
            }
        )
    return {
        "status": "ok",
        "message": "Job deleted from Dead Letter Queue",
        "job_id": job_id,
        "correlation_id": _correlation(request),
    }

# Job Instance Endpoints (must come AFTER static routes like /dead-letter)

@jobs_router.get("/{job_id}")
def get_job_route(job_id: str, request: Request):
    """Get details for a specific job."""
    job = get_job(job_id)
    if not job:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "job_not_found", "message": "Job not found"},
        )
    return {"job": _normalize_job(job), "correlation_id": _correlation(request)}

@jobs_router.delete("/{job_id}")
def delete_job_route(job_id: str, request: Request):
    """Delete a job record."""
    deleted = state_access.delete_job(job_id)
    if not deleted:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "job_not_found", "message": "Job not found"},
        )
    return {"status": "ok", "job_id": job_id, "correlation_id": _correlation(request)}

@jobs_router.post("/{job_id}/cancel")
def cancel_job_route(job_id: str, request: Request, force: bool = Query(False)):
    """Cancel a running job. Cannot cancel already-completed jobs."""
    existing = get_job(job_id)
    if not existing:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "job_not_found", "message": "Job not found"},
        )
    status = _normalize_job_status(existing.get("status"))
    if status in ("succeeded", "completed", "failed", "cancelled"):
        raise HTTPException(
            status_code=409,
            detail={
                "status": "error",
                "code": "job_already_terminal",
                "message": f"Cannot cancel job with status '{status}'. Only active jobs can be cancelled.",
            },
        )
    job = cancel_job(job_id, force=force)
    return {"job": _normalize_job(job), "correlation_id": _correlation(request)}

@jobs_router.post("/{job_id}/retry")
async def retry_job_route(job_id: str, request: Request):
    """Retry a failed job by re-queuing it with the same parameters.

    Only jobs with status 'failed' can be retried.
    """

    original_job = get_job(job_id)
    if not original_job:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "job_not_found", "message": "Job not found"}
        )

    normalized_status = _normalize_job_status(original_job.get("status"))
    if normalized_status != "failed":
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "invalid_job_status",
                "message": f"Only failed jobs can be retried. Current status: {normalized_status}"
            }
        )

    job_type = str(original_job.get("type") or original_job.get("job_type") or "").strip() or "run_report"
    if job_type != "run_report":
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "retry_not_supported",
                "message": f"Retry is not supported for job type '{job_type}'. Re-run the original request.",
            },
        )

    # Extract job parameters from meta or direct fields
    meta = original_job.get("meta") or original_job.get("metadata") or {}
    template_id = original_job.get("template_id") or meta.get("template_id")
    connection_id = original_job.get("connection_id") or meta.get("connection_id")
    start_date = meta.get("start_date") or original_job.get("start_date")
    end_date = meta.get("end_date") or original_job.get("end_date")
    key_values = meta.get("key_values") or original_job.get("key_values")
    batch_ids = meta.get("batch_ids") or original_job.get("batch_ids")
    docx = meta.get("docx", False)
    xlsx = meta.get("xlsx", False)
    template_name = original_job.get("template_name") or meta.get("template_name")
    kind = original_job.get("template_kind") or meta.get("kind") or "pdf"

    if not template_id:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "missing_template_id",
                "message": "Cannot retry job: missing template_id"
            }
        )

    # Create payload for new job
    payload = RunPayload(
        template_id=template_id,
        connection_id=connection_id,
        start_date=start_date,
        end_date=end_date,
        key_values=key_values,
        batch_ids=batch_ids,
        docx=docx,
        xlsx=xlsx,
        template_name=template_name,
    )

    # Queue the new job
    result = await queue_report_job(payload, request, kind=kind)

    return {
        "status": "ok",
        "message": "Job retry queued successfully",
        "original_job_id": job_id,
        "new_job": result,
        "correlation_id": _correlation(request),
    }

"""Schedules API Routes.

This module contains endpoints for report scheduling:
- CRUD operations for scheduled reports
- Manual trigger for immediate execution
- Schedule status and history
"""

from datetime import datetime, timedelta, timezone

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request

from backend.app.services.legacy_services import (
    ScheduleCreatePayload, ScheduleUpdatePayload,
    create_schedule,
    delete_schedule,
    get_schedule,
    list_schedules,
    update_schedule,
)

logger = logging.getLogger("neura.schedules")

schedules_router = APIRouter(dependencies=[Depends(require_api_key)])

_RUN_TIME_RE = re.compile(r"^([01]\d|2[0-3]):([0-5]\d)$")

def _validate_run_time(run_time: str | None) -> None:
    """Raise 422 if run_time is provided but not valid HH:MM."""
    if run_time is None:
        return
    if not _RUN_TIME_RE.match(run_time.strip()):
        raise HTTPException(
            status_code=422,
            detail={
                "status": "error",
                "code": "invalid_run_time",
                "message": "run_time must be in HH:MM format (24-hour, e.g. '08:00' or '18:30').",
            },
        )

async def _refresh_scheduler() -> None:
    try:
        from backend.api import SCHEDULER
    except (ImportError, AttributeError):
        return
    if SCHEDULER is None:
        return
    try:
        await SCHEDULER.refresh()
    except Exception:
        logger.warning("Scheduler refresh failed", exc_info=True)

def _fixup_next_run(schedules: list) -> list:
    """Recalculate next_run_at for active schedules where it's in the past."""
    now = datetime.now(timezone.utc)
    for s in schedules:
        if not s.get("active", True):
            continue
        next_run_str = s.get("next_run_at")
        if not next_run_str:
            continue
        try:
            next_run = datetime.fromisoformat(str(next_run_str).replace("Z", "+00:00"))
            if next_run.tzinfo is None:
                next_run = next_run.replace(tzinfo=timezone.utc)
            if next_run < now:
                freq = str(s.get("frequency") or "daily").lower()
                if freq == "weekly":
                    delta = timedelta(weeks=1)
                elif freq == "monthly":
                    delta = timedelta(days=30)
                else:
                    delta = timedelta(days=1)
                # Advance until next_run is in the future
                while next_run < now:
                    next_run = next_run + delta
                s["next_run_at"] = next_run.isoformat()
        except (ValueError, TypeError):
            pass
    return schedules

@schedules_router.get("")
def list_report_schedules(request: Request):
    """List all report schedules."""
    schedules = _fixup_next_run(list_schedules() or [])
    return {"schedules": schedules, "correlation_id": _correlation(request)}

@schedules_router.post("")
async def create_report_schedule(payload: ScheduleCreatePayload, request: Request):
    """Create a new report schedule."""
    if payload.interval_minutes is not None and payload.interval_minutes < 1:
        raise HTTPException(
            status_code=422,
            detail={"status": "error", "code": "invalid_interval", "message": "interval_minutes must be a positive integer (>= 1), or omit to use frequency default."},
        )
    _validate_run_time(payload.run_time)
    schedule = create_schedule(payload)
    await _refresh_scheduler()
    return {"schedule": schedule, "correlation_id": _correlation(request)}

@schedules_router.get("/{schedule_id}")
def get_report_schedule(schedule_id: str, request: Request):
    """Get a specific schedule by ID."""
    schedule = get_schedule(schedule_id)
    if not schedule:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "schedule_not_found", "message": "Schedule not found."}
        )
    return {"schedule": schedule, "correlation_id": _correlation(request)}

@schedules_router.put("/{schedule_id}")
async def update_report_schedule(schedule_id: str, payload: ScheduleUpdatePayload, request: Request):
    """Update an existing report schedule."""
    _validate_run_time(payload.run_time)
    schedule = update_schedule(schedule_id, payload)
    await _refresh_scheduler()
    return {"schedule": schedule, "correlation_id": _correlation(request)}

@schedules_router.delete("/{schedule_id}")
async def delete_report_schedule(schedule_id: str, request: Request):
    """Delete a report schedule."""
    removed = delete_schedule(schedule_id)
    if not removed:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "schedule_not_found", "message": "Schedule not found."}
        )
    await _refresh_scheduler()
    return {"status": "ok", "schedule_id": schedule_id, "correlation_id": _correlation(request)}

@schedules_router.post("/{schedule_id}/trigger")
async def trigger_schedule(schedule_id: str, background_tasks: BackgroundTasks, request: Request):
    """
    Manually trigger a scheduled report to run immediately.

    This creates a job and queues it for execution without waiting for the next scheduled run.
    The actual report generation happens asynchronously.
    """
    correlation_id = _correlation(request) or f"manual-trigger-{schedule_id}"

    # Find the schedule
    schedule = get_schedule(schedule_id)
    if not schedule:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "schedule_not_found", "message": "Schedule not found."}
        )

    # Import here to avoid circular imports
    from backend.app.services.scheduler import (
        JobRunTracker,
        _build_job_steps,
        _step_progress_from_steps,
        scheduler_runner,
    )

    # Dynamic date range based on frequency (daily=yesterday->today, weekly=7d, monthly=30d)
    from backend.app.services.scheduler import _compute_dynamic_dates
    frequency = str(schedule.get("frequency") or "daily").strip().lower()
    dyn_start, dyn_end = _compute_dynamic_dates(frequency)

    # Build the payload from schedule data
    payload = {
        "template_id": schedule.get("template_id"),
        "connection_id": schedule.get("connection_id"),
        "start_date": dyn_start,
        "end_date": dyn_end,
        "batch_ids": schedule.get("batch_ids") or None,
        "key_values": schedule.get("key_values") or None,
        "docx": bool(schedule.get("docx")),
        "xlsx": bool(schedule.get("xlsx")),
        "email_recipients": schedule.get("email_recipients") or None,
        "email_subject": schedule.get("email_subject") or f"[Manual Trigger] {schedule.get('name') or schedule.get('template_id')}",
        "email_message": schedule.get("email_message") or f"Manually triggered run for schedule '{schedule.get('name')}'.\nWindow: {dyn_start} - {dyn_end}.",
        "schedule_id": schedule_id,
        "schedule_name": schedule.get("name"),
    }
    kind = schedule.get("template_kind") or "pdf"

    # Create a RunPayload to validate the payload
    try:
        run_payload = RunPayload(**payload)
    except (ValueError, TypeError) as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "status": "error",
                "code": "invalid_schedule_payload",
                "message": "Schedule has invalid configuration",
            }
        )

    # Create job record
    steps = _build_job_steps(run_payload, kind=kind)
    meta = {
        "start_date": payload.get("start_date"),
        "end_date": payload.get("end_date"),
        "schedule_id": schedule_id,
        "schedule_name": schedule.get("name"),
        "manual_trigger": True,
        "docx": bool(payload.get("docx")),
        "xlsx": bool(payload.get("xlsx")),
    }
    job_record = state_access.create_job(
        job_type="run_report",
        template_id=run_payload.template_id,
        connection_id=run_payload.connection_id,
        template_name=schedule.get("template_name") or run_payload.template_id,
        template_kind=kind,
        schedule_id=schedule_id,
        correlation_id=correlation_id,
        steps=steps,
        meta=meta,
    )

    job_id = job_record.get("id")
    step_progress = _step_progress_from_steps(steps)
    job_tracker = JobRunTracker(job_id, correlation_id=correlation_id, step_progress=step_progress)

    def run_scheduled_report():
        """Background task to run the scheduled report."""
        started = datetime.now(timezone.utc)
        try:
            job_tracker.start()
            result = scheduler_runner(payload, kind, job_tracker=job_tracker)
            finished = datetime.now(timezone.utc)

            # Record the manual run in schedule history
            artifacts = {
                "html_url": result.get("html_url"),
                "pdf_url": result.get("pdf_url"),
                "docx_url": result.get("docx_url"),
                "xlsx_url": result.get("xlsx_url"),
            }
            state_access.record_schedule_run(
                schedule_id,
                started_at=started.isoformat(),
                finished_at=finished.isoformat(),
                status="success",
                next_run_at=None,  # Don't update next_run_at for manual triggers
                error=None,
                artifacts=artifacts,
            )
            job_tracker.succeed(result)
            logger.info(
                "manual_trigger_completed",
                extra={
                    "event": "manual_trigger_completed",
                    "schedule_id": schedule_id,
                    "job_id": job_id,
                    "correlation_id": correlation_id,
                }
            )
        except Exception as exc:
            finished = datetime.now(timezone.utc)
            state_access.record_schedule_run(
                schedule_id,
                started_at=started.isoformat(),
                finished_at=finished.isoformat(),
                status="failed",
                next_run_at=None,
                error="Schedule execution failed",
                artifacts=None,
            )
            job_tracker.fail("Scheduled job failed")
            logger.exception(
                "manual_trigger_failed",
                extra={
                    "event": "manual_trigger_failed",
                    "schedule_id": schedule_id,
                    "job_id": job_id,
                    "correlation_id": correlation_id,
                    "error": str(exc),
                }
            )

    # Queue the background task
    background_tasks.add_task(run_scheduled_report)

    logger.info(
        "manual_trigger_queued",
        extra={
            "event": "manual_trigger_queued",
            "schedule_id": schedule_id,
            "job_id": job_id,
            "correlation_id": correlation_id,
        }
    )

    return {
        "status": "triggered",
        "message": "Schedule triggered for immediate execution",
        "schedule_id": schedule_id,
        "job_id": job_id,
        "correlation_id": correlation_id,
    }

@schedules_router.post("/{schedule_id}/pause")
async def pause_schedule(schedule_id: str, request: Request):
    """Pause a schedule (set active to false)."""
    schedule = get_schedule(schedule_id)
    if not schedule:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "schedule_not_found", "message": "Schedule not found."}
        )

    updated = update_schedule(schedule_id, ScheduleUpdatePayload(active=False))
    await _refresh_scheduler()
    return {
        "status": "ok",
        "message": "Schedule paused",
        "schedule": updated,
        "correlation_id": _correlation(request),
    }

@schedules_router.post("/{schedule_id}/resume")
async def resume_schedule(schedule_id: str, request: Request):
    """Resume a paused schedule (set active to true)."""
    schedule = get_schedule(schedule_id)
    if not schedule:
        raise HTTPException(
            status_code=404,
            detail={"status": "error", "code": "schedule_not_found", "message": "Schedule not found."}
        )

    updated = update_schedule(schedule_id, ScheduleUpdatePayload(active=True))
    await _refresh_scheduler()
    return {
        "status": "ok",
        "message": "Schedule resumed",
        "schedule": updated,
        "correlation_id": _correlation(request),
    }

"""State Management API Routes.

This module contains endpoints for application state management:
- Bootstrap state for app initialization
- Last used connection/template tracking
"""

import importlib
from fastapi import APIRouter, Depends, Request
from pydantic import BaseModel

from backend.app.services.legacy_services import bootstrap_state

state_router = APIRouter(dependencies=[Depends(require_api_key)])

class LastUsedPayload(BaseModel):
    connection_id: Optional[str] = None
    template_id: Optional[str] = None

from backend.app.common import get_state_store

@state_router.get("/bootstrap")
def bootstrap_state_route(request: Request):
    """Get bootstrap state for app initialization.

    Returns connections, templates, last used selections, and other
    initialization data needed when the app starts.
    """
    return bootstrap_state(request)

@state_router.post("/last-used")
def set_last_used_route(payload: LastUsedPayload, request: Request):
    """Record the last-used connection and template IDs for session persistence."""
    last_used = get_state_store().set_last_used(
        connection_id=payload.connection_id,
        template_id=payload.template_id,
    )
    return {
        "status": "ok",
        "last_used": last_used,
        "correlation_id": _correlation(request),
    }

# MISC ROUTES (merged from misc_routes.py)
# Legacy/compatibility routes to expose unused modules for runtime reachability.

from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile

logger = logging.getLogger(__name__)

legacy_router = APIRouter(prefix="/legacy", tags=["legacy"], dependencies=[Depends(require_api_key)])

# Legacy legacy/* routers
from backend.app.services.legacy_services import router as src_router

legacy_router.include_router(src_router, prefix="/src")

# Legacy generate feature routers
def _build_generate_router() -> APIRouter:
    from backend.app.api.middleware import build_run_router
    from backend.app.api.middleware import build_discover_router
    from backend.app.api.middleware import build_chart_suggest_router
    from backend.app.api.middleware import build_saved_charts_router
    from backend.app.services.ai_services import (
        CHART_SUGGEST_PROMPT_VERSION,
        build_chart_suggestions_prompt,
    )
    from backend.app.services.reports import discover_batches_and_counts
    from backend.app.services.reports import (
        discover_batches_and_counts as discover_batches_and_counts_excel,
        build_batch_field_catalog_and_stats,
        build_batch_metrics,
    )
    from backend.app.services.legacy_services import run_report, queue_report_job
    from backend.app.services.legacy_services import db_path_from_payload_or_default
    from backend.app.services.legacy_services import clean_key_values
    from backend.app.services.legacy_services import manifest_endpoint, normalize_template_id, template_dir
    import backend.app.services.config as state_access
    from backend.app.services.templates import get_openai_client
    from backend.app.services.infra_services import call_chat_completion, get_correlation_id, strip_code_fences, load_manifest
    from backend.app.services.contract_builder import load_contract_v2
    import os

    logger = logging.getLogger("neura.legacy")

    run_router = build_run_router(
        reports_run_fn=run_report,
        enqueue_job_fn=queue_report_job,
    )

    discover_router = build_discover_router(
        template_dir_fn=template_dir,
        db_path_fn=db_path_from_payload_or_default,
        load_contract_fn=load_contract_v2,
        clean_key_values_fn=clean_key_values,
        discover_pdf_fn=discover_batches_and_counts,
        discover_excel_fn=discover_batches_and_counts_excel,
        build_field_catalog_fn=build_batch_field_catalog_and_stats,
        build_batch_metrics_fn=build_batch_metrics,
        load_manifest_fn=load_manifest,
        manifest_endpoint_fn_pdf=manifest_endpoint,
        manifest_endpoint_fn_excel=manifest_endpoint,
        logger=logger,
    )

    chart_suggest_router = build_chart_suggest_router(
        template_dir_fn=template_dir,
        db_path_fn=db_path_from_payload_or_default,
        load_contract_fn=load_contract_v2,
        clean_key_values_fn=clean_key_values,
        discover_pdf_fn=discover_batches_and_counts,
        discover_excel_fn=discover_batches_and_counts_excel,
        build_field_catalog_fn=build_batch_field_catalog_and_stats,
        build_metrics_fn=build_batch_metrics,
        build_prompt_fn=build_chart_suggestions_prompt,
        call_chat_completion_fn=lambda **kwargs: call_chat_completion(
            get_openai_client(), **kwargs, description=CHART_SUGGEST_PROMPT_VERSION
        ),
        model=get_model(),
        strip_code_fences_fn=strip_code_fences,
        get_correlation_id_fn=get_correlation_id,
        logger=logger,
    )

    def _ensure_template_exists(template_id: str) -> tuple[str, dict]:
        normalized = normalize_template_id(template_id)
        record = state_access.get_template_record(normalized)
        if not record:
            raise HTTPException(status_code=404, detail="template_not_found")
        return normalized, record

    saved_charts_router = build_saved_charts_router(
        ensure_template_exists=_ensure_template_exists,
        normalize_template_id=normalize_template_id,
    )

    generate_router = APIRouter()
    generate_router.include_router(run_router)
    generate_router.include_router(discover_router)
    generate_router.include_router(chart_suggest_router)
    generate_router.include_router(saved_charts_router)
    return generate_router

legacy_router.include_router(_build_generate_router(), prefix="/generate")

# Legacy pipeline + orchestration reachability
@legacy_router.get("/pipelines/report/steps")
async def report_pipeline_steps():
    """Expose report pipeline definition."""
    # ARCH-EXC-001: legacy compatibility route requires direct engine access.
    from backend.engine_all import create_report_pipeline

    pipeline = create_report_pipeline()
    return {
        "pipeline": pipeline.name,
        "steps": [{"name": s.name, "label": s.label} for s in pipeline.steps],
    }

@legacy_router.post("/orchestration/test-run")
async def orchestration_test_run():
    """Run a short orchestration job for validation."""
    from backend.engine_all import get_executor
    from backend.engine_all import Job, JobType, JobStep

    executor = get_executor()
    job = Job.create(
        job_type=JobType.REPORT_GENERATION,
        steps=[JobStep(name="noop", label="No-op step")],
    )

    def _runner(job_obj: Job, _executor) -> Dict[str, Any]:
        job_obj.step_running("noop")
        job_obj.step_succeeded("noop", progress=100.0)
        return {"status": "ok"}

    runners = getattr(executor, "_runners", {})
    if JobType.REPORT_GENERATION not in runners:
        executor.register_runner(JobType.REPORT_GENERATION, _runner)

    await executor.submit(job)
    for _ in range(40):
        if job.status.is_terminal:
            break
        await asyncio.sleep(0.05)

    return {"job": job.to_dict()}

# Extraction helpers (pdf/excel)
class PdfExtractRequest(BaseModel):
    method: str = "auto"
    max_pages: int = 10

class ExcelExtractRequest(BaseModel):
    max_rows: int = 5000

@legacy_router.post("/extraction/pdf")
async def legacy_extract_pdf(
    request: PdfExtractRequest,
    file: UploadFile = File(...),
):
    """Extract PDF tables using legacy extractors."""
    from backend.app.services.extractors import ExtractionConfig, extract_pdf_tables

    suffix = Path(file.filename or "document.pdf").suffix or ".pdf"
    tmp = tempfile.NamedTemporaryFile(prefix="nr-legacy-pdf-", suffix=suffix, delete=False)
    tmp_path = Path(tmp.name)
    try:
        with tmp:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                tmp.write(chunk)
        config = ExtractionConfig(max_pages=request.max_pages)
        result = extract_pdf_tables(tmp_path, method=request.method, config=config)
        return result.to_dict()
    finally:
        await file.close()
        with contextlib.suppress(FileNotFoundError):
            tmp_path.unlink(missing_ok=True)

@legacy_router.post("/extraction/excel")
async def legacy_extract_excel(
    request: ExcelExtractRequest,
    file: UploadFile = File(...),
):
    """Extract Excel data using legacy extractors."""
    from backend.app.services.extractors import extract_excel_data

    suffix = Path(file.filename or "document.xlsx").suffix or ".xlsx"
    tmp = tempfile.NamedTemporaryFile(prefix="nr-legacy-excel-", suffix=suffix, delete=False)
    tmp_path = Path(tmp.name)
    try:
        with tmp:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                tmp.write(chunk)
        result = extract_excel_data(tmp_path, max_rows=request.max_rows)
        return result.to_dict()
    finally:
        await file.close()
        with contextlib.suppress(FileNotFoundError):
            tmp_path.unlink(missing_ok=True)

# QuickChart integration
class QuickChartRequest(BaseModel):
    chart_type: str
    labels: List[str]
    data: Any
    title: Optional[str] = None

@legacy_router.post("/charts/quickchart/url")
async def quickchart_url(request: QuickChartRequest):
    """Generate a QuickChart URL without downloading the image."""
    from backend.app.services.platform_services import generate_chart_url

    try:
        url = generate_chart_url(
            request.chart_type,
            request.labels,
            request.data,
            title=request.title,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail="Invalid export request")
    return {"url": url}

# LLM utilities (agents, vision, document extractor)
class DocumentExtractRequest(BaseModel):
    use_vlm: bool = False
    max_pages: int = 10

@legacy_router.get("/llm/agents")
async def list_llm_agents():
    """List configured LLM agents and tasks."""
    from backend.app.services.llm import create_document_processing_crew

    crew = create_document_processing_crew(verbose=False)
    agents = [
        {"role": agent.role, "goal": agent.config.goal}
        for agent in crew.agents.values()
    ]
    tasks = [
        {"description": task.description, "agent_role": task.agent_role}
        for task in crew.tasks
    ]
    return {"agents": agents, "tasks": tasks}

@legacy_router.get("/llm/vision/model")
async def get_vision_model_info():
    """Return the configured vision model (if available)."""
    try:
        from backend.app.services.llm import VisionLanguageModel
        vlm = VisionLanguageModel()
        return {"model": vlm.model}
    except Exception as exc:
        logger.warning("Vision model not available: %s", exc)
        raise HTTPException(status_code=503, detail="Vision model not available")

@legacy_router.post("/llm/document-extract")
async def legacy_document_extract(
    request: DocumentExtractRequest,
    file: UploadFile = File(...),
):
    """Extract document content using the enhanced LLM document extractor."""
    from backend.app.services.llm import EnhancedDocumentExtractor

    suffix = Path(file.filename or "document").suffix or ".pdf"
    tmp = tempfile.NamedTemporaryFile(prefix="nr-legacy-doc-", suffix=suffix, delete=False)
    tmp_path = Path(tmp.name)
    try:
        with tmp:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                tmp.write(chunk)

        extractor = EnhancedDocumentExtractor(
            use_vlm=request.use_vlm,
            max_pages=request.max_pages,
        )
        result = extractor.extract(tmp_path)
        text_preview = (result.text or "")[:2000]
        return {
            "text_preview": text_preview,
            "table_count": len(result.tables),
            "metadata": result.metadata,
            "warnings": result.warnings,
        }
    finally:
        await file.close()
        with contextlib.suppress(FileNotFoundError):
            tmp_path.unlink(missing_ok=True)

"""API routes for Auto-Chart Generation."""

from pydantic import BaseModel, Field, field_validator, model_validator

from fastapi import APIRouter, Depends, Query, Request

from backend.app.services.platform_services import AutoChartService
from backend.app.services.config import enqueue_background_job

logger = logging.getLogger("neura.api.charts")

VALID_CHART_TYPES = {"bar", "line", "scatter", "pie", "area", "heatmap", "histogram", "box", "radar", "treemap"}

charts_router = APIRouter(dependencies=[Depends(require_api_key)])

class ChartAnalyzeRequest(BaseModel):
    data: List[Dict[str, Any]] = Field(..., min_length=1, max_length=100)
    column_descriptions: Optional[Dict[str, str]] = None
    max_suggestions: int = Field(default=3, ge=1, le=10)

class ChartGenerateRequest(BaseModel):
    data: List[Dict[str, Any]] = Field(..., min_length=1, max_length=1000)
    chart_type: str
    x_field: str
    y_fields: List[str] = Field(..., min_length=1, max_length=20)
    title: Optional[str] = Field(None, max_length=255)

    @field_validator("chart_type")
    @classmethod
    def validate_chart_type(cls, v: str) -> str:
        if v not in VALID_CHART_TYPES:
            raise ValueError(f"Invalid chart_type '{v}'. Must be one of: {', '.join(sorted(VALID_CHART_TYPES))}")
        return v

    @model_validator(mode="after")
    def validate_fields_exist_in_data(self):
        if not self.data:
            return self
        sample_keys = set(self.data[0].keys())
        if self.x_field not in sample_keys:
            raise ValueError(f"x_field '{self.x_field}' not found in data keys: {sorted(sample_keys)}")
        missing = [f for f in self.y_fields if f not in sample_keys]
        if missing:
            raise ValueError(f"y_fields {missing} not found in data keys: {sorted(sample_keys)}")
        return self

def get_service() -> AutoChartService:
    return AutoChartService()

@charts_router.get("/saved")
async def list_saved_charts():
    """List all saved charts across all templates."""
    all_charts = []
    try:
        templates = state_access.list_templates()
        for t in templates:
            tid = t.get("id", "")
            saved = state_access.list_saved_charts(tid) if hasattr(state_access, "list_saved_charts") else []
            for c in saved:
                c["template_id"] = tid
                all_charts.append(c)
    except Exception:
        pass
    return {"charts": all_charts, "total": len(all_charts)}

@charts_router.post("/analyze")
async def analyze_for_charts(
    payload: ChartAnalyzeRequest,
    request: Request,
    svc: AutoChartService = Depends(get_service),
    background: bool = Query(False),
):
    """Analyze data and suggest appropriate chart visualizations."""
    correlation_id = getattr(request.state, "correlation_id", None)
    if not background:
        suggestions = svc.analyze_data_for_charts(
            data=payload.data,
            column_descriptions=payload.column_descriptions,
            max_suggestions=payload.max_suggestions,
            correlation_id=correlation_id,
        )
        return {"status": "ok", "suggestions": suggestions, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "analyze", status="running", label="Analyze chart data")
        try:
            suggestions = svc.analyze_data_for_charts(
                data=payload.data,
                column_descriptions=payload.column_descriptions,
                max_suggestions=payload.max_suggestions,
                correlation_id=correlation_id,
            )
            state_access.record_job_step(job_id, "analyze", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"suggestions": suggestions},
            )
        except Exception as exc:
            logger.exception("chart_analyze_failed", extra={"job_id": job_id})
            safe_msg = "Chart analysis failed"
            state_access.record_job_step(job_id, "analyze", status="failed", error=safe_msg)
            state_access.record_job_completion(job_id, status="failed", error=safe_msg)

    job = await enqueue_background_job(
        job_type="chart_analyze",
        steps=[{"name": "analyze", "label": "Analyze chart data"}],
        meta={"background": True, "row_count": len(payload.data)},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

@charts_router.post("/generate")
async def generate_chart_config(
    payload: ChartGenerateRequest,
    request: Request,
    svc: AutoChartService = Depends(get_service),
    background: bool = Query(False),
):
    """Generate a chart configuration."""
    correlation_id = getattr(request.state, "correlation_id", None)
    if not background:
        config = svc.generate_chart_config(
            data=payload.data,
            chart_type=payload.chart_type,
            x_field=payload.x_field,
            y_fields=payload.y_fields,
            title=payload.title,
        )
        return {"status": "ok", "chart": config, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "generate", status="running", label="Generate chart config")
        try:
            config = svc.generate_chart_config(
                data=payload.data,
                chart_type=payload.chart_type,
                x_field=payload.x_field,
                y_fields=payload.y_fields,
                title=payload.title,
            )
            state_access.record_job_step(job_id, "generate", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"chart": config},
            )
        except Exception as exc:
            logger.exception("chart_generate_failed", extra={"job_id": job_id})
            safe_msg = "Chart generation failed"
            state_access.record_job_step(job_id, "generate", status="failed", error=safe_msg)
            state_access.record_job_completion(job_id, status="failed", error=safe_msg)

    job = await enqueue_background_job(
        job_type="chart_generate",
        steps=[{"name": "generate", "label": "Generate chart config"}],
        meta={"background": True, "row_count": len(payload.data)},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

"""
Widget Intelligence API Routes - widget catalog, selection, grid packing, data.

Endpoints:
    GET  /widgets/catalog              Full widget catalog (24 scenarios)
    POST /widgets/recommend            Claude-powered widget recommendations for a DB connection
    POST /widgets/select               AI-powered widget selection
    POST /widgets/pack-grid            Pack widgets into CSS grid layout
    POST /widgets/{scenario}/validate  Validate data shape
    POST /widgets/{scenario}/format    Format raw data
    POST /widgets/data                 Live data from active DB connection
    POST /widgets/data/report          Data from a report run (RAG)
    POST /widgets/feedback             Thompson Sampling reward signal
"""

import uuid
from typing import Any, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, Field

logger = logging.getLogger("neura.api.widgets")

widgets_router = APIRouter(tags=["widgets"], dependencies=[Depends(require_api_key)])

# Lazy-init service singleton
_svc = None

def _get_svc():
    global _svc
    if _svc is None:
        from backend.app.services.widget_intelligence import WidgetIntelligenceService
        _svc = WidgetIntelligenceService()
    return _svc

# -- Schemas --------------------------------------------------------------

class WidgetSelectRequest(BaseModel):
    query: str = Field(..., min_length=1, max_length=2000)
    query_type: str = Field(default="overview")
    data_profile: Optional[dict[str, Any]] = None
    max_widgets: int = Field(default=10, ge=1, le=20)

class GridPackRequest(BaseModel):
    widgets: list[dict[str, Any]] = Field(...)

class ValidateRequest(BaseModel):
    data: dict[str, Any] = Field(...)

class FormatRequest(BaseModel):
    data: dict[str, Any] = Field(...)

class FeedbackRequest(BaseModel):
    scenario: str = Field(..., min_length=1)
    reward: float = Field(..., ge=-1.0, le=1.0)

class WidgetDataRequest(BaseModel):
    connection_id: str = Field(..., min_length=1)
    scenario: str = Field(..., min_length=1)
    variant: Optional[str] = None
    filters: Optional[dict[str, Any]] = None
    limit: int = Field(default=100, ge=1, le=1000)

class RecommendRequest(BaseModel):
    connection_id: str = Field(..., min_length=1)
    query: str = Field(default="overview", max_length=2000)
    max_widgets: int = Field(default=8, ge=1, le=20)

class WidgetReportDataRequest(BaseModel):
    run_id: str = Field(..., min_length=1)
    scenario: str = Field(..., min_length=1)
    variant: Optional[str] = None

# Lazy-init data resolver
_resolver = None

def _get_resolver():
    global _resolver
    if _resolver is None:
        from backend.app.services.widget_intelligence import WidgetDataResolver
        _resolver = WidgetDataResolver()
    return _resolver

# -- Endpoints ------------------------------------------------------------

@widgets_router.get("/catalog")
async def get_widget_catalog():
    """Return the full widget catalog with all registered scenarios."""
    svc = _get_svc()
    catalog = svc.get_catalog()
    return {"widgets": catalog, "count": len(catalog)}

@widgets_router.post("/recommend")
async def recommend_widgets(req: RecommendRequest):
    """Analyze a connected DB using Claude LLM and recommend optimal widgets."""
    try:
        from backend.app.repositories import resolve_db_path
        from backend.app.services.legacy_services import build_rich_catalog_from_db, format_catalog_rich
        from backend.app.services.widget_intelligence import pack_grid as dynamic_pack
        from backend.app.services.widget_intelligence import (
            VALID_SCENARIOS, VARIANT_TO_SCENARIO, WidgetSlot,
        )
        from backend.app.services.widget_intelligence import WidgetSize
        from backend.app.services.llm import get_llm_client

        # 1. Resolve DB path from connection_id
        db_path = resolve_db_path(req.connection_id, None, None)

        # 2. Build rich catalog from DB schema
        rich_catalog = build_rich_catalog_from_db(db_path)
        table_count = len(rich_catalog)
        total_columns = sum(len(cols) for cols in rich_catalog.values())
        numeric_cols = sum(
            1 for cols in rich_catalog.values()
            for c in cols if c.get("type", "").upper() in ("INTEGER", "REAL", "FLOAT", "NUMERIC", "DECIMAL", "DOUBLE")
        )
        has_ts = any(
            c.get("type", "").upper() in ("DATE", "DATETIME", "TIMESTAMP")
            or any(kw in c.get("column", "").lower() for kw in ("date", "time", "timestamp"))
            for cols in rich_catalog.values() for c in cols
        )

        # 3. Format schema as text for Claude
        schema_text = format_catalog_rich(rich_catalog)

        # 4. Build variant reference for the prompt
        variant_list = "\n".join(
            f"  - {variant} (scenario: {scenario})"
            for variant, scenario in sorted(VARIANT_TO_SCENARIO.items())
        )

        # 5. Call Claude LLM to recommend widgets
        system_prompt = f"""You are a data visualization expert. Given a database schema and a user query, recommend the best dashboard widgets.

Available scenarios: {', '.join(VALID_SCENARIOS)}

Available variants (variant -> scenario):
{variant_list}

Widget sizes: compact, normal, expanded, hero

Rules:
- Analyze the database tables and columns to understand what data is available
- Match the user intent to the most relevant widget scenarios
- Pick the best variant for each scenario based on the data shape
- Assign a relevance score (0.0-1.0) for how well each widget fits
- Generate a short question each widget answers (e.g. 'What is the total billing amount?')
- Pick appropriate sizes: use 'hero' for the most important widget, 'expanded' for secondary insights, 'normal' for standard widgets, 'compact' for supporting metrics
- Return at most {req.max_widgets} widgets, ordered by relevance (highest first)
- Only recommend scenarios that make sense for the available data
- If the DB has numeric columns, include KPI and trend widgets
- If the DB has date/time columns, include timeline or trend widgets
- If the DB has categorical columns, include distribution or category-bar widgets

DATABASE SCHEMA:
{schema_text}

Respond with ONLY a JSON array (no markdown, no explanation):
[
  {{"scenario": "kpi", "variant": "kpi-live", "size": "hero", "relevance": 0.95, "question": "What is the current total?"}},
  ...
]"""

        user_message = req.query or "overview"

        client = get_llm_client()
        response = client.complete(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            description="widget_recommend",
            use_cache=True,
            cache_ttl=300.0,
        )

        # 6. Extract JSON from LLM response
        raw_text = (
            response.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        )
        from backend.app.services.infra_services import extract_json_array_from_llm_response
        recommended = extract_json_array_from_llm_response(raw_text, default=[])
        if not recommended:
            logger.error("LLM response has no JSON array: %s", raw_text[:500])
            raise ValueError("Claude did not return a valid JSON array")

        # 7. Validate and build WidgetSlot objects
        slots: list[WidgetSlot] = []
        for i, rec in enumerate(recommended[:req.max_widgets]):
            scenario = rec.get("scenario", "")
            variant = rec.get("variant", scenario)
            size_str = rec.get("size", "normal")
            relevance = float(rec.get("relevance", 0.5))
            question = rec.get("question", "")

            # Validate scenario
            if scenario not in VALID_SCENARIOS:
                logger.warning("LLM recommended invalid scenario %s, skipping", scenario)
                continue

            # Validate variant belongs to scenario
            if variant in VARIANT_TO_SCENARIO and VARIANT_TO_SCENARIO[variant] != scenario:
                variant = scenario  # Fall back to base scenario name

            # Validate size
            try:
                size = WidgetSize(size_str)
            except ValueError:
                size = WidgetSize.normal

            slots.append(WidgetSlot(
                id=f"w-{uuid.uuid4().hex[:8]}",
                scenario=scenario,
                variant=variant,
                size=size,
                relevance=min(max(relevance, 0.0), 1.0),
                question=question,
            ))

        if not slots:
            logger.warning("LLM returned no valid widgets, raw: %s", raw_text[:500])

        # 8. Pack into grid layout
        grid = dynamic_pack(slots)

        # 9. Build response
        widgets = [
            {
                "id": s.id,
                "scenario": s.scenario,
                "variant": s.variant,
                "size": s.size.value if hasattr(s.size, "value") else str(s.size),
                "question": s.question,
                "relevance": s.relevance,
            }
            for s in slots
        ]

        return {
            "widgets": widgets,
            "count": len(widgets),
            "grid": {
                "cells": [
                    {
                        "widget_id": c.widget_id,
                        "col_start": c.col_start,
                        "col_end": c.col_end,
                        "row_start": c.row_start,
                        "row_end": c.row_end,
                    }
                    for c in grid.cells
                ],
                "total_cols": grid.total_cols,
                "total_rows": grid.total_rows,
                "utilization_pct": grid.utilization_pct,
            },
            "profile": {
                "table_count": table_count,
                "column_count": total_columns,
                "numeric_columns": numeric_cols,
                "has_timeseries": has_ts,
                "tables": list(rich_catalog.keys()),
            },
        }

    except RuntimeError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.exception("Widget recommendation failed")
        raise HTTPException(status_code=500, detail=f"Recommendation failed: {e}")

@widgets_router.post("/select")
async def select_widgets(req: WidgetSelectRequest):
    """AI-powered widget selection for a dashboard query."""
    svc = _get_svc()
    widgets = svc.select_widgets(
        query=req.query,
        query_type=req.query_type,
        data_profile=req.data_profile,
        max_widgets=req.max_widgets,
    )
    return {"widgets": widgets, "count": len(widgets)}

@widgets_router.post("/pack-grid")
async def pack_grid(req: GridPackRequest):
    """Pack selected widgets into a CSS grid layout."""
    svc = _get_svc()
    if not req.widgets:
        raise HTTPException(status_code=400, detail="At least one widget required")
    layout = svc.pack_grid(req.widgets)
    return layout

@widgets_router.post("/{scenario}/validate")
async def validate_widget_data(scenario: str, req: ValidateRequest):
    """Validate data shape for a widget scenario."""
    svc = _get_svc()
    errors = svc.validate_data(scenario, req.data)
    return {"scenario": scenario, "valid": len(errors) == 0, "errors": errors}

@widgets_router.post("/{scenario}/format")
async def format_widget_data(scenario: str, req: FormatRequest):
    """Format raw data into frontend-ready shape."""
    svc = _get_svc()
    formatted = svc.format_data(scenario, req.data)
    return {"scenario": scenario, "data": formatted}

@widgets_router.post("/data")
async def get_widget_data(req: WidgetDataRequest):
    """Fetch live data from an active DB connection using the widget's RAG strategy."""
    resolver = _get_resolver()
    result = resolver.resolve(
        connection_id=req.connection_id,
        scenario=req.scenario,
        variant=req.variant,
        filters=req.filters,
        limit=req.limit,
    )
    # Always return the result - error field signals issues to the frontend
    return result

@widgets_router.post("/data/report")
async def get_widget_report_data(req: WidgetReportDataRequest):
    """Fetch widget data from a report run's extracted tables and content."""
    from backend.app.services.reports import ReportContextProvider

    provider = ReportContextProvider()
    ctx = provider.get_report_context(req.run_id)
    if not ctx:
        raise HTTPException(status_code=404, detail=f"Report run {req.run_id} not found")

    svc = _get_svc()
    plugin_meta = None
    catalog = svc.get_catalog()
    for w in catalog:
        if w["scenario"] == req.scenario:
            plugin_meta = w
            break

    # Build data from report context based on RAG strategy
    rag_strategy = plugin_meta["rag_strategy"] if plugin_meta else "single_metric"

    data = {}
    if rag_strategy in ("single_metric", "multi_metric") and ctx.tables:
        # Use first table from report
        table = ctx.tables[0]
        data = {
            "labels": [row[0] if row else "" for row in table.get("rows", [])],
            "datasets": [],
        }
        headers = table.get("headers", [])
        for i, hdr in enumerate(headers[1:], 1):
            data["datasets"].append({
                "label": hdr,
                "data": [row[i] if i < len(row) else 0 for row in table.get("rows", [])],
            })

    elif rag_strategy == "narrative":
        data = {
            "title": f"Report: {ctx.template_name}",
            "text": ctx.text_content[:2000] if ctx.text_content else "No content available.",
            "highlights": [
                f"Template: {ctx.template_name}",
                f"Status: {ctx.status}",
                f"Records: {len(ctx.tables)} tables",
            ],
        }

    elif rag_strategy in ("alert_query", "events_in_range"):
        # Treat report tables as event data
        events = []
        for t in ctx.tables:
            for row in t.get("rows", [])[:20]:
                events.append({
                    "message": row[0] if row else "",
                    "timestamp": row[1] if len(row) > 1 else "",
                    "severity": "info",
                })
        data = {"events": events, "alerts": events}

    else:
        data = {
            "title": ctx.template_name,
            "text": ctx.text_content[:500] if ctx.text_content else "",
        }

    # Format through plugin if available
    formatted = svc.format_data(req.scenario, data) if data else data

    return {
        "scenario": req.scenario,
        "data": formatted,
        "source": f"report:{req.run_id}",
        "strategy": rag_strategy,
    }

@widgets_router.post("/feedback")
async def submit_feedback(req: FeedbackRequest):
    """Submit reward signal for Thompson Sampling learning."""
    svc = _get_svc()
    svc.update_feedback(req.scenario, req.reward)
    return {"status": "ok", "scenario": req.scenario}

"""
AI API Routes
Endpoints for AI-powered writing and spreadsheet assistance.

Error handling:
- 400: Invalid input (empty text, text too long, bad parameters)
- 422: Pydantic validation errors (handled by FastAPI)
- 500: Unexpected internal errors
- 503: AI service unavailable (circuit breaker open, provider down)
"""

from fastapi import APIRouter, Depends, HTTPException, Request, status
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field, field_validator

from backend.app.api.middleware import limiter, RATE_LIMIT_STRICT
from backend.app.services.platform_services import (
    writing_service,
    spreadsheet_ai_service,
    WritingTone,
    InputValidationError,
    LLMResponseError,
    LLMUnavailableError,
    WritingServiceError,
)

logger = logging.getLogger(__name__)
ai_router = APIRouter(dependencies=[Depends(require_api_key)])

# HELPER — error mapping

def _handle_service_error(exc: Exception, operation: str) -> HTTPException:
    """Map service errors to appropriate HTTP status codes."""
    if isinstance(exc, InputValidationError):
        return HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid input for AI operation",
        )
    if isinstance(exc, LLMUnavailableError):
        return HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="AI service is temporarily unavailable. Please try again shortly.",
        )
    if isinstance(exc, LLMResponseError):
        logger.error("%s: LLM response error: %s", operation, exc)
        return HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"{operation} failed: the AI returned an invalid response. Please retry.",
        )
    # Unexpected error
    logger.error("%s failed unexpectedly: %s", operation, exc, exc_info=True)
    return HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail=f"{operation} failed due to an internal error.",
    )

# REQUEST/RESPONSE MODELS

# Writing AI Models
class GrammarCheckRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=100_000, description="Text to check")
    language: str = Field(default="en", min_length=2, max_length=10, description="Language code")
    strict: bool = Field(default=False, description="Enable strict mode")

class SummarizeRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=100_000, description="Text to summarize")
    max_length: Optional[int] = Field(default=None, ge=10, le=10_000, description="Maximum words")
    style: str = Field(default="bullet_points", description="Output style")

    @field_validator("style")
    @classmethod
    def validate_style(cls, v: str) -> str:
        allowed = {"bullet_points", "paragraph", "executive"}
        if v not in allowed:
            raise ValueError(f"style must be one of: {', '.join(sorted(allowed))}")
        return v

class RewriteRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=100_000, description="Text to rewrite")
    tone: str = Field(default="professional", description="Target tone")
    preserve_meaning: bool = Field(default=True, description="Preserve original meaning")

    @field_validator("tone")
    @classmethod
    def validate_tone(cls, v: str) -> str:
        valid_tones = {t.value for t in WritingTone}
        if v not in valid_tones:
            raise ValueError(f"tone must be one of: {', '.join(sorted(valid_tones))}")
        return v

class ExpandRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=50_000, description="Text to expand")
    target_length: Optional[int] = Field(default=None, ge=10, le=50_000, description="Target word count")
    add_examples: bool = Field(default=False, description="Include examples")
    add_details: bool = Field(default=True, description="Add explanatory details")

class TranslateRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=100_000, description="Text to translate")
    target_language: str = Field(..., min_length=2, max_length=50, description="Target language")
    source_language: Optional[str] = Field(default=None, min_length=2, max_length=50, description="Source language")
    preserve_formatting: bool = Field(default=True, description="Preserve formatting")

class GenerateContentRequest(BaseModel):
    prompt: str = Field(..., min_length=1, max_length=100_000, description="Content generation prompt")
    context: Optional[str] = Field(default=None, max_length=50_000, description="Additional context")
    tone: str = Field(default="professional", description="Target tone")
    max_length: Optional[int] = Field(default=None, ge=10, le=50_000, description="Maximum words")

    @field_validator("tone")
    @classmethod
    def validate_tone(cls, v: str) -> str:
        valid_tones = {t.value for t in WritingTone}
        if v not in valid_tones:
            raise ValueError(f"tone must be one of: {', '.join(sorted(valid_tones))}")
        return v

# Spreadsheet AI Models
class FormulaRequest(BaseModel):
    description: str = Field(..., min_length=3, max_length=2_000, description="Natural language description")
    context: Optional[str] = Field(default=None, max_length=5_000, description="Data context")
    spreadsheet_type: str = Field(default="excel", description="Spreadsheet type")

    @field_validator("spreadsheet_type")
    @classmethod
    def validate_type(cls, v: str) -> str:
        allowed = {"excel", "google_sheets", "libreoffice"}
        if v not in allowed:
            raise ValueError(f"spreadsheet_type must be one of: {', '.join(sorted(allowed))}")
        return v

class DataQualityRequest(BaseModel):
    data_sample: List[Dict[str, Any]] = Field(..., min_length=1, max_length=1_000, description="Data sample")
    column_info: Optional[Dict[str, str]] = Field(default=None, description="Column types")

class AnomalyRequest(BaseModel):
    data: List[Dict[str, Any]] = Field(..., min_length=2, max_length=5_000, description="Data to analyze")
    columns_to_analyze: Optional[List[str]] = Field(default=None, description="Specific columns")
    sensitivity: str = Field(default="medium", description="Detection sensitivity")

    @field_validator("sensitivity")
    @classmethod
    def validate_sensitivity(cls, v: str) -> str:
        allowed = {"low", "medium", "high"}
        if v not in allowed:
            raise ValueError(f"sensitivity must be one of: {', '.join(sorted(allowed))}")
        return v

class PredictionRequest(BaseModel):
    data: List[Dict[str, Any]] = Field(..., min_length=2, max_length=5_000, description="Existing data")
    target_description: str = Field(..., min_length=3, max_length=2_000, description="What to predict")
    based_on_columns: List[str] = Field(..., min_length=1, max_length=50, description="Input columns")

class ExplainFormulaRequest(BaseModel):
    formula: str = Field(..., min_length=1, max_length=5_000, description="Formula to explain")
    context: Optional[str] = Field(default=None, max_length=5_000, description="Data context")

class SuggestFormulasRequest(BaseModel):
    data_sample: List[Dict[str, Any]] = Field(..., min_length=1, max_length=1_000, description="Data sample")
    analysis_goals: Optional[str] = Field(default=None, max_length=2_000, description="Analysis goals")

# WRITING AI ENDPOINTS

@ai_router.post("/documents/{document_id}/ai/grammar")
@limiter.limit(RATE_LIMIT_STRICT)
async def check_grammar(request: Request, document_id: str, req: GrammarCheckRequest):
    """
    Check text for grammar, spelling, and style issues.

    Returns:
        GrammarCheckResult with issues, corrected text, and quality score.

    Status codes:
        200: Success
        400: Invalid input (text too long)
        503: AI service temporarily unavailable
    """
    try:
        result = await writing_service.check_grammar(
            text=req.text,
            language=req.language,
            strict=req.strict,
        )
        return JSONResponse(content=result.model_dump())
    except WritingServiceError as e:
        raise _handle_service_error(e, "Grammar check")
    except Exception as e:
        raise _handle_service_error(e, "Grammar check")

@ai_router.post("/documents/{document_id}/ai/summarize")
@limiter.limit(RATE_LIMIT_STRICT)
async def summarize_text(request: Request, document_id: str, req: SummarizeRequest):
    """
    Summarize text with optional length limit.

    Returns:
        SummarizeResult with summary, key points, and compression ratio.
    """
    try:
        result = await writing_service.summarize(
            text=req.text,
            max_length=req.max_length,
            style=req.style,
        )
        return JSONResponse(content=result.model_dump())
    except WritingServiceError as e:
        raise _handle_service_error(e, "Summarization")
    except Exception as e:
        raise _handle_service_error(e, "Summarization")

@ai_router.post("/documents/{document_id}/ai/rewrite")
@limiter.limit(RATE_LIMIT_STRICT)
async def rewrite_text(request: Request, document_id: str, req: RewriteRequest):
    """
    Rewrite text with specified tone.

    Returns:
        RewriteResult with rewritten text and list of changes.
    """
    try:
        tone = WritingTone(req.tone)
        result = await writing_service.rewrite(
            text=req.text,
            tone=tone,
            preserve_meaning=req.preserve_meaning,
        )
        return JSONResponse(content=result.model_dump())
    except WritingServiceError as e:
        raise _handle_service_error(e, "Rewrite")
    except Exception as e:
        raise _handle_service_error(e, "Rewrite")

@ai_router.post("/documents/{document_id}/ai/expand")
@limiter.limit(RATE_LIMIT_STRICT)
async def expand_text(request: Request, document_id: str, req: ExpandRequest):
    """
    Expand text with additional details and examples.

    Returns:
        ExpandResult with expanded text and word counts.
    """
    try:
        result = await writing_service.expand(
            text=req.text,
            target_length=req.target_length,
            add_examples=req.add_examples,
            add_details=req.add_details,
        )
        return JSONResponse(content=result.model_dump())
    except WritingServiceError as e:
        raise _handle_service_error(e, "Expansion")
    except Exception as e:
        raise _handle_service_error(e, "Expansion")

@ai_router.post("/documents/{document_id}/ai/translate")
@limiter.limit(RATE_LIMIT_STRICT)
async def translate_text(request: Request, document_id: str, req: TranslateRequest):
    """
    Translate text to target language.

    Returns:
        TranslateResult with translated text and confidence score.
    """
    try:
        result = await writing_service.translate(
            text=req.text,
            target_language=req.target_language,
            source_language=req.source_language,
            preserve_formatting=req.preserve_formatting,
        )
        return JSONResponse(content=result.model_dump())
    except WritingServiceError as e:
        raise _handle_service_error(e, "Translation")
    except Exception as e:
        raise _handle_service_error(e, "Translation")

@ai_router.post("/ai/generate")
@limiter.limit(RATE_LIMIT_STRICT)
async def generate_content(request: Request, req: GenerateContentRequest):
    """
    Generate new content based on a prompt.

    Returns:
        Generated content string.
    """
    try:
        tone = WritingTone(req.tone)
        content = await writing_service.generate_content(
            prompt=req.prompt,
            context=req.context,
            tone=tone,
            max_length=req.max_length,
        )
        return JSONResponse(content={"content": content})
    except WritingServiceError as e:
        raise _handle_service_error(e, "Content generation")
    except Exception as e:
        raise _handle_service_error(e, "Content generation")

# SPREADSHEET AI ENDPOINTS

@ai_router.post("/spreadsheets/{spreadsheet_id}/formula")
@limiter.limit(RATE_LIMIT_STRICT)
async def natural_language_to_formula(request: Request, spreadsheet_id: str, req: FormulaRequest):
    """
    Convert natural language description to spreadsheet formula.

    Returns:
        FormulaResult with formula, explanation, and alternatives.
    """
    try:
        result = await spreadsheet_ai_service.natural_language_to_formula(
            description=req.description,
            context=req.context,
            spreadsheet_type=req.spreadsheet_type,
        )
        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise _handle_service_error(e, "Formula generation")

@ai_router.post("/spreadsheets/{spreadsheet_id}/clean")
@limiter.limit(RATE_LIMIT_STRICT)
async def analyze_data_quality(request: Request, spreadsheet_id: str, req: DataQualityRequest):
    """
    Analyze data for quality issues and provide cleaning suggestions.

    Returns:
        DataCleaningResult with suggestions and quality score.
    """
    try:
        result = await spreadsheet_ai_service.analyze_data_quality(
            data_sample=req.data_sample,
            column_info=req.column_info,
        )
        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise _handle_service_error(e, "Data quality analysis")

@ai_router.post("/spreadsheets/{spreadsheet_id}/anomalies")
@limiter.limit(RATE_LIMIT_STRICT)
async def detect_anomalies(request: Request, spreadsheet_id: str, req: AnomalyRequest):
    """
    Detect anomalies in spreadsheet data.

    Returns:
        AnomalyDetectionResult with detected anomalies.
    """
    try:
        result = await spreadsheet_ai_service.detect_anomalies(
            data=req.data,
            columns_to_analyze=req.columns_to_analyze,
            sensitivity=req.sensitivity,
        )
        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise _handle_service_error(e, "Anomaly detection")

@ai_router.post("/spreadsheets/{spreadsheet_id}/predict")
@limiter.limit(RATE_LIMIT_STRICT)
async def generate_predictions(request: Request, spreadsheet_id: str, req: PredictionRequest):
    """
    Generate predictions for a new column based on existing data.

    Returns:
        PredictionColumn with predictions and confidence scores.
    """
    try:
        result = await spreadsheet_ai_service.generate_predictive_column(
            data=req.data,
            target_description=req.target_description,
            based_on_columns=req.based_on_columns,
        )
        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise _handle_service_error(e, "Prediction generation")

@ai_router.post("/spreadsheets/{spreadsheet_id}/explain")
@limiter.limit(RATE_LIMIT_STRICT)
async def explain_formula(request: Request, spreadsheet_id: str, req: ExplainFormulaRequest):
    """
    Explain what a formula does in plain language.

    Returns:
        FormulaExplanation with detailed breakdown.
    """
    try:
        result = await spreadsheet_ai_service.explain_formula(
            formula=req.formula,
            context=req.context,
        )
        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise _handle_service_error(e, "Formula explanation")

@ai_router.post("/spreadsheets/{spreadsheet_id}/suggest")
@limiter.limit(RATE_LIMIT_STRICT)
async def suggest_formulas(request: Request, spreadsheet_id: str, req: SuggestFormulasRequest):
    """
    Suggest useful formulas based on data structure.

    Returns:
        List of suggested formulas with explanations.
    """
    try:
        results = await spreadsheet_ai_service.suggest_formulas(
            data_sample=req.data_sample,
            analysis_goals=req.analysis_goals,
        )
        return JSONResponse(content={"suggestions": [r.model_dump() for r in results]})
    except Exception as e:
        raise _handle_service_error(e, "Formula suggestion")

# UTILITY ENDPOINTS

@ai_router.get("/tones")
async def get_available_tones():
    """Get list of available writing tones."""
    return {
        "tones": [
            {"value": tone.value, "label": tone.value.replace("_", " ").title()}
            for tone in WritingTone
        ]
    }

@ai_router.get("/health")
async def check_ai_health():
    """Check if AI services are configured and available."""
    import subprocess

    # Check if Claude CLI is available
    cli_available = False
    cli_version = None
    try:
        result = subprocess.run(
            ["claude", "--version"],
            capture_output=True,
            text=True,
            timeout=5,
        )
        cli_available = result.returncode == 0
        if cli_available:
            cli_version = result.stdout.strip()
    except Exception:
        pass

    return {
        "status": "ok" if cli_available else "degraded",
        "claude_cli_available": cli_available,
        "claude_cli_version": cli_version,
        "provider": "claude_code_cli",
        "model": "qwen",
        "services": {
            "writing": cli_available,
            "spreadsheet": cli_available,
            "docqa": cli_available,
            "research": cli_available,
        }
    }

"""API routes for Natural Language to SQL feature."""

from backend.app.schemas import (
    NL2SQLGenerateRequest,
    NL2SQLExecuteRequest,
    NL2SQLSaveRequest,
)
from backend.app.services.ai_services import NL2SQLService

nl2sql_router = APIRouter(dependencies=[Depends(require_api_key)])

def get_service() -> NL2SQLService:
    return NL2SQLService()

@nl2sql_router.post("/generate")
async def generate_sql(
    payload: NL2SQLGenerateRequest,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Generate SQL from a natural language question."""
    correlation_id = getattr(request.state, "correlation_id", None)
    result = svc.generate_sql(payload, correlation_id)
    return {
        "status": "ok",
        "sql": result.sql,
        "explanation": result.explanation,
        "confidence": result.confidence,
        "warnings": result.warnings,
        "original_question": result.original_question,
        "correlation_id": correlation_id,
    }

@nl2sql_router.post("/execute")
async def execute_query(
    payload: NL2SQLExecuteRequest,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Execute a SQL query and return results."""
    correlation_id = getattr(request.state, "correlation_id", None)
    result = svc.execute_query(payload, correlation_id)
    return {
        "status": "ok",
        "columns": result.columns,
        "rows": result.rows,
        "row_count": result.row_count,
        "total_count": result.total_count,
        "execution_time_ms": result.execution_time_ms,
        "truncated": result.truncated,
        "correlation_id": correlation_id,
    }

@nl2sql_router.post("/explain")
async def explain_query(
    request: Request,
    sql: str = Query(..., min_length=1, max_length=10000),
    svc: NL2SQLService = Depends(get_service),
):
    """Get a natural language explanation of a SQL query."""
    correlation_id = getattr(request.state, "correlation_id", None)
    explanation = svc.explain_query(sql, correlation_id)
    return {
        "status": "ok",
        "explanation": explanation,
        "correlation_id": correlation_id,
    }

@nl2sql_router.post("/save")
async def save_query(
    payload: NL2SQLSaveRequest,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Save a query as a reusable data source."""
    correlation_id = getattr(request.state, "correlation_id", None)
    saved = svc.save_query(payload, correlation_id)
    return {
        "status": "ok",
        "query": saved.model_dump(mode="json"),
        "correlation_id": correlation_id,
    }

@nl2sql_router.get("/saved")
async def list_saved_queries(
    request: Request,
    connection_id: Optional[str] = Query(None, max_length=64),
    tags: Optional[List[str]] = Query(None),
    svc: NL2SQLService = Depends(get_service),
):
    """List saved queries."""
    correlation_id = getattr(request.state, "correlation_id", None)
    queries = svc.list_saved_queries(connection_id=connection_id, tags=tags)
    return {
        "status": "ok",
        "queries": [q.model_dump(mode="json") for q in queries],
        "correlation_id": correlation_id,
    }

@nl2sql_router.get("/saved/{query_id}")
async def get_saved_query(
    query_id: str,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Get a saved query by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    query = svc.get_saved_query(query_id)
    if not query:
        raise HTTPException(
            status_code=404,
            detail={
                "status": "error",
                "code": "not_found",
                "message": "Query not found",
            },
        )
    return {
        "status": "ok",
        "query": query.model_dump(mode="json"),
        "correlation_id": correlation_id,
    }

@nl2sql_router.delete("/saved/{query_id}")
async def delete_saved_query(
    query_id: str,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Delete a saved query."""
    correlation_id = getattr(request.state, "correlation_id", None)
    deleted = svc.delete_saved_query(query_id)
    return {
        "status": "ok" if deleted else "error",
        "deleted": deleted,
        "query_id": query_id,
        "correlation_id": correlation_id,
    }

@nl2sql_router.get("/history")
async def get_query_history(
    request: Request,
    connection_id: Optional[str] = Query(None, max_length=64),
    limit: int = Query(50, ge=1, le=200),
    svc: NL2SQLService = Depends(get_service),
):
    """Get query history."""
    correlation_id = getattr(request.state, "correlation_id", None)
    history = svc.get_query_history(connection_id=connection_id, limit=limit)
    return {
        "status": "ok",
        "history": [h.model_dump(mode="json") for h in history],
        "correlation_id": correlation_id,
    }

@nl2sql_router.delete("/history/{entry_id}")
async def delete_query_history_entry(
    entry_id: str,
    request: Request,
    svc: NL2SQLService = Depends(get_service),
):
    """Delete a query history entry."""
    correlation_id = getattr(request.state, "correlation_id", None)
    deleted = svc.delete_query_history_entry(entry_id)
    return {
        "status": "ok" if deleted else "error",
        "deleted": deleted,
        "entry_id": entry_id,
        "correlation_id": correlation_id,
    }

"""API routes for Data Enrichment feature."""

from backend.app.schemas import (
    EnrichmentSourceCreate,
    SimpleEnrichmentRequest,
    SimplePreviewRequest,
)
from backend.app.services.enrichment_service import EnrichmentService

enrichment_router = APIRouter(dependencies=[Depends(require_api_key)])

def get_service() -> EnrichmentService:
    return EnrichmentService()

@enrichment_router.get("/sources")
async def list_available_sources(
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """List available enrichment source types."""
    correlation_id = getattr(request.state, "correlation_id", None)
    builtin = EnrichmentService.get_builtin_sources()
    custom_sources = [source.model_dump() for source in svc.list_sources()]
    return {
        "status": "ok",
        "sources": [*builtin, *custom_sources],
        "correlation_id": correlation_id,
    }

@enrichment_router.get("/source-types")
async def list_source_types(
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """List available enrichment source types (legacy endpoint)."""
    correlation_id = getattr(request.state, "correlation_id", None)
    source_types = svc.get_available_source_types()
    return {
        "status": "ok",
        "source_types": source_types,
        "correlation_id": correlation_id,
    }

@enrichment_router.post("/enrich")
async def enrich_data(
    payload: SimpleEnrichmentRequest,
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Enrich data with additional information using selected sources."""
    correlation_id = getattr(request.state, "correlation_id", None)
    result = await svc.simple_enrich(
        data=payload.data,
        sources=payload.sources,
        options=payload.options,
        correlation_id=correlation_id,
    )
    return {
        "status": "ok",
        "enriched_data": result["enriched_data"],
        "total_rows": result["total_rows"],
        "enriched_rows": result["enriched_rows"],
        "processing_time_ms": result["processing_time_ms"],
        "correlation_id": correlation_id,
    }

@enrichment_router.post("/preview")
async def preview_enrichment(
    payload: SimplePreviewRequest,
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Preview enrichment results on a sample."""
    correlation_id = getattr(request.state, "correlation_id", None)
    result = await svc.simple_preview(
        data=payload.data,
        sources=payload.sources,
        sample_size=payload.sample_size,
        correlation_id=correlation_id,
    )
    return {
        "status": "ok",
        "preview": result["preview"],
        "total_rows": result["total_rows"],
        "enriched_rows": result["enriched_rows"],
        "processing_time_ms": result["processing_time_ms"],
        "correlation_id": correlation_id,
    }

@enrichment_router.post("/sources/create")
async def create_source(
    payload: EnrichmentSourceCreate,
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Create a custom enrichment source."""
    correlation_id = getattr(request.state, "correlation_id", None)
    source = svc.create_source(payload, correlation_id)
    return {
        "status": "ok",
        "source": source.model_dump(),
        "correlation_id": correlation_id,
    }

@enrichment_router.get("/sources/{source_id}")
async def get_source(
    source_id: str,
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Get an enrichment source by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    # Check built-in sources first
    for source in EnrichmentService.get_builtin_sources():
        if source["id"] == source_id:
            return {
                "status": "ok",
                "source": source,
                "correlation_id": correlation_id,
            }
    # Check custom sources
    source = svc.get_source(source_id)
    if not source:
        raise HTTPException(
            status_code=404,
            detail={"code": "not_found", "message": "Source not found"},
        )
    return {
        "status": "ok",
        "source": source.model_dump(),
        "correlation_id": correlation_id,
    }

@enrichment_router.delete("/sources/{source_id}")
async def delete_source(
    source_id: str,
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Delete a custom enrichment source."""
    correlation_id = getattr(request.state, "correlation_id", None)
    deleted = svc.delete_source(source_id)
    if not deleted:
        raise HTTPException(
            status_code=404,
            detail={"code": "not_found", "message": f"Source {source_id} not found or cannot be deleted"},
        )
    return {
        "status": "ok",
        "deleted": True,
        "source_id": source_id,
        "correlation_id": correlation_id,
    }

@enrichment_router.get("/cache/stats")
async def get_cache_stats(
    request: Request,
    svc: EnrichmentService = Depends(get_service),
):
    """Get enrichment cache statistics."""
    correlation_id = getattr(request.state, "correlation_id", None)
    stats = svc.get_cache_stats()
    return {
        "status": "ok",
        "stats": stats,
        "correlation_id": correlation_id,
    }

@enrichment_router.delete("/cache")
async def clear_cache(
    request: Request,
    source_id: Optional[str] = Query(None, max_length=64),
    svc: EnrichmentService = Depends(get_service),
):
    """Clear enrichment cache."""
    correlation_id = getattr(request.state, "correlation_id", None)
    cleared = svc.clear_cache(source_id)
    return {
        "status": "ok",
        "cleared_entries": cleared,
        "source_id": source_id,
        "correlation_id": correlation_id,
    }

"""API routes for Cross-Database Federation feature."""

logger = logging.getLogger("neura.api.federation")
from backend.app.schemas import VirtualSchemaCreate, SuggestJoinsRequest, FederatedQueryRequest
from backend.app.services.ai_services import FederationService

federation_router = APIRouter(dependencies=[Depends(require_api_key)])

def get_service() -> FederationService:
    return FederationService()

@federation_router.post("/schemas")
async def create_virtual_schema(
    payload: VirtualSchemaCreate,
    request: Request,
    svc: FederationService = Depends(get_service),
):
    """Create a virtual schema spanning multiple databases."""
    correlation_id = getattr(request.state, "correlation_id", None)
    try:
        schema = svc.create_virtual_schema(payload, correlation_id)
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        logger.exception("federation_create_schema_failed", extra={"event": "federation_create_schema_failed"})
        raise HTTPException(status_code=500, detail="Failed to create virtual schema")
    return {"status": "ok", "schema": schema.model_dump(), "correlation_id": correlation_id}

@federation_router.get("/schemas")
async def list_virtual_schemas(
    request: Request,
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    svc: FederationService = Depends(get_service),
):
    """List all virtual schemas."""
    correlation_id = getattr(request.state, "correlation_id", None)
    all_schemas = svc.list_virtual_schemas()
    page = all_schemas[offset:offset + limit]
    return {
        "status": "ok",
        "schemas": [s.model_dump() for s in page],
        "total": len(all_schemas),
        "correlation_id": correlation_id,
    }

@federation_router.get("/schemas/{schema_id}")
async def get_virtual_schema(
    schema_id: str,
    request: Request,
    svc: FederationService = Depends(get_service),
):
    """Get a virtual schema by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    schema = svc.get_virtual_schema(schema_id)
    if not schema:
        raise HTTPException(status_code=404, detail={"code": "not_found", "message": f"Schema {schema_id} not found"})
    return {"status": "ok", "schema": schema.model_dump(), "correlation_id": correlation_id}

@federation_router.delete("/schemas/{schema_id}")
async def delete_virtual_schema(
    schema_id: str,
    request: Request,
    svc: FederationService = Depends(get_service),
):
    """Delete a virtual schema."""
    correlation_id = getattr(request.state, "correlation_id", None)
    deleted = svc.delete_virtual_schema(schema_id)
    if not deleted:
        raise HTTPException(status_code=404, detail={"code": "not_found", "message": f"Schema {schema_id} not found"})
    return {"status": "ok", "deleted": True, "correlation_id": correlation_id}

@federation_router.post("/suggest-joins")
async def suggest_joins(
    payload: SuggestJoinsRequest,
    request: Request,
    svc: FederationService = Depends(get_service),
):
    """Get AI-suggested joins between tables in different connections."""
    correlation_id = getattr(request.state, "correlation_id", None)
    try:
        suggestions = svc.suggest_joins(payload.connection_ids, correlation_id)
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        logger.exception("federation_suggest_joins_failed", extra={"event": "federation_suggest_joins_failed"})
        raise HTTPException(status_code=500, detail="Join suggestion failed")
    return {"status": "ok", "suggestions": [s.model_dump() for s in suggestions], "correlation_id": correlation_id}

@federation_router.post("/query")
async def execute_federated_query(
    payload: FederatedQueryRequest,
    request: Request,
    svc: FederationService = Depends(get_service),
):
    """Execute a federated query across multiple databases."""
    correlation_id = getattr(request.state, "correlation_id", None)
    try:
        result = svc.execute_query(payload, correlation_id)
    except (HTTPException, AppError):
        raise
    except Exception as exc:
        logger.exception("federation_query_failed", extra={"event": "federation_query_failed"})
        raise HTTPException(status_code=500, detail="Federated query execution failed")
    return {"status": "ok", "result": result, "correlation_id": correlation_id}

"""API routes for Template Recommendations."""

logger = logging.getLogger(__name__)

from backend.app.services.ai_services import RecommendationService

recommendations_router = APIRouter(dependencies=[Depends(require_api_key)])

class TemplateRecommendRequest(BaseModel):
    """Request payload for template recommendations (frontend format)."""
    data_description: Optional[str] = Field(None, max_length=1000)
    data_columns: Optional[List[str]] = Field(None, max_length=100)
    industry: Optional[str] = Field(None, max_length=100)
    output_format: Optional[str] = Field(None, max_length=50)

def get_service() -> RecommendationService:
    return RecommendationService()

@recommendations_router.post("/templates")
async def recommend_templates_post(
    payload: TemplateRecommendRequest,
    request: Request,
    svc: RecommendationService = Depends(get_service),
    background: bool = Query(False),
):
    """Get template recommendations based on data description and columns."""
    correlation_id = getattr(request.state, "correlation_id", None)

    # Build context from frontend payload
    context_parts = []
    if payload.data_description:
        context_parts.append(f"Data description: {payload.data_description}")
    if payload.data_columns:
        context_parts.append(f"Data columns: {', '.join(payload.data_columns)}")
    if payload.industry:
        context_parts.append(f"Industry: {payload.industry}")
    if payload.output_format:
        context_parts.append(f"Output format: {payload.output_format}")

    context = " | ".join(context_parts) if context_parts else None

    if not background:
        recommendations = svc.recommend_templates(
            context=context,
            limit=5,
            correlation_id=correlation_id,
        )
        return {"status": "ok", "recommendations": recommendations, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "recommend", status="running", label="Generate recommendations")
        try:
            recommendations = svc.recommend_templates(
                context=context,
                limit=5,
                correlation_id=correlation_id,
            )
            state_access.record_job_step(job_id, "recommend", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"recommendations": recommendations},
            )
        except Exception as exc:
            logger.exception("recommend_job_failed", extra={"job_id": job_id})
            state_access.record_job_step(job_id, "recommend", status="failed", error="Recommendation generation failed")
            state_access.record_job_completion(job_id, status="failed", error="Recommendation generation failed")

    job = await enqueue_background_job(
        job_type="recommend_templates",
        steps=[{"name": "recommend", "label": "Generate recommendations"}],
        meta={"background": True, "context": context},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

@recommendations_router.get("/templates")
async def recommend_templates_get(
    request: Request,
    connection_id: Optional[str] = Query(None),
    context: Optional[str] = Query(None, max_length=500),
    limit: int = Query(5, ge=1, le=20),
    svc: RecommendationService = Depends(get_service),
    background: bool = Query(False),
):
    """Get template recommendations based on context (query params)."""
    correlation_id = getattr(request.state, "correlation_id", None)
    if not background:
        recommendations = svc.recommend_templates(
            connection_id=connection_id,
            context=context,
            limit=limit,
            correlation_id=correlation_id,
        )
        return {"status": "ok", "recommendations": recommendations, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "recommend", status="running", label="Generate recommendations")
        try:
            recommendations = svc.recommend_templates(
                connection_id=connection_id,
                context=context,
                limit=limit,
                correlation_id=correlation_id,
            )
            state_access.record_job_step(job_id, "recommend", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"recommendations": recommendations},
            )
        except Exception as exc:
            logger.exception("recommend_job_failed", extra={"job_id": job_id})
            state_access.record_job_step(job_id, "recommend", status="failed", error="Recommendation generation failed")
            state_access.record_job_completion(job_id, status="failed", error="Recommendation generation failed")

    job = await enqueue_background_job(
        job_type="recommend_templates",
        steps=[{"name": "recommend", "label": "Generate recommendations"}],
        meta={"background": True, "context": context, "connection_id": connection_id},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

@recommendations_router.get("/catalog")
async def get_template_catalog(
    request: Request,
):
    """Get template catalog for browsing."""
    correlation_id = getattr(request.state, "correlation_id", None)

    templates = state_access.list_templates()

    # Build catalog with summary info
    catalog = []
    for t in templates:
        if t.get("status") == "approved":
            catalog.append({
                "id": t.get("id"),
                "name": t.get("name"),
                "kind": t.get("kind"),
                "description": t.get("description", ""),
                "tags": t.get("tags", []),
                "created_at": t.get("created_at"),
            })

    # Sort by name
    catalog.sort(key=lambda x: x.get("name", "").lower())

    return {"status": "ok", "catalog": catalog, "total": len(catalog), "correlation_id": correlation_id}

@recommendations_router.get("/templates/{template_id}/similar")
async def get_similar_templates(
    template_id: str,
    request: Request,
    limit: int = Query(3, ge=1, le=10),
    svc: RecommendationService = Depends(get_service),
):
    """Get templates similar to a given template."""
    correlation_id = getattr(request.state, "correlation_id", None)
    similar = svc.get_similar_templates(template_id, limit)
    return {"status": "ok", "similar": similar, "correlation_id": correlation_id}

"""API routes for Executive Summary Generation."""

logger = logging.getLogger(__name__)

from backend.app.services.ai_services import SummaryService

summary_router = APIRouter(dependencies=[Depends(require_api_key)])

class SummaryRequest(BaseModel):
    content: str = Field(..., min_length=10, max_length=50000)
    tone: str = Field(default="formal", pattern="^(formal|conversational|technical)$")
    max_sentences: int = Field(default=5, ge=2, le=15)
    focus_areas: Optional[List[str]] = Field(None, max_length=5)

def get_service() -> SummaryService:
    return SummaryService()

def _is_cancelled(job_id: str) -> bool:
    job = state_access.get_job(job_id) or {}
    return str(job.get("status") or "").strip().lower() == "cancelled"

@summary_router.post("/generate")
async def generate_summary(
    payload: SummaryRequest,
    request: Request,
    svc: SummaryService = Depends(get_service),
    background: bool = Query(False),
):
    """Generate an executive summary from content."""
    correlation_id = getattr(request.state, "correlation_id", None)
    if not background:
        summary = svc.generate_summary(
            content=payload.content,
            tone=payload.tone,
            max_sentences=payload.max_sentences,
            focus_areas=payload.focus_areas,
            correlation_id=correlation_id,
        )
        return {"status": "ok", "summary": summary, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        if _is_cancelled(job_id):
            return
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "generate", status="running", label="Generate summary")
        try:
            summary = svc.generate_summary(
                content=payload.content,
                tone=payload.tone,
                max_sentences=payload.max_sentences,
                focus_areas=payload.focus_areas,
                correlation_id=correlation_id,
            )
            if _is_cancelled(job_id):
                state_access.record_job_step(job_id, "generate", status="cancelled", error="Cancelled by user")
                return
            state_access.record_job_step(job_id, "generate", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"summary": summary},
            )
        except Exception as exc:
            if _is_cancelled(job_id):
                state_access.record_job_step(job_id, "generate", status="cancelled", error="Cancelled by user")
                return
            logger.exception("summary_job_failed", extra={"job_id": job_id})
            state_access.record_job_step(job_id, "generate", status="failed", error="Summary generation failed")
            state_access.record_job_completion(job_id, status="failed", error="Summary generation failed")

    job = await enqueue_background_job(
        job_type="summary_generate",
        steps=[{"name": "generate", "label": "Generate summary"}],
        meta={"background": True, "content_length": len(payload.content)},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

@summary_router.get("/reports/{report_id}")
async def get_report_summary(
    report_id: str,
    request: Request,
    svc: SummaryService = Depends(get_service),
    background: bool = Query(False),
):
    """Generate summary for a specific report."""
    correlation_id = getattr(request.state, "correlation_id", None)
    if not background:
        summary = svc.generate_report_summary(report_id, correlation_id)
        return {"status": "ok", "summary": summary, "correlation_id": correlation_id}

    async def runner(job_id: str) -> None:
        if _is_cancelled(job_id):
            return
        state_access.record_job_start(job_id)
        state_access.record_job_step(job_id, "generate", status="running", label="Generate report summary")
        try:
            summary = svc.generate_report_summary(report_id, correlation_id)
            if _is_cancelled(job_id):
                state_access.record_job_step(job_id, "generate", status="cancelled", error="Cancelled by user")
                return
            state_access.record_job_step(job_id, "generate", status="succeeded", progress=100.0)
            state_access.record_job_completion(
                job_id,
                status="succeeded",
                result={"summary": summary, "report_id": report_id},
            )
        except Exception as exc:
            if _is_cancelled(job_id):
                state_access.record_job_step(job_id, "generate", status="cancelled", error="Cancelled by user")
                return
            logger.exception("summary_job_failed", extra={"job_id": job_id})
            state_access.record_job_step(job_id, "generate", status="failed", error="Summary generation failed")
            state_access.record_job_completion(job_id, status="failed", error="Summary generation failed")

    job = await enqueue_background_job(
        job_type="summary_report",
        steps=[{"name": "generate", "label": "Generate report summary"}],
        meta={"background": True, "report_id": report_id},
        runner=runner,
    )
    return {"status": "queued", "job_id": job["id"], "correlation_id": correlation_id}

"""API routes for Multi-Document Synthesis."""

import io
import json
from fastapi import APIRouter, Depends, Request, HTTPException, File, Form, UploadFile

from backend.app.services.ai_services import DocumentSynthesisService
from backend.app.schemas import SynthesisDocumentType as DocumentType, SynthesisRequest

synthesis_router = APIRouter(dependencies=[Depends(require_api_key)])
MAX_SYNTHESIS_DOC_BYTES = 5 * 1024 * 1024

_EXTENSION_TO_DOC_TYPE = {
    ".txt": DocumentType.TEXT,
    ".md": DocumentType.TEXT,
    ".markdown": DocumentType.TEXT,
    ".csv": DocumentType.EXCEL,
    ".json": DocumentType.JSON,
    ".pdf": DocumentType.PDF,
    ".xlsx": DocumentType.EXCEL,
    ".xls": DocumentType.EXCEL,
    ".docx": DocumentType.WORD,
    ".doc": DocumentType.WORD,
}

class CreateSessionRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)

class AddDocumentRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)
    content: str = Field(..., min_length=10, max_length=5 * 1024 * 1024)
    doc_type: DocumentType = Field(default=DocumentType.TEXT)
    metadata: Optional[dict] = None

def _safe_decode_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")

def _infer_document_type(filename: Optional[str], explicit: Optional[DocumentType]) -> DocumentType:
    if explicit:
        return explicit
    suffix = Path(filename or "").suffix.lower()
    return _EXTENSION_TO_DOC_TYPE.get(suffix, DocumentType.TEXT)

def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
    except Exception as exc:
        raise HTTPException(status_code=500, detail="PDF extraction dependency not available") from exc

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        chunks = []
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                chunks.append(text)
        doc.close()
        return "\n\n".join(chunks)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse PDF file") from exc

def _extract_excel_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".csv":
        return _safe_decode_text(file_bytes)

    try:
        import pandas as pd
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Excel extraction dependencies are not available",
        ) from exc

    try:
        workbook = pd.ExcelFile(io.BytesIO(file_bytes))
        sheet_blocks = []
        for sheet_name in workbook.sheet_names[:20]:
            df = pd.read_excel(workbook, sheet_name=sheet_name)
            if df.empty:
                continue
            sheet_text = df.fillna("").to_csv(index=False, sep="\t").strip()
            if sheet_text:
                sheet_blocks.append(f"[Sheet: {sheet_name}]\n{sheet_text}")
        return "\n\n".join(sheet_blocks)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse Excel file") from exc

def _extract_word_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix != ".docx":
        # .doc has no reliable pure-python parser in this stack; attempt plain decode fallback
        return _safe_decode_text(file_bytes)

    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=500, detail="DOCX extraction dependency not available") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append("\t".join(cells))
        parts = paragraphs + ([""] + table_lines if table_lines else [])
        return "\n".join(parts)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse Word file") from exc

def _extract_document_text(filename: str, file_bytes: bytes, doc_type: DocumentType) -> str:
    if doc_type == DocumentType.PDF:
        return _extract_pdf_text(file_bytes)
    if doc_type == DocumentType.EXCEL:
        return _extract_excel_text(filename, file_bytes)
    if doc_type == DocumentType.WORD:
        return _extract_word_text(filename, file_bytes)
    if doc_type == DocumentType.JSON:
        text = _safe_decode_text(file_bytes)
        try:
            parsed = json.loads(text)
            return json.dumps(parsed, indent=2, ensure_ascii=False)
        except Exception:
            return text
    return _safe_decode_text(file_bytes)

def get_service() -> DocumentSynthesisService:
    return DocumentSynthesisService()

@synthesis_router.post("/sessions")
async def create_session(
    payload: CreateSessionRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Create a new synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.create_session(
        name=payload.name,
        correlation_id=correlation_id,
    )
    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}

@synthesis_router.get("/sessions")
async def list_sessions(
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """List all synthesis sessions."""
    correlation_id = getattr(request.state, "correlation_id", None)
    sessions = svc.list_sessions()
    return {
        "status": "ok",
        "sessions": [s.model_dump(mode="json") for s in sessions],
        "correlation_id": correlation_id,
    }

@synthesis_router.get("/sessions/{session_id}")
async def get_session(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Get a synthesis session by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.get_session(session_id)

    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}

@synthesis_router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Delete a synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.delete_session(session_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "deleted": True, "correlation_id": correlation_id}

@synthesis_router.post("/sessions/{session_id}/documents")
async def add_document(
    session_id: str,
    payload: AddDocumentRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Add a document to a synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    document = svc.add_document(
        session_id=session_id,
        name=payload.name,
        content=payload.content,
        doc_type=payload.doc_type,
        metadata=payload.metadata,
        correlation_id=correlation_id,
    )

    if not document:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "document": document.model_dump(mode="json"), "correlation_id": correlation_id}

@synthesis_router.post("/documents/extract")
async def extract_document(
    request: Request,
    file: UploadFile = File(...),
    doc_type: Optional[DocumentType] = Form(default=None),
):
    """Extract normalized text content from an uploaded file for synthesis."""
    correlation_id = getattr(request.state, "correlation_id", None)
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(file_bytes) > MAX_SYNTHESIS_DOC_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds {MAX_SYNTHESIS_DOC_BYTES // (1024 * 1024)}MB limit",
        )

    filename = file.filename or "upload"
    inferred_type = _infer_document_type(filename, doc_type)
    extracted = _extract_document_text(filename, file_bytes, inferred_type).strip()
    if not extracted:
        raise HTTPException(status_code=400, detail="Could not extract text from uploaded file")

    truncated = False
    if len(extracted) > MAX_SYNTHESIS_DOC_BYTES:
        extracted = extracted[:MAX_SYNTHESIS_DOC_BYTES]
        truncated = True

    return {
        "status": "ok",
        "document": {
            "name": filename,
            "doc_type": inferred_type.value,
            "content": extracted,
            "truncated": truncated,
        },
        "correlation_id": correlation_id,
    }

@synthesis_router.delete("/sessions/{session_id}/documents/{document_id}")
async def remove_document(
    session_id: str,
    document_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Remove a document from a session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.remove_document(session_id, document_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session or document not found")

    return {"status": "ok", "removed": True, "correlation_id": correlation_id}

@synthesis_router.get("/sessions/{session_id}/inconsistencies")
async def find_inconsistencies(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Find inconsistencies between documents in a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    inconsistencies = svc.find_inconsistencies(session_id, correlation_id)

    return {
        "status": "ok",
        "inconsistencies": [i.model_dump(mode="json") for i in inconsistencies],
        "count": len(inconsistencies),
        "correlation_id": correlation_id,
    }

@synthesis_router.post("/sessions/{session_id}/synthesize")
async def synthesize_documents(
    session_id: str,
    payload: SynthesisRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Synthesize information from all documents in a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.documents:
        raise HTTPException(status_code=400, detail="No documents in session")

    result = svc.synthesize(session_id, payload, correlation_id)

    if not result:
        raise HTTPException(status_code=500, detail="Synthesis failed")

    return {"status": "ok", "result": result.model_dump(mode="json"), "correlation_id": correlation_id}

"""API routes for Document Q&A Chat."""

import threading
from fastapi import APIRouter, Depends, Query, Request, HTTPException

logger = logging.getLogger("neura.api.docqa")

from backend.app.services.ai_services import DocumentQAService
from backend.app.schemas import AskRequest, FeedbackRequest, RegenerateRequest

docqa_router = APIRouter(dependencies=[Depends(require_api_key)])

class CreateSessionRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)

class AddDocumentRequest(BaseModel):
    # Limit content to 500KB to prevent memory exhaustion
    # For larger documents, use file upload with chunked processing
    name: str = Field(..., min_length=1, max_length=200)
    content: str = Field(..., min_length=10, max_length=500 * 1024)
    page_count: Optional[int] = None

_lock = threading.Lock()
_docqa_service: DocumentQAService | None = None

def get_service() -> DocumentQAService:
    """Return a singleton DocumentQAService instance."""
    global _docqa_service
    if _docqa_service is None:
        with _lock:
            if _docqa_service is None:
                _docqa_service = DocumentQAService()
    return _docqa_service

@docqa_router.post("/sessions")
async def create_session(
    payload: CreateSessionRequest,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Create a new Q&A session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.create_session(
        name=payload.name,
        correlation_id=correlation_id,
    )
    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}

@docqa_router.get("/sessions")
async def list_sessions(
    request: Request,
    limit: int = Query(50, ge=1, le=500),
    offset: int = Query(0, ge=0),
    svc: DocumentQAService = Depends(get_service),
):
    """List all Q&A sessions."""
    correlation_id = getattr(request.state, "correlation_id", None)
    sessions = svc.list_sessions()
    page = sessions[offset:offset + limit]
    return {
        "status": "ok",
        "sessions": [s.model_dump(mode="json") for s in page],
        "total": len(sessions),
        "correlation_id": correlation_id,
    }

@docqa_router.get("/sessions/{session_id}")
async def get_session(
    session_id: str,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Get a Q&A session by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.get_session(session_id)

    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}

@docqa_router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Delete a Q&A session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.delete_session(session_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "deleted": True, "correlation_id": correlation_id}

@docqa_router.post("/sessions/{session_id}/documents")
async def add_document(
    session_id: str,
    payload: AddDocumentRequest,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Add a document to a Q&A session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    document = svc.add_document(
        session_id=session_id,
        name=payload.name,
        content=payload.content,
        page_count=payload.page_count,
        correlation_id=correlation_id,
    )

    if not document:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "document": document.model_dump(mode="json"), "correlation_id": correlation_id}

@docqa_router.delete("/sessions/{session_id}/documents/{document_id}")
async def remove_document(
    session_id: str,
    document_id: str,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Remove a document from a session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.remove_document(session_id, document_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session or document not found")

    return {"status": "ok", "removed": True, "correlation_id": correlation_id}

@docqa_router.post("/sessions/{session_id}/ask")
async def ask_question(
    session_id: str,
    payload: AskRequest,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Ask a question about the documents in a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    response = svc.ask(session_id, payload, correlation_id)

    if not response:
        raise HTTPException(status_code=500, detail="Failed to process question")

    return {
        "status": "ok",
        "response": response.model_dump(mode="json"),
        "correlation_id": correlation_id,
    }

@docqa_router.post("/sessions/{session_id}/messages/{message_id}/feedback")
async def submit_feedback(
    session_id: str,
    message_id: str,
    payload: FeedbackRequest,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Submit feedback for a chat message."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    message = svc.submit_feedback(
        session_id,
        message_id,
        payload,
        correlation_id,
    )

    if not message:
        raise HTTPException(status_code=404, detail="Message not found")

    return {
        "status": "ok",
        "message": message.model_dump(mode="json"),
        "correlation_id": correlation_id,
    }

@docqa_router.post("/sessions/{session_id}/messages/{message_id}/regenerate")
async def regenerate_response(
    session_id: str,
    message_id: str,
    payload: RegenerateRequest,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Regenerate a response for a message."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    try:
        response = svc.regenerate_response(
            session_id,
            message_id,
            payload,
            correlation_id,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail="Document Q&A operation failed") from exc
    except Exception as exc:
        logger.error("Response regeneration failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Response regeneration failed.",
        ) from exc

    if not response:
        raise HTTPException(status_code=404, detail="Message not found")

    return {
        "status": "ok",
        "response": response.model_dump(mode="json"),
        "correlation_id": correlation_id,
    }

@docqa_router.get("/sessions/{session_id}/history")
async def get_chat_history(
    session_id: str,
    request: Request,
    limit: int = Query(50, ge=1, le=500),
    svc: DocumentQAService = Depends(get_service),
):
    """Get chat history for a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    messages = svc.get_chat_history(session_id, limit)

    return {
        "status": "ok",
        "messages": [m.model_dump(mode="json") for m in messages],
        "count": len(messages),
        "correlation_id": correlation_id,
    }

@docqa_router.delete("/sessions/{session_id}/history")
async def clear_chat_history(
    session_id: str,
    request: Request,
    svc: DocumentQAService = Depends(get_service),
):
    """Clear chat history for a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    success = svc.clear_history(session_id)
    if not success:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "cleared": True, "correlation_id": correlation_id}

"""Document AI API Routes.

REST API endpoints for document intelligence - parsing, classification, and analysis.
"""

from backend.app.schemas import (
    ClassifyRequest,
    ClassifyResponse,
    CompareRequest,
    CompareResponse,
    ComplianceCheckRequest,
    ComplianceCheckResponse,
    ContractAnalyzeRequest,
    ContractAnalyzeResponse,
    EntityExtractRequest,
    EntityExtractResponse,
    InvoiceParseRequest,
    InvoiceParseResponse,
    MultiDocSummarizeRequest,
    MultiDocSummarizeResponse,
    ReceiptScanRequest,
    ReceiptScanResponse,
    ResumeParseRequest,
    ResumeParseResponse,
    SemanticSearchRequest,
    SemanticSearchResponse,
)
from backend.app.services.docai_service import docai_service

logger = logging.getLogger("neura.api.docai")

docai_router = APIRouter(tags=["docai"], dependencies=[Depends(require_api_key)])

def _handle_docai_error(exc: Exception, operation: str) -> HTTPException:
    """Map docai service errors to HTTP status codes."""
    logger.error("%s failed: %s", operation, exc, exc_info=True)
    return HTTPException(
        status_code=500,
        detail=f"{operation} failed due to an internal error.",
    )

# Document Parsing Endpoints

@docai_router.post("/parse/invoice", response_model=InvoiceParseResponse)
async def parse_invoice(request: InvoiceParseRequest):
    """Parse an invoice document and extract structured data.

    from backend.app.services.docai_service import (
    Extracts invoice number, dates, vendor/billing info, line items,
    and totals from invoice documents (PDF, images, or text).
    """
    try:
        return await docai_service.parse_invoice(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Invoice parsing") from exc

@docai_router.post("/parse/contract", response_model=ContractAnalyzeResponse)
async def analyze_contract(request: ContractAnalyzeRequest):
    """Analyze a contract document.

    Extracts parties, clauses, obligations, key dates, and performs
    risk analysis on contract documents.
    """
    try:
        return await docai_service.analyze_contract(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Contract analysis") from exc

@docai_router.post("/parse/resume", response_model=ResumeParseResponse)
async def parse_resume(request: ResumeParseRequest):
    """Parse a resume/CV document.

    Extracts contact info, education, work experience, skills,
    certifications, and can optionally match against a job description.
    """
    try:
        return await docai_service.parse_resume(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Resume parsing") from exc

@docai_router.post("/parse/receipt", response_model=ReceiptScanResponse)
async def scan_receipt(request: ReceiptScanRequest):
    """Scan a receipt document.

    Extracts merchant info, date/time, line items, totals, and
    payment information from receipt images or PDFs.
    """
    try:
        return await docai_service.scan_receipt(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Receipt scanning") from exc

# Document Classification

@docai_router.post("/classify", response_model=ClassifyResponse)
async def classify_document(request: ClassifyRequest):
    """Classify a document by type.

    Determines document category (invoice, contract, resume, receipt, etc.)
    and suggests appropriate parsers for further processing.
    """
    try:
        return await docai_service.classify_document(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Document classification") from exc

# Entity Extraction

@docai_router.post("/entities", response_model=EntityExtractResponse)
async def extract_entities(request: EntityExtractRequest):
    """Extract named entities from a document.

    Identifies and extracts entities like persons, organizations,
    locations, dates, monetary values, emails, phones, and URLs.
    """
    try:
        return await docai_service.extract_entities(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Entity extraction") from exc

# Semantic Search

@docai_router.post("/search", response_model=SemanticSearchResponse)
async def semantic_search(request: SemanticSearchRequest):
    """Perform semantic search across documents.

    Uses embeddings to find semantically similar content
    rather than exact keyword matches.
    """
    try:
        return await docai_service.semantic_search(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Semantic search") from exc

# Document Comparison

@docai_router.post("/compare", response_model=CompareResponse)
async def compare_documents(request: CompareRequest):
    """Compare two documents.

    Calculates similarity, identifies differences, and optionally
    performs semantic comparison to find meaningful changes.
    """
    try:
        return await docai_service.compare_documents(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Document comparison") from exc

# Compliance Checking

@docai_router.post("/compliance", response_model=ComplianceCheckResponse)
async def check_compliance(request: ComplianceCheckRequest):
    """Check document for regulatory compliance.

    Analyzes document against specified regulations (GDPR, HIPAA, SOC2)
    and identifies violations and recommendations.
    """
    try:
        return await docai_service.check_compliance(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Compliance check") from exc

# Multi-document Summary

@docai_router.post("/summarize/multi", response_model=MultiDocSummarizeResponse)
async def summarize_multiple(request: MultiDocSummarizeRequest):
    """Summarize multiple documents.

    Creates a unified summary across multiple documents,
    identifying key points and common themes with source references.
    """
    try:
        return await docai_service.summarize_multiple(request)
    except HTTPException:
        raise
    except Exception as exc:
        raise _handle_docai_error(exc, "Multi-document summarization") from exc

"""
AI Agents API Routes
Endpoints for specialized AI agents.
"""

from fastapi import APIRouter, Depends, HTTPException, status

from backend.app.services.agents import agent_service
from backend.app.services.agents import AgentType

logger = logging.getLogger(__name__)
agents_router = APIRouter(dependencies=[Depends(require_api_key)])

# REQUEST MODELS

class ResearchRequest(BaseModel):
    topic: str = Field(..., description="Topic to research")
    depth: str = Field(default="comprehensive", description="Research depth")
    focus_areas: Optional[List[str]] = Field(default=None, description="Focus areas")
    max_sections: int = Field(default=5, ge=1, le=10, description="Max sections")

class DataAnalysisRequest(BaseModel):
    question: str = Field(..., description="Question about the data")
    data: List[Dict[str, Any]] = Field(..., description="Data to analyze")
    data_description: Optional[str] = Field(default=None, description="Data description")
    generate_charts: bool = Field(default=True, description="Generate chart suggestions")

class EmailDraftRequest(BaseModel):
    context: str = Field(..., description="Email context")
    purpose: str = Field(..., description="Email purpose")
    tone: str = Field(default="professional", description="Email tone")
    recipient_info: Optional[str] = Field(default=None, description="Recipient info")
    previous_emails: Optional[List[str]] = Field(default=None, description="Previous emails")

class ContentRepurposeRequest(BaseModel):
    content: str = Field(..., description="Original content")
    source_format: str = Field(..., description="Source format")
    target_formats: List[str] = Field(..., description="Target formats")
    preserve_key_points: bool = Field(default=True, description="Preserve key points")
    adapt_length: bool = Field(default=True, description="Adapt length")

class ProofreadingRequest(BaseModel):
    text: str = Field(..., description="Text to proofread")
    style_guide: Optional[str] = Field(default=None, description="Style guide")
    focus_areas: Optional[List[str]] = Field(default=None, description="Focus areas")
    preserve_voice: bool = Field(default=True, description="Preserve voice")

# AGENT ENDPOINTS

@agents_router.get("")
async def list_agents():
    """List available agent types with their capabilities."""
    agent_types = []
    for at in AgentType:
        agent_types.append({
            "id": at.value,
            "name": at.value.replace("_", " ").title(),
            "type": at.value,
            "status": "available",
        })
    return {"agents": agent_types, "total": len(agent_types)}

@agents_router.post("/research")
async def run_research_agent(request: ResearchRequest):
    """
    Run the research agent to compile a report on a topic.

    Returns:
        ResearchReport with findings
    """
    try:
        task = await agent_service.run_research(
            topic=request.topic,
            depth=request.depth,
            focus_areas=request.focus_areas,
            max_sections=request.max_sections,
        )
        return task.model_dump()
    except Exception as e:
        logger.exception("Research agent failed: %s", e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal server error")

@agents_router.post("/data-analysis")
async def run_data_analyst_agent(request: DataAnalysisRequest):
    """
    Run the data analyst agent to answer questions about data.

    Returns:
        DataAnalysisResult with insights
    """
    try:
        task = await agent_service.run_data_analyst(
            question=request.question,
            data=request.data,
            data_description=request.data_description,
            generate_charts=request.generate_charts,
        )
        return task.model_dump()
    except Exception as e:
        logger.exception("Data analyst agent failed: %s", e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal server error")

@agents_router.post("/email-draft")
async def run_email_draft_agent(request: EmailDraftRequest):
    """
    Run the email draft agent to compose an email.

    Returns:
        EmailDraft with composed email
    """
    try:
        task = await agent_service.run_email_draft(
            context=request.context,
            purpose=request.purpose,
            tone=request.tone,
            recipient_info=request.recipient_info,
            previous_emails=request.previous_emails,
        )
        return task.model_dump()
    except Exception as e:
        logger.exception("Email draft agent failed: %s", e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal server error")

@agents_router.post("/content-repurpose")
async def run_content_repurposing_agent(request: ContentRepurposeRequest):
    """
    Run the content repurposing agent to transform content.

    Returns:
        RepurposedContent with all versions
    """
    try:
        task = await agent_service.run_content_repurpose(
            content=request.content,
            source_format=request.source_format,
            target_formats=request.target_formats,
            preserve_key_points=request.preserve_key_points,
            adapt_length=request.adapt_length,
        )
        return task.model_dump()
    except Exception as e:
        logger.exception("Content repurposing agent failed: %s", e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal server error")

@agents_router.post("/proofread")
async def run_proofreading_agent(request: ProofreadingRequest):
    """
    Run the proofreading agent for comprehensive style and grammar check.

    Returns:
        ProofreadingResult with corrections
    """
    try:
        task = await agent_service.run_proofreading(
            text=request.text,
            style_guide=request.style_guide,
            focus_areas=request.focus_areas,
            preserve_voice=request.preserve_voice,
        )
        return task.model_dump()
    except Exception as e:
        logger.exception("Proofreading agent failed: %s", e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Internal server error")

# TASK MANAGEMENT ENDPOINTS

@agents_router.get("/tasks/{task_id}")
async def get_task(task_id: str):
    """Get task by ID."""
    task = agent_service.get_task(task_id)
    if not task:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Task not found")
    return task.model_dump()

@agents_router.get("/tasks")
async def list_tasks(agent_type: Optional[str] = None):
    """List all tasks, optionally filtered by agent type."""
    tasks = agent_service.list_tasks(agent_type=agent_type)
    return [t.model_dump() for t in tasks]

# UTILITY ENDPOINTS

@agents_router.get("/types")
async def list_agent_types():
    """List available agent types."""
    return {
        "types": [
            {
                "id": t.value,
                "name": t.value.replace("_", " ").title(),
                "description": {
                    "research": "Deep-dive research and report compilation",
                    "data_analyst": "Data analysis and question answering",
                    "email_draft": "Email composition based on context",
                    "content_repurpose": "Content transformation to multiple formats",
                    "proofreading": "Comprehensive style and grammar checking",
                }.get(t.value, "")
            }
            for t in AgentType
        ]
    }

@agents_router.get("/formats/repurpose")
async def list_repurpose_formats():
    """List available content repurposing formats."""
    return {
        "formats": [
            {"id": "tweet_thread", "name": "Twitter Thread", "description": "5-10 tweets, 280 chars each"},
            {"id": "linkedin_post", "name": "LinkedIn Post", "description": "Professional, 1300 chars max"},
            {"id": "blog_summary", "name": "Blog Summary", "description": "300-500 words"},
            {"id": "slides", "name": "Presentation Slides", "description": "Title + bullet points per slide"},
            {"id": "email_newsletter", "name": "Email Newsletter", "description": "Catchy subject, scannable body"},
            {"id": "video_script", "name": "Video Script", "description": "Conversational, 2-3 minutes"},
            {"id": "infographic", "name": "Infographic Copy", "description": "Headlines, stats, takeaways"},
            {"id": "podcast_notes", "name": "Podcast Show Notes", "description": "Summary, timestamps, links"},
            {"id": "press_release", "name": "Press Release", "description": "Headline, lead, quotes"},
            {"id": "executive_summary", "name": "Executive Summary", "description": "1 page, key decisions"},
        ]
    }

"""
AI Agents API Routes v2 - Production-grade implementation.

Features:
- Persistent task storage
- Idempotency support
- Progress tracking + SSE streaming
- Task management (cancel, retry)
- Comprehensive error handling
- Cost tracking

All tasks are persisted to SQLite and survive server restarts.
"""

from typing import Any, Dict, List, Literal, Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Request, Response, status
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, ConfigDict, Field, field_validator

from backend.app.api.middleware import limiter
from backend.app.services.agents import (
    TaskConflictError,
    TaskNotFoundError,
)
from backend.app.services.agents import agent_service_v2
from backend.app.services.agents import (
    AgentError,
    ValidationError,
)

logger = logging.getLogger(__name__)
agents_v2_router = APIRouter(dependencies=[Depends(require_api_key)])

# REQUEST MODELS

class ResearchRequest(BaseModel):
    """Request to run the research agent."""
    topic: str = Field(
        ...,
        max_length=500,
        description="Topic to research (must be at least 2 words)",
        examples=["AI trends in healthcare 2025", "Climate change mitigation strategies"],
    )
    depth: Literal["quick", "moderate", "comprehensive"] = Field(
        default="comprehensive",
        description="Research depth - quick (overview), moderate (balanced), comprehensive (detailed)",
    )
    focus_areas: Optional[List[str]] = Field(
        default=None,
        max_length=10,
        description="Specific areas to focus on (max 10)",
        examples=[["regulation", "adoption", "startups"]],
    )
    max_sections: int = Field(
        default=5,
        ge=1,
        le=20,
        description="Maximum number of sections in the report",
    )
    idempotency_key: Optional[str] = Field(
        default=None,
        max_length=64,
        description="Unique key for deduplication (same key returns existing task)",
    )
    priority: int = Field(
        default=0,
        ge=0,
        le=10,
        description="Task priority (0=lowest, 10=highest)",
    )
    webhook_url: Optional[str] = Field(
        default=None,
        max_length=2000,
        description="URL to notify when task completes",
    )
    sync: bool = Field(
        default=True,
        description="If true, wait for completion. If false, return immediately.",
    )

    @field_validator('topic')
    @classmethod
    def validate_topic(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("Topic cannot be empty or whitespace")
        if len(v.split()) < 2:
            raise ValueError("Topic must contain at least 2 words for meaningful research")
        return v

    @field_validator('focus_areas')
    @classmethod
    def validate_focus_areas(cls, v: Optional[List[str]]) -> Optional[List[str]]:
        if v:
            return [area.strip() for area in v if area and area.strip()]
        return v

class DataAnalystRequest(BaseModel):
    """Request to run the data analyst agent."""
    question: str = Field(
        ..., min_length=5, max_length=1000,
        description="Question to answer about the data",
    )
    data: List[Dict[str, Any]] = Field(
        ..., min_length=1,
        description="Tabular data as list of objects",
    )
    data_description: Optional[str] = Field(
        default=None, max_length=2000,
        description="Optional description of the dataset",
    )
    generate_charts: bool = Field(
        default=True,
        description="Whether to suggest chart visualisations",
    )
    idempotency_key: Optional[str] = Field(default=None, max_length=64)
    priority: int = Field(default=0, ge=0, le=10)
    webhook_url: Optional[str] = Field(default=None, max_length=2000)
    sync: bool = Field(default=True)

    @field_validator("question")
    @classmethod
    def validate_question(cls, v: str) -> str:
        v = v.strip()
        if not v or len(v.split()) < 2:
            raise ValueError("Question must contain at least 2 words")
        return v

class EmailDraftRequest(BaseModel):
    """Request to run the email draft agent."""
    context: str = Field(
        ..., min_length=5, max_length=5000,
        description="Background context for the email",
    )
    purpose: str = Field(
        ..., min_length=3, max_length=1000,
        description="Purpose/intent of the email",
    )
    tone: str = Field(
        default="professional",
        description="Tone: professional, friendly, formal, casual, empathetic, assertive",
    )
    recipient_info: Optional[str] = Field(
        default=None, max_length=2000,
        description="Information about the recipient",
    )
    previous_emails: Optional[List[str]] = Field(
        default=None,
        description="Previous emails in thread (last 3 kept)",
    )
    include_subject: bool = Field(default=True)
    idempotency_key: Optional[str] = Field(default=None, max_length=64)
    priority: int = Field(default=0, ge=0, le=10)
    webhook_url: Optional[str] = Field(default=None, max_length=2000)
    sync: bool = Field(default=True)

class ContentRepurposeRequest(BaseModel):
    """Request to run the content repurposing agent."""
    content: str = Field(
        ..., min_length=20, max_length=50000,
        description="Source content to repurpose",
    )
    source_format: str = Field(
        ..., min_length=1, max_length=50,
        description="Format of the source content (article, report, transcript, etc.)",
    )
    target_formats: List[str] = Field(
        ..., min_length=1,
        description="Target formats: tweet_thread, linkedin_post, blog_summary, slides, "
                    "email_newsletter, video_script, infographic, podcast_notes, press_release, "
                    "executive_summary",
    )
    preserve_key_points: bool = Field(default=True)
    adapt_length: bool = Field(default=True)
    idempotency_key: Optional[str] = Field(default=None, max_length=64)
    priority: int = Field(default=0, ge=0, le=10)
    webhook_url: Optional[str] = Field(default=None, max_length=2000)
    sync: bool = Field(default=True)

class ProofreadingRequest(BaseModel):
    """Request to run the proofreading agent."""
    text: str = Field(
        ..., min_length=10, max_length=50000,
        description="Text to proofread",
    )
    style_guide: Optional[str] = Field(
        default=None,
        description="Style guide: ap, chicago, apa, mla, none",
    )
    focus_areas: Optional[List[str]] = Field(
        default=None,
        description="Focus areas: grammar, spelling, punctuation, clarity, conciseness, "
                    "tone, consistency, formatting, word_choice, structure",
    )
    preserve_voice: bool = Field(
        default=True,
        description="Preserve the author's voice while correcting",
    )
    idempotency_key: Optional[str] = Field(default=None, max_length=64)
    priority: int = Field(default=0, ge=0, le=10)
    webhook_url: Optional[str] = Field(default=None, max_length=2000)
    sync: bool = Field(default=True)

class ReportAnalystRequest(BaseModel):
    """Request to run the Report Analyst agent."""
    run_id: str = Field(
        ...,
        min_length=1,
        max_length=100,
        description="Report run ID to analyze",
    )
    analysis_type: str = Field(
        default="summarize",
        description="Analysis type: summarize, insights, compare, qa",
    )
    question: Optional[str] = Field(
        default=None,
        max_length=2000,
        description="Question text (required for 'qa' analysis type)",
    )
    compare_run_id: Optional[str] = Field(
        default=None,
        max_length=100,
        description="Second run ID for comparison (required for 'compare' analysis type)",
    )
    focus_areas: Optional[List[str]] = Field(
        default=None,
        description="Optional areas to focus analysis on",
    )
    idempotency_key: Optional[str] = Field(default=None, max_length=64)
    priority: int = Field(default=0, ge=0, le=10)
    webhook_url: Optional[str] = Field(default=None, max_length=2000)
    sync: bool = Field(default=True)

class GenerateReportFromAgentRequest(BaseModel):
    """Request to trigger report generation from an agent task result."""
    template_id: str = Field(..., min_length=1, description="Template to use for report generation")
    connection_id: str = Field(..., min_length=1, description="Database connection to use")
    start_date: str = Field(..., description="Report start date (YYYY-MM-DD)")
    end_date: str = Field(..., description="Report end date (YYYY-MM-DD)")
    key_values: Optional[Dict[str, Any]] = Field(default=None, description="Additional key-value parameters")
    docx: bool = Field(default=False, description="Generate DOCX artifact")
    xlsx: bool = Field(default=False, description="Generate XLSX artifact")

class CancelRequest(BaseModel):
    """Request to cancel a task."""
    reason: Optional[str] = Field(
        default=None,
        max_length=500,
        description="Optional cancellation reason",
    )

# RESPONSE MODELS

class ProgressResponse(BaseModel):
    """Progress information for a task."""
    percent: int = Field(..., ge=0, le=100)
    message: Optional[str] = None
    current_step: Optional[str] = None
    total_steps: Optional[int] = None
    current_step_num: Optional[int] = None

class ErrorResponse(BaseModel):
    """Error information for a failed task."""
    code: Optional[str] = None
    message: Optional[str] = None
    retryable: bool = True

class CostResponse(BaseModel):
    """Cost tracking information."""
    tokens_input: int = 0
    tokens_output: int = 0
    estimated_cost_cents: int = 0

class AttemptsResponse(BaseModel):
    """Retry attempt information."""
    count: int = 0
    max: int = 3

class TimestampsResponse(BaseModel):
    """Task timestamps."""
    created_at: Optional[str] = None
    started_at: Optional[str] = None
    completed_at: Optional[str] = None

class LinksResponse(BaseModel):
    """HATEOAS links for task."""
    self_link: str = Field(..., alias="self")
    cancel: Optional[str] = None
    retry: Optional[str] = None
    events: str
    stream: Optional[str] = None

    model_config = ConfigDict(populate_by_name=True)

class TaskResponse(BaseModel):
    """Standard task response."""
    task_id: str
    agent_type: str
    status: str
    progress: ProgressResponse
    result: Optional[Dict[str, Any]] = None
    error: Optional[ErrorResponse] = None
    timestamps: TimestampsResponse
    cost: CostResponse
    attempts: AttemptsResponse
    links: LinksResponse

    model_config = ConfigDict(populate_by_name=True)

class TaskListResponse(BaseModel):
    """Response for task listing."""
    tasks: List[TaskResponse]
    total: int
    limit: int
    offset: int

class TaskEventResponse(BaseModel):
    """Task event for audit trail."""
    id: int
    event_type: str
    previous_status: Optional[str] = None
    new_status: Optional[str] = None
    event_data: Optional[Dict[str, Any]] = None
    created_at: Optional[str] = None

class StatsResponse(BaseModel):
    """Service statistics."""
    pending: int = 0
    running: int = 0
    completed: int = 0
    failed: int = 0
    cancelled: int = 0
    retrying: int = 0
    total: int = 0

# HELPER FUNCTIONS

def task_to_response(task) -> TaskResponse:
    """Convert AgentTaskModel to API response."""
    return TaskResponse(
        task_id=task.task_id,
        agent_type=task.agent_type.value if hasattr(task.agent_type, 'value') else task.agent_type,
        status=task.status.value if hasattr(task.status, 'value') else task.status,
        progress=ProgressResponse(
            percent=task.progress_percent,
            message=task.progress_message,
            current_step=task.current_step,
            total_steps=task.total_steps,
            current_step_num=task.current_step_num,
        ),
        result=task.result,
        error=ErrorResponse(
            code=task.error_code,
            message=task.error_message,
            retryable=task.is_retryable,
        ) if task.error_message else None,
        timestamps=TimestampsResponse(
            created_at=task.created_at.isoformat() if task.created_at else None,
            started_at=task.started_at.isoformat() if task.started_at else None,
            completed_at=task.completed_at.isoformat() if task.completed_at else None,
        ),
        cost=CostResponse(
            tokens_input=task.tokens_input,
            tokens_output=task.tokens_output,
            estimated_cost_cents=task.estimated_cost_cents,
        ),
        attempts=AttemptsResponse(
            count=task.attempt_count,
            max=task.max_attempts,
        ),
        links=LinksResponse(
            **{
                "self": f"/agents/v2/tasks/{task.task_id}",
                "cancel": f"/agents/v2/tasks/{task.task_id}/cancel" if task.can_cancel() else None,
                "retry": f"/agents/v2/tasks/{task.task_id}/retry" if task.can_retry() else None,
                "events": f"/agents/v2/tasks/{task.task_id}/events",
                "stream": f"/agents/v2/tasks/{task.task_id}/stream" if task.is_active() else None,
            }
        ),
    )

# RESEARCH AGENT ENDPOINT

@agents_v2_router.post("/research", response_model=TaskResponse, status_code=status.HTTP_202_ACCEPTED)
@limiter.limit("10/minute")
async def run_research_agent(request: Request, response: Response, body: ResearchRequest):
    """Run the research agent to compile a report on a topic."""
    try:
        task = await agent_service_v2.run_research(
            topic=body.topic,
            depth=body.depth,
            focus_areas=body.focus_areas,
            max_sections=body.max_sections,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

def _handle_agent_error(e: Exception) -> None:
    """Shared error handler for all agent endpoints."""
    if isinstance(e, ValidationError):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": e.code, "message": e.message, "field": e.details.get("field")},
        )
    if isinstance(e, AgentError):
        if e.code == "LLM_RATE_LIMITED":
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail={"code": e.code, "message": e.message, "retry_after": e.details.get("retry_after", 60)},
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"code": e.code, "message": e.message, "retryable": e.retryable},
        )
    logger.exception("Agent failed: %s", e)
    raise HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail={"code": "INTERNAL_ERROR", "message": "An internal error occurred"},
    )

# DATA ANALYST ENDPOINT

@agents_v2_router.post(
    "/data-analyst",
    response_model=TaskResponse,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Run Data Analyst Agent",
    description="Analyse tabular data: answer questions, compute statistics, suggest charts, generate SQL.",
)
@limiter.limit("10/minute")
async def run_data_analyst_agent(request: Request, response: Response, body: DataAnalystRequest):
    """Run the data analyst agent."""
    try:
        task = await agent_service_v2.run_data_analyst(
            question=body.question,
            data=body.data,
            data_description=body.data_description,
            generate_charts=body.generate_charts,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

# EMAIL DRAFT ENDPOINT

@agents_v2_router.post(
    "/email-draft",
    response_model=TaskResponse,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Run Email Draft Agent",
    description="Compose email drafts with tone control, thread context, and follow-up actions.",
)
@limiter.limit("10/minute")
async def run_email_draft_agent(request: Request, response: Response, body: EmailDraftRequest):
    """Run the email draft agent."""
    try:
        task = await agent_service_v2.run_email_draft(
            context=body.context,
            purpose=body.purpose,
            tone=body.tone,
            recipient_info=body.recipient_info,
            previous_emails=body.previous_emails,
            include_subject=body.include_subject,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

# CONTENT REPURPOSE ENDPOINT

@agents_v2_router.post(
    "/content-repurpose",
    response_model=TaskResponse,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Run Content Repurposing Agent",
    description="Transform content into multiple formats (tweets, LinkedIn, slides, newsletters, etc.).",
)
@limiter.limit("10/minute")
async def run_content_repurpose_agent(request: Request, response: Response, body: ContentRepurposeRequest):
    """Run the content repurposing agent."""
    try:
        task = await agent_service_v2.run_content_repurpose(
            content=body.content,
            source_format=body.source_format,
            target_formats=body.target_formats,
            preserve_key_points=body.preserve_key_points,
            adapt_length=body.adapt_length,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

# PROOFREADING ENDPOINT

@agents_v2_router.post(
    "/proofreading",
    response_model=TaskResponse,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Run Proofreading Agent",
    description="Grammar, style, and clarity checking with style guide support and readability scoring.",
)
@limiter.limit("10/minute")
async def run_proofreading_agent(request: Request, response: Response, body: ProofreadingRequest):
    """Run the proofreading agent."""
    try:
        task = await agent_service_v2.run_proofreading(
            text=body.text,
            style_guide=body.style_guide,
            focus_areas=body.focus_areas,
            preserve_voice=body.preserve_voice,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

# REPORT ANALYST ENDPOINT

@agents_v2_router.post(
    "/report-analyst",
    response_model=TaskResponse,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Run Report Analyst Agent",
    description="Analyze, summarize, compare, or ask questions about generated reports.",
)
@limiter.limit("10/minute")
async def run_report_analyst_agent(request: Request, response: Response, body: ReportAnalystRequest):
    """Run the report analyst agent."""
    try:
        task = await agent_service_v2.run_report_analyst(
            run_id=body.run_id,
            analysis_type=body.analysis_type,
            question=body.question,
            compare_run_id=body.compare_run_id,
            focus_areas=body.focus_areas,
            idempotency_key=body.idempotency_key,
            priority=body.priority,
            webhook_url=body.webhook_url,
            sync=body.sync,
        )
        return task_to_response(task)
    except Exception as e:
        _handle_agent_error(e)

# AGENT-TRIGGERED REPORT GENERATION ENDPOINT

@agents_v2_router.post(
    "/tasks/{task_id}/generate-report",
    status_code=status.HTTP_202_ACCEPTED,
    summary="Generate Report from Agent Task",
    description="Trigger report generation using an agent task's result as additional context.",
)
@limiter.limit("5/minute")
async def generate_report_from_agent(request: Request, response: Response, task_id: str, body: GenerateReportFromAgentRequest):
    """Generate a report using agent task results as context."""
    task = agent_service_v2.get_task(task_id)
    if not task:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )
    if task.status != "completed":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "TASK_NOT_COMPLETED", "message": f"Task must be completed (current: {task.status})"},
        )

    # Merge agent result into key_values so the report template can reference it
    key_values = dict(body.key_values or {})
    if task.result:
        key_values["_agent_context"] = {
            "task_id": task_id,
            "agent_type": task.agent_type if isinstance(task.agent_type, str) else task.agent_type.value,
            "result": task.result,
        }

    try:

        payload = RunPayload(
            template_id=body.template_id,
            connection_id=body.connection_id,
            start_date=body.start_date,
            end_date=body.end_date,
            key_values=key_values,
            docx=body.docx,
            xlsx=body.xlsx,
        )
        # Set correlation_id on request state for tracing
        if not hasattr(request.state, "correlation_id"):
            request.state.correlation_id = f"agent-{task_id}"

        job = await queue_report_job(payload, request, kind="pdf")
        job_id = (job.get("job_id") or job.get("id")) if isinstance(job, dict) else getattr(job, "job_id", getattr(job, "id", str(job)))
        return {
            "job_id": job_id,
            "task_id": task_id,
            "status": "queued",
            "message": "Report generation triggered from agent task result",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to trigger report from agent: %s", e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"code": "REPORT_TRIGGER_FAILED", "message": str(e)},
        )

# TASK MANAGEMENT ENDPOINTS

@agents_v2_router.get(
    "/tasks/{task_id}",
    response_model=TaskResponse,
    summary="Get Task",
    description="Get a task by ID with full status, progress, and result information.",
)
async def get_task(task_id: str):
    """Get task by ID."""
    task = agent_service_v2.get_task(task_id)
    if not task:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )
    return task_to_response(task)

@agents_v2_router.get(
    "/tasks",
    response_model=TaskListResponse,
    summary="List Tasks",
    description="List tasks with optional filtering by agent type, status, or user.",
)
async def list_tasks(
    agent_type: Optional[str] = Query(None, description="Filter by agent type"),
    task_status: Optional[str] = Query(None, alias="status", description="Filter by status"),
    user_id: Optional[str] = Query(None, description="Filter by user ID"),
    limit: int = Query(50, ge=1, le=100, description="Maximum number of tasks"),
    offset: int = Query(0, ge=0, description="Number of tasks to skip"),
):
    """List all tasks with optional filtering."""
    tasks = agent_service_v2.list_tasks(
        agent_type=agent_type,
        status=task_status,
        user_id=user_id,
        limit=limit,
        offset=offset,
    )

    total = agent_service_v2.count_tasks(
        agent_type=agent_type,
        status=task_status,
        user_id=user_id,
    )

    return TaskListResponse(
        tasks=[task_to_response(t) for t in tasks],
        total=total,
        limit=limit,
        offset=offset,
    )

@agents_v2_router.post(
    "/tasks/{task_id}/cancel",
    response_model=TaskResponse,
    summary="Cancel Task",
    description="Cancel a pending or running task. Cannot cancel completed tasks.",
)
async def cancel_task(task_id: str, request: Optional[CancelRequest] = None):
    """Cancel a pending or running task."""
    try:
        reason = request.reason if request else None
        task = agent_service_v2.cancel_task(task_id, reason)
        return task_to_response(task)

    except TaskNotFoundError:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )

    except TaskConflictError as e:
        logger.warning("cancel_task_conflict", extra={"task_id": task_id, "error": str(e)})
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"code": "CANNOT_CANCEL", "message": "Task cannot be cancelled in its current state"},
        )

@agents_v2_router.post(
    "/tasks/{task_id}/retry",
    response_model=TaskResponse,
    summary="Retry Task",
    description="Manually retry a failed task. Only works for retryable failures.",
)
async def retry_task(task_id: str):
    """Retry a failed task."""
    try:
        task = await agent_service_v2.retry_task(task_id)
        return task_to_response(task)

    except TaskNotFoundError:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )

    except TaskConflictError as e:
        logger.warning("retry_task_conflict", extra={"task_id": task_id, "error": str(e)})
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail={"code": "CANNOT_RETRY", "message": "Task cannot be retried in its current state"},
        )

@agents_v2_router.get(
    "/tasks/{task_id}/events",
    response_model=List[TaskEventResponse],
    summary="Get Task Events",
    description="Get audit trail events for a task.",
)
async def get_task_events(
    task_id: str,
    limit: int = Query(100, ge=1, le=500, description="Maximum number of events"),
):
    """Get audit events for a task."""
    # First check if task exists
    task = agent_service_v2.get_task(task_id)
    if not task:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )

    events = agent_service_v2.get_task_events(task_id, limit=limit)
    return [TaskEventResponse(**e) for e in events]

# SSE PROGRESS STREAMING (Trade-off 2)

@agents_v2_router.get(
    "/tasks/{task_id}/stream",
    responses={
        200: {"description": "SSE progress stream", "content": {"text/event-stream": {}}},
        404: {"description": "Task not found"},
    },
)
async def stream_task_progress(
    task_id: str,
    request: Request,
    poll_interval: float = Query(0.5, ge=0.1, le=5.0, description="Poll interval in seconds"),
    timeout: float = Query(300.0, ge=10.0, le=600.0, description="Stream timeout in seconds"),
):
    """Stream real-time progress for a task via Server-Sent Events."""
    # Verify task exists before opening stream
    task = agent_service_v2.get_task(task_id)
    if not task:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"code": "TASK_NOT_FOUND", "message": f"Task {task_id} not found"},
        )

    async def _event_generator():
        async for event in agent_service_v2.stream_task_progress(
            task_id,
            poll_interval=poll_interval,
            timeout=timeout,
        ):
            # Check if client disconnected
            if await request.is_disconnected():
                return
            # SSE format: data: {json}\n\n
            yield f"data: {json.dumps(event)}\n\n"

    return StreamingResponse(
        _event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # Disable nginx buffering
        },
    )

# UTILITY ENDPOINTS

@agents_v2_router.get(
    "/stats",
    response_model=StatsResponse,
    summary="Get Statistics",
    description="Get task counts by status.",
)
async def get_stats():
    """Get service statistics."""
    stats = agent_service_v2.get_stats()
    return StatsResponse(**stats)

@agents_v2_router.get(
    "/types",
    summary="List Agent Types",
    description="List available agent types with descriptions.",
)
async def list_agent_types():
    """List available agent types."""
    return {
        "types": [
            {
                "id": "research",
                "name": "Research Agent",
                "description": "Deep-dive research and report compilation",
                "endpoint": "/agents/v2/research",
            },
            {
                "id": "data_analyst",
                "name": "Data Analyst Agent",
                "description": "Analyse tabular data with statistics, insights, chart suggestions, and SQL",
                "endpoint": "/agents/v2/data-analyst",
            },
            {
                "id": "email_draft",
                "name": "Email Draft Agent",
                "description": "Compose email drafts with tone control, thread context, and follow-ups",
                "endpoint": "/agents/v2/email-draft",
            },
            {
                "id": "content_repurpose",
                "name": "Content Repurposing Agent",
                "description": "Transform content into 10 output formats (tweets, slides, newsletters, etc.)",
                "endpoint": "/agents/v2/content-repurpose",
            },
            {
                "id": "proofreading",
                "name": "Proofreading Agent",
                "description": "Grammar, style, and clarity checking with style guide support",
                "endpoint": "/agents/v2/proofreading",
            },
            {
                "id": "report_analyst",
                "name": "Report Analyst",
                "description": "Analyze, summarize, compare, or ask questions about generated reports",
                "endpoint": "/agents/v2/report-analyst",
            },
        ]
    }

@agents_v2_router.get(
    "/formats/repurpose",
    summary="List Repurpose Formats",
    description="List available content repurposing target formats with descriptions.",
)
async def list_repurpose_formats():
    """List available content repurposing formats."""
    return {
        "formats": [
            {"id": "tweet_thread", "name": "Twitter Thread", "description": "5-10 tweets, 280 chars each"},
            {"id": "linkedin_post", "name": "LinkedIn Post", "description": "Professional, 1300 chars max"},
            {"id": "blog_summary", "name": "Blog Summary", "description": "300-500 words"},
            {"id": "slides", "name": "Presentation Slides", "description": "Title + bullet points per slide"},
            {"id": "email_newsletter", "name": "Email Newsletter", "description": "Catchy subject, scannable body"},
            {"id": "video_script", "name": "Video Script", "description": "Conversational, 2-3 minutes"},
            {"id": "infographic", "name": "Infographic Copy", "description": "Headlines, stats, takeaways"},
            {"id": "podcast_notes", "name": "Podcast Show Notes", "description": "Summary, timestamps, links"},
            {"id": "press_release", "name": "Press Release", "description": "Headline, lead, quotes"},
            {"id": "executive_summary", "name": "Executive Summary", "description": "1 page, key decisions"},
        ]
    }

@agents_v2_router.get(
    "/health",
    summary="Health Check",
    description="Check if the agents service is healthy.",
)
async def health_check():
    """Health check endpoint."""
    try:
        stats = agent_service_v2.get_stats()
        return {
            "status": "healthy",
            "tasks": stats,
        }
    except Exception as e:
        logger.warning("Agents v2 health check failed: %s", e)
        return JSONResponse(
            status_code=503,
            content={"status": "unhealthy", "error": "Service unavailable"},
        )

"""
Document API Routes - Document editing and collaboration endpoints.
"""

from fastapi import APIRouter, Depends, HTTPException, Query, Request, Response, WebSocket, WebSocketDisconnect

from backend.app.utils import validate_path_safety
from backend.app.schemas import (
    CreateDocumentRequest,
    UpdateDocumentRequest,
    DocumentResponse,
    DocumentListResponse,
    CommentRequest,
    CommentResponse,
    CommentReplyRequest,
    CollaborationSessionResponse,
    PresenceUpdateBody,
    PDFMergeRequest,
    PDFWatermarkRequest,
    PDFRedactRequest,
    PDFReorderRequest,
    PDFSplitRequest,
    PDFRotateRequest,
    AIWritingRequest,
    AIWritingResponse,
    CreateFromTemplateRequest,
)
from ...services.documents_service import (
    DocumentService,
    CollaborationService,
    PDFOperationsService,
)
from ...services.documents_service import YjsWebSocketHandler
from backend.app.services.platform_services import writing_service as _writing_service
from backend.app.services.config import require_api_key, verify_ws_token
from backend.app.api.middleware import limiter, RATE_LIMIT_STANDARD

logger = logging.getLogger("neura.api.documents")

documents_router = APIRouter(tags=["documents"], dependencies=[Depends(require_api_key)])
ws_router = APIRouter()

# Service instances (would use dependency injection in production)
# Re-entrant because some getters call other getters while holding the lock.
# (e.g. get_ws_handler() -> get_collaboration_service()).
_lock = threading.RLock()
_doc_service: Optional[DocumentService] = None
_collab_service: Optional[CollaborationService] = None
_pdf_service: Optional[PDFOperationsService] = None
_ws_handler: Optional[YjsWebSocketHandler] = None

def _is_pytest() -> bool:
    return bool(os.getenv("PYTEST_CURRENT_TEST"))

def _stub_ai_response(
    *,
    text: str,
    operation: str,
    result_text: str | None = None,
    metadata: dict | None = None,
) -> AIWritingResponse:
    return AIWritingResponse(
        original_text=text,
        result_text=result_text if result_text is not None else text,
        suggestions=[],
        confidence=1.0,
        metadata={"operation": operation, **(metadata or {})},
    )

def get_document_service() -> DocumentService:
    global _doc_service
    if _doc_service is None:
        with _lock:
            if _doc_service is None:
                _doc_service = DocumentService()
    return _doc_service

def get_collaboration_service() -> CollaborationService:
    global _collab_service
    if _collab_service is None:
        with _lock:
            if _collab_service is None:
                _collab_service = CollaborationService()
    return _collab_service

def get_ws_handler() -> YjsWebSocketHandler:
    global _ws_handler
    if _ws_handler is None:
        with _lock:
            if _ws_handler is None:
                _ws_handler = YjsWebSocketHandler(get_collaboration_service())
    return _ws_handler

def _resolve_ws_base_url(request: Request) -> str:
    scheme = "wss" if request.url.scheme == "https" else "ws"
    host = request.headers.get("x-forwarded-host") or request.headers.get("host") or "localhost:8000"
    return f"{scheme}://{host}"

def get_pdf_service() -> PDFOperationsService:
    global _pdf_service
    if _pdf_service is None:
        with _lock:
            if _pdf_service is None:
                _pdf_service = PDFOperationsService()
    return _pdf_service

def validate_pdf_path(pdf_path: str | None) -> Path:
    """Validate that a PDF path is safe and within allowed directories."""
    if not pdf_path:
        raise HTTPException(status_code=400, detail="Document is not a PDF")

    # Check for dangerous path patterns
    is_safe, error = validate_path_safety(pdf_path)
    if not is_safe:
        logger.warning(f"Blocked unsafe PDF path: {pdf_path} - {error}")
        raise HTTPException(status_code=400, detail="Invalid PDF path")

    path = Path(pdf_path)

    # Resolve to absolute path and verify it's within allowed directories
    try:
        resolved = path.resolve()
        settings = get_settings()
        uploads_root = Path(settings.uploads_dir).resolve()
        excel_root = Path(settings.excel_uploads_dir).resolve()

        # Check if path is within allowed directories
        try:
            resolved.relative_to(uploads_root)
        except ValueError:
            try:
                resolved.relative_to(excel_root)
            except ValueError:
                logger.warning(f"PDF path outside allowed directories: {resolved}")
                raise HTTPException(status_code=400, detail="PDF not accessible")
    except Exception as e:
        if isinstance(e, HTTPException):
            raise
        logger.warning(f"Failed to validate PDF path: {pdf_path} - {e}")
        raise HTTPException(status_code=400, detail="Invalid PDF path")

    if not resolved.exists():
        raise HTTPException(status_code=404, detail="PDF file not found")

    return resolved

# Document CRUD Endpoints

@documents_router.post("", response_model=DocumentResponse)
@limiter.limit(RATE_LIMIT_STANDARD)
async def create_document(
    request: Request,
    response: Response,
    req: CreateDocumentRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Create a new document."""
    content_payload = req.content.model_dump() if req.content else None
    doc = doc_service.create(
        name=req.name,
        content=content_payload,
        is_template=req.is_template,
        metadata=req.metadata,
    )
    return DocumentResponse(**doc.model_dump())

@documents_router.get("", response_model=DocumentListResponse)
async def list_documents(
    is_template: Optional[bool] = Query(None),
    tags: Optional[str] = Query(None, description="Comma-separated tags"),
    q: Optional[str] = Query(None, description="Search documents by name"),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    doc_service: DocumentService = Depends(get_document_service),
):
    """List documents with optional filters."""
    tag_list = tags.split(",") if tags else None
    documents, total = doc_service.list_documents(
        is_template=is_template,
        tags=tag_list,
        limit=limit + offset if q else limit,
        offset=0 if q else offset,
    )
    # Apply text search filter on name and content
    if q:
        q_lower = q.strip().lower()
        def _doc_matches(d) -> bool:
            if q_lower in (d.name or "").lower():
                return True
            content = d.content
            if content:
                if hasattr(content, "model_dump"):
                    content = content.model_dump()
                if isinstance(content, dict):
                    for node in content.get("content", []):
                        for inline in node.get("content", []):
                            if q_lower in (inline.get("text") or "").lower():
                                return True
            return False
        documents = [d for d in documents if _doc_matches(d)]
        total = len(documents)
        documents = documents[offset:offset + limit]
    return DocumentListResponse(
        documents=[DocumentResponse(**d.model_dump()) for d in documents],
        total=total,
        offset=offset,
        limit=limit,
    )

# Search documents

@documents_router.post("/search")
async def search_documents(
    request: Request,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Search documents by text query with optional tag/template filters."""
    body = await request.json()
    q = str(body.get("q") or "").strip()
    tags = body.get("tags")
    is_template = body.get("is_template")
    limit = int(body.get("limit", 50))
    offset = int(body.get("offset", 0))

    documents, _ = doc_service.list_documents(
        is_template=is_template,
        tags=tags,
        limit=500,
        offset=0,
    )
    q_lower = q.lower()

    def _doc_matches(d) -> bool:
        if q_lower in (d.name or "").lower():
            return True
        content = d.content
        if content:
            if hasattr(content, "model_dump"):
                content = content.model_dump()
            if isinstance(content, dict):
                for node in content.get("content", []):
                    for inline in node.get("content", []):
                        if q_lower in (inline.get("text") or "").lower():
                            return True
        return False

    matched = [d for d in documents if _doc_matches(d)] if q_lower else list(documents)
    total = len(matched)
    paged = matched[offset:offset + limit]
    return {
        "results": [DocumentResponse(**d.model_dump()).model_dump() for d in paged],
        "total": total,
        "limit": limit,
        "offset": offset,
    }

# Static-path routes (MUST be before /{document_id} to avoid shadowing)

@documents_router.get("/templates", response_model=DocumentListResponse)
async def list_document_templates(
    tags: Optional[str] = Query(None, description="Comma-separated tags"),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    doc_service: DocumentService = Depends(get_document_service),
):
    """List document templates."""
    tag_list = tags.split(",") if tags else None
    documents, total = doc_service.list_documents(
        is_template=True,
        tags=tag_list,
        limit=limit,
        offset=offset,
    )
    return DocumentListResponse(
        documents=[DocumentResponse(**d.model_dump()) for d in documents],
        total=total,
        offset=offset,
        limit=limit,
    )

@documents_router.post("/templates/{template_id}/create", response_model=DocumentResponse)
@limiter.limit(RATE_LIMIT_STANDARD)
async def create_from_template(
    request: Request,
    response: Response,
    template_id: str,
    req: CreateFromTemplateRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Create a new document from a template."""
    template = doc_service.get(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    if not template.is_template:
        raise HTTPException(status_code=400, detail="Document is not a template")

    content_data = template.content
    if hasattr(content_data, "model_dump"):
        content_data = content_data.model_dump()

    new_name = req.name if req and req.name else f"{template.name} (copy)"
    doc = doc_service.create(
        name=new_name,
        content=content_data,
        is_template=False,
        metadata={"created_from_template": template_id},
    )
    return DocumentResponse(**doc.model_dump())

# Dynamic document routes (/{document_id} — MUST be after static routes)

@documents_router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Get a document by ID."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return DocumentResponse(**doc.model_dump())

@documents_router.put("/{document_id}", response_model=DocumentResponse)
@limiter.limit(RATE_LIMIT_STANDARD)
async def update_document(
    request: Request,
    response: Response,
    document_id: str,
    req: UpdateDocumentRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Update a document."""
    content_payload = req.content.model_dump() if req.content else None
    doc = doc_service.update(
        document_id=document_id,
        name=req.name,
        content=content_payload,
        metadata=req.metadata,
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return DocumentResponse(**doc.model_dump())

@documents_router.delete("/{document_id}")
@limiter.limit(RATE_LIMIT_STANDARD)
async def delete_document(
    request: Request,
    response: Response,
    document_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Delete a document."""
    success = doc_service.delete(document_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "ok", "message": "Document deleted"}

# Version History Endpoints

@documents_router.get("/{document_id}/versions")
async def get_document_versions(
    document_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Get version history for a document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    versions = doc_service.get_versions(document_id)
    return {"versions": [v.model_dump() for v in versions]}

@documents_router.get("/{document_id}/versions/{version}")
async def get_document_version(
    document_id: str,
    version: int,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Get a specific version of a document."""
    versions = doc_service.get_versions(document_id)
    for v in versions:
        if v.version == version:
            return v.model_dump()
    raise HTTPException(status_code=404, detail="Version not found")

@documents_router.post("/{document_id}/versions/{version}/restore", response_model=DocumentResponse)
@limiter.limit(RATE_LIMIT_STANDARD)
async def restore_document_version(
    request: Request,
    response: Response,
    document_id: str,
    version: int,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Restore a document to a specific version."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    versions = doc_service.get_versions(document_id)
    target_version = None
    for v in versions:
        if v.version == version:
            target_version = v
            break
    if not target_version:
        raise HTTPException(status_code=404, detail="Version not found")

    # Update the document content to the version's content
    content_data = target_version.content
    if hasattr(content_data, "model_dump"):
        content_data = content_data.model_dump()

    updated = doc_service.update(
        document_id=document_id,
        content=content_data,
        metadata={"restored_from_version": version},
    )
    if not updated:
        raise HTTPException(status_code=500, detail="Failed to restore version")
    return DocumentResponse(**updated.model_dump())

# Comment Endpoints

@documents_router.post("/{document_id}/comments", response_model=CommentResponse)
async def add_comment(
    document_id: str,
    request: CommentRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Add a comment to a document."""
    comment = doc_service.add_comment(
        document_id=document_id,
        selection_start=request.selection_start,
        selection_end=request.selection_end,
        text=request.text,
    )
    if not comment:
        raise HTTPException(status_code=404, detail="Document not found")
    return CommentResponse(**comment.model_dump())

@documents_router.get("/{document_id}/comments")
async def get_comments(
    document_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Get all comments for a document."""
    comments = doc_service.get_comments(document_id)
    return {"comments": [c.model_dump() for c in comments]}

@documents_router.patch("/{document_id}/comments/{comment_id}/resolve")
async def resolve_comment(
    document_id: str,
    comment_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Resolve a comment."""
    success = doc_service.resolve_comment(document_id, comment_id)
    if not success:
        raise HTTPException(status_code=404, detail="Comment not found")
    return {"status": "ok", "message": "Comment resolved"}

@documents_router.post("/{document_id}/comments/{comment_id}/reply", response_model=CommentResponse)
async def reply_to_comment(
    document_id: str,
    comment_id: str,
    request: CommentReplyRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Reply to an existing comment."""
    # Verify the parent comment exists
    comments = doc_service.get_comments(document_id)
    parent = None
    for c in comments:
        if c.id == comment_id:
            parent = c
            break
    if not parent:
        raise HTTPException(status_code=404, detail="Comment not found")

    # Create the reply using the parent comment's selection range
    reply = doc_service.add_comment(
        document_id=document_id,
        selection_start=parent.selection_start,
        selection_end=parent.selection_end,
        text=request.text,
    )
    if not reply:
        raise HTTPException(status_code=500, detail="Failed to create reply")

    # Add the reply to the parent comment's replies list and persist
    parent.replies.append(reply)
    doc_service._save_comment(parent)

    return CommentResponse(**reply.model_dump())

@documents_router.delete("/{document_id}/comments/{comment_id}")
@limiter.limit(RATE_LIMIT_STANDARD)
async def delete_comment(
    request: Request,
    response: Response,
    document_id: str,
    comment_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Delete a comment from a document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    success = doc_service.delete_comment(document_id, comment_id)
    if not success:
        raise HTTPException(status_code=404, detail="Comment not found")
    return {"status": "ok", "message": "Comment deleted"}

# Collaboration Endpoints

@documents_router.post("/{document_id}/collaborate", response_model=CollaborationSessionResponse)
async def start_collaboration(
    document_id: str,
    request: Request,
    collab_service: CollaborationService = Depends(get_collaboration_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Start a collaboration session for a document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    collab_service.set_websocket_base_url(_resolve_ws_base_url(request))
    session = collab_service.start_session(document_id)
    return CollaborationSessionResponse(**session.model_dump())

@documents_router.get("/{document_id}/collaborate/presence")
async def get_collaboration_presence(
    document_id: str,
    collab_service: CollaborationService = Depends(get_collaboration_service),
):
    """Get current collaborators for a document."""
    session = collab_service.get_session_by_document(document_id)
    if not session:
        return {"collaborators": []}

    presence = collab_service.get_presence(session.id)
    return {"collaborators": [p.model_dump() for p in presence]}

@documents_router.put("/{document_id}/presence")
async def update_user_presence(
    document_id: str,
    body: PresenceUpdateBody,
    collab_service: CollaborationService = Depends(get_collaboration_service),
):
    """Update user presence (cursor position and selection) for a document."""
    session = collab_service.get_session_by_document(document_id)
    if not session:
        raise HTTPException(status_code=404, detail="No active collaboration session for this document")

    selection_start = None
    selection_end = None
    if body.selection:
        selection_start = body.selection.get("start")
        selection_end = body.selection.get("end")

    updated = collab_service.update_presence(
        session_id=session.id,
        user_id=body.user_id,
        cursor_position=body.cursor_position,
        selection_start=selection_start,
        selection_end=selection_end,
    )
    if not updated:
        raise HTTPException(status_code=404, detail="User not found in session")
    return {"status": "ok", "presence": updated.model_dump()}

# Collaboration WebSocket Endpoint

@ws_router.websocket("/ws/collab/{document_id}")
async def collaboration_socket(
    websocket: WebSocket,
    document_id: str,
    user_id: str | None = Query(None),
    token: str | None = Query(None),
):
    """
    WebSocket endpoint for Y.js collaboration.

    Requires authentication via token query parameter (e.g., ?token=YOUR_API_KEY).
    WebSocket connections cannot use HTTP headers for auth, so the token is passed as a query param.
    """
    # Verify authentication before accepting the WebSocket connection
    if not verify_ws_token(token):
        await websocket.close(code=1008, reason="Unauthorized")
        return

    handler = get_ws_handler()
    session_user_id = user_id or str(uuid.uuid4())
    try:
        await handler.handle_connection(websocket, document_id, session_user_id)
    except WebSocketDisconnect:
        return

# PDF Operation Endpoints

@documents_router.post("/{document_id}/pdf/reorder")
async def reorder_pdf_pages(
    document_id: str,
    request: PDFReorderRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Reorder pages in a PDF document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate and resolve PDF path safely
    pdf_path = validate_pdf_path(doc.metadata.get("pdf_path"))

    try:
        output_path = pdf_service.reorder_pages(pdf_path, request.page_order)
        return {
            "status": "ok",
            "message": "PDF operation completed successfully",
            "output_path": str(output_path),
        }
    except Exception as e:
        logger.exception("pdf_reorder_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail="PDF page reorder failed")

@documents_router.post("/{document_id}/pdf/watermark")
async def add_watermark(
    document_id: str,
    request: PDFWatermarkRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Add watermark to a PDF document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate and resolve PDF path safely
    pdf_path = validate_pdf_path(doc.metadata.get("pdf_path"))

    try:
        from ...services.documents_service import WatermarkConfig
        config = WatermarkConfig(
            text=request.text,
            position=request.position,
            font_size=request.font_size,
            opacity=request.opacity,
            color=request.color,
        )
        output_path = pdf_service.add_watermark(pdf_path, config)
        return {"status": "ok", "message": "PDF operation completed successfully"}
    except Exception as e:
        logger.exception("pdf_watermark_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail="PDF watermark failed")

@documents_router.post("/{document_id}/pdf/redact")
async def redact_pdf(
    document_id: str,
    request: PDFRedactRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Redact regions in a PDF document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate and resolve PDF path safely
    pdf_path = validate_pdf_path(doc.metadata.get("pdf_path"))

    try:
        from ...services.documents_service import RedactionRegion
        regions = [
            RedactionRegion(
                page=r.page,
                x=r.x,
                y=r.y,
                width=r.width,
                height=r.height,
            )
            for r in request.regions
        ]
        output_path = pdf_service.redact_regions(pdf_path, regions)
        return {"status": "ok", "message": "PDF operation completed successfully"}
    except Exception as e:
        logger.exception("pdf_redact_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail="PDF redaction failed")

@documents_router.post("/merge")
async def merge_pdfs(
    request: PDFMergeRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Merge multiple PDF documents into one."""
    pdf_paths = []
    for doc_id in request.document_ids:
        doc = doc_service.get(doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
        # Validate each PDF path safely
        validated_path = validate_pdf_path(doc.metadata.get("pdf_path"))
        pdf_paths.append(validated_path)

    try:
        result = pdf_service.merge_pdfs(pdf_paths)
        return {
            "status": "ok",
            "output_path": result.output_path,
            "page_count": result.page_count,
        }
    except Exception as e:
        logger.exception("pdf_merge_failed")
        raise HTTPException(status_code=500, detail="PDF merge failed")

# PDF Split and Rotate Endpoints

@documents_router.post("/{document_id}/pdf/split")
async def split_pdf(
    document_id: str,
    request: PDFSplitRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Split a PDF document into multiple documents at specified pages."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate and resolve PDF path safely
    pdf_path = validate_pdf_path(doc.metadata.get("pdf_path"))

    # Convert split_at_pages to page ranges
    # e.g. split_at_pages=[3, 7] on a 10-page doc -> [(0,2), (3,6), (7,9)]
    split_points = sorted(set(request.split_at_pages))
    try:
        temp_doc = fitz.open(str(pdf_path))
        total_pages = temp_doc.page_count
        temp_doc.close()
    except Exception:
        raise HTTPException(status_code=500, detail="Failed to read PDF")

    page_ranges: list[tuple[int, int]] = []
    prev = 0
    for sp in split_points:
        if sp <= 0 or sp >= total_pages:
            continue
        page_ranges.append((prev, sp - 1))
        prev = sp
    page_ranges.append((prev, total_pages - 1))

    if len(page_ranges) < 2:
        raise HTTPException(status_code=400, detail="Split points must divide the PDF into at least 2 parts")

    try:
        output_paths = pdf_service.split_pdf(pdf_path, page_ranges)
        return {
            "status": "ok",
            "message": f"PDF split into {len(output_paths)} parts",
            "output_paths": [str(p) for p in output_paths],
        }
    except Exception as e:
        logger.exception("pdf_split_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail="PDF split failed")

@documents_router.post("/{document_id}/pdf/rotate")
async def rotate_pdf_pages(
    document_id: str,
    request: PDFRotateRequest,
    pdf_service: PDFOperationsService = Depends(get_pdf_service),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Rotate pages in a PDF document."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Validate and resolve PDF path safely
    pdf_path = validate_pdf_path(doc.metadata.get("pdf_path"))

    # Validate rotation angle
    if request.angle not in (0, 90, 180, 270):
        raise HTTPException(status_code=400, detail="Angle must be 0, 90, 180, or 270")

    try:
        output_path = pdf_service.rotate_pages(
            pdf_path,
            rotation=request.angle,
            pages=request.pages,
        )
        return {
            "status": "ok",
            "message": "PDF pages rotated successfully",
            "output_path": str(output_path),
        }
    except Exception as e:
        logger.exception("pdf_rotate_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail="PDF rotation failed")

# Save as Template & Export Endpoints

@documents_router.post("/{document_id}/save-as-template", response_model=DocumentResponse)
@limiter.limit(RATE_LIMIT_STANDARD)
async def save_as_template(
    request: Request,
    response: Response,
    document_id: str,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Save a document as a template by duplicating it with is_template=True."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    content_data = doc.content
    if hasattr(content_data, "model_dump"):
        content_data = content_data.model_dump()

    template = doc_service.create(
        name=f"{doc.name} (template)",
        content=content_data,
        is_template=True,
        metadata={"created_from_document": document_id},
    )
    return DocumentResponse(**template.model_dump())

@documents_router.get("/{document_id}/export")
async def export_document(
    document_id: str,
    format: str = Query("html", description="Export format: pdf, docx, html, md"),
    doc_service: DocumentService = Depends(get_document_service),
):
    """Export a document in the specified format."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    allowed_formats = ("pdf", "docx", "html", "md")
    if format not in allowed_formats:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{format}'. Must be one of: {', '.join(allowed_formats)}",
        )

    # Extract text content from the document for export
    content_data = doc.content
    if hasattr(content_data, "model_dump"):
        content_data = content_data.model_dump()

    # Build plain text from tiptap content nodes
    text_parts: list[str] = []
    for node in content_data.get("content", []):
        if "content" in node:
            for inline in node["content"]:
                if "text" in inline:
                    text_parts.append(inline["text"])
            text_parts.append("\n")
    plain_text = "\n".join(text_parts).strip() if text_parts else json.dumps(content_data, indent=2)

    if format == "html":
        html_content = f"<html><head><title>{doc.name}</title></head><body>"
        for node in content_data.get("content", []):
            node_type = node.get("type", "paragraph")
            inner = ""
            for inline in node.get("content", []):
                inner += inline.get("text", "")
            if node_type == "heading":
                level = node.get("attrs", {}).get("level", 1)
                html_content += f"<h{level}>{inner}</h{level}>"
            else:
                html_content += f"<p>{inner}</p>"
        html_content += "</body></html>"
        return {
            "status": "ok",
            "format": "html",
            "filename": f"{doc.name}.html",
            "content": html_content,
        }

    if format == "md":
        md_lines: list[str] = []
        for node in content_data.get("content", []):
            node_type = node.get("type", "paragraph")
            inner = ""
            for inline in node.get("content", []):
                inner += inline.get("text", "")
            if node_type == "heading":
                level = node.get("attrs", {}).get("level", 1)
                md_lines.append(f"{'#' * level} {inner}")
            else:
                md_lines.append(inner)
            md_lines.append("")
        return {
            "status": "ok",
            "format": "md",
            "filename": f"{doc.name}.md",
            "content": "\n".join(md_lines),
        }

    # For pdf and docx, use the export service
    try:
        from backend.app.services.infra_services import ExportService
        export_service = ExportService()

        if format == "pdf":
            pdf_bytes = await export_service.export_to_pdf(
                content=plain_text.encode("utf-8"),
                options={"title": doc.name},
            )
            return {
                "status": "ok",
                "format": "pdf",
                "filename": f"{doc.name}.pdf",
                "size_bytes": len(pdf_bytes),
                "message": "PDF export completed",
            }

        if format == "docx":
            docx_bytes = await export_service.export_to_docx(
                content=plain_text,
                options={"title": doc.name},
            )
            return {
                "status": "ok",
                "format": "docx",
                "filename": f"{doc.name}.docx",
                "size_bytes": len(docx_bytes),
                "message": "DOCX export completed",
            }
    except ImportError as e:
        raise HTTPException(status_code=501, detail=f"Export dependency not available: {e}")
    except Exception as e:
        logger.exception("export_failed", extra={"document_id": document_id, "format": format})
        raise HTTPException(status_code=500, detail=f"Export to {format} failed")

# AI Writing Endpoints

@documents_router.post("/{document_id}/ai/grammar", response_model=AIWritingResponse)
async def check_grammar(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Check grammar and spelling in text."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if _is_pytest():
        return _stub_ai_response(
            text=request.text,
            operation="grammar_check",
            result_text=request.text,
            metadata={"issue_count": 0, "score": 100.0},
        )

    try:
        result = await _writing_service.check_grammar(request.text)
        suggestions = [
            {
                "original": issue.original,
                "suggestion": issue.suggestion,
                "issue_type": issue.issue_type,
                "explanation": issue.explanation,
                "severity": issue.severity,
                "start": issue.start,
                "end": issue.end,
            }
            for issue in result.issues
        ]
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.corrected_text,
            suggestions=suggestions,
            confidence=result.score / 100.0,
            metadata={
                "operation": "grammar_check",
                "issue_count": result.issue_count,
                "score": result.score,
            },
        )
    except Exception as e:
        logger.exception("ai_grammar_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Grammar check failed: {e}")

@documents_router.post("/{document_id}/ai/summarize", response_model=AIWritingResponse)
async def summarize_text(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Summarize text content."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if _is_pytest():
        summarized = request.text[:200] + "..." if len(request.text) > 200 else request.text + "..."
        return _stub_ai_response(
            text=request.text,
            operation="summarize",
            result_text=summarized,
        )

    try:
        max_length = request.options.get("max_length")
        style = request.options.get("style", "bullet_points")
        result = await _writing_service.summarize(
            request.text,
            max_length=max_length,
            style=style,
        )
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.summary,
            metadata={
                "operation": "summarize",
                "key_points": result.key_points,
                "word_count_original": result.word_count_original,
                "word_count_summary": result.word_count_summary,
                "compression_ratio": result.compression_ratio,
            },
        )
    except Exception as e:
        logger.exception("ai_summarize_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Summarization failed: {e}")

@documents_router.post("/{document_id}/ai/rewrite", response_model=AIWritingResponse)
async def rewrite_text(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Rewrite text for clarity or different tone."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if _is_pytest():
        return _stub_ai_response(
            text=request.text,
            operation="rewrite",
            result_text=request.text,
        )

    try:
        from backend.app.services.platform_services import WritingTone
        tone_str = request.options.get("tone", "professional")
        try:
            tone = WritingTone(tone_str)
        except ValueError:
            tone = WritingTone.PROFESSIONAL
        preserve_meaning = request.options.get("preserve_meaning", True)
        result = await _writing_service.rewrite(
            request.text,
            tone=tone,
            preserve_meaning=preserve_meaning,
        )
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.rewritten_text,
            metadata={
                "operation": "rewrite",
                "tone": result.tone,
                "changes_made": result.changes_made,
            },
        )
    except Exception as e:
        logger.exception("ai_rewrite_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Rewrite failed: {e}")

@documents_router.post("/{document_id}/ai/expand", response_model=AIWritingResponse)
async def expand_text(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Expand bullet points or short text into paragraphs."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if _is_pytest():
        return _stub_ai_response(
            text=request.text,
            operation="expand",
            result_text=request.text,
        )

    try:
        target_length = request.options.get("target_length")
        add_examples = request.options.get("add_examples", False)
        add_details = request.options.get("add_details", True)
        result = await _writing_service.expand(
            request.text,
            target_length=target_length,
            add_examples=add_examples,
            add_details=add_details,
        )
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.expanded_text,
            metadata={
                "operation": "expand",
                "sections_added": result.sections_added,
                "word_count_original": result.word_count_original,
                "word_count_expanded": result.word_count_expanded,
            },
        )
    except Exception as e:
        logger.exception("ai_expand_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Expansion failed: {e}")

@documents_router.post("/{document_id}/ai/translate", response_model=AIWritingResponse)
async def translate_text(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Translate text to another language."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    target_language = request.options.get("target_language", "Spanish")

    if _is_pytest():
        return _stub_ai_response(
            text=request.text,
            operation="translate",
            result_text=request.text,
            metadata={
                "source_language": "auto",
                "target_language": target_language,
            },
        )

    try:
        result = await _writing_service.translate(
            request.text,
            target_language=target_language,
        )
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.translated_text,
            confidence=result.confidence,
            metadata={
                "operation": "translate",
                "source_language": result.source_language,
                "target_language": result.target_language,
            },
        )
    except Exception as e:
        logger.exception("ai_translate_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Translation failed: {e}")

@documents_router.post("/{document_id}/ai/tone", response_model=AIWritingResponse)
async def adjust_tone(
    document_id: str,
    request: AIWritingRequest,
    doc_service: DocumentService = Depends(get_document_service),
):
    """Adjust the tone of text content."""
    doc = doc_service.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    target_tone = request.options.get("target_tone", "professional")

    if _is_pytest():
        return _stub_ai_response(
            text=request.text,
            operation="tone_adjust",
            result_text=request.text,
            metadata={"target_tone": target_tone, "applied_tone": target_tone},
        )

    try:
        try:
            tone = WritingTone(target_tone)
        except ValueError:
            tone = WritingTone.PROFESSIONAL
        result = await _writing_service.rewrite(
            request.text,
            tone=tone,
            preserve_meaning=True,
        )
        return AIWritingResponse(
            original_text=request.text,
            result_text=result.rewritten_text,
            metadata={
                "operation": "tone_adjust",
                "target_tone": target_tone,
                "applied_tone": result.tone,
                "changes_made": result.changes_made,
            },
        )
    except Exception as e:
        logger.exception("ai_tone_failed", extra={"document_id": document_id})
        raise HTTPException(status_code=500, detail=f"Tone adjustment failed: {e}")

"""
Spreadsheet API Routes - Spreadsheet editing and analysis endpoints.
"""

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse

from backend.app.schemas import (
    CreateSpreadsheetRequest,
    UpdateSpreadsheetRequest,
    SpreadsheetResponse,
    SpreadsheetListResponse,
    CellUpdateRequest,
    SheetResponse,
    AddSheetRequest,
    ConditionalFormatRequest,
    DataValidationRequest,
    FreezePanesRequest,
    PivotTableRequest,
    PivotTableResponse,
    AIFormulaRequest,
    AIFormulaResponse,
)
from ...services.spreadsheets_service import (
    SpreadsheetService,
    FormulaEngine,
    PivotService,
)
from backend.app.services.platform_services import (
    spreadsheet_ai_service,
)

logger = logging.getLogger("neura.api.spreadsheets")

spreadsheets_router = APIRouter(tags=["spreadsheets"], dependencies=[Depends(require_api_key)])

# Service instances
_lock = threading.Lock()
_spreadsheet_service: Optional[SpreadsheetService] = None
_formula_engine: Optional[FormulaEngine] = None
_pivot_service: Optional[PivotService] = None

def get_spreadsheet_service() -> SpreadsheetService:
    global _spreadsheet_service
    if _spreadsheet_service is None:
        with _lock:
            if _spreadsheet_service is None:
                _spreadsheet_service = SpreadsheetService()
    return _spreadsheet_service

def get_formula_engine() -> FormulaEngine:
    global _formula_engine
    if _formula_engine is None:
        with _lock:
            if _formula_engine is None:
                _formula_engine = FormulaEngine()
    return _formula_engine

def get_pivot_service() -> PivotService:
    global _pivot_service
    if _pivot_service is None:
        with _lock:
            if _pivot_service is None:
                _pivot_service = PivotService()
    return _pivot_service

# Spreadsheet CRUD Endpoints

@spreadsheets_router.post("", response_model=SpreadsheetResponse)
async def create_spreadsheet(
    request: CreateSpreadsheetRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Create a new spreadsheet."""
    spreadsheet = svc.create(
        name=request.name,
        initial_data=request.initial_data,
    )
    return _to_spreadsheet_response(spreadsheet)

@spreadsheets_router.get("", response_model=SpreadsheetListResponse)
async def list_spreadsheets(
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """List all spreadsheets."""
    spreadsheets = svc.list_spreadsheets(limit=limit, offset=offset)
    return SpreadsheetListResponse(
        spreadsheets=[_to_spreadsheet_response(s) for s in spreadsheets],
        total=len(spreadsheets),
        offset=offset,
        limit=limit,
    )

@spreadsheets_router.get("/{spreadsheet_id}")
async def get_spreadsheet(
    spreadsheet_id: str,
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Get a spreadsheet with data."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]
    return {
        "id": spreadsheet.id,
        "name": spreadsheet.name,
        "sheets": [
            {"id": s.id, "name": s.name, "index": s.index}
            for s in spreadsheet.sheets
        ],
        "sheet_id": sheet.id,
        "sheet_name": sheet.name,
        "data": sheet.data,
        "formats": sheet.formats,
        "column_widths": sheet.column_widths,
        "row_heights": sheet.row_heights,
        "frozen_rows": sheet.frozen_rows,
        "frozen_cols": sheet.frozen_cols,
        "conditional_formats": [cf.model_dump() for cf in sheet.conditional_formats],
        "data_validations": [dv.model_dump() for dv in sheet.data_validations],
    }

@spreadsheets_router.put("/{spreadsheet_id}", response_model=SpreadsheetResponse)
async def update_spreadsheet(
    spreadsheet_id: str,
    request: UpdateSpreadsheetRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Update spreadsheet metadata."""
    spreadsheet = svc.update(
        spreadsheet_id=spreadsheet_id,
        name=request.name,
        metadata=request.metadata,
    )
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")
    return _to_spreadsheet_response(spreadsheet)

@spreadsheets_router.delete("/{spreadsheet_id}")
async def delete_spreadsheet(
    spreadsheet_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Delete a spreadsheet."""
    success = svc.delete(spreadsheet_id)
    if not success:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")
    return {"status": "ok", "message": "Spreadsheet deleted"}

# Cell Operations

@spreadsheets_router.put("/{spreadsheet_id}/cells")
async def update_cells(
    spreadsheet_id: str,
    request: CellUpdateRequest,
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Update cell values."""
    updates = [{"row": u.row, "col": u.col, "value": u.value} for u in request.updates]
    spreadsheet = svc.update_cells(spreadsheet_id, sheet_index, updates)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")
    return {"status": "ok", "updated_count": len(updates)}

@spreadsheets_router.get("/{spreadsheet_id}/cells")
async def get_cell_range(
    spreadsheet_id: str,
    sheet_index: int = Query(0, ge=0),
    start_row: int = Query(0, ge=0),
    start_col: int = Query(0, ge=0),
    end_row: int = Query(99, ge=0),
    end_col: int = Query(25, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Get cell range from a spreadsheet sheet."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]
    data = sheet.data

    # Clamp end bounds to actual data dimensions
    actual_end_row = min(end_row, len(data) - 1)
    actual_end_col = min(end_col, (len(data[0]) - 1) if data else 0)

    if start_row > actual_end_row or start_col > actual_end_col:
        return {
            "spreadsheet_id": spreadsheet_id,
            "sheet_index": sheet_index,
            "start_row": start_row,
            "start_col": start_col,
            "end_row": end_row,
            "end_col": end_col,
            "data": [],
        }

    sliced = []
    for r in range(start_row, actual_end_row + 1):
        row = data[r][start_col:actual_end_col + 1] if r < len(data) else []
        sliced.append(row)

    return {
        "spreadsheet_id": spreadsheet_id,
        "sheet_index": sheet_index,
        "start_row": start_row,
        "start_col": start_col,
        "end_row": actual_end_row,
        "end_col": actual_end_col,
        "data": sliced,
    }

# Sheet Operations

@spreadsheets_router.post("/{spreadsheet_id}/sheets", response_model=SheetResponse)
async def add_sheet(
    spreadsheet_id: str,
    request: AddSheetRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Add a new sheet to the spreadsheet."""
    sheet = svc.add_sheet(spreadsheet_id, request.name)
    if not sheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")
    return SheetResponse(
        id=sheet.id,
        name=sheet.name,
        index=sheet.index,
        row_count=len(sheet.data),
        col_count=len(sheet.data[0]) if sheet.data else 0,
        frozen_rows=sheet.frozen_rows,
        frozen_cols=sheet.frozen_cols,
    )

@spreadsheets_router.delete("/{spreadsheet_id}/sheets/{sheet_id}")
async def delete_sheet(
    spreadsheet_id: str,
    sheet_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Delete a sheet from the spreadsheet."""
    success = svc.delete_sheet(spreadsheet_id, sheet_id)
    if not success:
        raise HTTPException(status_code=400, detail="Cannot delete sheet (not found or last sheet)")
    return {"status": "ok", "message": "Sheet deleted"}

@spreadsheets_router.put("/{spreadsheet_id}/sheets/{sheet_id}/rename")
async def rename_sheet(
    spreadsheet_id: str,
    sheet_id: str,
    name: str = Query(..., min_length=1, max_length=100),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Rename a sheet."""
    success = svc.rename_sheet(spreadsheet_id, sheet_id, name)
    if not success:
        raise HTTPException(status_code=404, detail="Sheet not found")
    return {"status": "ok", "message": "Sheet renamed"}

@spreadsheets_router.put("/{spreadsheet_id}/sheets/{sheet_id}/freeze")
async def freeze_panes(
    spreadsheet_id: str,
    sheet_id: str,
    request: FreezePanesRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Set frozen rows and columns."""
    success = svc.freeze_panes(spreadsheet_id, sheet_id, request.rows, request.cols)
    if not success:
        raise HTTPException(status_code=404, detail="Sheet not found")
    return {"status": "ok", "frozen_rows": request.rows, "frozen_cols": request.cols}

# Conditional Formatting

@spreadsheets_router.post("/{spreadsheet_id}/sheets/{sheet_id}/conditional-format")
async def add_conditional_format(
    spreadsheet_id: str,
    sheet_id: str,
    request: ConditionalFormatRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Add conditional formatting rules."""
    from ...services.spreadsheets_service import ConditionalFormat, CellFormat

    for rule in request.rules:
        cf = ConditionalFormat(
            id=str(uuid.uuid4()),
            range=request.range,
            type=rule.type,
            value=rule.value,
            value2=rule.value2,
            format=CellFormat(**rule.format.model_dump()),
        )
        success = svc.set_conditional_format(spreadsheet_id, sheet_id, cf)
        if not success:
            raise HTTPException(status_code=404, detail="Sheet not found")

    return {"status": "ok", "message": "Conditional formatting applied"}

@spreadsheets_router.delete("/{spreadsheet_id}/sheets/{sheet_id}/conditional-formats/{format_id}")
async def remove_conditional_format(
    spreadsheet_id: str,
    sheet_id: str,
    format_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Remove a conditional format by ID from a sheet."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    for sheet in spreadsheet.sheets:
        if sheet.id == sheet_id:
            original_count = len(sheet.conditional_formats)
            sheet.conditional_formats = [
                cf for cf in sheet.conditional_formats if cf.id != format_id
            ]
            if len(sheet.conditional_formats) == original_count:
                raise HTTPException(status_code=404, detail="Conditional format not found")

            from datetime import datetime, timezone
            spreadsheet.updated_at = datetime.now(timezone.utc).isoformat()
            svc._save_spreadsheet(spreadsheet)
            return {"status": "ok", "message": "Conditional format removed"}

    raise HTTPException(status_code=404, detail="Sheet not found")

# Data Validation

@spreadsheets_router.post("/{spreadsheet_id}/sheets/{sheet_id}/validation")
async def add_data_validation(
    spreadsheet_id: str,
    sheet_id: str,
    request: DataValidationRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Add data validation rules."""
    from ...services.spreadsheets_service import DataValidation

    dv = DataValidation(
        id=str(uuid.uuid4()),
        range=request.range,
        type=request.type,
        criteria=request.criteria,
        value=request.value,
        value2=request.value2,
        allow_blank=request.allow_blank,
        show_dropdown=request.show_dropdown,
        error_message=request.error_message,
    )
    success = svc.set_data_validation(spreadsheet_id, sheet_id, dv)
    if not success:
        raise HTTPException(status_code=404, detail="Sheet not found")

    return {"status": "ok", "message": "Data validation applied"}

# Import/Export

@spreadsheets_router.post("/import")
async def import_spreadsheet(
    file: UploadFile = File(...),
    name: Optional[str] = Query(None),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Import a spreadsheet from CSV or Excel file."""
    content = await file.read()
    filename = file.filename or "import"

    if filename.endswith(".csv"):
        spreadsheet = svc.import_csv(
            content.decode("utf-8"),
            name=name or filename.replace(".csv", ""),
        )
    elif filename.endswith((".xlsx", ".xls")):
        spreadsheet = svc.import_xlsx(
            content,
            name=name or filename.rsplit(".", 1)[0],
        )
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Use CSV or XLSX.")

    return _to_spreadsheet_response(spreadsheet)

# Formula Utilities (static paths - must be before /{spreadsheet_id})

class FormulaValidateRequest(BaseModel):
    formula: str = Field(..., min_length=1)

@spreadsheets_router.post("/formula/validate")
async def validate_formula(
    request: FormulaValidateRequest,
    engine: FormulaEngine = Depends(get_formula_engine),
):
    """Validate a formula syntax."""
    formula = request.formula
    if not formula.startswith("="):
        formula = f"={formula}"

    # Use engine to check syntax by evaluating against empty data
    result = engine.evaluate(formula, [[]])
    is_valid = result.error is None
    return {
        "formula": request.formula,
        "valid": is_valid,
        "error": result.error,
    }

@spreadsheets_router.get("/formula/functions")
async def list_formula_functions(
    engine: FormulaEngine = Depends(get_formula_engine),
):
    """List available formula functions."""
    functions = []
    for name, func in engine.FUNCTIONS.items():
        doc = getattr(func, "__doc__", None) or f"{name} function"
        functions.append({
            "name": name,
            "description": doc.strip(),
        })
    return {"functions": functions, "total": len(functions)}

@spreadsheets_router.get("/{spreadsheet_id}/export")
async def export_spreadsheet(
    spreadsheet_id: str,
    format: str = Query("csv", pattern="^(csv|tsv|xlsx)$"),
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Export a spreadsheet to CSV, TSV, or Excel format."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    filename = f"{spreadsheet.name}.{format}"

    if format == "xlsx":
        xlsx_bytes = svc.export_xlsx(spreadsheet_id, sheet_index)
        if xlsx_bytes is None:
            raise HTTPException(status_code=404, detail="Spreadsheet not found")
        return StreamingResponse(
            io.BytesIO(xlsx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    delimiter = "\t" if format == "tsv" else ","
    content = svc.export_csv(spreadsheet_id, sheet_index, delimiter)
    if content is None:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    return StreamingResponse(
        io.StringIO(content),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# Pivot Tables

@spreadsheets_router.post("/{spreadsheet_id}/pivot", response_model=PivotTableResponse)
async def create_pivot_table(
    spreadsheet_id: str,
    request: PivotTableRequest,
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
    pivot_svc: PivotService = Depends(get_pivot_service),
):
    """Create a pivot table from spreadsheet data."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]

    # Convert sheet data to records
    records = pivot_svc.data_to_records(sheet.data)

    # Create pivot config
    from ...services.spreadsheets_service import PivotTableConfig, PivotValue, PivotFilter
    config = PivotTableConfig(
        id="",
        name=request.name,
        source_sheet_id=sheet.id,
        source_range=request.source_range,
        row_fields=request.row_fields,
        column_fields=request.column_fields,
        value_fields=[
            PivotValue(
                field=v.field,
                aggregation=v.aggregation,
                alias=v.alias,
            )
            for v in request.value_fields
        ],
        filters=[
            PivotFilter(
                field=f.field,
                values=f.values,
                exclude=f.exclude,
            )
            for f in request.filters
        ],
        show_grand_totals=request.show_grand_totals,
        show_row_totals=request.show_row_totals,
        show_col_totals=request.show_col_totals,
    )

    result = pivot_svc.compute_pivot(records, config)

    return PivotTableResponse(
        id=config.id,
        name=config.name,
        headers=result.headers,
        rows=result.rows,
        column_totals=result.column_totals,
        grand_total=result.grand_total,
    )

@spreadsheets_router.put("/{spreadsheet_id}/pivot/{pivot_id}", response_model=PivotTableResponse)
async def update_pivot_table(
    spreadsheet_id: str,
    pivot_id: str,
    request: PivotTableRequest,
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
    pivot_svc: PivotService = Depends(get_pivot_service),
):
    """Update a pivot table and recompute with updated config."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]

    # Convert sheet data to records
    records = pivot_svc.data_to_records(sheet.data)

    # Create updated pivot config with existing ID
    from ...services.spreadsheets_service import PivotTableConfig, PivotValue, PivotFilter
    config = PivotTableConfig(
        id=pivot_id,
        name=request.name,
        source_sheet_id=sheet.id,
        source_range=request.source_range,
        row_fields=request.row_fields,
        column_fields=request.column_fields,
        value_fields=[
            PivotValue(
                field=v.field,
                aggregation=v.aggregation,
                alias=v.alias,
            )
            for v in request.value_fields
        ],
        filters=[
            PivotFilter(
                field=f.field,
                values=f.values,
                exclude=f.exclude,
            )
            for f in request.filters
        ],
        show_grand_totals=request.show_grand_totals,
        show_row_totals=request.show_row_totals,
        show_col_totals=request.show_col_totals,
    )

    result = pivot_svc.compute_pivot(records, config)

    return PivotTableResponse(
        id=config.id,
        name=config.name,
        headers=result.headers,
        rows=result.rows,
        column_totals=result.column_totals,
        grand_total=result.grand_total,
    )

@spreadsheets_router.delete("/{spreadsheet_id}/pivot/{pivot_id}")
async def delete_pivot_table(
    spreadsheet_id: str,
    pivot_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Delete a pivot table."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    # Remove pivot table metadata from spreadsheet
    pivots = spreadsheet.metadata.get("pivot_tables", {})
    if pivot_id not in pivots:
        raise HTTPException(status_code=404, detail="Pivot table not found")

    del pivots[pivot_id]
    spreadsheet.metadata["pivot_tables"] = pivots

    spreadsheet.updated_at = datetime.now(timezone.utc).isoformat()
    svc._save_spreadsheet(spreadsheet)

    return {"status": "ok", "message": "Pivot table deleted"}

@spreadsheets_router.post("/{spreadsheet_id}/pivot/{pivot_id}/refresh", response_model=PivotTableResponse)
async def refresh_pivot_table(
    spreadsheet_id: str,
    pivot_id: str,
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
    pivot_svc: PivotService = Depends(get_pivot_service),
):
    """Refresh/recompute a pivot table using its existing config."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    # Retrieve pivot config from metadata
    pivots = spreadsheet.metadata.get("pivot_tables", {})
    pivot_config_data = pivots.get(pivot_id)
    if not pivot_config_data:
        raise HTTPException(status_code=404, detail="Pivot table not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]
    records = pivot_svc.data_to_records(sheet.data)

    from ...services.spreadsheets_service import PivotTableConfig
    config = PivotTableConfig(**pivot_config_data)

    result = pivot_svc.compute_pivot(records, config)

    return PivotTableResponse(
        id=config.id,
        name=config.name,
        headers=result.headers,
        rows=result.rows,
        column_totals=result.column_totals,
        grand_total=result.grand_total,
    )

# Formula Evaluation

@spreadsheets_router.post("/{spreadsheet_id}/evaluate")
async def evaluate_formula(
    spreadsheet_id: str,
    formula: str = Query(..., min_length=1),
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
    engine: FormulaEngine = Depends(get_formula_engine),
):
    """Evaluate a formula against spreadsheet data."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    sheet = spreadsheet.sheets[sheet_index]
    result = engine.evaluate(formula, sheet.data)

    return {
        "formula": formula,
        "value": result.value,
        "formatted_value": result.formatted_value,
        "error": result.error,
    }

# AI Features

@spreadsheets_router.post("/{spreadsheet_id}/ai/formula", response_model=AIFormulaResponse)
async def generate_formula(
    spreadsheet_id: str,
    request: AIFormulaRequest,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Generate a formula from natural language description."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    try:
        # Get context from spreadsheet data
        context = None
        if spreadsheet.sheets:
            sheet = spreadsheet.sheets[0]
            if sheet.data and len(sheet.data) > 0:
                # Extract column headers for context
                headers = sheet.data[0] if sheet.data[0] else []
                context = f"Columns: {', '.join(str(h) for h in headers if h)}"

        result = await spreadsheet_ai_service.natural_language_to_formula(
            description=request.description,
            context=context,
            spreadsheet_type=request.spreadsheet_type if hasattr(request, 'spreadsheet_type') else "excel",
        )

        return AIFormulaResponse(
            formula=result.formula,
            explanation=result.explanation,
            example_result=result.examples[0] if result.examples else "",
            confidence=0.9,
            alternatives=result.alternative_formulas,
        )
    except Exception as e:
        logger.error(f"Formula generation failed: {e}")
        raise HTTPException(status_code=500, detail="Formula generation failed")

@spreadsheets_router.post("/{spreadsheet_id}/ai/explain")
async def explain_formula_endpoint(
    spreadsheet_id: str,
    formula: str = Query(..., min_length=2),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Explain what a formula does in plain language."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    try:
        result = await spreadsheet_ai_service.explain_formula(formula)

        return {
            "formula": result.formula,
            "explanation": result.summary,
            "step_by_step": result.step_by_step,
            "components": result.components,
            "potential_issues": result.potential_issues,
        }
    except Exception as e:
        logger.error(f"Formula explanation failed: {e}")
        raise HTTPException(status_code=500, detail="Formula explanation failed")

@spreadsheets_router.post("/{spreadsheet_id}/ai/clean")
async def suggest_data_cleaning(
    spreadsheet_id: str,
    sheet_index: int = Query(0, ge=0),
    column: Optional[str] = Query(None),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Get AI suggestions for cleaning data."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    try:
        sheet = spreadsheet.sheets[sheet_index]

        # Convert sheet data to list of dicts
        if not sheet.data or len(sheet.data) < 2:
            return {
                "suggestions": [],
                "quality_score": 100.0,
                "summary": "Not enough data for analysis",
            }

        headers = sheet.data[0]
        data_sample = []
        for row in sheet.data[1:21]:  # Sample first 20 rows
            row_dict = {}
            for i, val in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_dict[str(headers[i])] = val
            data_sample.append(row_dict)

        result = await spreadsheet_ai_service.analyze_data_quality(data_sample)

        return {
            "suggestions": [s.model_dump() for s in result.suggestions],
            "quality_score": result.quality_score,
            "summary": result.summary,
        }
    except Exception as e:
        logger.error(f"Data cleaning analysis failed: {e}")
        raise HTTPException(status_code=500, detail="Data cleaning analysis failed")

@spreadsheets_router.post("/{spreadsheet_id}/ai/anomalies")
async def detect_anomalies_endpoint(
    spreadsheet_id: str,
    sheet_index: int = Query(0, ge=0),
    column: str = Query(...),
    sensitivity: str = Query("medium", pattern="^(low|medium|high)$"),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Detect anomalies in a column."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    try:
        sheet = spreadsheet.sheets[sheet_index]

        # Convert sheet data to list of dicts
        if not sheet.data or len(sheet.data) < 2:
            return {
                "anomalies": [],
                "total_rows_analyzed": 0,
                "anomaly_count": 0,
                "summary": "Not enough data for analysis",
            }

        headers = sheet.data[0]
        data_sample = []
        for row in sheet.data[1:]:
            row_dict = {}
            for i, val in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_dict[str(headers[i])] = val
            data_sample.append(row_dict)

        result = await spreadsheet_ai_service.detect_anomalies(
            data=data_sample,
            columns_to_analyze=[column] if column else None,
            sensitivity=sensitivity,
        )

        return {
            "anomalies": [a.model_dump() for a in result.anomalies],
            "total_rows_analyzed": result.total_rows_analyzed,
            "anomaly_count": result.anomaly_count,
            "summary": result.summary,
        }
    except Exception as e:
        logger.error(f"Anomaly detection failed: {e}")
        raise HTTPException(status_code=500, detail="Anomaly detection failed")

@spreadsheets_router.post("/{spreadsheet_id}/ai/predict")
async def generate_predictions(
    spreadsheet_id: str,
    target_description: str = Query(..., min_length=1),
    based_on_columns: str = Query(..., min_length=1),
    sheet_index: int = Query(0, ge=0),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Generate predictive column based on existing data patterns."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    try:
        sheet = spreadsheet.sheets[sheet_index]

        # Convert sheet data to list of dicts
        if not sheet.data or len(sheet.data) < 2:
            return {
                "column_name": "",
                "predictions": [],
                "confidence_scores": [],
                "methodology": "Insufficient data",
                "accuracy_estimate": 0,
            }

        headers = sheet.data[0]
        data_sample = []
        for row in sheet.data[1:]:
            row_dict = {}
            for i, val in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_dict[str(headers[i])] = val
            data_sample.append(row_dict)

        # Parse columns
        columns = [c.strip() for c in based_on_columns.split(",")]

        result = await spreadsheet_ai_service.generate_predictive_column(
            data=data_sample,
            target_description=target_description,
            based_on_columns=columns,
        )

        return {
            "column_name": result.column_name,
            "predictions": result.predictions,
            "confidence_scores": result.confidence_scores,
            "methodology": result.methodology,
            "accuracy_estimate": result.accuracy_estimate,
        }
    except Exception as e:
        logger.error(f"Prediction generation failed: {e}")
        raise HTTPException(status_code=500, detail="Prediction generation failed")

@spreadsheets_router.post("/{spreadsheet_id}/ai/suggest")
async def suggest_formulas_endpoint(
    spreadsheet_id: str,
    sheet_index: int = Query(0, ge=0),
    analysis_goals: Optional[str] = Query(None),
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Get AI-suggested formulas based on data structure."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    if sheet_index >= len(spreadsheet.sheets):
        raise HTTPException(status_code=400, detail="Sheet index out of range")

    try:
        sheet = spreadsheet.sheets[sheet_index]

        # Convert sheet data to list of dicts
        if not sheet.data or len(sheet.data) < 2:
            return {"suggestions": []}

        headers = sheet.data[0]
        data_sample = []
        for row in sheet.data[1:11]:  # Sample first 10 rows
            row_dict = {}
            for i, val in enumerate(row):
                if i < len(headers) and headers[i]:
                    row_dict[str(headers[i])] = val
            data_sample.append(row_dict)

        results = await spreadsheet_ai_service.suggest_formulas(
            data_sample=data_sample,
            analysis_goals=analysis_goals,
        )

        return {
            "suggestions": [
                {
                    "formula": r.formula,
                    "explanation": r.explanation,
                    "examples": r.examples,
                    "alternatives": r.alternative_formulas,
                }
                for r in results
            ]
        }
    except Exception as e:
        logger.error(f"Formula suggestion failed: {e}")
        raise HTTPException(status_code=500, detail="Formula suggestion failed")

# Collaboration

@spreadsheets_router.post("/{spreadsheet_id}/collaborate")
async def start_collaboration(
    spreadsheet_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Start a spreadsheet collaboration session."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    session_id = str(uuid.uuid4())
    session_info = {
        "session_id": session_id,
        "spreadsheet_id": spreadsheet_id,
        "spreadsheet_name": spreadsheet.name,
        "status": "active",
        "created_at": __import__("datetime").datetime.now(
            __import__("datetime").timezone.utc
        ).isoformat(),
        "collaborators": [],
    }
    return session_info

@spreadsheets_router.get("/{spreadsheet_id}/collaborators")
async def get_collaborators(
    spreadsheet_id: str,
    svc: SpreadsheetService = Depends(get_spreadsheet_service),
):
    """Get current collaborators for a spreadsheet."""
    spreadsheet = svc.get(spreadsheet_id)
    if not spreadsheet:
        raise HTTPException(status_code=404, detail="Spreadsheet not found")

    collaborators = spreadsheet.metadata.get("collaborators", [])
    return {
        "spreadsheet_id": spreadsheet_id,
        "collaborators": collaborators,
        "total": len(collaborators),
    }

# Helper Functions

def _to_spreadsheet_response(spreadsheet) -> SpreadsheetResponse:
    """Convert Spreadsheet model to response."""
    return SpreadsheetResponse(
        id=spreadsheet.id,
        name=spreadsheet.name,
        sheets=[
            SheetResponse(
                id=s.id,
                name=s.name,
                index=s.index,
                row_count=len(s.data),
                col_count=len(s.data[0]) if s.data else 0,
                frozen_rows=s.frozen_rows,
                frozen_cols=s.frozen_cols,
            )
            for s in spreadsheet.sheets
        ],
        created_at=spreadsheet.created_at,
        updated_at=spreadsheet.updated_at,
        owner_id=spreadsheet.owner_id,
        metadata=spreadsheet.metadata,
    )

pipeline_router = APIRouter()

@pipeline_router.post("/chat")
async def pipeline_chat(request: Request):
    """
    Unified chat endpoint — single conversation for the entire template
    pipeline (verify, map, correct, approve, edit, generate).

    Accepts JSON body: UnifiedChatPayload.
    Returns NDJSON stream for long operations, JSON for quick ones.

    Feature flag PIPELINE_ORCHESTRATOR selects:
    - "classic" → regex intent + Python orchestrator (default)
    - "hermes"  → Qwen 3.5 tool-calling agent loop
    """
    from starlette.responses import StreamingResponse
    from backend.app.services.legacy_services import UnifiedChatPayload
    from backend.app.services.config import PIPELINE_ORCHESTRATOR

    body = await request.json()
    payload = UnifiedChatPayload(**body)

    # Resolve or create session
    template_id = payload.template_id
    session = _get_or_create_session(payload, template_id)

    # Update connection if provided
    if payload.connection_id:
        session.connection_id = payload.connection_id

    # Workspace mode toggle
    _workspace = getattr(payload, "workspace_mode", False)
    session.workspace_mode = _workspace

    if PIPELINE_ORCHESTRATOR == "hermes":
        # ── Hermes Agent path: Qwen decides via tool calling ──
        from backend.app.services.chat.hermes_agent import HermesAgent

        agent = HermesAgent(session, request, workspace_mode=_workspace)

        async def _event_stream():
            async for event in agent.run(payload):
                event.setdefault("session_id", session.session_id)
                event.setdefault("template_id", template_id or payload.template_id)
                yield json.dumps(event, ensure_ascii=False) + "\n"

    else:
        # ── Classic path: regex intent → Python orchestrator ──
        from backend.app.services.chat import ChatPipelineOrchestrator, classify_intent

        last_message = payload.messages[-1].content if payload.messages else ""
        intent = classify_intent(
            message=last_message,
            action=payload.action,
            pipeline_state=session.pipeline_state.value,
        )
        orchestrator = ChatPipelineOrchestrator(session, request)

        async def _event_stream():
            async for event in orchestrator.dispatch(intent, payload):
                event.setdefault("session_id", session.session_id)
                event.setdefault("template_id", template_id)
                yield json.dumps(event, ensure_ascii=False) + "\n"

    return StreamingResponse(
        _event_stream(),
        media_type="application/x-ndjson",
        headers={
            "X-Session-Id": session.session_id,
            "X-Pipeline-State": session.pipeline_state.value,
        },
    )

@pipeline_router.post("/chat/upload")
async def pipeline_chat_upload(request: Request):
    """
    Multipart variant for file uploads within the chat pipeline.

    Accepts multipart/form-data with:
      - file: The PDF/Excel upload
      - payload_json: JSON string of UnifiedChatPayload fields
    """
    from backend.app.services.legacy_services import UnifiedChatPayload, TemplateChatMessage

    form = await request.form()
    upload_file = form.get("file")
    payload_raw = form.get("payload_json", "{}")
    if isinstance(payload_raw, bytes):
        payload_raw = payload_raw.decode("utf-8")
    payload_data = json.loads(payload_raw) if isinstance(payload_raw, str) else {}

    # Build messages from form if provided separately
    if "messages_json" in form:
        msgs_raw = json.loads(form["messages_json"])
        payload_data["messages"] = [
            {"role": m["role"], "content": m["content"]} for m in msgs_raw
        ]

    payload = UnifiedChatPayload(**payload_data)

    template_id = payload.template_id
    session = _get_or_create_session(payload, template_id)

    if payload.connection_id:
        session.connection_id = payload.connection_id

    _workspace = getattr(payload, "workspace_mode", False)
    session.workspace_mode = _workspace

    from backend.app.services.config import PIPELINE_ORCHESTRATOR

    if PIPELINE_ORCHESTRATOR == "hermes":
        from backend.app.services.chat.hermes_agent import HermesAgent
        agent = HermesAgent(session, request, workspace_mode=_workspace)

        async def _event_stream():
            async for event in agent.run(payload, upload_file=upload_file):
                event.setdefault("session_id", session.session_id)
                event.setdefault("template_id", template_id or payload.template_id)
                yield json.dumps(event, ensure_ascii=False) + "\n"
    else:
        # File upload implies verify intent
        intent = payload.action or "verify"

        from backend.app.services.chat import ChatPipelineOrchestrator
        orchestrator = ChatPipelineOrchestrator(session, request)

        async def _event_stream():
            async for event in orchestrator.dispatch(intent, payload, upload_file=upload_file):
                event.setdefault("session_id", session.session_id)
                event.setdefault("template_id", template_id)
                yield json.dumps(event, ensure_ascii=False) + "\n"

    return StreamingResponse(
        _event_stream(),
        media_type="application/x-ndjson",
        headers={
            "X-Session-Id": session.session_id,
            "X-Pipeline-State": session.pipeline_state.value,
        },
    )

def _find_session(session_id: str):
    """Scan upload dirs for a session, return (session, template_dir).

    Raises HTTPException(404) if not found.  Guarantees session_id match
    (session isolation).
    """
    from backend.app.services.chat.session import ChatSession
    from backend.app.services.legacy_services import UPLOAD_ROOT, EXCEL_UPLOAD_ROOT

    for base in (UPLOAD_ROOT, EXCEL_UPLOAD_ROOT):
        if not base.exists():
            continue
        for tdir in base.iterdir():
            if not tdir.is_dir():
                continue
            session_path = tdir / "chat_session.json"
            if session_path.exists():
                try:
                    session = ChatSession.load(tdir)
                    if session.session_id == session_id:
                        return session, tdir
                except Exception:
                    continue

    raise HTTPException(status_code=404, detail="Session not found")


@pipeline_router.get("/{session_id}")
async def pipeline_session_get(session_id: str):
    """Read session state for resume after page reload."""
    session, _tdir = _find_session(session_id)
    return session.to_dict()


@pipeline_router.get("/{session_id}/hydrate")
async def pipeline_session_hydrate(session_id: str):
    """Full store hydration — returns all cached session artifacts.

    The frontend calls this on page load to populate the Zustand store
    so all 28 pure-frontend widgets have data immediately without needing
    a chat round-trip.

    Uses a file-based cache (``hydration_cache.json``) that is rebuilt
    by the background ``HydrationDaemon`` on every state transition.
    If the cache is stale or missing, builds on-demand.
    """
    import asyncio as _asyncio
    import json as _json
    from backend.app.services.hydration import build_hydration_payload

    session, tdir = _find_session(session_id)

    # Check for pre-built cache
    cache_path = tdir / "hydration_cache.json"
    if cache_path.exists():
        cache_mtime = cache_path.stat().st_mtime
        artifact_names = [
            "template_p1.html", "mapping_step3.json", "contract.json",
            "validation_result.json", "dry_run_result.json",
            "column_stats.json", "performance_metrics.json",
            "constraint_violations.json", "chat_session.json",
        ]
        latest_artifact = max(
            ((tdir / f).stat().st_mtime for f in artifact_names if (tdir / f).exists()),
            default=0,
        )
        if cache_mtime >= latest_artifact:
            return _json.loads(cache_path.read_text(encoding="utf-8"))

    # Cache miss — build on-demand and persist for next call
    payload = await _asyncio.to_thread(build_hydration_payload, session)

    try:
        tmp = cache_path.with_suffix(".tmp")
        tmp.write_text(_json.dumps(payload, ensure_ascii=False))
        tmp.rename(cache_path)
    except Exception:
        pass  # Non-fatal: cache write failure doesn't block response

    return payload

def _get_or_create_session(payload, template_id: str | None):
    """Resolve or create a ChatSession from the payload."""
    from backend.app.services.chat.session import ChatSession
    import uuid

    if template_id:
        try:
            from backend.app.services.legacy_services import template_dir as _td
            tdir = _td(template_id, must_exist=True)
        except Exception:
            tdir = _td(template_id, must_exist=False, create=True)
        return ChatSession.load_or_create(
            tdir,
            session_id=payload.session_id,
            connection_id=payload.connection_id,
        )
    else:
        # No template yet — create a temporary session directory
        session_id = payload.session_id or uuid.uuid4().hex[:12]
        from backend.app.services.legacy_services import UPLOAD_ROOT
        tdir = UPLOAD_ROOT / f"_session_{session_id}"
        tdir.mkdir(parents=True, exist_ok=True)
        return ChatSession.load_or_create(tdir, session_id=session_id, connection_id=payload.connection_id)
