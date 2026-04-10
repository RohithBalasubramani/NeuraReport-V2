# mypy: ignore-errors
"""V1 Report Pipeline -- synced from production.

Contains: html_table_parser, date_utils, common_helpers, contract_adapter,
dataframe_pipeline, report_context, strategies, discovery_metrics,
discovery_excel, discovery, ReportGenerate (PDF), ReportGenerateExcel,
_pdf_worker, xlsx_export, docx_export.
"""
from __future__ import annotations


# ======================================================================
# html_table_parser
# ======================================================================

from html.parser import HTMLParser


class _SimpleTableParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.tables: list[list[list[str]]] = []
        self.thead_counts: list[int] = []
        self._table_depth = 0
        self._collecting = False
        self._current_table: list[list[str]] | None = None
        self._current_row: list[str] | None = None
        self._current_cell: list[str] | None = None
        self._in_thead = False
        self._thead_row_count = 0

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        if tag == "table":
            self._table_depth += 1
            if self._table_depth == 1:
                self._collecting = True
                self._current_table = []
                self._thead_row_count = 0
        elif tag == "thead" and self._collecting:
            self._in_thead = True
        elif tag in ("tbody", "tfoot") and self._collecting:
            self._in_thead = False
        elif self._collecting and tag == "tr":
            self._current_row = []
        elif self._collecting and tag in ("td", "th"):
            self._current_cell = []

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("td", "th") and self._collecting and self._current_row is not None:
            text = "".join(self._current_cell or []).strip()
            self._current_row.append(text)
            self._current_cell = None
        elif tag == "tr" and self._collecting:
            if self._current_row is not None:
                if any(cell.strip() for cell in self._current_row):
                    self._current_table.append(self._current_row[:])
                    if self._in_thead:
                        self._thead_row_count += 1
                self._current_row = None
        elif tag == "thead":
            self._in_thead = False
        elif tag == "table":
            if self._table_depth == 1 and self._collecting and self._current_table is not None:
                self.tables.append(self._current_table[:])
                self.thead_counts.append(self._thead_row_count)
                self._collecting = False
                self._current_table = None
            self._table_depth = max(0, self._table_depth - 1)

    def handle_data(self, data: str):
        if self._collecting and self._current_cell is not None:
            self._current_cell.append(data)

    def first_table(self) -> list[list[str]]:
        return self.tables[0] if self.tables else []


def _table_score(table: list[list[str]]) -> int:
    if not table:
        return 0
    row_count = len(table)
    max_cols = max((len(row) for row in table), default=0)
    multi_col_rows = sum(1 for row in table if sum(1 for cell in row if cell) >= 2)
    return (multi_col_rows or row_count) * max(1, max_cols)


def extract_first_table(html_text: str) -> list[list[str]]:
    tables = extract_tables(html_text, max_tables=None)
    if not tables:
        return []
    best_table = tables[0]
    best_score = _table_score(best_table)
    for table in tables[1:]:
        score = _table_score(table)
        if score > best_score:
            best_table = table
            best_score = score
    return best_table


def extract_tables(html_text: str, *, max_tables: int | None = None) -> list[list[list[str]]]:
    parser = _SimpleTableParser()
    parser.feed(html_text or "")
    normalized_tables: list[list[list[str]]] = []
    for table in parser.tables:
        if max_tables is not None and len(normalized_tables) >= max_tables:
            break
        normalized: list[list[str]] = []
        for row in table:
            cleaned = [(cell or "").strip() for cell in row]
            if any(cell for cell in cleaned):
                normalized.append(cleaned)
        if normalized:
            normalized_tables.append(normalized)
    return normalized_tables


def extract_tables_with_header_counts(html_text: str) -> list[tuple[list[list[str]], int]]:
    """Return tables paired with their <thead> row counts."""
    parser = _SimpleTableParser()
    parser.feed(html_text or "")
    result: list[tuple[list[list[str]], int]] = []
    for i, table in enumerate(parser.tables):
        normalized: list[list[str]] = []
        thead_total = parser.thead_counts[i] if i < len(parser.thead_counts) else 0
        skipped_empty = 0
        for row_idx, row in enumerate(table):
            cleaned = [(cell or "").strip() for cell in row]
            if any(cell for cell in cleaned):
                normalized.append(cleaned)
            elif row_idx < thead_total:
                skipped_empty += 1
        if normalized:
            result.append((normalized, max(0, thead_total - skipped_empty)))
    return result


__all__ = ["extract_first_table", "extract_tables", "extract_tables_with_header_counts"]


# ======================================================================
# date_utils
# ======================================================================

from pathlib import Path
from typing import Callable, Tuple

from backend.app.repositories import get_loader


def get_col_type(db_path, table: str, col: str) -> str:
    """
    Return the inferred column type (uppercased) for table.col or '' when unavailable.
    Uses the shared DataFrame loader's dtype map instead of SQLite PRAGMA calls.
    """
    if not col or not table:
        return ""
    try:
        from backend.app.services.legacy_services import get_loader_for_ref
        loader = get_loader_for_ref(db_path)
        return (loader.column_type(table, col) or "").upper()
    except Exception:
        return ""


def mk_between_pred_for_date(col: str, col_type: str) -> Tuple[str, Callable[[str, str], tuple]]:
    """
    Returns (predicate_sql, adapter) used to build BETWEEN date filters.
    The adapter receives (start, end) and returns a tuple of parameters.
    When the column is missing or unusable, the predicate degenerates to '1=1'
    and the adapter returns an empty tuple – preserving the existing fail-open behaviour.
    """
    if not col or not col_type:
        return "1=1", lambda _s, _e: tuple()

    t = col_type.upper()
    if "INT" in t:
        predicate = (
            f"(CASE WHEN ABS({col}) > 32503680000 THEN {col}/1000 ELSE {col} END) "
            f"BETWEEN strftime('%s', ?) AND strftime('%s', ?)"
        )
        return predicate, lambda start, end: (start, end)

    predicate = f"datetime({col}) BETWEEN datetime(?) AND datetime(?)"
    return predicate, lambda start, end: (start, end)


__all__ = ["get_col_type", "mk_between_pred_for_date"]


# ======================================================================
# common_helpers
# ======================================================================

import re
from datetime import datetime, timezone
from typing import Callable, Iterable

_TOKEN_REGEX_CACHE: dict[str, re.Pattern[str]] = {}
_TR_BLOCK_RE = re.compile(r"(?is)<tr\b[^>]*>.*?</tr>")
_BATCH_BLOCK_ANY_TAG = re.compile(
    r"(?is)"
    r"<(?P<tag>section|div|article|main|tbody|tr)\b"
    r'[^>]*\bclass\s*=\s*["\'][^"\']*\bbatch-block\b[^"\']*["\']'
    r"[^>]*>"
    r"(?P<inner>.*?)"
    r"</(?P=tag)>"
)


def _token_regex(token: str) -> re.Pattern[str]:
    cleaned = (token or "").strip()
    if not cleaned:
        raise ValueError("Token must be a non-empty string")
    cached = _TOKEN_REGEX_CACHE.get(cleaned)
    if cached is None:
        cached = re.compile(rf"\{{\{{?\s*{re.escape(cleaned)}\s*\}}\}}?")
        _TOKEN_REGEX_CACHE[cleaned] = cached
    return cached


def _segment_has_any_token(segment: str, tokens: Iterable[str]) -> bool:
    for token in tokens:
        if not token:
            continue
        if _token_regex(token).search(segment):
            return True
    return False


def _find_rowish_block(html_text: str, row_tokens: Iterable[str]) -> tuple[str, int, int] | None:
    candidate_tokens = [tok for tok in row_tokens if isinstance(tok, str) and tok.strip()]
    if not candidate_tokens:
        return None

    matches = [m for m in _TR_BLOCK_RE.finditer(html_text) if _segment_has_any_token(m.group(0), candidate_tokens)]
    if not matches:
        return None

    prototype = matches[0].group(0).strip()
    start_index = matches[0].start()
    end_index = matches[-1].end()
    return prototype, start_index, end_index


def _find_or_infer_batch_block(html_text: str) -> tuple[str, str, str]:
    """
    Return (full_match, tag_name, inner_html) of the repeating unit.
    Preference order:
      1) Any element with class="batch-block"
      2) First <tr> inside the first <tbody>
      3) First row-like <div> (class includes row|item|card)
      4) First large container (<section|main|div|article> under <body>)
    """
    m = _BATCH_BLOCK_ANY_TAG.search(html_text)
    if m:
        return m.group(0), m.group("tag").lower(), m.group("inner")

    m_tbody = re.search(r"(?is)<tbody\b[^>]*>(?P<body>.*?)</tbody>", html_text)
    if m_tbody:
        tbody = m_tbody.group("body")
        m_tr = re.search(r"(?is)<tr\b[^>]*>(?P<tr>.*?)</tr>", tbody)
        if m_tr:
            return m_tr.group(0), "tr", m_tr.group("tr")

    m_div = re.search(r"(?is)<div\b[^>]*\b(row|item|card)\b[^>]*>(?P<inner>.*?)</div>", html_text)
    if m_div:
        return m_div.group(0), "div", m_div.group("inner")

    m_body = re.search(r"(?is)<body\b[^>]*>(?P<body>.*?)</body>", html_text)
    if m_body:
        body = m_body.group("body")
        m_cont = re.search(r"(?is)<(section|main|div|article)\b[^>]*>(?P<inner>.*?)</\1>", body)
        if m_cont:
            return m_cont.group(0), m_cont.group(1).lower(), m_cont.group("inner")

    raise RuntimeError("No explicit batch-block and no suitable repeating unit could be inferred.")


def _select_prototype_block(html_text: str, row_tokens: Iterable[str]) -> tuple[str, int, int]:
    # Priority 1: BLOCK_REPEAT comment markers (most reliable — handles nested elements)
    block_repeat_pat = re.compile(
        r"<!--\s*BEGIN:BLOCK_REPEAT[^>]*-->\s*(.*?)\s*<!--\s*END:BLOCK_REPEAT\s*-->",
        re.DOTALL | re.IGNORECASE,
    )
    repeat_matches = list(block_repeat_pat.finditer(html_text))
    if repeat_matches:
        chosen = repeat_matches[0]
        if row_tokens:
            for m in repeat_matches:
                if _segment_has_any_token(m.group(1), row_tokens):
                    chosen = m
                    break
        prototype = chosen.group(1).strip()
        start0 = repeat_matches[0].start()
        end_last = repeat_matches[-1].end()
        return prototype, start0, end_last

    # Priority 2: CSS class batch-block (fallback for templates without BLOCK_REPEAT markers)
    explicit_blocks = list(_BATCH_BLOCK_ANY_TAG.finditer(html_text))
    if explicit_blocks:
        chosen_match = explicit_blocks[0]
        if row_tokens:
            for match in explicit_blocks:
                if _segment_has_any_token(match.group(0), row_tokens):
                    chosen_match = match
                    break
        prototype = chosen_match.group(0).strip()
        start0 = explicit_blocks[0].start()
        end_last = explicit_blocks[-1].end()
        return prototype, start0, end_last

    rowish = _find_rowish_block(html_text, row_tokens)
    if rowish:
        return rowish

    block_full, tag_name, _ = _find_or_infer_batch_block(html_text)

    # --- Additive: for header-only templates (no row_tokens), use the full
    #     containing element (e.g. entire <tbody>) instead of just one <tr>.
    #     This keeps all header-token <tr> rows inside the prototype block
    #     so they get filled together rather than being left in the shell. ---
    row_token_list = list(row_tokens) if row_tokens else []
    if not row_token_list and tag_name == "tr":
        m_tbody = re.search(r"(?is)<tbody\b[^>]*>.*?</tbody>", html_text)
        if m_tbody:
            block_full = m_tbody.group(0)

    start0 = html_text.find(block_full)
    if start0 < 0:
        raise RuntimeError("Inferred batch block could not be located in HTML via .find()")
    end_last = start0 + len(block_full)
    return block_full.strip(), start0, end_last


def _strip_found_block(html_text: str, block_full: str, block_tag: str) -> str:
    """Remove the found/inferred block once (used to build shell)."""
    return html_text.replace(block_full, "", 1)


def html_without_batch_blocks(html_text: str) -> str:
    """Legacy stripper kept for compatibility."""
    pat = re.compile(r'(?is)\s*<section\s+class=["\']batch-block["\']\s*>.*?</section>\s*')
    return pat.sub("", html_text)


def _raise_no_block(html: str, cause: Exception | None = None) -> None:
    """Build a short <section ...> preview and raise ValueError from here."""
    sec_tags = re.findall(r"(?is)<section\b[^>]*>", html)
    preview_lines = []
    for i, t in enumerate(sec_tags[:12]):
        snip = t[:140].replace("\n", " ")
        preview_lines.append(f'{i+1:02d}: {snip}{" ..." if len(t) > 140 else ""}')
    preview = "\n".join(preview_lines)
    msg = (
        "Could not find any <section class='batch-block'> blocks and no suitable fallback could be inferred.\n"
        "First few <section> tags present:\n" + preview
    )
    raise ValueError(msg) from cause


def _parse_date_like(value) -> datetime | None:
    if value is None:
        return None
    val = str(value).strip()
    if not val:
        return None

    iso_try = val.replace("Z", "+00:00")
    if " " in iso_try and "T" not in iso_try:
        iso_try = iso_try.replace(" ", "T", 1)
    try:
        return datetime.fromisoformat(iso_try)
    except ValueError:
        pass

    if re.fullmatch(r"\d{10,}", val):
        try:
            seconds = int(val)
            if len(val) > 10:
                scale = 10 ** (len(val) - 10)
                return datetime.fromtimestamp(seconds / scale, tz=timezone.utc)
            return datetime.fromtimestamp(seconds, tz=timezone.utc)
        except ValueError:
            pass

    try:
        from email.utils import parsedate_to_datetime
    except ImportError:  # pragma: no cover
        parsedate_to_datetime = None  # type: ignore

    if parsedate_to_datetime is not None:
        try:
            dt = parsedate_to_datetime(val)
            if dt:
                return dt if dt.tzinfo is None else dt.astimezone()
        except (TypeError, ValueError):
            pass

    candidates = [
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%d/%m/%Y",
        "%m/%d/%Y",
        "%d-%m-%Y",
        "%m-%d-%Y",
        "%d.%m.%Y",
        "%d %b %Y",
        "%d %B %Y",
        "%b %d %Y",
        "%B %d %Y",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y/%m/%d %H:%M",
        "%Y/%m/%d %H:%M:%S",
        "%d/%m/%Y %H:%M",
        "%d/%m/%Y %H:%M:%S",
        "%m/%d/%Y %H:%M",
        "%m/%d/%Y %H:%M:%S",
        "%d-%m-%Y %H:%M",
        "%d-%m-%Y %H:%M:%S",
        "%d.%m.%Y %H:%M",
        "%d.%m.%Y %H:%M:%S",
        "%d %b %Y %H:%M",
        "%d %b %Y %H:%M:%S",
        "%d %B %Y %H:%M",
        "%d %B %Y %H:%M:%S",
        "%b %d %Y %H:%M",
        "%b %d %Y %H:%M:%S",
    ]
    for fmt in candidates:
        try:
            return datetime.strptime(val, fmt)
        except ValueError:
            continue
    return None


def _has_time_component(raw_value, dt_obj: datetime | None) -> bool:
    if dt_obj and (dt_obj.hour or dt_obj.minute or dt_obj.second or dt_obj.microsecond):
        return True
    if raw_value is None:
        return False
    text = str(raw_value)
    if re.search(r"\d{1,2}:\d{2}", text):
        return True
    if re.search(r"\b(am|pm)\b", text, flags=re.IGNORECASE):
        return True
    if "T" in text or "t" in text:
        return True
    return False


def _format_for_token(token: str, dt_obj: datetime | None, include_time_default: bool = False) -> str:
    if not dt_obj:
        return ""

    token_lower = token.lower()
    token_clean = re.sub(r"[^a-z0-9]", "", token_lower)

    def _has(*needles: str) -> bool:
        return any(needle in token_clean for needle in needles)

    include_time = include_time_default or _has("time", "clock", "datetime", "timestamp")
    include_seconds = _has("second", "seconds", "sec", "timestamp", "precise", "fulltime")
    use_ampm = _has("ampm", "12h", "twelvehour")
    if include_seconds and not include_time:
        include_time = True
    if use_ampm and not include_time:
        include_time = True

    include_timezone = _has("timezone", "tz", "utc", "offset", "gmtoffset", "withtz", "withzone", "zulu")
    iso_like = _has("iso", "iso8601", "ymd", "rfc3339")
    rfc822_like = _has("rfc2822", "rfc822")
    http_like = _has("httpdate", "rfc7231")
    compact_like = _has("compact", "slug", "filename", "filestamp", "yyyymmdd", "numeric", "digits")
    us_like = _has("us", "usa", "mdy", "mmdd")
    dashed_like = _has("dash", "hyphen")
    long_like = _has("long", "verbose", "friendly", "pretty", "human")
    short_like = _has("short", "abbr", "mini", "brief")
    month_long_like = _has("monthname", "monthlong")
    month_short_like = _has("monthabbr", "monthshort")
    weekday_like = _has("weekday", "dayname")
    weekday_short = _has("weekdayshort", "weekdayabbr", "daynameshort")
    epoch_ms_like = _has("epochms", "millis", "milliseconds", "unixms")
    epoch_like = _has("epoch", "unixtime", "unix")

    dt_for_format = dt_obj
    if include_timezone and dt_for_format.tzinfo is None:
        try:
            dt_for_format = dt_for_format.astimezone()
        except ValueError:
            pass

    if epoch_ms_like:
        try:
            return str(int(dt_for_format.timestamp() * 1000))
        except (OSError, OverflowError, ValueError):
            pass
    if epoch_like:
        try:
            return str(int(dt_for_format.timestamp()))
        except (OSError, OverflowError, ValueError):
            pass

    if rfc822_like or http_like:
        try:
            from email.utils import format_datetime as _email_format_datetime

            base_dt = dt_for_format
            if base_dt.tzinfo is None:
                base_dt = base_dt.astimezone()
            return _email_format_datetime(base_dt)
        except Exception:
            pass

    if iso_like:
        dt_use = dt_for_format
        if include_time:
            timespec = "seconds" if include_seconds else "minutes"
            try:
                return dt_use.isoformat(timespec=timespec)
            except TypeError:
                return dt_use.isoformat()
        return dt_use.date().isoformat()

    if compact_like:
        date_part = dt_for_format.strftime("%Y%m%d")
        if include_time:
            if use_ampm:
                time_fmt = "%I%M%S%p" if include_seconds else "%I%M%p"
            else:
                time_fmt = "%H%M%S" if include_seconds else "%H%M"
            date_part = f"{date_part}_{dt_for_format.strftime(time_fmt)}"
        if include_timezone:
            tz = dt_for_format.strftime("%z")
            if tz:
                date_part = f"{date_part}{tz}"
        return date_part

    date_part = "%d/%m/%Y"
    if us_like:
        date_part = "%m/%d/%Y"
    elif dashed_like:
        date_part = "%d-%m-%Y"
    elif long_like:
        date_part = "%B %d, %Y"
    elif short_like or month_short_like:
        date_part = "%d %b %Y"
    elif month_long_like:
        date_part = "%d %B %Y"

    if weekday_like:
        prefix = "%a, " if weekday_short else "%A, "
        date_part = prefix + date_part

    fmt = date_part
    if include_time:
        if use_ampm:
            time_fmt = "%I:%M:%S %p" if include_seconds else "%I:%M %p"
        else:
            time_fmt = "%H:%M:%S" if include_seconds else "%H:%M"
        fmt = f"{fmt} {time_fmt}"
    if include_timezone:
        fmt = f"{fmt} %Z".strip()

    try:
        rendered = dt_for_format.strftime(fmt).strip()
        if not rendered and "%Z" in fmt:
            rendered = dt_for_format.strftime(fmt.replace("%Z", "%z")).strip()
        return rendered
    except Exception:
        if include_time:
            try:
                return dt_for_format.isoformat(timespec="seconds")
            except TypeError:
                return dt_for_format.isoformat()
        return dt_for_format.date().isoformat()


STYLE_OR_SCRIPT_RE = re.compile(r"(?is)(<style\b[^>]*>.*?</style>|<script\b[^>]*>.*?</script>)")


def _apply_outside_styles_scripts(html_in: str, transform_fn: Callable[[str], str]) -> str:
    parts = STYLE_OR_SCRIPT_RE.split(html_in)
    for i in range(len(parts)):
        if i % 2 == 0:
            parts[i] = transform_fn(parts[i])
    return "".join(parts)


def _sub_token_text(text: str, token: str, val: str) -> str:
    pat = re.compile(r"(\{\{\s*" + re.escape(token) + r"\s*\}\}|\{\s*" + re.escape(token) + r"\s*\})")
    return pat.sub(val, text)


def sub_token(html_in: str, token: str, val: str) -> str:
    return _apply_outside_styles_scripts(html_in, lambda txt: _sub_token_text(txt, token, val))


def _blank_known_tokens_text(text: str, tokens) -> str:
    for t in tokens:
        text = re.sub(r"\{\{\s*" + re.escape(t) + r"\s*\}\}", "", text)
        text = re.sub(r"\{\s*" + re.escape(t) + r"\s*\}", "", text)
    return text


def blank_known_tokens(html_in: str, tokens) -> str:
    return _apply_outside_styles_scripts(html_in, lambda txt: _blank_known_tokens_text(txt, tokens))


def _convert_css_length_to_mm(raw: str) -> float | None:
    if not raw:
        return None
    text = raw.strip().lower()
    if not text or text == "auto":
        return None
    m = re.match(r"([-+]?\d*\.?\d+)\s*(mm|cm|in|pt|pc|px)?", text)
    if not m:
        return None
    value = float(m.group(1))
    unit = m.group(2) or "px"
    if unit == "mm":
        return value
    if unit == "cm":
        return value * 10.0
    if unit in {"in", "inch", "inches"}:
        return value * 25.4
    if unit == "pt":
        return value * (25.4 / 72.0)
    if unit == "pc":
        return value * (25.4 / 6.0)
    if unit == "px":
        return value * (25.4 / 96.0)
    return None


def _parse_page_size_value(value: str) -> tuple[float, float] | None:
    if not value:
        return None
    text = value.strip().lower()
    if not text:
        return None
    size_map = {
        "a0": (841.0, 1189.0),
        "a1": (594.0, 841.0),
        "a2": (420.0, 594.0),
        "a3": (297.0, 420.0),
        "a4": (210.0, 297.0),
        "a5": (148.0, 210.0),
        "letter": (215.9, 279.4),
        "legal": (215.9, 355.6),
        "tabloid": (279.4, 431.8),
    }
    orientation = None
    tokens = [t for t in re.split(r"\s+", text) if t]
    size_tokens = tokens
    if tokens and tokens[-1] in {"portrait", "landscape"}:
        orientation = tokens[-1]
        size_tokens = tokens[:-1]
    if len(size_tokens) == 1 and size_tokens[0] in size_map:
        width_mm, height_mm = size_map[size_tokens[0]]
    elif len(size_tokens) >= 2:
        first = _convert_css_length_to_mm(size_tokens[0])
        second = _convert_css_length_to_mm(size_tokens[1])
        if first is None or second is None:
            return None
        width_mm, height_mm = first, second
    else:
        return None
    if orientation == "landscape":
        width_mm, height_mm = height_mm, width_mm
    return width_mm, height_mm


def _parse_margin_shorthand(value: str) -> tuple[float | None, float | None, float | None, float | None]:
    parts = [p for p in re.split(r"\s+", value.strip()) if p]
    values = [_convert_css_length_to_mm(p) for p in parts]
    if not values:
        return None, None, None, None
    if len(values) == 1:
        top = right = bottom = left = values[0]
    elif len(values) == 2:
        top = bottom = values[0]
        right = left = values[1]
    elif len(values) == 3:
        top = values[0]
        right = left = values[1]
        bottom = values[2]
    else:
        top, right, bottom, left = values[:4]
    return top, right, bottom, left


def _extract_page_metrics(html_in: str) -> dict[str, float]:
    default_width_mm, default_height_mm = 210.0, 297.0
    margin_top_mm = 0.0
    margin_bottom_mm = 0.0
    page_match = re.search(r"@page\b[^{}]*\{(?P<body>.*?)\}", html_in, re.IGNORECASE | re.DOTALL)
    if page_match:
        block = page_match.group("body")
        size_match = re.search(r"size\s*:\s*([^;]+);?", block, re.IGNORECASE)
        if size_match:
            parsed_size = _parse_page_size_value(size_match.group(1))
            if parsed_size:
                default_width_mm, default_height_mm = parsed_size
        margin_match = re.search(r"margin\s*:\s*([^;]+);?", block, re.IGNORECASE)
        if margin_match:
            mt, _, mb, _ = _parse_margin_shorthand(margin_match.group(1))
            if mt is not None:
                margin_top_mm = mt
            if mb is not None:
                margin_bottom_mm = mb
        for name, setter in (("margin-top", "top"), ("margin-bottom", "bottom")):
            specific = re.search(rf"{name}\s*:\s*([^;]+);?", block, re.IGNORECASE)
            if specific:
                as_mm = _convert_css_length_to_mm(specific.group(1))
                if as_mm is None:
                    continue
                if setter == "top":
                    margin_top_mm = as_mm
                else:
                    margin_bottom_mm = as_mm
    return {
        "page_width_mm": default_width_mm,
        "page_height_mm": default_height_mm,
        "margin_top_mm": max(margin_top_mm, 0.0),
        "margin_bottom_mm": max(margin_bottom_mm, 0.0),
    }


# ---------------------------------------------------------------------------
# Additive: date column auto-detection for contracts missing date_columns
# ---------------------------------------------------------------------------

def detect_date_column(db_path, table_name: str) -> str | None:
    """Auto-detect a likely date/timestamp column in *table_name*.

    Strategy:
      1. Name-based: columns containing 'date', 'timestamp', '_dt', '_ts'.
         Prefer 'date' over 'time' when multiple match.
      2. Value-based: sample up to 5 non-null TEXT values and check if they
         parse as dates via ``_parse_date_like``.
      3. Return the column name, or ``None`` if nothing matches.

    This is a safety-net fallback — contracts should specify date_columns
    explicitly whenever possible.
    """
    if not table_name:
        return None

    return _detect_date_column_df(db_path, table_name)


def _detect_date_column_df(db_path, table_name: str) -> str | None:
    """DataFrame-based date column detection — no SQL."""
    from pathlib import Path as _Path

    try:
        from backend.app.services.legacy_services import get_loader_for_ref
        loader = get_loader_for_ref(db_path)
        df = loader.frame(table_name)
    except Exception:
        return None

    if df is None or df.empty:
        return None

    _DATE_PATTERNS = ("date", "timestamp", "_dt", "_ts", "time")
    _PREFER_PATTERNS = ("date",)

    candidates: list[tuple[int, str]] = []
    for col in df.columns:
        col_lower = str(col).lower()
        if any(pat in col_lower for pat in _DATE_PATTERNS):
            priority = 0 if any(p in col_lower for p in _PREFER_PATTERNS) else 1
            candidates.append((priority, str(col)))

    if candidates:
        candidates.sort(key=lambda x: x[0])
        return candidates[0][1]

    # Strategy 2: value-based detection
    for col in df.columns:
        col_type = loader.column_type(table_name, str(col))
        if col_type in ("INTEGER", "REAL"):
            continue
        non_null = df[col].dropna()
        if non_null.empty:
            continue
        values = [str(v).strip() for v in non_null.head(5) if str(v).strip()]
        if not values:
            continue
        date_hits = sum(1 for v in values if _parse_date_like(v) is not None)
        if date_hits >= 3 or (date_hits == len(values) and len(values) >= 1):
            return str(col)

    return None


# ======================================================================
# contract_adapter
# ======================================================================

import logging
import re
from dataclasses import dataclass
from decimal import ROUND_HALF_UP, Decimal, InvalidOperation
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence, Tuple

logger = logging.getLogger(__name__)

_PARAM_RE = re.compile(r"PARAM:([A-Za-z0-9_]+)")
_DIRECT_COLUMN_RE = re.compile(r"^\s*(?P<table>[A-Za-z_][\w]*)\s*\.\s*(?P<column>[A-Za-z_][\w]*)\s*$")


def _ensure_mapping(value: Any) -> Dict[str, str]:
    if not isinstance(value, Mapping):
        return {}
    result: Dict[str, str] = {}
    for key, expr in value.items():
        if key is None:
            continue
        key_text = str(key).strip()
        if not key_text:
            continue
        expr_text = "" if expr is None else str(expr).strip()
        if not expr_text:
            continue
        result[key_text] = expr_text
    return result


def _ensure_mapping_mixed(value: Any) -> Dict[str, Any]:
    """Like _ensure_mapping but preserves dict/object values (for declarative ops)."""
    if not isinstance(value, Mapping):
        return {}
    result: Dict[str, Any] = {}
    for key, expr in value.items():
        if key is None:
            continue
        key_text = str(key).strip()
        if not key_text:
            continue
        if expr is None:
            continue
        # Preserve dicts (declarative ops) and numeric values as-is
        if isinstance(expr, (dict, int, float)):
            result[key_text] = expr
        else:
            expr_text = str(expr).strip()
            if expr_text:
                result[key_text] = expr_text
    return result


def _ensure_sequence(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, Iterable):
        return [str(item) for item in value]
    return [str(value)]


def format_decimal_str(value: Any, max_decimals: int = 3) -> str:
    """
    Format numeric values with rounding up to max_decimals and trim trailing zeros.
    Non-numeric values are returned as-is (converted to string).
    """
    if value is None:
        return ""

    decimal_value: Optional[Decimal] = None
    if isinstance(value, Decimal):
        decimal_value = value
    else:
        text = str(value).strip()
        if not text:
            return ""
        try:
            decimal_value = Decimal(text)
        except (InvalidOperation, ValueError):
            return str(value)

    if not decimal_value.is_finite():
        logger.warning(
            "format_decimal_non_finite value=%s coerced_to=0",
            value,
            extra={"event": "format_decimal_non_finite", "original": str(value)},
        )
        return "0"

    if max_decimals <= 0:
        rounded = decimal_value.quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    else:
        quantizer = Decimal("1").scaleb(-max_decimals)
        rounded = decimal_value.quantize(quantizer, rounding=ROUND_HALF_UP)

    formatted = format(rounded, "f")
    if "." in formatted:
        formatted = formatted.rstrip("0").rstrip(".")
    if formatted == "-0":
        formatted = "0"
    return formatted


def format_fixed_decimals(value: Any, decimals: int, max_decimals: int = 3) -> str:
    decimals = max(0, min(decimals, max_decimals))
    try:
        number = Decimal(str(value))
    except (InvalidOperation, ValueError, TypeError):
        return str(value)
    if not number.is_finite():
        logger.debug(
            "format_fixed_decimals_non_finite value=%s coerced_to=0",
            value,
            extra={"event": "format_fixed_decimals_non_finite", "original": str(value)},
        )
        number = Decimal(0)
    quantizer = Decimal("1").scaleb(-decimals) if decimals else Decimal("1")
    rounded = number.quantize(quantizer, rounding=ROUND_HALF_UP)
    if rounded == 0:
        rounded = Decimal(0).quantize(quantizer, rounding=ROUND_HALF_UP) if decimals else Decimal(0)
    formatted = format(rounded, "f")
    if rounded == 0 and formatted.startswith("-"):
        formatted = formatted[1:]
    return formatted


@dataclass(frozen=True)
class FormatterSpec:
    kind: str
    arg: Optional[str] = None

    @staticmethod
    def parse(raw: str | None) -> Optional["FormatterSpec"]:
        if not raw:
            return None
        text = raw.strip()
        if not text:
            return None
        m = re.match(r"([a-zA-Z0-9_]+)(?:\((.*)\))?$", text)
        if not m:
            return None
        kind = m.group(1).lower()
        arg = m.group(2)
        return FormatterSpec(kind=kind, arg=arg)


class ContractAdapter:
    """
    Convenience wrapper around a Step-5 contract payload.

    Exposes normalised accessors and helper methods that the discovery and
    report generation flows can rely on without duplicating parsing logic.
    """

    def __init__(self, contract: Mapping[str, Any] | None):
        self._raw = contract or {}

        tokens = self._raw.get("tokens") or {}
        self._scalar_tokens = _ensure_sequence(tokens.get("scalars"))
        self._row_tokens = _ensure_sequence(tokens.get("row_tokens"))
        self._total_tokens = _ensure_sequence(tokens.get("totals"))

        join_block = self._raw.get("join") or {}
        self._parent_table = str(join_block.get("parent_table") or "").strip()
        self._child_table = str(join_block.get("child_table") or "").strip()
        self._parent_key = join_block.get("parent_key")
        self._child_key = join_block.get("child_key")

        self._date_columns = _ensure_mapping(self._raw.get("date_columns"))

        filters = self._raw.get("filters") or {}
        self._required_filters = _ensure_mapping(filters.get("required"))
        self._optional_filters = _ensure_mapping(filters.get("optional"))

        self._pre_aggregate = self._raw.get("pre_aggregate") or {}
        self._group_aggregate = self._raw.get("group_aggregate") or {}
        self._post_aggregate = self._raw.get("post_aggregate") or {}
        self._reshape_rules = self._raw.get("reshape_rules") or []
        self._row_computed = _ensure_mapping_mixed(self._raw.get("row_computed"))
        self._totals_math = _ensure_mapping_mixed(self._raw.get("totals_math"))
        self._formatters_raw = _ensure_mapping(self._raw.get("formatters"))

        order_by_block = self._raw.get("order_by") or {}
        self._order_by_rows = _ensure_sequence(order_by_block.get("rows"))
        self._row_order = _ensure_sequence(self._raw.get("row_order"))

        self._totals_mapping = _ensure_mapping_mixed(self._raw.get("totals"))
        self._mapping = _ensure_mapping(self._raw.get("mapping"))
        if not self._parent_table:
            inferred_parent = self._infer_parent_table(self._mapping)
            if inferred_parent:
                self._parent_table = inferred_parent

        self._param_tokens = self._discover_param_tokens()
        self._formatter_cache: Dict[str, FormatterSpec | None] = {}

    # ------------------------------------------------------------------ #
    # Basic properties
    # ------------------------------------------------------------------ #
    @property
    def mapping(self) -> Dict[str, str]:
        return dict(self._mapping)

    @property
    def scalar_tokens(self) -> List[str]:
        return list(self._scalar_tokens)

    @property
    def row_tokens(self) -> List[str]:
        return list(self._row_tokens)

    @property
    def total_tokens(self) -> List[str]:
        return list(self._total_tokens)

    @property
    def parent_table(self) -> str:
        return self._parent_table

    @property
    def child_table(self) -> str:
        return self._child_table

    @property
    def parent_key(self) -> Any:
        return self._parent_key

    @property
    def child_key(self) -> Any:
        return self._child_key

    @property
    def date_columns(self) -> Dict[str, str]:
        return dict(self._date_columns)

    @property
    def required_filters(self) -> Dict[str, str]:
        return dict(self._required_filters)

    @property
    def optional_filters(self) -> Dict[str, str]:
        return dict(self._optional_filters)

    @property
    def reshape_rules(self) -> Sequence[Mapping[str, Any]]:
        return list(self._reshape_rules)

    @property
    def row_computed(self) -> Dict[str, str]:
        return dict(self._row_computed)

    @property
    def totals_math(self) -> Dict[str, str]:
        return dict(self._totals_math)

    @property
    def formatters(self) -> Dict[str, str]:
        return dict(self._formatters_raw)

    @property
    def order_by_rows(self) -> List[str]:
        return list(self._order_by_rows)

    @property
    def row_order(self) -> List[str]:
        return list(self._row_order)

    @property
    def totals_mapping(self) -> Dict[str, str]:
        return dict(self._totals_mapping)

    @property
    def param_tokens(self) -> List[str]:
        return sorted(self._param_tokens)

    # ------------------------------------------------------------------ #
    # Formatting helpers
    # ------------------------------------------------------------------ #
    def get_formatter_spec(self, token: str) -> Optional[FormatterSpec]:
        if token in self._formatter_cache:
            return self._formatter_cache[token]
        raw = self._formatters_raw.get(token)
        spec = FormatterSpec.parse(raw)
        self._formatter_cache[token] = spec
        return spec

    def format_value(self, token: str, value: Any) -> str:
        spec = self.get_formatter_spec(token)
        if spec is None:
            if value is None:
                return ""
            if isinstance(value, (int, float, Decimal)):
                return format_decimal_str(value)
            if isinstance(value, str):
                candidate = value.strip()
                if not candidate:
                    return ""
                lowered = candidate.lower()
                if "." in candidate or "e" in lowered:
                    return format_decimal_str(candidate)
                return value
            return str(value)
        kind = spec.kind
        if value is None:
            return ""

        if kind == "number":
            decimals = 0
            if spec.arg:
                try:
                    decimals = int(spec.arg.strip())
                except ValueError:
                    decimals = 0
            return format_fixed_decimals(value, decimals, max_decimals=3)

        if kind == "percent":
            decimals = 0
            if spec.arg:
                try:
                    decimals = int(spec.arg.strip())
                except ValueError:
                    decimals = 0
            try:
                text = str(value).strip()
                if text.endswith("%"):
                    text = text[:-1].strip()
                number = Decimal(text)
            except (InvalidOperation, ValueError, TypeError):
                return str(value)
            if not number.is_finite():
                number = Decimal(0)
            if abs(number) <= 1:
                number *= Decimal(100)
            formatted = format_fixed_decimals(number, decimals, max_decimals=3)
            return f"{formatted}%"

        if kind == "date":
            fmt = (spec.arg or "YYYY-MM-DD").strip()
            return self._format_date_like(value, fmt)

        return str(value)

    @staticmethod
    def _format_date_like(value: Any, fmt: str) -> str:
        from datetime import datetime

        if value is None:
            return ""
        text = str(value).strip()
        if not text:
            return ""
        dt: Optional[datetime] = None
        for pattern in (
            "%Y-%m-%d",
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%dT%H:%M:%S",
            "%d/%m/%Y",
            "%m/%d/%Y",
            "%d-%m-%Y %H:%M:%S",
            "%d-%m-%Y",
        ):
            try:
                dt = datetime.strptime(text, pattern)
                break
            except ValueError:
                continue
        if dt is None:
            try:
                dt = datetime.fromisoformat(text.replace("Z", "+00:00"))
            except ValueError:
                return text

        fmt_map = {
            # Date only
            "DD/MM/YYYY": "%d/%m/%Y",
            "YYYY-MM-DD": "%Y-%m-%d",
            "DD-MM-YYYY": "%d-%m-%Y",
            "MM/DD/YYYY": "%m/%d/%Y",
            # Date + time
            "YYYY-MM-DD HH:MM:SS": "%Y-%m-%d %H:%M:%S",
            "YYYY-MM-DD HH:MM": "%Y-%m-%d %H:%M",
            "DD/MM/YYYY HH:MM:SS": "%d/%m/%Y %H:%M:%S",
            "DD/MM/YYYY HH:MM": "%d/%m/%Y %H:%M",
            "MM/DD/YYYY HH:MM:SS": "%m/%d/%Y %H:%M:%S",
            "DD-MM-YYYY HH:MM:SS": "%d-%m-%Y %H:%M:%S",
            "DD-MM-YYYY HH:MM": "%d-%m-%Y %H:%M",
            # ISO 8601
            "ISO": "%Y-%m-%dT%H:%M:%S",
            "ISO8601": "%Y-%m-%dT%H:%M:%S",
        }
        fallback = "%Y-%m-%d %H:%M:%S" if "HH" in fmt.upper() else "%Y-%m-%d"
        pattern = fmt_map.get(fmt.upper(), fallback)
        return dt.strftime(pattern)

    @staticmethod
    def _infer_parent_table(mapping: Mapping[str, str]) -> Optional[str]:
        for expr in mapping.values():
            if not isinstance(expr, str):
                continue
            match = _DIRECT_COLUMN_RE.match(expr.strip())
            if not match:
                continue
            table_name = match.group("table").strip(' "`[]')
            if table_name and not table_name.lower().startswith("params"):
                return table_name
        return None

    # ------------------------------------------------------------------ #
    # DataFrame resolve methods
    # ------------------------------------------------------------------ #

    def _resolve_mapping_column(self, token: str) -> Optional[Tuple[str, str]]:
        """Return (table, column) from a token's mapping, or None if not a direct ref.

        Mappings like ``params.report_date`` are treated as parameter
        references (not real table.column lookups) and return ``None``.
        """
        expr = self._mapping.get(token, "")
        m = _DIRECT_COLUMN_RE.match(expr)
        if m:
            table = m.group("table")
            # params.xxx is a parameter reference, not a table.column
            if table.lower() == "params":
                return None
            return table, m.group("column")
        return None

    def _apply_date_filter_df(self, df, table: str, start_date: str | None, end_date: str | None):
        """Apply date range filter to a DataFrame using contract date_columns."""
        import pandas as pd

        date_col = self._date_columns.get(table.lower()) or self._date_columns.get(table)
        if not date_col or date_col not in df.columns:
            return df
        if not start_date and not end_date:
            return df
        start_dt = _parse_date_like(start_date) if start_date else None
        end_dt = _parse_date_like(end_date) if end_date else None
        if start_dt is None and end_dt is None:
            return df
        if end_dt is not None:
            end_dt = _snap_end_of_day(end_dt)

        dt_series = _coerce_datetime_series(df[date_col])
        mask = pd.Series(True, index=df.index)
        if start_dt:
            mask = mask & (dt_series >= start_dt)
        if end_dt:
            mask = mask & (dt_series <= end_dt)
        return df.loc[mask.fillna(False)]

    @staticmethod
    def _normalize_numeric_str(s: str) -> str:
        """Normalize numeric strings: '1.0' → '1', '2.00' → '2', 'abc' → 'abc'."""
        try:
            f = float(s)
            if f == int(f):
                return str(int(f))
            return str(f)
        except (ValueError, TypeError):
            return s

    def _apply_value_filters_df(self, df, value_filters: Dict[str, list]):
        """Apply equality filters from contract optional_filters."""
        import pandas as pd

        if not value_filters:
            return df
        mask = pd.Series(True, index=df.index)
        for filter_key, filter_values in value_filters.items():
            col_expr = self._optional_filters.get(filter_key, "")
            m = _DIRECT_COLUMN_RE.match(col_expr)
            col_name = m.group("column") if m else filter_key
            if col_name not in df.columns:
                continue
            normalized = [str(v).strip() for v in filter_values if str(v or "").strip()]
            if not normalized:
                continue
            # Normalize both sides for numeric comparison (e.g. "1" matches "1.0")
            norm_set = set(normalized) | {self._normalize_numeric_str(v) for v in normalized}
            series = df[col_name].astype(str).str.strip()
            norm_series = series.map(self._normalize_numeric_str)
            mask = mask & (series.isin(norm_set) | norm_series.isin(norm_set))
        return df.loc[mask.fillna(False)]

    def _apply_pre_aggregate_df(self, df):
        """Collapse time-series into one row per batch (first_per_run strategy).

        Reads ``pre_aggregate`` from the contract:
          batch_column   – column that identifies the batch (e.g. OIL_BACTH_COUNT)
          timestamp_column – used for ordering within each batch
          strategy       – currently only "first_per_run"
          skip_value     – batch_column value to exclude (e.g. 0)
        """
        pa = self._pre_aggregate
        batch_col = pa.get("batch_column", "")
        ts_col = pa.get("timestamp_column", "")
        strategy = pa.get("strategy", "")
        skip_value = pa.get("skip_value")

        if not batch_col or batch_col not in df.columns:
            return df
        if df.empty:
            return df

        # Filter out rows matching skip_value
        if skip_value is not None:
            mask = df[batch_col].astype(str).str.strip() != str(skip_value).strip()
            df = df.loc[mask]
            if df.empty:
                return df

        if strategy == "first_per_run":
            # Sort by timestamp if available, then keep first row per batch
            if ts_col and ts_col in df.columns:
                df = df.sort_values(ts_col, ascending=True)
                # Capture last timestamp per batch as end_timestamp_utc
                # (before dedup removes them)
                end_ts = df.groupby(batch_col, sort=False)[ts_col].transform("last")
                df = df.copy()
                df["end_timestamp_utc"] = end_ts
            df = df.drop_duplicates(subset=[batch_col], keep="first")

        logger.info("pre_aggregate applied: %s → %d rows", strategy, len(df))
        return df

    def _apply_group_aggregate_df(self, df, override_config=None):
        """Aggregate across all batches (e.g. sum) to collapse N batch rows → 1 row.

        Reads ``group_aggregate`` from the contract (or override_config):
          strategy – "sum" (only supported strategy currently)
          columns  – list of columns to aggregate
        Non-aggregated columns keep first row values.
        """
        import pandas as pd

        ga = override_config if override_config is not None else self._group_aggregate
        strategy = ga.get("strategy", "")
        agg_columns = ga.get("columns", [])

        if not strategy or df.empty:
            return df

        if strategy != "group_by" and not agg_columns:
            return df

        if strategy == "sum":
            # Resolve actual column names in the DataFrame
            agg_map = {}
            for col in agg_columns:
                actual = self._resolve_df_col(df, col)
                if actual:
                    agg_map[actual] = "sum"

            if not agg_map:
                return df

            # Coerce aggregation columns to numeric
            for col in agg_map:
                df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

            # Build aggregation dict: sum for specified cols, first for everything else
            full_agg = {}
            for col in df.columns:
                if col in agg_map:
                    full_agg[col] = "sum"
                else:
                    full_agg[col] = "first"

            result = df.groupby(lambda _: 0, sort=False).agg(full_agg)
            result = result.reset_index(drop=True)
            logger.info("group_aggregate applied: %s → %d rows (from %d)", strategy, len(result), len(df))
            return result

        if strategy == "group_by":
            # Group by specified column(s), aggregate others, add count column.
            # Contract: group_columns (list), aggregations (dict col→func), count_as (str)
            group_columns = ga.get("group_columns", [])
            aggregations = ga.get("aggregations", {})
            count_as = ga.get("count_as", "__count__")

            if not group_columns:
                logger.warning("group_by strategy missing group_columns, skipping")
                return df

            # Resolve actual column names
            resolved_groups = []
            for gc in group_columns:
                actual = self._resolve_df_col(df, gc)
                if actual:
                    resolved_groups.append(actual)
                elif gc in df.columns:
                    resolved_groups.append(gc)

            if not resolved_groups:
                logger.warning("group_by: none of %s found in DataFrame", group_columns)
                return df

            # Build aggregation map
            agg_map = {}
            for col, func in aggregations.items():
                actual = self._resolve_df_col(df, col) or col
                if actual in df.columns:
                    if func == "min":
                        df[actual] = pd.to_numeric(df[actual], errors="coerce")
                    agg_map[actual] = func

            # Default: first for all non-group columns not in aggregations
            full_agg = {}
            for col in df.columns:
                if col in resolved_groups:
                    continue
                if col in agg_map:
                    full_agg[col] = agg_map[col]
                else:
                    full_agg[col] = "first"

            result = df.groupby(resolved_groups, sort=True).agg(full_agg)

            # Add count column
            counts = df.groupby(resolved_groups, sort=True).size()
            result[count_as] = counts.values

            result = result.reset_index()
            logger.info("group_aggregate applied: %s → %d rows (from %d)", strategy, len(result), len(df))
            return result

        logger.warning("group_aggregate strategy %r not supported, skipping", strategy)
        return df

    @staticmethod
    def _resolve_df_col(df, col: str) -> str | None:
        """Resolve a column name against a DataFrame.

        Tries in order:
        1. Exact match
        2. Strip table prefix (table.column → column)
        3. Case-insensitive match
        4. Suffix-stripped match (row_ach_wt_kg ↔ row_ach_wt when _kg suffix differs)

        Returns the actual DataFrame column name, or None if not found.
        """
        # 1. Exact match
        if col in df.columns:
            return col
        # 2. Strip table prefix
        if "." in col:
            col = col.rsplit(".", 1)[1]
            if col in df.columns:
                return col
        # 3. Case-insensitive
        col_lower = col.lower()
        for actual in df.columns:
            if isinstance(actual, str) and actual.lower() == col_lower:
                return actual
        # 4. Suffix-stripped match (defense-in-depth for MELT alias mismatches)
        _COMMON_SUFFIXES = ("_kg", "_pct", "_sec", "_wt", "_amt", "_count", "_num")
        for suffix in _COMMON_SUFFIXES:
            if col_lower.endswith(suffix):
                stripped = col_lower[:-len(suffix)]
                for actual in df.columns:
                    if isinstance(actual, str) and actual.lower() == stripped:
                        return actual
            # Also try adding suffix to match longer column names
            for actual in df.columns:
                if isinstance(actual, str) and actual.lower().endswith(suffix):
                    actual_stripped = actual.lower()[:-len(suffix)]
                    if actual_stripped == col_lower:
                        return actual
        return None

    @staticmethod
    def _coerce_numeric(val):
        """Coerce a value or Series to numeric for arithmetic ops."""
        import pandas as pd
        if isinstance(val, pd.Series):
            return pd.to_numeric(val, errors="coerce").fillna(0)
        if isinstance(val, str):
            try:
                return float(val)
            except (ValueError, TypeError):
                return 0
        return val

    def _apply_declarative_op(self, df, op_spec) -> Any:
        """Interpret a declarative operation spec and return computed result.

        Supports both new-style dict ops and legacy SQL expression strings.
        """
        import pandas as pd

        if isinstance(op_spec, str):
            return self._interpret_legacy_computed(df, op_spec)

        if not isinstance(op_spec, dict):
            return None

        op = op_spec.get("op", "").lower()
        if op == "subtract":
            left = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("left", 0)))
            right = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("right", 0)))
            if left is None or right is None:
                return None
            return left - right
        elif op == "add":
            left = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("left", 0)))
            right = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("right", 0)))
            if left is None or right is None:
                return None
            return left + right
        elif op == "multiply":
            left = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("left", 0)))
            right = self._coerce_numeric(self._resolve_agg_or_col(df, op_spec.get("right", 0)))
            if left is None or right is None:
                return None
            return left * right
        elif op == "divide":
            num_spec = op_spec.get("numerator", op_spec.get("left", ""))
            den_spec = op_spec.get("denominator", op_spec.get("right", ""))
            num = self._coerce_numeric(self._resolve_agg_or_col(df, num_spec))
            den = self._coerce_numeric(self._resolve_agg_or_col(df, den_spec))
            if isinstance(den, (int, float)) and den == 0:
                return None
            if isinstance(num, pd.Series) and isinstance(den, pd.Series):
                return num / den.replace(0, float("nan"))
            return num / den if den else None
        elif op == "sum":
            col = op_spec.get("column", "")
            resolved = self._resolve_df_col(df, col)
            if resolved:
                return df[resolved].sum()
            return 0
        elif op == "mean":
            col = op_spec.get("column", "")
            resolved = self._resolve_df_col(df, col)
            if resolved:
                return df[resolved].mean()
            return 0
        elif op == "add_many":
            # Sum multiple columns row-wise: columns: ["col1", "col2", ...]
            cols = op_spec.get("columns", [])
            total = None
            for c in cols:
                rc = self._resolve_df_col(df, c)
                if rc:
                    series = pd.to_numeric(df[rc], errors="coerce").fillna(0)
                    total = series if total is None else total + series
            return total if total is not None else 0
        elif op == "count":
            col = op_spec.get("column", "")
            resolved = self._resolve_df_col(df, col)
            if resolved:
                return df[resolved].count()
            return len(df)
        elif op == "concat":
            cols = op_spec.get("columns", [])
            sep = op_spec.get("separator", " ")
            parts = []
            for c in cols:
                rc = self._resolve_df_col(df, c)
                if rc:
                    parts.append(df[rc].astype(str))
            if parts:
                result = parts[0]
                for p in parts[1:]:
                    result = result + sep + p
                return result
            return ""
        elif op == "format_date":
            col = op_spec.get("column", "")
            fmt = op_spec.get("format", "%Y-%m-%d")
            resolved = self._resolve_df_col(df, col)
            if resolved:
                dt_s = _coerce_datetime_series(df[resolved])
                return dt_s.dt.strftime(fmt).fillna("")
            return ""
        elif op == "format_number":
            col = op_spec.get("column", "")
            decimals = op_spec.get("decimals", 2)
            resolved = self._resolve_df_col(df, col)
            if resolved:
                return df[resolved].round(decimals)
            return 0
        elif op == "format_hms":
            col = op_spec.get("column", "")
            resolved = self._resolve_df_col(df, col)
            if resolved:
                total_sec = int(pd.to_numeric(df[resolved], errors="coerce").sum())
            else:
                total_sec = 0
            h, rem = divmod(abs(total_sec), 3600)
            m, s = divmod(rem, 60)
            sign = "-" if total_sec < 0 else ""
            return f"{sign}{h}:{m:02d}:{s:02d}"
        else:
            logger.warning("unknown_declarative_op", extra={"op": op})
            return None

    def _resolve_agg_or_col(self, df, spec) -> Any:
        """Resolve a spec that can be a column name string, nested op dict, or numeric literal."""
        if isinstance(spec, (int, float)):
            return spec
        if isinstance(spec, str):
            resolved = self._resolve_df_col(df, spec)
            if resolved:
                return df[resolved]
            return 0
        if isinstance(spec, dict):
            return self._apply_declarative_op(df, spec)
        return 0

    def _interpret_legacy_computed(self, df, expr: str) -> Any:
        """Regex-based fallback for old SQL expressions like SUM(col1) - SUM(col2)."""
        import pandas as pd

        # Try simple column reference: table.column
        m = _DIRECT_COLUMN_RE.match(expr.strip())
        if m:
            col = m.group("column")
            if col in df.columns:
                return df[col]
            return None

        # Try SUM(column)
        sum_match = re.match(r"^\s*SUM\s*\(\s*(?:\w+\.)?(\w+)\s*\)\s*$", expr, re.IGNORECASE)
        if sum_match:
            col = sum_match.group(1)
            if col in df.columns:
                return df[col].sum()
            return 0

        # Try SUM(a) - SUM(b)
        diff_match = re.match(
            r"^\s*SUM\s*\(\s*(?:\w+\.)?(\w+)\s*\)\s*-\s*SUM\s*\(\s*(?:\w+\.)?(\w+)\s*\)\s*$",
            expr, re.IGNORECASE,
        )
        if diff_match:
            col_a, col_b = diff_match.group(1), diff_match.group(2)
            a = df[col_a].sum() if col_a in df.columns else 0
            b = df[col_b].sum() if col_b in df.columns else 0
            return a - b

        # Try SUM(a) / NULLIF(SUM(b), 0) or similar ratio
        ratio_match = re.match(
            r"^\s*SUM\s*\(\s*(?:\w+\.)?(\w+)\s*\)\s*/\s*(?:NULLIF\s*\()?\s*SUM\s*\(\s*(?:\w+\.)?(\w+)\s*\)\s*(?:,\s*0\s*\))?\s*$",
            expr, re.IGNORECASE,
        )
        if ratio_match:
            col_a, col_b = ratio_match.group(1), ratio_match.group(2)
            a = df[col_a].sum() if col_a in df.columns else 0
            b = df[col_b].sum() if col_b in df.columns else 0
            return a / b if b else None

        logger.warning("uninterpretable_legacy_computed", extra={"expr": expr})
        return None

    def resolve_header_data(
        self,
        loader,
        params: Dict[str, Any],
        start_date: str | None = None,
        end_date: str | None = None,
    ) -> Dict[str, Any]:
        """Fetch scalar/header values via DataFrame access. Returns {token: value}."""
        result: Dict[str, Any] = {}
        header_tokens = _ensure_sequence(self._raw.get("header_tokens")) or self._scalar_tokens

        for token in header_tokens:
            # Check if it's a PARAM (PARAM:xxx format)
            mapping_expr = self._mapping.get(token, "")
            pm = _PARAM_RE.match(mapping_expr)
            if pm:
                param_name = pm.group(1)
                result[token] = params.get(param_name, "")
                continue

            # Check params.xxx dot-notation format
            dot_m = _DIRECT_COLUMN_RE.match(mapping_expr)
            if dot_m and dot_m.group("table").lower() == "params":
                param_name = dot_m.group("column")
                result[token] = params.get(param_name, "")
                continue

            ref = self._resolve_mapping_column(token)
            if not ref:
                result[token] = ""
                continue

            table, column = ref
            try:
                df = loader.frame(table)
            except Exception:
                result[token] = ""
                continue

            df = self._apply_date_filter_df(df, table, start_date, end_date)
            if df.empty:
                result[token] = ""
                continue
            # Exact match first, then case-insensitive fallback
            if column not in df.columns:
                col_lower = column.lower()
                matched = next(
                    (c for c in df.columns if isinstance(c, str) and c.lower() == col_lower),
                    None,
                )
                if matched:
                    column = matched
                else:
                    result[token] = ""
                    continue

            # Take first non-null value
            non_null = df[column].dropna()
            result[token] = str(non_null.iloc[0]) if not non_null.empty else ""

        return result

    def resolve_row_data(
        self,
        loader,
        params: Dict[str, Any],
        start_date: str | None = None,
        end_date: str | None = None,
        value_filters: Dict[str, list] | None = None,
    ):
        """Fetch row data via DataFrame filtering and return a DataFrame."""
        import pandas as pd

        row_tokens = _ensure_sequence(self._raw.get("row_tokens")) or self._row_tokens
        if not row_tokens:
            return pd.DataFrame()

        # Determine source table from first row token mapping reference,
        # falling back to parent_table.  Row tokens often live in the child
        # (detail) table, not the parent (header) table.
        source_table = None
        for tok in row_tokens:
            ref = self._resolve_mapping_column(tok)
            if ref:
                source_table = ref[0]
                break
        if not source_table:
            source_table = self._parent_table

        if not source_table:
            return pd.DataFrame()

        # Pre-filter at SQL level when date column is known to avoid loading
        # millions of rows into memory (e.g. 2.7M temperature rows).
        date_col = self._date_columns.get(source_table.lower()) or self._date_columns.get(source_table)
        try:
            if date_col and (start_date or end_date) and hasattr(loader, 'frame_date_filtered'):
                df = loader.frame_date_filtered(source_table, date_col, start_date, end_date).copy()
                logger.info("resolve_row_data: loaded %d rows from %s (SQL date-filtered)", len(df), source_table)
            else:
                df = loader.frame(source_table).copy()
                logger.info("resolve_row_data: loaded %d rows from %s (full)", len(df), source_table)
        except Exception:
            logger.exception("resolve_row_data: failed to load table %r", source_table)
            return pd.DataFrame()

        # Apply DataFrame-level date filter as safety net (handles timezone stripping, snap, etc.)
        df = self._apply_date_filter_df(df, source_table, start_date, end_date)
        logger.info("resolve_row_data: %d rows after date filter (start=%s end=%s)", len(df), start_date, end_date)

        # Apply value filters
        if value_filters:
            df = self._apply_value_filters_df(df, value_filters)

        # Apply pre_aggregate (collapse time-series into one row per batch)
        if self._pre_aggregate:
            df = self._apply_pre_aggregate_df(df)

        # Apply group_aggregate (sum across batches → single row)
        if self._group_aggregate:
            df = self._apply_group_aggregate_df(df)

        # Apply reshape rules if present
        # Store db_path for auto-discover in HOURLY_PIVOT
        self._db_path = getattr(loader, 'db_path', None)
        melt_alias_set: set[str] = set()
        if self._reshape_rules:
            df = self._apply_reshape_df(df, loader, source_table)
            logger.info("resolve_row_data: %d rows after reshape", len(df))
            if df.empty:
                return pd.DataFrame()
            # Build set of reshape alias column names for fallback resolution
            for rule in self._reshape_rules:
                for col_spec in rule.get("columns", []):
                    alias = col_spec.get("as", "")
                    if alias:
                        melt_alias_set.add(alias)
                # WINDOW_DIFF produces output_columns directly
                for out_col in rule.get("output_columns", []):
                    if out_col:
                        melt_alias_set.add(out_col)

        # Apply post_aggregate (runs AFTER reshape, e.g. group melted rows by material)
        if self._post_aggregate:
            df = self._apply_group_aggregate_df(df, self._post_aggregate)

        # Build result with mapped columns.
        # Add computed columns back to df so subsequent computations can reference them.
        result_cols: Dict[str, Any] = {}
        for tok in row_tokens:
            resolved = False
            short = tok.removeprefix("row_") if tok.startswith("row_") else tok

            # 1. Try row_computed first (format_date, concat, etc. take priority)
            if tok in self._row_computed:
                computed = self._apply_declarative_op(df, self._row_computed[tok])
                if computed is not None:
                    result_cols[tok] = computed
                    resolved = True
                    # Feed computed column back to df under both token name and
                    # short name so subsequent ops can reference it
                    try:
                        df[tok] = computed
                        if short != tok:
                            df[short] = computed
                    except Exception:
                        pass

            # 1b. INDEX mapping → 1-based serial number
            if not resolved and self._mapping.get(tok) == "INDEX":
                result_cols[tok] = list(range(1, len(df) + 1))
                resolved = True

            # 1c. Direct column match (MELT alias IS the token name)
            if not resolved and tok in df.columns:
                result_cols[tok] = df[tok].values
                resolved = True

            # 2. Try direct mapping resolution (table.column)
            if not resolved:
                ref = self._resolve_mapping_column(tok)
                if ref:
                    _, col = ref
                    if col in df.columns:
                        result_cols[tok] = df[col].values
                        resolved = True
                    elif melt_alias_set and (short in df.columns or tok in df.columns):
                        # After MELT, columns are named by alias (full or stripped)
                        col_to_use = short if short in df.columns else tok
                        result_cols[tok] = df[col_to_use].values
                        resolved = True
                    else:
                        # Case-insensitive column fallback
                        col_lower = col.lower()
                        for actual_col in df.columns:
                            if isinstance(actual_col, str) and actual_col.lower() == col_lower:
                                result_cols[tok] = df[actual_col].values
                                resolved = True
                                break

            # 3. Try MELT alias match by stripped name or full token name
            if not resolved and melt_alias_set and (short in df.columns or tok in df.columns):
                col_to_use = short if short in df.columns else tok
                result_cols[tok] = df[col_to_use].values
                resolved = True

            # 4. Fallback
            if not resolved:
                mapping_expr = self._mapping.get(tok, "")
                logger.warning("row_token_unresolved", extra={
                    "event": "row_token_unresolved",
                    "token": tok,
                    "mapping_expr": mapping_expr,
                    "available_columns": list(df.columns)[:20],
                })
                result_cols[tok] = ""

        try:
            result_df = pd.DataFrame(result_cols)
        except ValueError:
            # All scalar values (e.g. every token unresolved → "") — wrap in list
            result_df = pd.DataFrame({k: [v] for k, v in result_cols.items()})

        # Carry forward __batch_idx__ and metadata columns for BLOCK_REPEAT grouping
        if melt_alias_set and "__batch_idx__" in df.columns and len(df) == len(result_df):
            result_df["__batch_idx__"] = df["__batch_idx__"].values
            # Collect the set of all melted source columns
            _all_melted_src: set[str] = set()
            for rule in self._reshape_rules:
                for cs in rule.get("columns", []):
                    for f in cs.get("from", []):
                        if f != "INDEX" and "." in f:
                            _all_melted_src.add(f.split(".", 1)[1])
            for col in df.columns:
                if col.startswith("__") or col in _all_melted_src:
                    continue
                if col not in result_df.columns and col not in melt_alias_set:
                    result_df[f"__cf_{col}"] = df[col].values

        # Apply ordering
        order_cols = self._row_order or self._order_by_rows
        if order_cols:
            sort_by: list[str] = []
            ascending: list[bool] = []
            for clause in order_cols:
                parts = clause.strip().split()
                col_name = parts[0] if parts else ""
                if col_name in result_df.columns:
                    sort_by.append(col_name)
                    ascending.append(not (len(parts) > 1 and parts[1].upper() == "DESC"))
            if sort_by:
                result_df = result_df.sort_values(sort_by, ascending=ascending, ignore_index=True)

        return result_df

    def resolve_totals_data(self, rows_df) -> Dict[str, Any]:
        """Compute totals from the rows DataFrame using declarative specs."""
        import pandas as pd

        result: Dict[str, Any] = {}
        total_tokens = self._total_tokens
        totals_math = self._totals_math
        totals_mapping = self._totals_mapping

        # Add short aliases (strip row_ prefix) so totals_math can reference
        # "set_wt" even though column is "row_set_wt"
        if isinstance(rows_df, pd.DataFrame) and not rows_df.empty:
            for col in list(rows_df.columns):
                if col.startswith("row_"):
                    short = col.removeprefix("row_")
                    if short not in rows_df.columns:
                        rows_df = rows_df.copy()
                        rows_df[short] = rows_df[col]

        for tok in total_tokens or list(totals_math.keys()):
            if tok in totals_math:
                val = self._apply_declarative_op(rows_df, totals_math[tok])
                result[tok] = val if val is not None else ""
            elif tok in totals_mapping:
                val = self._apply_declarative_op(rows_df, totals_mapping[tok])
                result[tok] = val if val is not None else ""
            else:
                result[tok] = ""

        return result

    def _apply_reshape_df(self, df, loader, source_table: str):
        """Apply reshape rules (UNION_ALL / MELT) to produce long-form data."""
        import pandas as pd

        for rule in self._reshape_rules:
            strategy = str(rule.get("strategy", "")).upper()
            columns = rule.get("columns", [])

            if strategy == "UNION_ALL" and columns:
                frames: list[pd.DataFrame] = []
                # Each column spec: {"as": "alias", "from": ["col1", "col2", ...]}
                n_rows = len(columns[0].get("from", [])) if columns else 0
                for i in range(n_rows):
                    row_df = pd.DataFrame()
                    for col_spec in columns:
                        alias = col_spec.get("as", "")
                        sources = col_spec.get("from", [])
                        if i < len(sources):
                            src = sources[i]
                            # Source can be table.column or just column
                            if "." in src:
                                _, src_col = src.split(".", 1)
                            else:
                                src_col = src
                            if src_col in df.columns:
                                row_df[alias] = df[src_col].values
                            else:
                                row_df[alias] = ""
                    if not row_df.empty:
                        frames.append(row_df)

                if frames:
                    df = pd.concat(frames, ignore_index=True)
                    # Drop rows where all aliased columns are empty
                    alias_cols = [c.get("as", "") for c in columns if c.get("as")]
                    existing = [c for c in alias_cols if c in df.columns]
                    if existing:
                        df = df.dropna(subset=existing, how="all")
                        str_df = df[existing].fillna("").astype(str)
                        for c in str_df.columns:
                            str_df[c] = str_df[c].str.strip()
                        str_mask = (str_df != "").any(axis=1)
                        df = df.loc[str_mask].reset_index(drop=True)

            elif strategy == "MELT" and columns:
                # Melt: unpivot wide columns into long format.
                # Each column spec: {"as": "alias", "from": ["table.col1", "table.col2", ...]}
                # Special: {"as": "alias", "from": ["INDEX"]} → 1-based position index
                frames: list[pd.DataFrame] = []
                # Determine number of positions from the first non-INDEX column
                n_positions = 0
                for col_spec in columns:
                    froms = col_spec.get("from", [])
                    if froms and froms != ["INDEX"]:
                        n_positions = len(froms)
                        break

                # Tag each original row with a batch index for BLOCK_REPEAT grouping
                import numpy as np
                df["__batch_idx__"] = np.arange(len(df))

                for i in range(n_positions):
                    slice_df = pd.DataFrame()
                    for col_spec in columns:
                        alias = col_spec.get("as", "")
                        froms = col_spec.get("from", [])
                        if froms == ["INDEX"]:
                            # Generate 1-based index for each original row
                            slice_df[alias] = [i + 1] * len(df)
                        elif i < len(froms):
                            src = froms[i]
                            src_col = src.split(".", 1)[1] if "." in src else src
                            if src_col in df.columns:
                                slice_df[alias] = df[src_col].values
                            else:
                                # Not a column — treat as literal value (e.g. bin_number: "1", "2", ...)
                                slice_df[alias] = [src if src else ""] * len(df)
                    if not slice_df.empty:
                        # Carry forward non-melted columns (e.g. recipe_name, id, start_time)
                        melted_src_cols: set[str] = set()
                        for cs in columns:
                            for f in cs.get("from", []):
                                if f != "INDEX" and "." in f:
                                    melted_src_cols.add(f.split(".", 1)[1])
                        for orig_col in df.columns:
                            if orig_col not in melted_src_cols and orig_col not in slice_df.columns:
                                slice_df[orig_col] = df[orig_col].values
                        frames.append(slice_df)

                if frames:
                    df = pd.concat(frames, ignore_index=True)
                    # Drop rows where the primary alias (first non-INDEX) is null/empty
                    primary_alias = ""
                    for cs in columns:
                        if cs.get("from", []) != ["INDEX"]:
                            primary_alias = cs.get("as", "")
                            break
                    if primary_alias and primary_alias in df.columns:
                        mask = df[primary_alias].notna()
                        str_vals = df[primary_alias].astype(str).str.strip()
                        mask = mask & (str_vals != "") & (str_vals.str.lower() != "nan") & (str_vals.str.lower() != "none")
                        df = df.loc[mask].reset_index(drop=True)

            elif strategy == "SELECT" and columns:
                # SELECT: derive/rename columns, optionally group by + aggregate.
                # Each column spec: {"as": "alias", "from": ["table.col"]}
                # If from[0] is "date(table.col)" → extract date part.
                # If the rule has "group_by": true, group by derived columns
                # and SUM all numeric columns.

                for col_spec in columns:
                    alias = col_spec.get("as", "")
                    sources = col_spec.get("from", [])
                    if not alias or not sources:
                        continue
                    src = sources[0]
                    # Handle date() wrapper
                    date_match = re.match(r"date\((.+)\)", src, re.IGNORECASE)
                    if date_match:
                        inner = date_match.group(1)
                        src_col = inner.split(".", 1)[1] if "." in inner else inner
                        if src_col in df.columns:
                            dt_s = _coerce_datetime_series(df[src_col])
                            df[alias] = dt_s.dt.strftime("%Y-%m-%d").fillna("")
                    else:
                        src_col = src.split(".", 1)[1] if "." in src else src
                        if src_col in df.columns:
                            df[alias] = df[src_col].values

                # If group_by hint is present, group by derived alias columns
                group_by_aliases = rule.get("group_by")
                if group_by_aliases:
                    if isinstance(group_by_aliases, bool):
                        # Auto-detect: group by all non-numeric derived columns
                        group_by_aliases = [
                            cs.get("as", "") for cs in columns
                            if cs.get("as", "") in df.columns
                            and not pd.api.types.is_numeric_dtype(df[cs["as"]])
                        ]
                    existing_groups = [g for g in group_by_aliases if g in df.columns]
                    if existing_groups:
                        agg_map = {}
                        for col in df.columns:
                            if col in existing_groups:
                                continue
                            if pd.api.types.is_numeric_dtype(df[col]):
                                agg_map[col] = "sum"
                            else:
                                agg_map[col] = "first"
                        df = df.groupby(existing_groups, sort=True).agg(agg_map).reset_index()
                        logger.info("select_group_by applied → %d rows", len(df))

            elif strategy == "WINDOW_DIFF":
                # Detect run intervals from cumulative counter changes.
                # column_groups: [{machine_name, columns: [HRS, MIN, SEC]}]
                # Produces rows: machine_name, run_date, start_time, end_time, duration_sec, shift_no, total_time
                df = self._apply_window_diff(df, rule, loader)

            elif strategy == "HOURLY_PIVOT":
                # Transpose time-series: sensors become rows, hours become columns.
                # Config:
                #   timestamp_col: "timestamp_utc"
                #   sensors: [{col: "AI_1_PT_1", tag: "AI-1_PT-1", desc: "UF FEED", label: "FEED"}]
                #   hour_start: 6  (default 6 = 6AM)
                #   include_stats: true  (add MAX, MAX_DATETIME, MIN, MIN_DATETIME, AVG columns)
                #   divisor: 100  (divide raw values)
                df = self._apply_hourly_pivot(df, rule)

            elif strategy == "RUN_HOURS_DIFF":
                # Group by description, compute (last - first) total_seconds,
                # format as H:MM:SS. Produces one row per machine.
                # Config: group_col, seconds_col, timestamp_col
                df = self._apply_run_hours_diff(df, rule)

        return df

    def _apply_run_hours_diff(self, df: "pd.DataFrame", rule: dict) -> "pd.DataFrame":
        """Compute running hours diff (last - first) per machine group."""
        import pandas as pd

        group_col = rule.get("group_col", "description")
        seconds_col = rule.get("seconds_col", "total_seconds")
        ts_col = rule.get("timestamp_col", "timestamp_utc")

        if group_col not in df.columns or seconds_col not in df.columns:
            logger.warning("run_hours_diff: missing columns %s/%s", group_col, seconds_col)
            return df

        ts = _coerce_datetime_series(df[ts_col])
        df = df.copy()
        df["__ts_naive__"] = ts
        df = df.sort_values(["__ts_naive__"], ascending=True)

        rows = []
        sr = 0
        for desc, group in df.groupby(group_col, sort=True):
            sr += 1
            secs = pd.to_numeric(group[seconds_col], errors="coerce")
            first_s = secs.iloc[0] if not secs.empty else 0
            last_s = secs.iloc[-1] if not secs.empty else 0
            diff = abs(int(last_s - first_s))
            h, rem = divmod(diff, 3600)
            m, s = divmod(rem, 60)
            rows.append({
                "row_sr_no": str(sr),
                "row_description": str(desc),
                "row_running_hours": f"{h}h {m}m {s}s",
            })

        result = pd.DataFrame(rows)
        logger.info("run_hours_diff: %d groups → %d rows", len(rows), len(result))
        return result

    def _apply_hourly_pivot(self, df: "pd.DataFrame", rule: dict) -> "pd.DataFrame":
        """Transpose time-series: sensors → rows, hours → columns.

        Produces one row per sensor with 24 hourly columns (6AM..5AM) plus
        optional MAX/MIN/AVG stats.
        """
        import pandas as pd
        import numpy as np

        ts_col = rule.get("timestamp_col", "timestamp_utc")
        sensors = rule.get("sensors", [])
        hour_start = int(rule.get("hour_start", 6))
        include_stats = rule.get("include_stats", False)
        divisor = float(rule.get("divisor", 1))
        auto_discover = rule.get("auto_discover", False)

        # Auto-discover sensors from neuract__device_mappings table if enabled
        if (auto_discover or not sensors) and self._parent_table:
            try:
                import sqlite3 as _sq
                db_path = getattr(self, '_db_path', None)
                if db_path is None and hasattr(self, '_loader_ref') and self._loader_ref:
                    db_path = getattr(self._loader_ref, 'db_path', None)
                if db_path:
                    with _sq.connect(str(db_path), timeout=30) as _con:
                        _rows = _con.execute(
                            "SELECT field_key FROM neuract__device_mappings WHERE table_name = ? ORDER BY field_key",
                            (self._parent_table,)
                        ).fetchall()
                        if _rows:
                            # Also fetch the OPC address for description
                            _detail_rows = _con.execute(
                                "SELECT field_key, address FROM neuract__device_mappings WHERE table_name = ? ORDER BY field_key",
                                (self._parent_table,)
                            ).fetchall()
                            _addr_map = {r[0]: r[1] for r in _detail_rows}

                            discovered = []
                            for r in _rows:
                                col = r[0]
                                # Skip _TOTAL columns (totalizer/accumulator duplicates)
                                if col.endswith("_TOTAL"):
                                    continue
                                # Format tag: AI_10_RO1_ORP → AI-10_RO1 ORP
                                parts = col.split("_")
                                if len(parts) >= 3 and parts[1].isdigit():
                                    tag = f"{parts[0]}-{parts[1]}_{' '.join(parts[2:])}"
                                else:
                                    tag = col.replace("_", " ")
                                # Description from OPC address path
                                addr = _addr_map.get(col, "")
                                desc = ""
                                if addr:
                                    addr_parts = addr.split(".")
                                    if len(addr_parts) >= 2:
                                        desc = ".".join(addr_parts[-2:])
                                discovered.append({"col": col, "tag": tag, "desc": desc})

                            # Merge: keep existing sensor configs, add any missing
                            existing_cols = {s["col"] for s in sensors}
                            for d in discovered:
                                if d["col"] not in existing_cols and d["col"] != ts_col:
                                    sensors.append(d)
                            if not sensors:
                                sensors = [s for s in discovered if s["col"] != ts_col]
                            logger.info("hourly_pivot: auto-discovered %d sensors from device_mappings for %s", len(discovered), self._parent_table)
            except Exception as exc:
                logger.debug("hourly_pivot: auto-discover failed: %s", exc)

        if ts_col not in df.columns or not sensors:
            logger.warning("hourly_pivot: missing timestamp_col=%s or empty sensors", ts_col)
            return df

        ts = _coerce_datetime_series(df[ts_col])
        df = df.copy()
        df["__ts__"] = ts
        df["__hour__"] = ts.dt.hour

        # Hour labels: 6AM,7AM,...,5AM
        hour_labels = []
        for offset in range(24):
            h = (hour_start + offset) % 24
            if h == 0:
                hour_labels.append("12AM")
            elif h == 12:
                hour_labels.append("12PM")
            elif h < 12:
                hour_labels.append(f"{h}AM")
            else:
                hour_labels.append(f"{h - 12}PM")

        out_rows = []
        for idx, sensor in enumerate(sensors):
            col = sensor.get("col", "")
            tag = sensor.get("tag", col)
            desc = sensor.get("desc", "")
            label = sensor.get("label", "")

            if col not in df.columns:
                logger.warning("hourly_pivot: sensor column %s not found", col)
                continue

            raw = pd.to_numeric(df[col], errors="coerce")
            if divisor != 1:
                raw = raw / divisor

            row = {
                "row_sr_no": str(idx + 1),
                "row_tag_name": tag,
                "row_description": desc,
                "row_label": label,
            }

            if include_stats:
                valid = raw.dropna()
                if not valid.empty:
                    max_idx = valid.idxmax()
                    min_idx = valid.idxmin()
                    row["row_max"] = f"{valid.max():.2f}"
                    row["row_max_datetime"] = str(df.at[max_idx, ts_col])[:19] if pd.notna(max_idx) else ""
                    row["row_min"] = f"{valid.min():.2f}"
                    row["row_min_datetime"] = str(df.at[min_idx, ts_col])[:19] if pd.notna(min_idx) else ""
                    row["row_avg"] = f"{valid.mean():.2f}"
                else:
                    row.update({"row_max": "", "row_max_datetime": "", "row_min": "", "row_min_datetime": "", "row_avg": ""})

            for offset in range(24):
                h = (hour_start + offset) % 24
                mask = df["__hour__"] == h
                vals = raw[mask].dropna()
                col_key = f"row_h{offset}"
                row[col_key] = f"{vals.mean():.2f}" if not vals.empty else ""

            out_rows.append(row)

        if not out_rows:
            return df

        result = pd.DataFrame(out_rows)
        logger.info("hourly_pivot: %d sensors → %d rows × %d cols", len(sensors), len(result), len(result.columns))
        return result

    def _apply_window_diff(self, df, rule: dict, loader) -> "pd.DataFrame":
        """Detect machine run intervals from RUNHOURS cumulative counters."""
        import pandas as pd
        import numpy as np
        from datetime import datetime, timedelta

        column_groups = rule.get("column_groups", [])
        ts_col = rule.get("timestamp_column", "timestamp_utc")
        shift_boundaries = rule.get("shift_boundaries", {})

        if df.empty or not column_groups:
            return pd.DataFrame()

        # Ensure sorted by timestamp
        if ts_col in df.columns:
            df = df.sort_values(ts_col).reset_index(drop=True)

        intervals = []
        for group in column_groups:
            machine_name = group.get("machine_name", "")
            cols = group.get("columns", [])
            if not cols or not all(c in df.columns for c in cols):
                continue

            # Coerce to numeric
            for c in cols:
                df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

            # Build composite value: HRS*3600 + MIN*60 + SEC
            composite = df[cols[0]] * 3600
            if len(cols) > 1:
                composite = composite + df[cols[1]] * 60
            if len(cols) > 2:
                composite = composite + df[cols[2]]

            # Detect changes: where composite value differs from previous
            diffs = composite.diff().fillna(0)
            is_changing = diffs != 0

            # Find intervals where machine is running (consecutive changing rows)
            timestamps = pd.to_datetime(df[ts_col], errors="coerce")
            run_start = None

            for idx in range(len(df)):
                if is_changing.iloc[idx]:
                    if run_start is None:
                        # Start of a run interval — use PREVIOUS timestamp as start
                        run_start = idx - 1 if idx > 0 else idx
                else:
                    if run_start is not None:
                        # End of run interval
                        start_ts = timestamps.iloc[run_start]
                        end_ts = timestamps.iloc[idx - 1] if idx > 0 else timestamps.iloc[idx]
                        if pd.notna(start_ts) and pd.notna(end_ts):
                            dur = int((end_ts - start_ts).total_seconds())
                            if dur > 0:
                                intervals.append(self._make_interval_row(
                                    machine_name, start_ts, end_ts, dur, shift_boundaries
                                ))
                        run_start = None

            # Close open interval at end
            if run_start is not None:
                start_ts = timestamps.iloc[run_start]
                end_ts = timestamps.iloc[len(df) - 1]
                if pd.notna(start_ts) and pd.notna(end_ts):
                    dur = int((end_ts - start_ts).total_seconds())
                    if dur > 0:
                        intervals.append(self._make_interval_row(
                            machine_name, start_ts, end_ts, dur, shift_boundaries
                        ))

        if not intervals:
            return pd.DataFrame(columns=[
                "machine_name", "run_date", "start_time", "end_time",
                "duration_sec", "shift_no", "total_time"
            ])

        result = pd.DataFrame(intervals)
        logger.info("WINDOW_DIFF applied: %d intervals from %d machine groups", len(result), len(column_groups))
        return result

    @staticmethod
    def _make_interval_row(machine_name, start_ts, end_ts, duration_sec, shift_boundaries):
        """Create a single interval row dict."""
        # Format times
        run_date = start_ts.strftime("%Y-%m-%d")
        start_time = start_ts.strftime("%H:%M:%S")
        end_time = end_ts.strftime("%H:%M:%S")

        # Format total_time as HH:MM:SS
        hours = duration_sec // 3600
        minutes = (duration_sec % 3600) // 60
        seconds = duration_sec % 60
        total_time = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        # Determine shift number based on start_time hour
        shift_no = ""
        start_hour = start_ts.hour
        if shift_boundaries:
            if 6 <= start_hour < 14:
                shift_no = "1"
            elif 14 <= start_hour < 22:
                shift_no = "2"
            else:
                shift_no = "3"

        return {
            "machine_name": machine_name,
            "run_date": run_date,
            "start_time": start_time,
            "end_time": end_time,
            "duration_sec": duration_sec,
            "shift_no": shift_no,
            "total_time": total_time,
        }

    # ------------------------------------------------------------------ #
    # Internal helpers
    # ------------------------------------------------------------------ #
    def _discover_param_tokens(self) -> List[str]:
        tokens: set[str] = set()
        for expr in self._mapping.values():
            for match in _PARAM_RE.findall(expr):
                tokens.add(match)
            # Also detect params.xxx dot-notation
            m = _DIRECT_COLUMN_RE.match(str(expr).strip())
            if m and m.group("table").lower() == "params":
                tokens.add(m.group("column"))
        for expr in self._required_filters.values():
            tokens.update(_PARAM_RE.findall(expr))
        for expr in self._optional_filters.values():
            tokens.update(_PARAM_RE.findall(expr))
        return list(tokens)


# ======================================================================
# dataframe_pipeline
# ======================================================================

import logging
from typing import Any, Dict, List, Optional


logger = logging.getLogger(__name__)


class DataPipelineError(RuntimeError):
    """Raised when the data pipeline cannot produce valid results.

    This replaces the previous silent-failure pattern where exceptions
    were caught and empty data was returned, causing blank reports.
    """
    pass

_CF_PREFIX = "__cf_"


class DataFramePipeline:
    """Pure-DataFrame replacement for ``_run_generator_entrypoints()``."""

    def __init__(
        self,
        contract_adapter: ContractAdapter,
        loader,
        params: Dict[str, Any],
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
        value_filters: Optional[Dict[str, list]] = None,
    ) -> None:
        self.adapter = contract_adapter
        self.loader = loader
        self.params = params or {}
        self.start_date = start_date
        self.end_date = end_date
        self.value_filters = value_filters or {}

    def execute(self) -> Dict[str, list]:
        """Return ``{"header": [...], "rows": [...], "totals": [...]}``.

        The format is identical to what the SQL entrypoint runner produces so
        it can be used as a drop-in replacement in ``ReportGenerate.py``.

        When rows contain ``__batch_idx__`` (from MELT reshape), the result
        also includes a ``batches`` key — a list of per-batch dicts each
        having ``header``, ``rows``, and ``totals``.
        """
        header = self._resolve_header()
        rows_df = self._resolve_rows()
        totals = self._resolve_totals(rows_df)

        # Separate internal columns from user-visible row data
        row_cols = [c for c in (rows_df.columns if rows_df is not None else [])
                    if not c.startswith("__")]
        rows_list = (rows_df[row_cols].to_dict("records")
                     if rows_df is not None and not rows_df.empty and row_cols
                     else [])

        result: Dict[str, Any] = {
            "header": [header] if header else [],
            "rows": rows_list,
            "totals": [totals] if totals else [],
        }

        # --- Per-batch grouping for BLOCK_REPEAT ---
        if rows_df is not None and "__batch_idx__" in rows_df.columns and not rows_df.empty:
            batches = self._group_into_batches(rows_df, row_cols)
            if batches:
                result["batches"] = batches

        return result

    # -------------------------------------------------------------- #
    # Per-batch grouping
    # -------------------------------------------------------------- #
    def _group_into_batches(
        self, rows_df, row_cols: List[str]
    ) -> List[Dict[str, Any]]:
        """Group rows by ``__batch_idx__`` and return per-batch data."""
        import pandas as pd

        batches: List[Dict[str, Any]] = []

        # Extract carry-forward column names (prefixed with __cf_)
        cf_cols = [c for c in rows_df.columns if c.startswith(_CF_PREFIX)]

        for batch_idx, group in rows_df.groupby("__batch_idx__", sort=True):
            # Build batch-level header from carry-forward columns
            batch_header: Dict[str, Any] = {}
            if cf_cols and not group.empty:
                first = group.iloc[0]
                for col in cf_cols:
                    orig_name = col[len(_CF_PREFIX):]
                    val = first[col]
                    if pd.notna(val):
                        batch_header[orig_name] = val

            batch_header["__batch_number__"] = len(batches) + 1

            # Row data (only user-visible columns)
            batch_rows = group[row_cols].to_dict("records") if row_cols else []

            # Per-batch totals
            try:
                batch_totals = self.adapter.resolve_totals_data(group)
            except Exception:
                logger.exception("df_pipeline_batch_totals_failed batch=%s", batch_idx)
                batch_totals = {}

            batches.append({
                "header": batch_header,
                "rows": batch_rows,
                "totals": batch_totals,
            })

        logger.info("df_pipeline_grouped batches=%d total_rows=%d", len(batches), len(rows_df))
        return batches

    # -------------------------------------------------------------- #
    # Internal resolvers
    # -------------------------------------------------------------- #
    def _resolve_header(self) -> Dict[str, Any]:
        try:
            return self.adapter.resolve_header_data(
                self.loader,
                self.params,
                start_date=self.start_date,
                end_date=self.end_date,
            )
        except Exception as exc:
            logger.exception("df_pipeline_header_failed")
            raise DataPipelineError(f"Header resolution failed: {exc}") from exc

    def _resolve_rows(self):
        import pandas as pd

        try:
            result = self.adapter.resolve_row_data(
                self.loader,
                self.params,
                start_date=self.start_date,
                end_date=self.end_date,
                value_filters=self.value_filters,
            )
            if result is None or (isinstance(result, pd.DataFrame) and result.empty):
                logger.warning(
                    "df_pipeline_rows_empty — query returned no data",
                    extra={
                        "event": "df_pipeline_rows_empty",
                        "start_date": self.start_date,
                        "end_date": self.end_date,
                    },
                )
            return result
        except Exception as exc:
            logger.exception("df_pipeline_rows_failed")
            raise DataPipelineError(f"Row resolution failed: {exc}") from exc

    def _resolve_totals(self, rows_df) -> Dict[str, Any]:
        import pandas as pd

        if rows_df is None or rows_df.empty:
            return {}
        try:
            return self.adapter.resolve_totals_data(rows_df)
        except Exception as exc:
            # Totals failure is non-fatal: a report with rows but no totals
            # is still useful, whereas a report with no rows is not.
            logger.warning(
                "df_pipeline_totals_failed_degraded error=%s",
                exc,
                extra={"event": "df_pipeline_totals_failed_degraded", "error": str(exc)},
            )
            return {}


# ======================================================================
# report_context
# ======================================================================

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger("neura.reports.context")

# Maximum characters of text content to include (to fit LLM context windows)
DEFAULT_MAX_TEXT_CHARS = 60_000


@dataclass
class ReportContext:
    """All report data an agent needs to analyze a report."""

    run_id: str
    template_id: str
    template_name: str
    template_kind: str
    connection_id: Optional[str]
    connection_name: Optional[str]
    start_date: Optional[str]
    end_date: Optional[str]
    status: str
    created_at: Optional[str]

    html_content: str = ""
    text_content: str = ""
    tables: List[Dict[str, Any]] = field(default_factory=list)
    artifact_urls: Dict[str, Optional[str]] = field(default_factory=dict)
    key_values: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


class ReportContextProvider:
    """
    Service that reads report run records and artifacts, making them
    consumable by agents (text extraction, table parsing, etc.).
    """

    def __init__(self, max_text_chars: int = DEFAULT_MAX_TEXT_CHARS):
        self._max_text_chars = max_text_chars

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def get_report_context(self, run_id: str) -> Optional[ReportContext]:
        """
        Load full report context for a given run_id.

        Returns None if the run doesn't exist.
        """
        run = self._get_run(run_id)
        if not run:
            return None

        artifacts = run.get("artifacts") or {}
        html_content = ""
        text_content = ""
        tables: List[Dict[str, Any]] = []

        # Try to read the HTML artifact from disk
        html_url = artifacts.get("html_url")
        html_path = self._resolve_artifact_path(html_url)
        if html_path and html_path.is_file():
            try:
                raw_html = html_path.read_text(encoding="utf-8", errors="replace")
                html_content = raw_html
                text_content = self._extract_text_from_html(raw_html)
                tables = self._extract_tables_from_html(raw_html)
            except Exception as exc:
                logger.warning("Failed to read HTML artifact %s: %s", html_path, exc)
        else:
            if html_url:
                logger.warning(
                    "HTML artifact not found on disk for run %s: url=%s resolved=%s",
                    run_id, html_url, html_path,
                )
            else:
                logger.warning("No html_url in artifacts for run %s (keys: %s)", run_id, list(artifacts.keys()))

        # Truncate to fit LLM context
        if len(text_content) > self._max_text_chars:
            text_content = text_content[: self._max_text_chars] + f"\n\n[...truncated at {self._max_text_chars} chars]"

        return ReportContext(
            run_id=run_id,
            template_id=run.get("templateId") or "",
            template_name=run.get("templateName") or "",
            template_kind=run.get("templateKind") or "pdf",
            connection_id=run.get("connectionId"),
            connection_name=run.get("connectionName"),
            start_date=run.get("startDate"),
            end_date=run.get("endDate"),
            status=run.get("status") or "unknown",
            created_at=run.get("createdAt"),
            html_content=html_content,
            text_content=text_content,
            tables=tables,
            artifact_urls={
                "html_url": artifacts.get("html_url"),
                "pdf_url": artifacts.get("pdf_url"),
                "docx_url": artifacts.get("docx_url"),
                "xlsx_url": artifacts.get("xlsx_url"),
            },
            key_values=run.get("keyValues") or {},
            metadata={
                "batch_ids": run.get("batchIds") or [],
                "schedule_id": run.get("scheduleId"),
                "schedule_name": run.get("scheduleName"),
            },
        )

    def get_report_text(self, run_id: str) -> str:
        """Get plain text content of a report (stripped HTML)."""
        ctx = self.get_report_context(run_id)
        return ctx.text_content if ctx else ""

    def get_report_tables(self, run_id: str) -> List[Dict[str, Any]]:
        """Get extracted data tables from a report."""
        ctx = self.get_report_context(run_id)
        return ctx.tables if ctx else []

    def list_recent_reports(
        self,
        template_id: Optional[str] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """List recent report runs, optionally filtered by template."""
        from backend.app.repositories import state_store

        runs = state_store.list_report_runs(
            template_id=template_id,
            limit=limit,
        )
        return [
            {
                "run_id": r.get("id"),
                "template_id": r.get("templateId"),
                "template_name": r.get("templateName"),
                "template_kind": r.get("templateKind"),
                "status": r.get("status"),
                "created_at": r.get("createdAt"),
                "start_date": r.get("startDate"),
                "end_date": r.get("endDate"),
            }
            for r in runs
        ]

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _get_run(self, run_id: str) -> Optional[Dict[str, Any]]:
        """Fetch a report run record from the state store."""
        from backend.app.repositories import state_store
        return state_store.get_report_run(run_id)

    def _resolve_artifact_path(self, url: Optional[str]) -> Optional[Path]:
        """
        Convert an artifact URL (e.g. /uploads/template_id/file.html)
        back to a filesystem path.
        """
        if not url:
            return None

        from backend.app.services.config import UPLOAD_ROOT, EXCEL_UPLOAD_ROOT

        UPLOAD_ROOT_BASE = UPLOAD_ROOT.resolve()
        EXCEL_UPLOAD_ROOT_BASE = EXCEL_UPLOAD_ROOT.resolve()

        url = url.lstrip("/")

        if url.startswith("uploads/"):
            relative = url[len("uploads/"):]
            candidate = UPLOAD_ROOT_BASE / relative
        elif url.startswith("excel-uploads/"):
            relative = url[len("excel-uploads/"):]
            candidate = EXCEL_UPLOAD_ROOT_BASE / relative
        else:
            return None

        # Safety: ensure the relative path doesn't traverse upwards
        if ".." in str(relative):
            logger.warning("Artifact path contains traversal: %s", url)
            return None

        return candidate if candidate.is_file() else None

    def _extract_text_from_html(self, html: str) -> str:
        """Extract plain text from HTML using BeautifulSoup if available, else regex."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style elements
            for tag in soup(["script", "style", "head"]):
                tag.decompose()

            text = soup.get_text(separator="\n", strip=True)
            # Collapse multiple blank lines
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()
        except ImportError:
            # Fallback: regex-based stripping
            text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()

    def _extract_tables_from_html(self, html: str) -> List[Dict[str, Any]]:
        """
        Extract tables from HTML as list of dicts:
        [{"headers": [...], "rows": [[...], ...]}, ...]
        """
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            result = []

            for table_tag in soup.find_all("table"):
                headers = []
                rows = []

                # Extract headers from <thead> or first <tr> with <th>
                thead = table_tag.find("thead")
                if thead:
                    for th in thead.find_all("th"):
                        headers.append(th.get_text(strip=True))
                else:
                    first_row = table_tag.find("tr")
                    if first_row:
                        ths = first_row.find_all("th")
                        if ths:
                            headers = [th.get_text(strip=True) for th in ths]

                # Extract data rows
                tbody = table_tag.find("tbody") or table_tag
                for tr in tbody.find_all("tr"):
                    tds = tr.find_all("td")
                    if tds:
                        rows.append([td.get_text(strip=True) for td in tds])

                if headers or rows:
                    result.append({"headers": headers, "rows": rows})

            return result
        except ImportError:
            return []


# ======================================================================
# strategies
# ======================================================================

import logging
from dataclasses import dataclass
import importlib
from pathlib import Path
from typing import Optional

from backend.app.utils import StrategyRegistry
logger = logging.getLogger("neura.reports.strategies")


@dataclass
class RenderArtifacts:
    docx_path: Optional[Path]
    xlsx_path: Optional[Path]


class RenderStrategy:
    def render_docx(self, html_path: Path, pdf_path: Optional[Path], dest_tmp: Path, *, landscape: bool, font_scale: Optional[float]) -> Optional[Path]:
        if pdf_path and pdf_path.exists():
            try:
                pdf_result = pdf_file_to_docx(pdf_path, dest_tmp)
            except Exception:
                logger.exception("docx_pdf_convert_failed")
            else:
                if pdf_result:
                    return pdf_result
        try:
            api_mod = importlib.import_module("backend.api")
            html_to_docx = getattr(api_mod, "html_file_to_docx", html_file_to_docx)
        except Exception:
            html_to_docx = html_file_to_docx
        return html_to_docx(html_path, dest_tmp, landscape=landscape, body_font_scale=font_scale)

    def render_xlsx(self, html_path: Path, dest_tmp: Path) -> Optional[Path]:
        try:
            api_mod = importlib.import_module("backend.api")
            html_to_xlsx = getattr(api_mod, "html_file_to_xlsx", html_file_to_xlsx)
        except Exception:
            html_to_xlsx = html_file_to_xlsx
        return html_to_xlsx(html_path, dest_tmp)


class NotificationStrategy:
    def send(self, *, recipients: list[str], subject: str, body: str, attachments: list[Path]) -> bool:
        from backend.app.services.legacy_services import send_report_email
        return send_report_email(
            to_addresses=recipients,
            subject=subject,
            body=body,
            attachments=attachments,
        )


def build_render_strategy_registry() -> StrategyRegistry[RenderStrategy]:
    registry: StrategyRegistry[RenderStrategy] = StrategyRegistry(default_factory=RenderStrategy)
    registry.register("pdf", RenderStrategy())
    registry.register("excel", RenderStrategy())
    return registry


def build_notification_strategy_registry() -> StrategyRegistry[NotificationStrategy]:
    registry: StrategyRegistry[NotificationStrategy] = StrategyRegistry(default_factory=NotificationStrategy)
    registry.register("email", NotificationStrategy())
    return registry


# ======================================================================
# discovery_metrics
# ======================================================================

from typing import Any, Iterable, Mapping, Sequence

__all__ = [
    "build_batch_field_catalog_and_stats",
    "build_batch_metrics",
    "build_discovery_schema",
    "bin_numeric_metric",
    "group_metrics_by_field",
    "build_resample_support",
]


def _coerce_number(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    try:
        return float(value)
    except Exception:
        return 0.0


def _normalize_type(raw: Any) -> str:
    """
    Collapse loose field types coming from discovery into a stable set we can reason about.

    Currently we normalise into: "number", "datetime", or "string".
    """
    text = str(raw or "").strip().lower()
    if text in {"number", "numeric", "float", "double", "integer", "int"}:
        return "number"
    if text in {"datetime", "timestamp", "date", "time"}:
        return "datetime"
    if text in {"category", "categorical"}:
        return "string"
    return "string"


def build_batch_field_catalog_and_stats(
    batches: Sequence[Mapping[str, Any]],
    *,
    time_source: str | None = None,
    categorical_fields: Sequence[str] | None = None,
    numeric_fields: Sequence[str] | None = None,
    field_sources: Mapping[str, str] | None = None,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    rows_values: list[float] = []
    parent_values: list[float] = []
    for raw in batches:
        if not isinstance(raw, Mapping):
            continue
        rows_val = raw.get("rows")
        parent_val = raw.get("parent")
        try:
            rows_num = float(rows_val)
        except Exception:
            rows_num = 0.0
        try:
            parent_num = float(parent_val)
        except Exception:
            parent_num = 0.0
        rows_values.append(rows_num)
        parent_values.append(parent_num)

    def _basic_stats(values: list[float]) -> dict[str, float]:
        if not values:
            return {"min": 0.0, "max": 0.0, "avg": 0.0}
        total = sum(values)
        return {
            "min": float(min(values)),
            "max": float(max(values)),
            "avg": float(total / len(values)),
        }

    stats: dict[str, Any] = {
        "batch_count": len(batches),
        "rows_total": int(sum(rows_values)),
        "rows_stats": _basic_stats(rows_values),
        "parent_stats": _basic_stats(parent_values),
    }

    sources: dict[str, str] = {str(k): str(v) for k, v in (field_sources or {}).items() if str(k).strip()}
    field_catalog: list[dict[str, Any]] = []
    seen: set[str] = set()

    def _add_field(name: str, ftype: str, description: str, *, source: str | None = None):
        key = (name or "").strip()
        if not key or key in seen:
            return
        field_catalog.append(
            {
                "name": key,
                "type": (ftype or "unknown").strip(),
                "description": description,
                "source": source or sources.get(key) or "computed",
            }
        )
        seen.add(key)

    _add_field(
        "batch_index",
        "numeric",
        "1-based index of the batch in discovery order.",
        source="computed",
    )
    _add_field(
        "batch_id",
        "categorical",
        "Batch identifier (composite key from join keys).",
        source="computed",
    )
    _add_field(
        "rows",
        "numeric",
        "Number of child rows in this batch.",
        source=sources.get("rows") or "child_rows",
    )
    _add_field(
        "parent",
        "numeric",
        "Number of parent rows associated with this batch.",
        source=sources.get("parent") or "parent_rows",
    )
    _add_field(
        "rows_per_parent",
        "numeric",
        "Child rows divided by parent rows (if parent is zero, treat as rows).",
        source="computed",
    )

    if time_source:
        time_desc = f"Earliest timestamp per batch sourced from {time_source}."
    else:
        time_desc = "Earliest timestamp associated with the batch (if available)."
    _add_field(
        "time",
        "time",
        time_desc,
        source=time_source or sources.get("time") or "computed",
    )

    cat_fields: list[str] = []
    if categorical_fields:
        for field in categorical_fields:
            text = str(field or "").strip()
            if text and text not in cat_fields:
                cat_fields.append(text)
    if cat_fields:
        primary_cat = cat_fields[0]
        _add_field(
            "category",
            "categorical",
            f"Categorical label derived from key column '{primary_cat}'.",
            source=sources.get(primary_cat) or "computed",
        )
        for field in cat_fields:
            _add_field(
                field,
                "categorical",
                f"Key column '{field}' used to build the batch identifier.",
                source=sources.get(field) or "computed",
            )
    else:
        _add_field(
            "category",
            "categorical",
            "Categorical label derived from key columns (if available).",
            source="computed",
        )

    for field in numeric_fields or []:
        fname = (field or "").strip()
        if not fname or fname in seen:
            continue
        _add_field(
            fname,
            "numeric",
            f"Numeric measure '{fname}' derived from discovery results.",
            source=sources.get(fname) or "computed",
        )

    return field_catalog, stats


def build_batch_metrics(
    batches: Sequence[Mapping[str, Any]],
    batch_metadata: Mapping[str, Any] | None,
    *,
    limit: int | None = None,
    extra_fields: Sequence[str] | None = None,
) -> list[dict[str, Any]]:
    metrics: list[dict[str, Any]] = []
    metadata_lookup: Mapping[str, Any] = batch_metadata if isinstance(batch_metadata, Mapping) else {}
    iterable: Sequence[Mapping[str, Any]] = batches
    if limit is not None:
        iterable = list(iterable)[:limit]

    extras: list[str] = []
    if extra_fields is None:
        seen_extras: set[str] = set()
        for meta in metadata_lookup.values():
            if not isinstance(meta, Mapping):
                continue
            for key in meta.keys():
                key_text = str(key or "").strip()
                if not key_text or key_text in ("time", "category") or key_text in seen_extras:
                    continue
                seen_extras.add(key_text)
        extras = sorted(seen_extras)
    else:
        extras = []
        for field in extra_fields:
            name = str(field or "").strip()
            if not name or name in ("time", "category") or name in extras:
                continue
            extras.append(name)

    for idx, raw in enumerate(iterable, start=1):
        if not isinstance(raw, Mapping):
            continue
        batch_id = raw.get("id")
        rows_val = _coerce_number(raw.get("rows"))
        parent_val = _coerce_number(raw.get("parent"))
        safe_parent = parent_val if parent_val not in (None, 0) else 1.0
        entry: dict[str, Any] = {
            "batch_index": idx,
            "batch_id": str(batch_id) if batch_id is not None else str(idx),
            "rows": rows_val,
            "parent": parent_val,
            "rows_per_parent": rows_val / safe_parent if safe_parent else rows_val,
        }
        meta = metadata_lookup.get(str(batch_id)) if batch_id is not None else None
        if isinstance(meta, Mapping):
            for key in ("time", "category"):
                if key in meta:
                    entry[key] = meta[key]
            for extra in extras:
                if extra in meta:
                    entry[extra] = meta[extra]
        metrics.append(entry)
    return metrics


# Discovery payload schema surfaced to the frontend and chart helpers:
# {
#   "metrics": [
#     {"name": "rows", "type": "number", "description": "...", "bucketable": true},
#     ...
#   ],
#   "dimensions": [
#     {"name": "time", "type": "datetime", "kind": "temporal", "bucketable": true},
#     {"name": "category", "type": "string", "kind": "categorical", "bucketable": false},
#     {"name": "batch_index", "type": "number", "kind": "numeric", "bucketable": true},
#     ...
#   ],
#   "defaults": {"dimension": "time", "metric": "rows"}
# }
def build_discovery_schema(field_catalog: Iterable[Mapping[str, Any]]) -> dict[str, Any]:
    """
    Build a concise list of allowed metrics/dimensions from a field catalog.

    The schema is meant for the frontend to know which fields it can use for resampling
    (including numeric binning) and charting without guessing.
    """
    metrics: list[dict[str, Any]] = []
    dimensions: list[dict[str, Any]] = []
    seen_metric_names: set[str] = set()
    seen_dimension_names: set[str] = set()

    for field in field_catalog or []:
        name = str(field.get("name") or "").strip()
        if not name:
            continue
        normalized_type = _normalize_type(field.get("type"))
        description = str(field.get("description") or "").strip()
        base = {"name": name, "type": normalized_type}
        if description:
            base["description"] = description

        # Numeric fields can be treated both as metrics and as bucketable dimensions.
        if normalized_type == "number" and name not in seen_metric_names:
            metrics.append({**base, "bucketable": True})
            seen_metric_names.add(name)

        if name not in seen_dimension_names:
            kind = "categorical"
            bucketable = False
            if normalized_type == "datetime":
                kind = "temporal"
                bucketable = True
            elif normalized_type == "number":
                kind = "numeric"
                bucketable = True
            dimensions.append({**base, "kind": kind, "bucketable": bucketable})
            seen_dimension_names.add(name)

    def _pick_default_dimension() -> str:
        for preferred in ("time", "timestamp", "date"):
            if any(dim["name"] == preferred for dim in dimensions):
                return preferred
        for fallback in ("category", "batch_index"):
            if any(dim["name"] == fallback for dim in dimensions):
                return fallback
        return dimensions[0]["name"] if dimensions else "batch_index"

    def _pick_default_metric() -> str:
        for preferred in ("rows", "rows_per_parent", "parent"):
            if any(metric["name"] == preferred for metric in metrics):
                return preferred
        return metrics[0]["name"] if metrics else "rows"

    return {
        "metrics": metrics,
        "dimensions": dimensions,
        "defaults": {
            "dimension": _pick_default_dimension(),
            "metric": _pick_default_metric(),
        },
    }


def bin_numeric_metric(
    metrics: Sequence[Mapping[str, Any]],
    metric_name: str,
    *,
    bucket_count: int = 10,
    bucket_edges: Sequence[float] | None = None,
    value_range: tuple[float, float] | None = None,
) -> list[dict[str, Any]]:
    """
    Bucket a numeric metric into ranges and compute aggregates per bucket.

    Returns a list of buckets sorted by range:
      [
        {
          "bucket_index": 0,
          "start": <inclusive lower>,
          "end": <exclusive upper, inclusive for the last bucket>,
          "count": <rows in bucket>,
          "sum": <sum of metric>,
          "min": <min value>,
          "max": <max value>,
          "batch_ids": ["id_1", ...],
        },
        ...
      ]
    """
    if bucket_edges is not None:
        edges = sorted({float(v) for v in bucket_edges if v is not None})
        if len(edges) < 2:
            return []
    else:
        values = [_coerce_number(entry.get(metric_name)) for entry in metrics]
        if not values:
            return []
        min_value = float(values[0])
        max_value = float(values[0])
        for val in values[1:]:
            min_value = min(min_value, val)
            max_value = max(max_value, val)
        if value_range is not None:
            min_value = float(value_range[0])
            max_value = float(value_range[1])
        if max_value < min_value:
            min_value, max_value = max_value, min_value
        if bucket_count < 1:
            bucket_count = 1
        if max_value == min_value:
            edges = [min_value, max_value]
        else:
            step = (max_value - min_value) / bucket_count
            edges = [min_value + step * i for i in range(bucket_count)]
            edges.append(max_value)

    buckets: list[dict[str, Any]] = []
    for idx in range(len(edges) - 1):
        start = edges[idx]
        end = edges[idx + 1]
        buckets.append(
            {
                "bucket_index": idx,
                "start": start,
                "end": end,
                "count": 0,
                "sum": 0.0,
                "min": float("inf"),
                "max": float("-inf"),
                "batch_ids": [],
            }
        )

    for entry in metrics:
        value = _coerce_number(entry.get(metric_name))
        batch_id = entry.get("batch_id") or entry.get("id")
        target_idx = None
        for idx in range(len(edges) - 1):
            start = edges[idx]
            end = edges[idx + 1]
            is_last = idx == len(edges) - 2
            if (value >= start and value < end) or (is_last and value == end):
                target_idx = idx
                break
        if target_idx is None:
            continue
        bucket = buckets[target_idx]
        bucket["count"] += 1
        bucket["sum"] += value
        bucket["min"] = min(bucket["min"], value)
        bucket["max"] = max(bucket["max"], value)
        if batch_id is not None:
            batch_id_text = str(batch_id)
            if batch_id_text not in bucket["batch_ids"]:
                bucket["batch_ids"].append(batch_id_text)

    for bucket in buckets:
        if bucket["min"] == float("inf"):
            bucket["min"] = 0.0
        if bucket["max"] == float("-inf"):
            bucket["max"] = 0.0
    return buckets


def group_metrics_by_field(
    metrics: Sequence[Mapping[str, Any]],
    dimension_field: str,
    *,
    metric_field: str = "rows",
    aggregation: str = "sum",
) -> list[dict[str, Any]]:
    """
    Group metrics by a categorical field and aggregate a numeric metric.

    aggregation: one of ("sum", "avg", "min", "max", "count").
    """
    groups: dict[str, dict[str, Any]] = {}
    agg = aggregation.lower().strip()
    for entry in metrics:
        key_raw = entry.get(dimension_field)
        key = str(key_raw) if key_raw is not None else ""
        metric_value = _coerce_number(entry.get(metric_field))
        bucket = groups.setdefault(
            key,
            {
                "key": key,
                "label": key or "(empty)",
                "sum": 0.0,
                "count": 0,
                "min": float("inf"),
                "max": float("-inf"),
                "batch_ids": [],
            },
        )
        bucket["sum"] += metric_value
        bucket["count"] += 1
        bucket["min"] = min(bucket["min"], metric_value)
        bucket["max"] = max(bucket["max"], metric_value)
        batch_id = entry.get("batch_id") or entry.get("id")
        if batch_id is not None:
            batch_id_text = str(batch_id)
            if batch_id_text not in bucket["batch_ids"]:
                bucket["batch_ids"].append(batch_id_text)

    results: list[dict[str, Any]] = []
    for bucket in groups.values():
        value: float
        if agg == "avg":
            value = bucket["sum"] / bucket["count"] if bucket["count"] else 0.0
        elif agg == "min":
            value = 0.0 if bucket["min"] == float("inf") else bucket["min"]
        elif agg == "max":
            value = 0.0 if bucket["max"] == float("-inf") else bucket["max"]
        elif agg == "count":
            value = float(bucket["count"])
        else:
            value = bucket["sum"]
        results.append(
            {
                **bucket,
                "value": value,
            }
        )

    return sorted(results, key=lambda item: item["label"])


def build_resample_support(
    field_catalog: Iterable[Mapping[str, Any]],
    batch_metrics: Sequence[Mapping[str, Any]],
    *,
    schema: Mapping[str, Any] | None = None,
    default_metric: str | None = None,
    bucket_count: int = 10,
) -> dict[str, Any]:
    """
    Pre-compute numeric bins and categorical groups for the discovery payload.

    Shape:
      {
        "numeric_bins": {metric_name: <bin_numeric_metric output>},
        "category_groups": {dimension_name: <group_metrics_by_field output>}
      }
    """
    schema_data = schema if isinstance(schema, Mapping) else build_discovery_schema(field_catalog)
    numeric_bins: dict[str, list[dict[str, Any]]] = {}
    category_groups: dict[str, list[dict[str, Any]]] = {}
    defaults = schema_data.get("defaults") or {}
    metric_default = default_metric or defaults.get("metric") or "rows"

    for metric in schema_data.get("metrics", []):
        if not isinstance(metric, Mapping):
            continue
        name = str(metric.get("name") or "").strip()
        if not name:
            continue
        if not metric.get("bucketable"):
            continue
        if _normalize_type(metric.get("type")) != "number":
            continue
        numeric_bins[name] = bin_numeric_metric(batch_metrics, name, bucket_count=bucket_count)

    for dim in schema_data.get("dimensions", []):
        if not isinstance(dim, Mapping):
            continue
        name = str(dim.get("name") or "").strip()
        if not name:
            continue
        dim_kind = str(dim.get("kind") or dim.get("type") or "").lower()
        if dim_kind in {"categorical", "string"}:
            category_groups[name] = group_metrics_by_field(
                batch_metrics,
                name,
                metric_field=metric_default,
                aggregation="sum",
            )
        elif dim_kind in {"numeric", "number"} and dim.get("bucketable") and name not in numeric_bins:
            numeric_bins[name] = bin_numeric_metric(batch_metrics, name, bucket_count=bucket_count)

    return {
        "numeric_bins": numeric_bins,
        "category_groups": category_groups,
    }


# ======================================================================
# discovery_excel
# ======================================================================

import logging
import math
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Mapping, Sequence

import pandas as pd

logger = logging.getLogger(__name__)

from backend.app.repositories import SQLiteDataFrameLoader

try:  # pragma: no cover - compatibility shim
    from backend.app.services.contract_builder import build_or_load_contract_v2 as build_or_load_contract  # type: ignore
except Exception as exc:  # pragma: no cover

    def build_or_load_contract(*_args, _exc=exc, **_kwargs):  # type: ignore
        raise RuntimeError(
            "build_or_load_contract unavailable. Ensure contract_builder.build_or_load_contract_v2 exists."
        ) from _exc


_DATE_INPUT_FORMATS = (
    "%Y-%m-%d %H:%M:%S",
    "%Y-%m-%d %H:%M",
    "%Y-%m-%d",
    "%d/%m/%Y %H:%M:%S",
    "%d/%m/%Y %H:%M",
    "%d/%m/%Y",
    "%m/%d/%Y %H:%M:%S",
    "%m/%d/%Y %H:%M",
    "%m/%d/%Y",
)

_DIRECT_COLUMN_RE = re.compile(r"^\s*(?P<table>[A-Za-z_][\w]*)\s*\.\s*(?P<column>[A-Za-z_][\w]*)\s*$")


def _infer_primary_table(mapping_section: Mapping[str, str] | None) -> str | None:
    if not isinstance(mapping_section, Mapping):
        return None
    seen: list[str] = []
    for expr in mapping_section.values():
        if not isinstance(expr, str):
            continue
        match = _DIRECT_COLUMN_RE.match(expr.strip())
        if not match:
            continue
        table_name = match.group("table").strip(' "`[]')
        if not table_name or table_name.lower().startswith("params"):
            continue
        if table_name not in seen:
            seen.append(table_name)
    return seen[0] if seen else None


def _snap_end_of_day(dt: datetime) -> datetime:
    """Snap an end-date to the end of its specified precision.

    - Date-only  (00:00:00.000000) → 23:59:59.999999  (include whole day)
    - HH:MM only (ss=0, us=0)      → HH:MM:59.999999  (include whole minute)
    - HH:MM:SS   (us=0)            → HH:MM:SS.999999  (include whole second)
    - Already has microseconds      → unchanged

    This ensures ``2026-02-19 18:00`` includes records at 18:00:30, etc.
    """
    if dt.microsecond != 0:
        return dt
    if dt.hour == 0 and dt.minute == 0 and dt.second == 0:
        return dt.replace(hour=23, minute=59, second=59, microsecond=999999)
    if dt.second == 0:
        return dt.replace(second=59, microsecond=999999)
    return dt.replace(microsecond=999999)


def _parse_date_like(value) -> datetime | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    iso_try = text.replace("Z", "+00:00")
    if " " in iso_try and "T" not in iso_try:
        iso_try = iso_try.replace(" ", "T", 1)
    try:
        return datetime.fromisoformat(iso_try)
    except ValueError:
        pass
    for fmt in _DATE_INPUT_FORMATS:
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return None


def _stringify_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    text = str(value)
    return text.strip()


def _vectorized_stringify(series: pd.Series) -> pd.Series:
    """Vectorized equivalent of series.apply(_stringify_value)."""
    return series.fillna("").astype(str).str.strip()


def _vectorized_bid(df: pd.DataFrame, columns: list[str]) -> pd.Series:
    """Build batch ID from multiple columns, joined by '|' — vectorized."""
    parts = [df[col].fillna("").astype(str).str.strip() for col in columns]
    return parts[0].str.cat(parts[1:], sep="|")


def _coerce_datetime_series(series: pd.Series) -> pd.Series:
    if pd.api.types.is_datetime64_any_dtype(series):
        result = series
    elif pd.api.types.is_numeric_dtype(series):
        numeric = pd.to_numeric(series, errors="coerce")
        finite = numeric[~pd.isna(numeric)]
        max_abs = finite.abs().max() if not finite.empty else None
        if max_abs is not None and max_abs > 32503680000:
            numeric = numeric / 1000.0
        result = pd.to_datetime(numeric, unit="s", errors="coerce")
    else:
        text = series.astype(str).str.strip()
        result = pd.to_datetime(text, errors="coerce")
    if hasattr(result, "dt"):
        try:
            return result.dt.tz_localize(None)
        except (AttributeError, TypeError):
            return result
    return result


def _apply_date_filter(df: pd.DataFrame, column: str, start: str, end: str) -> pd.DataFrame:
    if df is None or df.empty or not column or column not in df.columns:
        return df
    start_dt = _parse_date_like(start)
    end_dt = _parse_date_like(end)
    if start_dt is None or end_dt is None:
        return df
    end_dt = _snap_end_of_day(end_dt)
    dt_series = _coerce_datetime_series(df[column])
    mask = (dt_series >= start_dt) & (dt_series <= end_dt)
    return df.loc[mask.fillna(False)]


def _normalize_key_value(raw_value: Any) -> list[str]:
    if isinstance(raw_value, (list, tuple, set)):
        values: list[str] = []
        for item in raw_value:
            text = str(item or "").strip()
            if text and text not in values:
                values.append(text)
        return values
    text = str(raw_value or "").strip()
    return [text] if text else []


def _apply_value_filters(df: pd.DataFrame, filters: list[tuple[str, list[str]]]) -> pd.DataFrame:
    if df is None or df.empty or not filters:
        return df
    mask = pd.Series(True, index=df.index)
    for column, values in filters:
        if column not in df.columns:
            continue
        normalized = [str(val).strip() for val in values if str(val or "").strip()]
        if not normalized:
            continue
        series = df[column]
        if pd.api.types.is_numeric_dtype(series):
            try:
                numeric_values = [float(val) for val in normalized]
            except ValueError:
                numeric_values = None
            if numeric_values is not None:
                cmp = pd.to_numeric(series, errors="coerce").isin(numeric_values)
                mask &= cmp.fillna(False)
                continue
        cmp = series.astype(str).str.strip().isin(normalized)
        mask &= cmp.fillna(False)
    return df.loc[mask]


def _build_batch_index(df: pd.DataFrame, key_columns: List[str], *, use_rowid: bool = False) -> tuple[list[str], Counter[str]]:
    if df is None or df.empty:
        return [], Counter()
    working = df.reset_index(drop=True).copy()
    columns = list(key_columns)
    if use_rowid:
        working["__rowid__"] = working.index + 1
        columns = ["__rowid__"]
    if not columns:
        working["__bid__"] = ""
    elif len(columns) == 1:
        col = columns[0]
        if col not in working.columns:
            working[col] = None
        working["__bid__"] = _vectorized_stringify(working[col])
    else:
        for col in columns:
            if col not in working.columns:
                working[col] = None
        working["__bid__"] = _vectorized_bid(working, columns)
    sort_cols = columns or ["__bid__"]
    working_sorted = working.sort_values(sort_cols, kind="mergesort")
    ordered_ids: list[str] = []
    seen: set[str] = set()
    for bid in working_sorted["__bid__"]:
        if bid not in seen:
            seen.add(bid)
            ordered_ids.append(bid)
    counts = Counter(working["__bid__"])
    return ordered_ids, counts


def _attach_batch_id(df: pd.DataFrame, key_columns: List[str], *, use_rowid: bool = False) -> pd.DataFrame:
    """
    Attach a "__bid__" column to the DataFrame using the join key columns.

    The logic mirrors _build_batch_index so later aggregations can group
    by the same batch ids without recomputing them.
    """
    if df is None or df.empty:
        return df

    working = df.reset_index(drop=True).copy()
    columns = list(key_columns)

    if use_rowid:
        working["__rowid__"] = working.index + 1
        columns = ["__rowid__"]

    if not columns:
        working["__bid__"] = ""
        return working

    for col in columns:
        if col not in working.columns:
            working[col] = None

    if len(columns) == 1:
        col = columns[0]
        working["__bid__"] = _vectorized_stringify(working[col])
    else:
        working["__bid__"] = _vectorized_bid(working, columns)

    return working


def _build_batch_metadata(
    df: pd.DataFrame,
    key_columns: List[str],
    *,
    date_column: str | None = None,
    use_rowid: bool = False,
    label_columns: Sequence[str] | None = None,
) -> Dict[str, Dict[str, object]]:
    """
    Derive lightweight per-batch metadata for resampling and charting.

    For each batch id we compute:
      - time:  representative timestamp (earliest in the batch) when a date
               column is available.
      - category: a human-readable label based on the first key column.
      - labels for each key column (first non-empty value per batch).

    The returned mapping is keyed by batch id (the same id produced by
    _build_batch_index).
    """
    if df is None or df.empty:
        return {}

    working = df.reset_index(drop=True).copy()
    columns = list(key_columns)

    if use_rowid:
        working["__rowid__"] = working.index + 1
        columns = ["__rowid__"]

    if not columns:
        working["__bid__"] = ""
    elif len(columns) == 1:
        col = columns[0]
        if col not in working.columns:
            working[col] = None
        working["__bid__"] = _vectorized_stringify(working[col])
    else:
        for col in columns:
            if col not in working.columns:
                working[col] = None
        working["__bid__"] = _vectorized_bid(working, columns)

    metadata: Dict[str, Dict[str, object]] = {}

    # Time dimension: earliest timestamp per batch, if a usable column exists.
    if date_column and date_column in working.columns:
        try:
            dt_series = _coerce_datetime_series(working[date_column])
        except Exception:  # pragma: no cover - defensive
            dt_series = None
        if dt_series is not None:
            working["_nr_time"] = dt_series
            grouped_time = working.groupby("__bid__")["_nr_time"].min()
            for bid, ts in grouped_time.items():
                # pandas uses NaT for missing values; treat as absent.
                if ts is None or pd.isna(ts):
                    continue
                try:
                    # Both pandas.Timestamp and datetime.datetime implement isoformat.
                    iso_value = ts.isoformat()
                except Exception:  # pragma: no cover - defensive
                    continue
                metadata.setdefault(bid, {})["time"] = iso_value

    label_cols = [col for col in (label_columns or key_columns or []) if col and not str(col).startswith("__")]
    category_source = label_cols[0] if label_cols else None
    if category_source is None and columns and not str(columns[0]).startswith("__"):
        category_source = columns[0]

    for col in label_cols:
        if col not in working.columns:
            continue
        label_field = f"_nr_label_{col}"
        working[label_field] = _vectorized_stringify(working[col])
        grouped = working.groupby("__bid__")[label_field].first()
        for bid, raw_val in grouped.items():
            text = _stringify_value(raw_val)
            if not text:
                continue
            meta = metadata.setdefault(bid, {})
            meta[col] = text
            if category_source == col and "category" not in meta:
                meta["category"] = text

    if category_source and category_source in working.columns:
        working["_nr_category"] = _vectorized_stringify(working[category_source])
        grouped_cat = working.groupby("__bid__")["_nr_category"].first()
        for bid, cat in grouped_cat.items():
            text = _stringify_value(cat)
            if not text:
                continue
            metadata.setdefault(bid, {}).setdefault("category", text)

    return metadata


def discover_batches_and_counts(
    *,
    db_path: Path,
    contract: dict,
    start_date: str,
    end_date: str,
    key_values: Mapping[str, Any] | None = None,
) -> dict:
    """
    Discover distinct batch IDs and count:
      - parent: number of parent rows (i.e., batches) per id in range
      - rows:   number of child rows per id in range
    Returns:
      {
        "batches": [{"id": "...", "parent": <int>, "rows": <int>}...],
        "batches_count": <int>,   # number of parent batches
        "rows_total": <int>       # sum of child rows
      }
    """
    adapter = ContractAdapter(contract)
    join_cfg = contract.get("join") or {}
    date_columns = adapter.date_columns or (contract.get("date_columns") or {})
    mapping_section = adapter.mapping or (contract.get("mapping") or {})

    parent_table = adapter.parent_table or (join_cfg.get("parent_table") or "").strip()
    child_table = adapter.child_table or (join_cfg.get("child_table") or "").strip()
    parent_key = adapter.parent_key if adapter.parent_key is not None else join_cfg.get("parent_key")
    child_key = adapter.child_key if adapter.child_key is not None else join_cfg.get("child_key")

    if not parent_table:
        inferred_parent = _infer_primary_table(mapping_section)
        if inferred_parent:
            parent_table = inferred_parent

    if not parent_table:
        raise ValueError("contract.join.parent_table is required for discovery")

    parent_date = (date_columns.get(parent_table) or "").strip()
    child_date = (date_columns.get(child_table) or "").strip() if child_table else ""

    def _split_keys(raw_keys: object) -> List[str]:
        """
        Normalise the join key field from the contract. Handles strings, comma-separated
        strings, iterables, and scalars (e.g., integers). Returns a list of non-empty strings.
        """
        if raw_keys is None:
            return []
        if isinstance(raw_keys, (list, tuple, set)):
            items = raw_keys
        else:
            text = str(raw_keys).strip()
            if not text:
                return []
            if isinstance(raw_keys, str):
                if "," in text:
                    items = text.split(",")
                elif "|" in text:
                    items = text.split("|")
                else:
                    items = [text]
            else:
                items = [text]
        result: List[str] = []
        for item in items:
            token = str(item).strip()
            if token:
                result.append(token)
        return result

    pcols = _split_keys(parent_key)
    ccols = _split_keys(child_key)
    use_rowid = False

    if not pcols:
        if ccols:
            pcols = list(ccols)
        else:
            use_rowid = True
            pcols = ["__rowid__"]

    has_child = bool(child_table and ccols)

    categorical_fields: list[str] = []
    for col in pcols:
        col_text = str(col or "").strip()
        if col_text and not col_text.startswith("__") and col_text not in categorical_fields:
            categorical_fields.append(col_text)
    if has_child:
        for col in ccols:
            col_text = str(col or "").strip()
            if col_text and not col_text.startswith("__") and col_text not in categorical_fields:
                categorical_fields.append(col_text)

    # Support pre-built loaders (MultiDataFrameLoader), PostgreSQL, and SQLite
    if hasattr(db_path, 'table_names') and callable(db_path.table_names):
        loader = db_path
    elif hasattr(db_path, 'is_postgresql') and db_path.is_postgresql:
        from backend.app.services.connection_utils import get_loader_for_ref
        loader = get_loader_for_ref(db_path)
    else:
        loader = SQLiteDataFrameLoader(db_path)

    if not isinstance(key_values, Mapping):
        key_values = {}
    parent_filters: list[tuple[str, Any]] = []
    child_filters: list[tuple[str, Any]] = []

    def _split_table_and_column(expr_text: str, fallback_table: str) -> tuple[str, str]:
        expr_text = expr_text.strip()
        if "." in expr_text:
            table_name, column_name = expr_text.split(".", 1)
        else:
            table_name, column_name = fallback_table, expr_text
        table_name = table_name.strip(' "`[]').lower()
        column_name = column_name.strip(' "`[]')
        return table_name, column_name

    def _append_filter_target(expr_text: str, value: Any, collection: list[tuple[str, Any]], expected_table: str):
        if not expr_text or not expected_table:
            return
        table_name, column_name = _split_table_and_column(expr_text, expected_table)
        if table_name != expected_table.lower():
            return
        if not column_name:
            return
        normalized = _normalize_key_value(value)
        if not normalized:
            return
        for idx, (col, existing_values) in enumerate(collection):
            if col == column_name:
                existing_list = _normalize_key_value(existing_values)
                merged = existing_list + [item for item in normalized if item not in existing_list]
                collection[idx] = (col, merged)
                break
        else:
            collection.append((column_name, normalized))

    if key_values:
        parent_table_lc = parent_table.lower()
        child_table_lc = child_table.lower() if child_table else ""
        for token, raw_value in key_values.items():
            if raw_value is None:
                continue
            expr = mapping_section.get(token)
            if not isinstance(expr, str):
                continue
            expr_text = expr.strip()
            if not expr_text or expr_text.upper().startswith("PARAM:"):
                continue
            if "." not in expr_text:
                continue
            table_name, column_name = expr_text.split(".", 1)
            table_name = table_name.strip(' "`[]')
            column_name = column_name.strip(' "`[]')
            if not column_name:
                continue
            table_key = table_name.lower()
            if table_key == parent_table_lc:
                values = _normalize_key_value(raw_value)
                if values:
                    parent_filters.append((column_name, values))
            if has_child and table_key == child_table_lc:
                values = _normalize_key_value(raw_value)
                if values:
                    child_filters.append((column_name, values))

    if isinstance(key_values, Mapping):
        for token, expr in adapter.required_filters.items():
            value = key_values.get(token)
            if value is None:
                continue
            _append_filter_target(expr, value, parent_filters, parent_table)
            if has_child:
                _append_filter_target(expr, value, child_filters, child_table)
        for token, expr in adapter.optional_filters.items():
            value = key_values.get(token)
            if value in (None, ""):
                continue
            _append_filter_target(expr, value, parent_filters, parent_table)
            if has_child:
                _append_filter_target(expr, value, child_filters, child_table)
    parent_filter_pairs: list[tuple[str, list[str]]] = [
        (col, _normalize_key_value(values)) for col, values in parent_filters
    ]
    child_filter_pairs: list[tuple[str, list[str]]] = [
        (col, _normalize_key_value(values)) for col, values in child_filters
    ]
    # Load tables with SQL-level date pre-filtering when a date column is known.
    # This avoids loading millions of rows into memory for large tables.
    try:
        if parent_date and (start_date or end_date) and hasattr(loader, 'frame_date_filtered'):
            parent_df = loader.frame_date_filtered(parent_table, parent_date, start_date, end_date).copy()
        else:
            parent_df = loader.frame(parent_table).copy()
    except Exception as exc:  # pragma: no cover - surfaced to caller
        raise RuntimeError(f"Failed to load parent table {parent_table!r}: {exc}") from exc

    child_df = None
    if has_child:
        try:
            if child_date and (start_date or end_date) and hasattr(loader, 'frame_date_filtered'):
                child_df = loader.frame_date_filtered(child_table, child_date, start_date, end_date).copy()
            else:
                child_df = loader.frame(child_table).copy()
        except Exception as exc:  # pragma: no cover - surfaced to caller
            raise RuntimeError(f"Failed to load child table {child_table!r}: {exc}") from exc

    # Apply DataFrame-level date filter as a safety net (handles timezone stripping, snap, etc.)
    parent_df = _apply_date_filter(parent_df, parent_date, start_date, end_date)
    parent_df = _apply_value_filters(parent_df, parent_filter_pairs)

    if has_child and child_df is not None:
        child_df = _apply_date_filter(child_df, child_date, start_date, end_date)
        child_df = _apply_value_filters(child_df, child_filter_pairs)

    parent_df = _attach_batch_id(parent_df, pcols, use_rowid=use_rowid)
    parent_ids, parent_counts = _build_batch_index(parent_df, pcols, use_rowid=use_rowid)
    if has_child and child_df is not None:
        child_df = _attach_batch_id(child_df, ccols, use_rowid=False)
        child_ids, child_counts = _build_batch_index(child_df, ccols)
    else:
        child_ids, child_counts = [], Counter()

    if not parent_ids and child_ids:
        parent_ids = list(child_ids)

    # Fast path: when there are a huge number of batches (e.g. 100K+ with
    # rowid-based keys), building per-batch metadata/aggregates/metrics is
    # extremely slow and the frontend cannot display them anyway.  Return a
    # lightweight response with counts and a truncated batch list.
    _DISCOVERY_BATCH_LIMIT = 10_000
    if len(parent_ids) > _DISCOVERY_BATCH_LIMIT:
        logger.info(
            "discovery_fast_path total_batches=%d limit=%d",
            len(parent_ids), _DISCOVERY_BATCH_LIMIT,
        )
        batches: List[Dict[str, object]] = []
        rows_total = 0
        for bid in parent_ids:
            parent_cnt = int(parent_counts.get(bid, 0))
            child_cnt = int(child_counts.get(bid, 0)) if has_child else parent_cnt
            rows_total += child_cnt
            if len(batches) < _DISCOVERY_BATCH_LIMIT:
                batches.append({"id": bid, "parent": parent_cnt, "rows": child_cnt, "selected": True})
        return {
            "batches": batches,
            "batches_count": len(parent_ids),
            "rows_total": rows_total,
            "batch_metadata": {},
            "field_catalog": [],
            "batch_metrics": [],
            "discovery_schema": {},
            "numeric_bins": {},
            "category_groups": {},
            "data_stats": {
                "batch_count": len(parent_ids),
                "rows_total": rows_total,
                "rows_stats": {},
                "parent_stats": {},
            },
            "_truncated": True,
        }

    # Lightweight metadata for resampling (time/category per batch id).
    batch_metadata: Dict[str, Dict[str, object]] = {}
    try:
        parent_meta = _build_batch_metadata(
            parent_df,
            pcols,
            date_column=parent_date or None,
            use_rowid=use_rowid,
            label_columns=pcols,
        )
        for bid, meta in parent_meta.items():
            batch_metadata.setdefault(bid, {}).update(meta)
    except Exception:  # pragma: no cover - defensive
        batch_metadata = {}

    if has_child and child_df is not None:
        try:
            child_meta = _build_batch_metadata(
                child_df,
                ccols,
                date_column=child_date or None,
                use_rowid=False,
                label_columns=ccols,
            )
            for bid, meta in child_meta.items():
                target = batch_metadata.setdefault(bid, {})
                if not isinstance(meta, Mapping):
                    continue
                for key, value in meta.items():
                    if value is None or (isinstance(value, str) and not str(value).strip()):
                        continue
                    if key == "time":
                        if "time" not in target:
                            target["time"] = value
                        continue
                    if key not in target:
                        target[key] = value
        except Exception:  # pragma: no cover - defensive
            # If metadata from child fails, keep whatever we have from parent.
            pass

    def _aggregate_numeric(df: pd.DataFrame | None, prefix: str) -> dict[str, dict[str, float]]:
        if df is None or df.empty or "__bid__" not in df.columns:
            return {}
        skip_cols = {"__bid__", "__rowid__"} | set(pcols) | set(ccols) | {parent_date, child_date}
        numeric_cols = [
            col for col in df.columns if col not in skip_cols and pd.api.types.is_numeric_dtype(df[col])
        ]
        if not numeric_cols:
            return {}
        # Vectorized: coerce all numeric columns at once, then groupby.sum()
        coerced = df[["__bid__"] + numeric_cols].copy()
        for col in numeric_cols:
            coerced[col] = pd.to_numeric(coerced[col], errors="coerce")
        summed = coerced.groupby("__bid__")[numeric_cols].sum()
        aggregates: dict[str, dict[str, float]] = {}
        for bid, row in summed.iterrows():
            entry = {f"{prefix}{col}": float(row[col]) for col in numeric_cols if pd.notna(row[col])}
            if entry:
                aggregates[str(bid)] = entry
        return aggregates

    def _aggregate_business_metrics(df: pd.DataFrame | None) -> dict[str, dict[str, float]]:
        if df is None or df.empty or "__bid__" not in df.columns:
            return {}
        metric_sources = {
            "revenue": "total_amount",
            "margin": "margin_amount",
            "cost": "cost_amount",
        }
        available = {name: col for name, col in metric_sources.items() if col in df.columns}
        if not available:
            return {}
        # Vectorized: coerce once, then groupby.agg()
        cols = list(available.values())
        coerced = df[["__bid__"] + cols].copy()
        for col in cols:
            coerced[col] = pd.to_numeric(coerced[col], errors="coerce")
        grouped = coerced.groupby("__bid__")
        summed = grouped[cols].sum()
        aggregates: dict[str, dict[str, float]] = {}
        revenue_col = metric_sources.get("revenue")
        has_revenue = revenue_col and revenue_col in cols
        mean_revenue = grouped[revenue_col].mean() if has_revenue else None
        for bid, row in summed.iterrows():
            entry: dict[str, float] = {}
            for metric_name, column in available.items():
                entry[metric_name] = float(row[column])
            if has_revenue and mean_revenue is not None:
                entry["avg_order_value"] = float(mean_revenue.loc[bid])
            if entry:
                aggregates[str(bid)] = entry
        return aggregates

    aggregated_metrics: dict[str, dict[str, float]] = {}
    parent_aggs = _aggregate_numeric(parent_df, "parent_")
    for bid, vals in parent_aggs.items():
        aggregated_metrics.setdefault(bid, {}).update(vals)
    if has_child and child_df is not None:
        child_aggs = _aggregate_numeric(child_df, "child_")
        for bid, vals in child_aggs.items():
            aggregated_metrics.setdefault(bid, {}).update(vals)
        business_aggs = _aggregate_business_metrics(child_df)
        for bid, vals in business_aggs.items():
            aggregated_metrics.setdefault(bid, {}).update(vals)

    for bid, metrics_map in aggregated_metrics.items():
        target = batch_metadata.setdefault(bid, {})
        for key, value in metrics_map.items():
            target[key] = value

    batches: List[Dict[str, object]] = []
    rows_total = 0

    for bid in parent_ids:
        parent_cnt = int(parent_counts.get(bid, 0))
        if has_child:
            child_cnt = int(child_counts.get(bid, 0))
        else:
            child_cnt = parent_cnt
        rows_total += child_cnt
        batches.append({"id": bid, "parent": parent_cnt, "rows": child_cnt})

    if parent_date:
        time_source = f"{parent_table}.{parent_date}" if parent_table else parent_date
    elif child_date:
        time_source = f"{child_table}.{child_date}" if child_table else child_date
    else:
        time_source = None

    business_metric_names = {"revenue", "margin", "avg_order_value", "cost"}
    present_business_metrics: set[str] = set()
    for metrics_map in aggregated_metrics.values():
        for key in metrics_map:
            if key in business_metric_names:
                present_business_metrics.add(key)

    field_sources: dict[str, str] = {
        "batch_index": "discovery_order",
        "batch_id": "composite_key",
        "rows_per_parent": "computed",
    }
    if parent_table:
        field_sources["parent"] = parent_table
    if child_table:
        field_sources["rows"] = child_table
    elif parent_table:
        field_sources["rows"] = parent_table
    if time_source:
        field_sources["time"] = time_source
    if categorical_fields:
        for col in categorical_fields:
            if col in pcols and parent_table:
                field_sources.setdefault(col, f"{parent_table}.{col}")
            elif col in ccols and child_table:
                field_sources.setdefault(col, f"{child_table}.{col}")
        primary = categorical_fields[0]
        if primary in field_sources:
            field_sources.setdefault("category", field_sources.get(primary, "computed"))
    if child_table:
        child_source_map = {
            "revenue": f"{child_table}.total_amount",
            "margin": f"{child_table}.margin_amount",
            "avg_order_value": f"{child_table}.total_amount",
            "cost": f"{child_table}.cost_amount",
        }
        for metric_name, source in child_source_map.items():
            if metric_name in present_business_metrics:
                field_sources.setdefault(metric_name, source)

    # Collect extra numeric fields for metrics/catalog.
    extra_numeric_fields: list[str] = []
    for metric_fields in aggregated_metrics.values():
        for field_name in metric_fields.keys():
            if field_name not in extra_numeric_fields:
                extra_numeric_fields.append(field_name)

    field_catalog, stats = build_batch_field_catalog_and_stats(
        batches,
        time_source=time_source,
        categorical_fields=categorical_fields,
        numeric_fields=["rows", "parent", "rows_per_parent", *extra_numeric_fields],
        field_sources=field_sources,
    )

    def _collect_metric_stats(metric_name: str) -> dict[str, float] | None:
        values: list[float] = []
        for metrics_map in aggregated_metrics.values():
            if metric_name not in metrics_map:
                continue
            try:
                values.append(float(metrics_map[metric_name]))
            except Exception:
                continue
        if not values:
            return None
        total_val = float(sum(values))
        return {
            "min": float(min(values)),
            "max": float(max(values)),
            "avg": float(total_val / len(values)),
            "total": total_val,
        }

    business_metric_stats: dict[str, dict[str, float]] = {}
    for metric_name in sorted(present_business_metrics):
        metric_stat = _collect_metric_stats(metric_name)
        if metric_stat:
            business_metric_stats[metric_name] = metric_stat
    if business_metric_stats:
        stats.setdefault("metrics_stats", {}).update(business_metric_stats)

    discovery_schema = build_discovery_schema(field_catalog)
    if isinstance(discovery_schema, dict):
        default_metric_candidates = ["revenue", "margin", "avg_order_value", "cost"]
        metrics_list = discovery_schema.get("metrics") or []
        defaults = discovery_schema.setdefault("defaults", {})
        for candidate in default_metric_candidates:
            if any(m.get("name") == candidate for m in metrics_list):
                defaults["metric"] = candidate
                break

    batch_metrics = build_batch_metrics(
        batches,
        batch_metadata,
        extra_fields=[*categorical_fields, *extra_numeric_fields],
    )
    resample_support = build_resample_support(
        field_catalog,
        batch_metrics,
        schema=discovery_schema,
        default_metric=discovery_schema.get("defaults", {}).get("metric"),
        bucket_count=10,
    )

    return {
        "batches": batches,
        "batches_count": len(batches),
        "rows_total": rows_total,
        "batch_metadata": batch_metadata,
        "field_catalog": field_catalog,
        "batch_metrics": batch_metrics,
        "discovery_schema": discovery_schema,
        "numeric_bins": resample_support["numeric_bins"],
        "category_groups": resample_support["category_groups"],
        "data_stats": stats,
    }


# ======================================================================
# discovery (legacy wrapper removed — discover_batches_and_counts is
# defined in the discovery_excel section above and serves both PDF and
# Excel pipelines)
# ======================================================================


# ======================================================================
# ReportGenerate
# ======================================================================

import asyncio
import contextlib
import json
import logging
import os
import re
import subprocess
import sys as _sys
from backend.app.repositories import SQLiteDataFrameLoader
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from decimal import Decimal, InvalidOperation
from itertools import product
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

logger = logging.getLogger(__name__)

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:
    import numpy as np
except ImportError:  # pragma: no cover
    np = None  # type: ignore

try:
    import cv2
except ImportError:  # pragma: no cover
    cv2 = None  # type: ignore

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore

try:
    from skimage.metrics import structural_similarity as ssim
except ImportError:  # pragma: no cover
    ssim = None  # type: ignore

try:
    from playwright.async_api import async_playwright
except ImportError:  # pragma: no cover
    async_playwright = None  # type: ignore


def _run_async(coro):
    """Run an async coroutine safely whether or not an event loop is running."""
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        loop = None
    if loop and loop.is_running():
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            return pool.submit(asyncio.run, coro).result()
    return asyncio.run(coro)


_PDF_WORKER_SCRIPT = str(Path(__file__).with_name("_pdf_worker.py"))


def _pdf_worker_mp_target(html_path: str, pdf_path: str, base_dir: str, pdf_scale: float | None) -> None:
    """Target function for multiprocessing.Process — runs _convert in a fresh process."""
    asyncio.run(_convert(
        html_path=html_path,
        pdf_path=pdf_path,
        base_dir=base_dir,
        pdf_scale=pdf_scale,
    ))


# Timeout for the PDF worker process (30 minutes — large chunked docs with 10M+ rows).
_PDF_PROCESS_TIMEOUT = int(os.environ.get("NEURA_PDF_PROCESS_TIMEOUT", "3600"))


def _html_to_pdf_subprocess(
    html_path: Path, pdf_path: Path, base_dir: Path, pdf_scale: float | None = None
) -> None:
    """Convert HTML to PDF by running Playwright in a dedicated subprocess.

    This avoids the SIGCHLD / asyncio event-loop conflict that occurs when
    ``asyncio.run()`` is called from a non-main thread inside uvicorn.

    In PyInstaller frozen mode, sys.executable is the bundled exe which
    cannot run .py scripts.  We use multiprocessing.Process instead so the
    PDF work runs in a separate OS process — freeing the GIL and preventing
    the main backend from stalling during large chunked renders.
    """
    # PyInstaller frozen mode: use multiprocessing.Process (requires freeze_support)
    if getattr(_sys, "frozen", False):
        import multiprocessing

        args = (
            str(html_path.resolve()),
            str(pdf_path.resolve()),
            str((base_dir or html_path.parent).resolve()),
            pdf_scale,
        )
        proc = multiprocessing.Process(
            target=_pdf_worker_mp_target,
            args=args,
            daemon=False,
        )
        proc.start()
        proc.join(timeout=_PDF_PROCESS_TIMEOUT)
        if proc.is_alive():
            logger.error("PDF worker process timed out after %ds, terminating", _PDF_PROCESS_TIMEOUT)
            proc.terminate()
            proc.join(timeout=10)
            raise RuntimeError(f"PDF worker process timed out after {_PDF_PROCESS_TIMEOUT}s")
        if proc.exitcode != 0:
            raise RuntimeError(f"PDF worker process failed with exit code {proc.exitcode}")
        return

    import json as _json

    args_json = _json.dumps({
        "html_path": str(html_path.resolve()),
        "pdf_path": str(pdf_path.resolve()),
        "base_dir": str((base_dir or html_path.parent).resolve()),
        "pdf_scale": pdf_scale,
    })

    env = {**os.environ}
    if "TMPDIR" not in env:
        home_tmp = Path.home() / ".tmp"
        if home_tmp.is_dir():
            env["TMPDIR"] = str(home_tmp)

    result = subprocess.run(
        [_sys.executable, _PDF_WORKER_SCRIPT, args_json],
        capture_output=True,
        text=True,
        timeout=_PDF_PROCESS_TIMEOUT,
        env=env,
    )
    if result.returncode != 0:
        stderr_tail = (result.stderr or "")[-2000:]
        raise RuntimeError(f"PDF subprocess failed:\n{stderr_tail}")


_BRAND_STYLE_RE = re.compile(r'<style\s+id="brand-kit-style"[^>]*>.*?</style>', re.DOTALL)


def _inject_brand_css(html: str, css_block: str) -> str:
    """Inject (or replace) a brand-kit ``<style>`` block into an HTML string.

    Strategy mirrors the Excel print-style injection in ReportGenerateExcel:
    replace an existing ``<style id="brand-kit-style">`` block if present,
    otherwise insert just before ``</head>``, or prepend to the document.
    """
    if _BRAND_STYLE_RE.search(html):
        return _BRAND_STYLE_RE.sub(css_block, html, count=1)
    head_close = re.search(r"(?is)</head>", html)
    if head_close:
        idx = head_close.start()
        return f"{html[:idx]}{css_block}{html[idx:]}"
    return f"{css_block}{html}"


_DATE_PARAM_START_ALIASES = {
    "start_ts_utc",
    "start_ts",
    "start_timestamp",
    "start_datetime",
    "start_date",
    "start_dt",
    "start_iso",
    "start_date_utc",
    "from_ts_utc",
    "from_ts",
    "from_timestamp",
    "from_datetime",
    "from_date",
    "from_dt",
    "from_iso",
    "from_date_utc",
    "range_start",
    "period_start",
}

_DATE_PARAM_END_ALIASES = {
    "end_ts_utc",
    "end_ts",
    "end_timestamp",
    "end_datetime",
    "end_date",
    "end_dt",
    "end_iso",
    "end_date_utc",
    "to_ts_utc",
    "to_ts",
    "to_timestamp",
    "to_datetime",
    "to_date",
    "to_dt",
    "to_iso",
    "to_date_utc",
    "range_end",
    "period_end",
}


# ======================================================
# ENTRYPOINT: DB-driven fill + PDF (no LLM here anymore)
# ======================================================
def fill_and_print(
    OBJ: dict,
    TEMPLATE_PATH: Path,
    DB_PATH: Path,
    OUT_HTML: Path,
    OUT_PDF: Path,
    START_DATE: str,
    END_DATE: str,
    batch_ids: list[str] | None = None,
    KEY_VALUES: dict | None = None,
    __force_single: bool = False,
    BRAND_KIT_ID: str | None = None,
):
    """
    DB-driven renderer:
      - Assumes TEMPLATE_PATH is already the *final shell* produced at Approve (auto_fill.py)
        containing a single prototype batch block.
      - Renders header tokens (parent row per batch), row repeater (child rows), totals, literals.
      - Writes OUT_HTML and prints OUT_PDF via Playwright.

    API contract preserved (same signature).
    """

    # ---- Guard required inputs ----
    for name in ("OBJ", "TEMPLATE_PATH", "DB_PATH", "START_DATE", "END_DATE"):
        if locals().get(name) is None:
            raise NameError(f"Missing required variable: `{name}`")

    # Ensure output dir exists
    OUT_DIR = OUT_HTML.parent
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    def _log_debug(*parts: object) -> None:
        message = " ".join(str(part) for part in parts)
        logger.debug(message)

    import time as _time
    _fp_start = _time.time()
    def _fp_progress(stage: str) -> None:
        elapsed = _time.time() - _fp_start
        print(f"[REPORT] {stage} ({elapsed:.1f}s)", flush=True)

    _fp_progress("fill_and_print START")
    _log_debug(
        "=== fill_and_print call ===",
        "force_single" if __force_single else "fanout_root",
        "KEY_VALUES raw=",
        KEY_VALUES or {},
    )

    # ---- Load the final shell HTML (created during Approve) ----
    from backend.app.services.infra_services import _fix_fixed_footers
    html = _fix_fixed_footers(TEMPLATE_PATH.read_text(encoding="utf-8"))

    # ---- Inject brand kit CSS if requested ----
    if BRAND_KIT_ID:
        try:
            from backend.app.services.ai_services import design_service
            brand_css = design_service.generate_brand_css_from_id(BRAND_KIT_ID)
            if brand_css:
                html = _inject_brand_css(html, brand_css)
                _log_debug("Brand kit CSS injected:", BRAND_KIT_ID)
        except Exception:
            logger.warning("Failed to inject brand kit CSS", exc_info=True)

    # Support pre-built loaders (MultiDataFrameLoader), PostgreSQL, and SQLite
    if hasattr(DB_PATH, 'table_names') and callable(DB_PATH.table_names):
        dataframe_loader = DB_PATH  # pre-built loader (e.g. MultiDataFrameLoader)
    elif hasattr(DB_PATH, 'is_postgresql') and DB_PATH.is_postgresql:
        from backend.app.services.connection_utils import get_loader_for_ref
        dataframe_loader = get_loader_for_ref(DB_PATH)
    else:
        dataframe_loader = SQLiteDataFrameLoader(DB_PATH)

    TOKEN_RE = re.compile(r"\{\{?\s*([A-Za-z0-9_\-\.]+)\s*\}\}?")
    TEMPLATE_TOKENS = {m.group(1) for m in TOKEN_RE.finditer(html)}

    # ---- Unpack contract ----
    OBJ = OBJ or {}
    contract_adapter = ContractAdapter(OBJ)
    param_token_set = {token for token in (contract_adapter.param_tokens or []) if token}

    PLACEHOLDER_TO_COL = contract_adapter.mapping

    # ---- Validate contract against live schema (detect drift) ----
    if PLACEHOLDER_TO_COL:
        _available_tables = set(dataframe_loader.table_names())
        _col_ref_re = re.compile(r"^([A-Za-z_][\w]*)\.([A-Za-z_][\w]*)$")
        _missing_refs: list[str] = []
        _columns_cache: dict[str, set[str]] = {}
        for _token, _col_ref in PLACEHOLDER_TO_COL.items():
            _m = _col_ref_re.match(str(_col_ref))
            if not _m:
                continue  # not a table.column ref — skip
            _tbl, _col = _m.group(1), _m.group(2)
            if _tbl not in _available_tables:
                _missing_refs.append(f"  {_token!r} -> {_col_ref!r} (table {_tbl!r} not found)")
                continue
            if _tbl not in _columns_cache:
                try:
                    # Use PRAGMA to get column names without loading data
                    if hasattr(dataframe_loader, 'column_names'):
                        _columns_cache[_tbl] = set(dataframe_loader.column_names(_tbl))
                    else:
                        _columns_cache[_tbl] = set(dataframe_loader.frame(_tbl).columns)
                except Exception:
                    _columns_cache[_tbl] = set()
            if _col not in _columns_cache[_tbl] and not _col.startswith("__"):
                _missing_refs.append(
                    f"  {_token!r} -> {_col_ref!r} (column {_col!r} not in table {_tbl!r})"
                )
        if _missing_refs:
            detail = "\n".join(_missing_refs)
            raise RuntimeError(
                f"Contract references columns that no longer exist in the database.\n"
                f"Re-approve the template mapping to fix this.\n\n"
                f"Missing references:\n{detail}"
            )

    join_raw = OBJ.get("join", {}) or {}
    JOIN = {
        "parent_table": contract_adapter.parent_table or join_raw.get("parent_table", ""),
        "child_table": contract_adapter.child_table or join_raw.get("child_table", ""),
        "parent_key": contract_adapter.parent_key or join_raw.get("parent_key", ""),
        "child_key": contract_adapter.child_key or join_raw.get("child_key", ""),
    }

    DATE_COLUMNS = contract_adapter.date_columns or (OBJ.get("date_columns", {}) or {})

    HEADER_TOKENS = contract_adapter.scalar_tokens or OBJ.get("header_tokens", [])
    ROW_TOKENS = contract_adapter.row_tokens or OBJ.get("row_tokens", [])
    TOTALS = contract_adapter.totals_mapping or OBJ.get("totals", {})
    # If totals is empty but totals_math has keys, use those as TOTALS markers
    if not TOTALS and contract_adapter.totals_math:
        TOTALS = {k: "COMPUTED" for k in contract_adapter.totals_math}
    ROW_ORDER = contract_adapter.row_order or OBJ.get("row_order", ["ROWID"])
    LITERALS = {
        str(token): "" if value is None else str(value) for token, value in (OBJ.get("literals", {}) or {}).items()
    }
    FORMATTERS = contract_adapter.formatters
    key_values_map: dict[str, list[str]] = {}
    if KEY_VALUES:
        for token, raw_value in KEY_VALUES.items():
            name = str(token or "").strip()
            if not name:
                continue
            values: list[str] = []
            if isinstance(raw_value, (list, tuple, set)):
                seen = set()
                for item in raw_value:
                    text = str(item or "").strip()
                    if text and text not in seen:
                        seen.add(text)
                        values.append(text)
            else:
                text = str(raw_value or "").strip()
                if text:
                    values = [text]
            if values:
                key_values_map[name] = values

    _DIRECT_COLUMN_RE = re.compile(r"^(?P<table>[A-Za-z_][\w]*)\.(?P<column>[A-Za-z_][\w]*)$")
    _SQL_IDENT_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")

    def _safe_ident(name: str) -> str:
        if _SQL_IDENT_RE.match(name):
            return name
        safe = str(name).replace('"', '""')
        return f'"{safe}"'

    def _resolve_token_column(token: str) -> tuple[str, str] | None:
        mapping_expr = PLACEHOLDER_TO_COL.get(token)
        if isinstance(mapping_expr, str):
            match = _DIRECT_COLUMN_RE.match(mapping_expr.strip())
            if match:
                return match.group("table"), match.group("column")
        required_filters = contract_adapter.required_filters
        optional_filters = contract_adapter.optional_filters
        filter_expr = (required_filters.get(token) or optional_filters.get(token) or "").strip()
        match = _DIRECT_COLUMN_RE.match(filter_expr)
        if match:
            return match.group("table"), match.group("column")
        return None

    def _canonicalize_case(table: str, column: str, raw_value: str) -> str:
        normalized_table = str(table or "").strip().lower()
        normalized_column = str(column or "").strip().lower()
        normalized_value = str(raw_value or "").strip()
        cache_key = (normalized_table, normalized_column, normalized_value.lower())
        if cache_key in _canonicalize_cache:
            return _canonicalize_cache[cache_key]
        canonical = normalized_value
        if not normalized_table or not normalized_column or not normalized_value:
            _canonicalize_cache[cache_key] = canonical
            return canonical
        try:
            frame = dataframe_loader.frame(table)
        except Exception:
            _canonicalize_cache[cache_key] = canonical
            return canonical
        if column not in frame.columns:
            _canonicalize_cache[cache_key] = canonical
            return canonical
        series = frame[column]
        try:
            matches = series.dropna().astype(str)
        except Exception:
            matches = series.dropna().apply(lambda v: str(v))
        lower_target = normalized_value.lower()
        mask = matches.str.lower() == lower_target
        filtered = matches[mask]
        if not filtered.empty:
            canonical = str(filtered.iloc[0])
        _canonicalize_cache[cache_key] = canonical
        return canonical

    _canonicalize_cache: dict[tuple[str, str, str], str] = {}

    for token, values in list(key_values_map.items()):
        resolved = _resolve_token_column(token)
        if not resolved:
            continue
        table_name, column_name = resolved
        if not table_name or not column_name:
            continue
        updated_values: list[str] = []
        changed = False
        for value in values:
            if not isinstance(value, str) or not value.strip():
                updated_values.append(value)
                continue
            canon = _canonicalize_case(table_name, column_name, value.strip())
            if canon != value:
                changed = True
            updated_values.append(canon)
        if changed:
            key_values_map[token] = updated_values

    for token, values in key_values_map.items():
        LITERALS[token] = ", ".join(values)

    alias_link_map: dict[str, str] = {}
    recipe_key_values = key_values_map.get("row_recipe_code")
    if recipe_key_values:
        alias_link_map = {
            "recipe_code": "row_recipe_code",
            "filter_recipe_code": "row_recipe_code",
        }
        literal_value = ", ".join(recipe_key_values)
        for alias in alias_link_map.keys():
            LITERALS[alias] = literal_value

    multi_key_selected = any(len(values) > 1 for values in key_values_map.values())

    def _first_alias_value(token: str) -> str | None:
        source = alias_link_map.get(token)
        if not source:
            return None
        return _first_key_value(key_values_map.get(source, []))

    def _apply_alias_params(target: dict[str, Any]) -> None:
        for alias in alias_link_map:
            if alias in target and str(target[alias] or "").strip():
                continue
            alias_value = _first_alias_value(alias)
            if alias_value is not None:
                target[alias] = alias_value

    _log_debug("Normalized key_values_map", key_values_map, "multi_key_selected", multi_key_selected)

    def _first_key_value(values: list[str]) -> str | None:
        for val in values:
            text = str(val or "").strip()
            if text:
                return text
        return None

    def _iter_key_combinations(values_map: dict[str, list[str]]) -> Iterable[dict[str, str]]:
        if not values_map:
            yield {}
            return
        tokens: list[str] = []
        value_lists: list[list[str]] = []
        for token, raw_values in values_map.items():
            unique: list[str] = []
            seen_local: set[str] = set()
            for val in raw_values:
                text = str(val or "").strip()
                if not text or text in seen_local:
                    continue
                seen_local.add(text)
                unique.append(text)
            if unique:
                tokens.append(token)
                value_lists.append(unique)
        if not tokens:
            yield {}
            return
        max_combos_raw = os.getenv("NEURA_REPORT_MAX_KEY_COMBINATIONS", "500")
        try:
            max_combos = int(max_combos_raw)
        except ValueError:
            max_combos = 500
        max_combos = max(1, max_combos)
        estimated = 1
        for values in value_lists:
            estimated *= max(1, len(values))
            if estimated > max_combos:
                raise ValueError(
                    f"Too many key combinations ({estimated} > {max_combos}). "
                    "Narrow key selections or reduce multi-select values."
                )
        for combo in product(*value_lists):
            yield {token: value for token, value in zip(tokens, combo)}

    _PLAYWRIGHT_ROW_FRIENDLY_LIMIT = 6000

    async def html_to_pdf_async(html_path: Path, pdf_path: Path, base_dir: Path, pdf_scale: float | None = None):
        if async_playwright is None:
            logger.warning("Playwright not available; skipping PDF generation.")
            return

        html_path_resolved = html_path.resolve()
        html_source = html_path_resolved.read_text(encoding="utf-8", errors="ignore")
        approx_row_count = html_source.lower().count("<tr")
        base_dir_resolved = (base_dir or html_path.parent).resolve()
        pdf_path_resolved = pdf_path.resolve()
        base_url = base_dir_resolved.as_uri()

        # Ensure Playwright uses a writable temp dir (avoids /tmp quota issues)
        if not os.environ.get("TMPDIR"):
            _fallback_tmp = Path.home() / ".tmp"
            if _fallback_tmp.is_dir():
                os.environ["TMPDIR"] = str(_fallback_tmp)

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            context = None
            try:
                context = await browser.new_context(base_url=base_url)
                page = await context.new_page()
                _pdf_timeout_ms = int(os.environ.get("NEURA_PDF_RENDER_TIMEOUT_MS", "600000"))
                page.set_default_timeout(_pdf_timeout_ms)
                await page.set_content(html_source, wait_until="load", timeout=_pdf_timeout_ms)
                await page.emulate_media(media="print")
                scale_value = pdf_scale or 1.0
                if not isinstance(scale_value, (int, float)):
                    scale_value = 1.0
                scale_value = max(0.1, min(float(scale_value), 2.0))
                try:
                    await page.pdf(
                        path=str(pdf_path_resolved),
                        format="A4",
                        print_background=True,
                        margin={"top": "10mm", "right": "10mm", "bottom": "10mm", "left": "10mm"},
                        prefer_css_page_size=True,
                        scale=scale_value,
                    )
                except Exception as exc:
                    if approx_row_count >= _PLAYWRIGHT_ROW_FRIENDLY_LIMIT:
                        raise RuntimeError(
                            (
                                "PDF rendering failed because the report contains "
                                f"approximately {approx_row_count:,} table rows, which exceeds the printable limit. "
                                "Please filter the data further or split the report into smaller chunks and try again."
                            )
                        ) from exc
                    raise
            finally:
                if context is not None:
                    await context.close()
                await browser.close()

    def _combine_html_documents(html_sections: list[str]) -> str:
        if not html_sections:
            return ""
        combined_body: list[str] = []
        doc_type = ""
        head_html = ""

        head_pattern = re.compile(r"(?is)<head\b[^>]*>(?P<head>.*)</head>")
        body_pattern = re.compile(r"(?is)<body\b[^>]*>(?P<body>.*)</body>")
        doctype_pattern = re.compile(r"(?is)^\s*<!DOCTYPE[^>]*>", re.MULTILINE)

        for idx, raw_html in enumerate(html_sections):
            text = raw_html or ""
            if idx == 0:
                doctype_match = doctype_pattern.search(text)
                if doctype_match:
                    doc_type = doctype_match.group(0).strip()
                    text = text[doctype_match.end() :]
                head_match = head_pattern.search(text)
                if head_match:
                    head_html = head_match.group(0).strip()
                body_match = body_pattern.search(text)
                if body_match:
                    section_body = body_match.group("body").strip()
                else:
                    section_body = text.strip()
                combined_body.append(f'<div class="nr-key-section" data-nr-section="1">\n{section_body}\n</div>')
            else:
                body_match = body_pattern.search(text)
                section = body_match.group("body").strip() if body_match else text.strip()
                combined_body.append(
                    f'<div class="nr-key-section" data-nr-section="{idx + 1}" style="page-break-before: always;">\n{section}\n</div>'
                )

        doc_lines = []
        if doc_type:
            doc_lines.append(doc_type)
        doc_lines.append("<html>")
        if head_html:
            doc_lines.append(head_html)
        doc_lines.append("<body>")
        doc_lines.append("\n\n".join(combined_body))
        doc_lines.append("</body>")
        doc_lines.append("</html>")
        return "\n".join(doc_lines)

    def _value_has_content(value: Any) -> bool:
        if value is None:
            return False
        if isinstance(value, (int, float, Decimal)):
            return value != 0
        text = str(value).strip()
        if not text:
            return False
        try:
            num = Decimal(text)
        except Exception:
            return True
        else:
            return num != 0

    def _row_has_significant_data(row: Mapping[str, Any], columns: list[str]) -> bool:
        return _row_has_any_data(row, (), columns)

    def _token_values_have_data(row: Mapping[str, Any], tokens: list[str]) -> bool:
        return _row_has_any_data(row, tokens, ())

    def _row_has_any_data(row: Mapping[str, Any], tokens: Sequence[str], columns: Sequence[str]) -> bool:
        for token in tokens:
            if not token:
                continue
            if _value_has_content(_value_for_token(row, token)):
                return True
        for col in columns:
            if not col:
                continue
            if _value_has_content(row.get(col)):
                return True
        for key, value in row.items():
            if not isinstance(key, str):
                continue
            if _is_counter_field(key):
                continue
            if _value_has_content(value):
                return True
        return False

    def _is_counter_field(name: str | None) -> bool:
        if not name:
            return False
        if not isinstance(name, str):
            name = str(name)
        normalized = re.sub(r"[^a-z0-9]", "", name.lower())
        if not normalized:
            return False
        if normalized in {
            "row",
            "rowid",
            "rowno",
            "rownum",
            "rownumber",
            "rowindex",
            "rowcounter",
            "srno",
            "sno",
        }:
            return True
        counter_markers = ("serial", "sequence", "seq", "counter")
        if any(marker in normalized for marker in counter_markers):
            return True
        # Exclude data fields that happen to end with counter-like suffixes
        # (e.g. row_bin_no is a bin identifier, row_recipe_no is a recipe ref)
        data_markers = ("bin", "recipe", "batch", "machine")
        if any(marker in normalized for marker in data_markers):
            return False
        counter_suffixes = (
            "slno",
            "srno",
            "sno",
            "snum",
            "snumber",
            "sl",
            "no",
            "num",
            "number",
            "idx",
            "index",
        )
        return any(normalized.endswith(suffix) and normalized.startswith("row") for suffix in counter_suffixes)

    def _reindex_serial_fields(rows: list[dict], tokens: Sequence[str], columns: Sequence[str]) -> None:
        serial_tokens = [tok for tok in tokens if _is_counter_field(tok)]
        serial_columns = [col for col in columns if _is_counter_field(col)]
        if not serial_tokens and not serial_columns:
            return

        # Skip fields whose existing values are non-numeric strings
        # (e.g. MELT-produced literals like "Scale-2", "Scale-3").
        def _has_non_numeric(field: str) -> bool:
            for row in rows:
                val = row.get(field)
                if val is None or isinstance(val, (int, float)):
                    continue
                try:
                    float(str(val))
                except (ValueError, TypeError):
                    return True
            return False

        serial_tokens = [t for t in serial_tokens if not _has_non_numeric(t)]
        serial_columns = [c for c in serial_columns if not _has_non_numeric(c)]
        if not serial_tokens and not serial_columns:
            return

        for idx, row in enumerate(rows, start=1):
            for tok in serial_tokens:
                row[tok] = idx
            for col in serial_columns:
                row[col] = idx

    def _fill_batch_level_tokens(
        block_html: str,
        batch_header: dict[str, Any],
        known_tokens: set[str],
    ) -> str:
        """Fill batch-level tokens from carry-forward data (BLOCK_REPEAT).

        Only replaces tokens OUTSIDE <tbody>...</tbody> to avoid clobbering
        row-level tokens that need per-row values.
        """
        batch_num = batch_header.get("__batch_number__", "")

        # Split block into: before-tbody, tbody, after-tbody
        tbody_pat = re.compile(r"(<tbody\b[^>]*>)(.*?)(</tbody>)", re.DOTALL | re.IGNORECASE)
        tbody_m = tbody_pat.search(block_html)

        if tbody_m:
            before = block_html[:tbody_m.start()]
            tbody_content = block_html[tbody_m.start():tbody_m.end()]
            after = block_html[tbody_m.end():]
        else:
            before = block_html
            tbody_content = ""
            after = ""

        # Fill tokens only in the before and after sections
        for section_html in [before, after]:
            for m in re.finditer(r"\{(\w+)\}", section_html):
                token = m.group(1)
                value = _match_batch_cf(token, batch_header, batch_num)
                if value is not None:
                    before = sub_token(before, token, format_token_value(token, value))
                    after = sub_token(after, token, format_token_value(token, value))

        return before + tbody_content + after

    def _match_batch_cf(
        token: str,
        cf: dict[str, Any],
        batch_num: Any,
    ) -> Any:
        """Match a batch-level token to carry-forward column data."""
        # 1. Direct match
        if token in cf:
            return cf[token]

        # 2. Sequential numbering tokens
        tok_low = token.lower()
        if tok_low in ("batch_no", "batch_number", "bth_no"):
            return batch_num

        # 3. Numbered suffix: start_time_1 → start_time, start_time_2 → end_time
        if tok_low.endswith("_1"):
            base = token[:-2]
            if base in cf:
                return cf[base]
        if tok_low.endswith("_2"):
            base = token[:-2]
            # Common pair: start_time_2 → end_time
            end_key = base.replace("start", "end")
            if end_key in cf:
                return cf[end_key]
            if base in cf:
                return cf[base]

        # 3b. Timestamp column fallback: start_time → timestamp_utc
        if tok_low in ("start_time",) and "timestamp_utc" in cf:
            ts = str(cf["timestamp_utc"])
            # Extract time portion from ISO timestamp (e.g. "2026-02-26T13:41:26+05:30" → "13:41:26+05:30")
            if "T" in ts:
                return ts.split("T", 1)[1]
            if " " in ts:
                return ts.split(" ", 1)[1]
            return ts

        # 3c. End time fallback: end_time → end_timestamp_utc (or timestamp_utc)
        if tok_low in ("end_time",):
            for et_col in ("end_timestamp_utc", "end_time"):
                if et_col in cf:
                    ts = str(cf[et_col])
                    if "T" in ts:
                        return ts.split("T", 1)[1]
                    if " " in ts:
                        return ts.split(" ", 1)[1]
                    return ts

        # 4. Date derivation: batch_date → date portion of start_time/end_time/timestamp_utc
        if "date" in tok_low:
            for dt_col in ("start_time", "end_time", "timestamp_utc"):
                if dt_col in cf:
                    dt_str = str(cf[dt_col])
                    if "T" in dt_str:
                        return dt_str.split("T")[0]
                    if " " in dt_str:
                        return dt_str.split(" ")[0]
                    return dt_str

        # 5. Duration computation from start_time and end_time (or timestamp_utc / end_timestamp_utc)
        if tok_low in ("duration_sec", "dur_sec", "duration"):
            from datetime import datetime
            def _try_parse_iso(s):
                if not s:
                    return None
                s = str(s)
                try:
                    return datetime.fromisoformat(s.replace("Z", "+00:00"))
                except (ValueError, TypeError):
                    pass
                for pat in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S"):
                    try:
                        return datetime.strptime(s, pat)
                    except ValueError:
                        continue
                return None

            st_raw = cf.get("start_time") or cf.get("timestamp_utc")
            et_raw = cf.get("end_time") or cf.get("end_timestamp_utc")
            st_dt = _try_parse_iso(st_raw)
            et_dt = _try_parse_iso(et_raw)
            if st_dt and et_dt:
                return int(abs((et_dt - st_dt).total_seconds()))

        # 6. Recipe aliases
        if tok_low in ("recipe_code", "recipe_no"):
            return cf.get("recipe_name", cf.get("id", ""))

        # 6b. batch_ prefix → strip prefix and try common aliases
        if tok_low.startswith("batch_"):
            suffix = tok_low[6:]  # "batch_recipe" → "recipe"
            _BATCH_ALIASES = {
                "recipe": ["recipe_name", "recipe_code", "recipe"],
                "no": ["id", "batch_id", "number", "batch_number"],
                "start": ["start_time", "batch_start", "timestamp_utc"],
                "end": ["end_time", "batch_end", "end_timestamp_utc"],
                "date": ["start_time", "batch_date"],
                "label": ["recipe_name", "batch_label"],
                "name": ["recipe_name"],
                "id": ["id", "batch_id"],
            }
            for alias in _BATCH_ALIASES.get(suffix, []):
                if alias in cf:
                    return cf[alias]
            # Also try direct suffix match against carry-forward keys
            for key, val in cf.items():
                if key.startswith("__"):
                    continue
                if suffix == key or (len(suffix) > 2 and suffix in key):
                    return val

        # 7. Fuzzy: token is a suffix of a cf column or vice-versa
        for key, val in cf.items():
            if key.startswith("__"):
                continue
            if tok_low.endswith(key.lower()) or key.lower().endswith(tok_low):
                return val

        # 7. Strip row_ prefix and try mapping to contract columns
        if tok_low.startswith("row_"):
            stripped = tok_low[4:]  # e.g. "row_recipe_code" -> "recipe_code"
            # Direct match on stripped name
            for key, val in cf.items():
                if key.startswith("__"):
                    continue
                if key.lower() == stripped:
                    return val
            # Common aliases: recipe_code -> recipe_name, batch_no -> id
            _aliases = {
                "recipe_code": ["recipe_name", "recipe"],
                "batch_no": ["id", "batch_id", "batch_number"],
                "start_time": ["start_time", "start_timestamp"],
                "end_time": ["end_time", "end_timestamp"],
            }
            for alias in _aliases.get(stripped, []):
                if alias in cf:
                    return cf[alias]

        return None

    _warned_tokens: set[str] = set()  # log each unresolved token only once per generation

    def _value_for_token(row: Mapping[str, Any], token: str) -> Any:
        def _sanitize(v: Any) -> Any:
            """Coerce NaN/NaT to None so downstream formatters get a clean value."""
            if v is None:
                return None
            try:
                import math
                if isinstance(v, float) and math.isnan(v):
                    return None
            except (TypeError, ValueError):
                pass
            return v

        if not token:
            return None
        if token in row:
            return _sanitize(row[token])
        normalized = str(token).lower()
        for key in row.keys():
            if isinstance(key, str) and key.lower() == normalized:
                return _sanitize(row[key])
        mapped = PLACEHOLDER_TO_COL.get(token)
        if mapped:
            col = _extract_col_name(mapped)
            if col:
                if col in row:
                    return _sanitize(row[col])
                for key in row.keys():
                    if isinstance(key, str) and key.lower() == col.lower():
                        return _sanitize(row[key])
        if token not in _warned_tokens:
            _warned_tokens.add(token)
            logger.warning(
                "token_unresolved token=%s available_keys=%s",
                token,
                list(row.keys())[:10],
                extra={"event": "token_unresolved", "token": token},
            )
        return None

    def _prune_placeholder_rows(rows: Sequence[Mapping[str, Any]], tokens: Sequence[str]) -> list[dict[str, Any]]:
        material_tokens = [tok for tok in tokens if tok and "material" in tok.lower()]
        pruned: list[dict[str, Any]] = []
        for row in rows:
            keep = True
            for tok in material_tokens:
                if not _value_has_content(_value_for_token(row, tok)):
                    keep = False
                    break
            if keep:
                pruned.append(dict(row))
        return pruned if pruned else [dict(row) for row in rows]

    def _filter_rows_for_render(
        rows: Sequence[Mapping[str, Any]],
        row_tokens_template: Sequence[str],
        row_columns: Sequence[str],
        *,
        treat_all_as_data: bool,
    ) -> list[dict[str, Any]]:
        if not rows:
            return []

        if treat_all_as_data:
            prepared = [dict(row) for row in rows]
        else:
            significant_tokens = [tok for tok in row_tokens_template if tok and not _is_counter_field(tok)]
            significant_columns = [col for col in row_columns if col and not _is_counter_field(col)]
            guard_rows = bool(significant_tokens or significant_columns)
            prepared: list[dict[str, Any]] = []
            for row in rows:
                if guard_rows and not _row_has_any_data(row, significant_tokens, significant_columns):
                    continue
                prepared.append(dict(row))

        if prepared:
            _reindex_serial_fields(prepared, row_tokens_template, row_columns)
        return prepared

    # Skip fanout when contract says to aggregate across batches
    _has_group_aggregate = bool((OBJ.get("group_aggregate") or {}).get("strategy"))
    if multi_key_selected and not __force_single and not _has_group_aggregate:
        html_sections: list[str] = []
        tmp_outputs: list[tuple[Path, Path]] = []
        try:
            for idx, combo in enumerate(_iter_key_combinations(key_values_map), start=1):
                selection: dict[str, str] = {token: value for token, value in combo.items()}
                for alias, source in alias_link_map.items():
                    if alias not in selection and source in selection:
                        selection[alias] = selection[source]
                _log_debug("Fanout iteration", idx, "selection", selection)
                tmp_html = OUT_HTML.with_name(f"{OUT_HTML.stem}__key{idx}.html")
                tmp_pdf = OUT_PDF.with_name(f"{OUT_PDF.stem}__key{idx}.pdf")
                result = fill_and_print(
                    OBJ=OBJ,
                    TEMPLATE_PATH=TEMPLATE_PATH,
                    DB_PATH=DB_PATH,
                    OUT_HTML=tmp_html,
                    OUT_PDF=tmp_pdf,
                    START_DATE=START_DATE,
                    END_DATE=END_DATE,
                    batch_ids=None,
                    KEY_VALUES=selection or None,
                    __force_single=True,
                )
                html_sections.append(Path(result["html_path"]).read_text(encoding="utf-8", errors="ignore"))
                tmp_outputs.append((Path(result["html_path"]), Path(result["pdf_path"])))

            if not html_sections:
                return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": False}

            combined_html = _combine_html_documents(html_sections)
            OUT_HTML.write_text(combined_html, encoding="utf-8")
            _html_to_pdf_subprocess(OUT_HTML, OUT_PDF, TEMPLATE_PATH.parent)
            return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": True}
        finally:
            for tmp_html_path, tmp_pdf_path in tmp_outputs:
                for path_sel in (tmp_html_path, tmp_pdf_path):
                    with contextlib.suppress(FileNotFoundError):
                        path_sel.unlink()

    def _get_literal_raw(token: str) -> str:
        if token not in LITERALS:
            return ""
        raw = LITERALS[token]
        return "" if raw is None else str(raw)

    def _literal_has_content(token: str) -> bool:
        return bool(_get_literal_raw(token).strip())

    def _first_nonempty_literal(tokens: Iterable[str]) -> tuple[str | None, str | None]:
        for tok in tokens:
            raw = _get_literal_raw(tok)
            if raw.strip():
                return tok, raw
        return None, None

    def _record_special_value(target: dict[str, str], token: str, value: str) -> None:
        existing_raw = _get_literal_raw(token)
        if existing_raw.strip():
            target[token] = existing_raw
        else:
            target[token] = value
            if token in LITERALS:
                LITERALS[token] = value

    def _filter_tokens_without_literal(tokens: set[str]) -> set[str]:
        return {tok for tok in tokens if not _literal_has_content(tok)}

    BEGIN_TAG = "<!-- BEGIN:BATCH (auto) -->"
    END_TAG = "<!-- END:BATCH (auto) -->"
    try:
        prototype_block, start0, end_last = _select_prototype_block(html, ROW_TOKENS)
    except Exception as exc:
        _raise_no_block(html, exc)
    shell_prefix = html[:start0] + BEGIN_TAG
    shell_suffix = END_TAG + html[end_last:]

    parent_table = JOIN.get("parent_table", "")
    parent_key = JOIN.get("parent_key", "")
    child_table = JOIN.get("child_table", "")
    child_key = JOIN.get("child_key", "")

    # --- Additive: auto-detect date columns if missing from contract ---
    for _tbl in (parent_table, child_table):
        if _tbl and _tbl not in DATE_COLUMNS:
            _auto = detect_date_column(DB_PATH, _tbl)
            if _auto:
                DATE_COLUMNS[_tbl] = _auto
                logger.info("date_column_auto_detected table=%s col=%s", _tbl, _auto)

    parent_date = DATE_COLUMNS.get(parent_table, "")
    child_date = DATE_COLUMNS.get(child_table, "")
    order_col = ROW_ORDER[0] if ROW_ORDER else "ROWID"
    if isinstance(order_col, str) and order_col.upper() != "ROWID":
        mapped_order = PLACEHOLDER_TO_COL.get(order_col, order_col)
        if isinstance(mapped_order, str):
            mapped_order = mapped_order.strip()
            if "." in mapped_order:
                mapped_order = mapped_order.split(".", 1)[1].strip()
            if mapped_order:
                order_col = mapped_order

    def _normalize_token_name(name: str) -> str:
        return re.sub(r"[^a-z0-9]", "", name.lower())

    token_index: dict[str, set[str]] = defaultdict(set)
    all_candidate_tokens = (
        set(TEMPLATE_TOKENS) | set(HEADER_TOKENS) | set(ROW_TOKENS) | set(TOTALS.keys()) | set(LITERALS.keys())
    )

    def _token_synonym_keys(norm: str) -> set[str]:
        """
        Generate lightweight normalization aliases so that abbreviated tokens like
        `pg_total` or `page_num` still map onto the same lookup keys as their
        longer forms without needing every variant enumerated manually.
        """
        if not norm:
            return set()
        aliases = {norm}
        replacements: tuple[tuple[str, str], ...] = (
            ("pg", "page"),
            ("num", "number"),
            ("no", "number"),
            ("cnt", "count"),
            ("ttl", "total"),
        )
        for src, dest in replacements:
            if src in norm and dest not in norm:
                aliases.add(norm.replace(src, dest))
        # Avoid generating implausible short aliases (e.g., converting a lone "no"
        # in tokens unrelated to pagination), but include a fallback where a token
        # is exactly "pg" so that later lookups on "page" resolve.
        if norm == "pg":
            aliases.add("page")
        return {alias for alias in aliases if alias}

    for tok in all_candidate_tokens:
        norm = _normalize_token_name(tok)
        for key in _token_synonym_keys(norm):
            token_index[key].add(tok)

    def _tokens_for_keys(keys: set[str]) -> set[str]:
        found: set[str] = set()
        for key in keys:
            found.update(token_index.get(key, set()))
        return found

    def _format_for_db(dt_obj: datetime | None, raw_value, include_time_default: bool) -> str:
        """
        Normalize input dates for SQLite bindings:
          - prefer ISO 8601 date or datetime strings
          - fall back to trimmed raw strings when parsing fails
        """
        if dt_obj:
            include_time = include_time_default or bool(
                dt_obj.hour or dt_obj.minute or dt_obj.second or dt_obj.microsecond
            )
            if include_time:
                return dt_obj.strftime("%Y-%m-%d %H:%M:%S")
            return dt_obj.strftime("%Y-%m-%d")
        return "" if raw_value is None else str(raw_value).strip()

    start_dt = _parse_date_like(START_DATE)
    end_dt = _parse_date_like(END_DATE)
    _IST = timezone(timedelta(hours=5, minutes=30))
    print_dt = datetime.now(_IST)

    start_has_time = _has_time_component(START_DATE, start_dt)
    end_has_time = _has_time_component(END_DATE, end_dt)

    START_DATE_KEYS = {"fromdate", "datefrom", "startdate", "periodstart", "rangefrom", "fromdt", "startdt", "fromdatetime", "startdatetime", "datetimefrom"}
    END_DATE_KEYS = {"todate", "dateto", "enddate", "periodend", "rangeto", "todt", "enddt", "todatetime", "enddatetime", "datetimeto"}
    PRINT_DATE_KEYS = {
        "printdate",
        "printedon",
        "printeddate",
        "generatedon",
        "generateddate",
        "rundate",
        "runon",
        "generatedat",
        "reportdate",
    }
    PRINT_TIME_KEYS = {
        "printtime",
        "printedat",
        "generatedtime",
        "runtime",
    }
    PAGE_NO_KEYS = {
        "page",
        "pageno",
        "pagenum",
        "pagenumber",
        "pageindex",
        "pageidx",
        "pagecurrent",
        "currentpage",
        "currpage",
        "pgno",
        "pgnum",
        "pgnumber",
        "pgindex",
        "pgcurrent",
    }
    PAGE_COUNT_KEYS = {
        "pagecount",
        "pagecounts",
        "totalpages",
        "pagestotal",
        "pages",
        "pagetotal",
        "totalpage",
        "pagecounttotal",
        "totalpagecount",
        "pagescount",
        "countpages",
        "lastpage",
        "finalpage",
        "maxpage",
        "pgtotal",
        "totalpg",
        "pgcount",
        "countpg",
        "pgs",
        "pgscount",
        "pgstotal",
        "totalpgs",
    }
    PAGE_LABEL_KEYS = {
        "pagelabel",
        "pageinfo",
        "pagesummary",
        "pagefooter",
        "pagefootertext",
        "pageindicator",
        "pagecaption",
        "pagefooterlabel",
        "pagetext",
        "pagefooterinfo",
        "pagehint",
    }

    special_values: dict[str, str] = {}

    start_tokens = _tokens_for_keys(START_DATE_KEYS)
    end_tokens = _tokens_for_keys(END_DATE_KEYS)
    print_tokens = _tokens_for_keys(PRINT_DATE_KEYS)
    print_time_tokens = _tokens_for_keys(PRINT_TIME_KEYS)
    page_number_tokens = _tokens_for_keys(PAGE_NO_KEYS)
    page_count_tokens = _tokens_for_keys(PAGE_COUNT_KEYS)
    page_label_tokens = _tokens_for_keys(PAGE_LABEL_KEYS)

    for tok in start_tokens:
        _record_special_value(
            special_values,
            tok,
            _format_for_token(tok, start_dt, include_time_default=start_has_time),
        )
    for tok in end_tokens:
        _record_special_value(
            special_values,
            tok,
            _format_for_token(tok, end_dt, include_time_default=end_has_time),
        )

    _, print_literal_value = _first_nonempty_literal(print_tokens)
    parsed_print_dt = _parse_date_like(print_literal_value) if print_literal_value else None
    print_dt_source = parsed_print_dt or print_dt
    print_has_time = _has_time_component(print_literal_value, parsed_print_dt)

    for tok in print_tokens:
        if print_literal_value and not parsed_print_dt:
            value = print_literal_value
        else:
            value = _format_for_token(tok, print_dt_source, include_time_default=print_has_time)
        _record_special_value(special_values, tok, value)

    for tok in print_time_tokens:
        _record_special_value(
            special_values, tok,
            print_dt_source.strftime("%I:%M %p") if print_dt_source else "",
        )

    page_number_tokens = _filter_tokens_without_literal(page_number_tokens)
    page_count_tokens = _filter_tokens_without_literal(page_count_tokens)
    page_label_tokens = _filter_tokens_without_literal(page_label_tokens)

    post_literal_specials = {tok: val for tok, val in special_values.items() if tok not in LITERALS}

    _ident_re = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")

    def qident(name: str) -> str:
        if _ident_re.match(name):
            return name
        safe = name.replace('"', '""')
        return f'"{safe}"'

    # ---- Composite-key helpers ----
    def _parse_key_cols(key_spec: str) -> list[str]:
        return [c.strip() for c in str(key_spec).split(",") if c and c.strip()]

    def _key_expr(cols: list[str]) -> str:
        parts = [f"COALESCE(CAST({qident(c)} AS TEXT),'')" for c in cols]
        if not parts:
            return "''"
        expr = parts[0]
        for p in parts[1:]:
            expr = f"{expr} || '|' || {p}"
        return expr

    def _split_bid(bid: str, n: int) -> list[str]:
        parts = str(bid).split("|")
        if len(parts) != n:
            raise ValueError(f"Composite key mismatch: expected {n} parts, got {len(parts)} in {bid!r}")
        return parts

    def _looks_like_composite_id(x: str, n: int) -> bool:
        return isinstance(x, str) and x.count("|") == (n - 1)

    pcols = _parse_key_cols(parent_key)
    ccols = _parse_key_cols(child_key)

    has_child = bool(child_table and ccols)
    parent_table_lc = parent_table.lower()
    child_table_lc = child_table.lower()
    parent_filter_map: dict[str, list[str]] = {}
    child_filter_map: dict[str, list[str]] = {}
    if key_values_map:
        for token, values in key_values_map.items():
            mapping_value = PLACEHOLDER_TO_COL.get(token)
            if not isinstance(mapping_value, str):
                continue
            target = mapping_value.strip()
            if not target or target.upper().startswith("PARAM:") or "." not in target:
                continue
            table_name, column_name = target.split(".", 1)
            table_name = table_name.strip(' "`[]')
            column_name = column_name.strip(' "`[]')
            if not column_name:
                continue
            table_key = table_name.lower()
            if table_key in (parent_table_lc, "header"):
                bucket = list(parent_filter_map.get(column_name, []))
                for val in values:
                    if val not in bucket:
                        bucket.append(val)
                if bucket:
                    parent_filter_map[column_name] = bucket
            if has_child and table_key in (child_table_lc, "rows"):
                bucket = list(child_filter_map.get(column_name, []))
                for val in values:
                    if val not in bucket:
                        bucket.append(val)
                if bucket:
                    child_filter_map[column_name] = bucket
    parent_filter_items = list(parent_filter_map.items())
    child_filter_items = list(child_filter_map.items())
    parent_filter_sqls: list[str] = []
    parent_filter_values: list[str] = []
    for col, values in parent_filter_items:
        normalized: list[str] = []
        for val in values:
            if not isinstance(val, str):
                continue
            text = val.strip()
            if text and text not in normalized:
                normalized.append(text)
        if not normalized:
            continue
        if len(normalized) == 1:
            parent_filter_sqls.append(f"{qident(col)} = ?")
        else:
            placeholders = ", ".join("?" for _ in normalized)
            parent_filter_sqls.append(f"{qident(col)} IN ({placeholders})")
        parent_filter_values.extend(normalized)
    parent_filter_values_tuple = tuple(parent_filter_values)

    child_filter_sqls: list[str] = []
    child_filter_values: list[str] = []
    for col, values in child_filter_items:
        normalized: list[str] = []
        for val in values:
            if not isinstance(val, str):
                continue
            text = val.strip()
            if text and text not in normalized:
                normalized.append(text)
        if not normalized:
            continue
        if len(normalized) == 1:
            child_filter_sqls.append(f"{qident(col)} = ?")
        else:
            placeholders = ", ".join("?" for _ in normalized)
            child_filter_sqls.append(f"{qident(col)} IN ({placeholders})")
        child_filter_values.extend(normalized)
    child_filter_values_tuple = tuple(child_filter_values)

    def _merge_predicate(base_sql: str, extras: list[str]) -> str:
        if not extras:
            return base_sql
        extras_joined = " AND ".join(extras)
        base_sql = (base_sql or "1=1").strip()
        return f"({base_sql}) AND {extras_joined}"

    # --- Date predicates and adapters (handle missing/invalid date columns)
    parent_type = get_col_type(DB_PATH, parent_table, parent_date)
    child_type = get_col_type(DB_PATH, child_table, child_date)
    parent_pred, adapt_parent = mk_between_pred_for_date(parent_date, parent_type)
    child_pred, adapt_child = mk_between_pred_for_date(child_date, child_type)
    parent_where_clause = _merge_predicate(parent_pred, parent_filter_sqls)
    child_where_clause = _merge_predicate(child_pred, child_filter_sqls) if has_child else child_pred
    db_start = _format_for_db(start_dt, START_DATE, start_has_time)
    db_end = _format_for_db(end_dt, END_DATE, end_has_time)
    PDATE = tuple(adapt_parent(db_start, db_end))  # () if 1=1
    CDATE = tuple(adapt_child(db_start, db_end)) if has_child else tuple()  # () if 1=1
    parent_params_all = tuple(PDATE) + parent_filter_values_tuple
    child_params_all = tuple(CDATE) + child_filter_values_tuple if has_child else tuple()

    sql_params: dict[str, object] = {
        "from_date": db_start,
        "to_date": db_end,
        "start_date": db_start,
        "end_date": db_end,
    }

    for token in contract_adapter.param_tokens:
        if token in ("from_date", "to_date", "start_date", "end_date"):
            continue
        if token in key_values_map:
            # Use comma-joined LITERALS for header display (multi-value support)
            if token in LITERALS and LITERALS[token]:
                sql_params[token] = LITERALS[token]
            else:
                first_value = _first_key_value(key_values_map[token])
                if first_value is not None:
                    sql_params[token] = first_value
        elif alias_link_map.get(token):
            alias_value = _first_alias_value(token)
            if alias_value is not None:
                sql_params[token] = alias_value
        elif token in LITERALS:
            sql_params[token] = LITERALS[token]
        elif token in special_values:
            sql_params[token] = special_values[token]
        else:
            sql_params.setdefault(token, "")

    _apply_alias_params(sql_params)

    def _apply_date_param_defaults(target: dict[str, object]) -> None:
        if not isinstance(target, dict):
            return

        def _inject(names: set[str], default_value: str) -> None:
            if not default_value:
                return
            for alias in names:
                if alias not in param_token_set and alias not in target:
                    continue
                current = target.get(alias)
                if isinstance(current, str):
                    if current.strip():
                        continue
                elif current not in (None, ""):
                    continue
                target[alias] = default_value

        _inject(_DATE_PARAM_START_ALIASES, db_start)
        _inject(_DATE_PARAM_END_ALIASES, db_end)

    _apply_date_param_defaults(sql_params)

    # Inject special values (report_date, generated_at, etc.) into sql_params
    # so that params.xxx mappings in contract_adapter can resolve them.
    for sv_key, sv_val in special_values.items():
        if sv_key not in sql_params:
            sql_params[sv_key] = sv_val

    # --- DataFrame pipeline (sole data path) ---

    df_value_filters: dict[str, list] = {}
    if key_values_map:
        for name, values in key_values_map.items():
            df_value_filters[name] = values

    pipeline = DataFramePipeline(
        contract_adapter=contract_adapter,
        loader=dataframe_loader,
        params=sql_params,
        start_date=sql_params.get("start_date") or sql_params.get("from_date"),
        end_date=sql_params.get("end_date") or sql_params.get("to_date"),
        value_filters=df_value_filters,
    )
    generator_results = pipeline.execute()

    BATCH_IDS = ["__DF_PIPELINE__"]

    _fp_progress(f"BATCH_IDS resolved: {len(BATCH_IDS or [])} batches")
    _log_debug("BATCH_IDS:", len(BATCH_IDS or []), (BATCH_IDS or [])[:20] if BATCH_IDS else [])
    # ---- Only touch tokens outside <style>/<script> ----
    def format_token_value(token: str, raw_value: Any) -> str:
        return contract_adapter.format_value(token, raw_value)

    def _fast_row_sub(template: str, tokens: list[str], col_lookup: dict[str, str],
                      rows: list[dict]) -> list[str]:
        """Single-pass token substitution for row templates (no <style>/<script>).

        Pre-compiles ONE regex for all tokens and replaces them in a single
        re.sub call per row — O(rows) instead of O(rows × tokens).
        """
        if not tokens:
            return [template] * len(rows)
        token_alts = "|".join(re.escape(t) for t in tokens)
        pat = re.compile(r"\{\{?\s*(" + token_alts + r")\s*\}\}?")
        results = []
        for r in rows:
            def _repl(m, _row=r):
                t = m.group(1)
                col = col_lookup.get(t)
                if not col:
                    return m.group(0)
                return format_token_value(t, _row.get(col))
            results.append(pat.sub(_repl, template))
        return results

    def _inject_page_counter_spans(
        html_in: str,
        page_tokens: set[str],
        count_tokens: set[str],
        label_tokens: set[str] | None = None,
    ) -> str:
        label_tokens = label_tokens or set()
        updated = html_in
        page_markup = '<span class="nr-page-number" data-nr-counter="page" aria-label="Current page number"></span>'
        count_markup = '<span class="nr-page-count" data-nr-counter="pages" aria-label="Total page count"></span>'

        for tok in page_tokens:
            updated = sub_token(updated, tok, page_markup)
        for tok in count_tokens:
            updated = sub_token(updated, tok, count_markup)
        for tok in label_tokens:
            if count_tokens:
                label_markup = (
                    f'<span class="nr-page-label" data-nr-counter-label="1">Page {page_markup} of {count_markup}</span>'
                )
            else:
                label_markup = f'<span class="nr-page-label" data-nr-counter-label="1">Page {page_markup}</span>'
            updated = sub_token(updated, tok, label_markup)

        if (page_tokens or count_tokens or label_tokens) and "nr-page-counter-style" not in updated:
            style_block = """
<style id="nr-page-counter-style">
  .nr-page-number,
  .nr-page-count { white-space: nowrap; font-variant-numeric: tabular-nums; }
  .nr-page-label { white-space: nowrap; }
  @media screen {
    .nr-page-number::after { content: attr(data-nr-screen); }
    .nr-page-count::after { content: attr(data-nr-total-pages); }
  }
  @media print {
    body { counter-reset: page; }
    .nr-page-number::after { content: counter(page); }
    .nr-page-count::after { content: counter(pages); }
    .nr-page-count[data-nr-total-pages]::after { content: attr(data-nr-total-pages); }
  }
</style>
"""
            if "</head>" in updated:
                updated = updated.replace("</head>", style_block + "</head>", 1)
            else:
                updated = style_block + updated

        if (page_tokens or count_tokens or label_tokens) and "nr-page-counter-script" not in updated:
            metrics = _extract_page_metrics(updated)
            metrics_json = json.dumps(metrics)
            script_template = """
<script id="nr-page-counter-script">
(function() {
  const METRICS = __NR_METRICS__;
  const PX_PER_MM = 96 / 25.4;
  const BREAK_VALUES = ['page', 'always', 'left', 'right'];
  const TRAILING_BREAK_SENTINEL = '__nr_trailing_break__';
  let lastPageNodes = [];
  let lastCountNodes = [];

  function isForcedBreak(value) {
    if (!value) return false;
    const normalized = String(value).toLowerCase().trim();
    if (!normalized) return false;
    return BREAK_VALUES.indexOf(normalized) !== -1;
  }

  function readBreakValue(style, which) {
    if (!style) return '';
    if (which === 'before') {
      return (
        style.getPropertyValue('break-before') ||
        style.getPropertyValue('page-break-before') ||
        style.breakBefore ||
        style.pageBreakBefore ||
        ''
      );
    }
    return (
      style.getPropertyValue('break-after') ||
      style.getPropertyValue('page-break-after') ||
      style.breakAfter ||
      style.pageBreakAfter ||
      ''
    );
  }

  function findNextElement(node) {
    if (!node) return null;
    let current = node;
    while (current) {
      if (current.nextElementSibling) return current.nextElementSibling;
      current = current.parentElement;
    }
    return null;
  }

  function resolveNodeOffset(node) {
    if (!node || typeof node.getBoundingClientRect !== 'function') return 0;
    const rect = node.getBoundingClientRect();
    const scrollY = typeof window !== 'undefined' ? window.scrollY || window.pageYOffset || 0 : 0;
    return Math.max(0, rect.top + scrollY);
  }

  function collectManualBreakAnchors(root) {
    if (!root || !root.ownerDocument) return [];
    const anchors = [];
    const seen = new Set();
    const showElement = typeof NodeFilter !== 'undefined' && NodeFilter.SHOW_ELEMENT ? NodeFilter.SHOW_ELEMENT : 1;
    const walker = root.ownerDocument.createTreeWalker(root, showElement);

    function pushAnchor(target) {
      if (!target) return;
      if (seen.has(target)) return;
      seen.add(target);
      anchors.push(target);
    }

    while (walker.nextNode()) {
      const element = walker.currentNode;
      const style = window.getComputedStyle ? window.getComputedStyle(element) : null;
      if (!style) continue;
      if (isForcedBreak(readBreakValue(style, 'before'))) {
        pushAnchor(element);
      }
      if (isForcedBreak(readBreakValue(style, 'after'))) {
        const next = findNextElement(element);
        pushAnchor(next || TRAILING_BREAK_SENTINEL);
      }
    }
    return anchors;
  }

  function buildPageStartOffsets(manualAnchors, usableHeightPx, contentHeight, totalPages) {
    const offsets = [0];
    const seenOffsets = new Set([0]);
    manualAnchors.forEach((anchor) => {
      let offset = null;
      if (anchor === TRAILING_BREAK_SENTINEL) {
        offset = contentHeight + usableHeightPx;
      } else if (anchor && typeof anchor.getBoundingClientRect === 'function') {
        offset = resolveNodeOffset(anchor);
      }
      if (offset == null || !Number.isFinite(offset)) {
        return;
      }
      const key = Math.round(offset * 1000) / 1000;
      if (seenOffsets.has(key)) return;
      seenOffsets.add(key);
      offsets.push(offset);
    });

    offsets.sort((a, b) => a - b);

    while (offsets.length < totalPages) {
      const last = offsets[offsets.length - 1];
      offsets.push(last + usableHeightPx);
    }

    return offsets;
  }

  function resolvePageIndexFromOffsets(offset, startOffsets) {
    if (!startOffsets || !startOffsets.length) return 0;
    let index = 0;
    for (let i = 0; i < startOffsets.length; i += 1) {
      if (offset >= startOffsets[i] - 0.5) {
        index = i;
      } else {
        break;
      }
    }
    return index;
  }

  function indexOfSection(node, sections) {
    if (!node || !sections || !sections.length) return -1;
    const target = typeof node.closest === 'function' ? node.closest('.nr-key-section') : null;
    if (!target) return -1;
    for (let i = 0; i < sections.length; i += 1) {
      if (sections[i] === target) {
        return i;
      }
    }
    return -1;
  }

  function assignScreenText(node, text, key) {
    if (!node) return;
    const stringText = text == null ? '' : String(text);
    if (node.getAttribute('aria-label') === null) {
      node.setAttribute('aria-label', key === 'count' ? 'Total page count' : 'Current page number');
    }
    node.setAttribute('data-nr-screen', stringText);
    if (key === 'count') {
      node.setAttribute('data-nr-total-pages', stringText);
    } else {
      node.setAttribute('data-nr-page-estimate', stringText);
    }
    node.setAttribute('data-nr-' + key + '-text', stringText);
    node.textContent = stringText;
  }

  function clearNodesForPrint() {
    const nodes = lastPageNodes.concat(lastCountNodes);
    nodes.forEach((node) => {
      if (!node) return;
      if (!node.hasAttribute('data-nr-print-cache')) {
        node.setAttribute('data-nr-print-cache', node.textContent || '');
      }
      node.textContent = '';
    });
  }

  function restoreNodesAfterPrint() {
    const nodes = lastPageNodes.concat(lastCountNodes);
    nodes.forEach((node) => {
      if (!node) return;
      const cached = node.getAttribute('data-nr-print-cache');
      if (cached != null) {
        const key = node.getAttribute('data-nr-counter') === 'pages' ? 'count' : 'page';
        const preferred = node.getAttribute('data-nr-' + key + '-text');
        node.textContent = preferred != null ? preferred : cached;
        node.removeAttribute('data-nr-print-cache');
      }
    });
  }

  function computeTotals() {
    try {
      const doc = document.documentElement;
      const body = document.body;
      if (!doc || !body) return;
      const usableHeightMm = Math.max(METRICS.page_height_mm - (METRICS.margin_top_mm + METRICS.margin_bottom_mm), 0.1);
      const usableHeightPx = usableHeightMm * PX_PER_MM;
      const contentHeight = Math.max(
        body.scrollHeight,
        body.offsetHeight,
        doc.scrollHeight,
        doc.offsetHeight
      );
      const contentPages = Math.max(1, Math.ceil(contentHeight / usableHeightPx));
      const manualAnchors = collectManualBreakAnchors(body);
      const manualPages = manualAnchors.length > 0 ? manualAnchors.length + 1 : 1;
      const totalPages = Math.max(contentPages, manualPages);
      const startOffsets = buildPageStartOffsets(manualAnchors, usableHeightPx, contentHeight, totalPages);
      const totalAsString = String(totalPages);
      doc.setAttribute('data-nr-total-pages', totalAsString);
      const countNodes = Array.from(document.querySelectorAll('[data-nr-counter="pages"]'));
      countNodes.forEach((node) => assignScreenText(node, totalAsString, 'count'));
      const pageNodes = Array.from(document.querySelectorAll('[data-nr-counter="page"]'));
      const sections = Array.from(document.querySelectorAll('.nr-key-section'));
      pageNodes.forEach((node) => {
        const sectionIndex = indexOfSection(node, sections);
        let pageIndex;
        if (sectionIndex >= 0) {
          pageIndex = sectionIndex;
        } else {
          const offset = resolveNodeOffset(node);
          pageIndex = resolvePageIndexFromOffsets(offset, startOffsets);
        }
        const pageNumber = Math.min(totalPages, Math.max(1, pageIndex + 1));
        assignScreenText(node, String(pageNumber), 'page');
      });
      lastPageNodes = pageNodes;
      lastCountNodes = countNodes;
    } catch (err) {
      if (typeof console !== 'undefined' && console.warn) {
        console.warn('nr-page-counter: unable to compute preview counters', err);
      }
    }
  }

  function scheduleCompute() {
    computeTotals();
    setTimeout(computeTotals, 180);
  }

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', () => {
      scheduleCompute();
    }, { once: true });
  } else {
    scheduleCompute();
  }

  if (typeof window !== 'undefined') {
    window.addEventListener('resize', computeTotals, { passive: true });
    window.addEventListener('beforeprint', clearNodesForPrint);
    window.addEventListener('afterprint', () => {
      restoreNodesAfterPrint();
      scheduleCompute();
    });
    if (typeof window.matchMedia === 'function') {
      const mediaQuery = window.matchMedia('print');
      if (typeof mediaQuery.addEventListener === 'function') {
        mediaQuery.addEventListener('change', (event) => {
          if (event.matches) {
            clearNodesForPrint();
          } else {
            restoreNodesAfterPrint();
            scheduleCompute();
          }
        });
      } else if (typeof mediaQuery.addListener === 'function') {
        mediaQuery.addListener((event) => {
          if (event.matches) {
            clearNodesForPrint();
          } else {
            restoreNodesAfterPrint();
            scheduleCompute();
          }
        });
      }
    }
  }
})();
</script>
"""
            script_block = script_template.replace("__NR_METRICS__", metrics_json)
            if "</body>" in updated:
                updated = updated.replace("</body>", script_block + "</body>", 1)
            else:
                updated = updated + script_block
        return updated

    # ---- Helpers to find tbody / row template (improved) ----
    def best_rows_tbody(inner_html: str, allowed_tokens: set):
        tbodys = list(re.finditer(r"(?is)<tbody\b[^>]*>(.*?)</tbody>", inner_html))
        best = (None, None, -1)  # (match, inner, hits)
        for m in tbodys:
            tin = m.group(1)
            hits = 0
            for trm in re.finditer(r"(?is)<tr\b[^>]*>.*?</tr>", tin):
                tr_html = trm.group(0)
                toks = re.findall(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", tr_html)
                flat = [a.strip() if a else b.strip() for (a, b) in toks]
                hits += sum(1 for t in flat if t in allowed_tokens)
            if hits > best[2]:
                best = (m, tin, hits)
        if best[0] is not None:
            return best[0], best[1]
        return (tbodys[0], tbodys[0].group(1)) if tbodys else (None, None)

    def find_row_template(tbody_inner: str, allowed_tokens: set):
        for m in re.finditer(r"(?is)<tr\b[^>]*>.*?</tr>", tbody_inner):
            tr_html = m.group(0)
            toks = re.findall(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", tr_html)
            flat = []
            for a, b in toks:
                if a:
                    flat.append(a.strip())
                if b:
                    flat.append(b.strip())
            flat = [t for t in flat if t in allowed_tokens]
            if flat:
                return tr_html, (m.start(0), m.end(0)), sorted(set(flat), key=len, reverse=True)
        return None, None, []

    def majority_table_for_tokens(tokens, mapping):
        from collections import Counter

        tbls = []
        for t in tokens:
            tc = mapping.get(t, "")
            if "." in tc:
                tbls.append(tc.split(".", 1)[0])
        return Counter(tbls).most_common(1)[0][0] if tbls else None

    # ---- Pre-compute minimal column sets ----
    _SQL_EXPR_CHARS = re.compile(r"[|+\-*/()']")

    def _extract_col_name(mapping_value: str | None) -> str | None:
        if not isinstance(mapping_value, str):
            return None
        target = mapping_value.strip()
        if "." not in target:
            return None
        # Skip SQL expressions (e.g. "table.col || ' ' || table.col2") — these
        # are handled by generator SQL, not by column prefetch.
        if _SQL_EXPR_CHARS.search(target):
            return None
        after_dot = target.split(".", 1)[1].strip()
        if not after_dot:
            return None
        col = re.split(r"[,)\s]", after_dot, 1)[0].strip()
        return col or None

    header_cols = sorted({col for t in HEADER_TOKENS for col in [_extract_col_name(PLACEHOLDER_TO_COL.get(t))] if col})
    row_cols = sorted({col for t in ROW_TOKENS for col in [_extract_col_name(PLACEHOLDER_TO_COL.get(t))] if col})

    totals_by_table = defaultdict(lambda: defaultdict(list))
    total_token_to_target = {}

    for token, raw_target in TOTALS.items():
        if isinstance(raw_target, dict):
            continue  # Declarative op spec — handled by DF pipeline, not SQL prefetch
        target = (raw_target or PLACEHOLDER_TO_COL.get(token, "")).strip()
        if not target or "." not in target:
            continue
        table_name, col_name = [part.strip() for part in target.split(".", 1)]
        if not table_name or not col_name:
            continue
        totals_by_table[table_name][col_name].append(token)
        total_token_to_target[token] = (table_name, col_name)

    def _coerce_total_value(raw):
        if raw is None:
            return None, "0"
        try:
            decimal_value = Decimal(str(raw).strip())
        except (InvalidOperation, ValueError, TypeError, AttributeError):
            return None, "0"
        if not decimal_value.is_finite():
            return None, "0"
        formatted = format_decimal_str(decimal_value, max_decimals=3) or "0"
        return float(decimal_value), formatted

    totals_accum = defaultdict(float)
    last_totals_per_token = {token: "0" for token in TOTALS}

    # ---- Render all batches ----
    rendered_blocks = []
    if generator_results is not None:
        # ── Detect batch-level tokens (present in block but unknown to contract) ──
        _all_block_tokens = set(re.findall(r"\{(\w+)\}", prototype_block))
        _known_tokens = set(ROW_TOKENS) | set(TOTALS.keys()) | set(HEADER_TOKENS)
        _batch_level_tokens = _all_block_tokens - _known_tokens

        _df_batches = generator_results.get("batches") or []
        _use_per_batch = bool(_df_batches)
        _top_rows = generator_results.get("rows") or []

        logger.info(
            "df_render_decision per_batch=%s batches=%d top_rows=%d batch_tokens=%s",
            _use_per_batch, len(_df_batches), len(_top_rows), _batch_level_tokens,
        )

        # ── Shared: pre-analyse the row template from the prototype once ──
        # Include both mapped tokens AND contract-declared row_tokens (MELT aliases, computed, etc.)
        allowed_row_tokens = (
            {t for t in PLACEHOLDER_TO_COL.keys() if t not in TOTALS}
            | set(ROW_TOKENS)
        ) - set(HEADER_TOKENS)
        header_rows = generator_results.get("header") or []
        header_row = header_rows[0] if header_rows else {}

        def _render_df_block(rows_data, totals_data, batch_header=None):
            """Render one block from DF pipeline data. Returns (html, had_rows)."""
            blk = prototype_block

            # (a) Fill global header tokens — but NOT tokens that exist in batch carry-forward
            # (batch-per-row tokens like batch_no, start_time should be filled by batch_header, not global header)
            batch_cf_keys = set()
            if batch_header:
                batch_cf_keys = {k for k in batch_header if not k.startswith("__")}
            for t in HEADER_TOKENS:
                if t in header_row and t not in batch_cf_keys:
                    blk = sub_token(blk, t, format_token_value(t, header_row[t]))

            # (b) Fill batch-level tokens from carry-forward data
            if batch_header:
                blk = _fill_batch_level_tokens(blk, batch_header, _known_tokens)

            # (c) Render rows
            filtered = []
            if rows_data:
                
                tbody_m, tbody_inner = best_rows_tbody(blk, allowed_row_tokens)
                
                if tbody_m and tbody_inner:
                    # --- "tbody" mode: row template inside <tbody> ---
                    row_template, row_span, rtt = find_row_template(tbody_inner, allowed_row_tokens)
                    if row_template and rtt:
                        rcols = [_extract_col_name(PLACEHOLDER_TO_COL.get(tok)) or "" for tok in rtt]
                        filtered = _filter_rows_for_render(rows_data, rtt, rcols, treat_all_as_data=bool(__force_single))
                        filtered = _prune_placeholder_rows(filtered, rtt)
                        if filtered:
                            parts = []
                            for row in filtered:
                                tr = row_template
                                for tok in rtt:
                                    tr = sub_token(tr, tok, format_token_value(tok, _value_for_token(row, tok)))
                                parts.append(tr)
                            new_inner = tbody_inner[:row_span[0]] + "\n".join(parts) + tbody_inner[row_span[1]:]
                            blk = blk[:tbody_m.start(1)] + new_inner + blk[tbody_m.end(1):]
                else:
                    # --- "tr" mode: prototype_block IS the row template ---
                    tr_toks = [
                        (m.group(1) or m.group(2)).strip()
                        for m in re.finditer(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", blk)
                    ]
                    rtt = [t for t in tr_toks if t in allowed_row_tokens]
                    if rtt:
                        rcols = [_extract_col_name(PLACEHOLDER_TO_COL.get(tok)) or "" for tok in rtt]
                        filtered = _filter_rows_for_render(rows_data, rtt, rcols, treat_all_as_data=bool(__force_single))
                        filtered = _prune_placeholder_rows(filtered, rtt)
                        if filtered:
                            parts = []
                            for row in filtered:
                                tr = blk
                                for tok in rtt:
                                    tr = sub_token(tr, tok, format_token_value(tok, _value_for_token(row, tok)))
                                parts.append(tr)
                            blk = "\n".join(parts)

            # (d) Fill totals
            if filtered and totals_data:
                for token in TOTALS:
                    value = totals_data.get(token)
                    formatted = format_token_value(token, value)
                    blk = sub_token(blk, token, formatted)
                    last_totals_per_token[token] = formatted
                    target = total_token_to_target.get(token)
                    if target:
                        fv, _ = _coerce_total_value(value)
                        if fv is not None:
                            totals_accum[target] = totals_accum.get(target, 0.0) + fv

            # (e) Blank out any remaining unfilled tokens (no DB data)
            for m in list(re.finditer(r"\{(\w+)\}", blk)):
                tok = m.group(1)
                if tok not in _known_tokens:
                    blk = sub_token(blk, tok, "")
            # Also blank UNRESOLVED totals that weren't filled
            for tok in TOTALS:
                if f"{{{tok}}}" in blk:
                    blk = sub_token(blk, tok, "")

            return blk, bool(filtered)

        # ── Render per-batch blocks ──
        if _use_per_batch:
            for batch_data in _df_batches:
                blk_html, had_rows = _render_df_block(
                    batch_data.get("rows", []),
                    batch_data.get("totals", {}),
                    batch_header=batch_data.get("header"),
                )
                if had_rows:
                    rendered_blocks.append(blk_html)
        else:
            # ── Single-block rendering (original path) ──
            rows_data = generator_results.get("rows") or []
            totals_data = (generator_results.get("totals") or [{}])[0]
            blk_html, had_rows = _render_df_block(rows_data, totals_data)
            if had_rows:
                rendered_blocks.append(blk_html)
            elif header_rows and HEADER_TOKENS:
                logger.debug("Generator: header-only block; appending without row data.")
                if ROW_TOKENS:
                    # Header-only block that expected rows — inject "no data" message
                    _no_data_msg = (
                        '<tr><td colspan="100" style="text-align:center;padding:20px;'
                        'color:#666;font-style:italic;">No data available for the '
                        'selected date range</td></tr>'
                    )
                    _tbody_re = re.compile(r'(<tbody[^>]*>)(.*?)(</tbody>)', re.DOTALL)
                    _tbody_match = _tbody_re.search(blk_html)
                    if _tbody_match:
                        blk_html = (
                            blk_html[:_tbody_match.start(2)]
                            + _no_data_msg
                            + blk_html[_tbody_match.end(2):]
                        )
                    else:
                        # No <tbody> — replace the block's row content
                        blk_html = re.sub(
                            r'<tr\b[^>]*>.*?</tr>',
                            _no_data_msg,
                            blk_html,
                            count=1,
                            flags=re.DOTALL,
                        )
                rendered_blocks.append(blk_html)
            elif ROW_TOKENS:
                # Had row tokens but no data and no headers — inject "no data" message
                logger.debug("Generator produced no row data; injecting empty-state message.")
                _no_data_msg = (
                    '<tr><td colspan="100" style="text-align:center;padding:20px;'
                    'color:#666;font-style:italic;">No data available for the '
                    'selected date range</td></tr>'
                )
                _tbody_re = re.compile(r'(<tbody[^>]*>)(.*?)(</tbody>)', re.DOTALL)
                _tbody_match = _tbody_re.search(blk_html)
                if _tbody_match:
                    blk_html = (
                        blk_html[:_tbody_match.start(2)]
                        + _no_data_msg
                        + blk_html[_tbody_match.end(2):]
                    )
                else:
                    blk_html = _no_data_msg
                rendered_blocks.append(blk_html)
            else:
                logger.debug("Generator produced no usable row data after filtering; skipping block.")
    # ---- Assemble full document ----
    rows_rendered = bool(rendered_blocks)
    if not rows_rendered:
        logger.debug("No rendered blocks generated for this selection.")

    html_multi = shell_prefix + "\n".join(rendered_blocks) + shell_suffix

    for tok, val in post_literal_specials.items():
        html_multi = sub_token(html_multi, tok, val if val is not None else "")

    if page_number_tokens or page_count_tokens or page_label_tokens:
        html_multi = _inject_page_counter_spans(html_multi, page_number_tokens, page_count_tokens, page_label_tokens)

    if total_token_to_target:
        overall_formatted = {}
        for (table_name, col_name), total in totals_accum.items():
            _, formatted = _coerce_total_value(total)
            overall_formatted[(table_name, col_name)] = formatted

        for token, target in total_token_to_target.items():
            table_name, col_name = target
            value = overall_formatted.get((table_name, col_name), last_totals_per_token.get(token, "0"))
            html_multi = sub_token(html_multi, token, value)

    # --- Fill remaining totals tokens from DF pipeline (outside batch blocks). ---
    if generator_results:
        totals_row = (generator_results.get("totals") or [{}])[0]
        if totals_row:
            for token in TOTALS:
                value = totals_row.get(token)
                if value is not None:
                    html_multi = sub_token(html_multi, token, format_token_value(token, value))

    # --- Fill remaining header tokens in the shell (outside batch blocks). ---
    if HEADER_TOKENS and generator_results:
        gen_header = (generator_results.get("header") or [{}])[0] if generator_results.get("header") else {}
        for t in HEADER_TOKENS:
            val = gen_header.get(t)
            if val is not None and str(val).strip():
                html_multi = sub_token(html_multi, t, format_token_value(t, val))

    # Apply literals globally
    for t, s in LITERALS.items():
        html_multi = sub_token(html_multi, t, s)

    # Blank any remaining known tokens
    ALL_KNOWN_TOKENS = set(HEADER_TOKENS) | set(ROW_TOKENS) | set(TOTALS.keys()) | set(LITERALS.keys())
    html_multi = blank_known_tokens(html_multi, ALL_KNOWN_TOKENS)

    # Strip internal BATCH markers — they are pipeline internals and must not leak into output
    html_multi = html_multi.replace(BEGIN_TAG, "").replace(END_TAG, "")

    # write to the path requested by the API
    _fp_progress("writing HTML output")
    OUT_HTML.write_text(html_multi, encoding="utf-8")
    _log_debug("Wrote HTML:", OUT_HTML)

    _fp_progress("starting PDF generation via Playwright")
    _html_to_pdf_subprocess(OUT_HTML, OUT_PDF, TEMPLATE_PATH.parent)
    _fp_progress("PDF generation complete")
    _log_debug("Wrote PDF via Playwright:", OUT_PDF)

    return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": rows_rendered}


# keep CLI usage (unchanged)
if __name__ == "__main__":
    print("Module ready for API integration. Call fill_and_print(...) from your FastAPI endpoint.")


# ======================================================================
# ReportGenerateExcel
# ======================================================================

import asyncio
import contextlib
import json
import logging
import os
import re
import sqlite3
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from decimal import Decimal, InvalidOperation
from itertools import product
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

logger = logging.getLogger(__name__)

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:
    import numpy as np
except ImportError:  # pragma: no cover
    np = None  # type: ignore

try:
    import cv2
except ImportError:  # pragma: no cover
    cv2 = None  # type: ignore

try:
    import fitz
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore

try:
    from skimage.metrics import structural_similarity as ssim
except ImportError:  # pragma: no cover
    ssim = None  # type: ignore

try:
    from playwright.async_api import async_playwright
except ImportError:  # pragma: no cover
    async_playwright = None  # type: ignore

import subprocess
import sys as _sys

_PDF_WORKER_SCRIPT = str(Path(__file__).with_name("_pdf_worker.py"))


def _run_async(coro):
    """Run an async coroutine safely whether or not an event loop is running."""
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        loop = None
    if loop and loop.is_running():
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            return pool.submit(asyncio.run, coro).result()
    return asyncio.run(coro)


def _pdf_worker_mp_target(html_path: str, pdf_path: str, base_dir: str, pdf_scale: float | None) -> None:
    """Target function for multiprocessing.Process — runs _convert in a fresh process."""
    asyncio.run(_convert(
        html_path=html_path,
        pdf_path=pdf_path,
        base_dir=base_dir,
        pdf_scale=pdf_scale,
    ))


# Timeout for the PDF worker process (30 minutes — large chunked docs with 10M+ rows).
_PDF_PROCESS_TIMEOUT = int(os.environ.get("NEURA_PDF_PROCESS_TIMEOUT", "3600"))


def _html_to_pdf_subprocess(
    html_path: Path, pdf_path: Path, base_dir: Path, pdf_scale: float | None = None
) -> None:
    """Convert HTML to PDF by running Playwright in a dedicated subprocess.

    This avoids the SIGCHLD / asyncio event-loop conflict that occurs when
    ``asyncio.run()`` is called from a non-main thread inside uvicorn.

    In PyInstaller frozen mode, sys.executable is the bundled exe which
    cannot run .py scripts.  We use multiprocessing.Process instead so the
    PDF work runs in a separate OS process — freeing the GIL and preventing
    the main backend from stalling during large chunked renders.
    """
    # PyInstaller frozen mode: use multiprocessing.Process (requires freeze_support)
    if getattr(_sys, "frozen", False):
        import multiprocessing

        args = (
            str(html_path.resolve()),
            str(pdf_path.resolve()),
            str((base_dir or html_path.parent).resolve()),
            pdf_scale,
        )
        proc = multiprocessing.Process(
            target=_pdf_worker_mp_target,
            args=args,
            daemon=False,
        )
        proc.start()
        proc.join(timeout=_PDF_PROCESS_TIMEOUT)
        if proc.is_alive():
            logger.error("PDF worker process timed out after %ds, terminating", _PDF_PROCESS_TIMEOUT)
            proc.terminate()
            proc.join(timeout=10)
            raise RuntimeError(f"PDF worker process timed out after {_PDF_PROCESS_TIMEOUT}s")
        if proc.exitcode != 0:
            raise RuntimeError(f"PDF worker process failed with exit code {proc.exitcode}")
        return

    import json as _json

    args_json = _json.dumps({
        "html_path": str(html_path.resolve()),
        "pdf_path": str(pdf_path.resolve()),
        "base_dir": str((base_dir or html_path.parent).resolve()),
        "pdf_scale": pdf_scale,
    })

    env = {**os.environ}
    # Ensure the project root is on PYTHONPATH so the worker can import backend.*
    project_root = str(Path(__file__).resolve().parents[3])
    env["PYTHONPATH"] = project_root + os.pathsep + env.get("PYTHONPATH", "")
    # Ensure Playwright uses a writable temp dir (some hosts quota-limit /tmp)
    if "TMPDIR" not in env:
        home_tmp = Path.home() / ".tmp"
        if home_tmp.is_dir():
            env["TMPDIR"] = str(home_tmp)

    result = subprocess.run(
        [_sys.executable, _PDF_WORKER_SCRIPT, args_json],
        capture_output=True,
        text=True,
        timeout=_PDF_PROCESS_TIMEOUT,
        env=env,
    )
    if result.returncode != 0:
        stderr_tail = (result.stderr or "")[-2000:]
        raise RuntimeError(f"PDF subprocess failed:\n{stderr_tail}")


from backend.app.repositories import SQLiteDataFrameLoader
_TABLE_RE = re.compile(r"(?is)<table\b[^>]*>(?P<body>.*?)</table>")
_ROW_RE = re.compile(r"(?is)<tr\b[^>]*>(?P<body>.*?)</tr>")
_CELL_RE = re.compile(r"(?is)<t(?:d|h)\b[^>]*>.*?</t(?:d|h)>")
_ROW_ONLY_RE = re.compile(r"(?is)<tr\b[^>]*>")
_COLSPAN_RE = re.compile(r'(?is)colspan\s*=\s*["\']?(\d+)["\']?')

_EXCEL_PRINT_STYLE_ID = "excel-print-sizing"
_EXCEL_PRINT_MARGIN_MM = 10
_DEFAULT_EXCEL_PRINT_SCALE = 0.82
_MIN_EXCEL_PRINT_SCALE = 0.4
_MAX_EXCEL_PRINT_SCALE = 0.97
_PAGE_HEIGHT_PX = 980
_BASE_ROW_HEIGHT_PX = 28
_EXCEL_STYLE_BLOCK_RE = re.compile(r'(?is)<style\b[^>]*id=["\']excel-print-sizing["\'][^>]*>.*?</style>')

_DATE_PARAM_START_ALIASES = {
    "start_ts_utc",
    "start_ts",
    "start_timestamp",
    "start_datetime",
    "start_date",
    "start_dt",
    "start_iso",
    "start_date_utc",
    "from_ts_utc",
    "from_ts",
    "from_timestamp",
    "from_datetime",
    "from_date",
    "from_dt",
    "from_iso",
    "from_date_utc",
    "range_start",
    "period_start",
}

_DATE_PARAM_END_ALIASES = {
    "end_ts_utc",
    "end_ts",
    "end_timestamp",
    "end_datetime",
    "end_date",
    "end_dt",
    "end_iso",
    "end_date_utc",
    "to_ts_utc",
    "to_ts",
    "to_timestamp",
    "to_datetime",
    "to_date",
    "to_dt",
    "to_iso",
    "to_date_utc",
    "range_end",
    "period_end",
}


def _clamp_excel_print_scale(scale: float | None) -> float:
    if scale is None or not isinstance(scale, (float, int)):
        return _DEFAULT_EXCEL_PRINT_SCALE
    return max(_MIN_EXCEL_PRINT_SCALE, min(float(scale), _MAX_EXCEL_PRINT_SCALE))


def _estimate_excel_print_scale(column_count: int) -> float:
    count = max(0, int(column_count))
    if count <= 8:
        return 0.98
    if count <= 12:
        return 0.9
    if count <= 16:
        return 0.82
    if count <= 20:
        return 0.74
    if count <= 24:
        return 0.68
    if count <= 30:
        return 0.6
    if count <= 36:
        return 0.52
    if count <= 44:
        return 0.48
    return 0.44


def _count_table_columns(html_text: str) -> int:
    max_cols = 0
    for table_match in _TABLE_RE.finditer(html_text or ""):
        table_body = table_match.group("body") or ""
        for row_match in _ROW_RE.finditer(table_body):
            row_html = row_match.group("body") or ""
            total = 0
            for cell_match in _CELL_RE.finditer(row_html):
                cell_html = cell_match.group(0)
                span_match = _COLSPAN_RE.search(cell_html)
                span = 1
                if span_match:
                    try:
                        span = max(1, int(span_match.group(1)))
                    except (TypeError, ValueError):
                        span = 1
                total += span
            if total:
                max_cols = max(max_cols, total)
        if max_cols:
            break
    return max_cols


def _count_table_rows(html_text: str) -> int:
    return len(_ROW_ONLY_RE.findall(html_text or ""))


def _estimate_rows_per_page(scale: float) -> int:
    clamped = _clamp_excel_print_scale(scale)
    if clamped <= 0:
        return 1
    capacity = int((_PAGE_HEIGHT_PX / _BASE_ROW_HEIGHT_PX) / clamped)
    return max(1, capacity)


def _inject_excel_print_styles(
    html_text: str,
    *,
    scale: float | None = None,
    rows_per_page: int | None = None,
) -> str:
    """
    Ensure Excel-generated reports include print-specific CSS that
    switches pages to landscape and downscales the rendered content.
    """
    scale_value = _clamp_excel_print_scale(scale)
    scale_str = f"{scale_value:.4f}".rstrip("0").rstrip(".")
    pagination_css = ""
    if rows_per_page and rows_per_page > 0:
        pagination_css = (
            "@media print {\n"
            f"  tbody tr:nth-of-type({rows_per_page}n+1):not(:first-child) {{\n"
            "    break-before: page;\n"
            "    page-break-before: always;\n"
            "  }\n"
            "}\n"
        )

    style_block = (
        f'\n<style id="{_EXCEL_PRINT_STYLE_ID}">\n'
        f":root {{ --excel-print-scale: {scale_str}; }}\n"
        "@page {\n"
        "  size: A4 landscape;\n"
        f"  margin: {_EXCEL_PRINT_MARGIN_MM}mm;\n"
        "}\n"
        "html, body {\n"
        "  margin: 0 auto;\n"
        "  padding: 0;\n"
        "  background: #fff;\n"
        "  max-width: calc(100% / var(--excel-print-scale));\n"
        "}\n"
        "body {\n"
        "  width: calc(100% / var(--excel-print-scale));\n"
        "  min-height: calc(100% / var(--excel-print-scale));\n"
        "  transform-origin: top left;\n"
        "}\n"
        "table {\n"
        "  width: 100% !important;\n"
        "  table-layout: fixed;\n"
        "  border-collapse: collapse;\n"
        "}\n"
        "th, td {\n"
        "  word-break: break-word;\n"
        "  white-space: normal;\n"
        "}\n"
        "@media screen {\n"
        "  body {\n"
        "    transform: scale(var(--excel-print-scale));\n"
        "  }\n"
        "}\n"
        "@media print {\n"
        "  body {\n"
        "    transform: scale(var(--excel-print-scale));\n"
        "    zoom: 1;\n"
        "  }\n"
        "}\n"
        f"{pagination_css}"
        "</style>\n"
    )

    if _EXCEL_STYLE_BLOCK_RE.search(html_text):
        return _EXCEL_STYLE_BLOCK_RE.sub(style_block, html_text, count=1)

    head_close = re.search(r"(?is)</head>", html_text)
    if head_close:
        idx = head_close.start()
        return f"{html_text[:idx]}{style_block}{html_text[idx:]}"

    return f"{style_block}{html_text}"


# ======================================================
# ENTRYPOINT: DB-driven fill + PDF for EXCEL templates
# ======================================================
def fill_and_print_excel(
    OBJ: dict,
    TEMPLATE_PATH: Path,
    DB_PATH: Path,
    OUT_HTML: Path,
    OUT_PDF: Path,
    START_DATE: str,
    END_DATE: str,
    batch_ids: list[str] | None = None,
    KEY_VALUES: dict | None = None,
    __force_single: bool = False,
    BRAND_KIT_ID: str | None = None,
):
    """
    DB-driven renderer:
      - Assumes TEMPLATE_PATH is already the *final shell* produced at Approve (auto_fill.py)
        containing a single prototype batch block.
      - Renders header tokens (parent row per batch), row repeater (child rows), totals, literals.
      - Writes OUT_HTML and prints OUT_PDF via Playwright.

    API contract preserved (same signature).
    """

    # ---- Guard required inputs ----
    for name in ("OBJ", "TEMPLATE_PATH", "DB_PATH", "START_DATE", "END_DATE"):
        if locals().get(name) is None:
            raise NameError(f"Missing required variable: `{name}`")

    # Ensure output dir exists
    OUT_DIR = OUT_HTML.parent
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    def _log_debug(*parts: object) -> None:
        message = " ".join(str(part) for part in parts)
        logger.debug(message)

    import time as _time
    _fp_start = _time.time()
    def _fp_progress(stage: str) -> None:
        elapsed = _time.time() - _fp_start
        print(f"[REPORT-EXCEL] {stage} ({elapsed:.1f}s)", flush=True)

    _fp_progress("fill_and_print START")
    _log_debug(
        "=== fill_and_print call ===",
        "force_single" if __force_single else "fanout_root",
        "KEY_VALUES raw=",
        KEY_VALUES or {},
    )

    # ---- Load the final shell HTML (created during Approve) ----
    from backend.app.services.infra_services import _fix_fixed_footers
    html = _fix_fixed_footers(TEMPLATE_PATH.read_text(encoding="utf-8"))

    # ---- Inject brand kit CSS if requested ----
    if BRAND_KIT_ID:
        try:
            from backend.app.services.ai_services import design_service
            brand_css = design_service.generate_brand_css_from_id(BRAND_KIT_ID)
            if brand_css:
                html = _inject_brand_css(html, brand_css)
                logger.debug("Brand kit CSS injected (Excel path): %s", BRAND_KIT_ID)
        except Exception:
            logger.warning("Failed to inject brand kit CSS (Excel path)", exc_info=True)

    # Support pre-built loaders (MultiDataFrameLoader), PostgreSQL, and SQLite
    if hasattr(DB_PATH, 'table_names') and callable(DB_PATH.table_names):
        dataframe_loader = DB_PATH
    elif hasattr(DB_PATH, 'is_postgresql') and DB_PATH.is_postgresql:
        from backend.app.services.connection_utils import get_loader_for_ref
        dataframe_loader = get_loader_for_ref(DB_PATH)
    else:
        dataframe_loader = SQLiteDataFrameLoader(DB_PATH)

    TOKEN_RE = re.compile(r"\{\{?\s*([A-Za-z0-9_\-\.]+)\s*\}\}?")
    TEMPLATE_TOKENS = {m.group(1) for m in TOKEN_RE.finditer(html)}

    # ---- Unpack contract ----
    OBJ = OBJ or {}
    contract_adapter = ContractAdapter(OBJ)
    PLACEHOLDER_TO_COL = contract_adapter.mapping
    param_token_set = {token for token in (contract_adapter.param_tokens or []) if token}

    join_raw = OBJ.get("join", {}) or {}
    JOIN = {
        "parent_table": contract_adapter.parent_table or join_raw.get("parent_table", ""),
        "child_table": contract_adapter.child_table or join_raw.get("child_table", ""),
        "parent_key": contract_adapter.parent_key or join_raw.get("parent_key", ""),
        "child_key": contract_adapter.child_key or join_raw.get("child_key", ""),
    }

    DATE_COLUMNS = contract_adapter.date_columns or (OBJ.get("date_columns", {}) or {})

    HEADER_TOKENS = contract_adapter.scalar_tokens or OBJ.get("header_tokens", [])
    ROW_TOKENS = contract_adapter.row_tokens or OBJ.get("row_tokens", [])
    row_token_count = sum(1 for tok in ROW_TOKENS if isinstance(tok, str) and tok.strip())
    TOTALS = contract_adapter.totals_mapping or OBJ.get("totals", {})
    ROW_ORDER = contract_adapter.row_order or OBJ.get("row_order", ["ROWID"])
    LITERALS = {
        str(token): "" if value is None else str(value) for token, value in (OBJ.get("literals", {}) or {}).items()
    }
    FORMATTERS = contract_adapter.formatters
    key_values_map: dict[str, list[str]] = {}
    if KEY_VALUES:
        for token, raw_value in KEY_VALUES.items():
            name = str(token or "").strip()
            if not name:
                continue
            values: list[str] = []
            if isinstance(raw_value, (list, tuple, set)):
                seen = set()
                for item in raw_value:
                    text = str(item or "").strip()
                    if text and text not in seen:
                        seen.add(text)
                        values.append(text)
            else:
                text = str(raw_value or "").strip()
                if text:
                    values = [text]
            if values:
                key_values_map[name] = values

    _DIRECT_COLUMN_RE = re.compile(r"^(?P<table>[A-Za-z_][\w]*)\.(?P<column>[A-Za-z_][\w]*)$")
    _SQL_IDENT_RE = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")

    def _safe_ident(name: str) -> str:
        if _SQL_IDENT_RE.match(name):
            return name
        safe = str(name).replace('"', '""')
        return f'"{safe}"'

    def _resolve_token_column(token: str) -> tuple[str, str] | None:
        mapping_expr = PLACEHOLDER_TO_COL.get(token)
        if isinstance(mapping_expr, str):
            match = _DIRECT_COLUMN_RE.match(mapping_expr.strip())
            if match:
                return match.group("table"), match.group("column")
        required_filters = contract_adapter.required_filters
        optional_filters = contract_adapter.optional_filters
        filter_expr = (required_filters.get(token) or optional_filters.get(token) or "").strip()
        match = _DIRECT_COLUMN_RE.match(filter_expr)
        if match:
            return match.group("table"), match.group("column")
        return None

    def _canonicalize_case(table: str, column: str, raw_value: str) -> str:
        normalized_table = str(table or "").strip().lower()
        normalized_column = str(column or "").strip().lower()
        normalized_value = str(raw_value or "").strip()
        cache_key = (normalized_table, normalized_column, normalized_value.lower())
        if cache_key in _canonicalize_cache:
            return _canonicalize_cache[cache_key]
        canonical = normalized_value
        if not normalized_table or not normalized_column or not normalized_value:
            _canonicalize_cache[cache_key] = canonical
            return canonical
        # Use a targeted SQL query instead of loading the entire table
        try:
            if hasattr(dataframe_loader, 'db_path'):
                import sqlite3 as _sqlite3
                quoted_t = table.replace('"', '""')
                quoted_c = column.replace('"', '""')
                with _sqlite3.connect(str(dataframe_loader.db_path), timeout=30) as _con:
                    row = _con.execute(
                        f'SELECT "{quoted_c}" FROM "{quoted_t}" WHERE "{quoted_c}" = ? COLLATE NOCASE LIMIT 1',
                        (normalized_value,)
                    ).fetchone()
                    if row and row[0] is not None:
                        canonical = str(row[0])
            else:
                # MultiDataFrameLoader or loaders without db_path — use DataFrame
                df = dataframe_loader.frame(table)
                if column in df.columns:
                    mask = df[column].astype(str).str.strip().str.lower() == normalized_value.lower()
                    matches = df.loc[mask, column]
                    if not matches.empty:
                        canonical = str(matches.iloc[0])
        except Exception:
            pass
        _canonicalize_cache[cache_key] = canonical
        return canonical

    _canonicalize_cache: dict[tuple[str, str, str], str] = {}

    for token, values in list(key_values_map.items()):
        resolved = _resolve_token_column(token)
        if not resolved:
            continue
        table_name, column_name = resolved
        if not table_name or not column_name:
            continue
        updated_values: list[str] = []
        changed = False
        for value in values:
            if not isinstance(value, str) or not value.strip():
                updated_values.append(value)
                continue
            canon = _canonicalize_case(table_name, column_name, value.strip())
            if canon != value:
                changed = True
            updated_values.append(canon)
        if changed:
            key_values_map[token] = updated_values

    for token, values in key_values_map.items():
        LITERALS[token] = ", ".join(values)

    alias_link_map: dict[str, str] = {}
    recipe_key_values = key_values_map.get("row_recipe_code")
    if recipe_key_values:
        alias_link_map = {
            "recipe_code": "row_recipe_code",
            "filter_recipe_code": "row_recipe_code",
        }
        literal_value = ", ".join(recipe_key_values)
        for alias in alias_link_map.keys():
            LITERALS[alias] = literal_value

    multi_key_selected = any(len(values) > 1 for values in key_values_map.values())

    def _first_alias_value(token: str) -> str | None:
        source = alias_link_map.get(token)
        if not source:
            return None
        return _first_key_value(key_values_map.get(source, []))

    def _apply_alias_params(target: dict[str, Any]) -> None:
        if not alias_link_map:
            return
        for alias in alias_link_map:
            if alias in target and str(target[alias] or "").strip():
                continue
            alias_value = _first_alias_value(alias)
            if alias_value is not None:
                target[alias] = alias_value

    _log_debug("Normalized key_values_map", key_values_map, "multi_key_selected", multi_key_selected)

    def _first_key_value(values: list[str]) -> str | None:
        for val in values:
            text = str(val or "").strip()
            if text:
                return text
        return None

    def _iter_key_combinations(values_map: dict[str, list[str]]) -> Iterable[dict[str, str]]:
        if not values_map:
            yield {}
            return
        tokens: list[str] = []
        value_lists: list[list[str]] = []
        for token, raw_values in values_map.items():
            unique: list[str] = []
            seen_local: set[str] = set()
            for val in raw_values:
                text = str(val or "").strip()
                if not text or text in seen_local:
                    continue
                seen_local.add(text)
                unique.append(text)
            if unique:
                tokens.append(token)
                value_lists.append(unique)
        if not tokens:
            yield {}
            return
        max_combos_raw = os.getenv("NEURA_REPORT_MAX_KEY_COMBINATIONS", "500")
        try:
            max_combos = int(max_combos_raw)
        except ValueError:
            max_combos = 500
        max_combos = max(1, max_combos)
        estimated = 1
        for values in value_lists:
            estimated *= max(1, len(values))
            if estimated > max_combos:
                raise ValueError(
                    f"Too many key combinations ({estimated} > {max_combos}). "
                    "Narrow key selections or reduce multi-select values."
                )
        for combo in product(*value_lists):
            yield {token: value for token, value in zip(tokens, combo)}

    _PLAYWRIGHT_ROW_FRIENDLY_LIMIT = 6000

    async def html_to_pdf_async(html_path: Path, pdf_path: Path, base_dir: Path, pdf_scale: float | None = None):
        if async_playwright is None:
            logger.warning("Playwright not available; skipping PDF generation.")
            return

        # Ensure TMPDIR exists (some systems lack a writable /tmp for Chromium)
        if not os.environ.get("TMPDIR"):
            _fallback_tmp = Path.home() / ".tmp"
            if _fallback_tmp.is_dir():
                os.environ["TMPDIR"] = str(_fallback_tmp)

        html_path_resolved = html_path.resolve()
        html_source = html_path_resolved.read_text(encoding="utf-8", errors="ignore")
        approx_row_count = html_source.lower().count("<tr")
        base_dir_resolved = (base_dir or html_path.parent).resolve()
        pdf_path_resolved = pdf_path.resolve()
        base_url = base_dir_resolved.as_uri()

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            context = None
            try:
                context = await browser.new_context(base_url=base_url)
                page = await context.new_page()
                _pdf_timeout_ms = int(os.environ.get("NEURA_PDF_RENDER_TIMEOUT_MS", "600000"))
                page.set_default_timeout(_pdf_timeout_ms)
                await page.set_content(html_source, wait_until="load", timeout=_pdf_timeout_ms)
                await page.emulate_media(media="print")
                scale_value = pdf_scale or 1.0
                if not isinstance(scale_value, (int, float)):
                    scale_value = 1.0
                scale_value = max(0.1, min(float(scale_value), 2.0))
                try:
                    await page.pdf(
                        path=str(pdf_path_resolved),
                        format="A4",
                        landscape=True,
                        print_background=True,
                        margin={"top": "10mm", "right": "10mm", "bottom": "10mm", "left": "10mm"},
                        prefer_css_page_size=True,
                        scale=scale_value,
                    )
                except Exception as exc:
                    if approx_row_count >= _PLAYWRIGHT_ROW_FRIENDLY_LIMIT:
                        raise RuntimeError(
                            (
                                "PDF rendering failed because the report contains "
                                f"approximately {approx_row_count:,} table rows, which exceeds the printable limit. "
                                "Please filter the data further or split the report into smaller chunks and try again."
                            )
                        ) from exc
                    raise
            finally:
                if context is not None:
                    await context.close()
                await browser.close()

    def _combine_html_documents(html_sections: list[str]) -> str:
        if not html_sections:
            return ""
        combined_body: list[str] = []
        doc_type = ""
        head_html = ""

        head_pattern = re.compile(r"(?is)<head\b[^>]*>(?P<head>.*)</head>")
        body_pattern = re.compile(r"(?is)<body\b[^>]*>(?P<body>.*)</body>")
        doctype_pattern = re.compile(r"(?is)^\s*<!DOCTYPE[^>]*>", re.MULTILINE)

        for idx, raw_html in enumerate(html_sections):
            text = raw_html or ""
            if idx == 0:
                doctype_match = doctype_pattern.search(text)
                if doctype_match:
                    doc_type = doctype_match.group(0).strip()
                    text = text[doctype_match.end() :]
                head_match = head_pattern.search(text)
                if head_match:
                    head_html = head_match.group(0).strip()
                body_match = body_pattern.search(text)
                if body_match:
                    section_body = body_match.group("body").strip()
                else:
                    section_body = text.strip()
                combined_body.append(f'<div class="nr-key-section" data-nr-section="1">\n{section_body}\n</div>')
            else:
                body_match = body_pattern.search(text)
                section = body_match.group("body").strip() if body_match else text.strip()
                combined_body.append(
                    f'<div class="nr-key-section" data-nr-section="{idx + 1}" style="page-break-before: always;">\n{section}\n</div>'
                )

        doc_lines = []
        if doc_type:
            doc_lines.append(doc_type)
        doc_lines.append("<html>")
        if head_html:
            doc_lines.append(head_html)
        doc_lines.append("<body>")
        doc_lines.append("\n\n".join(combined_body))
        doc_lines.append("</body>")
        doc_lines.append("</html>")
        return "\n".join(doc_lines)

    def _value_has_content(value: Any) -> bool:
        if value is None:
            return False
        if isinstance(value, (int, float, Decimal)):
            return value != 0
        text = str(value).strip()
        if not text:
            return False
        try:
            num = Decimal(text)
        except Exception:
            return True
        else:
            return num != 0

    def _row_has_significant_data(row: Mapping[str, Any], columns: list[str]) -> bool:
        return _row_has_any_data(row, (), columns)

    def _token_values_have_data(row: Mapping[str, Any], tokens: list[str]) -> bool:
        return _row_has_any_data(row, tokens, ())

    def _row_has_any_data(row: Mapping[str, Any], tokens: Sequence[str], columns: Sequence[str]) -> bool:
        for token in tokens:
            if not token:
                continue
            if _value_has_content(_value_for_token(row, token)):
                return True
        for col in columns:
            if not col:
                continue
            if _value_has_content(row.get(col)):
                return True
        for key, value in row.items():
            if not isinstance(key, str):
                continue
            if _is_counter_field(key):
                continue
            if _value_has_content(value):
                return True
        return False

    def _is_counter_field(name: str | None) -> bool:
        if not name:
            return False
        if not isinstance(name, str):
            name = str(name)
        normalized = re.sub(r"[^a-z0-9]", "", name.lower())
        if not normalized:
            return False
        if normalized in {
            "row",
            "rowid",
            "rowno",
            "rownum",
            "rownumber",
            "rowindex",
            "rowcounter",
            "srno",
            "sno",
        }:
            return True
        counter_markers = ("serial", "sequence", "seq", "counter")
        if any(marker in normalized for marker in counter_markers):
            return True
        # Exclude data fields that happen to end with counter-like suffixes
        # (e.g. row_bin_no is a bin identifier, row_recipe_no is a recipe ref)
        data_markers = ("bin", "recipe", "batch", "machine")
        if any(marker in normalized for marker in data_markers):
            return False
        counter_suffixes = (
            "slno",
            "srno",
            "sno",
            "snum",
            "snumber",
            "sl",
            "no",
            "num",
            "number",
            "idx",
            "index",
        )
        return any(normalized.endswith(suffix) and normalized.startswith("row") for suffix in counter_suffixes)

    def _reindex_serial_fields(rows: list[dict], tokens: Sequence[str], columns: Sequence[str]) -> None:
        serial_tokens = [tok for tok in tokens if tok and _is_counter_field(tok)]
        serial_columns = [col for col in columns if col and _is_counter_field(col)]
        if not serial_tokens and not serial_columns:
            return

        def _has_non_numeric(field: str) -> bool:
            for row in rows:
                val = row.get(field)
                if val is None or isinstance(val, (int, float)):
                    continue
                try:
                    float(str(val))
                except (ValueError, TypeError):
                    return True
            return False
        serial_tokens = [t for t in serial_tokens if not _has_non_numeric(t)]
        serial_columns = [c for c in serial_columns if not _has_non_numeric(c)]
        if not serial_tokens and not serial_columns:
            return
        for idx, row in enumerate(rows, start=1):
            for tok in serial_tokens:
                row[tok] = idx
            for col in serial_columns:
                row[col] = idx

    _warned_tokens: set[str] = set()  # log each unresolved token only once per generation

    def _value_for_token(row: Mapping[str, Any], token: str) -> Any:
        def _sanitize(v: Any) -> Any:
            if v is None:
                return None
            try:
                import math
                if isinstance(v, float) and math.isnan(v):
                    return None
            except (TypeError, ValueError):
                pass
            return v

        if not token:
            return None
        if token in row:
            return _sanitize(row[token])
        normalized = str(token).lower()
        for key in row.keys():
            if isinstance(key, str) and key.lower() == normalized:
                return _sanitize(row[key])
        mapped = PLACEHOLDER_TO_COL.get(token)
        if mapped:
            col = _extract_col_name(mapped)
            if col:
                if col in row:
                    return _sanitize(row[col])
                for key in row.keys():
                    if isinstance(key, str) and key.lower() == col.lower():
                        return _sanitize(row[key])
        if token not in _warned_tokens:
            _warned_tokens.add(token)
            logger.warning(
                "token_unresolved token=%s available_keys=%s",
                token,
                list(row.keys())[:10],
                extra={"event": "token_unresolved", "token": token},
            )
        return None

    def _prune_placeholder_rows(rows: Sequence[Mapping[str, Any]], tokens: Sequence[str]) -> list[dict[str, Any]]:
        material_tokens = [tok for tok in tokens if tok and "material" in tok.lower()]
        pruned: list[dict[str, Any]] = []
        for row in rows:
            keep = True
            for tok in material_tokens:
                if not _value_has_content(_value_for_token(row, tok)):
                    keep = False
                    break
            if keep:
                pruned.append(dict(row))
        return pruned if pruned else [dict(row) for row in rows]

    def _filter_rows_for_render(
        rows: Sequence[Mapping[str, Any]],
        row_tokens_template: Sequence[str],
        row_columns: Sequence[str],
        *,
        treat_all_as_data: bool,
    ) -> list[dict[str, Any]]:
        if not rows:
            return []

        if treat_all_as_data:
            prepared = [dict(row) for row in rows]
        else:
            significant_tokens = [tok for tok in row_tokens_template if tok and not _is_counter_field(tok)]
            significant_columns = [col for col in row_columns if col and not _is_counter_field(col)]
            prepared: list[dict[str, Any]] = []
            for row in rows:
                if significant_tokens or significant_columns:
                    if not _row_has_any_data(row, significant_tokens, significant_columns):
                        continue
                prepared.append(dict(row))

        if prepared:
            _reindex_serial_fields(prepared, row_tokens_template, row_columns)
        return prepared

    if multi_key_selected and not __force_single:
        html_sections: list[str] = []
        tmp_outputs: list[tuple[Path, Path]] = []
        try:
            for idx, combo in enumerate(_iter_key_combinations(key_values_map), start=1):
                selection: dict[str, str] = {token: value for token, value in combo.items()}
                if alias_link_map:
                    for alias, source in alias_link_map.items():
                        if alias not in selection and source in selection:
                            selection[alias] = selection[source]
                _log_debug("Fanout iteration", idx, "selection", selection)
                tmp_html = OUT_HTML.with_name(f"{OUT_HTML.stem}__key{idx}.html")
                tmp_pdf = OUT_PDF.with_name(f"{OUT_PDF.stem}__key{idx}.pdf")
                result = fill_and_print_excel(
                    OBJ=OBJ,
                    TEMPLATE_PATH=TEMPLATE_PATH,
                    DB_PATH=DB_PATH,
                    OUT_HTML=tmp_html,
                    OUT_PDF=tmp_pdf,
                    START_DATE=START_DATE,
                    END_DATE=END_DATE,
                    batch_ids=None,
                    KEY_VALUES=selection or None,
                    __force_single=True,
                )
                html_sections.append(Path(result["html_path"]).read_text(encoding="utf-8", errors="ignore"))
                tmp_outputs.append((Path(result["html_path"]), Path(result["pdf_path"])))

            if not html_sections:
                return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": False}

            combined_html = _combine_html_documents(html_sections)
            column_count = max(row_token_count, _count_table_columns(combined_html))
            excel_print_scale = _estimate_excel_print_scale(column_count)
            row_count = _count_table_rows(combined_html)
            rows_per_page = _estimate_rows_per_page(excel_print_scale)
            if row_count <= rows_per_page:
                rows_per_page = None
            combined_html = _inject_excel_print_styles(
                combined_html,
                scale=excel_print_scale,
                rows_per_page=rows_per_page,
            )
            OUT_HTML.write_text(combined_html, encoding="utf-8")
            _run_async(
                html_to_pdf_async(
                    OUT_HTML,
                    OUT_PDF,
                    TEMPLATE_PATH.parent,
                    pdf_scale=excel_print_scale,
                )
            )
            return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": True}
        finally:
            for tmp_html_path, tmp_pdf_path in tmp_outputs:
                for path_sel in (tmp_html_path, tmp_pdf_path):
                    with contextlib.suppress(FileNotFoundError):
                        path_sel.unlink()

    def _get_literal_raw(token: str) -> str:
        if token not in LITERALS:
            return ""
        raw = LITERALS[token]
        return "" if raw is None else str(raw)

    def _literal_has_content(token: str) -> bool:
        return bool(_get_literal_raw(token).strip())

    def _first_nonempty_literal(tokens: Iterable[str]) -> tuple[str | None, str | None]:
        for tok in tokens:
            raw = _get_literal_raw(tok)
            if raw.strip():
                return tok, raw
        return None, None

    def _record_special_value(target: dict[str, str], token: str, value: str) -> None:
        existing_raw = _get_literal_raw(token)
        if existing_raw.strip():
            target[token] = existing_raw
        else:
            target[token] = value
            if token in LITERALS:
                LITERALS[token] = value

    def _filter_tokens_without_literal(tokens: set[str]) -> set[str]:
        return {tok for tok in tokens if not _literal_has_content(tok)}

    BEGIN_TAG = "<!-- BEGIN:BATCH (auto) -->"
    END_TAG = "<!-- END:BATCH (auto) -->"
    try:
        prototype_block, start0, end_last = _select_prototype_block(html, ROW_TOKENS)
    except Exception as exc:
        _raise_no_block(html, exc)
    shell_prefix = html[:start0] + BEGIN_TAG
    shell_suffix = END_TAG + html[end_last:]

    parent_table = JOIN.get("parent_table", "")
    parent_key = JOIN.get("parent_key", "")
    child_table = JOIN.get("child_table", "")
    child_key = JOIN.get("child_key", "")

    # --- Additive: auto-detect date columns if missing from contract ---
    for _tbl in (parent_table, child_table):
        if _tbl and _tbl not in DATE_COLUMNS:
            _auto = detect_date_column(DB_PATH, _tbl)
            if _auto:
                DATE_COLUMNS[_tbl] = _auto
                logger.info("date_column_auto_detected table=%s col=%s", _tbl, _auto)

    parent_date = DATE_COLUMNS.get(parent_table, "")
    child_date = DATE_COLUMNS.get(child_table, "")
    order_col = ROW_ORDER[0] if ROW_ORDER else "ROWID"
    if isinstance(order_col, str) and order_col.upper() != "ROWID":
        mapped_order = PLACEHOLDER_TO_COL.get(order_col, order_col)
        if isinstance(mapped_order, str):
            mapped_order = mapped_order.strip()
            if "." in mapped_order:
                mapped_order = mapped_order.split(".", 1)[1].strip()
            if mapped_order:
                order_col = mapped_order

    def _normalize_token_name(name: str) -> str:
        return re.sub(r"[^a-z0-9]", "", name.lower())

    token_index: dict[str, set[str]] = defaultdict(set)
    all_candidate_tokens = (
        set(TEMPLATE_TOKENS) | set(HEADER_TOKENS) | set(ROW_TOKENS) | set(TOTALS.keys()) | set(LITERALS.keys())
    )

    def _token_synonym_keys(norm: str) -> set[str]:
        """
        Generate lightweight normalization aliases so that abbreviated tokens like
        `pg_total` or `page_num` still map onto the same lookup keys as their
        longer forms without needing every variant enumerated manually.
        """
        if not norm:
            return set()
        aliases = {norm}
        replacements: tuple[tuple[str, str], ...] = (
            ("pg", "page"),
            ("num", "number"),
            ("no", "number"),
            ("cnt", "count"),
            ("ttl", "total"),
        )
        for src, dest in replacements:
            if src in norm and dest not in norm:
                aliases.add(norm.replace(src, dest))
        # Avoid generating implausible short aliases (e.g., converting a lone "no"
        # in tokens unrelated to pagination), but include a fallback where a token
        # is exactly "pg" so that later lookups on "page" resolve.
        if norm == "pg":
            aliases.add("page")
        return {alias for alias in aliases if alias}

    for tok in all_candidate_tokens:
        norm = _normalize_token_name(tok)
        for key in _token_synonym_keys(norm):
            token_index[key].add(tok)

    def _tokens_for_keys(keys: set[str]) -> set[str]:
        found: set[str] = set()
        for key in keys:
            found.update(token_index.get(key, set()))
        return found

    def _format_for_db(dt_obj: datetime | None, raw_value, include_time_default: bool) -> str:
        """
        Normalize input dates for SQLite bindings:
          - prefer ISO 8601 date or datetime strings
          - fall back to trimmed raw strings when parsing fails
        """
        if dt_obj:
            include_time = include_time_default or bool(
                dt_obj.hour or dt_obj.minute or dt_obj.second or dt_obj.microsecond
            )
            if include_time:
                return dt_obj.strftime("%Y-%m-%d %H:%M:%S")
            return dt_obj.strftime("%Y-%m-%d")
        return "" if raw_value is None else str(raw_value).strip()

    start_dt = _parse_date_like(START_DATE)
    end_dt = _parse_date_like(END_DATE)
    _IST = timezone(timedelta(hours=5, minutes=30))
    print_dt = datetime.now(_IST)

    start_has_time = _has_time_component(START_DATE, start_dt)
    end_has_time = _has_time_component(END_DATE, end_dt)

    START_DATE_KEYS = {"fromdate", "datefrom", "startdate", "periodstart", "rangefrom", "fromdt", "startdt", "fromdatetime", "startdatetime", "datetimefrom"}
    END_DATE_KEYS = {"todate", "dateto", "enddate", "periodend", "rangeto", "todt", "enddt", "todatetime", "enddatetime", "datetimeto"}
    PRINT_DATE_KEYS = {
        "printdate",
        "printedon",
        "printeddate",
        "generatedon",
        "generateddate",
        "rundate",
        "runon",
        "generatedat",
    }
    PRINT_TIME_KEYS = {
        "printtime",
        "printedat",
        "generatedtime",
        "runtime",
    }
    PAGE_NO_KEYS = {
        "page",
        "pageno",
        "pagenum",
        "pagenumber",
        "pageindex",
        "pageidx",
        "pagecurrent",
        "currentpage",
        "currpage",
        "pgno",
        "pgnum",
        "pgnumber",
        "pgindex",
        "pgcurrent",
    }
    PAGE_COUNT_KEYS = {
        "pagecount",
        "pagecounts",
        "totalpages",
        "pagestotal",
        "pages",
        "pagetotal",
        "totalpage",
        "pagecounttotal",
        "totalpagecount",
        "pagescount",
        "countpages",
        "lastpage",
        "finalpage",
        "maxpage",
        "pgtotal",
        "totalpg",
        "pgcount",
        "countpg",
        "pgs",
        "pgscount",
        "pgstotal",
        "totalpgs",
    }
    PAGE_LABEL_KEYS = {
        "pagelabel",
        "pageinfo",
        "pagesummary",
        "pagefooter",
        "pagefootertext",
        "pageindicator",
        "pagecaption",
        "pagefooterlabel",
        "pagetext",
        "pagefooterinfo",
        "pagehint",
    }

    special_values: dict[str, str] = {}

    start_tokens = _tokens_for_keys(START_DATE_KEYS)
    end_tokens = _tokens_for_keys(END_DATE_KEYS)
    print_tokens = _tokens_for_keys(PRINT_DATE_KEYS)
    print_time_tokens = _tokens_for_keys(PRINT_TIME_KEYS)
    page_number_tokens = _tokens_for_keys(PAGE_NO_KEYS)
    page_count_tokens = _tokens_for_keys(PAGE_COUNT_KEYS)
    page_label_tokens = _tokens_for_keys(PAGE_LABEL_KEYS)

    for tok in start_tokens:
        _record_special_value(
            special_values,
            tok,
            _format_for_token(tok, start_dt, include_time_default=start_has_time),
        )
    for tok in end_tokens:
        _record_special_value(
            special_values,
            tok,
            _format_for_token(tok, end_dt, include_time_default=end_has_time),
        )

    _, print_literal_value = _first_nonempty_literal(print_tokens)
    parsed_print_dt = _parse_date_like(print_literal_value) if print_literal_value else None
    print_dt_source = parsed_print_dt or print_dt
    print_has_time = _has_time_component(print_literal_value, parsed_print_dt)

    for tok in print_tokens:
        if print_literal_value and not parsed_print_dt:
            value = print_literal_value
        else:
            value = _format_for_token(tok, print_dt_source, include_time_default=print_has_time)
        _record_special_value(special_values, tok, value)

    for tok in print_time_tokens:
        _record_special_value(special_values, tok, print_dt.strftime("%I:%M %p") if print_dt else "")

    page_number_tokens = _filter_tokens_without_literal(page_number_tokens)
    page_count_tokens = _filter_tokens_without_literal(page_count_tokens)
    page_label_tokens = _filter_tokens_without_literal(page_label_tokens)

    post_literal_specials = {tok: val for tok, val in special_values.items() if tok not in LITERALS}

    _ident_re = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*$")

    def qident(name: str) -> str:
        if _ident_re.match(name):
            return name
        safe = name.replace('"', '""')
        return f'"{safe}"'

    # ---- Composite-key helpers ----
    def _parse_key_cols(key_spec: str) -> list[str]:
        return [c.strip() for c in str(key_spec).split(",") if c and c.strip()]

    def _key_expr(cols: list[str]) -> str:
        parts = [f"COALESCE(CAST({qident(c)} AS TEXT),'')" for c in cols]
        if not parts:
            return "''"
        expr = parts[0]
        for p in parts[1:]:
            expr = f"{expr} || '|' || {p}"
        return expr

    def _split_bid(bid: str, n: int) -> list[str]:
        parts = str(bid).split("|")
        if len(parts) != n:
            raise ValueError(f"Composite key mismatch: expected {n} parts, got {len(parts)} in {bid!r}")
        return parts

    def _looks_like_composite_id(x: str, n: int) -> bool:
        return isinstance(x, str) and x.count("|") == (n - 1)

    pcols = _parse_key_cols(parent_key)
    ccols = _parse_key_cols(child_key)

    has_child = bool(child_table and ccols)
    parent_table_lc = parent_table.lower()
    child_table_lc = child_table.lower()
    parent_filter_map: dict[str, list[str]] = {}
    child_filter_map: dict[str, list[str]] = {}
    if key_values_map:
        for token, values in key_values_map.items():
            mapping_value = PLACEHOLDER_TO_COL.get(token)
            if not isinstance(mapping_value, str):
                continue
            target = mapping_value.strip()
            if not target or target.upper().startswith("PARAM:") or "." not in target:
                continue
            table_name, column_name = target.split(".", 1)
            table_name = table_name.strip(' "`[]')
            column_name = column_name.strip(' "`[]')
            if not column_name:
                continue
            table_key = table_name.lower()
            if table_key in (parent_table_lc, "header"):
                bucket = list(parent_filter_map.get(column_name, []))
                for val in values:
                    if val not in bucket:
                        bucket.append(val)
                if bucket:
                    parent_filter_map[column_name] = bucket
            if has_child and table_key in (child_table_lc, "rows"):
                bucket = list(child_filter_map.get(column_name, []))
                for val in values:
                    if val not in bucket:
                        bucket.append(val)
                if bucket:
                    child_filter_map[column_name] = bucket
    parent_filter_items = list(parent_filter_map.items())
    child_filter_items = list(child_filter_map.items())
    parent_filter_sqls: list[str] = []
    parent_filter_values: list[str] = []
    for col, values in parent_filter_items:
        normalized: list[str] = []
        for val in values:
            if not isinstance(val, str):
                continue
            text = val.strip()
            if text and text not in normalized:
                normalized.append(text)
        if not normalized:
            continue
        if len(normalized) == 1:
            parent_filter_sqls.append(f"{qident(col)} = ?")
        else:
            placeholders = ", ".join("?" for _ in normalized)
            parent_filter_sqls.append(f"{qident(col)} IN ({placeholders})")
        parent_filter_values.extend(normalized)
    parent_filter_values_tuple = tuple(parent_filter_values)

    child_filter_sqls: list[str] = []
    child_filter_values: list[str] = []
    for col, values in child_filter_items:
        normalized: list[str] = []
        for val in values:
            if not isinstance(val, str):
                continue
            text = val.strip()
            if text and text not in normalized:
                normalized.append(text)
        if not normalized:
            continue
        if len(normalized) == 1:
            child_filter_sqls.append(f"{qident(col)} = ?")
        else:
            placeholders = ", ".join("?" for _ in normalized)
            child_filter_sqls.append(f"{qident(col)} IN ({placeholders})")
        child_filter_values.extend(normalized)
    child_filter_values_tuple = tuple(child_filter_values)

    def _merge_predicate(base_sql: str, extras: list[str]) -> str:
        if not extras:
            return base_sql
        extras_joined = " AND ".join(extras)
        base_sql = (base_sql or "1=1").strip()
        return f"({base_sql}) AND {extras_joined}"

    # --- Date predicates and adapters (handle missing/invalid date columns)
    parent_type = get_col_type(DB_PATH, parent_table, parent_date)
    child_type = get_col_type(DB_PATH, child_table, child_date)
    parent_pred, adapt_parent = mk_between_pred_for_date(parent_date, parent_type)
    child_pred, adapt_child = mk_between_pred_for_date(child_date, child_type)
    parent_where_clause = _merge_predicate(parent_pred, parent_filter_sqls)
    child_where_clause = _merge_predicate(child_pred, child_filter_sqls) if has_child else child_pred
    db_start = _format_for_db(start_dt, START_DATE, start_has_time)
    db_end = _format_for_db(end_dt, END_DATE, end_has_time)
    PDATE = tuple(adapt_parent(db_start, db_end))  # () if 1=1
    CDATE = tuple(adapt_child(db_start, db_end)) if has_child else tuple()  # () if 1=1
    parent_params_all = tuple(PDATE) + parent_filter_values_tuple
    child_params_all = tuple(CDATE) + child_filter_values_tuple if has_child else tuple()

    sql_params: dict[str, object] = {
        "from_date": db_start,
        "to_date": db_end,
        "start_date": db_start,
        "end_date": db_end,
    }

    for token in contract_adapter.param_tokens:
        if token in ("from_date", "to_date", "start_date", "end_date"):
            continue
        if token in key_values_map:
            first_value = _first_key_value(key_values_map[token])
            if first_value is not None:
                sql_params[token] = first_value
        elif alias_link_map.get(token):
            alias_value = _first_alias_value(token)
            if alias_value is not None:
                sql_params[token] = alias_value
        elif token in LITERALS:
            sql_params[token] = LITERALS[token]
        elif token in special_values:
            sql_params[token] = special_values[token]
        else:
            sql_params.setdefault(token, "")

    _apply_alias_params(sql_params)

    def _apply_date_param_defaults(target: dict[str, object]) -> None:
        if not isinstance(target, dict):
            return

        def _inject(names: set[str], default_value: str) -> None:
            if not default_value:
                return
            for alias in names:
                if alias not in param_token_set and alias not in target:
                    continue
                current = target.get(alias)
                if isinstance(current, str):
                    if current.strip():
                        continue
                elif current not in (None, ""):
                    continue
                target[alias] = default_value

        _inject(_DATE_PARAM_START_ALIASES, db_start)
        _inject(_DATE_PARAM_END_ALIASES, db_end)

    _apply_date_param_defaults(sql_params)

    # --- DataFrame pipeline (sole data path) ---

    df_value_filters: dict[str, list] = {}
    if key_values_map:
        for name, values in key_values_map.items():
            df_value_filters[name] = values

    pipeline = DataFramePipeline(
        contract_adapter=contract_adapter,
        loader=dataframe_loader,
        params=sql_params,
        start_date=sql_params.get("start_date") or sql_params.get("from_date"),
        end_date=sql_params.get("end_date") or sql_params.get("to_date"),
        value_filters=df_value_filters,
    )
    generator_results = pipeline.execute()

    BATCH_IDS = ["__DF_PIPELINE__"]

    _log_debug("BATCH_IDS:", len(BATCH_IDS or []), (BATCH_IDS or [])[:20] if BATCH_IDS else [])
    # ---- Only touch tokens outside <style>/<script> ----
    def format_token_value(token: str, raw_value: Any) -> str:
        return contract_adapter.format_value(token, raw_value)


    def _inject_page_counter_spans(
        html_in: str,
        page_tokens: set[str],
        count_tokens: set[str],
        label_tokens: set[str] | None = None,
    ) -> str:
        tokens_to_remove = (
            {tok for tok in page_tokens if tok}
            | {tok for tok in count_tokens if tok}
            | {tok for tok in (label_tokens or set()) if tok}
        )
        if not tokens_to_remove:
            return html_in

        token_pattern = "|".join(re.escape(tok) for tok in tokens_to_remove)
        placeholder_pattern = rf"(?:\{{\{{\s*(?:{token_pattern})\s*\}}\}}|\{{\s*(?:{token_pattern})\s*\}})"
        placeholder_re = re.compile(placeholder_pattern)
        element_re = re.compile(
            rf"(?is)<(?P<tag>[a-z0-9:_-]+)(?:\s[^>]*)?>(?:(?!</(?P=tag)>).)*?{placeholder_pattern}(?:(?!</(?P=tag)>).)*?</(?P=tag)>"
        )
        line_re = re.compile(rf"(?im)^[^\n]*{placeholder_pattern}[^\n]*\n?")

        def _remove_blocks(text: str) -> str:
            while True:
                text, count = element_re.subn("", text)
                if count == 0:
                    break
            text = line_re.sub("", text)
            text = placeholder_re.sub("", text)
            return text

        return _remove_blocks(html_in)

    # ---- Helpers to find tbody / row template (improved) ----
    def best_rows_tbody(inner_html: str, allowed_tokens: set):
        tbodys = list(re.finditer(r"(?is)<tbody\b[^>]*>(.*?)</tbody>", inner_html))
        best = (None, None, -1)  # (match, inner, hits)
        for m in tbodys:
            tin = m.group(1)
            hits = 0
            for trm in re.finditer(r"(?is)<tr\b[^>]*>.*?</tr>", tin):
                tr_html = trm.group(0)
                toks = re.findall(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", tr_html)
                flat = [a.strip() if a else b.strip() for (a, b) in toks]
                hits += sum(1 for t in flat if t in allowed_tokens)
            if hits > best[2]:
                best = (m, tin, hits)
        if best[0] is not None:
            return best[0], best[1]
        return (tbodys[0], tbodys[0].group(1)) if tbodys else (None, None)

    def find_row_template(tbody_inner: str, allowed_tokens: set):
        for m in re.finditer(r"(?is)<tr\b[^>]*>.*?</tr>", tbody_inner):
            tr_html = m.group(0)
            toks = re.findall(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", tr_html)
            flat = []
            for a, b in toks:
                if a:
                    flat.append(a.strip())
                if b:
                    flat.append(b.strip())
            flat = [t for t in flat if t in allowed_tokens]
            if flat:
                return tr_html, (m.start(0), m.end(0)), sorted(set(flat), key=len, reverse=True)
        return None, None, []

    def majority_table_for_tokens(tokens, mapping):
        from collections import Counter

        tbls = []
        for t in tokens:
            tc = mapping.get(t, "")
            if "." in tc:
                tbls.append(tc.split(".", 1)[0])
        return Counter(tbls).most_common(1)[0][0] if tbls else None

    # ---- Pre-compute minimal column sets ----
    def _extract_col_name(mapping_value: str | None) -> str | None:
        if not isinstance(mapping_value, str):
            return None
        target = mapping_value.strip()
        if "." not in target:
            return None
        # Skip SQL expressions (e.g. "table.col || ' ' || table.col2")
        if _SQL_EXPR_CHARS.search(target):
            return None
        after_dot = target.split(".", 1)[1].strip()
        if not after_dot:
            return None
        col = re.split(r"[,)\s]", after_dot, 1)[0].strip()
        return col or None

    header_cols = sorted({col for t in HEADER_TOKENS for col in [_extract_col_name(PLACEHOLDER_TO_COL.get(t))] if col})
    row_cols = sorted({col for t in ROW_TOKENS for col in [_extract_col_name(PLACEHOLDER_TO_COL.get(t))] if col})

    totals_by_table = defaultdict(lambda: defaultdict(list))
    total_token_to_target = {}

    for token, raw_target in TOTALS.items():
        if isinstance(raw_target, dict):
            continue  # Declarative op spec — handled by DF pipeline, not SQL prefetch
        target = (raw_target or PLACEHOLDER_TO_COL.get(token, "")).strip()
        if not target or "." not in target:
            continue
        table_name, col_name = [part.strip() for part in target.split(".", 1)]
        if not table_name or not col_name:
            continue
        totals_by_table[table_name][col_name].append(token)
        total_token_to_target[token] = (table_name, col_name)

    def _coerce_total_value(raw):
        if raw is None:
            return None, "0"
        try:
            decimal_value = Decimal(str(raw).strip())
        except (InvalidOperation, ValueError, TypeError, AttributeError):
            return None, "0"
        if not decimal_value.is_finite():
            return None, "0"
        formatted = format_decimal_str(decimal_value, max_decimals=3) or "0"
        return float(decimal_value), formatted

    totals_accum = defaultdict(float)
    last_totals_per_token = {token: "0" for token in TOTALS}

    child_totals_cols = {col: list(tokens) for col, tokens in totals_by_table.get(child_table, {}).items()}
    # ---- Render all batches ----
    rendered_blocks = []
    generator_header_replacements: dict[str, str] | None = None
    if generator_results is not None:
        block_html = prototype_block

        header_rows = generator_results.get("header") or []
        header_row = header_rows[0] if header_rows else {}
        header_token_values: dict[str, str] = {}
        for t in HEADER_TOKENS:
            if t in header_row:
                value = header_row[t]
                formatted_value = format_token_value(t, value)
                block_html = sub_token(block_html, t, formatted_value)
                header_token_values[t] = formatted_value
        if header_token_values:
            generator_header_replacements = header_token_values

        allowed_row_tokens = {t for t in PLACEHOLDER_TO_COL.keys() if t not in TOTALS} - set(HEADER_TOKENS)
        rows_data = generator_results.get("rows") or []
        filtered_rows: list[dict[str, Any]] = []
        row_tokens_in_template: list[str] = []

        if rows_data:
            tbody_m, tbody_inner = best_rows_tbody(block_html, allowed_row_tokens)
            if tbody_m and tbody_inner:
                row_template, row_span, row_tokens_in_template = find_row_template(tbody_inner, allowed_row_tokens)
                if row_template and row_tokens_in_template:
                    row_columns_template = [
                        _extract_col_name(PLACEHOLDER_TO_COL.get(tok)) or "" for tok in row_tokens_in_template
                    ]
                    render_columns = list(row_tokens_in_template)
                    filtered_rows = _filter_rows_for_render(
                        rows_data,
                        row_tokens_in_template,
                        render_columns,
                        treat_all_as_data=bool(__force_single),
                    )
                    filtered_rows = _prune_placeholder_rows(filtered_rows, row_tokens_in_template)
                    if __force_single:
                        _log_debug(
                            f"[multi-debug] generator rows: total={len(rows_data)}, filtered={len(filtered_rows)}, key_values={KEY_VALUES}"
                        )
                    if filtered_rows:
                        parts: list[str] = []
                        for row in filtered_rows:
                            tr = row_template
                            for tok in row_tokens_in_template:
                                val = _value_for_token(row, tok)
                                tr = sub_token(tr, tok, format_token_value(tok, val))
                            parts.append(tr)
                        new_tbody_inner = tbody_inner[: row_span[0]] + "\n".join(parts) + tbody_inner[row_span[1] :]
                        block_html = block_html[: tbody_m.start(1)] + new_tbody_inner + block_html[tbody_m.end(1) :]
            else:
                tr_tokens = [
                    m.group(1) or m.group(2)
                    for m in re.finditer(r"\{\{\s*([^}\n]+?)\s*\}\}|\{\s*([^}\n]+?)\s*\}", block_html)
                ]
                row_tokens_in_template = [t.strip() for t in tr_tokens if t and t.strip() in allowed_row_tokens]
                if row_tokens_in_template:
                    row_columns_template = [
                        _extract_col_name(PLACEHOLDER_TO_COL.get(tok)) or "" for tok in row_tokens_in_template
                    ]
                    render_columns = list(row_tokens_in_template)
                    filtered_rows = _filter_rows_for_render(
                        rows_data,
                        row_tokens_in_template,
                        render_columns,
                        treat_all_as_data=bool(__force_single),
                    )
                    filtered_rows = _prune_placeholder_rows(filtered_rows, row_tokens_in_template)
                    if __force_single:
                        _log_debug(
                            f"[multi-debug] generator rows (no tbody): total={len(rows_data)}, filtered={len(filtered_rows)}, key_values={KEY_VALUES}"
                        )
                    if filtered_rows:
                        parts = []
                        for row in filtered_rows:
                            tr = prototype_block
                            for tok in row_tokens_in_template:
                                val = _value_for_token(row, tok)
                                tr = sub_token(tr, tok, format_token_value(tok, val))
                            parts.append(tr)
                        block_html = "\n".join(parts)

        if filtered_rows:
            totals_row = (generator_results.get("totals") or [{}])[0]
            for token in TOTALS:
                value = totals_row.get(token)
                formatted = format_token_value(token, value)
                block_html = sub_token(block_html, token, formatted)
                last_totals_per_token[token] = formatted
                target = total_token_to_target.get(token)
                if target:
                    fv, _formatted = _coerce_total_value(value)
                    if fv is not None:
                        totals_accum[target] = totals_accum.get(target, 0.0) + fv

            rendered_blocks.append(block_html)
        else:
            _log_debug("Generator SQL produced no usable row data after filtering; skipping block.")

    # ---- Assemble full document ----
    rows_rendered = bool(rendered_blocks)
    if not rows_rendered:
        _log_debug("No rendered blocks generated for this selection.")

    html_multi = shell_prefix + "\n".join(rendered_blocks) + shell_suffix

    # Substitute totals into the assembled document (tfoot may be in shell_suffix)
    for token, formatted in last_totals_per_token.items():
        html_multi = sub_token(html_multi, token, formatted)

    for tok, val in post_literal_specials.items():
        html_multi = sub_token(html_multi, tok, val if val is not None else "")

    if page_number_tokens or page_count_tokens or page_label_tokens:
        html_multi = _inject_page_counter_spans(html_multi, page_number_tokens, page_count_tokens, page_label_tokens)

    if total_token_to_target:
        overall_formatted = {}
        for (table_name, col_name), total in totals_accum.items():
            _, formatted = _coerce_total_value(total)
            overall_formatted[(table_name, col_name)] = formatted

        for token, target in total_token_to_target.items():
            table_name, col_name = target
            value = overall_formatted.get((table_name, col_name), last_totals_per_token.get(token, "0"))
            html_multi = sub_token(html_multi, token, value)

    if generator_header_replacements:
        for token, value in generator_header_replacements.items():
            html_multi = sub_token(html_multi, token, value)

    # Apply literals globally
    for t, s in LITERALS.items():
        html_multi = sub_token(html_multi, t, s)

    # Blank any remaining known tokens
    ALL_KNOWN_TOKENS = set(HEADER_TOKENS) | set(ROW_TOKENS) | set(TOTALS.keys()) | set(LITERALS.keys())
    html_multi = blank_known_tokens(html_multi, ALL_KNOWN_TOKENS)

    # Strip internal BATCH markers — they are pipeline internals and must not leak into output
    html_multi = html_multi.replace(BEGIN_TAG, "").replace(END_TAG, "")

    column_count = max(row_token_count, _count_table_columns(html_multi))
    excel_print_scale = _estimate_excel_print_scale(column_count)
    row_count = _count_table_rows(html_multi)
    rows_per_page = _estimate_rows_per_page(excel_print_scale)
    if row_count <= rows_per_page:
        rows_per_page = None
    html_multi = _inject_excel_print_styles(
        html_multi,
        scale=excel_print_scale,
        rows_per_page=rows_per_page,
    )

    # write to the path requested by the API
    _fp_progress("writing HTML output")
    OUT_HTML.write_text(html_multi, encoding="utf-8")
    _log_debug("Wrote HTML:", OUT_HTML)

    _fp_progress("starting PDF generation via Playwright")
    # Use subprocess isolation for Playwright to avoid SIGCHLD / asyncio
    # event-loop conflicts when running inside uvicorn's thread pool.
    _html_to_pdf_subprocess(
        OUT_HTML,
        OUT_PDF,
        TEMPLATE_PATH.parent,
        pdf_scale=excel_print_scale,
    )
    _log_debug("Wrote PDF via Playwright subprocess:", OUT_PDF)

    return {"html_path": str(OUT_HTML), "pdf_path": str(OUT_PDF), "rows_rendered": rows_rendered}


# keep CLI usage (unchanged)
if __name__ == "__main__":
    print("Module ready for API integration. Call fill_and_print_excel(...) from your FastAPI endpoint.")


# ======================================================================
# _pdf_worker
# ======================================================================

#!/usr/bin/env python3
"""Standalone Playwright PDF worker — runs in its own process.

This script is invoked as a subprocess by ReportGenerateExcel to isolate
Playwright's Chromium browser from the main uvicorn event loop.  Running
Playwright in its own process avoids the SIGCHLD / asyncio-subprocess
conflict that occurs when ``asyncio.run()`` is called from a non-main
thread inside the web server.

Usage:
    python _pdf_worker.py <json-args>

The JSON argument must contain:
    html_path   – absolute path to the source HTML file
    pdf_path    – absolute path for the output PDF
    base_dir    – absolute path used as Playwright's base_url
    pdf_scale   – optional float (0.1 – 2.0, default 1.0)
"""

import asyncio
import json
import os
import re
import sys
import tempfile
from pathlib import Path

# Configurable timeout for Playwright page operations (default: 5 minutes).
# Large reports (10M+ rows) generate huge HTML that takes time to render.
_PDF_RENDER_TIMEOUT_MS = int(os.environ.get("NEURA_PDF_RENDER_TIMEOUT_MS", "600000"))

# Maximum number of <tr> rows before we switch to chunked PDF generation.
_CHUNK_THRESHOLD = int(os.environ.get("NEURA_PDF_CHUNK_THRESHOLD", "8000"))
# Rows per chunk when chunking.
_CHUNK_SIZE = int(os.environ.get("NEURA_PDF_CHUNK_SIZE", "5000"))


def _count_tr(html: str) -> int:
    """Fast count of <tr> tags in the HTML."""
    return html.lower().count("<tr")


def _split_html_chunks(html: str, chunk_size: int) -> list[str]:
    """Split a large HTML table into multiple smaller HTML documents.

    Preserves <head>, header elements, <thead>, and <tfoot> in each chunk.
    Splits only the <tbody> rows.
    """
    # Extract everything before <tbody> and after </tbody>
    tbody_match = re.search(r"(<tbody[^>]*>)(.*?)(</tbody>)", html, re.DOTALL | re.IGNORECASE)
    if not tbody_match:
        return [html]

    pre_tbody = html[: tbody_match.start(2)]   # everything up to and including <tbody>
    tbody_content = tbody_match.group(2)
    post_tbody = html[tbody_match.end(2):]      # </tbody> and everything after

    # Split tbody content into individual <tr>...</tr> blocks
    rows = re.findall(r"<tr[\s>].*?</tr>", tbody_content, re.DOTALL | re.IGNORECASE)
    if not rows:
        return [html]

    chunks = []
    for i in range(0, len(rows), chunk_size):
        chunk_rows = rows[i : i + chunk_size]
        chunk_html = pre_tbody + "\n".join(chunk_rows) + post_tbody
        chunks.append(chunk_html)

    return chunks


async def _render_single(page, pdf_path: str, scale: float) -> None:
    """Render a single page to PDF."""
    await page.emulate_media(media="print")
    await page.pdf(
        path=pdf_path,
        format="A4",
        landscape=True,
        print_background=True,
        margin={"top": "10mm", "right": "10mm", "bottom": "10mm", "left": "10mm"},
        prefer_css_page_size=True,
        scale=scale,
    )


def _merge_pdfs(pdf_paths: list[str], output_path: str) -> None:
    """Merge multiple PDF files into one using pikepdf (or PyPDF2 fallback)."""
    try:
        import pikepdf
        merged = pikepdf.Pdf.new()
        for p in pdf_paths:
            src = pikepdf.open(p)
            merged.pages.extend(src.pages)
        merged.save(output_path)
        merged.close()
    except ImportError:
        from PyPDF2 import PdfMerger
        merger = PdfMerger()
        for p in pdf_paths:
            merger.append(p)
        merger.write(output_path)
        merger.close()


async def _launch_browser(p):
    """Launch a Chromium-based browser, preferring system Chrome/Edge.

    Strategy:
    1. Try system Edge (pre-installed on all Windows 10/11)
    2. Try system Chrome (commonly installed)
    3. Fall back to Playwright's own Chromium (downloaded at first launch)
    """
    launch_args = [
        "--disable-dev-shm-usage",
        "--disable-gpu",
        "--no-sandbox",
    ]

    # Try system browsers first — no download needed, no AV issues
    for channel in ("msedge", "chrome"):
        try:
            browser = await p.chromium.launch(channel=channel, args=launch_args)
            print(f"[pdf_worker] Using system browser: {channel}", file=sys.stderr)
            return browser
        except Exception:
            continue

    # Fall back to Playwright's bundled/downloaded Chromium
    browser = await p.chromium.launch(args=launch_args)
    print("[pdf_worker] Using Playwright Chromium", file=sys.stderr)
    return browser


async def _convert(html_path: str, pdf_path: str, base_dir: str, pdf_scale: float | None = None) -> None:
    from playwright.async_api import async_playwright

    html_source = Path(html_path).read_text(encoding="utf-8", errors="ignore")
    base_url = Path(base_dir).resolve().as_uri()

    scale_value = pdf_scale or 1.0
    if not isinstance(scale_value, (int, float)):
        scale_value = 1.0
    scale_value = max(0.1, min(float(scale_value), 2.0))

    tr_count = _count_tr(html_source)
    needs_chunking = tr_count > _CHUNK_THRESHOLD

    async with async_playwright() as p:
        browser = await _launch_browser(p)

        if not needs_chunking:
            # Standard single-pass rendering
            context = await browser.new_context(base_url=base_url)
            try:
                page = await context.new_page()
                page.set_default_timeout(_PDF_RENDER_TIMEOUT_MS)
                await page.set_content(html_source, wait_until="load", timeout=_PDF_RENDER_TIMEOUT_MS)
                await _render_single(page, pdf_path, scale_value)
            finally:
                await context.close()
                await browser.close()
            return

        # Chunked rendering for large documents
        print(f"[pdf_worker] Large document ({tr_count} rows), using chunked rendering", file=sys.stderr)
        chunks = _split_html_chunks(html_source, _CHUNK_SIZE)
        print(f"[pdf_worker] Split into {len(chunks)} chunks", file=sys.stderr)

        tmp_dir = tempfile.mkdtemp(prefix="neura_pdf_chunks_")
        chunk_paths: list[str] = []

        try:
            for idx, chunk_html in enumerate(chunks):
                chunk_pdf = os.path.join(tmp_dir, f"chunk_{idx:04d}.pdf")
                context = await browser.new_context(base_url=base_url)
                try:
                    page = await context.new_page()
                    page.set_default_timeout(_PDF_RENDER_TIMEOUT_MS)
                    await page.set_content(chunk_html, wait_until="load", timeout=_PDF_RENDER_TIMEOUT_MS)
                    await _render_single(page, chunk_pdf, scale_value)
                    chunk_paths.append(chunk_pdf)
                    print(f"[pdf_worker] Chunk {idx + 1}/{len(chunks)} done", file=sys.stderr)
                finally:
                    await context.close()

            # Merge all chunks into the final PDF
            _merge_pdfs(chunk_paths, pdf_path)
            print(f"[pdf_worker] Merged {len(chunk_paths)} chunks into {pdf_path}", file=sys.stderr)
        finally:
            await browser.close()
            # Cleanup temp files
            for cp in chunk_paths:
                try:
                    os.unlink(cp)
                except OSError:
                    pass
            try:
                os.rmdir(tmp_dir)
            except OSError:
                pass


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: _pdf_worker.py <json-args>", file=sys.stderr)
        sys.exit(1)

    args = json.loads(sys.argv[1])
    asyncio.run(
        _convert(
            html_path=args["html_path"],
            pdf_path=args["pdf_path"],
            base_dir=args["base_dir"],
            pdf_scale=args.get("pdf_scale"),
        )
    )


if __name__ == "__main__":
    main()


# ======================================================================
# xlsx_export
# ======================================================================

import logging
from pathlib import Path
from typing import Optional

# Prefer xlsxwriter (streaming, constant memory) over openpyxl (in-memory).
_xlsxwriter = None
try:
    import xlsxwriter as _xlsxwriter  # type: ignore
except ImportError:
    _xlsxwriter = None

# Fallback: openpyxl (loads entire workbook into RAM — OOM on large reports).
_openpyxl = None
try:
    import openpyxl as _openpyxl  # type: ignore
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # type: ignore
    from openpyxl.utils import get_column_letter  # type: ignore
    from openpyxl.worksheet.table import Table, TableStyleInfo  # type: ignore
except ImportError:  # pragma: no cover
    _openpyxl = None
    Alignment = None  # type: ignore
    Border = None  # type: ignore
    Font = None  # type: ignore
    PatternFill = None  # type: ignore
    Side = None  # type: ignore
    get_column_letter = None  # type: ignore
    Table = None  # type: ignore
    TableStyleInfo = None  # type: ignore


logger = logging.getLogger("neura.reports.xlsx")


# ---------------------------------------------------------------------------
# Shared HTML → rows parsing
# ---------------------------------------------------------------------------

def _table_score(table: list[list[str]]) -> int:
    if not table:
        return 0
    row_count = len(table)
    max_cols = max((len(row) for row in table), default=0)
    multi_col_rows = sum(1 for row in table if sum(1 for cell in row if cell.strip()) >= 2)
    return (multi_col_rows or row_count) * max(1, max_cols)


def _select_best_table_index(tables: list[list[list[str]]]) -> int:
    best_idx = 0
    best_score = -1
    for idx, table in enumerate(tables):
        score = _table_score(table)
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx


def _looks_like_total_row(row: list[str]) -> bool:
    for cell in row:
        if cell and cell.strip().lower() == "total":
            return True
    return False


def _parse_html_to_rows(html_text: str):
    """Parse HTML and return (rows, data_row_positions, preface_ranges, data_header_row_idx, data_thead_count)."""
    tables_with_meta = extract_tables_with_header_counts(html_text)
    tables = [t for t, _ in tables_with_meta]
    thead_counts = [c for _, c in tables_with_meta]
    rows: list[list[str]] = []
    data_row_positions: list[int] = []
    preface_ranges: list[tuple[int, int]] = []
    data_header_row_idx: int | None = None

    if tables:
        best_idx = _select_best_table_index(tables)
        data_thead_count = thead_counts[best_idx] if best_idx < len(thead_counts) else 1

        for idx, table in enumerate(tables):
            is_data_table = idx == best_idx
            if not table:
                continue
            serial_counter = 0
            table_start_idx = len(rows) + 1
            header_rows = data_thead_count if is_data_table else 1
            for row_idx_in_table, row in enumerate(table):
                clean_row = [(cell or "").strip() for cell in row]
                if is_data_table and row_idx_in_table >= header_rows:
                    if _looks_like_total_row(clean_row):
                        if clean_row:
                            clean_row[0] = ""
                    else:
                        serial_counter += 1
                        if not clean_row[0]:
                            clean_row[0] = str(serial_counter)
                elif is_data_table and row_idx_in_table == header_rows - 1:
                    data_header_row_idx = len(rows) + 1
                rows.append(clean_row)
                if is_data_table and row_idx_in_table >= header_rows:
                    data_row_positions.append(len(rows))
            table_end_idx = len(rows)
            if (not is_data_table) and table_end_idx >= table_start_idx:
                preface_ranges.append((table_start_idx, table_end_idx))
            rows.append([])

        while rows and not rows[-1]:
            rows.pop()
    else:
        rows = [[line.strip()] for line in html_text.splitlines() if line.strip()]
        if not rows:
            rows = [["Report output unavailable"]]

    # Pad all rows to the same width so preface/header rows span the full
    # sheet width in Excel (prevents narrow preface rows when data has many columns).
    max_cols = max((len(r) for r in rows if r), default=1)
    for i, row in enumerate(rows):
        if row and len(row) < max_cols:
            rows[i] = row + [""] * (max_cols - len(row))

    return rows, data_row_positions, preface_ranges, data_header_row_idx


# ---------------------------------------------------------------------------
# xlsxwriter-based export (streaming, constant memory)
# ---------------------------------------------------------------------------

def _html_file_to_xlsx_xlsxwriter(html_path: Path, output_path: Path) -> Optional[Path]:
    """Export HTML to XLSX using xlsxwriter (streaming writer, constant memory)."""
    html_text = html_path.read_text(encoding="utf-8", errors="ignore")
    rows, data_row_positions, preface_ranges, data_header_row_idx = _parse_html_to_rows(html_text)

    data_start = data_row_positions[0] if data_row_positions else None
    data_end = data_row_positions[-1] if data_row_positions else None
    data_max_cols = (
        max(len(rows[idx - 1]) for idx in data_row_positions)
        if data_row_positions
        else max((len(r) for r in rows if r), default=1)
    )
    data_max_cols = max(1, data_max_cols)
    sheet_width = max((len(r) for r in rows if r), default=data_max_cols)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    wb = _xlsxwriter.Workbook(str(output_path), {"constant_memory": False})
    ws = wb.add_worksheet("Report")

    # Pre-define format objects (xlsxwriter requires this)
    fmt_default = wb.add_format({
        "text_wrap": True,
        "valign": "top",
        "align": "left",
    })
    fmt_bold = wb.add_format({
        "bold": True,
        "text_wrap": True,
        "valign": "top",
        "align": "left",
    })
    fmt_title = wb.add_format({
        "bold": True,
        "font_size": 14,
        "bg_color": "#BDD7EE",
        "text_wrap": True,
        "valign": "vcenter",
        "align": "center",
    })
    fmt_preface = wb.add_format({
        "bold": True,
        "font_size": 11,
        "bg_color": "#D9E1F2",
        "text_wrap": True,
        "valign": "vcenter",
        "align": "center",
    })
    fmt_data_header = wb.add_format({
        "bold": True,
        "font_color": "#FFFFFF",
        "bg_color": "#2F75B5",
        "text_wrap": True,
        "valign": "vcenter",
        "align": "center",
    })
    fmt_border = wb.add_format({
        "border": 1,
        "border_color": "#C0C0C0",
        "text_wrap": True,
        "valign": "top",
        "align": "left",
    })

    # Build set lookups for fast row classification
    preface_row_set: set[int] = set()
    first_preface_row = preface_ranges[0][0] if preface_ranges else None
    for start_idx, end_idx in preface_ranges:
        for ri in range(start_idx, end_idx + 1):
            preface_row_set.add(ri)

    # Track max column widths for auto-sizing
    col_widths: dict[int, int] = {}

    for r_idx, row in enumerate(rows, start=1):
        if not row:
            continue

        # Determine format for this row
        is_preface = r_idx in preface_row_set
        is_data_header = r_idx == data_header_row_idx
        is_title = is_preface and r_idx == first_preface_row

        if is_title:
            row_fmt = fmt_title
        elif is_preface:
            row_fmt = fmt_preface
        elif is_data_header:
            row_fmt = fmt_data_header
        elif r_idx == 1:
            row_fmt = fmt_bold
        else:
            row_fmt = fmt_border

        # For preface title/subtitle rows with only 1 non-empty cell,
        # merge across all data columns so the text spans the full width.
        non_empty = [i for i, v in enumerate(row) if v and str(v).strip()]
        if (is_title or (is_preface and len(non_empty) == 1)) and data_max_cols > 1:
            val = row[non_empty[0]] if non_empty else ""
            ws.merge_range(r_idx - 1, 0, r_idx - 1, data_max_cols - 1, val, row_fmt)
        else:
            for c_idx, value in enumerate(row):
                ws.write(r_idx - 1, c_idx, value, row_fmt)
        # Track width
        for c_idx, value in enumerate(row):
            text_len = len(str(value)) if value else 0
            old = col_widths.get(c_idx, 0)
            if text_len > old:
                col_widths[c_idx] = text_len

    # Set column widths
    for c_idx, max_len in col_widths.items():
        width = min(120, max(12, max_len + 2))
        ws.set_column(c_idx, c_idx, width)

    # Freeze panes
    if data_header_row_idx:
        freeze_row = data_header_row_idx  # 1-based → 0-based is data_header_row_idx - 1, but freeze expects row below
        ws.freeze_panes(freeze_row, 0)
    elif data_start:
        ws.freeze_panes(data_start, 0)
    elif len(rows) > 1:
        ws.freeze_panes(1, 0)

    # Autofilter on the data range
    if data_header_row_idx and data_end:
        ws.autofilter(data_header_row_idx - 1, 0, data_end - 1, data_max_cols - 1)
    elif len(rows) > 0:
        ws.autofilter(0, 0, len(rows) - 1, sheet_width - 1)

    try:
        wb.close()
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "xlsx_export_save_failed",
            extra={
                "event": "xlsx_export_save_failed",
                "html_path": str(html_path),
                "xlsx_path": str(output_path),
                "error": str(exc),
                "engine": "xlsxwriter",
            },
        )
        return None

    logger.info(
        "xlsx_export_success",
        extra={
            "event": "xlsx_export_success",
            "html_path": str(html_path),
            "xlsx_path": str(output_path),
            "engine": "xlsxwriter",
        },
    )
    return output_path


# ---------------------------------------------------------------------------
# openpyxl-based export (fallback — loads everything into RAM)
# ---------------------------------------------------------------------------

def _auto_column_widths(worksheet) -> None:
    if get_column_letter is None:  # pragma: no cover
        return
    for col_idx in range(1, worksheet.max_column + 1):
        letter = get_column_letter(col_idx)
        max_len = 0
        for cell in worksheet[letter]:
            value = cell.value
            if value is None:
                continue
            text = str(value)
            max_len = max(max_len, len(text))
        width = min(120, max(12, max_len + 2))
        worksheet.column_dimensions[letter].width = width


def _html_file_to_xlsx_openpyxl(html_path: Path, output_path: Path) -> Optional[Path]:
    """Export HTML to XLSX using openpyxl (in-memory — may OOM on large reports)."""
    html_text = html_path.read_text(encoding="utf-8", errors="ignore")
    rows, data_row_positions, preface_ranges, data_header_row_idx = _parse_html_to_rows(html_text)

    data_start = data_row_positions[0] if data_row_positions else None
    data_end = data_row_positions[-1] if data_row_positions else None
    data_max_cols = (
        max(len(rows[idx - 1]) for idx in data_row_positions)
        if data_row_positions
        else max((len(r) for r in rows if r), default=1)
    )
    data_max_cols = max(1, data_max_cols)
    sheet_width = max((len(r) for r in rows if r), default=data_max_cols)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    wb = _openpyxl.Workbook()
    ws = wb.active
    ws.title = "Report"

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            if Alignment is not None:
                cell.alignment = Alignment(
                    horizontal="left",
                    vertical="top",
                    wrap_text=True,
                )
            if Font is not None and r_idx == 1:
                cell.font = Font(bold=True)

        if len(row) == 0:
            continue

    if (
        PatternFill is not None
        and Alignment is not None
        and Font is not None
        and sheet_width > 0
        and preface_ranges
    ):
        header_fill = PatternFill("solid", fgColor="D9E1F2")
        title_fill = PatternFill("solid", fgColor="BDD7EE")
        first_preface_row = preface_ranges[0][0]
        for start_idx, end_idx in preface_ranges:
            for row_idx in range(start_idx, end_idx + 1):
                row_data = rows[row_idx - 1] if 0 <= row_idx - 1 < len(rows) else []
                if not row_data or not any(cell for cell in row_data):
                    continue
                ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=sheet_width)
                cell = ws.cell(row=row_idx, column=1)
                cell_text = cell.value or ""
                is_title_row = row_idx == first_preface_row and cell_text and cell_text.upper() == cell_text
                cell.alignment = Alignment(
                    horizontal="center" if is_title_row else "left",
                    vertical="center",
                    wrap_text=True,
                )
                cell.font = Font(bold=True, size=14 if is_title_row else 11)
                cell.fill = title_fill if is_title_row else header_fill

    if (
        PatternFill is not None
        and Alignment is not None
        and Font is not None
        and data_header_row_idx is not None
        and data_max_cols > 0
    ):
        data_header_fill = PatternFill("solid", fgColor="2F75B5")
        for col_idx in range(1, data_max_cols + 1):
            cell = ws.cell(row=data_header_row_idx, column=col_idx)
            cell.fill = data_header_fill
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    if Border is not None and Side is not None:
        thin = Side(style="thin", color="FFC0C0C0")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for row in ws.iter_rows(
            min_row=1,
            max_row=ws.max_row,
            min_col=1,
            max_col=ws.max_column,
        ):
            for cell in row:
                cell.border = border

    if data_header_row_idx:
        freeze_row = data_header_row_idx + 1
    else:
        freeze_row = (data_start + 1) if data_start else 2
    if ws.max_row >= freeze_row:
        ws.freeze_panes = f"A{freeze_row}"
    elif ws.max_row > 1:
        ws.freeze_panes = "A2"

    if Table is not None and TableStyleInfo is not None and ws.max_column > 0 and data_end is not None:
        table_cols = max(1, min(data_max_cols, ws.max_column))
        table_top = data_header_row_idx if data_header_row_idx is not None else data_start
        if table_top is not None and table_top <= data_end:
            ref = f"A{table_top}:{ws.cell(row=data_end, column=table_cols).coordinate}"
            table = Table(displayName="ReportTable", ref=ref)
            table.tableStyleInfo = TableStyleInfo(
                name="TableStyleMedium9",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False,
            )
            ws.add_table(table)
        else:
            ws.auto_filter.ref = f"A1:{ws.cell(row=ws.max_row, column=ws.max_column).coordinate}"
    else:
        ws.auto_filter.ref = f"A1:{ws.cell(row=ws.max_row, column=ws.max_column).coordinate}"

    _auto_column_widths(ws)

    try:
        wb.save(output_path)
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "xlsx_export_save_failed",
            extra={
                "event": "xlsx_export_save_failed",
                "html_path": str(html_path),
                "xlsx_path": str(output_path),
                "error": str(exc),
                "engine": "openpyxl",
            },
        )
        return None

    logger.info(
        "xlsx_export_success",
        extra={
            "event": "xlsx_export_success",
            "html_path": str(html_path),
            "xlsx_path": str(output_path),
            "engine": "openpyxl",
        },
    )
    return output_path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def html_file_to_xlsx(html_path: Path, output_path: Path) -> Optional[Path]:
    """Convert an HTML report file to XLSX.

    Uses xlsxwriter (streaming, constant memory) when available.
    Falls back to openpyxl if xlsxwriter is not installed.
    """
    if _xlsxwriter is not None:
        return _html_file_to_xlsx_xlsxwriter(html_path, output_path)
    if _openpyxl is not None:
        logger.info("xlsx_export_using_openpyxl_fallback")
        return _html_file_to_xlsx_openpyxl(html_path, output_path)
    logger.warning(
        "xlsx_export_unavailable",
        extra={
            "event": "xlsx_export_unavailable",
            "reason": "neither xlsxwriter nor openpyxl installed",
            "html_path": str(html_path),
        },
    )
    return None


# ======================================================================
# docx_export
# ======================================================================

import contextlib
import logging
import re
import time
from io import BytesIO
from pathlib import Path
from typing import Iterable, Optional

try:
    from html2docx import html2docx  # type: ignore
except ImportError:  # pragma: no cover
    html2docx = None  # type: ignore

try:  # pragma: no cover - only exercised when python-docx is available
    from docx import Document  # type: ignore
    from docx.enum.section import WD_ORIENT  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Mm, Pt  # type: ignore
except ImportError:  # pragma: no cover
    Document = None  # type: ignore
    WD_ORIENT = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    Mm = None  # type: ignore
    Pt = None  # type: ignore

try:  # pragma: no cover - optional dependency for PDF conversion
    from pdf2docx import Converter  # type: ignore
except ImportError:  # pragma: no cover
    Converter = None  # type: ignore

try:
    from lxml import etree
    from lxml import html as lxml_html  # type: ignore
except ImportError:  # pragma: no cover
    etree = None  # type: ignore
    lxml_html = None  # type: ignore


logger = logging.getLogger("neura.reports.docx")

_BODY_TAG_RE = re.compile(r"(?is)<body\b(?P<attrs>[^>]*)>", re.MULTILINE)
_STYLE_ATTR_RE = re.compile(r'(?is)(style\s*=\s*)(["\'])(?P<value>.*?)\2')
_STYLE_BLOCK_RE = re.compile(r"(?is)<style\b[^>]*>.*?</style>")
_SCRIPT_BLOCK_RE = re.compile(r"(?is)<script\b[^>]*>.*?</script>")
_TITLE_RE = re.compile(r"(?is)<title\b[^>]*>(?P<value>.*?)</title>")
_TAG_RE = re.compile(r"(?is)<[^>]+>")
_WHITESPACE_RE = re.compile(r"\s+")

_A4_LANDSCAPE_WIDTH_MM = 297
_A4_LANDSCAPE_HEIGHT_MM = 210
_MAX_FALLBACK_ROWS = 500
_MAX_FALLBACK_TABLES = 8


def _inject_body_style(html_text: str, style_rule: str) -> str:
    """Attach/extend a style attribute on the <body> tag without disturbing the markup."""
    match = _BODY_TAG_RE.search(html_text)
    if not match:
        return f'<body style="{style_rule}">{html_text}</body>'

    attrs = match.group("attrs") or ""
    new_attrs = attrs

    def _style_repl(style_match: re.Match[str]) -> str:
        existing = (style_match.group("value") or "").strip().rstrip(";")
        merged = "; ".join(filter(None, [existing, style_rule]))
        return f"{style_match.group(1)}{style_match.group(2)}{merged}{style_match.group(2)}"

    if _STYLE_ATTR_RE.search(attrs):
        new_attrs = _STYLE_ATTR_RE.sub(_style_repl, attrs, count=1)
    else:
        spacer = "" if attrs.endswith(" ") or not attrs else " "
        new_attrs = f'{attrs}{spacer}style="{style_rule}"'

    new_tag = f"<body{new_attrs}>"
    return f"{html_text[: match.start()]}{new_tag}{html_text[match.end():]}"


def _apply_body_font_scale(html_text: str, scale: float | None) -> str:
    if not scale or scale <= 0:
        return html_text

    clamped = max(0.5, min(scale, 1.0))
    percent = round(clamped * 100, 1)
    style_rule = f"font-size: {percent}%; line-height: 1.15;"
    return _inject_body_style(html_text, style_rule)


def _append_inline_style(node, style: str) -> None:
    if node is None or not style:
        return
    existing = (node.get("style") or "").strip()
    if existing and not existing.endswith(";"):
        existing = f"{existing};"
    parts = [part.strip() for part in (existing.rstrip(";"), style.strip()) if part and part.strip()]
    node.set("style", "; ".join(parts))


def _inline_report_styles(html_text: str) -> str:
    if not html_text:
        return html_text
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return html_text

    def _set_style(xpath: str, style: str) -> None:
        if not style:
            return
        for node in document.xpath(xpath):
            _append_inline_style(node, style)

    _set_style(
        "//body", "font-family: 'Times New Roman', serif; font-size: 12px; color: #000; line-height: 1.2; margin: 0;"
    )
    _set_style("//div[@id='report-header']", "margin-top: 0; page-break-inside: avoid;")
    _set_style(
        "//div[@id='report-header']//table",
        "width: 100%; border: 1px solid #000; border-collapse: collapse; table-layout: fixed;",
    )
    _set_style("//div[@id='report-header']//td", "vertical-align: top; padding: 1.6mm 2.4mm; border: 1px solid #000;")
    _set_style(
        "//div[contains(concat(' ', normalize-space(@class), ' '), ' title-wrap ')]",
        "margin: 3mm 0 2.5mm 0; text-align: center; font-weight: bold; font-size: 18px; border-top: 1px solid #000; border-bottom: 1px solid #000; padding: 2mm 0;",
    )
    _set_style("//table[@id='data-table']", "width: 100%; border-collapse: collapse; table-layout: fixed;")
    _set_style(
        "//table[@id='data-table']//th | //table[@id='data-table']//td", "border: 1px solid #000; padding: 1mm 2.2mm;"
    )
    _set_style(
        "//table[@id='data-table']//thead//th",
        "text-align: center; font-weight: bold; white-space: nowrap; padding: 1.6mm 2.4mm;",
    )
    _set_style(
        "//table[@id='data-table']//td[contains(concat(' ', normalize-space(@class), ' '), ' num ')]",
        "text-align: right;",
    )
    _set_style("//tfoot[@id='report-totals']//td", "font-weight: bold; border-top: 1.2px solid #000;")
    _set_style(
        "//tfoot[@id='report-totals']//td[contains(concat(' ', normalize-space(@class), ' '), ' label ')]",
        "text-align: left;",
    )
    _set_style(
        "//footer[@id='report-footer']",
        "font-size: 11px; color: #000; display: flex; justify-content: space-between; align-items: center;",
    )
    _set_style(
        "//footer[@id='report-footer']//div[contains(concat(' ', normalize-space(@class), ' '), ' page ')]",
        "text-align: center;",
    )

    return etree.tostring(document, encoding="unicode", method="html")


def _extract_report_title(html_text: str) -> str:
    if not html_text:
        return ""
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return _extract_html_title(html_text)
    title_nodes = document.xpath("//*[contains(concat(' ', normalize-space(@class), ' '), ' title-wrap ')]")
    if title_nodes:
        return _normalize_whitespace(title_nodes[0].text_content())
    return _extract_html_title(html_text)


def _extract_footer_brand(html_text: str) -> str:
    if not html_text:
        return ""
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return ""
    brand_nodes = document.xpath(
        "//footer[@id='report-footer']//div[contains(concat(' ', normalize-space(@class), ' '), ' brand ')]"
    )
    if brand_nodes:
        return _normalize_whitespace(brand_nodes[0].text_content())
    return ""


def _configure_document_layout(document, *, body_font_scale: float | None = None) -> None:
    if document is None:
        return
    if Mm is not None:
        try:
            section = document.sections[0]
        except Exception:
            section = None
        if section is not None:
            section.left_margin = Mm(16)
            section.right_margin = Mm(16)
            section.top_margin = Mm(14)
            section.bottom_margin = Mm(14)


def _infer_numeric_columns(header_row: list[str]) -> set[int]:
    numeric_columns: set[int] = set()
    tokens = ("wt", "weight", "error", "%", "kg", "total", "qty")
    for idx, cell in enumerate(header_row or []):
        text = (cell or "").lower()
        if idx == 0:
            continue
        if any(token in text for token in tokens):
            numeric_columns.add(idx)
    return numeric_columns


def _column_widths(max_columns: int, *, ratios: Optional[Iterable[float]], document) -> list[float] | None:
    if max_columns <= 0 or document is None or Mm is None:
        return None
    try:
        section = document.sections[0]
        available = section.page_width - section.left_margin - section.right_margin
    except Exception:
        return None
    if available <= 0:
        return None
    ratio_list = list(ratios or [])
    if ratio_list and len(ratio_list) < max_columns:
        last = ratio_list[-1]
        ratio_list.extend([last] * (max_columns - len(ratio_list)))
    if not ratio_list:
        ratio_list = [1.0] * max_columns
    total = sum(ratio_list) or 1.0
    return [available * (value / total) for value in ratio_list[:max_columns]]


def _write_docx_table(
    document,
    rows: list[list[str]],
    *,
    header_rows: int = 0,
    column_widths: Optional[Iterable[float]] = None,
    numeric_columns: Optional[set[int]] = None,
) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    widths = _column_widths(max_cols, ratios=column_widths, document=document)

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.text = _normalize_whitespace(value)
            if WD_ALIGN_PARAGRAPH is not None:
                if numeric_columns and c_idx in (numeric_columns or set()) and r_idx >= header_rows:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif r_idx < header_rows:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = r_idx < header_rows

    if widths:
        for c_idx, width in enumerate(widths):
            for cell in table.columns[c_idx].cells:
                cell.width = width


def _strip_style_blocks(html_text: str) -> str:
    cleaned = _STYLE_BLOCK_RE.sub("", html_text or "")
    cleaned = _SCRIPT_BLOCK_RE.sub("", cleaned)
    return cleaned


def _normalize_whitespace(text: str) -> str:
    return _WHITESPACE_RE.sub(" ", text or "").strip()


def _strip_html_to_text(html_text: str) -> str:
    return _normalize_whitespace(_TAG_RE.sub(" ", html_text or ""))


def _extract_html_title(html_text: str) -> str:
    match = _TITLE_RE.search(html_text or "")
    if not match:
        return ""
    return _strip_html_to_text(match.group("value"))


def _extract_section_nodes(html_text: str) -> list[etree._Element]:
    try:
        document = lxml_html.fromstring(html_text or "")
    except Exception:
        return []
    sections = document.xpath("//div[contains(concat(' ', normalize-space(@class), ' '), ' nr-key-section ')]")
    if sections:
        return sections
    body = document.xpath("//body")
    return body or [document]


def _fallback_docx_from_tables(html_text: str, output_path: Path, *, body_font_scale: float | None) -> Optional[Path]:
    if Document is None:  # pragma: no cover
        return None

    sections = _extract_section_nodes(html_text)
    if not sections:
        try:
            sections = [lxml_html.fromstring(html_text or "<div></div>")]
        except Exception:
            sections = []
    if not sections:
        return None

    try:
        document = Document()  # type: ignore[call-arg]
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_fallback_init_failed",
            extra={
                "event": "docx_fallback_init_failed",
                "error": str(exc),
            },
        )
        return None

    _configure_document_layout(document, body_font_scale=body_font_scale)

    for index, section_node in enumerate(sections):
        if index > 0:
            document.add_page_break()

        section_html = etree.tostring(section_node, encoding="unicode", method="html")
        title_text = _extract_report_title(section_html) or _extract_report_title(html_text)
        if title_text:
            paragraph = document.add_paragraph(title_text)
            if WD_ALIGN_PARAGRAPH is not None:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

        tables = extract_tables(section_html, max_tables=_MAX_FALLBACK_TABLES)
        if not tables:
            plain_text = _strip_html_to_text(section_html)
            document.add_paragraph(plain_text or "Report data unavailable.")
            continue

        header_rows = tables[0] if len(tables) >= 1 else []
        data_rows = tables[1] if len(tables) >= 2 else []
        extra_tables = tables[2:] if len(tables) > 2 else []

        if header_rows:
            capped_header = header_rows[:_MAX_FALLBACK_ROWS]
            _write_docx_table(
                document,
                capped_header,
                header_rows=0,
                column_widths=[2.5, 1.5],
            )
            document.add_paragraph("")

        if data_rows:
            capped_rows = data_rows[:_MAX_FALLBACK_ROWS]
            numeric_columns = _infer_numeric_columns(capped_rows[0] if capped_rows else [])
            _write_docx_table(
                document,
                capped_rows,
                header_rows=1,
                numeric_columns=numeric_columns,
            )
            document.add_paragraph("")

        for rows in extra_tables:
            if not rows:
                continue
            capped_rows = rows[:_MAX_FALLBACK_ROWS]
            _write_docx_table(document, capped_rows, header_rows=1)
            document.add_paragraph("")

        brand_text = _extract_footer_brand(section_html) or _extract_footer_brand(html_text)
        if brand_text:
            footer_paragraph = document.add_paragraph(brand_text)
            if WD_ALIGN_PARAGRAPH is not None:
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    try:
        document.save(output_path)
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_fallback_save_failed",
            extra={
                "event": "docx_fallback_save_failed",
                "docx_path": str(output_path),
                "error": str(exc),
            },
        )
        return None
    return output_path


def _clamp_body_scale(scale: float | None) -> float:
    if not scale or not isinstance(scale, (int, float)):
        return 1.0
    return max(0.3, min(float(scale), 1.0))


def _enforce_landscape_layout(docx_path: Path, *, margin_mm: float = 10.0) -> None:
    if Document is None or WD_ORIENT is None or Mm is None:  # pragma: no cover
        logger.debug(
            "docx_landscape_skipped",
            extra={
                "event": "docx_landscape_skipped",
                "reason": "python-docx unavailable",
                "docx_path": str(docx_path),
            },
        )
        return

    try:
        document = Document(docx_path)  # type: ignore[call-arg]
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_landscape_open_failed",
            extra={
                "event": "docx_landscape_open_failed",
                "docx_path": str(docx_path),
                "error": str(exc),
            },
        )
        return

    width = Mm(_A4_LANDSCAPE_WIDTH_MM)
    height = Mm(_A4_LANDSCAPE_HEIGHT_MM)
    margin = Mm(margin_mm)

    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE  # type: ignore[assignment]
        section.page_width = width
        section.page_height = height
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    try:
        document.save(docx_path)
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_landscape_save_failed",
            extra={
                "event": "docx_landscape_save_failed",
                "docx_path": str(docx_path),
                "error": str(exc),
            },
        )


def html_file_to_docx(
    html_path: Path,
    output_path: Path,
    *,
    landscape: bool = False,
    body_font_scale: float | None = None,
) -> Optional[Path]:
    """
    Convert an HTML file into a DOCX document using html2docx.

    Parameters
    ----------
    html_path:
        Source HTML document to convert.
    output_path:
        Where the generated DOCX should be written.
    landscape:
        When True, enforces an A4 landscape layout inside the resulting DOCX.
    body_font_scale:
        Optional percentage (0-1] used to downscale the body font size before conversion.

    Returns the output path on success, or None when conversion is unavailable.
    """
    if lxml_html is None or etree is None:
        logger.warning("lxml not available — DOCX export disabled")
        return None

    html_text = html_path.read_text(encoding="utf-8", errors="ignore")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    html_with_inline = _inline_report_styles(html_text)
    html_for_docx = _apply_body_font_scale(_strip_style_blocks(html_with_inline), body_font_scale)

    if html2docx is not None:  # pragma: no branch
        try:
            title_text = _extract_html_title(html_for_docx) or "Report"
            buffer: BytesIO = html2docx(html_for_docx, title_text)  # type: ignore[call-arg]
        except Exception as exc:  # pragma: no cover
            logger.exception(
                "docx_export_html2docx_failed",
                extra={
                    "event": "docx_export_html2docx_failed",
                    "html_path": str(html_path),
                    "docx_path": str(output_path),
                    "error": str(exc),
                },
            )
        else:
            with output_path.open("wb") as handle:
                handle.write(buffer.getvalue())
            if landscape:
                _enforce_landscape_layout(output_path)
            logger.info(
                "docx_export_success",
                extra={
                    "event": "docx_export_success",
                    "html_path": str(html_path),
                    "docx_path": str(output_path),
                    "landscape": landscape,
                    "font_scale": body_font_scale,
                    "strategy": "html2docx",
                },
            )
            return output_path

    structured = _fallback_docx_from_tables(html_text, output_path, body_font_scale=body_font_scale)
    if structured:
        if landscape:
            _enforce_landscape_layout(output_path)
        logger.info(
            "docx_export_success",
            extra={
                "event": "docx_export_success",
                "html_path": str(html_path),
                "docx_path": str(output_path),
                "landscape": landscape,
                "font_scale": body_font_scale,
                "strategy": "structured",
            },
        )
        return output_path

    logger.warning(
        "docx_export_unavailable",
        extra={
            "event": "docx_export_unavailable",
            "reason": "python-docx unavailable" if Document is None else "html2docx not installed",
            "html_path": str(html_path),
        },
    )
    return None


def pdf_file_to_docx(
    pdf_path: Path,
    output_path: Path,
    *,
    start_page: int = 0,
    end_page: int | None = None,
) -> Optional[Path]:
    """
    Convert an already-rendered PDF into DOCX using pdf2docx for near-carbon-copy layout.
    Returns None when conversion is unavailable or fails so callers can fall back to HTML export.
    """
    if Converter is None:  # pragma: no cover
        logger.debug(
            "docx_pdf_convert_skipped",
            extra={
                "event": "docx_pdf_convert_skipped",
                "reason": "pdf2docx unavailable",
                "pdf_path": str(pdf_path),
            },
        )
        return None

    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        logger.warning(
            "docx_pdf_convert_missing_pdf",
            extra={
                "event": "docx_pdf_convert_missing_pdf",
                "pdf_path": str(pdf_path),
            },
        )
        return None

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        converter = Converter(str(pdf_path))
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_pdf_convert_open_failed",
            extra={
                "event": "docx_pdf_convert_open_failed",
                "pdf_path": str(pdf_path),
                "error": str(exc),
            },
        )
        return None

    method = "single-threaded"
    t0 = time.monotonic()
    try:
        # Try multi-processing first (distributes pages across CPU cores,
        # ~4-8x faster for large PDFs).  Falls back to single-threaded if
        # multiprocessing fails (e.g. inside a PyInstaller frozen exe).
        try:
            converter.convert(
                str(output_path),
                start=start_page,
                end=end_page,
                multi_processing=True,
                cpu_count=0,  # 0 = use all available CPUs
            )
            method = "multi-processing"
        except Exception as mp_exc:
            logger.warning(
                "docx_pdf_multiprocessing_failed",
                extra={
                    "event": "docx_pdf_multiprocessing_failed",
                    "error": str(mp_exc),
                    "pdf_path": str(pdf_path),
                },
            )
            with contextlib.suppress(Exception):
                converter.close()
            converter = Converter(str(pdf_path))
            converter.convert(str(output_path), start=start_page, end=end_page)
            method = "single-threaded-fallback"
    except Exception as exc:  # pragma: no cover
        elapsed = time.monotonic() - t0
        logger.warning(
            "docx_pdf_convert_failed",
            extra={
                "event": "docx_pdf_convert_failed",
                "pdf_path": str(pdf_path),
                "docx_path": str(output_path),
                "method": method,
                "elapsed_sec": round(elapsed, 1),
                "error": str(exc),
            },
        )
        return None
    finally:
        with contextlib.suppress(Exception):
            converter.close()

    elapsed = time.monotonic() - t0
    logger.info(
        "docx_pdf_convert_success",
        extra={
            "event": "docx_pdf_convert_success",
            "pdf_path": str(pdf_path),
            "docx_path": str(output_path),
            "method": method,
            "elapsed_sec": round(elapsed, 1),
            "start_page": start_page,
            "end_page": end_page,
        },
    )
    return output_path


# ======================================================================
# Namespace shims for backwards compatibility
# ======================================================================
# Legacy code imports ReportGenerate.fill_and_print and
# ReportGenerateExcel.fill_and_print as module attributes.

import types as _types

ReportGenerate = _types.SimpleNamespace(fill_and_print=fill_and_print)
ReportGenerateExcel = _types.SimpleNamespace(fill_and_print=fill_and_print_excel)


# ── V2-only: Enhanced report pipeline (not in V1) ──
async def run_enhanced_report(
    template_id: str,
    connection_id: str,
    filters: Optional[Dict[str, Any]] = None,
    batch_values: Optional[Dict[str, Any]] = None,
    sse_bridge=None,
) -> Dict[str, Any]:
    """
    Enhanced report generation via LangGraph pipeline.

    This is the V2 entry point — checks feature flags and delegates
    to the appropriate pipeline path with quality evaluation.

    Args:
        template_id: Template to generate report from.
        connection_id: Database connection to use.
        filters: Optional filter conditions.
        batch_values: Optional batch field values.
        sse_bridge: Optional PipelineSSEBridge for SSE streaming.

    Returns:
        Dict with pdf_path, quality_score, metadata, etc.
    """
    from backend.app.services.infra_services import get_v2_config
    from backend.app.services.infra_services import get_event_bus, PipelineStageEvent

    cfg = get_v2_config()
    bus = get_event_bus()
    start_time = time.time()

    # V2 feature flag: disable SSE bridge when streaming is not enabled
    if not cfg.enable_sse_streaming:
        sse_bridge = None

    # Emit pipeline start event
    bus.publish(PipelineStageEvent(
        run_id=f"report-{template_id[:8]}",
        stage="pipeline",
        status="started",
        metadata={"template_id": template_id, "connection_id": connection_id},
    ))

    try:
        from backend.app.services.pipeline_combined import (
            run_report_pipeline,
            _langgraph_available,
        )

        if not _langgraph_available:
            logger.warning("LangGraph not available, cannot run enhanced pipeline")
            return {"success": False, "fallback": True, "reason": "langgraph_unavailable"}

        # Run the pipeline with SSE callback
        def sse_callback(event_type: str, stage: str, **data):
            """Callback for pipeline nodes to emit SSE events."""
            if sse_bridge:
                if event_type == "stage_start":
                    sse_bridge.emit_stage_start(stage, **data)
                elif event_type == "stage_complete":
                    sse_bridge.emit_stage_complete(stage, **data)
                elif event_type == "stage_retry":
                    sse_bridge.emit_stage_retry(stage, data.get("attempt", 1), data.get("reason", ""))
                elif event_type == "error":
                    sse_bridge.emit_error(stage, data.get("error", ""))

            # Also publish to EventBus
            bus.publish(PipelineStageEvent(
                run_id=f"report-{template_id[:8]}",
                stage=stage,
                status=event_type.replace("stage_", ""),
                metadata=data,
            ))

        result = await run_report_pipeline(
            template_id=template_id,
            connection_id=connection_id,
            filters=filters,
            batch_values=batch_values,
        )

        quality_score = None

        # Quality loop wrapper (if enabled)
        if cfg.enable_quality_loop:
            quality_score = await _evaluate_quality(result, sse_bridge)

            if quality_score is not None and quality_score < cfg.quality_threshold:
                # Selective retry (BFI pattern): only re-run failed stages
                if cfg.pipeline_selective_retry:
                    logger.info(
                        "report_quality_retry_selective",
                        extra={
                            "template_id": template_id,
                            "score": quality_score,
                            "threshold": cfg.quality_threshold,
                        },
                    )
                    result = await _selective_retry(
                        result, template_id, connection_id, sse_bridge
                    )
                    quality_score = await _evaluate_quality(result, sse_bridge)

        duration_ms = (time.time() - start_time) * 1000

        # Finish SSE stream
        if sse_bridge:
            sse_bridge.finish(
                success=True,
                summary={
                    "template_id": template_id,
                    "duration_ms": round(duration_ms, 1),
                    "quality_score": quality_score,
                    "phases_completed": result.get("checkpoints", []),
                },
            )

        bus.publish(PipelineStageEvent(
            run_id=f"report-{template_id[:8]}",
            stage="pipeline",
            status="completed",
            duration_ms=duration_ms,
        ))

        return {
            "success": True,
            "result": result,
            "quality_score": quality_score,
            "duration_ms": round(duration_ms, 1),
            "pipeline_mode": "langgraph",
        }

    except Exception as exc:
        logger.exception(
            "report_enhanced_pipeline_failed",
            extra={"template_id": template_id},
        )
        if sse_bridge:
            sse_bridge.emit_error("pipeline", str(exc))
            sse_bridge.finish(success=False)

        bus.publish(PipelineStageEvent(
            run_id=f"report-{template_id[:8]}",
            stage="pipeline",
            status="failed",
            error=str(exc),
        ))

        return {
            "success": False,
            "fallback": True,
            "reason": "pipeline_exception",
            "error": str(exc),
        }


async def _evaluate_quality(result: Dict[str, Any], sse_bridge=None) -> Optional[float]:
    """Evaluate report output quality using QualityEvaluator."""
    try:
        from backend.app.services.quality_service import QualityEvaluator

        evaluator = QualityEvaluator()
        content = result.get("rendered_html", "")
        if not content:
            return None

        score = await evaluator.evaluate(content, content_type="report")
        quality_score = score.overall_score

        if sse_bridge:
            sse_bridge.emit_quality_score(
                quality_score,
                dimensions=score.dimension_scores if hasattr(score, "dimension_scores") else {},
            )

        return quality_score
    except Exception:
        logger.warning("Quality evaluation failed", exc_info=True)
        return None


async def _selective_retry(
    prev_result: Dict[str, Any],
    template_id: str,
    connection_id: str,
    sse_bridge=None,
) -> Dict[str, Any]:
    """
    Selective retry — only re-run execute_queries + render_html.

    BFI pattern: saves ~85% of pipeline work by preserving
    template verification, schema analysis, and mapping extraction.
    """
    try:
        from backend.app.services.pipeline_combined import (
            _execute_queries,
            _render_html,
        )

        if sse_bridge:
            sse_bridge.emit_stage_retry("execute_queries", attempt=2, reason="quality_below_threshold")

        # Re-run only the data stages
        state = dict(prev_result)
        state["retry_count"] = state.get("retry_count", 0) + 1

        state = _execute_queries(state)
        state = _render_html(state)

        return state
    except Exception:
        logger.warning("Selective retry failed, returning original result", exc_info=True)
        return prev_result
