"""Document AI Service.

Main service for document intelligence - parsing, classification, and analysis.
"""
from __future__ import annotations

import asyncio
import base64
import json
import logging
import re
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

from backend.app.schemas import (
    ClassifyRequest,
    ClassifyResponse,
    CompareRequest,
    CompareResponse,
    ComplianceCheckRequest,
    ComplianceCheckResponse,
    ComplianceRule,
    ComplianceViolation,
    ContractAnalyzeRequest,
    ContractAnalyzeResponse,
    DiffType,
    DocumentCategory,
    DocumentDiff,
    EntityExtractRequest,
    EntityExtractResponse,
    EntityType,
    ExtractedEntity,
    InvoiceParseRequest,
    InvoiceParseResponse,
    MultiDocSummarizeRequest,
    MultiDocSummarizeResponse,
    ReceiptScanRequest,
    ReceiptScanResponse,
    ResumeParseRequest,
    ResumeParseResponse,
    RiskLevel,
    SearchResult,
    SemanticSearchRequest,
    SemanticSearchResponse,
    SummarySource,
)
# contract_analyzer defined later in this file
# invoice_parser defined later in this file
# receipt_scanner defined later in this file
# resume_parser defined later in this file

logger = logging.getLogger(__name__)

# Shared text extraction mixin for all parsers
class _TextExtractorMixin:
    """Shared OCR/PDF text extraction used by InvoiceParser, ReceiptScanner, ContractAnalyzer, ResumeParser."""
    _ocr_available: bool = False

    def _check_ocr(self) -> bool:
        try:
            from backend.app.services.llm import get_llm_config
            config = get_llm_config()
            return bool(config.vision_enabled and config.vision_model)
        except Exception:
            return False

    async def _extract_text_from_request(self, request) -> str:
        """Extract text from a request with content or file_path attributes."""
        if getattr(request, 'content', None):
            content = base64.b64decode(request.content)
            return await self._extract_from_bytes(content)
        file_path = getattr(request, 'file_path', None)
        if file_path:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                return await self._extract_from_pdf(path)
            elif suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
                return await self._extract_from_image(path)
            else:
                return path.read_text(encoding="utf-8")
        return ""

    async def _extract_from_bytes(self, content: bytes) -> str:
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                return text
        except Exception:
            pass
        if self._ocr_available:
            try:
                from backend.app.services.infra_services import ocr_extract
                ocr_text = ocr_extract(content)
                if ocr_text:
                    return ocr_text
            except Exception:
                pass
        return content.decode("utf-8", errors="ignore")

    async def _extract_from_pdf(self, path: Path) -> str:
        try:
            import fitz
            doc = fitz.open(str(path))
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            return ""

    async def _extract_from_image(self, path: Path) -> str:
        if not self._ocr_available:
            return ""
        try:
            from backend.app.services.infra_services import ocr_extract_from_file
            return ocr_extract_from_file(path) or ""
        except Exception:
            return ""

class DocAIService:
    """Main Document AI service orchestrating all document intelligence features."""

    def __init__(self) -> None:
        """Initialize the DocAI service."""
        self._nlp_available = self._check_nlp()
        self._embeddings_available = self._check_embeddings()
        self._llm_client = None

    # ---- LLM integration (mirrors SpreadsheetAIService pattern) ----

    def _get_llm_client(self):
        if self._llm_client is None:
            from backend.app.services.llm import get_llm_client
            self._llm_client = get_llm_client()
        return self._llm_client

    async def _call_llm(
        self,
        system_prompt: str,
        user_prompt: str,
        max_tokens: int = 2000,
        description: str = "docai",
    ) -> str:
        client = self._get_llm_client()
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]
        response = await asyncio.to_thread(
            client.complete,
            messages=messages,
            description=description,
            max_tokens=max_tokens,
        )
        content = (
            response.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        )
        return content

    @staticmethod
    def _parse_llm_json(raw: str) -> Any:
        text = raw.strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*\n?", "", text, count=1)
            text = re.sub(r"\n?```\s*$", "", text, count=1)
        return json.loads(text)

    def _check_nlp(self) -> bool:
        """Check if spaCy AND a language model are available."""
        try:
            import spacy
            spacy.load("en_core_web_sm")
            return True
        except Exception:
            return False

    def _check_embeddings(self) -> bool:
        """Check if embedding libraries are available."""
        try:
            from sentence_transformers import SentenceTransformer  # noqa: F401
            return True
        except ImportError:
            return False

    # Document Parsing Methods

    async def parse_invoice(self, request: InvoiceParseRequest) -> InvoiceParseResponse:
        """Parse an invoice document."""
        return await invoice_parser.parse(request)

    async def analyze_contract(
        self, request: ContractAnalyzeRequest
    ) -> ContractAnalyzeResponse:
        """Analyze a contract document."""
        return await contract_analyzer.analyze(request)

    async def parse_resume(self, request: ResumeParseRequest) -> ResumeParseResponse:
        """Parse a resume document."""
        return await resume_parser.parse(request)

    async def scan_receipt(self, request: ReceiptScanRequest) -> ReceiptScanResponse:
        """Scan a receipt document."""
        return await receipt_scanner.scan(request)

    # Document Classification

    async def classify_document(self, request: ClassifyRequest) -> ClassifyResponse:
        """Classify a document by type."""
        start_time = time.time()
        text = await self._extract_text(request.file_path, request.content)

        parser_map = {
            DocumentCategory.INVOICE: ["invoice_parser"],
            DocumentCategory.CONTRACT: ["contract_analyzer"],
            DocumentCategory.RESUME: ["resume_parser"],
            DocumentCategory.RECEIPT: ["receipt_scanner"],
        }

        # Try LLM classification first, fall back to keyword matching
        try:
            result = await self._classify_with_llm(text)
            best_category = result["category"]
            best_score = result["confidence"]
            all_scores = {best_category.value: best_score}
        except Exception:
            logger.debug("LLM classification unavailable, falling back to keywords")
            scores = self._calculate_category_scores(text.lower())
            best_category = max(scores, key=scores.get)
            best_score = scores[best_category]
            all_scores = {k.value: v for k, v in scores.items()}

        suggested_parsers = parser_map.get(best_category, [])
        processing_time_ms = int((time.time() - start_time) * 1000)

        return ClassifyResponse(
            category=best_category,
            confidence=best_score,
            all_scores=all_scores,
            suggested_parsers=suggested_parsers,
            processing_time_ms=processing_time_ms,
        )

    async def _classify_with_llm(self, text: str) -> Dict[str, Any]:
        """Classify document using LLM."""
        categories = ", ".join(c.value for c in DocumentCategory)
        system_prompt = (
            "You are a document classifier. Classify the document into exactly one "
            f"of these categories: {categories}. "
            "Return ONLY a JSON object with keys: category, confidence (0-1), reasoning."
        )
        user_prompt = f"Classify this document:\n\n{text[:4000]}"

        raw = await self._call_llm(system_prompt, user_prompt, max_tokens=500, description="docai_classify")
        parsed = self._parse_llm_json(raw)

        cat_str = str(parsed.get("category", "other")).lower()
        cat_map = {c.value: c for c in DocumentCategory}
        category = cat_map.get(cat_str, DocumentCategory.OTHER)
        confidence = min(max(float(parsed.get("confidence", 0.5)), 0.0), 1.0)

        return {"category": category, "confidence": confidence}

    def _calculate_category_scores(self, text: str) -> Dict[DocumentCategory, float]:
        """Calculate classification scores for each category.

        Uses LLM-based semantic classification with keyword heuristic fallback.
        """
        # Try LLM-based classification first
        try:
            return self._classify_document_agent_sync(text)
        except Exception:
            logger.debug("docai_agent_classify_fallback_to_keywords", exc_info=True)

        return self._calculate_category_scores_keyword(text)

    def _classify_document_agent_sync(self, text: str) -> Dict[DocumentCategory, float]:
        """Use LLM to classify document into categories with confidence scores."""
        client = self._get_llm_client()
        categories = [c.value for c in DocumentCategory]
        prompt = (
            "Classify this document into categories with confidence scores (0.0-1.0).\n"
            f"Categories: {', '.join(categories)}\n\n"
            f"Document excerpt:\n{text[:3000]}\n\n"
            'Return ONLY valid JSON: {"category_scores": {"category_name": score, ...}}'
        )
        resp = client.complete(
            messages=[{"role": "user", "content": prompt}],
            description="docai_agent_classify",
            max_tokens=500,
        )
        content = (
            resp.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        ) if isinstance(resp, dict) else str(resp)
        parsed = self._parse_llm_json(content)
        raw_scores = parsed.get("category_scores", {})

        cat_map = {c.value: c for c in DocumentCategory}
        scores: Dict[DocumentCategory, float] = {}
        for cat_str, score in raw_scores.items():
            cat = cat_map.get(cat_str.lower())
            if cat is not None:
                scores[cat] = min(max(float(score), 0.0), 1.0)

        # Ensure OTHER has a default
        if DocumentCategory.OTHER not in scores:
            scores[DocumentCategory.OTHER] = 0.1

        logger.info("docai_agent_classify_success", extra={"categories": len(scores)})
        return scores

    def _calculate_category_scores_keyword(self, text: str) -> Dict[DocumentCategory, float]:
        """Keyword-based fallback classification."""
        category_keywords = {
            DocumentCategory.INVOICE: [
                "invoice", "bill", "amount due", "payment terms", "due date",
                "invoice number", "line item", "subtotal", "tax"
            ],
            DocumentCategory.CONTRACT: [
                "agreement", "party", "whereas", "hereby", "shall", "term",
                "termination", "governing law", "indemnif", "liability"
            ],
            DocumentCategory.RESUME: [
                "experience", "education", "skills", "employment", "bachelor",
                "master", "university", "objective", "profile", "linkedin"
            ],
            DocumentCategory.RECEIPT: [
                "receipt", "total", "cash", "credit", "thank you", "store",
                "change", "payment", "transaction"
            ],
            DocumentCategory.REPORT: [
                "report", "analysis", "findings", "conclusion", "executive summary",
                "methodology", "results", "recommendations"
            ],
            DocumentCategory.LETTER: [
                "dear", "sincerely", "regards", "yours truly", "to whom it may concern"
            ],
            DocumentCategory.FORM: [
                "please fill", "required field", "signature", "date of birth",
                "applicant", "checkbox", "form"
            ],
            DocumentCategory.PRESENTATION: [
                "slide", "presentation", "agenda", "overview", "key points"
            ],
            DocumentCategory.SPREADSHEET: [
                "total", "sum", "average", "column", "row", "cell"
            ],
        }

        scores: Dict[DocumentCategory, float] = {}
        for category, keywords in category_keywords.items():
            matches = sum(1 for kw in keywords if kw in text)
            scores[category] = min(matches / len(keywords), 1.0)

        scores[DocumentCategory.OTHER] = 0.1
        return scores

    # Entity Extraction

    async def extract_entities(
        self, request: EntityExtractRequest
    ) -> EntityExtractResponse:
        """Extract named entities from text."""
        start_time = time.time()

        # Get text from various sources
        if request.text:
            text = request.text
        else:
            text = await self._extract_text(request.file_path, request.content)

        entities: List[ExtractedEntity] = []
        entity_counts: Dict[str, int] = {}

        # Use spaCy if available, then LLM, then regex
        if self._nlp_available:
            entities = await self._extract_with_spacy(text, request.entity_types)
        else:
            try:
                entities = await self._extract_with_llm(text, request.entity_types)
            except Exception as e:
                logger.warning("LLM entity extraction failed: %s, falling back to regex", e)
                entities = self._extract_with_regex(text, request.entity_types)

        # Count entities by type
        for entity in entities:
            type_name = entity.entity_type.value
            entity_counts[type_name] = entity_counts.get(type_name, 0) + 1

        processing_time_ms = int((time.time() - start_time) * 1000)

        return EntityExtractResponse(
            entities=entities,
            entity_counts=entity_counts,
            processing_time_ms=processing_time_ms,
        )

    async def _extract_with_spacy(
        self, text: str, entity_types: Optional[List[EntityType]]
    ) -> List[ExtractedEntity]:
        """Extract entities using spaCy."""
        import spacy

        try:
            nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Model not installed — delegate to LLM or regex
            try:
                return await self._extract_with_llm(text, entity_types)
            except Exception:
                return self._extract_with_regex(text, entity_types)

        doc = nlp(text)
        entities: List[ExtractedEntity] = []

        # Map spaCy labels to our entity types
        label_map = {
            "PERSON": EntityType.PERSON,
            "ORG": EntityType.ORGANIZATION,
            "GPE": EntityType.LOCATION,
            "LOC": EntityType.LOCATION,
            "DATE": EntityType.DATE,
            "MONEY": EntityType.MONEY,
            "PERCENT": EntityType.PERCENTAGE,
            "PRODUCT": EntityType.PRODUCT,
            "EVENT": EntityType.EVENT,
        }

        for ent in doc.ents:
            entity_type = label_map.get(ent.label_)
            if not entity_type:
                continue

            if entity_types and entity_type not in entity_types:
                continue

            entities.append(ExtractedEntity(
                text=ent.text,
                entity_type=entity_type,
                start=ent.start_char,
                end=ent.end_char,
                confidence=0.85,  # spaCy doesn't provide confidence scores
            ))

        return entities

    def _extract_with_regex(
        self, text: str, entity_types: Optional[List[EntityType]]
    ) -> List[ExtractedEntity]:
        """Extract entities using regex patterns."""
        import re
        entities: List[ExtractedEntity] = []

        patterns = {
            EntityType.EMAIL: r"[\w.+-]+@[\w-]+\.[\w.-]+",
            EntityType.PHONE: r"\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
            EntityType.URL: r"https?://[\w.-]+(?:/[\w./-]*)?",
            EntityType.MONEY: r"\$[\d,]+(?:\.\d{2})?",
            EntityType.PERCENTAGE: r"\d+(?:\.\d+)?\s*(?:%|mg/L|m3/hr|MLD|mV|ppm)\b",
            EntityType.DATE: r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2}|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}(?:,\s*\d{4})?)\b",
            EntityType.PRODUCT: r"\b[A-Z]{2,4}[-_]\d{3,4}(?:[-_]\d+)?\b",
        }

        for entity_type, pattern in patterns.items():
            if entity_types and entity_type not in entity_types:
                continue

            for match in re.finditer(pattern, text):
                entities.append(ExtractedEntity(
                    text=match.group(0),
                    entity_type=entity_type,
                    start=match.start(),
                    end=match.end(),
                    confidence=0.9,
                ))

        return entities

    async def _extract_with_llm(
        self, text: str, entity_types: Optional[List[EntityType]]
    ) -> List[ExtractedEntity]:
        """Extract entities using the LLM client."""
        type_hint = ""
        if entity_types:
            type_hint = f" Focus on these types: {', '.join(t.value for t in entity_types)}."

        system_prompt = (
            "You are a named-entity recognition engine. "
            "Return ONLY a JSON array of objects with keys: "
            "text, entity_type, start, end.  "
            "entity_type must be one of: person, organization, location, date, "
            "money, percentage, email, phone, url, product, event."
        )
        user_prompt = f"Extract all named entities from this text.{type_hint}\n\n{text[:4000]}"

        raw = await self._call_llm(system_prompt, user_prompt, max_tokens=1500, description="docai_entities")
        if not raw:
            return self._extract_with_regex(text, entity_types)

        parsed = self._parse_llm_json(raw)
        if not isinstance(parsed, list):
            parsed = parsed.get("entities", []) if isinstance(parsed, dict) else []

        type_map = {t.value: t for t in EntityType}
        entities: List[ExtractedEntity] = []
        for item in parsed:
            et = type_map.get(str(item.get("entity_type", "")).lower())
            if not et:
                continue
            if entity_types and et not in entity_types:
                continue
            entities.append(ExtractedEntity(
                text=str(item.get("text", "")),
                entity_type=et,
                start=int(item.get("start", 0)),
                end=int(item.get("end", 0)),
                confidence=0.8,
            ))
        return entities

    # Semantic Search

    async def semantic_search(
        self, request: SemanticSearchRequest
    ) -> SemanticSearchResponse:
        """Perform semantic search across documents."""
        start_time = time.time()
        results: List[SearchResult] = []

        if not self._embeddings_available:
            # Fallback to keyword search
            results = await self._keyword_search(
                request.query, request.document_ids, request.top_k
            )
        else:
            results = await self._embedding_search(
                request.query, request.document_ids, request.top_k, request.threshold
            )

        processing_time_ms = int((time.time() - start_time) * 1000)

        return SemanticSearchResponse(
            query=request.query,
            results=results,
            total_results=len(results),
            processing_time_ms=processing_time_ms,
        )

    async def _keyword_search(
        self, query: str, document_ids: Optional[List[str]], top_k: int
    ) -> List[SearchResult]:
        """Simple keyword-based search fallback."""
        # This would integrate with the document library
        # For now, return empty results
        return []

    async def _embedding_search(
        self,
        query: str,
        document_ids: Optional[List[str]],
        top_k: int,
        threshold: float,
    ) -> List[SearchResult]:
        """Search using embeddings."""
        from sentence_transformers import SentenceTransformer

        model = SentenceTransformer("all-MiniLM-L6-v2")
        query_embedding = model.encode(query)

        # This would integrate with the document library and vector store
        # For now, return empty results
        return []

    # Document Comparison

    async def compare_documents(self, request: CompareRequest) -> CompareResponse:
        """Compare two documents for differences."""
        start_time = time.time()

        text_a = await self._extract_text(
            request.document_a_path, request.document_a_content
        )
        text_b = await self._extract_text(
            request.document_b_path, request.document_b_content
        )

        # Try LLM-powered comparison, fall back to Jaccard + difflib
        try:
            result = await self._compare_with_llm(text_a, text_b)
            similarity = result["similarity_score"]
            summary = result["summary"]
        except Exception:
            logger.debug("LLM comparison unavailable, falling back to Jaccard")
            similarity = self._calculate_similarity(text_a, text_b)
            summary = None

        differences = self._find_differences(text_a, text_b)

        if summary is None:
            summary = self._generate_comparison_summary(
                similarity, len(differences), text_a, text_b
            )

        significant = [d.modified_text or d.original_text for d in differences
                       if d.significance == "high"][:5]

        processing_time_ms = int((time.time() - start_time) * 1000)

        return CompareResponse(
            similarity_score=similarity,
            differences=differences,
            summary=summary,
            significant_changes=significant,
            processing_time_ms=processing_time_ms,
        )

    async def _compare_with_llm(self, text_a: str, text_b: str) -> Dict[str, Any]:
        """Compare documents using LLM for semantic similarity."""
        system_prompt = (
            "You are a document comparison engine. Compare the two documents and "
            "return ONLY a JSON object with keys: similarity_score (float 0-1), summary (string)."
        )
        user_prompt = (
            f"Document A:\n{text_a[:2000]}\n\n"
            f"Document B:\n{text_b[:2000]}"
        )

        raw = await self._call_llm(system_prompt, user_prompt, max_tokens=500, description="docai_compare")
        parsed = self._parse_llm_json(raw)

        score = min(max(float(parsed.get("similarity_score", 0.0)), 0.0), 1.0)
        summary = str(parsed.get("summary", ""))

        return {"similarity_score": score, "summary": summary}

    def _calculate_similarity(self, text_a: str, text_b: str) -> float:
        """Calculate text similarity using Jaccard index."""
        words_a = set(text_a.lower().split())
        words_b = set(text_b.lower().split())

        if not words_a or not words_b:
            return 0.0

        intersection = len(words_a & words_b)
        union = len(words_a | words_b)

        return intersection / union if union > 0 else 0.0

    def _find_differences(
        self, text_a: str, text_b: str
    ) -> List[DocumentDiff]:
        """Find differences between two texts."""
        import difflib

        differ = difflib.Differ()
        lines_a = text_a.split("\n")
        lines_b = text_b.split("\n")

        diff = list(differ.compare(lines_a, lines_b))
        differences: List[DocumentDiff] = []

        for i, line in enumerate(diff):
            if line.startswith("- "):
                differences.append(DocumentDiff(
                    diff_type=DiffType.DELETION,
                    original_text=line[2:],
                    significance="medium" if len(line) > 50 else "low",
                ))
            elif line.startswith("+ "):
                differences.append(DocumentDiff(
                    diff_type=DiffType.ADDITION,
                    modified_text=line[2:],
                    significance="medium" if len(line) > 50 else "low",
                ))

        return differences[:50]  # Limit to 50 differences

    def _generate_comparison_summary(
        self,
        similarity: float,
        diff_count: int,
        text_a: str,
        text_b: str,
    ) -> str:
        """Generate a summary of document comparison."""
        if similarity > 0.9:
            level = "nearly identical"
        elif similarity > 0.7:
            level = "substantially similar"
        elif similarity > 0.5:
            level = "moderately similar"
        else:
            level = "significantly different"

        return (
            f"Documents are {level} with {similarity:.0%} similarity. "
            f"Found {diff_count} differences between the two versions."
        )

    # Compliance Checking

    async def check_compliance(
        self, request: ComplianceCheckRequest
    ) -> ComplianceCheckResponse:
        """Check document for compliance with regulations."""
        start_time = time.time()
        text = await self._extract_text(request.file_path, request.content)
        text_lower = text.lower()

        violations: List[ComplianceViolation] = []
        warnings: List[str] = []
        recommendations: List[str] = []

        # Define compliance rules for each regulation
        rules = self._get_compliance_rules(request.regulations)

        for rule in rules:
            violation = self._check_rule(text_lower, rule)
            if violation:
                violations.append(violation)

        # Generate recommendations
        if violations:
            recommendations.append("Review and address all violations before proceeding")
            if any(v.severity == RiskLevel.CRITICAL for v in violations):
                recommendations.append("Critical violations require immediate attention")

        processing_time_ms = int((time.time() - start_time) * 1000)

        return ComplianceCheckResponse(
            compliant=len(violations) == 0,
            violations=violations,
            warnings=warnings,
            recommendations=recommendations,
            checked_regulations=request.regulations,
            processing_time_ms=processing_time_ms,
        )

    def _get_compliance_rules(
        self, regulations: List[str]
    ) -> List[ComplianceRule]:
        """Get compliance rules for specified regulations."""
        all_rules = {
            "GDPR": [
                ComplianceRule(
                    rule_id="gdpr_1",
                    name="Personal Data Processing",
                    description="Must specify lawful basis for processing personal data",
                    regulation="GDPR",
                ),
                ComplianceRule(
                    rule_id="gdpr_2",
                    name="Data Subject Rights",
                    description="Must include provisions for data subject rights",
                    regulation="GDPR",
                ),
                ComplianceRule(
                    rule_id="gdpr_3",
                    name="Data Retention",
                    description="Must specify data retention periods",
                    regulation="GDPR",
                ),
            ],
            "HIPAA": [
                ComplianceRule(
                    rule_id="hipaa_1",
                    name="PHI Protection",
                    description="Must include PHI protection requirements",
                    regulation="HIPAA",
                ),
                ComplianceRule(
                    rule_id="hipaa_2",
                    name="Business Associate",
                    description="Must include Business Associate Agreement for third parties",
                    regulation="HIPAA",
                ),
            ],
            "SOC2": [
                ComplianceRule(
                    rule_id="soc2_1",
                    name="Security Controls",
                    description="Must reference security controls",
                    regulation="SOC2",
                ),
            ],
        }

        rules: List[ComplianceRule] = []
        for reg in regulations:
            rules.extend(all_rules.get(reg.upper(), []))

        return rules

    def _check_rule(
        self, text: str, rule: ComplianceRule
    ) -> Optional[ComplianceViolation]:
        """Check if a rule is violated in the text.

        Uses LLM-based semantic analysis with keyword fallback.
        """
        # Try agent-based compliance check first
        try:
            return self._check_rule_agent(text, rule)
        except Exception:
            logger.debug("compliance_agent_fallback_to_keywords", exc_info=True)

        return self._check_rule_keyword(text, rule)

    def _check_rule_agent(
        self, text: str, rule: ComplianceRule
    ) -> Optional[ComplianceViolation]:
        """Agent-based compliance analysis with semantic understanding."""
        client = self._get_llm_client()
        prompt = (
            f"Check if this document complies with this regulation rule:\n"
            f"Rule: {rule.name}\n"
            f"Description: {rule.description}\n"
            f"Regulation: {rule.regulation}\n\n"
            f"Document excerpt:\n{text[:4000]}\n\n"
            'Return ONLY valid JSON: {"compliant": true/false, "evidence": "quote or reason", "confidence": 0.0-1.0}'
        )
        resp = client.complete(
            messages=[{"role": "user", "content": prompt}],
            description="docai_compliance_check",
            max_tokens=500,
        )
        content = (
            resp.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        ) if isinstance(resp, dict) else str(resp)
        parsed = self._parse_llm_json(content)

        compliant = parsed.get("compliant", True)
        confidence = float(parsed.get("confidence", 0.5))

        # Fall back to keyword check for low-confidence results
        if confidence < 0.5:
            logger.debug("compliance_agent_low_confidence", extra={"rule": rule.rule_id, "confidence": confidence})
            return self._check_rule_keyword(text, rule)

        if not compliant:
            evidence = parsed.get("evidence", f"Missing required language for {rule.name}")
            return ComplianceViolation(
                rule=rule,
                location="Document-wide",
                description=str(evidence),
                severity=RiskLevel.MEDIUM,
                remediation=f"Add language addressing {rule.description}",
            )

        return None

    def _check_rule_keyword(
        self, text: str, rule: ComplianceRule
    ) -> Optional[ComplianceViolation]:
        """Keyword-based compliance checking fallback."""
        required_keywords = {
            "gdpr_1": ["lawful basis", "legitimate interest", "consent"],
            "gdpr_2": ["right to access", "right to erasure", "data subject"],
            "gdpr_3": ["retention", "data retention", "delete"],
            "hipaa_1": ["protected health information", "phi", "health information"],
            "hipaa_2": ["business associate", "baa"],
            "soc2_1": ["security", "controls", "audit"],
        }

        keywords = required_keywords.get(rule.rule_id, [])
        has_keywords = any(kw in text for kw in keywords)

        if not has_keywords:
            return ComplianceViolation(
                rule=rule,
                location="Document-wide",
                description=f"Missing required language for {rule.name}",
                severity=RiskLevel.MEDIUM,
                remediation=f"Add language addressing {rule.description}",
            )

        return None

    # Multi-document Summarization

    async def summarize_multiple(
        self, request: MultiDocSummarizeRequest
    ) -> MultiDocSummarizeResponse:
        """Summarize multiple documents."""
        start_time = time.time()

        # This would integrate with the document library
        # For now, return a placeholder response
        summary = "Multi-document summary feature requires document library integration."
        key_points: List[str] = []
        common_themes: List[str] = []
        sources: List[SummarySource] = []

        processing_time_ms = int((time.time() - start_time) * 1000)

        return MultiDocSummarizeResponse(
            summary=summary,
            key_points=key_points,
            common_themes=common_themes,
            sources=sources,
            document_count=len(request.document_ids),
            processing_time_ms=processing_time_ms,
        )

    # Helper Methods

    async def _extract_text(
        self, file_path: Optional[str], content: Optional[str]
    ) -> str:
        """Extract text from file path or content."""
        if content:
            try:
                decoded = base64.b64decode(content)
                # Try to decode as text
                return decoded.decode("utf-8")
            except Exception:
                # Try PDF extraction
                try:
                    import fitz
                    doc = fitz.open(stream=decoded, filetype="pdf")
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    return text
                except Exception:
                    return ""

        if file_path:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            suffix = path.suffix.lower()
            if suffix == ".pdf":
                try:
                    import fitz
                    doc = fitz.open(str(path))
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                    return text
                except ImportError:
                    return ""
            elif suffix == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(path))
                    return "\n".join([p.text for p in doc.paragraphs])
                except ImportError:
                    return ""
            else:
                return path.read_text(encoding="utf-8")

        return ""

# Singleton instance
docai_service = DocAIService()

"""Invoice Parser Service.

Extracts structured data from invoice documents using OCR and pattern matching.
"""
# from __future__ import annotations (already at top)

import base64
import re
import time
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from backend.app.schemas import (
    ConfidenceLevel,
    InvoiceAddress,
    InvoiceLineItem,
    InvoiceParseRequest,
    InvoiceParseResponse,
)

class InvoiceParser(_TextExtractorMixin):
    """Parser for extracting data from invoice documents."""

    INVOICE_NUMBER_PATTERNS = [
        r"(?:invoice|inv|bill)\s*(?:#|no\.?|number)\s*[:\s]*([A-Z0-9-]+)",
        r"(?:reference|ref)\s*(?:#|no\.?|number)\s*[:\s]*([A-Z0-9-]+)",
    ]

    DATE_PATTERNS = [
        r"(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})",
        r"(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
        r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2},?\s+\d{4})",
    ]

    AMOUNT_PATTERNS = [
        r"(?m)^\s*(?:grand\s+total|total|amount\s+due|balance\s+due)\b\s*[:\s]*(?:[$]|[A-Z]{3})?\s*([\d,]+\.?\d*)",
        r"(?:[$]|[A-Z]{3})\s*([\d,]+\.?\d*)\s*(?:total|due)\b",
    ]

    def __init__(self) -> None:
        self._ocr_available = self._check_ocr()

    async def parse(self, request: InvoiceParseRequest) -> InvoiceParseResponse:
        """Parse an invoice document."""
        start_time = time.time()
        text = await self._extract_text_from_request(request)

        # Extract fields
        invoice_number = self._extract_invoice_number(text)
        invoice_date, due_date = self._extract_dates(text)
        vendor, bill_to = self._extract_addresses(text) if request.extract_addresses else (None, None)
        line_items = self._extract_line_items(text) if request.extract_line_items else []
        subtotal, tax_total, total = self._extract_amounts(text)
        currency = self._detect_currency(text)
        payment_terms = self._extract_payment_terms(text)

        confidence = self._calculate_confidence(
            invoice_number, invoice_date, total, line_items
        )

        processing_time_ms = int((time.time() - start_time) * 1000)

        return InvoiceParseResponse(
            invoice_number=invoice_number,
            invoice_date=invoice_date,
            due_date=due_date,
            vendor=vendor,
            bill_to=bill_to,
            line_items=line_items,
            subtotal=subtotal,
            tax_total=tax_total,
            total=total,
            currency=currency,
            payment_terms=payment_terms,
            raw_text=text[:5000] if text else None,
            confidence_score=confidence,
            processing_time_ms=processing_time_ms,
        )

    # _extract_text, _extract_from_bytes, _extract_from_pdf, _extract_from_image inherited from _TextExtractorMixin

    def _extract_invoice_number(self, text: str) -> Optional[str]:
        """Extract invoice number from text."""
        text_lower = text.lower()
        for pattern in self.INVOICE_NUMBER_PATTERNS:
            match = re.search(pattern, text_lower, re.IGNORECASE)
            if match:
                return match.group(1).upper()
        return None

    def _extract_dates(self, text: str) -> Tuple[Optional[datetime], Optional[datetime]]:
        """Extract invoice date and due date from text."""
        dates: List[datetime] = []

        for pattern in self.DATE_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                try:
                    date_str = match.group(1)
                    parsed = self._parse_date(date_str)
                    if parsed:
                        dates.append(parsed)
                except Exception:
                    continue

        # Sort dates - typically first is invoice date, later is due date
        dates = sorted(set(dates))

        invoice_date = dates[0] if dates else None
        due_date = dates[-1] if len(dates) > 1 else None

        return invoice_date, due_date

    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Parse a date string into datetime."""
        formats = [
            "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d",
            "%m-%d-%Y", "%d-%m-%Y", "%Y-%m-%d",
            "%B %d, %Y", "%b %d, %Y",
            "%m/%d/%y", "%d/%m/%y",
        ]

        for fmt in formats:
            try:
                return datetime.strptime(date_str.strip(), fmt)
            except ValueError:
                continue
        return None

    def _extract_addresses(
        self, text: str
    ) -> Tuple[Optional[InvoiceAddress], Optional[InvoiceAddress]]:
        """Extract vendor and billing addresses."""
        # Simple heuristic: look for address-like blocks
        vendor = None
        bill_to = None

        # Look for "From" or "Vendor" section
        from_match = re.search(
            r"(?:from|vendor|seller|company)[:\s]*\n(.+?)(?:\n\n|bill\s*to|ship\s*to)",
            text, re.IGNORECASE | re.DOTALL
        )
        if from_match:
            vendor = self._parse_address_block(from_match.group(1))

        # Look for "Bill To" section
        bill_match = re.search(
            r"(?:bill\s*to|customer|client)[:\s]*\n(.+?)(?:\n\n|ship\s*to|items|description)",
            text, re.IGNORECASE | re.DOTALL
        )
        if bill_match:
            bill_to = self._parse_address_block(bill_match.group(1))

        return vendor, bill_to

    def _parse_address_block(self, text: str) -> InvoiceAddress:
        """Parse an address text block."""
        lines = [line.strip() for line in text.strip().split("\n") if line.strip()]

        name = lines[0] if lines else None
        street = lines[1] if len(lines) > 1 else None

        # Extract email
        email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
        email = email_match.group(0) if email_match else None

        # Extract phone
        phone_match = re.search(r"[\d\s\-\+\(\)]{10,}", text)
        phone = phone_match.group(0).strip() if phone_match else None

        return InvoiceAddress(
            name=name,
            street=street,
            email=email,
            phone=phone,
        )

    def _extract_line_items(self, text: str) -> List[InvoiceLineItem]:
        """Extract line items from invoice."""
        items: List[InvoiceLineItem] = []

        # Look for tabular data pattern
        # Description | Qty | Price | Amount
        line_pattern = r"(.+?)\s+(\d+(?:\.\d+)?)\s+\$?([\d,]+\.?\d*)\s+\$?([\d,]+\.?\d*)"

        for match in re.finditer(line_pattern, text):
            try:
                description = match.group(1).strip()
                quantity = float(match.group(2))
                unit_price = float(match.group(3).replace(",", ""))
                amount = float(match.group(4).replace(",", ""))

                if description and len(description) > 2:
                    items.append(InvoiceLineItem(
                        description=description,
                        quantity=quantity,
                        unit_price=unit_price,
                        amount=amount,
                        confidence=ConfidenceLevel.MEDIUM,
                    ))
            except (ValueError, IndexError):
                continue

        return items

    def _extract_amounts(self, text: str) -> Tuple[Optional[float], Optional[float], Optional[float]]:
        """Extract subtotal, tax, and total amounts."""
        subtotal = None
        tax = None
        total = None

        # Subtotal
        subtotal_match = re.search(
            r"(?:sub\s*total|subtotal)\s*[:\s]*[$€£¥]?\s*([\d,]+\.?\d*)",
            text, re.IGNORECASE
        )
        if subtotal_match:
            subtotal = float(subtotal_match.group(1).replace(",", ""))

        # Tax
        tax_match = re.search(
            r"(?:tax|vat|gst)\s*[:\s]*[$€£¥]?\s*([\d,]+\.?\d*)",
            text, re.IGNORECASE
        )
        if tax_match:
            tax = float(tax_match.group(1).replace(",", ""))

        # Total
        for pattern in self.AMOUNT_PATTERNS:
            total_match = re.search(pattern, text, re.IGNORECASE)
            if total_match:
                total = float(total_match.group(1).replace(",", ""))
                break

        return subtotal, tax, total

    def _detect_currency(self, text: str) -> str:
        """Detect currency from text."""
        text_upper = text.upper()
        if "$" in text or re.search(r"\bUSD\b", text_upper):
            return "USD"
        if re.search(r"\bEUR\b", text_upper):
            return "EUR"
        if re.search(r"\bGBP\b", text_upper):
            return "GBP"
        if re.search(r"\b(?:JPY|CNY)\b", text_upper):
            return "JPY"
        return "USD"
    def _extract_payment_terms(self, text: str) -> Optional[str]:
        """Extract payment terms from invoice."""
        terms_match = re.search(
            r"(?:payment\s*terms?|terms?)[:\s]*(.+?)(?:\n|$)",
            text, re.IGNORECASE
        )
        if terms_match:
            return terms_match.group(1).strip()[:100]

        # Look for "Net XX" patterns
        net_match = re.search(r"(net\s*\d+)", text, re.IGNORECASE)
        if net_match:
            return net_match.group(1)

        return None

    def _calculate_confidence(
        self,
        invoice_number: Optional[str],
        invoice_date: Optional[datetime],
        total: Optional[float],
        line_items: List[InvoiceLineItem],
    ) -> float:
        """Calculate overall confidence score."""
        score = 0.0

        if invoice_number:
            score += 0.25
        if invoice_date:
            score += 0.25
        if total is not None:
            score += 0.25
        if line_items:
            score += 0.25

        return score

# Singleton instance
invoice_parser = InvoiceParser()

"""Receipt Scanner Service.

Extracts structured data from receipt images and documents.
"""
# from __future__ import annotations (already at top)

import base64
import re
import time
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from backend.app.schemas import (
    ReceiptItem,
    ReceiptScanRequest,
    ReceiptScanResponse,
)

class ReceiptScanner(_TextExtractorMixin):
    """Scanner for extracting data from receipt documents."""

    CATEGORY_KEYWORDS = {
        "Groceries": ["grocery", "supermarket", "food", "produce", "dairy", "meat"],
        "Restaurant": ["restaurant", "cafe", "diner", "bar", "grill", "pizza", "sushi"],
        "Gas Station": ["gas", "fuel", "petroleum", "shell", "exxon", "chevron", "bp"],
        "Retail": ["store", "shop", "mart", "outlet", "boutique"],
        "Pharmacy": ["pharmacy", "drug", "cvs", "walgreens", "rite aid"],
        "Electronics": ["electronics", "best buy", "apple", "tech"],
        "Office Supplies": ["office", "staples", "depot"],
        "Hardware": ["hardware", "home depot", "lowes", "ace"],
    }

    # Item category keywords
    ITEM_CATEGORIES = {
        "Food": ["milk", "bread", "eggs", "cheese", "meat", "chicken", "beef", "fish",
                 "fruit", "vegetable", "rice", "pasta", "cereal"],
        "Beverage": ["water", "soda", "juice", "coffee", "tea", "beer", "wine"],
        "Household": ["paper", "towel", "soap", "detergent", "cleaner", "tissue"],
        "Personal Care": ["shampoo", "toothpaste", "deodorant", "lotion"],
        "Snacks": ["chips", "candy", "cookies", "chocolate", "snack"],
    }

    def __init__(self) -> None:
        self._ocr_available = self._check_ocr()

    async def scan(self, request: ReceiptScanRequest) -> ReceiptScanResponse:
        """Scan a receipt document."""
        start_time = time.time()
        text = await self._extract_text_from_request(request)

        # Extract merchant information
        merchant_name = self._extract_merchant_name(text)
        merchant_address = self._extract_merchant_address(text)
        merchant_phone = self._extract_merchant_phone(text)

        # Extract date and time
        date, time_str = self._extract_datetime(text)

        # Extract line items
        items = self._extract_items(text, request.categorize_items)

        # Extract amounts
        subtotal = self._extract_subtotal(text)
        tax = self._extract_tax(text)
        tip = self._extract_tip(text)
        total = self._extract_total(text, subtotal, tax, tip, items)

        # Extract payment info
        payment_method, card_last_four = self._extract_payment_info(text)

        # Detect currency
        currency = self._detect_currency(text)

        # Categorize receipt
        category = self._categorize_receipt(text, merchant_name)

        confidence = self._calculate_confidence(
            merchant_name, date, total, items
        )

        processing_time_ms = int((time.time() - start_time) * 1000)

        return ReceiptScanResponse(
            merchant_name=merchant_name,
            merchant_address=merchant_address,
            merchant_phone=merchant_phone,
            date=date,
            time=time_str,
            items=items,
            subtotal=subtotal,
            tax=tax,
            tip=tip,
            total=total,
            payment_method=payment_method,
            card_last_four=card_last_four,
            currency=currency,
            category=category,
            raw_text=text[:3000] if text else None,
            confidence_score=confidence,
            processing_time_ms=processing_time_ms,
        )

    # Extraction methods inherited from _TextExtractorMixin

    def _extract_merchant_name(self, text: str) -> Optional[str]:
        """Extract merchant name from receipt."""
        # Usually the merchant name is in the first few lines
        lines = text.strip().split("\n")[:5]

        for line in lines:
            line = line.strip()
            # Skip lines that look like addresses or phone numbers
            if "@" in line or re.search(r"\d{3}[-.\s]?\d{3}", line):
                continue
            if re.search(r"\d+\s+\w+\s+(?:st|ave|rd|blvd|dr|ln)", line, re.IGNORECASE):
                continue

            # Take first substantial line
            if 3 < len(line) < 50:
                return line

        return None

    def _extract_merchant_address(self, text: str) -> Optional[str]:
        """Extract merchant address from receipt."""
        # Look for address pattern
        address_match = re.search(
            r"(\d+\s+[\w\s]+(?:st|street|ave|avenue|rd|road|blvd|boulevard|dr|drive|ln|lane)[.,]?\s*"
            r"[\w\s]*,?\s*[A-Z]{2}\s*\d{5})",
            text, re.IGNORECASE
        )
        if address_match:
            return address_match.group(1)

        return None

    def _extract_merchant_phone(self, text: str) -> Optional[str]:
        """Extract merchant phone from receipt."""
        phone_match = re.search(
            r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
            text
        )
        if phone_match:
            return phone_match.group(0)

        return None

    def _extract_datetime(
        self, text: str
    ) -> Tuple[Optional[datetime], Optional[str]]:
        """Extract date and time from receipt."""
        date_obj = None
        time_str = None

        # Date patterns
        date_patterns = [
            (r"(\d{1,2})/(\d{1,2})/(\d{2,4})", "%m/%d/%Y"),
            (r"(\d{1,2})-(\d{1,2})-(\d{2,4})", "%m-%d-%Y"),
            (r"(\w{3})\s+(\d{1,2}),?\s+(\d{4})", "%b %d %Y"),
        ]

        for pattern, fmt in date_patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    date_str = match.group(0)
                    # Handle 2-digit years
                    if len(date_str.split("/")[-1]) == 2:
                        fmt = fmt.replace("%Y", "%y")
                    date_obj = datetime.strptime(date_str.replace(",", ""), fmt)
                    break
                except ValueError:
                    continue

        # Time pattern
        time_match = re.search(r"(\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?)", text, re.IGNORECASE)
        if time_match:
            time_str = time_match.group(1)

        return date_obj, time_str

    def _extract_items(
        self, text: str, categorize: bool
    ) -> List[ReceiptItem]:
        """Extract line items from receipt."""
        items: List[ReceiptItem] = []

        # Look for patterns like:
        # ITEM NAME          $9.99
        # ITEM NAME    2@$1.50    $3.00
        item_patterns = [
            # Name followed by price
            r"([A-Z][A-Z\s]{2,30})\s+\$?([\d]+\.[\d]{2})\s*$",
            # Name with quantity and price
            r"([A-Z][A-Z\s]{2,30})\s+(\d+)\s*@\s*\$?([\d]+\.[\d]{2})\s+\$?([\d]+\.[\d]{2})",
            # Simple name and price on same line
            r"^([^$\d]{3,30})\s+\$?([\d]+\.[\d]{2})",
        ]

        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if not line or len(line) < 5:
                continue

            # Skip header/footer lines
            if any(skip in line.lower() for skip in
                   ["subtotal", "total", "tax", "cash", "credit", "change", "thank you"]):
                continue

            for pattern in item_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    groups = match.groups()

                    if len(groups) >= 4:
                        # With quantity
                        name = groups[0].strip()
                        quantity = float(groups[1])
                        unit_price = float(groups[2])
                        total_price = float(groups[3])
                    elif len(groups) >= 2:
                        # Without quantity
                        name = groups[0].strip()
                        quantity = 1.0
                        total_price = float(groups[1])
                        unit_price = total_price
                    else:
                        continue

                    # Categorize if requested
                    category = None
                    if categorize:
                        category = self._categorize_item(name)

                    items.append(ReceiptItem(
                        name=name,
                        quantity=quantity,
                        unit_price=unit_price,
                        total_price=total_price,
                        category=category,
                    ))
                    break

        return items

    def _categorize_item(self, name: str) -> Optional[str]:
        """Categorize an item based on its name."""
        name_lower = name.lower()

        for category, keywords in self.ITEM_CATEGORIES.items():
            for keyword in keywords:
                if keyword in name_lower:
                    return category

        return None

    def _extract_subtotal(self, text: str) -> Optional[float]:
        """Extract subtotal from receipt."""
        subtotal_match = re.search(
            r"(?:sub\s*total|subtotal)\s*:?\s*\$?([\d,]+\.[\d]{2})",
            text, re.IGNORECASE
        )
        if subtotal_match:
            return float(subtotal_match.group(1).replace(",", ""))

        return None

    def _extract_tax(self, text: str) -> Optional[float]:
        """Extract tax from receipt."""
        tax_match = re.search(
            r"(?:tax|sales\s+tax|vat)\s*:?\s*\$?([\d,]+\.[\d]{2})",
            text, re.IGNORECASE
        )
        if tax_match:
            return float(tax_match.group(1).replace(",", ""))

        return None

    def _extract_tip(self, text: str) -> Optional[float]:
        """Extract tip from receipt."""
        tip_match = re.search(
            r"(?:tip|gratuity)\s*:?\s*\$?([\d,]+\.[\d]{2})",
            text, re.IGNORECASE
        )
        if tip_match:
            return float(tip_match.group(1).replace(",", ""))

        return None

    def _extract_total(
        self,
        text: str,
        subtotal: Optional[float],
        tax: Optional[float],
        tip: Optional[float],
        items: List[ReceiptItem],
    ) -> float:
        """Extract or calculate total from receipt."""
        # Try to find total directly
        total_patterns = [
            r"(?m)^\s*(?:grand\s+)?total\b\s*:?\s*\$?([\d,]+\.[\d]{2})",
            r"(?m)^\s*amount\s+due\b\s*:?\s*\$?([\d,]+\.[\d]{2})",
            r"(?m)^\s*balance\s+due\b\s*:?\s*\$?([\d,]+\.[\d]{2})",
        ]

        for pattern in total_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return float(match.group(1).replace(",", ""))

        # Calculate from components
        if subtotal:
            total = subtotal
            if tax:
                total += tax
            if tip:
                total += tip
            return total

        # Sum items
        if items:
            return sum(item.total_price for item in items)

        return 0.0

    def _extract_payment_info(
        self, text: str
    ) -> Tuple[Optional[str], Optional[str]]:
        """Extract payment method and card info."""
        payment_method = None
        card_last_four = None

        # Payment method
        if re.search(r"(?:visa|mastercard|amex|american\s+express|discover)", text, re.IGNORECASE):
            payment_method = "Credit Card"
        elif re.search(r"(?:debit|debit\s+card)", text, re.IGNORECASE):
            payment_method = "Debit Card"
        elif re.search(r"(?:cash|cash\s+payment)", text, re.IGNORECASE):
            payment_method = "Cash"
        elif re.search(r"(?:apple\s+pay|google\s+pay|paypal)", text, re.IGNORECASE):
            payment_method = "Digital Payment"

        # Card last four digits
        card_match = re.search(r"(?:\*+|x+)(\d{4})", text, re.IGNORECASE)
        if card_match:
            card_last_four = card_match.group(1)

        return payment_method, card_last_four

    def _detect_currency(self, text: str) -> str:
        """Detect currency from receipt."""
        text_upper = text.upper()
        if "$" in text or re.search(r"\bUSD\b", text_upper):
            return "USD"
        if re.search(r"\bEUR\b", text_upper):
            return "EUR"
        if re.search(r"\bGBP\b", text_upper):
            return "GBP"
        if re.search(r"\bJPY\b", text_upper):
            return "JPY"
        if re.search(r"\bCAD\b", text_upper) or "C$" in text:
            return "CAD"
        return "USD"
    def _categorize_receipt(
        self, text: str, merchant_name: Optional[str]
    ) -> Optional[str]:
        """Categorize the receipt based on content."""
        text_lower = text.lower()
        merchant_lower = (merchant_name or "").lower()

        for category, keywords in self.CATEGORY_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower or keyword in merchant_lower:
                    return category

        return None

    def _calculate_confidence(
        self,
        merchant_name: Optional[str],
        date: Optional[datetime],
        total: float,
        items: List[ReceiptItem],
    ) -> float:
        """Calculate overall confidence score."""
        score = 0.0

        if merchant_name:
            score += 0.25
        if date:
            score += 0.25
        if total > 0:
            score += 0.25
        if items:
            score += 0.25

        return score

# Singleton instance
receipt_scanner = ReceiptScanner()

"""Contract Analyzer Service.

Analyzes contract documents for clauses, risks, and obligations.
"""
# from __future__ import annotations (already at top)

import base64
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from backend.app.schemas import (
    ConfidenceLevel,
    ContractAnalyzeRequest,
    ContractAnalyzeResponse,
    ContractClause,
    ContractClauseType,
    ContractObligation,
    ContractParty,
    RiskLevel,
)

class ContractAnalyzer(_TextExtractorMixin):
    """Analyzer for extracting data from contract documents."""

    # Clause detection patterns
    CLAUSE_PATTERNS: Dict[ContractClauseType, List[str]] = {
        ContractClauseType.TERMINATION: [
            r"(?:termination|terminate|cancel)",
            r"(?:end\s+of\s+agreement|expir)",
        ],
        ContractClauseType.INDEMNIFICATION: [
            r"(?:indemnif|hold\s+harmless|defend\s+and\s+indemnify)",
        ],
        ContractClauseType.LIMITATION_OF_LIABILITY: [
            r"(?:limitation\s+of\s+liability|limit\s+liability|maximum\s+liability)",
            r"(?:aggregate\s+liability|cap\s+on\s+damages)",
        ],
        ContractClauseType.CONFIDENTIALITY: [
            r"(?:confidential|non-disclosure|nda|proprietary\s+information)",
        ],
        ContractClauseType.INTELLECTUAL_PROPERTY: [
            r"(?:intellectual\s+property|ip\s+rights|patent|trademark|copyright)",
        ],
        ContractClauseType.FORCE_MAJEURE: [
            r"(?:force\s+majeure|act\s+of\s+god|unforeseeable\s+circumstances)",
        ],
        ContractClauseType.GOVERNING_LAW: [
            r"(?:governing\s+law|applicable\s+law|jurisdiction)",
        ],
        ContractClauseType.DISPUTE_RESOLUTION: [
            r"(?:dispute\s+resolution|arbitration|mediation)",
        ],
        ContractClauseType.ASSIGNMENT: [
            r"(?:assignment|transfer\s+of\s+rights|assignable)",
        ],
        ContractClauseType.PAYMENT: [
            r"(?:payment\s+terms|payment\s+schedule|fee|compensation|invoice)",
        ],
        ContractClauseType.WARRANTY: [
            r"(?:warrant|representation|guarantee)",
        ],
        ContractClauseType.INSURANCE: [
            r"(?:insurance|coverage|policy)",
        ],
    }

    # Risk indicators
    RISK_INDICATORS = {
        RiskLevel.CRITICAL: [
            r"(?:unlimited\s+liability|personal\s+guarantee|waive\s+all\s+rights)",
            r"(?:automatic\s+renewal|perpetual\s+license)",
        ],
        RiskLevel.HIGH: [
            r"(?:exclusive\s+rights|non-compete|non-solicitation)",
            r"(?:termination\s+for\s+convenience|terminate\s+at\s+any\s+time)",
        ],
        RiskLevel.MEDIUM: [
            r"(?:change\s+in\s+control|material\s+breach|cure\s+period)",
        ],
    }

    def __init__(self) -> None:
        self._ocr_available = self._check_ocr()
        self._nlp_available = self._check_nlp()

    def _check_nlp(self) -> bool:
        try:
            import spacy  # noqa: F401
            return True
        except Exception:
            return False

    async def analyze(self, request: ContractAnalyzeRequest) -> ContractAnalyzeResponse:
        """Analyze a contract document."""
        start_time = time.time()
        text = await self._extract_text_from_request(request)

        # Extract contract metadata
        title = self._extract_title(text)
        contract_type = self._detect_contract_type(text)
        effective_date, expiration_date = self._extract_dates(text)
        parties = self._extract_parties(text)

        # Extract clauses
        clauses = self._extract_clauses(text, request.analyze_risks)

        # Extract obligations
        obligations = (
            self._extract_obligations(text, parties)
            if request.extract_obligations
            else []
        )

        # Extract key dates
        key_dates = self._extract_key_dates(text)

        # Extract monetary values
        total_value, currency = self._extract_value(text)

        # Calculate risk summary
        risk_summary, overall_risk = self._assess_risks(clauses, text)

        # Generate recommendations
        recommendations = self._generate_recommendations(clauses, risk_summary)

        # Generate summary
        summary = self._generate_summary(
            title, contract_type, parties, effective_date, total_value
        )

        confidence = self._calculate_confidence(
            title, parties, clauses, effective_date
        )

        processing_time_ms = int((time.time() - start_time) * 1000)

        return ContractAnalyzeResponse(
            title=title,
            contract_type=contract_type,
            effective_date=effective_date,
            expiration_date=expiration_date,
            parties=parties,
            clauses=clauses,
            obligations=obligations,
            key_dates=key_dates,
            total_value=total_value,
            currency=currency,
            risk_summary=risk_summary,
            overall_risk_level=overall_risk,
            recommendations=recommendations,
            summary=summary,
            confidence_score=confidence,
            processing_time_ms=processing_time_ms,
        )

    async def _extract_text_from_request(self, request) -> str:
        """Override: also supports .docx for contracts."""
        if getattr(request, 'file_path', None):
            path = Path(request.file_path)
            if path.suffix.lower() == ".docx":
                return await self._extract_from_docx(path)
        return await super()._extract_text_from_request(request)

    async def _extract_from_docx(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return ""

    def _extract_title(self, text: str) -> Optional[str]:
        """Extract contract title."""
        # Look for title in first few lines
        lines = text.strip().split("\n")[:10]

        for line in lines:
            line = line.strip()
            # Look for "Agreement" or "Contract" in line
            if re.search(r"(?:agreement|contract)", line, re.IGNORECASE):
                # Clean up the title
                title = re.sub(r"^\d+[\.\)]\s*", "", line)
                title = title.strip()
                if 10 < len(title) < 200:
                    return title

        return None

    def _detect_contract_type(self, text: str) -> Optional[str]:
        """Detect the type of contract."""
        text_lower = text.lower()

        contract_types = {
            "Employment Agreement": ["employment", "employee", "employer", "hire"],
            "Non-Disclosure Agreement": ["non-disclosure", "nda", "confidential"],
            "Service Agreement": ["service agreement", "services", "contractor"],
            "License Agreement": ["license", "licensing", "licensed"],
            "Sales Agreement": ["sale", "purchase", "buyer", "seller"],
            "Lease Agreement": ["lease", "tenant", "landlord", "rental"],
            "Partnership Agreement": ["partnership", "partners", "joint venture"],
            "Consulting Agreement": ["consulting", "consultant"],
            "Subscription Agreement": ["subscription", "subscriber"],
            "Master Service Agreement": ["master service", "msa"],
        }

        for contract_type, keywords in contract_types.items():
            if any(keyword in text_lower for keyword in keywords):
                return contract_type

        return "General Agreement"

    def _extract_dates(
        self, text: str
    ) -> Tuple[Optional[datetime], Optional[datetime]]:
        """Extract effective and expiration dates."""
        effective_date = None
        expiration_date = None

        # Effective date patterns
        effective_match = re.search(
            r"(?:effective\s+(?:date|as\s+of)|dated\s+as\s+of|commenc)"
            r"[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\w+\s+\d{1,2},?\s+\d{4})",
            text, re.IGNORECASE
        )
        if effective_match:
            effective_date = self._parse_date(effective_match.group(1))

        # Expiration date patterns
        expiration_match = re.search(
            r"(?:expir|terminat|end\s+date|valid\s+until)"
            r"[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\w+\s+\d{1,2},?\s+\d{4})",
            text, re.IGNORECASE
        )
        if expiration_match:
            expiration_date = self._parse_date(expiration_match.group(1))

        return effective_date, expiration_date

    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Parse a date string into datetime."""
        formats = [
            "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d",
            "%m-%d-%Y", "%d-%m-%Y", "%Y-%m-%d",
            "%B %d, %Y", "%b %d, %Y", "%B %d %Y",
        ]

        for fmt in formats:
            try:
                return datetime.strptime(date_str.strip(), fmt)
            except ValueError:
                continue
        return None

    def _extract_parties(self, text: str) -> List[ContractParty]:
        """Extract parties to the contract."""
        parties: List[ContractParty] = []

        # Look for party definitions
        party_pattern = r'"([^"]+)"(?:\s*,?\s*a\s+)?(?:(\w+)(?:\s+organized|\s+incorporated)?)?'

        # Also look for "between X and Y" pattern
        between_match = re.search(
            r"between\s+(.+?)\s+(?:and|&)\s+(.+?)(?:\.|,|\n)",
            text[:2000], re.IGNORECASE
        )

        if between_match:
            party1 = between_match.group(1).strip()
            party2 = between_match.group(2).strip()

            # Clean party names
            party1 = re.sub(r'^["\'](.*)["\']$', r"\1", party1)
            party2 = re.sub(r'^["\'](.*)["\']$', r"\1", party2)

            parties.append(ContractParty(name=party1, role="Party A"))
            parties.append(ContractParty(name=party2, role="Party B"))

        return parties

    def _extract_clauses(
        self, text: str, analyze_risks: bool
    ) -> List[ContractClause]:
        """Extract and analyze clauses from the contract."""
        clauses: List[ContractClause] = []

        # Split text into sections by numbered headers
        sections = re.split(r"\n\s*(?:\d+[\.\)]\s*|\(?[a-z]\)\s*)", text)

        for i, section in enumerate(sections):
            if len(section) < 50:
                continue

            # Detect clause type
            clause_type = self._classify_clause(section)
            if clause_type == ContractClauseType.OTHER:
                continue

            # Extract title
            lines = section.strip().split("\n")
            title = lines[0][:100] if lines else f"Section {i}"

            # Analyze risk if requested
            risk_level = RiskLevel.INFORMATIONAL
            risk_explanation = None
            suggestions: List[str] = []

            if analyze_risks:
                risk_level, risk_explanation = self._analyze_clause_risk(section)
                if risk_level in [RiskLevel.HIGH, RiskLevel.CRITICAL]:
                    suggestions = self._generate_clause_suggestions(clause_type, section)

            clauses.append(ContractClause(
                clause_type=clause_type,
                title=title,
                text=section[:2000],
                risk_level=risk_level,
                risk_explanation=risk_explanation,
                suggestions=suggestions,
                confidence=ConfidenceLevel.MEDIUM,
            ))

        return clauses

    def _classify_clause(self, text: str) -> ContractClauseType:
        """Classify a clause based on its content."""
        text_lower = text.lower()

        for clause_type, patterns in self.CLAUSE_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return clause_type

        return ContractClauseType.OTHER

    def _analyze_clause_risk(
        self, text: str
    ) -> Tuple[RiskLevel, Optional[str]]:
        """Analyze risk level of a clause."""
        text_lower = text.lower()

        for risk_level, patterns in self.RISK_INDICATORS.items():
            for pattern in patterns:
                match = re.search(pattern, text_lower)
                if match:
                    return risk_level, f"Contains potentially risky language: '{match.group(0)}'"

        return RiskLevel.INFORMATIONAL, None

    def _generate_clause_suggestions(
        self, clause_type: ContractClauseType, text: str
    ) -> List[str]:
        """Generate improvement suggestions for a clause."""
        suggestions: List[str] = []

        if clause_type == ContractClauseType.LIMITATION_OF_LIABILITY:
            if "unlimited" in text.lower():
                suggestions.append("Consider adding a liability cap")
            suggestions.append("Review whether the limitation scope is appropriate")

        elif clause_type == ContractClauseType.TERMINATION:
            if "convenience" in text.lower():
                suggestions.append("Consider adding notice period requirements")
            suggestions.append("Ensure termination rights are mutual")

        elif clause_type == ContractClauseType.INDEMNIFICATION:
            suggestions.append("Review scope of indemnification obligations")
            suggestions.append("Consider adding carve-outs for negligence")

        return suggestions

    def _extract_obligations(
        self, text: str, parties: List[ContractParty]
    ) -> List[ContractObligation]:
        """Extract obligations from the contract."""
        obligations: List[ContractObligation] = []

        # Look for "shall" and "must" patterns
        obligation_pattern = r"([A-Z][^.]*?)\s+(?:shall|must|agrees?\s+to|is\s+obligated\s+to)\s+([^.]+)"

        for match in re.finditer(obligation_pattern, text):
            subject = match.group(1).strip()
            action = match.group(2).strip()

            # Try to match subject to a party
            party_name = subject
            for party in parties:
                if party.name.lower() in subject.lower():
                    party_name = party.name
                    break

            if len(action) > 10:
                obligations.append(ContractObligation(
                    party=party_name,
                    obligation=action[:500],
                ))

        return obligations[:20]  # Limit to 20 obligations

    def _extract_key_dates(self, text: str) -> Dict[str, datetime]:
        """Extract key dates from the contract."""
        key_dates: Dict[str, datetime] = {}

        date_labels = [
            "effective date",
            "start date",
            "end date",
            "expiration date",
            "termination date",
            "renewal date",
            "payment date",
            "delivery date",
        ]

        for label in date_labels:
            pattern = rf"{label}[:\s]*(\d{{1,2}}[-/]\d{{1,2}}[-/]\d{{2,4}})"
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                parsed = self._parse_date(match.group(1))
                if parsed:
                    key_dates[label.replace(" ", "_")] = parsed

        return key_dates

    def _extract_value(self, text: str) -> Tuple[Optional[float], Optional[str]]:
        """Extract monetary value from contract."""
        # Look for total value patterns
        value_pattern = r"(?:total\s+)?(?:value|amount|price|fee|compensation)[:\s]*[$€£]?\s*([\d,]+(?:\.\d{2})?)"
        match = re.search(value_pattern, text, re.IGNORECASE)

        if match:
            value = float(match.group(1).replace(",", ""))
            currency = "USD"
            if "€" in text[:match.start() + 50]:
                currency = "EUR"
            elif "£" in text[:match.start() + 50]:
                currency = "GBP"
            return value, currency

        return None, None

    def _assess_risks(
        self, clauses: List[ContractClause], text: str
    ) -> Tuple[Dict[str, int], RiskLevel]:
        """Assess overall contract risks."""
        risk_counts = {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
            "informational": 0,
        }

        for clause in clauses:
            risk_counts[clause.risk_level.value] += 1

        # Determine overall risk
        if risk_counts["critical"] > 0:
            overall = RiskLevel.CRITICAL
        elif risk_counts["high"] > 2:
            overall = RiskLevel.HIGH
        elif risk_counts["high"] > 0 or risk_counts["medium"] > 3:
            overall = RiskLevel.MEDIUM
        else:
            overall = RiskLevel.LOW

        return risk_counts, overall

    def _generate_recommendations(
        self, clauses: List[ContractClause], risk_summary: Dict[str, int]
    ) -> List[str]:
        """Generate recommendations based on analysis."""
        recommendations: List[str] = []

        # Check for missing important clauses
        found_types = {c.clause_type for c in clauses}
        important_types = [
            ContractClauseType.LIMITATION_OF_LIABILITY,
            ContractClauseType.TERMINATION,
            ContractClauseType.CONFIDENTIALITY,
            ContractClauseType.GOVERNING_LAW,
        ]

        for clause_type in important_types:
            if clause_type not in found_types:
                recommendations.append(
                    f"Consider adding a {clause_type.value.replace('_', ' ')} clause"
                )

        # Risk-based recommendations
        if risk_summary.get("critical", 0) > 0:
            recommendations.append("Critical risk items require immediate legal review")

        if risk_summary.get("high", 0) > 0:
            recommendations.append("High risk items should be negotiated before signing")

        return recommendations

    def _generate_summary(
        self,
        title: Optional[str],
        contract_type: Optional[str],
        parties: List[ContractParty],
        effective_date: Optional[datetime],
        total_value: Optional[float],
    ) -> str:
        """Generate a contract summary."""
        parts = []

        if title:
            parts.append(f"This is a {title}")
        elif contract_type:
            parts.append(f"This is a {contract_type}")

        if len(parties) >= 2:
            parts.append(f"between {parties[0].name} and {parties[1].name}")

        if effective_date:
            parts.append(f"effective {effective_date.strftime('%B %d, %Y')}")

        if total_value:
            parts.append(f"with a total value of ${total_value:,.2f}")

        return " ".join(parts) + "." if parts else "Contract analysis complete."

    def _calculate_confidence(
        self,
        title: Optional[str],
        parties: List[ContractParty],
        clauses: List[ContractClause],
        effective_date: Optional[datetime],
    ) -> float:
        """Calculate overall confidence score."""
        score = 0.0

        if title:
            score += 0.2
        if len(parties) >= 2:
            score += 0.3
        if len(clauses) >= 3:
            score += 0.3
        if effective_date:
            score += 0.2

        return min(score, 1.0)

# Singleton instance
contract_analyzer = ContractAnalyzer()

"""Resume Parser Service.

Extracts structured data from resume/CV documents.
"""
# from __future__ import annotations (already at top)

import base64
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from backend.app.schemas import (
    Education,
    ResumeParseRequest,
    ResumeParseResponse,
    WorkExperience,
)

class ResumeParser(_TextExtractorMixin):
    """Parser for extracting data from resume documents."""

    # Common section headers
    SECTION_HEADERS = {
        "experience": [
            r"(?:work\s+)?experience",
            r"employment\s+history",
            r"professional\s+experience",
            r"work\s+history",
        ],
        "education": [
            r"education",
            r"academic\s+background",
            r"qualifications",
        ],
        "skills": [
            r"skills",
            r"technical\s+skills",
            r"core\s+competencies",
            r"expertise",
        ],
        "certifications": [
            r"certifications?",
            r"licenses?",
            r"credentials",
        ],
        "summary": [
            r"(?:professional\s+)?summary",
            r"(?:career\s+)?objective",
            r"profile",
            r"about\s+me",
        ],
    }

    # Common skills keywords
    TECHNICAL_SKILLS = [
        "python", "javascript", "java", "c++", "c#", "ruby", "go", "rust",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "redis",
        "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
        "react", "angular", "vue", "node.js", "django", "flask",
        "machine learning", "deep learning", "ai", "data science",
        "git", "ci/cd", "agile", "scrum", "jira",
    ]

    def __init__(self) -> None:
        self._ocr_available = self._check_ocr()
        self._nlp_available = self._check_nlp()

    def _check_nlp(self) -> bool:
        """Check if NLP libraries are available."""
        try:
            import spacy  # noqa: F401
            return True
        except Exception:
            return False

    async def parse(self, request: ResumeParseRequest) -> ResumeParseResponse:
        """Parse a resume document."""
        start_time = time.time()
        text = await self._extract_text_from_request(request)

        # Extract contact information
        name = self._extract_name(text)
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        location = self._extract_location(text)
        linkedin_url = self._extract_linkedin(text)
        github_url = self._extract_github(text)
        portfolio_url = self._extract_portfolio(text)

        # Extract summary
        summary = self._extract_summary(text)

        # Extract sections
        education = self._extract_education(text)
        experience = self._extract_experience(text)

        # Extract skills
        skills = (
            self._extract_skills(text)
            if request.extract_skills
            else []
        )

        # Extract certifications and languages
        certifications = self._extract_certifications(text)
        languages = self._extract_languages(text)

        # Calculate total experience
        total_years = self._calculate_total_experience(experience)

        # Job matching if requested
        job_match_score = None
        job_match_details = None
        if request.match_job_description:
            job_match_score, job_match_details = self._match_job(
                text, skills, experience, request.match_job_description
            )

        confidence = self._calculate_confidence(
            name, email, education, experience
        )

        processing_time_ms = int((time.time() - start_time) * 1000)

        return ResumeParseResponse(
            name=name,
            email=email,
            phone=phone,
            location=location,
            linkedin_url=linkedin_url,
            github_url=github_url,
            portfolio_url=portfolio_url,
            summary=summary,
            education=education,
            experience=experience,
            skills=skills,
            certifications=certifications,
            languages=languages,
            total_years_experience=total_years,
            job_match_score=job_match_score,
            job_match_details=job_match_details,
            raw_text=text[:5000] if text else None,
            confidence_score=confidence,
            processing_time_ms=processing_time_ms,
        )

    async def _extract_text_from_request(self, request) -> str:
        """Override: also supports .docx for resumes."""
        if getattr(request, 'file_path', None):
            path = Path(request.file_path)
            if path.suffix.lower() == ".docx":
                return await self._extract_from_docx(path)
        return await super()._extract_text_from_request(request)

    async def _extract_from_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            return ""

    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from resume."""
        # Usually the name is in the first few lines
        lines = text.strip().split("\n")[:5]

        for line in lines:
            line = line.strip()
            # Skip lines that look like headers or contact info
            if "@" in line or re.search(r"\d{3}[-.\s]?\d{3}", line):
                continue
            if len(line) < 3 or len(line) > 50:
                continue

            # Check if line looks like a name (mostly letters and spaces)
            if re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$", line):
                return line

        # Fallback: first non-empty line
        for line in lines:
            line = line.strip()
            if line and len(line) < 50:
                return line

        return None

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume."""
        email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
        return email_match.group(0) if email_match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from resume."""
        # Various phone formats
        phone_patterns = [
            r"\+?1?[-.\s]?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})",
            r"\+?\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,4}",
        ]

        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)

        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location from resume."""
        # Look for city, state pattern
        location_match = re.search(
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z]{2})\b",
            text[:1000]
        )
        if location_match:
            return f"{location_match.group(1)}, {location_match.group(2)}"

        return None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL from resume."""
        url_match = re.search(
            r"(https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9-]+)",
            text, re.IGNORECASE
        )
        if url_match:
            return url_match.group(1)

        bare_match = re.search(r"linkedin\.com/in/([a-zA-Z0-9-]+)", text, re.IGNORECASE)
        if bare_match:
            return f"https://linkedin.com/in/{bare_match.group(1)}"

        label_match = re.search(r"linkedin:\s*([a-zA-Z0-9-]+)", text, re.IGNORECASE)
        if label_match:
            return f"https://linkedin.com/in/{label_match.group(1)}"
        return None

    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL from resume."""
        url_match = re.search(
            r"(https?://(?:www\.)?github\.com/[a-zA-Z0-9-]+)",
            text, re.IGNORECASE
        )
        if url_match:
            return url_match.group(1)

        bare_match = re.search(r"github\.com/([a-zA-Z0-9-]+)", text, re.IGNORECASE)
        if bare_match:
            return f"https://github.com/{bare_match.group(1)}"

        label_match = re.search(r"github:\s*([a-zA-Z0-9-]+)", text, re.IGNORECASE)
        if label_match:
            return f"https://github.com/{label_match.group(1)}"
        return None

    def _extract_portfolio(self, text: str) -> Optional[str]:
        """Extract portfolio URL from resume."""
        url_match = re.search(
            r"(?:portfolio|website|blog)[:\s]*(?:https?://)?([a-zA-Z0-9.-]+\.[a-z]{2,})",
            text, re.IGNORECASE
        )
        if url_match:
            return f"https://{url_match.group(1)}"
        return None

    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary from resume."""
        # Find summary section
        for pattern in self.SECTION_HEADERS["summary"]:
            match = re.search(
                rf"{pattern}\s*:?\s*\n(.+?)(?=\n\s*\n|\n\s*(?:experience|education|skills|certif|certifications)\b|$)",
                text, re.IGNORECASE | re.DOTALL
            )
            if match:
                summary = match.group(1).strip()
                if len(summary) > 20:
                    return summary[:1000]

        return None

    def _extract_education(self, text: str) -> List[Education]:
        """Extract education entries from resume."""
        education_list: List[Education] = []

        # Find education section
        education_section = ""
        for pattern in self.SECTION_HEADERS["education"]:
            match = re.search(
                rf"{pattern}\s*:?\s*\n(.+?)(?:\n\n|experience|skills|certif|$)",
                text, re.IGNORECASE | re.DOTALL
            )
            if match:
                education_section = match.group(1)
                break

        if not education_section:
            return education_list

        # Look for degree patterns
        degree_patterns = [
            r"(Bachelor|Master|PhD|Doctor|Associate|MBA|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)[^,\n]*",
        ]

        # Split into entries by double newline or bullet
        entries = re.split(r"\n\n|\n•|\n-", education_section)

        for entry in entries:
            if len(entry) < 10:
                continue

            # Extract institution
            institution = None
            lines = entry.strip().split("\n")
            if lines:
                institution = lines[0].strip()

            # Extract degree
            degree = None
            for pattern in degree_patterns:
                degree_match = re.search(pattern, entry, re.IGNORECASE)
                if degree_match:
                    degree = degree_match.group(0)
                    break

            # Extract dates
            date_match = re.search(r"(\d{4})\s*[-–]\s*(\d{4}|present)", entry, re.IGNORECASE)
            start_date = date_match.group(1) if date_match else None
            end_date = date_match.group(2) if date_match else None

            if institution:
                education_list.append(Education(
                    institution=institution,
                    degree=degree,
                    start_date=start_date,
                    end_date=end_date,
                ))

        return education_list[:5]  # Limit to 5 entries

    def _extract_experience(self, text: str) -> List[WorkExperience]:
        """Extract work experience entries from resume."""
        experience_list: List[WorkExperience] = []

        # Find experience section
        experience_section = ""
        for pattern in self.SECTION_HEADERS["experience"]:
            match = re.search(
                rf"{pattern}\s*:?\s*\n(.+?)(?:education|skills|certif|$)",
                text, re.IGNORECASE | re.DOTALL
            )
            if match:
                experience_section = match.group(1)
                break

        if not experience_section:
            return experience_list

        # Split into entries
        entries = re.split(r"\n\n\n|\n\n(?=[A-Z])", experience_section)

        for entry in entries:
            if len(entry) < 20:
                continue

            lines = entry.strip().split("\n")
            if not lines:
                continue

            # Extract company and title
            company = None
            title = None

            # First line usually has company or title
            first_line = lines[0].strip()
            if first_line:
                company = first_line

            # Second line might have the other
            if len(lines) > 1:
                second_line = lines[1].strip()
                if re.search(r"(?:manager|engineer|developer|analyst|director|specialist|coordinator)", second_line, re.IGNORECASE):
                    title = second_line
                elif not company:
                    company = second_line

            # Try to find title in entry
            if not title:
                title_match = re.search(
                    r"((?:Senior\s+)?(?:Software|Data|Product|Project|Marketing|Sales|HR|UX|UI|Full\s*Stack|Front\s*End|Back\s*End)\s+"
                    r"(?:Engineer|Developer|Designer|Manager|Analyst|Specialist|Director|Coordinator|Consultant))",
                    entry, re.IGNORECASE
                )
                if title_match:
                    title = title_match.group(1)

            # Extract dates
            date_match = re.search(
                r"(\w+\s+\d{4}|\d{4})\s*[-–]\s*(\w+\s+\d{4}|\d{4}|present|current)",
                entry, re.IGNORECASE
            )
            start_date = date_match.group(1) if date_match else None
            end_date = date_match.group(2) if date_match else None
            is_current = bool(end_date and re.search(r"present|current", end_date, re.IGNORECASE))

            # Extract achievements (bullet points)
            achievements: List[str] = []
            for line in lines:
                line = line.strip()
                if line.startswith(("•", "-", "*", "○")):
                    achievement = line.lstrip("•-*○").strip()
                    if len(achievement) > 10:
                        achievements.append(achievement[:200])

            if company or title:
                experience_list.append(WorkExperience(
                    company=company or "Unknown Company",
                    title=title or "Position",
                    start_date=start_date,
                    end_date=end_date,
                    is_current=is_current,
                    achievements=achievements[:5],
                ))

        return experience_list[:10]  # Limit to 10 entries

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume."""
        skills: List[str] = []
        text_lower = text.lower()

        # Find skills section
        skills_section = ""
        for pattern in self.SECTION_HEADERS["skills"]:
            match = re.search(
                rf"{pattern}\s*:?\s*\n(.+?)(?:\n\n|experience|education|certif|$)",
                text, re.IGNORECASE | re.DOTALL
            )
            if match:
                skills_section = match.group(1)
                break

        # Extract from skills section
        if skills_section:
            # Split by commas, bullets, or pipes
            skill_items = re.split(r"[,•|\n]", skills_section)
            for item in skill_items:
                item = item.strip().strip("-•*")
                if 2 < len(item) < 50:
                    skills.append(item)

        # Also look for known technical skills throughout the document
        for skill in self.TECHNICAL_SKILLS:
            if skill.lower() in text_lower and skill not in [s.lower() for s in skills]:
                skills.append(skill.title() if len(skill) > 3 else skill.upper())

        return list(set(skills))[:30]  # Deduplicate and limit

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume."""
        certifications: List[str] = []

        # Find certifications section
        for pattern in self.SECTION_HEADERS["certifications"]:
            match = re.search(
                rf"{pattern}\s*:?\s*\n(.+?)(?:\n\n|skills|experience|education|$)",
                text, re.IGNORECASE | re.DOTALL
            )
            if match:
                section = match.group(1)
                # Split by newlines or bullets
                items = re.split(r"\n|•|-", section)
                for item in items:
                    item = item.strip()
                    if 5 < len(item) < 100:
                        certifications.append(item)

        # Look for common certification patterns
        cert_patterns = [
            r"(?:AWS|Azure|Google[ \t]+Cloud|GCP)[ \t]+Certified[ \t]+[^,\n]+",
            r"(?:PMP|CISSP|CISM|CEH|CompTIA|CCNA|CCNP)[^,\n]*",
            r"Certified[ \t]+[A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+)*",
        ]

        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)

        return list(set(certifications))[:10]

    def _extract_languages(self, text: str) -> List[str]:
        """Extract spoken languages from resume."""
        languages: List[str] = []

        # Look for languages section or mentions
        lang_match = re.search(
            r"(?:languages?|fluent\s+in|speaks?)\s*:?\s*([^\n]+)",
            text, re.IGNORECASE
        )
        if lang_match:
            lang_text = lang_match.group(1)
            # Split by commas or "and"
            items = re.split(r"[,;]|\band\b", lang_text)
            for item in items:
                item = item.strip()
                # Filter to likely language names
                if re.match(r"^[A-Z][a-z]+(?:\s+\([^)]+\))?$", item):
                    languages.append(item)

        return languages[:5]

    def _calculate_total_experience(
        self, experience: List[WorkExperience]
    ) -> Optional[float]:
        """Calculate total years of experience."""
        total_months = 0
        current_year = datetime.now(timezone.utc).year

        for exp in experience:
            start_year = None
            end_year = None

            if exp.start_date:
                year_match = re.search(r"(\d{4})", exp.start_date)
                if year_match:
                    start_year = int(year_match.group(1))

            if exp.end_date:
                if re.search(r"present|current", exp.end_date, re.IGNORECASE):
                    end_year = current_year
                else:
                    year_match = re.search(r"(\d{4})", exp.end_date)
                    if year_match:
                        end_year = int(year_match.group(1))

            if start_year and end_year:
                total_months += (end_year - start_year) * 12

        return round(total_months / 12, 1) if total_months > 0 else None

    def _match_job(
        self,
        text: str,
        skills: List[str],
        experience: List[WorkExperience],
        job_description: str,
    ) -> Tuple[float, Dict[str, Any]]:
        """Match resume against a job description."""
        job_lower = job_description.lower()
        text_lower = text.lower()

        # Skill match
        job_skills = []
        for skill in self.TECHNICAL_SKILLS:
            if skill.lower() in job_lower:
                job_skills.append(skill)

        matching_skills = [s for s in skills if s.lower() in job_lower]
        skill_score = len(matching_skills) / max(len(job_skills), 1) if job_skills else 0

        # Keyword match
        job_keywords = set(re.findall(r"\b\w{4,}\b", job_lower))
        resume_keywords = set(re.findall(r"\b\w{4,}\b", text_lower))
        keyword_overlap = len(job_keywords & resume_keywords)
        keyword_score = min(keyword_overlap / max(len(job_keywords), 1), 1.0)

        # Experience match (simplified)
        years_required = None
        years_match = re.search(r"(\d+)\+?\s*years?", job_description, re.IGNORECASE)
        if years_match:
            years_required = int(years_match.group(1))

        total_years = self._calculate_total_experience(experience)
        exp_score = 1.0
        if years_required and total_years:
            exp_score = min(total_years / years_required, 1.0)

        # Overall score
        overall_score = (skill_score * 0.4 + keyword_score * 0.3 + exp_score * 0.3)

        details = {
            "skill_score": round(skill_score, 2),
            "keyword_score": round(keyword_score, 2),
            "experience_score": round(exp_score, 2),
            "matching_skills": matching_skills,
            "required_skills": job_skills,
            "years_required": years_required,
            "years_found": total_years,
        }

        return round(overall_score, 2), details

    def _calculate_confidence(
        self,
        name: Optional[str],
        email: Optional[str],
        education: List[Education],
        experience: List[WorkExperience],
    ) -> float:
        """Calculate overall confidence score."""
        score = 0.0

        if name:
            score += 0.2
        if email:
            score += 0.2
        if education:
            score += 0.3
        if experience:
            score += 0.3

        return min(score, 1.0)

# Singleton instance
resume_parser = ResumeParser()
